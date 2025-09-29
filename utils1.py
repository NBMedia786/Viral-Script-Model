

# from __future__ import annotations

# import os
# import re
# import json
# import zipfile
# import xml.etree.ElementTree as ET
# from typing import Any, Dict, List, Optional, Iterable

# # PDF
# try:
#     from PyPDF2 import PdfReader
# except Exception:  # pragma: no cover
#     PdfReader = None  # type: ignore

# # DOCX reading/writing
# from docx import Document
# from docx.oxml.table import CT_Tbl
# from docx.oxml.text.paragraph import CT_P
# from docx.text.paragraph import Paragraph
# from docx.table import Table
# from docx.shared import Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH

# # -----------------------------------------------------------------------------#
# # Constants
# # -----------------------------------------------------------------------------#

# BEGIN_JSON_TOKEN = "BEGIN_JSON"
# END_JSON_TOKEN = "END_JSON"

# # Final display order (7 parameters)
# PARAM_ORDER: List[str] = [
#     "Suspense Building",
#     "Language/Tone",
#     "Intro + Main Hook/Cliffhanger",
#     "Story Structure + Flow",
#     "Pacing",
#     "Mini-Hooks (30–60s)",
#     "Outro (Ending)",
# ]

# # Areas of Improvement (AOI) canonical schema (simplified, unlimited items)
# AOI_KEYS: List[str] = ["quote_verbatim", "issue", "fix", "why_this_helps"]

# # -----------------------------------------------------------------------------#
# # Sanitizer (shared with UI intentions)
# # -----------------------------------------------------------------------------#

# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )

# def sanitize_editor_text(s: Optional[str]) -> str:
#     """Remove Decision/Score boilerplate, leading numeric/bullet labels, emojis, and tidy whitespace."""
#     if not s:
#         return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•|\*)\s*', '', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# # -----------------------------------------------------------------------------#
# # Helpers: normalization (preserve original spacing as much as possible)
# # -----------------------------------------------------------------------------#

# def _strip_bom(s: str) -> str:
#     return s.lstrip("\ufeff") if s and s.startswith("\ufeff") else s

# def _normalize_text(s: str) -> str:
#     """
#     Preserve author formatting as much as we reasonably can:
#     - unify newlines, strip BOM/nbspace
#     - DO NOT collapse multiple spaces/newlines aggressively
#     """
#     if not s:
#         return ""
#     s = _strip_bom(s)
#     s = s.replace("\r\n", "\n").replace("\r", "\n")
#     s = s.replace("\xa0", " ")  # non-breaking space
#     return s.strip()

# # -----------------------------------------------------------------------------#
# # DOCX extraction (paragraphs + tables + text boxes)
# # -----------------------------------------------------------------------------#

# def _iter_block_items(document: Document):
#     """Yield paragraphs and tables in document order (tables appear where seen)."""
#     parent = document.element.body
#     for child in parent.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# def _paragraph_text_with_breaks(p: Paragraph) -> str:
#     """
#     Extract paragraph text preserving soft line breaks (w:br) between runs.
#     python-docx's .text loses explicit <w:br/> breaks; we reconstruct them.
#     """
#     parts: List[str] = []
#     for run in p.runs:
#         # append text first
#         if run.text:
#             parts.append(run.text)
#         # detect <w:br/> under this run
#         for br in run._r.xpath(".//w:br"):
#             parts.append("\n")
#     # Collapse multiple \n from consecutive <w:br> into single newlines
#     txt = "".join(parts)
#     txt = re.sub(r'\n{3,}', '\n\n', txt)
#     return txt

# def _text_from_paragraph(p: Paragraph) -> str:
#     return _paragraph_text_with_breaks(p) or ""

# def _text_from_table(tbl: Table) -> str:
#     """
#     Flatten a table by rows; keep cell text order; add line breaks between rows.
#     Also include nested tables within cells if present.
#     """
#     lines: List[str] = []
#     for row in tbl.rows:
#         row_cells: List[str] = []
#         for cell in row.cells:
#             # paragraphs (with soft breaks)
#             cell_bits: List[str] = []
#             for para in cell.paragraphs:
#                 cell_bits.append(_paragraph_text_with_breaks(para))
#             # nested tables
#             for nt in cell._tc.iterchildren():
#                 if isinstance(nt, CT_Tbl):
#                     t = Table(nt, cell._parent)
#                     nested = _text_from_table(t)
#                     if nested:
#                         cell_bits.append(nested)
#             row_cells.append("\n".join([cl for cl in cell_bits if cl]))
#         # join cells with spacing; then add a row break
#         line = "  ".join([c for c in row_cells if c])
#         lines.append(line)
#     return "\n".join([ln for ln in lines if ln.strip()])

# def _extract_docx_in_order(docx_path: str) -> str:
#     """Extract paragraphs and tables in visual order with minimal normalization."""
#     doc = Document(docx_path)
#     chunks: List[str] = []
#     for block in _iter_block_items(doc):
#         if isinstance(block, Paragraph):
#             t = _text_from_paragraph(block)
#             if t is not None:
#                 chunks.append(t)
#         elif isinstance(block, Table):
#             t = _text_from_table(block)
#             if t is not None:
#                 chunks.append(t)
#     text = "\n".join(chunks)
#     return _normalize_text(text)

# def _extract_textboxes_from_docx(docx_path: str) -> str:
#     """
#     Grab text inside text boxes / shapes (w:txbxContent) straight from the DOCX XML.
#     Preserve paragraph-level breaks inside the textbox.
#     """
#     try:
#         with zipfile.ZipFile(docx_path) as z:
#             xml = z.read("word/document.xml")
#     except Exception:
#         return ""
#     root = ET.fromstring(xml)
#     ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
#     bits: List[str] = []
#     for txbx in root.findall(".//w:txbxContent", ns):
#         paras = txbx.findall(".//w:p", ns)
#         p_texts: List[str] = []
#         for p in paras:
#             # collect runs and breaks inside this textbox paragraph
#             texts: List[str] = []
#             for t in p.findall(".//w:t", ns):
#                 texts.append(t.text or "")
#             # approximate: insert newline between w:p's
#             p_texts.append("".join(texts))
#         if p_texts:
#             bits.append("\n".join(p_texts))
#     return _normalize_text("\n\n".join(bits))

# # -----------------------------------------------------------------------------#
# # PDF extraction
# # -----------------------------------------------------------------------------#

# def _load_pdf_text(path: str) -> str:
#     """Basic PDF text extraction via PyPDF2/PdfReader."""
#     if PdfReader is None:
#         return ""
#     try:
#         reader = PdfReader(path)
#         pages_text: List[str] = []
#         for page in reader.pages:
#             try:
#                 t = page.extract_text() or ""
#             except Exception:
#                 t = ""
#             pages_text.append(t)
#         return _normalize_text("\n\n".join(pages_text))
#     except Exception:
#         return ""

# # -----------------------------------------------------------------------------#
# # Public: load_script_file
# # -----------------------------------------------------------------------------#

# def load_script_file(path: str) -> str:
#     """
#     Load a script from txt/docx/pdf.
#     DOCX: grabs paragraphs + tables in visual order and appends textbox content.
#     PDF: PyPDF2.
#     """
#     ext = os.path.splitext(path)[1].lower()

#     if ext == ".txt":
#         with open(path, "r", encoding="utf-8", errors="ignore") as f:
#             return _normalize_text(f.read())

#     if ext == ".docx":
#         main = _extract_docx_in_order(path)
#         # Always try to append textboxes/shapes after main content (if any distinct content found)
#         tbx = _extract_textboxes_from_docx(path)
#         if tbx and tbx not in main:
#             main = (main + ("\n\n" if main else "") + tbx).strip()
#         return main

#     if ext == ".pdf":
#         return _load_pdf_text(path)

#     return ""

# # -----------------------------------------------------------------------------#
# # JSON extraction from model output
# # -----------------------------------------------------------------------------#

# def _extract_between_tokens(text: str, start_token: str, end_token: str) -> Optional[str]:
#     if not text:
#         return None
#     i = text.find(start_token)
#     j = text.rfind(end_token)
#     if i == -1 or j == -1 or j <= i:
#         return None
#     return text[i + len(start_token) : j].strip()

# def _extract_fenced_block(text: str) -> Optional[str]:
#     if not text:
#         return None
#     m = re.search(r"```json\s*(.*?)\s*```", text, flags=re.DOTALL | re.IGNORECASE)
#     if m:
#         return m.group(1).strip()
#     m = re.search(r"```\s*(.*?)\s*```", text, flags=re.DOTALL)
#     if m:
#         return m.group(1).strip()
#     return None

# def _extract_balanced_json(text: str) -> Optional[str]:
#     if not text:
#         return None
#     start = text.find("{")
#     if start == -1:
#         return None
#     depth = 0
#     for i in range(start, len(text)):
#         ch = text[i]
#         if ch == "{":
#             depth += 1
#         elif ch == "}":
#             depth -= 1
#             if depth == 0:
#                 return text[start : i + 1]
#     return None

# def extract_review_json(model_output: str) -> Optional[Dict[str, Any]]:
#     """
#     Extract and parse the structured JSON from the model output.
#     Priority:
#       1) BETWEEN BEGIN_JSON / END_JSON
#       2) First ```json fenced block
#       3) First triple-fenced block
#       4) Balanced-brace fallback
#     """
#     if not model_output:
#         return None

#     candidates: List[str] = []

#     between = _extract_between_tokens(model_output, BEGIN_JSON_TOKEN, END_JSON_TOKEN)
#     if between:
#         candidates.append(between)

#     fenced = _extract_fenced_block(model_output)
#     if fenced:
#         candidates.append(fenced)

#     balanced = _extract_balanced_json(model_output)
#     if balanced:
#         candidates.append(balanced)

#     for s in candidates:
#         try:
#             return json.loads(s)
#         except Exception:
#             # trailing comma healing
#             s2 = re.sub(r",\s*([\]}])", r"\1", s)
#             try:
#                 return json.loads(s2)
#             except Exception:
#                 continue
#     return None

# # -----------------------------------------------------------------------------#
# # AOI utilities (normalize to the 4-field schema used by the new UI)
# # -----------------------------------------------------------------------------#

# def _clean_str(x: Any) -> str:
#     if x is None:
#         return ""
#     return str(x).strip()

# def _coerce_aois(block: Dict[str, Any]) -> List[Dict[str, str]]:
#     """
#     Normalize a parameter's AOIs to the simplified schema:
#       {quote_verbatim, issue, fix, why_this_helps}
#     - Accepts any incoming shape; pulls/renames common legacy keys if needed.
#     - Filters out empty entries (no quote & no issue/fix).
#     """
#     raw = block.get("areas_of_improvement") or []
#     if not isinstance(raw, Iterable):
#         return []
#     out: List[Dict[str, str]] = []
#     for item in raw:
#         if not isinstance(item, dict):
#             continue
#         q = _clean_str(item.get("quote_verbatim") or item.get("quote") or item.get("line") or "")
#         issue = _clean_str(item.get("issue") or "")
#         # map legacy 'edit_suggestion' → 'fix'
#         fix = _clean_str(item.get("fix") or item.get("edit_suggestion") or "")
#         why = _clean_str(item.get("why_this_helps") or item.get("why") or "")
#         if not (q or issue or fix or why):
#             continue
#         out.append({
#             "quote_verbatim": q[:240] if q else "",
#             "issue": issue,
#             "fix": fix,
#             "why_this_helps": why,
#         })
#     return out

# def normalize_review_payload(data: Dict[str, Any]) -> Dict[str, Any]:
#     """
#     OPTIONAL helper: normalize AOIs across all parameters to the canonical 4-field schema
#     AND sanitize explanation/weakness/suggestion/summary + AOI fields for clean output.
#     Returns the same dict (mutated in place) for convenience.
#     """
#     if not isinstance(data, dict):
#         return data

#     # sanitize top-level fields if present
#     for k in ("strengths", "weaknesses", "suggestions", "drop_off_risks"):
#         if isinstance(data.get(k), list):
#             data[k] = [sanitize_editor_text(x) for x in data[k]]

#     if isinstance(data.get("viral_quotient"), str):
#         data["viral_quotient"] = sanitize_editor_text(data["viral_quotient"])

#     per = data.get("per_parameter") or {}
#     if isinstance(per, dict):
#         for _, block in per.items():
#             if not isinstance(block, dict):
#                 continue
#             # normalize AOIs
#             block["areas_of_improvement"] = _coerce_aois(block)
#             # sanitize common fields
#             for fld in ("explanation", "weakness", "suggestion", "summary"):
#                 if isinstance(block.get(fld), str):
#                     block[fld] = sanitize_editor_text(block[fld])
#             # sanitize AOIs
#             aois = block.get("areas_of_improvement") or []
#             for a in aois:
#                 if not isinstance(a, dict): 
#                     continue
#                 for fld in ("quote_verbatim", "issue", "fix", "why_this_helps"):
#                     if isinstance(a.get(fld), str):
#                         a[fld] = sanitize_editor_text(a[fld])
#     return data

# # -----------------------------------------------------------------------------#
# # Optional: DOCX renderer (AOIs-first; raw extractions hidden)
# # -----------------------------------------------------------------------------#

# def _add_heading(doc: Document, text: str, size: int = 14, bold: bool = True):
#     p = doc.add_paragraph()
#     run = p.add_run(text)
#     run.bold = bold
#     run.font.size = Pt(size)

# def _add_bullets(doc: Document, items: List[str]):
#     items = items or []
#     for it in items:
#         doc.add_paragraph(f"• {it}")

# def _add_param_block(doc: Document, title: str, block: Dict[str, Any]):
#     _add_heading(doc, title, size=12, bold=True)

#     # AOIs first (primary editorial output for fixes)
#     aois = _coerce_aois(block)
#     if aois:
#         doc.add_paragraph("Areas of Improvement:")
#         for i, a in enumerate(aois, start=1):
#             q = a.get("quote_verbatim") or ""
#             issue = a.get("issue") or ""
#             fix = a.get("fix") or ""
#             why = a.get("why_this_helps") or ""
#             if q:
#                 doc.add_paragraph(f"  {i}. Line: {q}")
#             if issue:
#                 doc.add_paragraph(f"     • Issue: {issue}")
#             if fix:
#                 doc.add_paragraph(f"     • Fix: {fix}")
#             if why:
#                 doc.add_paragraph(f"     • Why: {why}")

#     # Raw legacy extractions intentionally hidden now

#     sc = block.get("score", "")
#     if sc != "":
#         doc.add_paragraph(f"Score: {sc}/10")
#     if block.get("explanation"):
#         doc.add_paragraph(f"Explanation: {block['explanation']}")
#     if block.get("weakness"):
#         doc.add_paragraph(f"Weakness: {block['weakness']}")
#     if block.get("suggestion"):
#         doc.add_paragraph(f"Suggestion: {block['suggestion']}")
#     if block.get("summary"):
#         doc.add_paragraph(f"Summary: {block['summary']}")
#     doc.add_paragraph("")

# def save_review_docx_claude_style_from_json(
#     data: Dict[str, Any],
#     out_path: str,
#     title: Optional[str] = None,
# ):
#     # Safe dir creation for local filenames
#     dirpath = os.path.dirname(out_path)
#     if dirpath:
#         os.makedirs(dirpath, exist_ok=True)

#     # Normalize & sanitize for clean export
#     normalize_review_payload(data)

#     doc = Document()

#     title_text = title or "Viral Script Review"
#     p = doc.add_paragraph()
#     run = p.add_run(title_text)
#     run.bold = True
#     run.font.size = Pt(16)
#     p.alignment = WD_ALIGN_PARAGRAPH.LEFT
#     doc.add_paragraph("")

#     _add_heading(doc, "Parameter Analysis", size=14)
#     per_param: Dict[str, Any] = data.get("per_parameter", {}) or {}
#     for param in PARAM_ORDER:
#         block = per_param.get(param, {}) or {}
#         _add_param_block(doc, param, block)

#     scores: Dict[str, Any] = data.get("scores", {}) or {}
#     _add_heading(doc, "Scoring Table", size=14)
#     table = doc.add_table(rows=len(PARAM_ORDER) + 1, cols=2)
#     table.style = "Table Grid"
#     table.rows[0].cells[0].text = "Parameter"
#     table.rows[0].cells[1].text = "Score (1–10)"
#     for i, p_name in enumerate(PARAM_ORDER, start=1):
#         table.rows[i].cells[0].text = p_name
#         table.rows[i].cells[1].text = str(scores.get(p_name, ""))

#     doc.add_paragraph("")
#     overall = data.get("overall_rating", "—")
#     _add_heading(doc, f"Overall Rating: {overall}/10", size=14)

#     doc.add_paragraph("")
#     _add_heading(doc, "Strengths", size=13)
#     _add_bullets(doc, data.get("strengths") or [])
#     doc.add_paragraph("")
#     _add_heading(doc, "Weaknesses", size=13)
#     _add_bullets(doc, data.get("weaknesses") or [])
#     doc.add_paragraph("")
#     _add_heading(doc, "Suggestions", size=13)
#     _add_bullets(doc, data.get("suggestions") or [])
#     doc.add_paragraph("")
#     _add_heading(doc, "Drop-off Risks", size=13)
#     _add_bullets(doc, data.get("drop_off_risks") or [])
#     doc.add_paragraph("")
#     _add_heading(doc, "Viral Quotient", size=13)
#     vq = data.get("viral_quotient", "")
#     if vq:
#         doc.add_paragraph(vq)

#     doc.save(out_path)

















































##################################################










# # utils1.py — helpers shared by app_grammerly_ui.py and run_review_multi.py

# from __future__ import annotations

# import os, re, json, zipfile, xml.etree.ElementTree as ET
# from typing import Any, Dict, List, Optional, Iterable

# # Optional PDF support
# try:
#     from PyPDF2 import PdfReader
# except Exception:
#     PdfReader = None  # type: ignore

# # DOCX helpers
# from docx import Document
# from docx.oxml.table import CT_Tbl
# from docx.oxml.text.paragraph import CT_P
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # -------------------- UI order for parameters -------------------- #
# PARAM_ORDER: List[str] = [
#     "Suspense Building",
#     "Language/Tone",
#     "Intro + Main Hook/Cliffhanger",
#     "Story Structure + Flow",
#     "Pacing",
#     "Mini-Hooks (30–60s)",
#     "Outro (Ending)",
# ]

# # -------------------- Sanitizers -------------------- #
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )

# def sanitize_editor_text(s: Optional[str]) -> str:
#     if not s:
#         return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•|\*)\s*', '', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# def _strip_bom(s: str) -> str:
#     return s.lstrip("\ufeff") if s and s.startswith("\ufeff") else s

# def _normalize_text(s: str) -> str:
#     if not s:
#         return ""
#     s = _strip_bom(s)
#     s = s.replace("\r\n", "\n").replace("\r", "\n")
#     s = s.replace("\xa0", " ")
#     return s.strip()

# # -------------------- DOCX extraction -------------------- #
# def _iter_block_items(document: Document):
#     parent = document.element.body
#     for child in parent.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# def _paragraph_text_with_breaks(p: Paragraph) -> str:
#     parts: List[str] = []
#     for run in p.runs:
#         if run.text:
#             parts.append(run.text)
#         for _ in run._r.xpath(".//w:br"):
#             parts.append("\n")
#     txt = "".join(parts)
#     txt = re.sub(r'\n{3,}', '\n\n', txt)
#     return txt

# def _text_from_paragraph(p: Paragraph) -> str:
#     return _paragraph_text_with_breaks(p) or ""

# def _text_from_table(tbl: Table) -> str:
#     lines: List[str] = []
#     for row in tbl.rows:
#         row_cells: List[str] = []
#         for cell in row.cells:
#             cell_bits: List[str] = []
#             for para in cell.paragraphs:
#                 cell_bits.append(_paragraph_text_with_breaks(para))
#             for nt in cell._tc.iterchildren():
#                 if isinstance(nt, CT_Tbl):
#                     t = Table(nt, cell._parent)
#                     nested = _text_from_table(t)
#                     if nested:
#                         cell_bits.append(nested)
#             row_cells.append("\n".join([cl for cl in cell_bits if cl]))
#         line = "  ".join([c for c in row_cells if c])
#         lines.append(line)
#     return "\n".join([ln for ln in lines if ln.strip()])

# def _extract_docx_in_order(docx_path: str) -> str:
#     doc = Document(docx_path)
#     chunks: List[str] = []
#     for block in _iter_block_items(doc):
#         if isinstance(block, Paragraph):
#             t = _text_from_paragraph(block)
#             if t is not None:
#                 chunks.append(t)
#         elif isinstance(block, Table):
#             t = _text_from_table(block)
#             if t is not None:
#                 chunks.append(t)
#     text = "\n".join(chunks)
#     return _normalize_text(text)

# def _extract_textboxes_from_docx(docx_path: str) -> str:
#     try:
#         with zipfile.ZipFile(docx_path) as z:
#             xml = z.read("word/document.xml")
#     except Exception:
#         return ""
#     root = ET.fromstring(xml)
#     ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
#     bits: List[str] = []
#     for txbx in root.findall(".//w:txbxContent", ns):
#         paras = txbx.findall(".//w:p", ns)
#         p_texts: List[str] = []
#         for p in paras:
#             texts: List[str] = []
#             for t in p.findall(".//w:t", ns):
#                 texts.append(t.text or "")
#             p_texts.append("".join(texts))
#         if p_texts:
#             bits.append("\n".join(p_texts))
#     return _normalize_text("\n\n".join(bits))

# # -------------------- PDF extraction -------------------- #
# def _load_pdf_text(path: str) -> str:
#     if PdfReader is None:
#         return ""
#     try:
#         reader = PdfReader(path)
#         pages_text: List[str] = []
#         for page in reader.pages:
#             try:
#                 t = page.extract_text() or ""
#             except Exception:
#                 t = ""
#             pages_text.append(t)
#         return _normalize_text("\n\n".join(pages_text))
#     except Exception:
#         return ""

# # -------------------- Public file loader -------------------- #
# def load_script_file(path: str) -> str:
#     ext = os.path.splitext(path)[1].lower()
#     if ext == ".txt":
#         with open(path, "r", encoding="utf-8", errors="ignore") as f:
#             return _normalize_text(f.read())
#     if ext == ".docx":
#         main = _extract_docx_in_order(path)
#         tbx = _extract_textboxes_from_docx(path)
#         if tbx and tbx not in main:
#             main = (main + ("\n\n" if main else "") + tbx).strip()
#         return main
#     if ext == ".pdf":
#         return _load_pdf_text(path)
#     return ""

# # -------------------- JSON extraction helpers -------------------- #
# BEGIN_JSON_TOKEN = "BEGIN_JSON"
# END_JSON_TOKEN = "END_JSON"

# def _extract_between_tokens(text: str, start_token: str, end_token: str) -> Optional[str]:
#     if not text:
#         return None
#     i = text.find(start_token)
#     j = text.rfind(end_token)
#     if i == -1 or j == -1 or j <= i:
#         return None
#     return text[i + len(start_token) : j].strip()

# def _extract_fenced_block(text: str) -> Optional[str]:
#     if not text:
#         return None
#     m = re.search(r"```json\s*(.*?)\s*```", text, flags=re.DOTALL | re.IGNORECASE)
#     if m:
#         return m.group(1).strip()
#     m = re.search(r"```\s*(.*?)\s*```", text, flags=re.DOTALL)
#     if m:
#         return m.group(1).strip()
#     return None

# def _extract_balanced_json(text: str) -> Optional[str]:
#     if not text:
#         return None
#     start = text.find("{")
#     if start == -1:
#         return None
#     depth = 0
#     for i in range(start, len(text)):
#         ch = text[i]
#         if ch == "{":
#             depth += 1
#         elif ch == "}":
#             depth -= 1
#             if depth == 0:
#                 return text[start : i + 1]
#     return None

# def extract_review_json(model_output: str) -> Optional[Dict[str, Any]]:
#     if not model_output:
#         return None
#     candidates: List[str] = []
#     between = _extract_between_tokens(model_output, BEGIN_JSON_TOKEN, END_JSON_TOKEN)
#     if between:
#         candidates.append(between)
#     fenced = _extract_fenced_block(model_output)
#     if fenced:
#         candidates.append(fenced)
#     balanced = _extract_balanced_json(model_output)
#     if balanced:
#         candidates.append(balanced)
#     for s in candidates:
#         try:
#             return json.loads(s)
#         except Exception:
#             s2 = re.sub(r",\s*([\]}])", r"\1", s)
#             try:
#                 return json.loads(s2)
#             except Exception:
#                 continue
#     return None

# # -------------------- AOI normalization for UI -------------------- #
# AOI_KEYS: List[str] = ["quote_verbatim", "issue", "fix", "why_this_helps"]

# def _clean_str(x: Any) -> str:
#     if x is None:
#         return ""
#     return str(x).strip()

# def _coerce_aois(block: Dict[str, Any]) -> List[Dict[str, str]]:
#     raw = block.get("areas_of_improvement") or []
#     if not isinstance(raw, Iterable):
#         return []
#     out: List[Dict[str, str]] = []
#     for item in raw:
#         if not isinstance(item, dict):
#             continue
#         q = _clean_str(item.get("quote_verbatim") or item.get("quote") or item.get("line") or "")
#         issue = _clean_str(item.get("issue") or "")
#         fix = _clean_str(item.get("fix") or item.get("edit_suggestion") or "")
#         why = _clean_str(item.get("why_this_helps") or item.get("why") or "")
#         if not (q or issue or fix or why):
#             continue
#         out.append({
#             "quote_verbatim": q[:240] if q else "",
#             "issue": issue,
#             "fix": fix,
#             "why_this_helps": why,
#         })
#     return out

# def normalize_review_payload(data: Dict[str, Any]) -> Dict[str, Any]:
#     if not isinstance(data, dict):
#         return data

#     for k in ("strengths", "weaknesses", "suggestions", "drop_off_risks"):
#         if isinstance(data.get(k), list):
#             data[k] = [sanitize_editor_text(x) for x in data[k]]

#     if isinstance(data.get("viral_quotient"), str):
#         data["viral_quotient"] = sanitize_editor_text(data["viral_quotient"])

#     per = data.get("per_parameter") or {}
#     if isinstance(per, dict):
#         for _, block in per.items():
#             if not isinstance(block, dict):
#                 continue
#             block["areas_of_improvement"] = _coerce_aois(block)
#             for fld in ("explanation", "weakness", "suggestion", "summary"):
#                 if isinstance(block.get(fld), str):
#                     block[fld] = sanitize_editor_text(block[fld])
#             aois = block.get("areas_of_improvement") or []
#             for a in aois:
#                 if not isinstance(a, dict): 
#                     continue
#                 for fld in ("quote_verbatim", "issue", "fix", "why_this_helps"):
#                     if isinstance(a.get(fld), str):
#                         a[fld] = sanitize_editor_text(a[fld])
#     return data




##########################################


###Current Working version









# # utils1.py — helpers shared by app_grammerly_ui.py and run_review_multi.py

# from __future__ import annotations

# import os, re, json, zipfile, xml.etree.ElementTree as ET
# from typing import Any, Dict, List, Optional, Iterable

# # Optional PDF support
# try:
#     from PyPDF2 import PdfReader
# except Exception:
#     PdfReader = None  # type: ignore

# # DOCX helpers
# from docx import Document
# from docx.oxml.table import CT_Tbl
# from docx.oxml.text.paragraph import CT_P
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # -------------------- UI order for parameters -------------------- #
# PARAM_ORDER: List[str] = [
#     "Suspense Building",
#     "Language/Tone",
#     "Intro + Main Hook/Cliffhanger",
#     "Story Structure + Flow",
#     "Pacing",
#     "Mini-Hooks (30–60s)",
#     "Outro (Ending)",
#     "Grammar & Spelling",  # <-- NEW parameter appears in UI
# ]

# # -------------------- Sanitizers -------------------- #
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )

# def sanitize_editor_text(s: Optional[str]) -> str:
#     if not s:
#         return ""
#     t = str(s)
#     # Remove accidental "Decision:" or "Score: 8" echoes
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     # Strip leading list bullets/numbers from each line
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•|\*)\s*', '', t, flags=re.M)
#     # Remove emojis and normalize whitespace
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# def _strip_bom(s: str) -> str:
#     return s.lstrip("\ufeff") if s and s.startswith("\ufeff") else s

# def _normalize_text(s: str) -> str:
#     if not s:
#         return ""
#     s = _strip_bom(s)
#     s = s.replace("\r\n", "\n").replace("\r", "\n")
#     s = s.replace("\xa0", " ")
#     return s.strip()

# # -------------------- DOCX extraction -------------------- #
# def _iter_block_items(document: Document):
#     parent = document.element.body
#     for child in parent.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# def _paragraph_text_with_breaks(p: Paragraph) -> str:
#     parts: List[str] = []
#     for run in p.runs:
#         if run.text:
#             parts.append(run.text)
#         for _ in run._r.xpath(".//w:br"):
#             parts.append("\n")
#     txt = "".join(parts)
#     txt = re.sub(r'\n{3,}', '\n\n', txt)
#     return txt

# def _text_from_paragraph(p: Paragraph) -> str:
#     return _paragraph_text_with_breaks(p) or ""

# def _text_from_table(tbl: Table) -> str:
#     lines: List[str] = []
#     for row in tbl.rows:
#         row_cells: List[str] = []
#         for cell in row.cells:
#             cell_bits: List[str] = []
#             for para in cell.paragraphs:
#                 cell_bits.append(_paragraph_text_with_breaks(para))
#             for nt in cell._tc.iterchildren():
#                 if isinstance(nt, CT_Tbl):
#                     t = Table(nt, cell._parent)
#                     nested = _text_from_table(t)
#                     if nested:
#                         cell_bits.append(nested)
#             row_cells.append("\n".join([cl for cl in cell_bits if cl]))
#         line = "  ".join([c for c in row_cells if c])
#         lines.append(line)
#     return "\n".join([ln for ln in lines if ln.strip()])

# def _extract_docx_in_order(docx_path: str) -> str:
#     doc = Document(docx_path)
#     chunks: List[str] = []
#     for block in _iter_block_items(doc):
#         if isinstance(block, Paragraph):
#             t = _text_from_paragraph(block)
#             if t is not None:
#                 chunks.append(t)
#         elif isinstance(block, Table):
#             t = _text_from_table(block)
#             if t is not None:
#                 chunks.append(t)
#     text = "\n".join(chunks)
#     return _normalize_text(text)

# def _extract_textboxes_from_docx(docx_path: str) -> str:
#     try:
#         with zipfile.ZipFile(docx_path) as z:
#             xml = z.read("word/document.xml")
#     except Exception:
#         return ""
#     root = ET.fromstring(xml)
#     ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
#     bits: List[str] = []
#     for txbx in root.findall(".//w:txbxContent", ns):
#         paras = txbx.findall(".//w:p", ns)
#         p_texts: List[str] = []
#         for p in paras:
#             texts: List[str] = []
#             for t in p.findall(".//w:t", ns):
#                 texts.append(t.text or "")
#             p_texts.append("".join(texts))
#         if p_texts:
#             bits.append("\n".join(p_texts))
#     return _normalize_text("\n\n".join(bits))

# # -------------------- PDF extraction -------------------- #
# def _load_pdf_text(path: str) -> str:
#     if PdfReader is None:
#         return ""
#     try:
#         reader = PdfReader(path)
#         pages_text: List[str] = []
#         for page in reader.pages:
#             try:
#                 t = page.extract_text() or ""
#             except Exception:
#                 t = ""
#             pages_text.append(t)
#         return _normalize_text("\n\n".join(pages_text))
#     except Exception:
#         return ""

# # -------------------- Public file loader -------------------- #
# def load_script_file(path: str) -> str:
#     ext = os.path.splitext(path)[1].lower()
#     if ext == ".txt":
#         with open(path, "r", encoding="utf-8", errors="ignore") as f:
#             return _normalize_text(f.read())
#     if ext == ".docx":
#         main = _extract_docx_in_order(path)
#         tbx = _extract_textboxes_from_docx(path)
#         if tbx and tbx not in main:
#             main = (main + ("\n\n" if main else "") + tbx).strip()
#         return main
#     if ext == ".pdf":
#         return _load_pdf_text(path)
#     return ""

# # -------------------- JSON extraction helpers -------------------- #
# BEGIN_JSON_TOKEN = "BEGIN_JSON"
# END_JSON_TOKEN = "END_JSON"

# def _extract_between_tokens(text: str, start_token: str, end_token: str) -> Optional[str]:
#     if not text:
#         return None
#     i = text.find(start_token)
#     j = text.rfind(end_token)
#     if i == -1 or j == -1 or j <= i:
#         return None
#     return text[i + len(start_token) : j].strip()

# def _extract_fenced_block(text: str) -> Optional[str]:
#     if not text:
#         return None
#     m = re.search(r"```json\s*(.*?)\s*```", text, flags=re.DOTALL | re.IGNORECASE)
#     if m:
#         return m.group(1).strip()
#     m = re.search(r"```\s*(.*?)\s*```", text, flags=re.DOTALL)
#     if m:
#         return m.group(1).strip()
#     return None

# def _extract_balanced_json(text: str) -> Optional[str]:
#     if not text:
#         return None
#     start = text.find("{")
#     if start == -1:
#         return None
#     depth = 0
#     for i in range(start, len(text)):
#         ch = text[i]
#         if ch == "{":
#             depth += 1
#         elif ch == "}":
#             depth -= 1
#             if depth == 0:
#                 return text[start : i + 1]
#     return None

# def extract_review_json(model_output: str) -> Optional[Dict[str, Any]]:
#     if not model_output:
#         return None
#     candidates: List[str] = []
#     between = _extract_between_tokens(model_output, BEGIN_JSON_TOKEN, END_JSON_TOKEN)
#     if between:
#         candidates.append(between)
#     fenced = _extract_fenced_block(model_output)
#     if fenced:
#         candidates.append(fenced)
#     balanced = _extract_balanced_json(model_output)
#     if balanced:
#         candidates.append(balanced)
#     for s in candidates:
#         try:
#             return json.loads(s)
#         except Exception:
#             s2 = re.sub(r",\s*([\]}])", r"\1", s)
#             try:
#                 return json.loads(s2)
#             except Exception:
#                 continue
#     return None

# # -------------------- AOI normalization for UI -------------------- #
# AOI_KEYS: List[str] = ["quote_verbatim", "issue", "fix", "why_this_helps"]

# def _clean_str(x: Any) -> str:
#     if x is None:
#         return ""
#     return str(x).strip()

# def _coerce_aois(block: Dict[str, Any]) -> List[Dict[str, str]]:
#     raw = block.get("areas_of_improvement") or []
#     if not isinstance(raw, Iterable):
#         return []
#     out: List[Dict[str, str]] = []
#     for item in raw:
#         if not isinstance(item, dict):
#             continue
#         q = _clean_str(item.get("quote_verbatim") or item.get("quote") or item.get("line") or "")
#         issue = _clean_str(item.get("issue") or "")
#         fix = _clean_str(item.get("fix") or item.get("edit_suggestion") or "")
#         why = _clean_str(item.get("why_this_helps") or item.get("why") or "")
#         if not (q or issue or fix or why):
#             continue
#         out.append({
#             "quote_verbatim": q[:240] if q else "",
#             "issue": issue,
#             "fix": fix,
#             "why_this_helps": why,
#         })
#     return out

# def normalize_review_payload(data: Dict[str, Any]) -> Dict[str, Any]:
#     if not isinstance(data, dict):
#         return data

#     # Top-level lists
#     for k in ("strengths", "weaknesses", "suggestions", "drop_off_risks"):
#         if isinstance(data.get(k), list):
#             data[k] = [sanitize_editor_text(x) for x in data[k]]

#     if isinstance(data.get("viral_quotient"), str):
#         data["viral_quotient"] = sanitize_editor_text(data["viral_quotient"])

#     per = data.get("per_parameter") or {}
#     if isinstance(per, dict):
#         for _, block in per.items():
#             if not isinstance(block, dict):
#                 continue

#             # AOIs normalized
#             block["areas_of_improvement"] = _coerce_aois(block)

#             # Standard text fields
#             for fld in ("explanation", "weakness", "suggestion", "summary"):
#                 if isinstance(block.get(fld), str):
#                     block[fld] = sanitize_editor_text(block[fld])

#             # Optional tense metadata for Grammar & Spelling (passed through + sanitized)
#             for fld in ("dominant_tense", "tense_target", "tense_consistency"):
#                 if isinstance(block.get(fld), str):
#                     block[fld] = sanitize_editor_text(block[fld])

#             # Sanitize AOIs fields
#             aois = block.get("areas_of_improvement") or []
#             for a in aois:
#                 if not isinstance(a, dict):
#                     continue
#                 for fld in ("quote_verbatim", "issue", "fix", "why_this_helps"):
#                     if isinstance(a.get(fld), str):
#                         a[fld] = sanitize_editor_text(a[fld])

#     return data











##################################



















# # utils1.py — helpers shared by app_grammarly_ui.py and run_review_multi.py

# from __future__ import annotations
# import os, re, json, zipfile, xml.etree.ElementTree as ET
# from typing import Any, Dict, List, Optional, Iterable

# # Optional PDF support
# try:
#     from PyPDF2 import PdfReader
# except Exception:
#     PdfReader = None  # type: ignore

# # DOCX helpers
# from docx import Document
# from docx.oxml.table import CT_Tbl
# from docx.oxml.text.paragraph import CT_P
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # -------------------- UI order for parameters -------------------- #
# PARAM_ORDER: List[str] = [
#     "Suspense Building",
#     "Language/Tone",
#     "Intro + Main Hook/Cliffhanger",
#     "Story Structure + Flow",
#     "Pacing",
#     "Mini-Hooks (30–60s)",
#     "Outro (Ending)",
#     "Grammar & Spelling",  # NEW parameter
# ]

# # -------------------- Sanitizers -------------------- #
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE,
# )

# def sanitize_editor_text(s: Optional[str]) -> str:
#     if not s:
#         return ""
#     t = str(s)
#     # Remove accidental "Decision:" or "Score: 8" echoes
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     # Strip bullets/numbers
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•|\*)\s*', '', t, flags=re.M)
#     # Remove emojis and compress whitespace
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# def _strip_bom(s: str) -> str:
#     return s.lstrip("\ufeff") if s and s.startswith("\ufeff") else s

# def _normalize_text(s: str) -> str:
#     if not s:
#         return ""
#     s = _strip_bom(s)
#     s = s.replace("\r\n", "\n").replace("\r", "\n")
#     s = s.replace("\xa0", " ")
#     s = re.sub(r"[ \t]+", " ", s)
#     s = re.sub(r"\n{3,}", "\n\n", s)
#     return s.strip()

# # -------------------- DOCX extraction -------------------- #
# def _iter_block_items(document: Document):
#     parent = document.element.body
#     for child in parent.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# def _paragraph_text_with_breaks(p: Paragraph) -> str:
#     parts: List[str] = []
#     for run in p.runs:
#         if run.text:
#             parts.append(run.text)
#         for _ in run._r.xpath(".//w:br"):
#             parts.append("\n")
#     txt = "".join(parts)
#     return re.sub(r"\n{3,}", "\n\n", txt)

# def _text_from_paragraph(p: Paragraph) -> str:
#     return _paragraph_text_with_breaks(p) or ""

# def _text_from_table(tbl: Table) -> str:
#     lines: List[str] = []
#     for row in tbl.rows:
#         row_cells: List[str] = []
#         for cell in row.cells:
#             cell_bits: List[str] = []
#             for para in cell.paragraphs:
#                 cell_bits.append(_paragraph_text_with_breaks(para))
#             for nt in cell._tc.iterchildren():
#                 if isinstance(nt, CT_Tbl):
#                     t = Table(nt, cell._parent)
#                     nested = _text_from_table(t)
#                     if nested:
#                         cell_bits.append(nested)
#             row_cells.append("\n".join([cl for cl in cell_bits if cl]))
#         line = "  ".join([c for c in row_cells if c])
#         lines.append(line)
#     return "\n".join([ln for ln in lines if ln.strip()])

# def _extract_docx_in_order(docx_path: str) -> str:
#     doc = Document(docx_path)
#     chunks: List[str] = []
#     for block in _iter_block_items(doc):
#         if isinstance(block, Paragraph):
#             t = _text_from_paragraph(block)
#             if t is not None:
#                 chunks.append(t)
#         elif isinstance(block, Table):
#             t = _text_from_table(block)
#             if t is not None:
#                 chunks.append(t)
#     return _normalize_text("\n".join(chunks))

# def _extract_textboxes_from_docx(docx_path: str) -> str:
#     try:
#         with zipfile.ZipFile(docx_path) as z:
#             xml = z.read("word/document.xml")
#     except Exception:
#         return ""
#     root = ET.fromstring(xml)
#     ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
#     bits: List[str] = []
#     for txbx in root.findall(".//w:txbxContent", ns):
#         paras = txbx.findall(".//w:p", ns)
#         p_texts: List[str] = []
#         for p in paras:
#             texts: List[str] = [t.text or "" for t in p.findall(".//w:t", ns)]
#             p_texts.append("".join(texts))
#         if p_texts:
#             bits.append("\n".join(p_texts))
#     return _normalize_text("\n\n".join(bits))

# # -------------------- LEFT-COLUMN extractor -------------------- #
# def _text_from_cell_paragraphs(cell) -> str:
#     parts: List[str] = []
#     for i, p in enumerate(cell.paragraphs):
#         t = "".join(r.text or "" for r in p.runs)
#         parts.append(t)
#         if i != len(cell.paragraphs) - 1:
#             parts.append("\n")
#     return "".join(parts).strip()

# def _table_looks_like_narration_visuals(tbl: Table) -> bool:
#     try:
#         if len(tbl.rows) == 0 or len(tbl.columns) < 2:
#             return False
#     except Exception:
#         return False
#     try:
#         hdr0 = _text_from_cell_paragraphs(tbl.rows[0].cells[0]).lower()
#         hdr1 = _text_from_cell_paragraphs(tbl.rows[0].cells[1]).lower()
#         if ("narration" in hdr0 and ("visual" in hdr1 or "clip" in hdr1 or "time" in hdr1)):
#             return True
#     except Exception:
#         pass
#     import re
#     link_re = re.compile(r'(https?://|youtube\.com|drive\.google|\.mp4\b|\b\d{1,2}:\d{2}\b)')
#     right_hits = left_hits = 0
#     for r in tbl.rows[: min(6, len(tbl.rows))]:
#         if len(r.cells) < 2:
#             continue
#         left = _text_from_cell_paragraphs(r.cells[0])
#         right = _text_from_cell_paragraphs(r.cells[1])
#         if left.strip():
#             left_hits += 1
#         if link_re.search(right):
#             right_hits += 1
#     return right_hits >= 2 and left_hits >= 2

# def _text_from_table_first_column(tbl: Table) -> str:
#     lines: List[str] = []
#     seen_tc_ids = set()
#     for row in tbl.rows:
#         if not row.cells:
#             continue
#         c0 = row.cells[0]
#         tc_id = id(c0._tc)
#         if tc_id in seen_tc_ids:
#             continue
#         seen_tc_ids.add(tc_id)
#         cell_text = _text_from_cell_paragraphs(c0)
#         if cell_text.strip():
#             lines.append(cell_text.strip())
#     return "\n\n".join(lines)

# ##new 

# def _has_two_plus_col_table(doc: Document) -> bool:
#     """Return True if the DOCX has at least one table with >= 2 columns."""
#     try:
#         for blk in _iter_block_items(doc):
#             if isinstance(blk, Table):
#                 # Robust even with merged cells
#                 try:
#                     if len(blk.columns) >= 2:
#                         return True
#                 except Exception:
#                     # Fallback: check first row
#                     if blk.rows and len(blk.rows[0].cells) >= 2:
#                         return True
#         return False
#     except Exception:
#         return False

# # def extract_left_column_script_or_default(docx_path: str) -> tuple[str, bool]:
# #     doc = Document(docx_path)
# #     looks_two_col = any(
# #         isinstance(blk, Table) and _table_looks_like_narration_visuals(blk)
# #         for blk in _iter_block_items(doc)
# #     )
# #     if not looks_two_col:
# #         return _extract_docx_in_order(docx_path), False

# #     chunks: List[str] = []
# #     for blk in _iter_block_items(doc):
# #         if isinstance(blk, Paragraph):
# #             t = _text_from_paragraph(blk)
# #             if t.strip():
# #                 chunks.append(t.strip())
# #         elif isinstance(blk, Table):
# #             if _table_looks_like_narration_visuals(blk):
# #                 left = _text_from_table_first_column(blk)
# #                 if left.strip():
# #                     chunks.append(left.strip())
# #             else:
# #                 left = _text_from_table_first_column(blk)
# #                 if left.strip():
# #                     chunks.append(left.strip())
# #     return _normalize_text("\n\n".join(chunks)), True

# def extract_left_column_script_or_default(docx_path: str) -> tuple[str, bool]:
#     """
#     Always take paragraph text as-is, and for ANY table take ONLY the first column.
#     Returns (text, True) to indicate the left-column extractor was used.
#     """
#     doc = Document(docx_path)
#     chunks: List[str] = []

#     for blk in _iter_block_items(doc):
#         if isinstance(blk, Paragraph):
#             t = _text_from_paragraph(blk)
#             if t.strip():
#                 chunks.append(t.strip())
#         elif isinstance(blk, Table):
#             left = _text_from_table_first_column(blk)
#             if left.strip():
#                 chunks.append(left.strip())

#     return _normalize_text("\n\n".join(chunks)), True

# # -------------------- PDF extraction -------------------- #
# def _load_pdf_text(path: str) -> str:
#     if PdfReader is None:
#         return ""
#     try:
#         reader = PdfReader(path)
#         pages_text: List[str] = []
#         for page in reader.pages:
#             try:
#                 t = page.extract_text() or ""
#             except Exception:
#                 t = ""
#             pages_text.append(t)
#         return _normalize_text("\n\n".join(pages_text))
#     except Exception:
#         return ""

# # -------------------- Public loader -------------------- #
# def load_script_file(path: str) -> str:
#     ext = os.path.splitext(path)[1].lower()
#     if ext == ".txt":
#         with open(path, "r", encoding="utf-8", errors="ignore") as f:
#             return _normalize_text(f.read())
#     if ext == ".docx":
#         main, _ = extract_left_column_script_or_default(path)
#         tbx = _extract_textboxes_from_docx(path)
#         if tbx and tbx not in main:
#             main = (main + ("\n\n" if main else "") + tbx).strip()
#         return main
#     if ext == ".pdf":
#         return _load_pdf_text(path)
#     return ""

# # -------------------- JSON extraction helpers -------------------- #
# BEGIN_JSON_TOKEN = "BEGIN_JSON"
# END_JSON_TOKEN = "END_JSON"

# def _extract_between_tokens(text: str, start_token: str, end_token: str) -> Optional[str]:
#     if not text:
#         return None
#     i, j = text.find(start_token), text.rfind(end_token)
#     if i == -1 or j == -1 or j <= i:
#         return None
#     return text[i + len(start_token) : j].strip()

# def _extract_fenced_block(text: str) -> Optional[str]:
#     if not text:
#         return None
#     m = re.search(r"```json\s*(.*?)\s*```", text, flags=re.DOTALL | re.I)
#     if m:
#         return m.group(1).strip()
#     m = re.search(r"```\s*(.*?)\s*```", text, flags=re.DOTALL)
#     return m.group(1).strip() if m else None

# def _extract_balanced_json(text: str) -> Optional[str]:
#     if not text:
#         return None
#     start = text.find("{")
#     if start == -1:
#         return None
#     depth = 0
#     for i, ch in enumerate(text[start:], start):
#         if ch == "{":
#             depth += 1
#         elif ch == "}":
#             depth -= 1
#             if depth == 0:
#                 return text[start : i + 1]
#     return None

# def extract_review_json(model_output: str) -> Optional[Dict[str, Any]]:
#     if not model_output:
#         return None
#     candidates: List[str] = []
#     for fn in (_extract_between_tokens, _extract_fenced_block, _extract_balanced_json):
#         val = fn(model_output, BEGIN_JSON_TOKEN, END_JSON_TOKEN) if fn == _extract_between_tokens else fn(model_output)
#         if val:
#             candidates.append(val)
#     for s in candidates:
#         try:
#             return json.loads(s)
#         except Exception:
#             s2 = re.sub(r",\s*([\]}])", r"\1", s)
#             try:
#                 return json.loads(s2)
#             except Exception:
#                 continue
#     return None

# # -------------------- AOI normalization -------------------- #
# AOI_KEYS: List[str] = ["quote_verbatim", "issue", "fix", "why_this_helps"]

# def _clean_str(x: Any) -> str:
#     return str(x).strip() if x else ""

# def _coerce_aois(block: Dict[str, Any]) -> List[Dict[str, str]]:
#     raw = block.get("areas_of_improvement") or []
#     if not isinstance(raw, Iterable):
#         return []
#     out: List[Dict[str, str]] = []
#     for item in raw:
#         if not isinstance(item, dict):
#             continue
#         q = _clean_str(item.get("quote_verbatim") or item.get("quote") or item.get("line") or "")
#         issue = _clean_str(item.get("issue"))
#         fix = _clean_str(item.get("fix") or item.get("edit_suggestion"))
#         why = _clean_str(item.get("why_this_helps") or item.get("why"))
#         if not (q or issue or fix or why):
#             continue
#         out.append({
#             "quote_verbatim": q[:240] if q else "",
#             "issue": issue,
#             "fix": fix,
#             "why_this_helps": why,
#         })
#     return out

# def normalize_review_payload(data: Dict[str, Any]) -> Dict[str, Any]:
#     if not isinstance(data, dict):
#         return data
#     for k in ("strengths", "weaknesses", "suggestions", "drop_off_risks"):
#         if isinstance(data.get(k), list):
#             data[k] = [sanitize_editor_text(x) for x in data[k]]
#     if isinstance(data.get("viral_quotient"), str):
#         data["viral_quotient"] = sanitize_editor_text(data["viral_quotient"])
#     per = data.get("per_parameter") or {}
#     if isinstance(per, dict):
#         for _, block in per.items():
#             if not isinstance(block, dict):
#                 continue
#             block["areas_of_improvement"] = _coerce_aois(block)
#             for fld in ("explanation", "weakness", "suggestion", "summary"):
#                 if isinstance(block.get(fld), str):
#                     block[fld] = sanitize_editor_text(block[fld])
#             for fld in ("dominant_tense", "tense_target", "tense_consistency"):
#                 if isinstance(block.get(fld), str):
#                     block[fld] = sanitize_editor_text(block[fld])
#             for a in block.get("areas_of_improvement", []):
#                 for fld in ("quote_verbatim", "issue", "fix", "why_this_helps"):
#                     if isinstance(a.get(fld), str):
#                         a[fld] = sanitize_editor_text(a[fld])
#     return data








#####################################





# utils1.py — helpers shared by app_grammarly_ui.py and run_review_multi.py

from __future__ import annotations
import os, re, json, zipfile, xml.etree.ElementTree as ET
from typing import Any, Dict, List, Optional, Iterable, Tuple

# Optional PDF support
try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None  # type: ignore

# DOCX helpers
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.table import Table

# -------------------- UI order for parameters -------------------- #
PARAM_ORDER: List[str] = [
    "Suspense Building",
    "Language/Tone",
    "Intro + Main Hook/Cliffhanger",
    "Story Structure + Flow",
    "Pacing",
    "Mini-Hooks (30–60s)",
    "Outro (Ending)",
    "Grammar & Spelling",  # NEW parameter
]

# -------------------- Sanitizers -------------------- #
_EMOJI_RE = re.compile(
    r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
    flags=re.UNICODE,
)

def sanitize_editor_text(s: Optional[str]) -> str:
    if not s:
        return ""
    t = str(s)
    # Remove accidental "Decision:" or "Score: 8" echoes
    t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
    t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
    # Strip bullets/numbers
    t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•|\*)\s*', '', t, flags=re.M)
    # Remove emojis and compress whitespace
    t = _EMOJI_RE.sub('', t)
    t = re.sub(r'[ \t]+', ' ', t)
    t = re.sub(r'\n{3,}', '\n\n', t)
    return t.strip()

def _strip_bom(s: str) -> str:
    return s.lstrip("\ufeff") if s and s.startswith("\ufeff") else s

def _normalize_text(s: str) -> str:
    if not s:
        return ""
    s = _strip_bom(s)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = s.replace("\xa0", " ")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()

# -------------------- DOCX traversal & base extraction -------------------- #
def _iter_block_items(document: Document):
    parent = document.element.body
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)

def _paragraph_text_with_breaks(p: Paragraph) -> str:
    """
    Paragraph text that preserves hard line breaks (<w:br/>).
    """
    parts: List[str] = []
    for run in p.runs:
        # run.text (no <w:br/>)
        if run.text:
            parts.append(run.text)
        # add newlines for each <w:br/> in the run
        try:
            brs = run._r.findall(".//w:br", run._r.nsmap)  # type: ignore[attr-defined]
            if brs:
                parts.append("\n" * len(brs))
        except Exception:
            pass
    txt = "".join(parts)
    return re.sub(r"\n{3,}", "\n\n", txt)

def _text_from_paragraph(p: Paragraph) -> str:
    return _paragraph_text_with_breaks(p) or ""

def _text_from_table(tbl: Table) -> str:
    """
    Generic table-to-text (used by full flatten fallback).
    """
    lines: List[str] = []
    for row in tbl.rows:
        row_cells: List[str] = []
        for cell in row.cells:
            cell_bits: List[str] = []
            for para in cell.paragraphs:
                cell_bits.append(_paragraph_text_with_breaks(para))
            # nested tables
            for nt in cell._tc.iterchildren():
                if isinstance(nt, CT_Tbl):
                    t = Table(nt, cell._parent)
                    nested = _text_from_table(t)
                    if nested:
                        cell_bits.append(nested)
            row_cells.append("\n".join([cl for cl in cell_bits if cl]))
        line = "  ".join([c for c in row_cells if c])
        lines.append(line)
    return "\n".join([ln for ln in lines if ln.strip()])

def _extract_docx_in_order(docx_path: str) -> str:
    """
    Flatten the whole document in reading order (paragraphs + all tables).
    """
    doc = Document(docx_path)
    chunks: List[str] = []
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            t = _text_from_paragraph(block)
            if t is not None:
                chunks.append(t)
        elif isinstance(block, Table):
            t = _text_from_table(block)
            if t is not None:
                chunks.append(t)
    return _normalize_text("\n".join(chunks))

# -------------------- Text boxes (shapes) -------------------- #
def _extract_textboxes_from_docx(docx_path: str) -> str:
    """
    Reads w:txbxContent from document.xml to capture text in drawing shapes/textboxes.
    """
    try:
        with zipfile.ZipFile(docx_path) as z:
            xml = z.read("word/document.xml")
    except Exception:
        return ""
    root = ET.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    bits: List[str] = []
    for txbx in root.findall(".//w:txbxContent", ns):
        paras = txbx.findall(".//w:p", ns)
        p_texts: List[str] = []
        for p in paras:
            texts: List[str] = [t.text or "" for t in p.findall(".//w:t", ns)]
            p_texts.append("".join(texts))
        if p_texts:
            bits.append("\n".join(p_texts))
    return _normalize_text("\n\n".join(bits))

# -------------------- LEFT-COLUMN (VO) extractor -------------------- #
def _run_text_with_breaks(run) -> str:
    """
    Return run text + explicit '\n' for each <w:br/>.
    """
    t = run.text or ""
    try:
        brs = run._r.findall(".//w:br", run._r.nsmap)  # type: ignore[attr-defined]
        if brs:
            t = t + ("\n" * len(brs))
    except Exception:
        pass
    return t

def _text_from_cell_paragraphs(cell) -> str:
    """
    Cell text preserving paragraph boundaries and hard line breaks.
    """
    parts: List[str] = []
    for i, p in enumerate(cell.paragraphs):
        run_txt = "".join(_run_text_with_breaks(r) for r in p.runs)
        parts.append(run_txt)
        if i != len(cell.paragraphs) - 1:
            parts.append("\n")
    # Keep internal structure; trim trailing newlines only
    return "".join(parts).strip("\n")

def _table_col_count(tbl: Table) -> int:
    try:
        return len(tbl.rows[0].cells)
    except Exception:
        return 0

def _is_vo_visuals_header_row(tbl: Table) -> bool:
    """
    Header row like: [Voice Over | Visuals] (case-insensitive, variants supported).
    """
    try:
        if _table_col_count(tbl) < 2:
            return False
        first = _text_from_cell_paragraphs(tbl.rows[0].cells[0]).strip().lower()
        second = _text_from_cell_paragraphs(tbl.rows[0].cells[1]).strip().lower()
        return (
            ("voice over" in first or first == "vo" or "narration" in first)
            and ("visual" in second or "visuals" in second or "clip" in second or "time" in second)
        )
    except Exception:
        return False

def _table_looks_like_vo(tbl: Table) -> bool:
    """
    Heuristics to decide if a table participates in VO:
      - 2+ columns with a VO|VISUALS header, or
      - Left/only column has real content in at least 2 of the first 6 rows.
    """
    cols = _table_col_count(tbl)
    if cols == 0:
        return False
    if cols >= 2 and _is_vo_visuals_header_row(tbl):
        return True

    nonempty = 0
    sample_rows = min(6, len(tbl.rows))
    for r in range(sample_rows):
        try:
            c0 = tbl.rows[r].cells[0]
            if _text_from_cell_paragraphs(c0).strip():
                nonempty += 1
        except Exception:
            continue
    return nonempty >= 2

def _extract_table_first_or_only_col(tbl: Table) -> str:
    """
    Gather VO text from a table:
      - if >= 2 columns: take column 0, skipping header row if VO|VISUALS
      - if 1 column: take that column (VO-only sections)
    """
    cols = _table_col_count(tbl)
    if cols == 0:
        return ""
    out: List[str] = []
    start_row = 1 if (cols >= 2 and _is_vo_visuals_header_row(tbl)) else 0
    for r in range(start_row, len(tbl.rows)):
        cell = tbl.rows[r].cells[0]
        t = _text_from_cell_paragraphs(cell)
        if t.strip():
            out.append(t.strip())
    return "\n\n".join(out).strip()

def _has_two_plus_col_table(doc: Document) -> bool:
    """Return True if the DOCX has at least one table with >= 2 columns."""
    try:
        for blk in _iter_block_items(doc):
            if isinstance(blk, Table):
                try:
                    if len(blk.columns) >= 2:
                        return True
                except Exception:
                    if blk.rows and len(blk.rows[0].cells) >= 2:
                        return True
        return False
    except Exception:
        return False

def extract_left_column_script_or_default(docx_path: str) -> Tuple[str, bool]:
    """
    If the document contains any VO-looking table, return ONLY the VO text
    (first column for 2+ columns, the only column for 1 column) from ALL tables,
    preserving hard line breaks. Otherwise, return the whole document flattened.

    Returns:
      (text, used_left_extractor_flag)
    """
    doc = Document(docx_path)

    # Detect VO layout anywhere in the doc
    any_vo = any(isinstance(blk, Table) and _table_looks_like_vo(blk) for blk in _iter_block_items(doc))
    if not any_vo:
        # Not a VO layout → fall back to full flatten
        return _extract_docx_in_order(docx_path), False

    # VO mode: aggregate VO from every table (handles VO-only 1-column sections too)
    chunks: List[str] = []
    for blk in _iter_block_items(doc):
        if isinstance(blk, Table):
            coltext = _extract_table_first_or_only_col(blk)
            if coltext:
                chunks.append(coltext)

    if not chunks:
        # Defensive fallback if the heuristics failed
        return _extract_docx_in_order(docx_path), False

    text = "\n\n".join(chunks)
    text = text.replace("\r\n", "\n")
    # Trim right spaces per line but preserve structure
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip(), True

# -------------------- PDF extraction -------------------- #
def _load_pdf_text(path: str) -> str:
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(path)
        pages_text: List[str] = []
        for page in reader.pages:
            try:
                t = page.extract_text() or ""
            except Exception:
                t = ""
            pages_text.append(t)
        return _normalize_text("\n\n".join(pages_text))
    except Exception:
        return ""

# -------------------- Public loader -------------------- #
def load_script_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return _normalize_text(f.read())
    if ext == ".docx":
        main, _ = extract_left_column_script_or_default(path)
        tbx = _extract_textboxes_from_docx(path)
        if tbx and tbx not in main:
            main = (main + ("\n\n" if main else "") + tbx).strip()
        return main
    if ext == ".pdf":
        return _load_pdf_text(path)
    return ""

# -------------------- JSON extraction helpers -------------------- #
BEGIN_JSON_TOKEN = "BEGIN_JSON"
END_JSON_TOKEN = "END_JSON"

def _extract_between_tokens(text: str, start_token: str, end_token: str) -> Optional[str]:
    if not text:
        return None
    i, j = text.find(start_token), text.rfind(end_token)
    if i == -1 or j == -1 or j <= i:
        return None
    return text[i + len(start_token) : j].strip()

def _extract_fenced_block(text: str) -> Optional[str]:
    if not text:
        return None
    m = re.search(r"```json\s*(.*?)\s*```", text, flags=re.DOTALL | re.I)
    if m:
        return m.group(1).strip()
    m = re.search(r"```\s*(.*?)\s*```", text, flags=re.DOTALL)
    return m.group(1).strip() if m else None

def _extract_balanced_json(text: str) -> Optional[str]:
    if not text:
        return None
    start = text.find("{")
    if start == -1:
        return None
    depth = 0
    for i, ch in enumerate(text[start:], start):
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[start : i + 1]
    return None

def extract_review_json(model_output: str) -> Optional[Dict[str, Any]]:
    if not model_output:
        return None
    candidates: List[str] = []
    for fn in (_extract_between_tokens, _extract_fenced_block, _extract_balanced_json):
        val = fn(model_output, BEGIN_JSON_TOKEN, END_JSON_TOKEN) if fn == _extract_between_tokens else fn(model_output)
        if val:
            candidates.append(val)
    for s in candidates:
        try:
            return json.loads(s)
        except Exception:
            s2 = re.sub(r",\s*([\]}])", r"\1", s)
            try:
                return json.loads(s2)
            except Exception:
                continue
    return None

# -------------------- AOI normalization -------------------- #
AOI_KEYS: List[str] = ["quote_verbatim", "issue", "fix", "why_this_helps"]

def _clean_str(x: Any) -> str:
    return str(x).strip() if x else ""

def _coerce_aois(block: Dict[str, Any]) -> List[Dict[str, str]]:
    raw = block.get("areas_of_improvement") or []
    if not isinstance(raw, Iterable):
        return []
    out: List[Dict[str, str]] = []
    for item in raw:
        if not isinstance(item, dict):
            continue
        q = _clean_str(item.get("quote_verbatim") or item.get("quote") or item.get("line") or "")
        issue = _clean_str(item.get("issue"))
        fix = _clean_str(item.get("fix") or item.get("edit_suggestion"))
        why = _clean_str(item.get("why_this_helps") or item.get("why"))
        if not (q or issue or fix or why):
            continue
        out.append({
            "quote_verbatim": q[:240] if q else "",
            "issue": issue,
            "fix": fix,
            "why_this_helps": why,
        })
    return out

def normalize_review_payload(data: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(data, dict):
        return data
    for k in ("strengths", "weaknesses", "suggestions", "drop_off_risks"):
        if isinstance(data.get(k), list):
            data[k] = [sanitize_editor_text(x) for x in data[k]]
    if isinstance(data.get("viral_quotient"), str):
        data["viral_quotient"] = sanitize_editor_text(data["viral_quotient"])
    per = data.get("per_parameter") or {}
    if isinstance(per, dict):
        for _, block in per.items():
            if not isinstance(block, dict):
                continue
            block["areas_of_improvement"] = _coerce_aois(block)
            for fld in ("explanation", "weakness", "suggestion", "summary"):
                if isinstance(block.get(fld), str):
                    block[fld] = sanitize_editor_text(block[fld])
            for fld in ("dominant_tense", "tense_target", "tense_consistency"):
                if isinstance(block.get(fld), str):
                    block[fld] = sanitize_editor_text(block[fld])
            for a in block.get("areas_of_improvement", []):
                for fld in ("quote_verbatim", "issue", "fix", "why_this_helps"):
                    if isinstance(a.get(fld), str):
                        a[fld] = sanitize_editor_text(a[fld])
    return data
