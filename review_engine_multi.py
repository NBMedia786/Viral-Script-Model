

# from __future__ import annotations

# import os
# import re
# import json
# import time
# from typing import Any, Dict, List, Optional

# from dotenv import load_dotenv
# from langchain_google_genai import ChatGoogleGenerativeAI
# from pydantic import BaseModel, conint  # NEW: for structured aggregator output

# # NEW: use the shared normalizer/sanitizer so artifacts are cleaned before UI
# from utils1 import normalize_review_payload

# # -----------------------------------------------------------------------------#
# # Env / model
# # -----------------------------------------------------------------------------#
# load_dotenv()

# def _require_api_key() -> None:
#     if not os.getenv("GOOGLE_API_KEY"):
#         raise EnvironmentError("GOOGLE_API_KEY not found in environment/.env")

# def _make_llm(temperature: float) -> ChatGoogleGenerativeAI:
#     model = os.getenv("GEMINI_MODEL") or "gemini-2.5-flash"
#     # Deterministic-ish defaults; keep kwargs minimal for Gemini
#     return ChatGoogleGenerativeAI(model=model, temperature=temperature, top_p=0.0, top_k=1)

# # -----------------------------------------------------------------------------#
# # File helpers
# # -----------------------------------------------------------------------------#
# def _strip_bom(s: str) -> str:
#     return s.lstrip("\ufeff") if s and s.startswith("\ufeff") else s

# def _read_text(path: str) -> str:
#     with open(path, "r", encoding="utf-8") as f:
#         return _strip_bom(f.read())

# def _load_prompt(prompts_dir: str, n: int) -> str:
#     """
#     Load 'n.yaml' or 'n.yml'. If the file uses the 'content: |' convention,
#     extract the content block; otherwise, return the full file text.
#     """
#     for ext in ("yaml", "yml"):
#         p = os.path.join(prompts_dir, f"{n}.{ext}")
#         if os.path.isfile(p):
#             raw = _read_text(p).strip()
#             # Try to capture 'content: |' block; if not present, use whole file
#             m = re.search(r"^\s*content\s*:\s*\|?\s*(.*)$", raw, flags=re.S | re.I)
#             txt = (m.group(1) if m else raw).replace("\r\n", "\n")
#             return txt.strip()
#     raise FileNotFoundError(f"Prompt {n} (yaml/yml) not found in {os.path.abspath(prompts_dir)}")

# def _inject(template: str, **kwargs: str) -> str:
#     out = template
#     for k, v in kwargs.items():
#         out = out.replace("{" + k + "}", v)
#     return out

# # -----------------------------------------------------------------------------#
# # JSON extraction (for specialists that return plain JSON)
# # -----------------------------------------------------------------------------#
# _BEGIN = "BEGIN_JSON"
# _END = "END_JSON"

# def _between_tokens(text: str, start: str, end: str) -> Optional[str]:
#     i = text.find(start)
#     j = text.rfind(end)
#     if i == -1 or j == -1 or j <= i:
#         return None
#     return text[i + len(start): j].strip()

# def _fenced(text: str) -> Optional[str]:
#     m = re.search(r"```json\s*(.*?)\s*```", text, flags=re.S | re.I)
#     if m: return m.group(1).strip()
#     m = re.search(r"```\s*(.*?)\s*```", text, flags=re.S)
#     if m: return m.group(1).strip()
#     return None

# def _balanced(text: str) -> Optional[str]:
#     start = text.find("{")
#     if start == -1: return None
#     depth = 0
#     for i in range(start, len(text)):
#         ch = text[i]
#         if ch == "{":
#             depth += 1
#         elif ch == "}":
#             depth -= 1
#             if depth == 0:
#                 return text[start:i+1]
#     return None

# def _parse_json(text: str) -> Dict[str, Any]:
#     """
#     Try multiple strategies to extract valid JSON. Tolerates trailing commas.
#     """
#     for candidate in filter(None, (_between_tokens(text, _BEGIN, _END),
#                                   _fenced(text), _balanced(text))):
#         try:
#             return json.loads(candidate)
#         except Exception:
#             try:
#                 fixed = re.sub(r",\s*([\]}])", r"\1", candidate)
#                 return json.loads(fixed)
#             except Exception:
#                 continue
#     raise ValueError("Could not parse JSON from model output")

# # -----------------------------------------------------------------------------#
# # Conversions → legacy block for UI
# # -----------------------------------------------------------------------------#
# def _coerce_explanation(raw: Dict[str, Any]) -> str:
#     """
#     Accept 'explanation' (string) or 'explanation_bullets' (list[str]).
#     """
#     explanation = raw.get("explanation", "")
#     if isinstance(explanation, str) and explanation.strip():
#         return explanation.strip()
#     bullets = raw.get("explanation_bullets", [])
#     if isinstance(bullets, list) and bullets:
#         parts = [str(b).strip().rstrip(".") + "." for b in bullets if str(b).strip()]
#         return " ".join(parts)[:1400]
#     return ""

# def _to_legacy_param_block(raw: Dict[str, Any]) -> Dict[str, Any]:
#     """
#     Normalize a specialist block to the legacy shape expected by the UI.
#     - Raw 'extractions' are intentionally suppressed (we only use AOIs for inline highlights).
#     - AOIs pass through untouched (4-field schema, unlimited).
#     - 'summary' (plain human note) is passed through for right-panel display.
#     """
#     score = int(raw.get("score", 0) or 0)
#     explanation = _coerce_explanation(raw)
#     weakness = str(raw.get("weakness", "") or "").strip() or "Not present"

#     # Keep single suggestion for legacy UI AND preserve full list if present
#     suggestion = str(raw.get("suggestion", "") or "").strip()
#     suggestions_list: List[str] = []
#     if not suggestion and isinstance(raw.get("suggestions"), list) and raw["suggestions"]:
#         suggestions_list = [str(s).strip() for s in raw["suggestions"] if str(s).strip()]
#         suggestion = suggestions_list[0] if suggestions_list else ""
#     elif isinstance(raw.get("suggestions"), list):
#         suggestions_list = [str(s).strip() for s in raw["suggestions"] if str(s).strip()]
#     suggestion = suggestion or "Not present"

#     # Hide raw extractions in UI—AOIs only
#     ex: List[str] = []
#     aoi = raw.get("areas_of_improvement") or []

#     block: Dict[str, Any] = {
#         "extractions": ex,  # kept in schema but always empty
#         "score": score,
#         "explanation": explanation,
#         "weakness": weakness,
#         "suggestion": suggestion,
#         "areas_of_improvement": aoi,
#         "summary": str(raw.get("summary", "") or "").strip(),
#     }
#     if suggestions_list:
#         block["suggestions_list"] = suggestions_list
#     return block

# # -----------------------------------------------------------------------------#
# # Display-name mapping for UI & aggregator
# # -----------------------------------------------------------------------------#
# DISPLAY_BY_INDEX = {
#     1: "Suspense Building",
#     2: "Language/Tone",
#     3: "Intro + Main Hook/Cliffhanger",
#     4: "Story Structure + Flow",
#     5: "Pacing",
#     6: "Mini-Hooks (30–60s)",
#     7: "Outro (Ending)",
#     # 8 = global preamble (loaded, not called)
#     9: "Overall Summary (Aggregator)",
# }

# # -----------------------------------------------------------------------------#
# # LLM invoke with retries
# # -----------------------------------------------------------------------------#
# def _invoke_with_retries(llm: ChatGoogleGenerativeAI, prompt: str, tries: int = 3, base_delay: float = 0.8):
#     last_err = None
#     for k in range(tries):
#         try:
#             return llm.invoke(prompt)
#         except Exception as e:
#             last_err = e
#             if k < tries - 1:
#                 time.sleep(base_delay * (2 ** k))
#             else:
#                 raise
#     raise last_err  # type: ignore

# # -----------------------------------------------------------------------------#
# # Aggregator structured schema (model returns this)
# # -----------------------------------------------------------------------------#
# class AggregatorAll(BaseModel):
#     overall_rating: conint(ge=1, le=10) 
#     strengths: List[str]
#     weaknesses: List[str]
#     suggestions: List[str]
#     drop_off_risks: List[str]
#     viral_quotient: str

# # -----------------------------------------------------------------------------#
# # Core runner
# # -----------------------------------------------------------------------------#
# def run_review_multi(
#     script_text: str,
#     prompts_dir: str = "prompts",
#     temperature: float = 0.0,  # default locked to 0.0
#     include_commentary: bool = False,  # kept for API parity; ignored
# ) -> str:
#     """
#     Execute prompts 1..7 and 9, prepend prompts/8.yaml as a global preamble to each call,
#     convert to legacy shape, and return BEGIN_JSON ... END_JSON for the UI.
#     """
#     _require_api_key()
#     llm = _make_llm(temperature)

#     # Load global preamble (8.yaml) if present
#     try:
#         global_preamble = _load_prompt(prompts_dir, 8).strip()
#         if global_preamble:
#             global_preamble += "\n\n"
#     except FileNotFoundError:
#         global_preamble = ""

#     # 1..7 specialists
#     scores: Dict[str, int] = {}
#     per_parameter: Dict[str, Dict[str, Any]] = {}

#     for i in range(1, 7 + 1):
#         name = DISPLAY_BY_INDEX[i]
#         tmpl = _load_prompt(prompts_dir, i)
#         prompt_body = _inject(tmpl, script=script_text)
#         prompt = f"{global_preamble}{prompt_body}"

#         try:
#             resp = _invoke_with_retries(llm, prompt)
#             raw_text = getattr(resp, "content", "") or ""
#             data = _parse_json(raw_text)
#         except Exception as e:
#             short = (str(e) or "unknown").strip()
#             raise RuntimeError(f"JSON parse failed on prompt {i} ({name}). Error: {short}")

#         block = _to_legacy_param_block(data)
#         scores[name] = int(block.get("score", 0))
#         per_parameter[name] = block

#     # Build evidence for aggregator (legacy shape only)
#     evidence = {"scores": scores, "per_parameter": per_parameter}
#     evidence_json = json.dumps(evidence, ensure_ascii=False)

#     # 9 = merged meta-synthesis + aggregator (MODEL decides overall_rating)
#     tmpl9 = _load_prompt(prompts_dir, 9)
#     prompt9_body = _inject(
#         tmpl9,
#         evidence_json=evidence_json,
#         script=script_text,
#     )
#     prompt9 = f"{global_preamble}{prompt9_body}"

#     # Ask for structured output (enforces keys and types)
#     try:
#         llm_aggr = _make_llm(temperature).with_structured_output(AggregatorAll)
#         agg: AggregatorAll = llm_aggr.invoke(prompt9)
#     except Exception as e:
#         short = (str(e) or "unknown").strip()
#         raise RuntimeError(f"Aggregator failed on prompt 9. Error: {short}")

#     # Build final payload: use model-decided overall_rating
#     final_payload: Dict[str, Any] = {
#         "scores": scores,
#         "per_parameter": per_parameter,
#         "overall_rating": int(agg.overall_rating),
#         "strengths": agg.strengths,
#         "weaknesses": agg.weaknesses,
#         "suggestions": agg.suggestions,
#         "drop_off_risks": agg.drop_off_risks,
#         "viral_quotient": agg.viral_quotient,
#     }

#     # Normalize & sanitize everything before returning
#     final_payload = normalize_review_payload(final_payload)

#     return _wrap_json(final_payload)

# def _wrap_json(payload: Dict[str, Any]) -> str:
#     return f"{_BEGIN}\n{json.dumps(payload, ensure_ascii=False)}\n{_END}\n"




















# # run_review_multi.py — S3-first prompt loader (Runpod), no local prompts required
# from __future__ import annotations

# import os
# import re
# import json
# import time
# from typing import Any, Dict, List, Optional

# from dotenv import load_dotenv
# from langchain_google_genai import ChatGoogleGenerativeAI
# from pydantic import BaseModel, conint  # for structured aggregator output

# # Use the shared normalizer/sanitizer so artifacts are cleaned before UI
# from utils1 import normalize_review_payload

# # -----------------------------
# # Runpod S3 client (boto3)
# # -----------------------------
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError

# load_dotenv()

# _RP_ENDPOINT = os.getenv("RUNPOD_S3_ENDPOINT", "").strip()
# _RP_BUCKET   = os.getenv("RUNPOD_S3_BUCKET", "").strip()
# _RP_REGION   = os.getenv("RUNPOD_S3_REGION", "").strip()

# def _s3_enabled() -> bool:
#     return bool(_RP_ENDPOINT and _RP_BUCKET)

# _S3_CLIENT = None
# def _s3_client():
#     global _S3_CLIENT
#     if _S3_CLIENT is not None:
#         return _S3_CLIENT
#     if not _s3_enabled():
#         return None
#     _S3_CLIENT = boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         region_name=_RP_REGION or None,
#         config=Config(s3={"addressing_style": "path"})
#     )
#     return _S3_CLIENT

# def _read_s3_text(key: str) -> Optional[str]:
#     """Return object body as UTF-8 text or None."""
#     cli = _s3_client()
#     if not cli:
#         return None
#     try:
#         obj = cli.get_object(Bucket=_RP_BUCKET, Key=key)
#         return obj["Body"].read().decode("utf-8", errors="ignore")
#     except ClientError:
#         return None
#     except Exception:
#         return None

# def _join_s3_key(prefix: str, filename: str) -> str:
#     return f"{prefix.rstrip('/')}/{filename.lstrip('/')}"

# # -----------------------------------------------------------------------------#
# # Env / model
# # -----------------------------------------------------------------------------#
# def _require_api_key() -> None:
#     if not os.getenv("GOOGLE_API_KEY"):
#         raise EnvironmentError("GOOGLE_API_KEY not found in environment/.env")

# def _make_llm(temperature: float) -> ChatGoogleGenerativeAI:
#     """
#     Create a Gemini chat LLM with minimal, deterministic-ish defaults.
#     """
#     model = os.getenv("GEMINI_MODEL") or "gemini-2.5-flash"
#     return ChatGoogleGenerativeAI(
#         model=model,
#         temperature=temperature,
#         top_p=0.0,
#         top_k=1,
#     )

# # -----------------------------------------------------------------------------#
# # Prompt loaders (S3-first; local fallback)
# # -----------------------------------------------------------------------------#
# def _strip_bom(s: str) -> str:
#     return s.lstrip("\ufeff") if s and s.startswith("\ufeff") else s

# def _read_text_local(path: str) -> str:
#     with open(path, "r", encoding="utf-8") as f:
#         return _strip_bom(f.read())

# def _extract_content_block(raw: str) -> str:
#     """
#     If the YAML uses:
#       content: |
#         ...
#     extract that block; else return whole file.
#     """
#     raw = raw.strip()
#     m = re.search(r"^\s*content\s*:\s*\|?\s*(.*)$", raw, flags=re.S | re.I)
#     txt = (m.group(1) if m else raw).replace("\r\n", "\n")
#     return txt.strip()

# def _load_prompt_s3(prompts_prefix: str, n: int) -> Optional[str]:
#     """
#     Try S3: {prompts_prefix}/{n}.yaml then {n}.yml
#     Returns the prompt text or None if not found.
#     """
#     if not _s3_enabled():
#         return None
#     for ext in ("yaml", "yml"):
#         key = _join_s3_key(prompts_prefix, f"{n}.{ext}")
#         txt = _read_s3_text(key)
#         if txt:
#             return _extract_content_block(txt)
#     return None

# def _load_prompt_local(prompts_dir: str, n: int) -> Optional[str]:
#     """
#     Local fallback if S3 not configured.
#     """
#     for ext in ("yaml", "yml"):
#         p = os.path.join(prompts_dir, f"{n}.{ext}")
#         if os.path.isfile(p):
#             raw = _read_text_local(p)
#             return _extract_content_block(raw)
#     return None

# def _load_prompt(prompts_dir_or_prefix: str, n: int) -> str:
#     """
#     Unified loader: prefer S3 (Runpod) if configured; else local.
#     `prompts_dir_or_prefix` should be something like "Scriptmodel/prompts".
#     """
#     # If the caller accidentally passes an s3:// url, strip it to a key prefix:
#     prefix = prompts_dir_or_prefix
#     if prefix.startswith("s3://"):
#         # s3://bucket/... => strip "s3://<bucket>/" so we keep only the key
#         parts = prefix.split("/", 3)
#         prefix = parts[3] if len(parts) > 3 else ""

#     # S3 first
#     txt = _load_prompt_s3(prefix, n)
#     if txt:
#         return txt

#     # Local fallback (only if S3 off or missing file)
#     txt = _load_prompt_local(prompts_dir_or_prefix, n)
#     if txt:
#         return txt

#     where = f"S3({_RP_BUCKET}:{prefix})" if _s3_enabled() else os.path.abspath(prompts_dir_or_prefix)
#     raise FileNotFoundError(f"Prompt {n} (yaml/yml) not found in {where}")

# def _inject(template: str, **kwargs: str) -> str:
#     out = template
#     for k, v in kwargs.items():
#         out = out.replace("{" + k + "}", v)
#     return out

# # -----------------------------------------------------------------------------#
# # JSON extraction (for specialists that return plain JSON)
# # -----------------------------------------------------------------------------#
# _BEGIN = "BEGIN_JSON"
# _END = "END_JSON"

# def _between_tokens(text: str, start: str, end: str) -> Optional[str]:
#     i = text.find(start)
#     j = text.rfind(end)
#     if i == -1 or j == -1 or j <= i:
#         return None
#     return text[i + len(start): j].strip()

# def _fenced(text: str) -> Optional[str]:
#     m = re.search(r"```json\s*(.*?)\s*```", text, flags=re.S | re.I)
#     if m: return m.group(1).strip()
#     m = re.search(r"```\s*(.*?)\s*```", text, flags=re.S)
#     if m: return m.group(1).strip()
#     return None

# def _balanced(text: str) -> Optional[str]:
#     start = text.find("{")
#     if start == -1: return None
#     depth = 0
#     for i in range(start, len(text)):
#         ch = text[i]
#         if ch == "{":
#             depth += 1
#         elif ch == "}":
#             depth -= 1
#             if depth == 0:
#                 return text[start:i+1]
#     return None

# def _parse_json(text: str) -> Dict[str, Any]:
#     """
#     Try multiple strategies to extract valid JSON. Tolerates trailing commas.
#     """
#     for candidate in filter(None, (_between_tokens(text, _BEGIN, _END),
#                                   _fenced(text), _balanced(text))):
#         try:
#             return json.loads(candidate)
#         except Exception:
#             try:
#                 fixed = re.sub(r",\s*([\]}])", r"\1", candidate)
#                 return json.loads(fixed)
#             except Exception:
#                 continue
#     raise ValueError("Could not parse JSON from model output")

# # -----------------------------------------------------------------------------#
# # Conversions → legacy block for UI
# # -----------------------------------------------------------------------------#
# def _coerce_explanation(raw: Dict[str, Any]) -> str:
#     """
#     Accept 'explanation' (string) or 'explanation_bullets' (list[str]).
#     """
#     explanation = raw.get("explanation", "")
#     if isinstance(explanation, str) and explanation.strip():
#         return explanation.strip()
#     bullets = raw.get("explanation_bullets", [])
#     if isinstance(bullets, list) and bullets:
#         parts = [str(b).strip().rstrip(".") + "." for b in bullets if str(b).strip()]
#         return " ".join(parts)[:1400]
#     return ""

# def _to_legacy_param_block(raw: Dict[str, Any]) -> Dict[str, Any]:
#     """
#     Normalize a specialist block to the legacy shape expected by the UI.
#     - Raw 'extractions' are intentionally suppressed (we only use AOIs for inline highlights).
#     - AOIs pass through untouched (4-field schema, unlimited).
#     - 'summary' (plain human note) is passed through for right-panel display.
#     """
#     score = int(raw.get("score", 0) or 0)
#     explanation = _coerce_explanation(raw)
#     weakness = str(raw.get("weakness", "") or "").strip() or "Not present"

#     # Keep single suggestion for legacy UI AND preserve full list if present
#     suggestion = str(raw.get("suggestion", "") or "").strip()
#     suggestions_list: List[str] = []
#     if not suggestion and isinstance(raw.get("suggestions"), list) and raw["suggestions"]:
#         suggestions_list = [str(s).strip() for s in raw["suggestions"] if str(s).strip()]
#         suggestion = suggestions_list[0] if suggestions_list else ""
#     elif isinstance(raw.get("suggestions"), list):
#         suggestions_list = [str(s).strip() for s in raw["suggestions"] if str(s).strip()]
#     suggestion = suggestion or "Not present"

#     # Hide raw extractions in UI—AOIs only
#     ex: List[str] = []
#     aoi = raw.get("areas_of_improvement") or []

#     block: Dict[str, Any] = {
#         "extractions": ex,  # kept in schema but always empty
#         "score": score,
#         "explanation": explanation,
#         "weakness": weakness,
#         "suggestion": suggestion,
#         "areas_of_improvement": aoi,
#         "summary": str(raw.get("summary", "") or "").strip(),
#     }
#     if suggestions_list:
#         block["suggestions_list"] = suggestions_list
#     return block

# # -----------------------------------------------------------------------------#
# # Display-name mapping for UI & aggregator
# # -----------------------------------------------------------------------------#
# DISPLAY_BY_INDEX = {
#     1: "Suspense Building",
#     2: "Language/Tone",
#     3: "Intro + Main Hook/Cliffhanger",
#     4: "Story Structure + Flow",
#     5: "Pacing",
#     6: "Mini-Hooks (30–60s)",
#     7: "Outro (Ending)",
#     # 8 = global preamble (loaded, not called)
#     9: "Overall Summary (Aggregator)",
# }

# # -----------------------------------------------------------------------------#
# # LLM invoke with retries
# # -----------------------------------------------------------------------------#
# def _invoke_with_retries(llm: ChatGoogleGenerativeAI, prompt: str, tries: int = 3, base_delay: float = 0.8):
#     last_err = None
#     for k in range(tries):
#         try:
#             return llm.invoke(prompt)
#         except Exception as e:
#             last_err = e
#             if k < tries - 1:
#                 time.sleep(base_delay * (2 ** k))
#             else:
#                 raise
#     raise last_err  # type: ignore

# # -----------------------------------------------------------------------------#
# # Aggregator structured schema (model returns this)
# # -----------------------------------------------------------------------------#
# class AggregatorAll(BaseModel):
#     overall_rating: conint(ge=1, le=10)
#     strengths: List[str]
#     weaknesses: List[str]
#     suggestions: List[str]
#     drop_off_risks: List[str]
#     viral_quotient: str

# # -----------------------------------------------------------------------------#
# # Core runner
# # -----------------------------------------------------------------------------#
# def run_review_multi(
#     script_text: str,
#     prompts_dir: str = "Scriptmodel/prompts",  # <-- treat as S3 prefix now
#     temperature: float = 0.0,
#     include_commentary: bool = False,  # kept for API parity; ignored
# ) -> str:
#     """
#     Execute prompts 1..7 and 9, prepend prompts/8.yaml as a global preamble to each call,
#     convert to legacy shape, and return BEGIN_JSON ... END_JSON for the UI.

#     NOTE: `prompts_dir` is treated as an S3 prefix (e.g., "Scriptmodel/prompts") when
#     RUNPOD_S3_* env vars are set. If S3 is not configured, it falls back to local files.
#     """
#     _require_api_key()
#     llm = _make_llm(temperature)

#     # Load global preamble (8.yaml) if present
#     try:
#         global_preamble = _load_prompt(prompts_dir, 8).strip()
#         if global_preamble:
#             global_preamble += "\n\n"
#     except FileNotFoundError:
#         global_preamble = ""

#     # 1..7 specialists
#     scores: Dict[str, int] = {}
#     per_parameter: Dict[str, Dict[str, Any]] = {}

#     for i in range(1, 7 + 1):
#         name = DISPLAY_BY_INDEX[i]
#         tmpl = _load_prompt(prompts_dir, i)
#         prompt_body = _inject(tmpl, script=script_text)
#         prompt = f"{global_preamble}{prompt_body}"

#         try:
#             resp = _invoke_with_retries(llm, prompt)
#             raw_text = getattr(resp, "content", "") or ""
#             data = _parse_json(raw_text)
#         except Exception as e:
#             short = (str(e) or "unknown").strip()
#             raise RuntimeError(f"JSON parse failed on prompt {i} ({name}). Error: {short}")

#         block = _to_legacy_param_block(data)
#         scores[name] = int(block.get("score", 0))
#         per_parameter[name] = block

#     # Build evidence for aggregator (legacy shape only)
#     evidence = {"scores": scores, "per_parameter": per_parameter}
#     evidence_json = json.dumps(evidence, ensure_ascii=False)

#     # 9 = merged meta-synthesis + aggregator (MODEL decides overall_rating)
#     tmpl9 = _load_prompt(prompts_dir, 9)
#     prompt9_body = _inject(
#         tmpl9,
#         evidence_json=evidence_json,
#         script=script_text,
#     )
#     prompt9 = f"{global_preamble}{prompt9_body}"

#     # Ask for structured output (enforces keys and types)
#     try:
#         llm_aggr = _make_llm(temperature).with_structured_output(AggregatorAll)
#         agg: AggregatorAll = llm_aggr.invoke(prompt9)
#     except Exception as e:
#         short = (str(e) or "unknown").strip()
#         raise RuntimeError(f"Aggregator failed on prompt 9. Error: {short}")

#     # Build final payload: use model-decided overall_rating
#     final_payload: Dict[str, Any] = {
#         "scores": scores,
#         "per_parameter": per_parameter,
#         "overall_rating": int(agg.overall_rating),
#         "strengths": agg.strengths,
#         "weaknesses": agg.weaknesses,
#         "suggestions": agg.suggestions,
#         "drop_off_risks": agg.drop_off_risks,
#         "viral_quotient": agg.viral_quotient,
#     }

#     # Normalize & sanitize everything before returning
#     final_payload = normalize_review_payload(final_payload)

#     return _wrap_json(final_payload)

# def _wrap_json(payload: Dict[str, Any]) -> str:
#     return f"{_BEGIN}\n{json.dumps(payload, ensure_ascii=False)}\n{_END}\n"


























# # run_review_multi.py — S3-first prompt loader (Runpod), no local prompts required
# from __future__ import annotations

# import os
# import re
# import json
# import time
# from typing import Any, Dict, List, Optional, Tuple
# from collections import Counter

# from dotenv import load_dotenv
# from langchain_google_genai import ChatGoogleGenerativeAI
# from pydantic import BaseModel, conint  # for structured aggregator output

# # Use the shared normalizer/sanitizer so artifacts are cleaned before UI
# from utils1 import normalize_review_payload

# # -----------------------------
# # Runpod S3 client (boto3)
# # -----------------------------
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError

# # Optional: docx loader (safe fallback if not installed)
# try:
#     import docx  # python-docx
# except Exception:
#     docx = None

# load_dotenv()

# _RP_ENDPOINT = os.getenv("RUNPOD_S3_ENDPOINT", "").strip()
# _RP_BUCKET   = os.getenv("RUNPOD_S3_BUCKET", "").strip()
# _RP_REGION   = os.getenv("RUNPOD_S3_REGION", "").strip()

# def _s3_enabled() -> bool:
#     return bool(_RP_ENDPOINT and _RP_BUCKET)

# _S3_CLIENT = None
# def _s3_client():
#     global _S3_CLIENT
#     if _S3_CLIENT is not None:
#         return _S3_CLIENT
#     if not _s3_enabled():
#         return None
#     _S3_CLIENT = boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         region_name=_RP_REGION or None,
#         config=Config(s3={"addressing_style": "path"})
#     )
#     return _S3_CLIENT

# def _read_s3_text(key: str) -> Optional[str]:
#     """Return object body as UTF-8 text or None."""
#     cli = _s3_client()
#     if not cli:
#         return None
#     try:
#         obj = cli.get_object(Bucket=_RP_BUCKET, Key=key)
#         return obj["Body"].read().decode("utf-8", errors="ignore")
#     except ClientError:
#         return None
#     except Exception:
#         return None

# def _join_s3_key(prefix: str, filename: str) -> str:
#     return f"{prefix.rstrip('/')}/{filename.lstrip('/')}"

# # -----------------------------------------------------------------------------#
# # Env / model
# # -----------------------------------------------------------------------------#
# def _require_api_key() -> None:
#     if not os.getenv("GOOGLE_API_KEY"):
#         raise EnvironmentError("GOOGLE_API_KEY not found in environment/.env")

# def _make_llm(temperature: float) -> ChatGoogleGenerativeAI:
#     """
#     Create a Gemini chat LLM with minimal, deterministic-ish defaults.
#     """
#     model = os.getenv("GEMINI_MODEL") or "gemini-2.5-flash"
#     return ChatGoogleGenerativeAI(
#         model=model,
#         temperature=temperature,
#         top_p=0.0,
#         top_k=1,
#     )

# # -----------------------------------------------------------------------------#
# # Prompt loaders (S3-first; local fallback)
# # -----------------------------------------------------------------------------#
# def _strip_bom(s: str) -> str:
#     return s.lstrip("\ufeff") if s and s.startswith("\ufeff") else s

# def _read_text_local(path: str) -> str:
#     with open(path, "r", encoding="utf-8") as f:
#         return _strip_bom(f.read())

# def _extract_content_block(raw: str) -> str:
#     """
#     If the YAML uses:
#       content: |
#         ...
#     extract that block; else return whole file.
#     """
#     raw = raw.strip()
#     m = re.search(r"^\s*content\s*:\s*\|?\s*(.*)$", raw, flags=re.S | re.I)
#     txt = (m.group(1) if m else raw).replace("\r\n", "\n")
#     return txt.strip()

# def _load_prompt_s3(prompts_prefix: str, n: int) -> Optional[str]:
#     """
#     Try S3: {prompts_prefix}/{n}.yaml then {n}.yml
#     Returns the prompt text or None if not found.
#     """
#     if not _s3_enabled():
#         return None
#     for ext in ("yaml", "yml"):
#         key = _join_s3_key(prompts_prefix, f"{n}.{ext}")
#         txt = _read_s3_text(key)
#         if txt:
#             return _extract_content_block(txt)
#     return None

# def _load_prompt_local(prompts_dir: str, n: int) -> Optional[str]:
#     """
#     Local fallback if S3 not configured.
#     """
#     for ext in ("yaml", "yml"):
#         p = os.path.join(prompts_dir, f"{n}.{ext}")
#         if os.path.isfile(p):
#             raw = _read_text_local(p)
#             return _extract_content_block(raw)
#     return None

# def _load_prompt(prompts_dir_or_prefix: str, n: int) -> str:
#     """
#     Unified loader: prefer S3 (Runpod) if configured; else local.
#     `prompts_dir_or_prefix` should be something like "Scriptmodel/prompts".
#     """
#     # If the caller accidentally passes an s3:// url, strip it to a key prefix:
#     prefix = prompts_dir_or_prefix
#     if prefix.startswith("s3://"):
#         # s3://bucket/... => strip "s3://<bucket>/" so we keep only the key
#         parts = prefix.split("/", 3)
#         prefix = parts[3] if len(parts) > 3 else ""

#     # S3 first
#     txt = _load_prompt_s3(prefix, n)
#     if txt:
#         return txt

#     # Local fallback (only if S3 off or missing file)
#     txt = _load_prompt_local(prompts_dir_or_prefix, n)
#     if txt:
#         return txt

#     where = f"S3({_RP_BUCKET}:{prefix})" if _s3_enabled() else os.path.abspath(prompts_dir_or_prefix)
#     raise FileNotFoundError(f"Prompt {n} (yaml/yml) not found in {where}")

# # -----------------------------------------------------------------------------#
# # Template rendering (supports {var} and {{ var }})
# # -----------------------------------------------------------------------------#
# _VAR_DBL = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")
# _VAR_SGL = re.compile(r"\{([a-zA-Z0-9_]+)\}")

# def _render_template(template: str, variables: Dict[str, str]) -> str:
#     if not variables:
#         return template
#     # replace {{ var }}
#     def _r1(m: re.Match) -> str:
#         key = m.group(1)
#         return str(variables.get(key, ""))
#     out = _VAR_DBL.sub(_r1, template)

#     # replace {var}
#     def _r2(m: re.Match) -> str:
#         key = m.group(1)
#         return str(variables.get(key, ""))
#     out = _VAR_SGL.sub(_r2, out)
#     return out

# # -----------------------------------------------------------------------------#
# # JSON extraction (for specialists that return plain JSON)
# # -----------------------------------------------------------------------------#
# _BEGIN = "BEGIN_JSON"
# _END = "END_JSON"

# def _between_tokens(text: str, start: str, end: str) -> Optional[str]:
#     i = text.find(start)
#     j = text.rfind(end)
#     if i == -1 or j == -1 or j <= i:
#         return None
#     return text[i + len(start): j].strip()

# def _fenced(text: str) -> Optional[str]:
#     m = re.search(r"```json\s*(.*?)\s*```", text, flags=re.S | re.I)
#     if m: return m.group(1).strip()
#     m = re.search(r"```\s*(.*?)\s*```", text, flags=re.S)
#     if m: return m.group(1).strip()
#     return None

# def _balanced(text: str) -> Optional[str]:
#     start = text.find("{")
#     if start == -1: return None
#     depth = 0
#     for i in range(start, len(text)):
#         ch = text[i]
#         if ch == "{":
#             depth += 1
#         elif ch == "}":
#             depth -= 1
#             if depth == 0:
#                 return text[start:i+1]
#     return None

# def _parse_json(text: str) -> Dict[str, Any]:
#     """
#     Try multiple strategies to extract valid JSON. Tolerates trailing commas.
#     """
#     for candidate in filter(None, (_between_tokens(text, _BEGIN, _END),
#                                   _fenced(text), _balanced(text))):
#         try:
#             return json.loads(candidate)
#         except Exception:
#             try:
#                 fixed = re.sub(r",\s*([\]}])", r"\1", candidate)
#                 return json.loads(fixed)
#             except Exception:
#                 continue
#     raise ValueError("Could not parse JSON from model output")

# # -----------------------------------------------------------------------------#
# # Conversions → legacy block for UI
# # -----------------------------------------------------------------------------#
# def _coerce_explanation(raw: Dict[str, Any]) -> str:
#     """
#     Accept 'explanation' (string) or 'explanation_bullets' (list[str]).
#     """
#     explanation = raw.get("explanation", "")
#     if isinstance(explanation, str) and explanation.strip():
#         return explanation.strip()
#     bullets = raw.get("explanation_bullets", [])
#     if isinstance(bullets, list) and bullets:
#         parts = [str(b).strip().rstrip(".") + "." for b in bullets if str(b).strip()]
#         return " ".join(parts)[:1400]
#     return ""

# def _to_legacy_param_block(raw: Dict[str, Any]) -> Dict[str, Any]:
#     """
#     Normalize a specialist block to the legacy shape expected by the UI.
#     - Raw 'extractions' are intentionally suppressed (we only use AOIs for inline highlights).
#     - AOIs pass through untouched (4-field schema, unlimited).
#     - 'summary' (plain human note) is passed through for right-panel display.
#     """
#     score = int(raw.get("score", 0) or 0)
#     explanation = _coerce_explanation(raw)
#     weakness = str(raw.get("weakness", "") or "").strip() or "Not present"

#     # Keep single suggestion for legacy UI AND preserve full list if present
#     suggestion = str(raw.get("suggestion", "") or "").strip()
#     suggestions_list: List[str] = []
#     if not suggestion and isinstance(raw.get("suggestions"), list) and raw["suggestions"]:
#         suggestions_list = [str(s).strip() for s in raw["suggestions"] if str(s).strip()]
#         suggestion = suggestions_list[0] if suggestions_list else ""
#     elif isinstance(raw.get("suggestions"), list):
#         suggestions_list = [str(s).strip() for s in raw["suggestions"] if str(s).strip()]
#     suggestion = suggestion or "Not present"

#     # Hide raw extractions in UI—AOIs only
#     ex: List[str] = []
#     aoi = raw.get("areas_of_improvement") or []

#     block: Dict[str, Any] = {
#         "extractions": ex,  # kept in schema but always empty
#         "score": score,
#         "explanation": explanation,
#         "weakness": weakness,
#         "suggestion": suggestion,
#         "areas_of_improvement": aoi,
#         "summary": str(raw.get("summary", "") or "").strip(),
#     }
#     if suggestions_list:
#         block["suggestions_list"] = suggestions_list
#     return block

# # -----------------------------------------------------------------------------#
# # Display-name mapping for UI & aggregator
# # -----------------------------------------------------------------------------#
# DISPLAY_BY_INDEX = {
#     1: "Suspense Building",
#     2: "Language/Tone",
#     3: "Intro + Main Hook/Cliffhanger",
#     4: "Story Structure + Flow",
#     5: "Pacing",
#     6: "Mini-Hooks (30–60s)",
#     7: "Outro (Ending)",
#     # 8 = global preamble (loaded, not called)
#     9: "Overall Summary (Aggregator)",
# }

# # -----------------------------------------------------------------------------#
# # LLM invoke with retries
# # -----------------------------------------------------------------------------#
# def _invoke_with_retries(llm: ChatGoogleGenerativeAI, prompt: str, tries: int = 3, base_delay: float = 0.8):
#     last_err = None
#     for k in range(tries):
#         try:
#             return llm.invoke(prompt)
#         except Exception as e:
#             last_err = e
#             if k < tries - 1:
#                 time.sleep(base_delay * (2 ** k))
#             else:
#                 raise
#     raise last_err  # type: ignore

# # -----------------------------------------------------------------------------#
# # Aggregator structured schema (model returns this)
# # -----------------------------------------------------------------------------#
# class AggregatorAll(BaseModel):
#     overall_rating: conint(ge=1, le=10)
#     strengths: List[str]
#     weaknesses: List[str]
#     suggestions: List[str]
#     drop_off_risks: List[str]
#     viral_quotient: str

# # -----------------------------------------------------------------------------#
# # DNA/excerpts + style windows loaders
# # -----------------------------------------------------------------------------#
# def _load_signature_excerpts() -> Tuple[str, str, List[str]]:
#     """
#     Returns (dna_profile, dna_excerpts_text, excerpts_list)
#     """
#     import yaml
#     path = os.getenv("SIGNATURE_EXCERPTS_PATH", "reference/signature_excerpts.yaml")
#     dna_profile, excerpts_list = "", []
#     try:
#         if os.path.isfile(path):
#             data = yaml.safe_load(_read_text_local(path)) or {}
#             dna_profile = str(data.get("dna_profile", "") or "")
#             excerpts_list = [str(x) for x in (data.get("excerpts") or []) if str(x).strip()]
#     except Exception:
#         pass
#     dna_excerpts_text = "\n".join(f"- {x}" for x in excerpts_list) if excerpts_list else ""
#     return dna_profile, dna_excerpts_text, excerpts_list

# def _load_docx_text(path: str) -> str:
#     if not path or not os.path.isfile(path):
#         return ""
#     if docx is None:
#         return ""
#     try:
#         d = docx.Document(path)
#         return "\n".join(p.text for p in d.paragraphs if p.text and p.text.strip())
#     except Exception:
#         return ""

# def _chunk_text(text: str, target_chars: int = 2500, overlap_chars: int = 250) -> List[str]:
#     if not text:
#         return []
#     paras = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
#     chunks, buf = [], ""
#     for p in paras:
#         if len(buf) + len(p) + 2 <= target_chars:
#             buf += (("\n\n" if buf else "") + p)
#         else:
#             if buf:
#                 chunks.append(buf)
#                 tail = buf[-overlap_chars:] if overlap_chars else ""
#                 buf = (tail + "\n\n" + p) if tail else p
#             else:
#                 for i in range(0, len(p), target_chars):
#                     chunks.append(p[i:i+target_chars])
#                 buf = ""
#     if buf:
#         chunks.append(buf)
#     return chunks

# _WORD_RE = re.compile(r"[a-zA-Z0-9']+")

# def _tf(text: str) -> Counter:
#     words = [w.lower() for w in _WORD_RE.findall(text)]
#     return Counter(words)

# def _cosine(a: Counter, b: Counter) -> float:
#     if not a or not b:
#         return 0.0
#     common = set(a.keys()) & set(b.keys())
#     num = sum(a[w] * b[w] for w in common)
#     denom1 = sum(v*v for v in a.values()) ** 0.5
#     denom2 = sum(v*v for v in b.values()) ** 0.5
#     if denom1 == 0 or denom2 == 0:
#         return 0.0
#     return num / (denom1 * denom2)

# def _extract_probes(text: str, max_chars_each: int = 1800, n_probes: int = 3) -> List[str]:
#     if not text:
#         return []
#     L = len(text)
#     anchors = [0, L // 2, max(L - max_chars_each, 0)]
#     return [text[a:a+max_chars_each] for a in anchors[:n_probes]]

# def _trim_excerpt(s: str, max_chars: int = 500) -> str:
#     s = re.sub(r"\s+", " ", s or "").strip()
#     return s[:max_chars]

# def _get_style_windows(signature_text: str,
#                        user_script: str,
#                        k_per_probe: int = 2,
#                        max_windows: int = 4,
#                        max_chars: int = 500) -> List[str]:
#     """
#     Lightweight TF cosine retrieval over paragraph-chunks of the signature script.
#     Returns a list of '- excerpt...' lines (already prefixed) for direct template use.
#     """
#     if not signature_text or not user_script:
#         return []
#     chunks = _chunk_text(signature_text, target_chars=2500, overlap_chars=250)
#     if not chunks:
#         return []
#     chunk_tfs = [ _tf(c) for c in chunks ]
#     probes = _extract_probes(user_script, max_chars_each=1800, n_probes=3)
#     scored: List[Tuple[int, float]] = []
#     seen = set()
#     for p in probes:
#         q = _tf(p)
#         sims = [(i, _cosine(q, ct)) for i, ct in enumerate(chunk_tfs)]
#         sims.sort(key=lambda x: x[1], reverse=True)
#         for i, sim in sims[:k_per_probe]:
#             if i not in seen:
#                 scored.append((i, sim))
#                 seen.add(i)
#     scored.sort(key=lambda x: x[1], reverse=True)
#     top = scored[:max_windows]
#     return [ "- " + _trim_excerpt(chunks[i], max_chars=max_chars) for (i, _s) in top ]

# # -----------------------------------------------------------------------------#
# # Core runner
# # -----------------------------------------------------------------------------#
# def run_review_multi(
#     script_text: str,
#     prompts_dir: str = "Scriptmodel/prompts",  # <-- treat as S3 prefix now
#     temperature: float = 0.0,
#     include_commentary: bool = False,  # kept for API parity; ignored
# ) -> str:
#     """
#     Execute prompts 1..7 and 9, prepend prompts/8.yaml as a global preamble to each call,
#     convert to legacy shape, and return BEGIN_JSON ... END_JSON for the UI.

#     NOTE: `prompts_dir` is treated as an S3 prefix (e.g., "Scriptmodel/prompts") when
#     RUNPOD_S3_* env vars are set. If S3 is not configured, it falls back to local files.
#     """
#     _require_api_key()
#     llm = _make_llm(temperature)

#     # ---- Load DNA + excerpts (data layer)
#     dna_profile, dna_excerpts_text, _excerpts_list = _load_signature_excerpts()

#     # ---- Optional: load full signature and build style windows (cheap retrieval)
#     sig_path = os.getenv("SIGNATURE_SCRIPT_PATH", "reference/signature_script.docx")
#     signature_text = _load_docx_text(sig_path)
#     max_windows = int(os.getenv("STYLE_WINDOWS_PER_CALL", "4") or "4")
#     max_chars   = int(os.getenv("STYLE_WINDOW_MAX_CHARS", "500") or "500")
#     signature_style_windows_list = _get_style_windows(
#         signature_text,
#         script_text,
#         k_per_probe=2,
#         max_windows=max_windows,
#         max_chars=max_chars,
#     )
#     signature_style_windows = "\n".join(signature_style_windows_list)

#     # ---- Base variables for all templates
#     base_vars: Dict[str, str] = {
#         "script": script_text,
#         "dna_profile": dna_profile,
#         "dna_excerpts_text": dna_excerpts_text,
#         "signature_style_windows": signature_style_windows,
#     }

#     # Load global preamble (8.yaml) if present and render with DNA/anchors
#     try:
#         global_preamble_tmpl = _load_prompt(prompts_dir, 8).strip()
#         if global_preamble_tmpl:
#             global_preamble = _render_template(global_preamble_tmpl, base_vars) + "\n\n"
#         else:
#             global_preamble = ""
#     except FileNotFoundError:
#         global_preamble = ""

#     # 1..7 specialists
#     scores: Dict[str, int] = {}
#     per_parameter: Dict[str, Dict[str, Any]] = {}

#     for i in range(1, 7 + 1):
#         name = DISPLAY_BY_INDEX[i]
#         tmpl = _load_prompt(prompts_dir, i)
#         prompt_body = _render_template(tmpl, base_vars)
#         prompt = f"{global_preamble}{prompt_body}"

#         try:
#             resp = _invoke_with_retries(llm, prompt)
#             raw_text = getattr(resp, "content", "") or ""
#             data = _parse_json(raw_text)
#         except Exception as e:
#             short = (str(e) or "unknown").strip()
#             raise RuntimeError(f"JSON parse failed on prompt {i} ({name}). Error: {short}")

#         block = _to_legacy_param_block(data)
#         scores[name] = int(block.get("score", 0))
#         per_parameter[name] = block

#     # Build evidence for aggregator (legacy shape only)
#     evidence = {"scores": scores, "per_parameter": per_parameter}
#     evidence_json = json.dumps(evidence, ensure_ascii=False)

#     # 9 = merged meta-synthesis + aggregator (MODEL decides overall_rating)
#     tmpl9 = _load_prompt(prompts_dir, 9)
#     prompt9_body = _render_template(
#         tmpl9,
#         {
#             **base_vars,
#             "evidence_json": evidence_json,
#         }
#     )
#     prompt9 = f"{global_preamble}{prompt9_body}"

#     # Ask for structured output (enforces keys and types)
#     try:
#         llm_aggr = _make_llm(temperature).with_structured_output(AggregatorAll)
#         agg: AggregatorAll = llm_aggr.invoke(prompt9)
#     except Exception as e:
#         short = (str(e) or "unknown").strip()
#         raise RuntimeError(f"Aggregator failed on prompt 9. Error: {short}")

#     # Build final payload: use model-decided overall_rating
#     final_payload: Dict[str, Any] = {
#         "scores": scores,
#         "per_parameter": per_parameter,
#         "overall_rating": int(agg.overall_rating),
#         "strengths": agg.strengths,
#         "weaknesses": agg.weaknesses,
#         "suggestions": agg.suggestions,
#         "drop_off_risks": agg.drop_off_risks,
#         "viral_quotient": agg.viral_quotient,
#     }

#     # Normalize & sanitize everything before returning
#     final_payload = normalize_review_payload(final_payload)

#     return _wrap_json(final_payload)

# def _wrap_json(payload: Dict[str, Any]) -> str:
#     return f"{_BEGIN}\n{json.dumps(payload, ensure_ascii=False)}\n{_END}\n"




#######################################################

##last working code 







# # run_review_multi.py — S3-first prompt loader (Runpod), no local prompts required
# from __future__ import annotations

# import os
# import re
# import json
# import time
# from typing import Any, Dict, List, Optional, Tuple
# from collections import Counter

# from dotenv import load_dotenv
# from langchain_google_genai import ChatGoogleGenerativeAI
# from pydantic import BaseModel, conint  # for structured aggregator output

# # Use the shared normalizer/sanitizer so artifacts are cleaned before UI
# from utils1 import normalize_review_payload

# # -----------------------------
# # Runpod S3 client (boto3)
# # -----------------------------
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError

# # Optional: docx loader (safe fallback if not installed)
# try:
#     import docx  # python-docx
# except Exception:
#     docx = None

# load_dotenv()

# _RP_ENDPOINT = os.getenv("RUNPOD_S3_ENDPOINT", "").strip()
# _RP_BUCKET   = os.getenv("RUNPOD_S3_BUCKET", "").strip()
# _RP_REGION   = os.getenv("RUNPOD_S3_REGION", "").strip()

# def _s3_enabled() -> bool:
#     return bool(_RP_ENDPOINT and _RP_BUCKET)

# _S3_CLIENT = None
# def _s3_client():
#     global _S3_CLIENT
#     if _S3_CLIENT is not None:
#         return _S3_CLIENT
#     if not _s3_enabled():
#         return None
#     _S3_CLIENT = boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         region_name=_RP_REGION or None,
#         config=Config(s3={"addressing_style": "path"})
#     )
#     return _S3_CLIENT

# def _read_s3_text(key: str) -> Optional[str]:
#     """Return object body as UTF-8 text or None."""
#     cli = _s3_client()
#     if not cli:
#         return None
#     try:
#         obj = cli.get_object(Bucket=_RP_BUCKET, Key=key)
#         return obj["Body"].read().decode("utf-8", errors="ignore")
#     except ClientError:
#         return None
#     except Exception:
#         return None

# def _join_s3_key(prefix: str, filename: str) -> str:
#     return f"{prefix.rstrip('/')}/{filename.lstrip('/')}"

# # -----------------------------------------------------------------------------#
# # Env / model
# # -----------------------------------------------------------------------------#
# def _require_api_key() -> None:
#     if not os.getenv("GOOGLE_API_KEY"):
#         raise EnvironmentError("GOOGLE_API_KEY not found in environment/.env")

# def _make_llm(temperature: float) -> ChatGoogleGenerativeAI:
#     """
#     Create a Gemini chat LLM with minimal, deterministic-ish defaults.
#     """
#     model = os.getenv("GEMINI_MODEL") or "gemini-2.5-flash"
#     return ChatGoogleGenerativeAI(
#         model=model,
#         temperature=temperature,
#         top_p=0.0,
#         top_k=1,
#     )

# # -----------------------------------------------------------------------------#
# # Prompt loaders (S3-first; local fallback)
# # -----------------------------------------------------------------------------#
# def _strip_bom(s: str) -> str:
#     return s.lstrip("\ufeff") if s and s.startswith("\ufeff") else s

# def _read_text_local(path: str) -> str:
#     with open(path, "r", encoding="utf-8") as f:
#         return _strip_bom(f.read())

# def _extract_content_block(raw: str) -> str:
#     """
#     If the YAML uses:
#       content: |
#         ...
#     extract that block; else return whole file.
#     """
#     raw = raw.strip()
#     m = re.search(r"^\s*content\s*:\s*\|?\s*(.*)$", raw, flags=re.S | re.I)
#     txt = (m.group(1) if m else raw).replace("\r\n", "\n")
#     return txt.strip()

# def _load_prompt_s3(prompts_prefix: str, n: int) -> Optional[str]:
#     """
#     Try S3: {prompts_prefix}/{n}.yaml then {n}.yml
#     Returns the prompt text or None if not found.
#     """
#     if not _s3_enabled():
#         return None
#     for ext in ("yaml", "yml"):
#         key = _join_s3_key(prompts_prefix, f"{n}.{ext}")
#         txt = _read_s3_text(key)
#         if txt:
#             return _extract_content_block(txt)
#     return None

# def _load_prompt_local(prompts_dir: str, n: int) -> Optional[str]:
#     """
#     Local fallback if S3 not configured.
#     """
#     for ext in ("yaml", "yml"):
#         p = os.path.join(prompts_dir, f"{n}.{ext}")
#         if os.path.isfile(p):
#             raw = _read_text_local(p)
#             return _extract_content_block(raw)
#     return None

# def _load_prompt(prompts_dir_or_prefix: str, n: int) -> str:
#     """
#     Unified loader: prefer S3 (Runpod) if configured; else local.
#     `prompts_dir_or_prefix` should be something like "Scriptmodel/prompts".
#     """
#     # If the caller accidentally passes an s3:// url, strip it to a key prefix:
#     prefix = prompts_dir_or_prefix
#     if prefix.startswith("s3://"):
#         # s3://bucket/... => strip "s3://<bucket>/" so we keep only the key
#         parts = prefix.split("/", 3)
#         prefix = parts[3] if len(parts) > 3 else ""

#     # S3 first
#     txt = _load_prompt_s3(prefix, n)
#     if txt:
#         return txt

#     # Local fallback (only if S3 off or missing file)
#     txt = _load_prompt_local(prompts_dir_or_prefix, n)
#     if txt:
#         return txt

#     where = f"S3({_RP_BUCKET}:{prefix})" if _s3_enabled() else os.path.abspath(prompts_dir_or_prefix)
#     raise FileNotFoundError(f"Prompt {n} (yaml/yml) not found in {where}")

# # -----------------------------------------------------------------------------#
# # Template rendering (supports {var} and {{ var }})
# # -----------------------------------------------------------------------------#
# _VAR_DBL = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")
# _VAR_SGL = re.compile(r"\{([a-zA-Z0-9_]+)\}")

# def _render_template(template: str, variables: Dict[str, str]) -> str:
#     if not variables:
#         return template
#     # replace {{ var }}
#     def _r1(m: re.Match) -> str:
#         key = m.group(1)
#         return str(variables.get(key, ""))
#     out = _VAR_DBL.sub(_r1, template)

#     # replace {var}
#     def _r2(m: re.Match) -> str:
#         key = m.group(1)
#         return str(variables.get(key, ""))
#     out = _VAR_SGL.sub(_r2, out)
#     return out

# # -----------------------------------------------------------------------------#
# # JSON extraction (for specialists that return plain JSON)
# # -----------------------------------------------------------------------------#
# _BEGIN = "BEGIN_JSON"
# _END = "END_JSON"

# def _between_tokens(text: str, start: str, end: str) -> Optional[str]:
#     i = text.find(start)
#     j = text.rfind(end)
#     if i == -1 or j == -1 or j <= i:
#         return None
#     return text[i + len(start): j].strip()

# def _fenced(text: str) -> Optional[str]:
#     m = re.search(r"```json\s*(.*?)\s*```", text, flags=re.S | re.I)
#     if m: return m.group(1).strip()
#     m = re.search(r"```\s*(.*?)\s*```", text, flags=re.S)
#     if m: return m.group(1).strip()
#     return None

# def _balanced(text: str) -> Optional[str]:
#     start = text.find("{")
#     if start == -1: return None
#     depth = 0
#     for i in range(start, len(text)):
#         ch = text[i]
#         if ch == "{":
#             depth += 1
#         elif ch == "}":
#             depth -= 1
#             if depth == 0:
#                 return text[start:i+1]
#     return None

# def _parse_json(text: str) -> Dict[str, Any]:
#     """
#     Try multiple strategies to extract valid JSON. Tolerates trailing commas.
#     """
#     for candidate in filter(None, (_between_tokens(text, _BEGIN, _END),
#                                   _fenced(text), _balanced(text))):
#         try:
#             return json.loads(candidate)
#         except Exception:
#             try:
#                 fixed = re.sub(r",\s*([\]}])", r"\1", candidate)
#                 return json.loads(fixed)
#             except Exception:
#                 continue
#     raise ValueError("Could not parse JSON from model output")

# # -----------------------------------------------------------------------------#
# # Conversions → legacy block for UI
# # -----------------------------------------------------------------------------#
# def _coerce_explanation(raw: Dict[str, Any]) -> str:
#     """
#     Accept 'explanation' (string) or 'explanation_bullets' (list[str]).
#     """
#     explanation = raw.get("explanation", "")
#     if isinstance(explanation, str) and explanation.strip():
#         return explanation.strip()
#     bullets = raw.get("explanation_bullets", [])
#     if isinstance(bullets, list) and bullets:
#         parts = [str(b).strip().rstrip(".") + "." for b in bullets if str(b).strip()]
#         return " ".join(parts)[:1400]
#     return ""

# def _to_legacy_param_block(raw: Dict[str, Any]) -> Dict[str, Any]:
#     """
#     Normalize a specialist block to the legacy shape expected by the UI.
#     - Raw 'extractions' are intentionally suppressed (we only use AOIs for inline highlights).
#     - AOIs pass through untouched (4-field schema, unlimited).
#     - 'summary' (plain human note) is passed through for right-panel display.
#     - For 'Grammar & Spelling', also pass through tense metadata if present.
#     """
#     score = int(raw.get("score", 0) or 0)
#     explanation = _coerce_explanation(raw)
#     weakness = str(raw.get("weakness", "") or "").strip() or "Not present"

#     # Keep single suggestion for legacy UI AND preserve full list if present
#     suggestion = str(raw.get("suggestion", "") or "").strip()
#     suggestions_list: List[str] = []
#     if not suggestion and isinstance(raw.get("suggestions"), list) and raw["suggestions"]:
#         suggestions_list = [str(s).strip() for s in raw["suggestions"] if str(s).strip()]
#         suggestion = suggestions_list[0] if suggestions_list else ""
#     elif isinstance(raw.get("suggestions"), list):
#         suggestions_list = [str(s).strip() for s in raw["suggestions"] if str(s).strip()]
#     suggestion = suggestion or "Not present"

#     # Hide raw extractions in UI—AOIs only
#     ex: List[str] = []
#     aoi = raw.get("areas_of_improvement") or []

#     block: Dict[str, Any] = {
#         "extractions": ex,  # kept in schema but always empty
#         "score": score,
#         "explanation": explanation,
#         "weakness": weakness,
#         "suggestion": suggestion,
#         "areas_of_improvement": aoi,
#         "summary": str(raw.get("summary", "") or "").strip(),
#     }

#     # Optional pass-through fields for Grammar & Spelling (if present)
#     for k in ("dominant_tense", "tense_consistency", "tense_target"):
#         if k in raw:
#             block[k] = raw[k]

#     if suggestions_list:
#         block["suggestions_list"] = suggestions_list
#     return block

# # -----------------------------------------------------------------------------#
# # Display-name mapping for UI & aggregator
# # -----------------------------------------------------------------------------#
# DISPLAY_BY_INDEX = {
#     1: "Suspense Building",
#     2: "Language/Tone",
#     3: "Intro + Main Hook/Cliffhanger",
#     4: "Story Structure + Flow",
#     5: "Pacing",
#     6: "Mini-Hooks (30–60s)",
#     7: "Outro (Ending)",
#     # 8 = global preamble (loaded, not called)
#     9: "Overall Summary (Aggregator)",
#     # New specialist:
#     11: "Grammar & Spelling",
# }

# # Indices to run (specialists only; 8 is preamble, 9 is aggregator)
# SPECIALIST_INDEXES: List[int] = [1, 2, 3, 4, 5, 6, 7, 11]

# # -----------------------------------------------------------------------------#
# # LLM invoke with retries
# # -----------------------------------------------------------------------------#
# def _invoke_with_retries(llm: ChatGoogleGenerativeAI, prompt: str, tries: int = 3, base_delay: float = 0.8):
#     last_err = None
#     for k in range(tries):
#         try:
#             return llm.invoke(prompt)
#         except Exception as e:
#             last_err = e
#             if k < tries - 1:
#                 time.sleep(base_delay * (2 ** k))
#             else:
#                 raise
#     raise last_err  # type: ignore

# # -----------------------------------------------------------------------------#
# # Aggregator structured schema (model returns this)
# # -----------------------------------------------------------------------------#
# class AggregatorAll(BaseModel):
#     overall_rating: conint(ge=1, le=10)
#     strengths: List[str]
#     weaknesses: List[str]
#     suggestions: List[str]
#     drop_off_risks: List[str]
#     viral_quotient: str

# # -----------------------------------------------------------------------------#
# # DNA/excerpts + style windows loaders
# # -----------------------------------------------------------------------------#
# def _load_signature_excerpts() -> Tuple[str, str, List[str]]:
#     """
#     Returns (dna_profile, dna_excerpts_text, excerpts_list)
#     """
#     import yaml
#     path = os.getenv("SIGNATURE_EXCERPTS_PATH", "reference/signature_excerpts.yaml")
#     dna_profile, excerpts_list = "", []
#     try:
#         if os.path.isfile(path):
#             data = yaml.safe_load(_read_text_local(path)) or {}
#             dna_profile = str(data.get("dna_profile", "") or "")
#             excerpts_list = [str(x) for x in (data.get("excerpts") or []) if str(x).strip()]
#     except Exception:
#         pass
#     dna_excerpts_text = "\n".join(f"- {x}" for x in excerpts_list) if excerpts_list else ""
#     return dna_profile, dna_excerpts_text, excerpts_list

# def _load_docx_text(path: str) -> str:
#     if not path or not os.path.isfile(path):
#         return ""
#     if docx is None:
#         return ""
#     try:
#         d = docx.Document(path)
#         return "\n".join(p.text for p in d.paragraphs if p.text and p.text.strip())
#     except Exception:
#         return ""

# def _chunk_text(text: str, target_chars: int = 2500, overlap_chars: int = 250) -> List[str]:
#     if not text:
#         return []
#     paras = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
#     chunks, buf = [], ""
#     for p in paras:
#         if len(buf) + len(p) + 2 <= target_chars:
#             buf += (("\n\n" if buf else "") + p)
#         else:
#             if buf:
#                 chunks.append(buf)
#                 tail = buf[-overlap_chars:] if overlap_chars else ""
#                 buf = (tail + "\n\n" + p) if tail else p
#             else:
#                 for i in range(0, len(p), target_chars):
#                     chunks.append(p[i:i+target_chars])
#                 buf = ""
#     if buf:
#         chunks.append(buf)
#     return chunks

# _WORD_RE = re.compile(r"[a-zA-Z0-9']+")

# def _tf(text: str) -> Counter:
#     words = [w.lower() for w in _WORD_RE.findall(text)]
#     return Counter(words)

# def _cosine(a: Counter, b: Counter) -> float:
#     if not a or not b:
#         return 0.0
#     common = set(a.keys()) & set(b.keys())
#     num = sum(a[w] * b[w] for w in common)
#     denom1 = sum(v*v for v in a.values()) ** 0.5
#     denom2 = sum(v*v for v in b.values()) ** 0.5
#     if denom1 == 0 or denom2 == 0:
#         return 0.0
#     return num / (denom1 * denom2)

# def _extract_probes(text: str, max_chars_each: int = 1800, n_probes: int = 3) -> List[str]:
#     if not text:
#         return []
#     L = len(text)
#     anchors = [0, L // 2, max(L - max_chars_each, 0)]
#     return [text[a:a+max_chars_each] for a in anchors[:n_probes]]

# def _trim_excerpt(s: str, max_chars: int = 500) -> str:
#     s = re.sub(r"\s+", " ", s or "").strip()
#     return s[:max_chars]

# def _get_style_windows(signature_text: str,
#                        user_script: str,
#                        k_per_probe: int = 2,
#                        max_windows: int = 4,
#                        max_chars: int = 500) -> List[str]:
#     """
#     Lightweight TF cosine retrieval over paragraph-chunks of the signature script.
#     Returns a list of '- excerpt...' lines (already prefixed) for direct template use.
#     """
#     if not signature_text or not user_script:
#         return []
#     chunks = _chunk_text(signature_text, target_chars=2500, overlap_chars=250)
#     if not chunks:
#         return []
#     chunk_tfs = [ _tf(c) for c in chunks ]
#     probes = _extract_probes(user_script, max_chars_each=1800, n_probes=3)
#     scored: List[Tuple[int, float]] = []
#     seen = set()
#     for p in probes:
#         q = _tf(p)
#         sims = [(i, _cosine(q, ct)) for i, ct in enumerate(chunk_tfs)]
#         sims.sort(key=lambda x: x[1], reverse=True)
#         for i, sim in sims[:k_per_probe]:
#             if i not in seen:
#                 scored.append((i, sim))
#                 seen.add(i)
#     scored.sort(key=lambda x: x[1], reverse=True)
#     top = scored[:max_windows]
#     return [ "- " + _trim_excerpt(chunks[i], max_chars=max_chars) for (i, _s) in top ]

# # -----------------------------------------------------------------------------#
# # Core runner
# # -----------------------------------------------------------------------------#
# def run_review_multi(
#     script_text: str,
#     prompts_dir: str = "Scriptmodel/prompts/",  # <-- treat as S3 prefix now
#     temperature: float = 0.0,
#     include_commentary: bool = False,  # kept for API parity; ignored
# ) -> str:
#     """
#     Execute prompts 1..7 and 11, prepend prompts/8.yaml as a global preamble to each call,
#     then 9 as the aggregator. Convert to legacy shape, and return BEGIN_JSON ... END_JSON.

#     NOTE: `prompts_dir` is treated as an S3 prefix (e.g., "Scriptmodel/prompts") when
#     RUNPOD_S3_* env vars are set. If S3 is not configured, it falls back to local files.
#     """
#     _require_api_key()
#     llm = _make_llm(temperature)

#     # ---- Load DNA + excerpts (data layer)
#     dna_profile, dna_excerpts_text, _excerpts_list = _load_signature_excerpts()

#     # ---- Optional: load full signature and build style windows (cheap retrieval)
#     sig_path = os.getenv("SIGNATURE_SCRIPT_PATH", "reference/signature_script.docx")
#     signature_text = _load_docx_text(sig_path)
#     max_windows = int(os.getenv("STYLE_WINDOWS_PER_CALL", "4") or "4")
#     max_chars   = int(os.getenv("STYLE_WINDOW_MAX_CHARS", "500") or "500")
#     signature_style_windows_list = _get_style_windows(
#         signature_text,
#         script_text,
#         k_per_probe=2,
#         max_windows=max_windows,
#         max_chars=max_chars,
#     )
#     signature_style_windows = "\n".join(signature_style_windows_list)

#     # ---- Base variables for all templates
#     base_vars: Dict[str, str] = {
#         "script": script_text,
#         "dna_profile": dna_profile,
#         "dna_excerpts_text": dna_excerpts_text,
#         "signature_style_windows": signature_style_windows,
#     }

#     # Load global preamble (8.yaml) if present and render with DNA/anchors
#     try:
#         global_preamble_tmpl = _load_prompt(prompts_dir, 8).strip()
#         if global_preamble_tmpl:
#             global_preamble = _render_template(global_preamble_tmpl, base_vars) + "\n\n"
#         else:
#             global_preamble = ""
#     except FileNotFoundError:
#         global_preamble = ""

#     # 1..7 + 11 specialists
#     scores: Dict[str, int] = {}
#     per_parameter: Dict[str, Dict[str, Any]] = {}

#     for i in SPECIALIST_INDEXES:
#         name = DISPLAY_BY_INDEX[i]
#         tmpl = _load_prompt(prompts_dir, i)
#         prompt_body = _render_template(tmpl, base_vars)
#         prompt = f"{global_preamble}{prompt_body}"

#         try:
#             resp = _invoke_with_retries(llm, prompt)
#             raw_text = getattr(resp, "content", "") or ""
#             data = _parse_json(raw_text)
#         except Exception as e:
#             short = (str(e) or "unknown").strip()
#             raise RuntimeError(f"JSON parse failed on prompt {i} ({name}). Error: {short}")

#         block = _to_legacy_param_block(data)
#         scores[name] = int(block.get("score", 0))
#         per_parameter[name] = block

#     # Build evidence for aggregator (legacy shape only)
#     evidence = {"scores": scores, "per_parameter": per_parameter}
#     evidence_json = json.dumps(evidence, ensure_ascii=False)

#     # 9 = merged meta-synthesis + aggregator (MODEL decides overall_rating)
#     tmpl9 = _load_prompt(prompts_dir, 9)
#     prompt9_body = _render_template(
#         tmpl9,
#         {
#             **base_vars,
#             "evidence_json": evidence_json,
#         }
#     )
#     prompt9 = f"{global_preamble}{prompt9_body}"

#     # Ask for structured output (enforces keys and types)
#     try:
#         llm_aggr = _make_llm(temperature).with_structured_output(AggregatorAll)
#         agg: AggregatorAll = llm_aggr.invoke(prompt9)
#     except Exception as e:
#         short = (str(e) or "unknown").strip()
#         raise RuntimeError(f"Aggregator failed on prompt 9. Error: {short}")

#     # Build final payload: use model-decided overall_rating
#     final_payload: Dict[str, Any] = {
#         "scores": scores,
#         "per_parameter": per_parameter,
#         "overall_rating": int(agg.overall_rating),
#         "strengths": agg.strengths,
#         "weaknesses": agg.weaknesses,
#         "suggestions": agg.suggestions,
#         "drop_off_risks": agg.drop_off_risks,
#         "viral_quotient": agg.viral_quotient,
#     }

#     # Normalize & sanitize everything before returning
#     final_payload = normalize_review_payload(final_payload)

#     return _wrap_json(final_payload)

# def _wrap_json(payload: Dict[str, Any]) -> str:
#     return f"{_BEGIN}\n{json.dumps(payload, ensure_ascii=False)}\n{_END}\n"








##########################################



##History 








# review_engine_multi.py — S3-only (Runpod) prompt loader with manifest support and retries
from __future__ import annotations

import os
import re
import json
import time
import tempfile
from typing import Any, Dict, List, Optional, Tuple
from collections import Counter

# Optional Streamlit secrets read (works even when imported from Streamlit app)
try:
    import streamlit as st  # type: ignore
except Exception:  # pragma: no cover
    st = None  # type: ignore

from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from pydantic import BaseModel, conint  # for structured aggregator output

# Use the shared normalizer/sanitizer so artifacts are cleaned before UI
from utils1 import normalize_review_payload

# -----------------------------
# Runpod S3 client (boto3) — S3 ONLY
# -----------------------------
import boto3
from botocore.config import Config
from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError, BotoCoreError

load_dotenv()

def _get_env(key: str, default: str = "") -> str:
    # prefer Streamlit secrets if available, then env var
    if st is not None:
        try:
            v2 = st.secrets.get(key)  # type: ignore[attr-defined]
            if isinstance(v2, str) and v2.strip():
                return v2.strip()
        except Exception:
            pass
    v = os.getenv(key, "")
    return v.strip() if v else (default or "").strip()

_RP_ENDPOINT = _get_env("RUNPOD_S3_ENDPOINT")
_RP_BUCKET   = _get_env("RUNPOD_S3_BUCKET")
_RP_REGION   = _get_env("RUNPOD_S3_REGION") or _get_env("AWS_DEFAULT_REGION") or "us-east-1"

_AK = _get_env("AWS_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY")
_SK = _get_env("AWS_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_KEY")
_ST = _get_env("AWS_SESSION_TOKEN")  # optional

def _require_s3():
    missing = []
    if not _RP_ENDPOINT: missing.append("RUNPOD_S3_ENDPOINT")
    if not _RP_BUCKET:   missing.append("RUNPOD_S3_BUCKET")
    if not _AK:          missing.append("AWS_ACCESS_KEY_ID/RUNPOD_S3_ACCESS_KEY_ID")
    if not _SK:          missing.append("AWS_SECRET_ACCESS_KEY/RUNPOD_S3_SECRET_ACCESS_KEY")
    if missing:
        raise RuntimeError("S3 required but not configured. Missing: " + ", ".join(missing))

_S3 = None
def _s3():
    global _S3
    if _S3 is not None:
        return _S3
    _require_s3()
    _S3 = boto3.client(
        "s3",
        endpoint_url=_RP_ENDPOINT,
        region_name=_RP_REGION,
        aws_access_key_id=_AK,
        aws_secret_access_key=_SK,
        aws_session_token=_ST or None,
        config=Config(
            signature_version="s3v4",
            s3={"addressing_style": "path"},
            retries={"max_attempts": 10, "mode": "adaptive"},
            read_timeout=90,
            connect_timeout=15,
            tcp_keepalive=True,
        ),
    )
    return _S3

def _s3_get_bytes(key: str, tries: int = 6) -> Optional[bytes]:
    key = key.lstrip("/")
    last = None
    for i in range(tries):
        try:
            resp = _s3().get_object(Bucket=_RP_BUCKET, Key=key)
            return resp["Body"].read()
        except (EndpointConnectionError, BotoCoreError, ClientError) as e:
            status = getattr(e, "response", {}).get("ResponseMetadata", {}).get("HTTPStatusCode")
            # retry on network or 5xx
            if isinstance(e, EndpointConnectionError) or status in (500, 502, 503, 504):
                last = e
                time.sleep(min(2 ** i, 12))
                continue
            return None
        except Exception as e:  # pragma: no cover
            last = e
            time.sleep(min(2 ** i, 12))
    return None

def _s3_get_text(key: str, tries: int = 6) -> Optional[str]:
    b = _s3_get_bytes(key, tries=tries)
    if b is None:
        return None
    try:
        return b.decode("utf-8", errors="ignore")
    except Exception:
        return None

def _s3_list(prefix: str) -> List[str]:
    out: List[str] = []
    token = None
    pfx = prefix.rstrip("/") + "/"
    while True:
        kw = {"Bucket": _RP_BUCKET, "Prefix": pfx, "MaxKeys": 1000}
        if token: kw["ContinuationToken"] = token
        try:
            resp = _s3().list_objects_v2(**kw)
        except (ClientError, EndpointConnectionError, NoCredentialsError, BotoCoreError):
            break  # No List permission or network issue -> treat as empty
        for c in resp.get("Contents", []):
            out.append(c["Key"])
        token = resp.get("NextContinuationToken")
        if not token:
            break
    return out

def _join_key(prefix: str, filename: str) -> str:
    return f"{prefix.rstrip('/')}/{filename.lstrip('/')}"

# -----------------------------------------------------------------------------#
# Env / model
# -----------------------------------------------------------------------------#
def _require_api_key() -> None:
    if not os.getenv("GOOGLE_API_KEY"):
        raise EnvironmentError("GOOGLE_API_KEY not found in environment/.env")

def _make_llm(temperature: float) -> ChatGoogleGenerativeAI:
    model = os.getenv("GEMINI_MODEL") or "gemini-2.5-flash"
    return ChatGoogleGenerativeAI(
        model=model,
        temperature=temperature,
        top_p=0.0,
        top_k=1,
    )

# -----------------------------------------------------------------------------#
# Prompt loaders (S3-only, with optional manifest; no local fallback)
# -----------------------------------------------------------------------------#
def _strip_bom(s: str) -> str:
    return s.lstrip("\ufeff") if s and s.startswith("\ufeff") else s

def _extract_content_block(raw: str) -> str:
    """
    If the YAML contains a top-level `content: |`, return that block;
    otherwise return the whole file. This lets you keep metadata in the YAML.
    """
    raw = raw.strip()
    # naive but practical: find 'content:' at start-of-line
    m = re.search(r"^\s*content\s*:\s*\|?\s*(.*)$", raw, flags=re.S | re.I)
    txt = (m.group(1) if m else raw).replace("\r\n", "\n")
    return txt.strip()

def _read_prompts_manifest(prompts_prefix: str) -> Dict[str, str]:
    """
    Reads prompts manifest JSON at <prompts_prefix>/_manifest.json, if present.
    Accepts any of these formats:
      {"1":"1_intro.yaml","2":"2_tone.yaml"}
      {"map":{"1":"1_intro.yaml", ...}}
      [{"index":1,"file":"1_intro.yaml"}, ...]
    Returns a dict { "1": "1_intro.yaml", ... } with string keys.
    """
    key = _join_key(prompts_prefix, "_manifest.json")
    txt = _s3_get_text(key)
    if not txt:
        return {}
    try:
        j = json.loads(txt)
    except Exception:
        return {}
    if isinstance(j, dict):
        if "map" in j and isinstance(j["map"], dict):
            j = j["map"]
        return {str(k): str(v) for k, v in j.items() if isinstance(v, str) and v}
    if isinstance(j, list):
        out: Dict[str, str] = {}
        for item in j:
            try:
                idx = str(item.get("index"))
                fil = str(item.get("file"))
                if idx and fil:
                    out[idx] = fil
            except Exception:
                continue
        return out
    return {}

def _candidate_prompt_keys(prompts_prefix: str, n: int) -> List[str]:
    # Try exact numeric filenames first (no List required)
    cands = [
        _join_key(prompts_prefix, f"{n}.yaml"),
        _join_key(prompts_prefix, f"{n}.yml"),
        _join_key(prompts_prefix, f"{n:02}.yaml"),
        _join_key(prompts_prefix, f"{n:02}.yml"),
    ]
    # From manifest (no List required)
    man = _read_prompts_manifest(prompts_prefix)
    man_name = man.get(str(n))
    if man_name:
        cands.insert(0, _join_key(prompts_prefix, man_name))
    return cands

def _discover_prompt_via_list(prompts_prefix: str, n: int) -> Optional[str]:
    # Requires ListBucket; search for keys like "n_*.yaml|yml" and zero-padded variants
    keys = _s3_list(prompts_prefix)
    if not keys:
        return None
    nstr = str(n)
    n2 = f"{n:02}"
    best: Optional[str] = None
    for k in keys:
        base = os.path.basename(k).lower()
        if (base.startswith(nstr + "_") or base.startswith(n2 + "_")) and (base.endswith(".yaml") or base.endswith(".yml")):
            best = k
            break
    return best

def _load_prompt(prompts_prefix: str, n: int) -> str:
    """
    Unified S3 loader for prompt index n.
    Search order (no List permission required):
      1) {n}.yaml / {n}.yml / {n:02}.yaml / {n:02}.yml
      2) manifest _manifest.json -> map[n] => filename
      3) If List is allowed: first key matching "n_*.yaml|yml" or "0n_*.yaml|yml"
    """
    if prompts_prefix.startswith("s3://"):
        parts = prompts_prefix.split("/", 3)
        # ignore bucket from s3:// if provided; we use configured bucket instead
        prompts_prefix = parts[3] if len(parts) > 3 else ""

    tried: List[str] = []
    for key in _candidate_prompt_keys(prompts_prefix, n):
        tried.append(key)
        txt = _s3_get_text(key)
        if txt:
            return _extract_content_block(_strip_bom(txt))

    # optional discovery via List
    k = _discover_prompt_via_list(prompts_prefix, n)
    if k:
        tried.append(k)
        txt = _s3_get_text(k)
        if txt:
            return _extract_content_block(_strip_bom(txt))

    raise FileNotFoundError(
        "Prompt {n} not found. Tried keys:\n- ".format(n=n) + "\n- ".join(tried) +
        "\nUpload a YAML to one of those paths or add _manifest.json under the prompts prefix."
    )

# -----------------------------------------------------------------------------#
# Template rendering (supports {var} and {{ var }})
# -----------------------------------------------------------------------------#
_VAR_DBL = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")
_VAR_SGL = re.compile(r"\{([a-zA-Z0-9_]+)\}")

def _render_template(template: str, variables: Dict[str, str]) -> str:
    if not variables:
        return template
    # replace {{ var }}
    def _r1(m: re.Match) -> str:  # type: ignore
        key = m.group(1)
        return str(variables.get(key, ""))
    out = _VAR_DBL.sub(_r1, template)

    # replace {var}
    def _r2(m: re.Match) -> str:  # type: ignore
        key = m.group(1)
        return str(variables.get(key, ""))
    out = _VAR_SGL.sub(_r2, out)
    return out

# -----------------------------------------------------------------------------#
# JSON extraction (for specialists that return plain JSON)
# -----------------------------------------------------------------------------#
_BEGIN = "BEGIN_JSON"
_END = "END_JSON"

def _between_tokens(text: str, start: str, end: str) -> Optional[str]:
    i = text.find(start)
    j = text.rfind(end)
    if i == -1 or j == -1 or j <= i:
        return None
    return text[i + len(start): j].strip()

def _fenced(text: str) -> Optional[str]:
    m = re.search(r"```json\s*(.*?)\s*```", text, flags=re.S | re.I)
    if m: return m.group(1).strip()
    m = re.search(r"```\s*(.*?)\s*```", text, flags=re.S)
    if m: return m.group(1).strip()
    return None

def _balanced(text: str) -> Optional[str]:
    start = text.find("{")
    if start == -1: return None
    depth = 0
    for i in range(start, len(text)):
        ch = text[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[start:i+1]
    return None

def _parse_json(text: str) -> Dict[str, Any]:
    """
    Try multiple strategies to extract valid JSON. Tolerates trailing commas.
    """
    for candidate in filter(None, (_between_tokens(text, _BEGIN, _END),
                                  _fenced(text), _balanced(text))):
        try:
            return json.loads(candidate)
        except Exception:
            try:
                fixed = re.sub(r",\s*([\]}])", r"\1", candidate)
                return json.loads(fixed)
            except Exception:
                continue
    raise ValueError("Could not parse JSON from model output")

# -----------------------------------------------------------------------------#
# Conversions → legacy block for UI
# -----------------------------------------------------------------------------#
def _coerce_explanation(raw: Dict[str, Any]) -> str:
    explanation = raw.get("explanation", "")
    if isinstance(explanation, str) and explanation.strip():
        return explanation.strip()
    bullets = raw.get("explanation_bullets", [])
    if isinstance(bullets, list) and bullets:
        parts = [str(b).strip().rstrip(".") + "." for b in bullets if str(b).strip()]
        return " ".join(parts)[:1400]
    return ""

def _to_legacy_param_block(raw: Dict[str, Any]) -> Dict[str, Any]:
    score = int(raw.get("score", 0) or 0)
    explanation = _coerce_explanation(raw)
    weakness = str(raw.get("weakness", "") or "").strip() or "Not present"

    suggestion = str(raw.get("suggestion", "") or "").strip()
    suggestions_list: List[str] = []
    if not suggestion and isinstance(raw.get("suggestions"), list) and raw["suggestions"]:
        suggestions_list = [str(s).strip() for s in raw["suggestions"] if str(s).strip()]
        suggestion = suggestions_list[0] if suggestions_list else ""
    elif isinstance(raw.get("suggestions"), list):
        suggestions_list = [str(s).strip() for s in raw["suggestions"] if str(s).strip()]
    suggestion = suggestion or "Not present"

    ex: List[str] = []
    aoi = raw.get("areas_of_improvement") or []

    block: Dict[str, Any] = {
        "extractions": ex,
        "score": score,
        "explanation": explanation,
        "weakness": weakness,
        "suggestion": suggestion,
        "areas_of_improvement": aoi,
        "summary": str(raw.get("summary", "") or "").strip(),
    }

    for k in ("dominant_tense", "tense_consistency", "tense_target"):
        if k in raw:
            block[k] = raw[k]

    if suggestions_list:
        block["suggestions_list"] = suggestions_list
    return block

# -----------------------------------------------------------------------------#
# Display-name mapping for UI & aggregator
# -----------------------------------------------------------------------------#
DISPLAY_BY_INDEX = {
    1: "Suspense Building",
    2: "Language/Tone",
    3: "Intro + Main Hook/Cliffhanger",
    4: "Story Structure + Flow",
    5: "Pacing",
    6: "Mini-Hooks (30–60s)",
    7: "Outro (Ending)",
    # 8 = global preamble (loaded, not called)
    9: "Overall Summary (Aggregator)",
    11: "Grammar & Spelling",
}

SPECIALIST_INDEXES: List[int] = [1, 2, 3, 4, 5, 6, 7, 11]

# -----------------------------------------------------------------------------#
# LLM invoke with retries
# -----------------------------------------------------------------------------#
def _invoke_with_retries(llm: ChatGoogleGenerativeAI, prompt: str, tries: int = 3, base_delay: float = 0.8):
    last_err = None
    for k in range(tries):
        try:
            return llm.invoke(prompt)
        except Exception as e:
            last_err = e
            if k < tries - 1:
                time.sleep(base_delay * (2 ** k))
            else:
                raise
    raise last_err  # type: ignore

# -----------------------------------------------------------------------------#
# Aggregator structured schema (model returns this)
# -----------------------------------------------------------------------------#
class AggregatorAll(BaseModel):
    overall_rating: conint(ge=1, le=10)
    strengths: List[str]
    weaknesses: List[str]
    suggestions: List[str]
    drop_off_risks: List[str]
    viral_quotient: str

# -----------------------------------------------------------------------------#
# Signature refs (S3) & style windows
# -----------------------------------------------------------------------------#
try:
    import yaml  # type: ignore
except Exception:  # pragma: no cover
    yaml = None  # type: ignore

try:
    import docx  # type: ignore
except Exception:  # pragma: no cover
    docx = None  # type: ignore

def _load_signature_excerpts_s3() -> Tuple[str, str, List[str]]:
    """
    Loads signature excerpts from S3 key set via env SIGNATURE_EXCERPTS_KEY (default Scriptmodel/reference/signature_excerpts.yaml).
    Returns (dna_profile, dna_excerpts_text, excerpts_list)
    """
    key = _get_env("SIGNATURE_EXCERPTS_KEY", "Scriptmodel/reference/signature_excerpts.yaml")
    txt = _s3_get_text(key) if key else None
    dna_profile, excerpts_list = "", []
    if txt and yaml is not None:
        try:
            data = yaml.safe_load(txt) or {}
            dna_profile = str(data.get("dna_profile", "") or "")
            excerpts_list = [str(x) for x in (data.get("excerpts") or []) if str(x).strip()]
        except Exception:
            pass
    dna_excerpts_text = "\n".join(f"- {x}" for x in excerpts_list) if excerpts_list else ""
    return dna_profile, dna_excerpts_text, excerpts_list

def _load_docx_text_s3() -> str:
    """
    Loads signature script DOCX from S3 key SIGNATURE_SCRIPT_KEY (default Scriptmodel/reference/signature_script.docx).
    Returns plain text (paragraph-joined) or empty string on failure.
    """
    key = _get_env("SIGNATURE_SCRIPT_KEY", "Scriptmodel/reference/signature_script.docx")
    if not key or docx is None:
        return ""
    b = _s3_get_bytes(key)
    if not b:
        return ""
    try:
        # write to temp, parse with python-docx
        fd, tmp = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        with open(tmp, "wb") as f:
            f.write(b)
        d = docx.Document(tmp)  # type: ignore
        try:
            os.remove(tmp)
        except Exception:
            pass
        return "\n".join(p.text for p in d.paragraphs if p.text and p.text.strip())
    except Exception:
        return ""

def _chunk_text(text: str, target_chars: int = 2500, overlap_chars: int = 250) -> List[str]:
    if not text:
        return []
    paras = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    chunks, buf = [], ""
    for p in paras:
        if len(buf) + len(p) + 2 <= target_chars:
            buf += (("\n\n" if buf else "") + p)
        else:
            if buf:
                chunks.append(buf)
                tail = buf[-overlap_chars:] if overlap_chars else ""
                buf = (tail + "\n\n" + p) if tail else p
            else:
                for i in range(0, len(p), target_chars):
                    chunks.append(p[i:i+target_chars])
                buf = ""
    if buf:
        chunks.append(buf)
    return chunks

_WORD_RE = re.compile(r"[a-zA-Z0-9']+")

def _tf(text: str) -> Counter:
    words = [w.lower() for w in _WORD_RE.findall(text)]
    return Counter(words)

def _cosine(a: Counter, b: Counter) -> float:
    if not a or not b:
        return 0.0
    common = set(a.keys()) & set(b.keys())
    num = sum(a[w] * b[w] for w in common)
    denom1 = sum(v*v for v in a.values()) ** 0.5
    denom2 = sum(v*v for v in b.values()) ** 0.5
    if denom1 == 0 or denom2 == 0:
        return 0.0
    return num / (denom1 * denom2)

def _extract_probes(text: str, max_chars_each: int = 1800, n_probes: int = 3) -> List[str]:
    if not text:
        return []
    L = len(text)
    anchors = [0, L // 2, max(L - max_chars_each, 0)]
    return [text[a:a+max_chars_each] for a in anchors[:n_probes]]

def _trim_excerpt(s: str, max_chars: int = 500) -> str:
    s = re.sub(r"\s+", " ", s or "").strip()
    return s[:max_chars]

def _get_style_windows(signature_text: str,
                       user_script: str,
                       k_per_probe: int = 2,
                       max_windows: int = 4,
                       max_chars: int = 500) -> List[str]:
    if not signature_text or not user_script:
        return []
    chunks = _chunk_text(signature_text, target_chars=2500, overlap_chars=250)
    if not chunks:
        return []
    chunk_tfs = [ _tf(c) for c in chunks ]
    probes = _extract_probes(user_script, max_chars_each=1800, n_probes=3)
    scored: List[Tuple[int, float]] = []
    seen = set()
    for p in probes:
        q = _tf(p)
        sims = [(i, _cosine(q, ct)) for i, ct in enumerate(chunk_tfs)]
        sims.sort(key=lambda x: x[1], reverse=True)
        for i, sim in sims[:k_per_probe]:
            if i not in seen:
                scored.append((i, sim))
                seen.add(i)
    scored.sort(key=lambda x: x[1], reverse=True)
    top = scored[:max_windows]
    return [ "- " + _trim_excerpt(chunks[i], max_chars=max_chars) for (i, _s) in top ]

# -----------------------------------------------------------------------------#
# Core runner
# -----------------------------------------------------------------------------#
def run_review_multi(
    script_text: str,
    prompts_dir: str = "Scriptmodel/prompts",  # S3 prefix only
    temperature: float = 0.0,
    include_commentary: bool = False,  # kept for API parity; ignored
) -> str:
    """
    Execute prompts 1..7 and 11, prepend prompts/8.yaml as a global preamble to each call,
    then 9 as the aggregator. Convert to legacy shape, and return BEGIN_JSON ... END_JSON.

    NOTE: `prompts_dir` is treated as an S3 prefix (e.g., "Scriptmodel/prompts").
    There is NO local fallback in this S3-only build.
    """
    _require_api_key()
    # Ensure S3 is configured
    _require_s3()

    llm = _make_llm(temperature)

    # ---- Load DNA + excerpts from S3 (optional)
    dna_profile, dna_excerpts_text, _excerpts_list = _load_signature_excerpts_s3()

    # ---- Optional: load full signature (S3) and build style windows
    signature_text = _load_docx_text_s3()
    max_windows = int(os.getenv("STYLE_WINDOWS_PER_CALL", "4") or "4")
    max_chars   = int(os.getenv("STYLE_WINDOW_MAX_CHARS", "500") or "500")
    signature_style_windows_list = _get_style_windows(
        signature_text,
        script_text,
        k_per_probe=2,
        max_windows=max_windows,
        max_chars=max_chars,
    )
    signature_style_windows = "\n".join(signature_style_windows_list)

    # ---- Base variables for all templates
    base_vars: Dict[str, str] = {
        "script": script_text,
        "dna_profile": dna_profile,
        "dna_excerpts_text": dna_excerpts_text,
        "signature_style_windows": signature_style_windows,
    }

    # Load global preamble (8.yaml) if present and render with DNA/anchors
    try:
        global_preamble_tmpl = _load_prompt(prompts_dir, 8).strip()
        if global_preamble_tmpl:
            global_preamble = _render_template(global_preamble_tmpl, base_vars) + "\n\n"
        else:
            global_preamble = ""
    except FileNotFoundError:
        global_preamble = ""

    # 1..7 + 11 specialists
    scores: Dict[str, int] = {}
    per_parameter: Dict[str, Dict[str, Any]] = {}

    for i in SPECIALIST_INDEXES:
        name = DISPLAY_BY_INDEX[i]
        try:
            tmpl = _load_prompt(prompts_dir, i)
        except FileNotFoundError as e:
            raise FileNotFoundError(f"{e}\nTip: put '{i}.yaml' (or set _manifest.json) under s3://{_RP_BUCKET}/{prompts_dir}/")

        prompt_body = _render_template(tmpl, base_vars)
        prompt = f"{global_preamble}{prompt_body}"

        try:
            resp = _invoke_with_retries(llm, prompt)
            raw_text = getattr(resp, "content", "") or ""
            data = _parse_json(raw_text)
        except Exception as e:
            short = (str(e) or "unknown").strip()
            raise RuntimeError(f"JSON parse failed on prompt {i} ({name}). Error: {short}")

        block = _to_legacy_param_block(data)
        scores[name] = int(block.get("score", 0))
        per_parameter[name] = block

    # Build evidence for aggregator (legacy shape only)
    evidence = {"scores": scores, "per_parameter": per_parameter}
    evidence_json = json.dumps(evidence, ensure_ascii=False)

    # 9 = merged meta-synthesis + aggregator (MODEL decides overall_rating)
    try:
        tmpl9 = _load_prompt(prompts_dir, 9)
    except FileNotFoundError as e:
        raise FileNotFoundError(f"{e}\nTip: ensure '9.yaml' exists under s3://{_RP_BUCKET}/{prompts_dir}/")

    prompt9_body = _render_template(
        tmpl9,
        {
            **base_vars,
            "evidence_json": evidence_json,
        }
    )
    prompt9 = f"{global_preamble}{prompt9_body}"

    # Ask for structured output (enforces keys and types)
    try:
        llm_aggr = _make_llm(temperature).with_structured_output(AggregatorAll)
        agg: AggregatorAll = llm_aggr.invoke(prompt9)
    except Exception as e:
        short = (str(e) or "unknown").strip()
        raise RuntimeError(f"Aggregator failed on prompt 9. Error: {short}")

    # Build final payload: use model-decided overall_rating
    final_payload: Dict[str, Any] = {
        "scores": scores,
        "per_parameter": per_parameter,
        "overall_rating": int(agg.overall_rating),
        "strengths": agg.strengths,
        "weaknesses": agg.weaknesses,
        "suggestions": agg.suggestions,
        "drop_off_risks": agg.drop_off_risks,
        "viral_quotient": agg.viral_quotient,
    }

    # Normalize & sanitize everything before returning
    final_payload = normalize_review_payload(final_payload)

    return _wrap_json(final_payload)

def _wrap_json(payload: Dict[str, Any]) -> str:
    return f"{_BEGIN}\n{json.dumps(payload, ensure_ascii=False)}\n{_END}\n"













#######################################















# review_engine_multi.py — S3-only (Runpod) prompt loader with manifest support and retries
from __future__ import annotations

import os
import re
import json
import time
import tempfile
from typing import Any, Dict, List, Optional, Tuple
from collections import Counter

# Optional Streamlit secrets read (works even when imported from Streamlit app)
try:
    import streamlit as st  # type: ignore
except Exception:  # pragma: no cover
    st = None  # type: ignore

from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from pydantic import BaseModel, conint  # for structured aggregator output

# Use the shared normalizer/sanitizer so artifacts are cleaned before UI
from utils1 import normalize_review_payload

# -----------------------------
# Runpod S3 client (boto3) — S3 ONLY
# -----------------------------
import boto3
from botocore.config import Config
from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError, BotoCoreError

load_dotenv()

# -----------------------------------------------------------------------------#
# Streamlit-aware secrets/env getter
# -----------------------------------------------------------------------------#
def _sget(key: str, default: Optional[str] = None, section: Optional[str] = None) -> str:
    """
    Resolve config values in this priority:
      1) os.environ[key]
      2) st.secrets[section][key] (if section provided and available)
      3) st.secrets[key]
      4) default
    Returns a trimmed string ("" if nothing found).
    Works safely even if Streamlit isn't installed/running.
    """
    # 1) environment
    v = os.getenv(key)
    if v not in (None, ""):
        return str(v).strip()

    # 2/3) streamlit secrets
    if st is not None:
        try:
            if section:
                sect = st.secrets.get(section, {})  # type: ignore[attr-defined]
                if isinstance(sect, dict):
                    v2 = sect.get(key)
                    if isinstance(v2, str) and v2.strip():
                        return v2.strip()
            v3 = st.secrets.get(key)  # type: ignore[attr-defined]
            if isinstance(v3, str) and v3.strip():
                return v3.strip()
        except Exception:
            pass

    # 4) default
    return (default or "").strip()

# -----------------------------------------------------------------------------#
# Env / S3 config (via _sget)
# -----------------------------------------------------------------------------#
_RP_ENDPOINT = _sget("RUNPOD_S3_ENDPOINT")
_RP_BUCKET   = _sget("RUNPOD_S3_BUCKET")
_RP_REGION   = _sget("RUNPOD_S3_REGION") or _sget("AWS_DEFAULT_REGION") or "us-east-1"

_AK = _sget("AWS_ACCESS_KEY_ID") or _sget("RUNPOD_S3_ACCESS_KEY_ID") or _sget("RUNPOD_S3_ACCESS_KEY")
_SK = _sget("AWS_SECRET_ACCESS_KEY") or _sget("RUNPOD_S3_SECRET_ACCESS_KEY") or _sget("RUNPOD_S3_SECRET_KEY")
_ST = _sget("AWS_SESSION_TOKEN")  # optional

def _require_s3():
    missing = []
    if not _RP_ENDPOINT: missing.append("RUNPOD_S3_ENDPOINT")
    if not _RP_BUCKET:   missing.append("RUNPOD_S3_BUCKET")
    if not _AK:          missing.append("AWS_ACCESS_KEY_ID/RUNPOD_S3_ACCESS_KEY_ID")
    if not _SK:          missing.append("AWS_SECRET_ACCESS_KEY/RUNPOD_S3_SECRET_ACCESS_KEY")
    if missing:
        raise RuntimeError("S3 required but not configured. Missing: " + ", ".join(missing))

_S3 = None
def _s3():
    global _S3
    if _S3 is not None:
        return _S3
    _require_s3()
    _S3 = boto3.client(
        "s3",
        endpoint_url=_RP_ENDPOINT,
        region_name=_RP_REGION,
        aws_access_key_id=_AK,
        aws_secret_access_key=_SK,
        aws_session_token=_ST or None,
        config=Config(
            signature_version="s3v4",
            s3={"addressing_style": "path"},
            retries={"max_attempts": 10, "mode": "adaptive"},
            read_timeout=90,
            connect_timeout=15,
            tcp_keepalive=True,
        ),
    )
    return _S3

def _s3_get_bytes(key: str, tries: int = 6) -> Optional[bytes]:
    key = key.lstrip("/")
    last = None
    for i in range(tries):
        try:
            resp = _s3().get_object(Bucket=_RP_BUCKET, Key=key)
            return resp["Body"].read()
        except (EndpointConnectionError, BotoCoreError, ClientError) as e:
            status = getattr(e, "response", {}).get("ResponseMetadata", {}).get("HTTPStatusCode")
            # retry on network or 5xx
            if isinstance(e, EndpointConnectionError) or status in (500, 502, 503, 504):
                last = e
                time.sleep(min(2 ** i, 12))
                continue
            return None
        except Exception as e:  # pragma: no cover
            last = e
            time.sleep(min(2 ** i, 12))
    return None

def _s3_get_text(key: str, tries: int = 6) -> Optional[str]:
    b = _s3_get_bytes(key, tries=tries)
    if b is None:
        return None
    try:
        return b.decode("utf-8", errors="ignore")
    except Exception:
        return None

def _s3_list(prefix: str) -> List[str]:
    out: List[str] = []
    token = None
    pfx = prefix.rstrip("/") + "/"
    while True:
        kw = {"Bucket": _RP_BUCKET, "Prefix": pfx, "MaxKeys": 1000}
        if token: kw["ContinuationToken"] = token
        try:
            resp = _s3().list_objects_v2(**kw)
        except (ClientError, EndpointConnectionError, NoCredentialsError, BotoCoreError):
            break  # No List permission or network issue -> treat as empty
        for c in resp.get("Contents", []):
            out.append(c["Key"])
        token = resp.get("NextContinuationToken")
        if not token:
            break
    return out

def _join_key(prefix: str, filename: str) -> str:
    return f"{prefix.rstrip('/')}/{filename.lstrip('/')}"

# -----------------------------------------------------------------------------#
# Env / model
# -----------------------------------------------------------------------------#
def _require_api_key() -> None:
    if not _sget("GOOGLE_API_KEY"):
        raise EnvironmentError("GOOGLE_API_KEY not found (env or Streamlit secrets)")

def _make_llm(temperature: float) -> ChatGoogleGenerativeAI:
    model = _sget("GEMINI_MODEL", "gemini-2.5-flash") or "gemini-2.5-flash"
    return ChatGoogleGenerativeAI(
        model=model,
        temperature=temperature,
        top_p=0.0,
        top_k=1,
    )

# -----------------------------------------------------------------------------#
# Prompt loaders (S3-only, with optional manifest; no local fallback)
# -----------------------------------------------------------------------------#
def _strip_bom(s: str) -> str:
    return s.lstrip("\ufeff") if s and s.startswith("\ufeff") else s

def _extract_content_block(raw: str) -> str:
    """
    If the YAML contains a top-level `content: |`, return that block;
    otherwise return the whole file. This lets you keep metadata in the YAML.
    """
    raw = raw.strip()
    # naive but practical: find 'content:' at start-of-line
    m = re.search(r"^\s*content\s*:\s*\|?\s*(.*)$", raw, flags=re.S | re.I)
    txt = (m.group(1) if m else raw).replace("\r\n", "\n")
    return txt.strip()

def _read_prompts_manifest(prompts_prefix: str) -> Dict[str, str]:
    """
    Reads prompts manifest JSON at <prompts_prefix>/_manifest.json, if present.
    Accepts any of these formats:
      {"1":"1_intro.yaml","2":"2_tone.yaml"}
      {"map":{"1":"1_intro.yaml", ...}}
      [{"index":1,"file":"1_intro.yaml"}, ...]
    Returns a dict { "1": "1_intro.yaml", ... } with string keys.
    """
    key = _join_key(prompts_prefix, "_manifest.json")
    txt = _s3_get_text(key)
    if not txt:
        return {}
    try:
        j = json.loads(txt)
    except Exception:
        return {}
    if isinstance(j, dict):
        if "map" in j and isinstance(j["map"], dict):
            j = j["map"]
        return {str(k): str(v) for k, v in j.items() if isinstance(v, str) and v}
    if isinstance(j, list):
        out: Dict[str, str] = {}
        for item in j:
            try:
                idx = str(item.get("index"))
                fil = str(item.get("file"))
                if idx and fil:
                    out[idx] = fil
            except Exception:
                continue
        return out
    return {}

def _candidate_prompt_keys(prompts_prefix: str, n: int) -> List[str]:
    # Try exact numeric filenames first (no List required)
    cands = [
        _join_key(prompts_prefix, f"{n}.yaml"),
        _join_key(prompts_prefix, f"{n}.yml"),
        _join_key(prompts_prefix, f"{n:02}.yaml"),
        _join_key(prompts_prefix, f"{n:02}.yml"),
    ]
    # From manifest (no List required)
    man = _read_prompts_manifest(prompts_prefix)
    man_name = man.get(str(n))
    if man_name:
        cands.insert(0, _join_key(prompts_prefix, man_name))
    return cands

def _discover_prompt_via_list(prompts_prefix: str, n: int) -> Optional[str]:
    # Requires ListBucket; search for keys like "n_*.yaml|yml" and zero-padded variants
    keys = _s3_list(prompts_prefix)
    if not keys:
        return None
    nstr = str(n)
    n2 = f"{n:02}"
    best: Optional[str] = None
    for k in keys:
        base = os.path.basename(k).lower()
        if (base.startswith(nstr + "_") or base.startswith(n2 + "_")) and (base.endswith(".yaml") or base.endswith(".yml")):
            best = k
            break
    return best

def _load_prompt(prompts_prefix: str, n: int) -> str:
    """
    Unified S3 loader for prompt index n.
    Search order (no List permission required):
      1) {n}.yaml / {n}.yml / {n:02}.yaml / {n:02}.yml
      2) manifest _manifest.json -> map[n] => filename
      3) If List is allowed: first key matching "n_*.yaml|yml" or "0n_*.yaml|yml"
    """
    if prompts_prefix.startswith("s3://"):
        parts = prompts_prefix.split("/", 3)
        # ignore bucket from s3:// if provided; we use configured bucket instead
        prompts_prefix = parts[3] if len(parts) > 3 else ""

    tried: List[str] = []
    for key in _candidate_prompt_keys(prompts_prefix, n):
        tried.append(key)
        txt = _s3_get_text(key)
        if txt:
            return _extract_content_block(_strip_bom(txt))

    # optional discovery via List
    k = _discover_prompt_via_list(prompts_prefix, n)
    if k:
        tried.append(k)
        txt = _s3_get_text(k)
        if txt:
            return _extract_content_block(_strip_bom(txt))

    raise FileNotFoundError(
        "Prompt {n} not found. Tried keys:\n- ".format(n=n) + "\n- ".join(tried) +
        "\nUpload a YAML to one of those paths or add _manifest.json under the prompts prefix."
    )

# -----------------------------------------------------------------------------#
# Template rendering (supports {var} and {{ var }})
# -----------------------------------------------------------------------------#
_VAR_DBL = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")
_VAR_SGL = re.compile(r"\{([a-zA-Z0-9_]+)\}")

def _render_template(template: str, variables: Dict[str, str]) -> str:
    if not variables:
        return template
    # replace {{ var }}
    def _r1(m: re.Match) -> str:  # type: ignore
        key = m.group(1)
        return str(variables.get(key, ""))
    out = _VAR_DBL.sub(_r1, template)

    # replace {var}
    def _r2(m: re.Match) -> str:  # type: ignore
        key = m.group(1)
        return str(variables.get(key, ""))
    out = _VAR_SGL.sub(_r2, out)
    return out

# -----------------------------------------------------------------------------#
# JSON extraction (for specialists that return plain JSON)
# -----------------------------------------------------------------------------#
_BEGIN = "BEGIN_JSON"
_END = "END_JSON"

def _between_tokens(text: str, start: str, end: str) -> Optional[str]:
    i = text.find(start)
    j = text.rfind(end)
    if i == -1 or j == -1 or j <= i:
        return None
    return text[i + len(start): j].strip()

def _fenced(text: str) -> Optional[str]:
    m = re.search(r"```json\s*(.*?)\s*```", text, flags=re.S | re.I)
    if m: return m.group(1).strip()
    m = re.search(r"```\s*(.*?)\s*```", text, flags=re.S)
    if m: return m.group(1).strip()
    return None

def _balanced(text: str) -> Optional[str]:
    start = text.find("{")
    if start == -1: return None
    depth = 0
    for i in range(start, len(text)):
        ch = text[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[start:i+1]
    return None

def _parse_json(text: str) -> Dict[str, Any]:
    """
    Try multiple strategies to extract valid JSON. Tolerates trailing commas.
    """
    for candidate in filter(None, (_between_tokens(text, _BEGIN, _END),
                                  _fenced(text), _balanced(text))):
        try:
            return json.loads(candidate)
        except Exception:
            try:
                fixed = re.sub(r",\s*([\]}])", r"\1", candidate)
                return json.loads(fixed)
            except Exception:
                continue
    raise ValueError("Could not parse JSON from model output")

# -----------------------------------------------------------------------------#
# Conversions → legacy block for UI
# -----------------------------------------------------------------------------#
def _coerce_explanation(raw: Dict[str, Any]) -> str:
    explanation = raw.get("explanation", "")
    if isinstance(explanation, str) and explanation.strip():
        return explanation.strip()
    bullets = raw.get("explanation_bullets", [])
    if isinstance(bullets, list) and bullets:
        parts = [str(b).strip().rstrip(".") + "." for b in bullets if str(b).strip()]
        return " ".join(parts)[:1400]
    return ""

def _to_legacy_param_block(raw: Dict[str, Any]) -> Dict[str, Any]:
    score = int(raw.get("score", 0) or 0)
    explanation = _coerce_explanation(raw)
    weakness = str(raw.get("weakness", "") or "").strip() or "Not present"

    suggestion = str(raw.get("suggestion", "") or "").strip()
    suggestions_list: List[str] = []
    if not suggestion and isinstance(raw.get("suggestions"), list) and raw["suggestions"]:
        suggestions_list = [str(s).strip() for s in raw["suggestions"] if str(s).strip()]
        suggestion = suggestions_list[0] if suggestions_list else ""
    elif isinstance(raw.get("suggestions"), list):
        suggestions_list = [str(s).strip() for s in raw["suggestions"] if str(s).strip()]
    suggestion = suggestion or "Not present"

    ex: List[str] = []
    aoi = raw.get("areas_of_improvement") or []

    block: Dict[str, Any] = {
        "extractions": ex,
        "score": score,
        "explanation": explanation,
        "weakness": weakness,
        "suggestion": suggestion,
        "areas_of_improvement": aoi,
        "summary": str(raw.get("summary", "") or "").strip(),
    }

    for k in ("dominant_tense", "tense_consistency", "tense_target"):
        if k in raw:
            block[k] = raw[k]

    if suggestions_list:
        block["suggestions_list"] = suggestions_list
    return block

# -----------------------------------------------------------------------------#
# Display-name mapping for UI & aggregator
# -----------------------------------------------------------------------------#
DISPLAY_BY_INDEX = {
    1: "Suspense Building",
    2: "Language/Tone",
    3: "Intro + Main Hook/Cliffhanger",
    4: "Story Structure + Flow",
    5: "Pacing",
    6: "Mini-Hooks (30–60s)",
    7: "Outro (Ending)",
    # 8 = global preamble (loaded, not called)
    9: "Overall Summary (Aggregator)",
    11: "Grammar & Spelling",
}

SPECIALIST_INDEXES: List[int] = [1, 2, 3, 4, 5, 6, 7, 11]

# -----------------------------------------------------------------------------#
# LLM invoke with retries
# -----------------------------------------------------------------------------#
def _invoke_with_retries(llm: ChatGoogleGenerativeAI, prompt: str, tries: int = 3, base_delay: float = 0.8):
    last_err = None
    for k in range(tries):
        try:
            return llm.invoke(prompt)
        except Exception as e:
            last_err = e
            if k < tries - 1:
                time.sleep(base_delay * (2 ** k))
            else:
                raise
    raise last_err  # type: ignore

# -----------------------------------------------------------------------------#
# Aggregator structured schema (model returns this)
# -----------------------------------------------------------------------------#
class AggregatorAll(BaseModel):
    overall_rating: conint(ge=1, le=10)
    strengths: List[str]
    weaknesses: List[str]
    suggestions: List[str]
    drop_off_risks: List[str]
    viral_quotient: str

# -----------------------------------------------------------------------------#
# Signature refs (S3) & style windows
# -----------------------------------------------------------------------------#
try:
    import yaml  # type: ignore
except Exception:  # pragma: no cover
    yaml = None  # type: ignore

try:
    import docx  # type: ignore
except Exception:  # pragma: no cover
    docx = None  # type: ignore

def _load_signature_excerpts_s3() -> Tuple[str, str, List[str]]:
    """
    Loads signature excerpts from S3 key set via env SIGNATURE_EXCERPTS_KEY (default Scriptmodel/reference/signature_excerpts.yaml).
    Returns (dna_profile, dna_excerpts_text, excerpts_list)
    """
    key = _sget("SIGNATURE_EXCERPTS_KEY", "Scriptmodel/reference/signature_excerpts.yaml")
    txt = _s3_get_text(key) if key else None
    dna_profile, excerpts_list = "", []
    if txt and yaml is not None:
        try:
            data = yaml.safe_load(txt) or {}
            dna_profile = str(data.get("dna_profile", "") or "")
            excerpts_list = [str(x) for x in (data.get("excerpts") or []) if str(x).strip()]
        except Exception:
            pass
    dna_excerpts_text = "\n".join(f"- {x}" for x in excerpts_list) if excerpts_list else ""
    return dna_profile, dna_excerpts_text, excerpts_list

def _load_docx_text_s3() -> str:
    """
    Loads signature script DOCX from S3 key SIGNATURE_SCRIPT_KEY (default Scriptmodel/reference/signature_script.docx).
    Returns plain text (paragraph-joined) or empty string on failure.
    """
    key = _sget("SIGNATURE_SCRIPT_KEY", "Scriptmodel/reference/signature_script.docx")
    if not key or docx is None:
        return ""
    b = _s3_get_bytes(key)
    if not b:
        return ""
    try:
        # write to temp, parse with python-docx
        fd, tmp = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        with open(tmp, "wb") as f:
            f.write(b)
        d = docx.Document(tmp)  # type: ignore
        try:
            os.remove(tmp)
        except Exception:
            pass
        return "\n".join(p.text for p in d.paragraphs if p.text and p.text.strip())
    except Exception:
        return ""

def _chunk_text(text: str, target_chars: int = 2500, overlap_chars: int = 250) -> List[str]:
    if not text:
        return []
    paras = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    chunks, buf = [], ""
    for p in paras:
        if len(buf) + len(p) + 2 <= target_chars:
            buf += (("\n\n" if buf else "") + p)
        else:
            if buf:
                chunks.append(buf)
                tail = buf[-overlap_chars:] if overlap_chars else ""
                buf = (tail + "\n\n" + p) if tail else p
            else:
                for i in range(0, len(p), target_chars):
                    chunks.append(p[i:i+target_chars])
                buf = ""
    if buf:
        chunks.append(buf)
    return chunks

_WORD_RE = re.compile(r"[a-zA-Z0-9']+")

def _tf(text: str) -> Counter:
    words = [w.lower() for w in _WORD_RE.findall(text)]
    return Counter(words)

def _cosine(a: Counter, b: Counter) -> float:
    if not a or not b:
        return 0.0
    common = set(a.keys()) & set(b.keys())
    num = sum(a[w] * b[w] for w in common)
    denom1 = sum(v*v for v in a.values()) ** 0.5
    denom2 = sum(v*v for v in b.values()) ** 0.5
    if denom1 == 0 or denom2 == 0:
        return 0.0
    return num / (denom1 * denom2)

def _extract_probes(text: str, max_chars_each: int = 1800, n_probes: int = 3) -> List[str]:
    if not text:
        return []
    L = len(text)
    anchors = [0, L // 2, max(L - max_chars_each, 0)]
    return [text[a:a+max_chars_each] for a in anchors[:n_probes]]

def _trim_excerpt(s: str, max_chars: int = 500) -> str:
    s = re.sub(r"\s+", " ", s or "").strip()
    return s[:max_chars]

def _get_style_windows(signature_text: str,
                       user_script: str,
                       k_per_probe: int = 2,
                       max_windows: int = 4,
                       max_chars: int = 500) -> List[str]:
    if not signature_text or not user_script:
        return []
    chunks = _chunk_text(signature_text, target_chars=2500, overlap_chars=250)
    if not chunks:
        return []
    chunk_tfs = [ _tf(c) for c in chunks ]
    probes = _extract_probes(user_script, max_chars_each=1800, n_probes=3)
    scored: List[Tuple[int, float]] = []
    seen = set()
    for p in probes:
        q = _tf(p)
        sims = [(i, _cosine(q, ct)) for i, ct in enumerate(chunk_tfs)]
        sims.sort(key=lambda x: x[1], reverse=True)
        for i, sim in sims[:k_per_probe]:
            if i not in seen:
                scored.append((i, sim))
                seen.add(i)
    scored.sort(key=lambda x: x[1], reverse=True)
    top = scored[:max_windows]
    return [ "- " + _trim_excerpt(chunks[i], max_chars=max_chars) for (i, _s) in top ]

# -----------------------------------------------------------------------------#
# Core runner
# -----------------------------------------------------------------------------#
def run_review_multi(
    script_text: str,
    prompts_dir: str = "Scriptmodel/prompts",  # S3 prefix only
    temperature: float = 0.0,
    include_commentary: bool = False,  # kept for API parity; ignored
) -> str:
    """
    Execute prompts 1..7 and 11, prepend prompts/8.yaml as a global preamble to each call,
    then 9 as the aggregator. Convert to legacy shape, and return BEGIN_JSON ... END_JSON.

    NOTE: `prompts_dir` is treated as an S3 prefix (e.g., "Scriptmodel/prompts").
    There is NO local fallback in this S3-only build.
    """
    _require_api_key()
    # Ensure S3 is configured
    _require_s3()

    llm = _make_llm(temperature)

    # ---- Load DNA + excerpts from S3 (optional)
    dna_profile, dna_excerpts_text, _excerpts_list = _load_signature_excerpts_s3()

    # ---- Optional: load full signature (S3) and build style windows
    signature_text = _load_docx_text_s3()
    max_windows = int(_sget("STYLE_WINDOWS_PER_CALL", "4") or "4")
    max_chars   = int(_sget("STYLE_WINDOW_MAX_CHARS", "500") or "500")
    signature_style_windows_list = _get_style_windows(
        signature_text,
        script_text,
        k_per_probe=2,
        max_windows=max_windows,
        max_chars=max_chars,
    )
    signature_style_windows = "\n".join(signature_style_windows_list)

    # ---- Base variables for all templates
    base_vars: Dict[str, str] = {
        "script": script_text,
        "dna_profile": dna_profile,
        "dna_excerpts_text": dna_excerpts_text,
        "signature_style_windows": signature_style_windows,
    }

    # Load global preamble (8.yaml) if present and render with DNA/anchors
    try:
        global_preamble_tmpl = _load_prompt(prompts_dir, 8).strip()
        if global_preamble_tmpl:
            global_preamble = _render_template(global_preamble_tmpl, base_vars) + "\n\n"
        else:
            global_preamble = ""
    except FileNotFoundError:
        global_preamble = ""

    # 1..7 + 11 specialists
    scores: Dict[str, int] = {}
    per_parameter: Dict[str, Dict[str, Any]] = {}

    for i in SPECIALIST_INDEXES:
        name = DISPLAY_BY_INDEX[i]
        try:
            tmpl = _load_prompt(prompts_dir, i)
        except FileNotFoundError as e:
            raise FileNotFoundError(f"{e}\nTip: put '{i}.yaml' (or set _manifest.json) under s3://{_RP_BUCKET}/{prompts_dir}/")

        prompt_body = _render_template(tmpl, base_vars)
        prompt = f"{global_preamble}{prompt_body}"

        try:
            resp = _invoke_with_retries(llm, prompt)
            raw_text = getattr(resp, "content", "") or ""
            data = _parse_json(raw_text)
        except Exception as e:
            short = (str(e) or "unknown").strip()
            raise RuntimeError(f"JSON parse failed on prompt {i} ({name}). Error: {short}")

        block = _to_legacy_param_block(data)
        scores[name] = int(block.get("score", 0))
        per_parameter[name] = block

    # Build evidence for aggregator (legacy shape only)
    evidence = {"scores": scores, "per_parameter": per_parameter}
    evidence_json = json.dumps(evidence, ensure_ascii=False)

    # 9 = merged meta-synthesis + aggregator (MODEL decides overall_rating)
    try:
        tmpl9 = _load_prompt(prompts_dir, 9)
    except FileNotFoundError as e:
        raise FileNotFoundError(f"{e}\nTip: ensure '9.yaml' exists under s3://{_RP_BUCKET}/{prompts_dir}/")

    prompt9_body = _render_template(
        tmpl9,
        {
            **base_vars,
            "evidence_json": evidence_json,
        }
    )
    prompt9 = f"{global_preamble}{prompt9_body}"

    # Ask for structured output (enforces keys and types)
    try:
        llm_aggr = _make_llm(temperature).with_structured_output(AggregatorAll)
        agg: AggregatorAll = llm_aggr.invoke(prompt9)
    except Exception as e:
        short = (str(e) or "unknown").strip()
        raise RuntimeError(f"Aggregator failed on prompt 9. Error: {short}")

    # Build final payload: use model-decided overall_rating
    final_payload: Dict[str, Any] = {
        "scores": scores,
        "per_parameter": per_parameter,
        "overall_rating": int(agg.overall_rating),
        "strengths": agg.strengths,
        "weaknesses": agg.weaknesses,
        "suggestions": agg.suggestions,
        "drop_off_risks": agg.drop_off_risks,
        "viral_quotient": agg.viral_quotient,
    }

    # Normalize & sanitize everything before returning
    final_payload = normalize_review_payload(final_payload)

    return _wrap_json(final_payload)

def _wrap_json(payload: Dict[str, Any]) -> str:
    return f"{_BEGIN}\n{json.dumps(payload, ensure_ascii=False)}\n{_END}\n"



