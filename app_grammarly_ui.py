

# import os, re, glob, json, tempfile, difflib, uuid, datetime
# from pathlib import Path
# from typing import Dict, Any, List, Tuple, Optional

# import streamlit as st
# import pandas as pd

# from utils1 import extract_review_json, PARAM_ORDER, load_script_file
# from review_engine_multi import run_review_multi

# # ---- DOCX rendering imports (already in requirements via python-docx) ----
# from docx import Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # ---------- Folders ----------
# SCRIPTS_DIR = "scripts"
# PROMPTS_DIR = "prompts"
# OUTPUT_DIR  = "outputs"
# HISTORY_DIR = os.path.join(OUTPUT_DIR, "_history")
# for p in (SCRIPTS_DIR, PROMPTS_DIR, OUTPUT_DIR, HISTORY_DIR):
#     Path(p).mkdir(parents=True, exist_ok=True)

# # ---------- Colors ----------
# PARAM_COLORS: Dict[str, str] = {
#     "Suspense Building":              "#ff6b6b",
#     "Language/Tone":                  "#6b8cff",
#     "Intro + Main Hook/Cliffhanger":  "#ffb86b",
#     "Story Structure + Flow":         "#a78bfa",
#     "Pacing":                         "#f43f5e",
#     "Mini-Hooks (30–60s)":            "#eab308",
#     "Outro (Ending)":                 "#8b5cf6",
# }

# # ---------- Config ----------
# STRICT_MATCH_ONLY = False  # set True to disable fuzzy sentence fallback entirely

# # ---------- App config ----------
# st.set_page_config(page_title="M7 — Grammarly UI", page_icon="🕵️", layout="wide")

# # ---------- Header patch & Recents card CSS ----------
# def render_app_title():
#     st.markdown(
#         '<h1 class="app-title">Viral Script Reviewer</h1>',
#         unsafe_allow_html=True
#     )
#     st.markdown("""
#     <style>
#     :root { --sep:#e5e7eb; }

#     /* Push content below Streamlit sticky header (covers multiple versions) */
#     .stApp .block-container { padding-top: 4.25rem !important; }

#     /* App title */
#     .app-title{
#       font-weight: 700;
#       font-size: 2.1rem;
#       line-height: 1.3;
#       margin: 0 0 1rem 0;
#       padding-left: 40px !important;
#       padding-top: .25rem !important;
#       white-space: normal;
#       word-break: break-word;
#       hyphens: auto;
#       overflow: visible;
#       position: relative !important;
#       z-index: 10 !important;
#     }
#     [data-testid="collapsedControl"] { z-index: 6 !important; }
#     header[data-testid="stHeader"], .stAppHeader {
#       background: transparent !important;
#       box-shadow: none !important;
#     }
#     @media (min-width: 992px){
#       .app-title { padding-left: 0 !important; }
#     }

#     /* vertical dividers between columns */
#     div[data-testid="column"]:nth-of-type(1){position:relative;}
#     div[data-testid="column"]:nth-of-type(1)::after{
#       content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);
#     }
#     div[data-testid="column"]:nth-of-type(2){position:relative;}
#     div[data-testid="column"]:nth-of-type(2)::after{
#       content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);
#     }

#     /* parameter chips */
#     .param-row{display:flex;gap:8px;overflow-x:auto;padding:0 2px 8px;}
#     .param-chip{padding:6px 10px;border-radius:999px;font-size:13px;white-space:nowrap;border:1px solid transparent;cursor:pointer;}

#     /* center script container — page scrolls; no inner scroll */
#     .docxwrap{
#       border:1px solid #eee;
#       border-radius:8px;
#       padding:16px 14px 18px;
#     }

#     /* Typography similar to Word */
#     .docxwrap .h1, .docxwrap .h2, .docxwrap .h3 { font-weight:700; margin: 10px 0 6px; }
#     .docxwrap .h1 { font-size: 1.3rem; border-bottom: 2px solid #000; padding-bottom: 4px; }
#     .docxwrap .h2 { font-size: 1.15rem; border-bottom: 1px solid #000; padding-bottom: 3px; }
#     .docxwrap .h3 { font-size: 1.05rem; }
#     .docxwrap p { margin: 10px 0; line-height: 1.7; font-family: ui-serif, Georgia, "Times New Roman", serif; }

#     /* inline styles */
#     .docxwrap strong { font-weight: 700; }
#     .docxwrap em { font-style: italic; }
#     .docxwrap u { text-decoration: underline; }

#     /* DOCX-like table rendering */
#     .docxwrap table { border-collapse: collapse; width: 100%; margin: 12px 0; }
#     .docxwrap th, .docxwrap td {
#       border: 1px solid #bbb;
#       padding: 8px;
#       vertical-align: top;
#       font-family: ui-serif, Georgia, "Times New Roman", serif;
#       line-height: 1.6;
#     }

#     /* AOI highlight */
#     .docxwrap mark{
#       padding:0 2px;
#       border-radius:3px;
#       border:1px solid rgba(0,0,0,.12);
#     }

#     /* Recents: nicer card borders */
#     .rec-card{
#       border:1px solid #e5e7eb;
#       border-radius:10px;
#       padding:14px 16px;
#       margin: 10px 0 16px 0;
#       box-shadow: 0 1px 0 rgba(0,0,0,.02);
#       background:#fff;
#     }
#     .rec-title{font-weight:600; margin-bottom:.25rem;}
#     .rec-meta{color:#6b7280; font-size:12.5px; margin-bottom:.4rem;}
#     .rec-row{display:flex; align-items:center; justify-content:space-between; gap:12px;}
#     </style>
#     """, unsafe_allow_html=True)

# render_app_title()

# # ---------- Session ----------
# for key, default in [
#     ("review_ready", False),
#     ("script_text", ""),
#     ("base_stem", ""),
#     ("data", None),
#     ("spans_by_param", {}),
#     ("param_choice", None),
#     ("source_docx_path", None),
#     ("heading_ranges", []),  # exact offsets of DOCX headings for suppression
#     ("flattened_docx_path", None),   # temp flattened DOCX path (if created)
#     ("flatten_used", False),         # whether we flattened for this run
#     ("ui_mode", "home"),             # "home" | "review" | "recents"
# ]:
#     st.session_state.setdefault(key, default)

# # ---------- Sanitizer (for editor meta, not for the script body) ----------
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )
# def _sanitize_editor_text(s: Optional[str]) -> str:
#     if not s:
#         return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•)\s*', '', t, flags=re.M)
#     t = re.sub(r'^\s*[-*]\s+', '• ', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# # ---------- DOCX traversal ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# # ---------- Auto-flatten helpers ----------
# def _docx_contains_tables(path: str) -> bool:
#     doc = Document(path)
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Table):
#             return True
#     return False

# def _copy_paragraph(dest_doc: Document, src_para: Paragraph):
#     """Copy a paragraph with inline bold/italic/underline."""
#     p = dest_doc.add_paragraph()
#     try:
#         if src_para.style and src_para.style.name:
#             p.style = src_para.style.name
#     except Exception:
#         pass
#     for run in src_para.runs:
#         r = p.add_run(run.text or "")
#         r.bold = run.bold
#         r.italic = run.italic
#         r.underline = run.underline
#     return p

# def flatten_docx_tables_to_longtext(source_path: str) -> str:
#     """
#     Create a paragraph-only DOCX from a tabular DOCX:
#     - Preserves order and inline styles
#     - De-dupes merged cells
#     - Adds a blank line between rows and after each table
#     """
#     src = Document(source_path)
#     new = Document()

#     for blk in _iter_docx_blocks(src):
#         if isinstance(blk, Paragraph):
#             _copy_paragraph(new, blk)
#         else:  # Table
#             seen_tc_ids = set()
#             for row in blk.rows:
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     if tc_id in seen_tc_ids:
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     for p in cell.paragraphs:
#                         _copy_paragraph(new, p)
#                 new.add_paragraph("")  # blank line between rows
#             new.add_paragraph("")      # extra blank line after table

#     fd, tmp_path = tempfile.mkstemp(suffix=".docx")
#     os.close(fd)
#     new.save(tmp_path)
#     return tmp_path

# # ---------- merged-cell–aware, heading-in-table aware builder ----------
# def build_docx_text_with_meta(docx_path: str) -> Tuple[str, List[Tuple[int,int]]]:
#     """
#     Build plain text (offset-aligned with HTML rendering) AND collect exact
#     character-offset ranges for Heading paragraphs (including those inside TABLE cells)
#     to suppress AOI highlights there. Skips duplicated merged cells so offsets stay stable.
#     """
#     doc = Document(docx_path)
#     out: List[str] = []
#     heading_ranges: List[Tuple[int,int]] = []
#     current_offset = 0

#     def _append_and_advance(s: str):
#         nonlocal current_offset
#         out.append(s)
#         current_offset += len(s)

#     seen_tc_ids: set = set()

#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             para_text = "".join(run.text or "" for run in blk.runs)
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             if sty.startswith("heading"):
#                 start = current_offset
#                 end   = start + len(para_text)
#                 heading_ranges.append((start, end))
#             _append_and_advance(para_text)
#             _append_and_advance("\n")  # paragraph separator

#         else:  # Table
#             for row in blk.rows:
#                 row_cell_tcs = []
#                 for cell in row.cells:
#                     tc = cell._tc
#                     tc_id = id(tc)
#                     row_cell_tcs.append((tc_id, cell))

#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     if tc_id in seen_tc_ids:
#                         if idx != len(row_cell_tcs) - 1:
#                             _append_and_advance("\t")
#                         continue

#                     seen_tc_ids.add(tc_id)
#                     cell_text_parts: List[str] = []
#                     for i, p in enumerate(cell.paragraphs):
#                         t = "".join(r.text or "" for r in p.runs)
#                         sty = (p.style.name or "").lower() if p.style else ""
#                         if sty.startswith("heading"):
#                             hs = current_offset + sum(len(x) for x in cell_text_parts)
#                             he = hs + len(t)
#                             heading_ranges.append((hs, he))
#                         cell_text_parts.append(t)
#                         if i != len(cell.paragraphs) - 1:
#                             cell_text_parts.append("\n")
#                     cell_text = "".join(cell_text_parts)
#                     _append_and_advance(cell_text)

#                     if idx != len(row_cell_tcs) - 1:
#                         _append_and_advance("\t")

#                 _append_and_advance("\n")
#             _append_and_advance("\n")

#     return "".join(out), heading_ranges

# def _wrap_inline(safe_text: str, run) -> str:
#     out = safe_text
#     if getattr(run, "underline", False):
#         out = f"<u>{out}</u>"
#     if getattr(run, "italic", False):
#         out = f"<em>{out}</em>"
#     if getattr(run, "bold", False):
#         out = f"<strong>{out}</strong>"
#     return out

# # ---------- Invisible/bridge characters ----------
# _BRIDGE_CHARS = set("\u200b\u200c\u200d\u2060\ufeff\xa0\u00ad")  # ZW*, WJ, BOM, NBSP, soft hyphen

# # ---------- renderer with merged-cell parity ----------
# def render_docx_html_with_highlights(docx_path: str,
#                                      highlight_spans: List[Tuple[int,int,str,str]]) -> str:
#     """
#     Build HTML from DOCX runs and inject <mark> by global character offsets.
#     Keeps a single <mark> open across multiple runs until span ends.
#     Mirrors build_docx_text_with_meta() linearization.
#     """
#     doc = Document(docx_path)
#     spans = [s for s in highlight_spans if s[0] < s[1]]
#     spans.sort(key=lambda x: x[0])

#     cur_span = 0
#     current_offset = 0

#     def esc(s: str) -> str:
#         return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

#     def open_mark_if_needed(html_parts, mark_state, color, end):
#         if not mark_state["open"]:
#             html_parts.append(
#                 f'<mark style="background:{color}33;border:1px solid {color};border-radius:3px;padding:0 2px;">'
#             )
#             mark_state.update(open=True, end=end, color=color)

#     def close_mark_if_open(html_parts, mark_state):
#         if mark_state["open"]:
#             html_parts.append('</mark>')
#             mark_state.update(open=False, end=None, color=None)

#     def emit_run_text(run_text: str, run, html_parts: List[str], mark_state: Dict[str, Any]):
#         nonlocal cur_span, current_offset
#         t = run_text or ""
#         i = 0
#         while i < len(t):
#             next_start, next_end, color = None, None, None
#             if cur_span < len(spans):
#                 next_start, next_end, color, _aid = spans[cur_span]

#             if not mark_state["open"]:
#                 if cur_span >= len(spans) or current_offset + (len(t) - i) <= next_start:
#                     chunk = t[i:]
#                     html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += len(chunk)
#                     break
#                 if current_offset < next_start:
#                     take = next_start - current_offset
#                     chunk = t[i:i+take]
#                     html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take
#                     continue
#                 open_mark_if_needed(html_parts, mark_state, color, next_end)
#             else:
#                 take = min(mark_state["end"] - current_offset, len(t) - i)
#                 if take > 0:
#                     chunk = t[i:i+take]
#                     html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take
#                 if current_offset >= mark_state["end"]:
#                     close_mark_if_open(html_parts, mark_state)
#                     cur_span += 1

#     html: List[str] = ['<div class="docxwrap">']
#     seen_tc_ids: set = set()

#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             mark_state = {"open": False, "end": None, "color": None}
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             open_tag = '<div class="h1">' if sty.startswith("heading 1") else \
#                        '<div class="h2">' if sty.startswith("heading 2") else \
#                        '<div class="h3">' if sty.startswith("heading 3") else "<p>"
#             close_tag = "</div>" if sty.startswith("heading") else "</p>"

#             html.append(open_tag)
#             for run in blk.runs:
#                 emit_run_text(run.text or "", run, html, mark_state)
#             close_mark_if_open(html, mark_state)
#             html.append(close_tag)
#             current_offset += 1  # '\n'

#         else:  # Table
#             html.append("<table>")
#             for row in blk.rows:
#                 html.append("<tr>")
#                 row_cell_tcs = [(id(cell._tc), cell) for cell in row.cells]

#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     html.append("<td>")
#                     if tc_id not in seen_tc_ids:
#                         seen_tc_ids.add(tc_id)
#                         for p_idx, p in enumerate(cell.paragraphs):
#                             mark_state = {"open": False, "end": None, "color": None}
#                             html.append("<div>")
#                             for run in p.runs:
#                                 emit_run_text(run.text or "", run, html, mark_state)
#                             close_mark_if_open(html, mark_state)
#                             html.append("</div>")
#                             if p_idx != len(cell.paragraphs) - 1:
#                                 current_offset += 1
#                     html.append("</td>")
#                     if idx != len(row_cell_tcs) - 1:
#                         current_offset += 1  # '\t'
#                 html.append("</tr>")
#                 current_offset += 1   # row '\n'
#             html.append("</table>")
#             current_offset += 1       # extra '\n'

#     html.append("</div>")
#     return "".join(html)

# # ---------- Highlight search (on plain text) ----------
# def _normalize_keep_len(s: str) -> str:
#     trans = {
#         "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
#         "\u2013": "-", "\u2014": "-",
#         "\xa0": " ",
#         "\u200b": " ", "\u200c": " ", "\u200d": " ", "\u2060": " ",
#         "\ufeff": " ", "\u00ad": " ",
#     }
#     return (s or "").translate(str.maketrans(trans))

# def _tokenize(s: str) -> List[str]:
#     return re.findall(r"\w+", (s or "").lower())

# def _iter_sentences_with_spans(text: str) -> List[Tuple[int,int,str]]:
#     spans = []
#     for m in re.finditer(r'[^.!?]+[.!?]+|\Z', text, flags=re.S):
#         s, e = m.start(), m.end()
#         seg = text[s:e]
#         if seg.strip():
#             spans.append((s, e, seg))
#     return spans

# def _squash_ws(s: str) -> str:
#     return re.sub(r"\s+", " ", s or "").strip()

# def _clean_quote_for_match(q: str) -> str:
#     if not q:
#         return ""
#     q = _normalize_keep_len(q).strip()
#     q = re.sub(r'^[\'"“”‘’\[\(\{<…\-\–\—\s]+', '', q)
#     q = re.sub(r'[\'"“”‘’\]\)\}>…\-\–\—\s]+$', '', q)
#     return _squash_ws(q)

# def _snap_and_bridge_to_word(text: str, start: int, end: int, max_bridge: int = 2) -> Tuple[int,int]:
#     n = len(text)
#     s, e = max(0, start), max(start, end)

#     def _is_invisible_ws(ch: str) -> bool:
#         return ch in _BRIDGE_CHARS

#     while s > 0:
#         prev = text[s-1]
#         cur  = text[s] if s < n else ""
#         if prev.isalnum() and cur.isalnum():
#             s -= 1
#             continue
#         j = s; brid = 0
#         while j < n and _is_invisible_ws(text[j]):
#             brid += 1; j += 1
#             if brid > max_bridge: break
#         if brid and (s-1) >= 0 and text[s-1].isalnum() and (j < n and text[j].isalnum()):
#             s -= 1; continue
#         break

#     while e < n:
#         prev = text[e-1] if e > 0 else ""
#         nxt  = text[e]
#         if prev.isalnum() and nxt.isalnum():
#             e += 1
#             continue
#         j = e; brid = 0
#         while j < n and _is_invisible_ws(text[j]):
#             brid += 1; j += 1
#             if brid > max_bridge: break
#         if brid and (e-1) >= 0 and text[e-1].isalnum() and (j < n and text[j].isalnum()):
#             e = j + 1; continue
#         break

#     while e < n and text[e] in ',"”’\')]}':
#         e += 1
#     return s, e

# def _heal_split_word_left(text: str, start: int) -> int:
#     i = start
#     if i <= 1 or i >= len(text): return start
#     if text[i-1] != " ": return start
#     j = i - 2
#     while j >= 0 and text[j].isalpha():
#         j -= 1
#     prev_token = text[j+1:i-1]
#     if len(prev_token) == 1:
#         return i - 2
#     return start

# def _overlaps_any(s: int, e: int, ranges: List[Tuple[int,int]]) -> bool:
#     for rs, re_ in ranges:
#         if e > rs and s < re_:
#             return True
#     return False

# def _fuzzy_window_span(tl: str, nl: str, start: int, w: int) -> Tuple[float, Optional[Tuple[int,int]]]:
#     window = tl[start:start+w]
#     sm = difflib.SequenceMatcher(a=nl, b=window)
#     blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
#     if not blocks:
#         return 0.0, None
#     coverage = sum(b.size for b in blocks) / max(1, len(nl))
#     first_b = min(blocks, key=lambda b: b.b)
#     last_b  = max(blocks, key=lambda b: b.b + b.size)
#     s = start + first_b.b
#     e = start + last_b.b + last_b.size
#     return coverage, (s, e)

# def find_span_smart(text: str, needle: str) -> Optional[Tuple[int,int]]:
#     if not text or not needle:
#         return None

#     t_orig = text
#     t_norm = _normalize_keep_len(text)
#     n_norm = _clean_quote_for_match(needle)
#     if not n_norm:
#         return None

#     tl = t_norm.lower()
#     nl = n_norm.lower()

#     i = tl.find(nl)
#     if i != -1:
#         s, e = _snap_and_bridge_to_word(t_orig, i, i + len(nl))
#         s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     m = re.search(re.escape(nl).replace(r"\ ", r"\s+"), tl, flags=re.IGNORECASE)
#     if m:
#         s, e = _snap_and_bridge_to_word(t_orig, m.start(), m.end())
#         s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     if not STRICT_MATCH_ONLY and len(nl) >= 12:
#         w = max(60, min(240, len(nl) + 80))
#         best_cov, best_span = 0.0, None
#         step = max(1, w // 2)
#         for start in range(0, max(1, len(tl) - w + 1), step):
#             cov, se = _fuzzy_window_span(tl, nl, start, w)
#             if cov > best_cov:
#                 best_cov, best_span = cov, se
#         if best_span and best_cov >= 0.65:
#             s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
#             if s > 0 and t_orig[s-1:s+1].lower() in {" the", " a", " an"}:
#                 s -= 1
#             s = _heal_split_word_left(t_orig, s)
#             return (s, e)

#     if not STRICT_MATCH_ONLY:
#         keys = [w for w in _tokenize(nl) if len(w) >= 4][:8]
#         if len(keys) >= 2:
#             kset = set(keys)
#             best_score, best_span = 0.0, None
#             for s, e, seg in _iter_sentences_with_spans(t_norm):
#                 toks = set(_tokenize(seg))
#                 ov = len(kset & toks)
#                 if ov == 0: continue
#                 score = ov / max(2, len(kset))
#                 length_pen = min(1.0, 120 / max(20, e - s))
#                 score *= (0.6 + 0.4 * length_pen)
#                 if score > best_score:
#                     best_score, best_span = score, (s, min(e, s + 400))
#             if best_span and best_score >= 0.35:
#                 s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
#                 s = _heal_split_word_left(t_orig, s)
#                 return (s, e)
#     return None

# def merge_overlaps(spans: List[Tuple[int,int,str,str]]) -> List[Tuple[int,int,str,str]]:
#     if not spans:
#         return []
#     spans.sort(key=lambda x: x[0])
#     out = [spans[0]]
#     for s,e,c,aid in spans[1:]:
#         ps,pe,pc,paid = out[-1]
#         if s <= pe and pc == c and e > pe:
#             out[-1] = (ps, e, pc, paid)
#         else:
#             out.append((s,e,c,aid))
#     return out

# _PUNCT_WS = set(" \t\r\n,.;:!?)('\"[]{}-—–…") | _BRIDGE_CHARS

# def merge_overlaps_and_adjacent(base_text: str,
#                                 spans: List[Tuple[int,int,str,str]],
#                                 max_gap: int = 2) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans = sorted(spans, key=lambda x: x[0])
#     out = [spans[0]]
#     for s, e, c, aid in spans[1:]:
#         ps, pe, pc, paid = out[-1]
#         if c == pc and s <= pe:
#             out[-1] = (ps, max(pe, e), pc, paid); continue
#         if c == pc and s - pe <= max_gap:
#             gap = base_text[max(0, pe):max(0, s)]
#             if all((ch in _PUNCT_WS) for ch in gap):
#                 out[-1] = (ps, e, pc, paid); continue
#         out.append((s, e, c, aid))
#     return out

# # ---------- Heading filters ----------
# def _is_heading_like(q: str) -> bool:
#     if not q: return True
#     s = q.strip()
#     if not re.search(r'[.!?]', s):
#         words = re.findall(r"[A-Za-z]+", s)
#         if 1 <= len(words) <= 7:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.8:
#                 return True
#         if s.lower() in {"introduction","voiceover","outro","epilogue","prologue","credits","title","hook","horrifying discovery","final decision"}:
#             return True
#         if len(s) <= 3: return True
#     return False

# def _is_heading_context(script_text: str, s: int, e: int) -> bool:
#     left = script_text.rfind("\n", 0, s) + 1
#     right = script_text.find("\n", e);  right = len(script_text) if right == -1 else right
#     line = script_text[left:right].strip()
#     if len(line) <= 70 and not re.search(r'[.!?]', line):
#         words = re.findall(r"[A-Za-z]+", line)
#         if 1 <= len(words) <= 8:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.7:
#                 return True
#     return False

# def _tighten_to_quote(script_text: str, span: Tuple[int,int], quote: str) -> Tuple[int,int]:
#     if not span or not quote:
#         return span
#     s, e = span
#     if e <= s or s < 0 or e > len(script_text):
#         return span

#     window = script_text[s:e]
#     win_norm = _normalize_keep_len(window).lower()
#     q_norm = _clean_quote_for_match(quote).lower()
#     if not q_norm:
#         return span

#     i = win_norm.find(q_norm)
#     if i == -1:
#         m = re.search(re.escape(q_norm).replace(r"\ ", r"\s+"), win_norm, flags=re.IGNORECASE)
#         if not m:
#             return span
#         i, j = m.start(), m.end()
#     else:
#         j = i + len(q_norm)

#     s2, e2 = s + i, s + j
#     s2, e2 = _snap_and_bridge_to_word(script_text, s2, e2)
#     s2 = _heal_split_word_left(script_text, s2)
#     if s2 >= s and e2 <= e and e2 > s2:
#         return (s2, e2)
#     return span

# def build_spans_by_param(script_text: str, data: dict, heading_ranges: Optional[List[Tuple[int,int]]] = None) -> Dict[str, List[Tuple[int,int,str,str]]]:
#     heading_ranges = heading_ranges or []
#     raw = (data or {}).get("per_parameter", {}) or {}
#     per: Dict[str, Dict[str, Any]] = {k:(v or {}) for k,v in raw.items()}
#     spans_map: Dict[str, List[Tuple[int,int,str,str]]] = {p: [] for p in PARAM_ORDER}
#     st.session_state["aoi_match_ranges"] = {}

#     for p in spans_map.keys():
#         color = PARAM_COLORS.get(p, "#ffd54f")
#         blk = per.get(p, {}) or {}
#         aois = blk.get("areas_of_improvement") or []
#         for idx, item in enumerate(aois, start=1):
#             raw_q = (item or {}).get("quote_verbatim", "") or ""
#             q = _sanitize_editor_text(raw_q)
#             clean = _clean_quote_for_match(re.sub(r"^[•\-\d\.\)\s]+", "", q).strip())
#             if not clean: continue
#             if _is_heading_like(clean): continue

#             pos = find_span_smart(script_text, clean)
#             if not pos: continue
#             pos = _tighten_to_quote(script_text, pos, raw_q)
#             s, e = pos

#             if heading_ranges and _overlaps_any(s, e, heading_ranges): continue
#             if _is_heading_context(script_text, s, e): continue

#             aid = f"{p.replace(' ','_')}-AOI-{idx}"
#             spans_map[p].append((s, e, color, aid))
#             st.session_state["aoi_match_ranges"][aid] = (s, e)
#     return spans_map

# # ---------- History (save + load + recents UI) ----------

# def _save_history_snapshot(
#     title: str,
#     data: dict,
#     script_text: str,
#     source_docx_path: Optional[str],
#     heading_ranges: List[Tuple[int,int]],
#     spans_by_param: Dict[str, List[Tuple[int,int,str,str]]],
#     aoi_match_ranges: Dict[str, Tuple[int,int]]
# ):
#     run_id = str(uuid.uuid4())
#     now = datetime.datetime.now()
#     created_at_iso = now.replace(microsecond=0).isoformat()  # local ISO
#     created_at_human = now.strftime("%Y-%m-%d %H:%M:%S")

#     blob = {
#         "run_id": run_id,
#         "title": title or "untitled",
#         "created_at": created_at_iso,              # always a STRING for stability
#         "created_at_human": created_at_human,
#         "overall_rating": (data or {}).get("overall_rating", ""),
#         "scores": (data or {}).get("scores", {}),
#         "data": data or {},
#         "script_text": script_text or "",
#         "source_docx_path": source_docx_path,      # may be None
#         "heading_ranges": heading_ranges or [],
#         "spans_by_param": spans_by_param or {},
#         "aoi_match_ranges": aoi_match_ranges or {},
#     }
#     out_fp = os.path.join(HISTORY_DIR, f"{created_at_iso.replace(':','-')}__{run_id}.json")
#     with open(out_fp, "w", encoding="utf-8") as f:
#         json.dump(blob, f, ensure_ascii=False, indent=2)

# def _load_all_history() -> List[dict]:
#     """
#     Load all saved runs from outputs/_history/*.json.
#     Backward-compatible:
#       - If created_at is numeric (epoch), convert to ISO string.
#       - If created_at is missing/invalid, use file mtime.
#       - Always return created_at as string so sorting never mixes str/float.
#     """
#     out = []
#     for fp in sorted(glob.glob(os.path.join(HISTORY_DIR, "*.json"))):
#         try:
#             with open(fp, "r", encoding="utf-8") as f:
#                 j = json.load(f)
#         except Exception:
#             continue

#         j.setdefault("_path", fp)

#         ca = j.get("created_at")
#         try:
#             if isinstance(ca, (int, float)):
#                 dt = datetime.datetime.utcfromtimestamp(float(ca))
#                 j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                 if not j.get("created_at_human"):
#                     j["created_at_human"] = dt.astimezone().strftime("%Y-%m-%d %H:%M:%S")
#             elif isinstance(ca, str) and ca:
#                 # already ISO-ish; keep
#                 pass
#             else:
#                 # fallback to file mtime
#                 mtime = os.path.getmtime(fp)
#                 dt = datetime.datetime.fromtimestamp(mtime)
#                 j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                 if not j.get("created_at_human"):
#                     j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#         except Exception:
#             j["created_at"] = str(ca or "")

#         out.append(j)

#     out.sort(key=lambda r: r.get("created_at") or "", reverse=True)
#     return out

# def _render_recents_centerpane():
#     """Center-pane browser for saved runs. No script preview shown."""
#     st.subheader("📄 Recents")

#     q = st.text_input("Filter by title…", "")

#     cols = st.columns([1, 4])
#     with cols[0]:
#         if st.button("← Back"):
#             st.session_state.ui_mode = "home"
#             st.rerun()

#     # list
#     recs = _load_all_history()
#     ql = q.strip().lower()
#     if ql:
#         recs = [r for r in recs if ql in (r.get("title","").lower())]

#     if not recs:
#         st.caption("No history yet."); st.stop()

#     # draw cards (no snippet/intro text)
#     for rec in recs:
#         run_id = rec.get("run_id")
#         title  = rec.get("title") or "(untitled)"
#         created_h = rec.get("created_at_human", "")
#         overall = rec.get("overall_rating", "")
#         with st.container():
#             st.markdown(f"""
#             <div class="rec-card">
#               <div class="rec-row">
#                 <div>
#                   <div class="rec-title">{title}</div>
#                   <div class="rec-meta">{created_h}</div>
#                   <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
#                 </div>
#                 <div>
#             """, unsafe_allow_html=True)

#             open_key = f"open_{run_id}"
#             st.button("Open", key=open_key)

#             st.markdown("</div></div></div>", unsafe_allow_html=True)

#             if st.session_state.get(open_key):
#                 # Load and restore this run
#                 path = rec.get("_path")
#                 if path and os.path.exists(path):
#                     try:
#                         with open(path, "r", encoding="utf-8") as f:
#                             jj = json.load(f)
#                     except Exception:
#                         st.error("Could not load this run file."); st.stop()

#                     st.session_state.script_text      = jj.get("script_text","")
#                     st.session_state.base_stem        = jj.get("title","untitled")
#                     st.session_state.data             = jj.get("data",{})
#                     st.session_state.heading_ranges   = jj.get("heading_ranges",[])
#                     st.session_state.spans_by_param   = jj.get("spans_by_param",{})
#                     st.session_state.param_choice     = None
#                     st.session_state.source_docx_path = jj.get("source_docx_path")
#                     st.session_state.review_ready     = True
#                     st.session_state["aoi_match_ranges"] = jj.get("aoi_match_ranges", {})
#                     st.session_state.ui_mode          = "review"
#                     st.rerun()

# # ---------- Sidebar ----------
# with st.sidebar:
#     st.subheader("Run Settings")
#     st.caption("Prompts dir"); st.code(os.path.abspath(PROMPTS_DIR))

#     if st.button("📁 Recents", use_container_width=True):
#         st.session_state.ui_mode = "recents"
#         st.rerun()

#     if st.button("🆕 New review", use_container_width=True):
#         # cleanup temp flattened docx if we created one
#         fp = st.session_state.get("flattened_docx_path")
#         if fp and os.path.exists(fp):
#             try: os.remove(fp)
#             except Exception: pass
#         for k in ["review_ready","script_text","base_stem","data","spans_by_param","param_choice",
#                   "source_docx_path","heading_ranges","flattened_docx_path","flatten_used"]:
#             st.session_state[k] = (False if k=="review_ready"
#                                    else "" if k in ("script_text","base_stem")
#                                    else {} if k=="spans_by_param"
#                                    else [] if k=="heading_ranges"
#                                    else None if k in ("source_docx_path","flattened_docx_path")
#                                    else False if k=="flatten_used"
#                                    else None)
#         st.session_state.ui_mode = "home"
#         st.rerun()

# # ---------- Input screen ----------
# def render_home():
#     st.subheader("🎬 Script Source")
#     tab_pick, tab_upload, tab_paste = st.tabs(["Pick from scripts/ folder", "Upload file", "Paste raw text"])
#     selected_path = None; uploaded_file = None; pasted_text = None

#     with tab_pick:
#         files = [f for f in glob.glob(os.path.join(SCRIPTS_DIR, "*")) if f.lower().endswith((".txt",".docx",".pdf"))]
#         rel = [os.path.relpath(f, SCRIPTS_DIR) for f in files]
#         choice = st.selectbox("Choose a file", options=["— Select —"] + rel, index=0)
#         if choice != "— Select —":
#             selected_path = os.path.join(SCRIPTS_DIR, choice)
#         if st.button("🔁 Refresh list"):
#             st.rerun()

#     with tab_upload:
#         up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
#         if up is not None:
#             suffix = os.path.splitext(up.name)[1].lower()
#             with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
#                 tmp.write(up.read()); tmp_path = tmp.name
#             uploaded_file = tmp_path

#     with tab_paste:
#         pasted_text = st.text_area("Paste VO / narration here", height=300)

#     if st.button("🚀 Run Review", type="primary", use_container_width=True):
#         base_stem = "pasted_script"; source_docx_path = None; heading_ranges = []

#         if selected_path:
#             base_stem = os.path.splitext(os.path.basename(selected_path))[0]
#             if selected_path.lower().endswith(".docx"):
#                 path_to_use = selected_path
#                 if _docx_contains_tables(path_to_use):
#                     flat = flatten_docx_tables_to_longtext(path_to_use)
#                     st.session_state.flattened_docx_path = flat
#                     st.session_state.flatten_used = True
#                     path_to_use = flat
#                 script_text, heading_ranges = build_docx_text_with_meta(path_to_use)
#                 source_docx_path = path_to_use
#             else:
#                 script_text = load_script_file(selected_path)

#         elif uploaded_file:
#             base_stem = "uploaded_script"
#             if uploaded_file.lower().endswith(".docx"):
#                 path_to_use = uploaded_file
#                 if _docx_contains_tables(path_to_use):
#                     flat = flatten_docx_tables_to_longtext(path_to_use)
#                     st.session_state.flattened_docx_path = flat
#                     st.session_state.flatten_used = True
#                     path_to_use = flat
#                 script_text, heading_ranges = build_docx_text_with_meta(path_to_use)
#                 source_docx_path = path_to_use
#             else:
#                 script_text = load_script_file(uploaded_file)

#         elif pasted_text and pasted_text.strip():
#             script_text = pasted_text.strip()
#         else:
#             st.warning("Select, upload, or paste a script first."); st.stop()

#         if len(script_text.strip()) < 50:
#             st.error("Extracted text looks too short. Check your file extraction."); st.stop()

#         with st.spinner("Running analysis…"):
#             try:
#                 review_text = run_review_multi(script_text=script_text, prompts_dir=PROMPTS_DIR, temperature=0.0)
#             finally:
#                 if uploaded_file and not source_docx_path:
#                     try: os.remove(uploaded_file)
#                     except Exception: pass

#         data = extract_review_json(review_text)
#         if not data:
#             st.error("JSON not detected in model output."); st.stop()

#         st.session_state.script_text      = script_text
#         st.session_state.base_stem        = base_stem
#         st.session_state.data             = data
#         st.session_state.heading_ranges   = heading_ranges
#         st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
#         st.session_state.param_choice     = None
#         st.session_state.source_docx_path = source_docx_path
#         st.session_state.review_ready     = True
#         st.session_state.ui_mode          = "review"

#         # Save snapshot for Recents
#         _save_history_snapshot(
#             title=base_stem,
#             data=data,
#             script_text=script_text,
#             source_docx_path=source_docx_path,
#             heading_ranges=heading_ranges,
#             spans_by_param=st.session_state.spans_by_param,
#             aoi_match_ranges=st.session_state.get("aoi_match_ranges", {})
#         )

#         st.rerun()

# # ---------- Results screen ----------
# def render_review():
#     script_text     = st.session_state.script_text
#     data            = st.session_state.data
#     spans_by_param  = st.session_state.spans_by_param
#     scores: Dict[str,int] = (data or {}).get("scores", {}) or {}
#     source_docx_path: Optional[str] = st.session_state.source_docx_path

#     left, center, right = st.columns([1.1, 2.7, 1.4], gap="large")

#     with left:
#         st.subheader("Final score")
#         ordered = [p for p in PARAM_ORDER if p in scores]
#         df = pd.DataFrame({"Parameter": ordered, "Score (1–10)": [scores.get(p, "") for p in ordered]})
#         st.dataframe(df, hide_index=True, use_container_width=True)
#         st.markdown(f'**Overall:** {data.get("overall_rating","—")}/10')
#         st.divider()

#         strengths = (data or {}).get("strengths") or []
#         if not strengths:
#             per = (data or {}).get("per_parameter", {}) or {}
#             best = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
#             for name, sc in best:
#                 if sc >= 8 and name in per:
#                     exp = _sanitize_editor_text((per[name] or {}).get("explanation", "") or "")
#                     first = re.split(r"(?<=[.!?])\s+", exp.strip())[0] if exp else f"Consistently strong {name.lower()}."
#                     strengths.append(f"{name}: {first}")
#                 if len(strengths) >= 3: break

#         def _bullets(title: str, items):
#             st.markdown(f"**{title}**")
#             for s in (items or []):
#                 if isinstance(s, str) and s.strip():
#                     st.write("• " + _sanitize_editor_text(s))
#             if not items:
#                 st.write("• —")

#         _bullets("Strengths", strengths)
#         _bullets("Weaknesses", data.get("weaknesses"))
#         _bullets("Suggestions", data.get("suggestions"))
#         _bullets("Drop-off Risks", data.get("drop_off_risks"))
#         st.markdown("**Viral Quotient**"); st.write(_sanitize_editor_text(data.get("viral_quotient","—")))

#     with right:
#         st.subheader("Parameters")
#         st.markdown('<div class="param-row">', unsafe_allow_html=True)
#         for p in [p for p in PARAM_ORDER if p in scores]:
#             if st.button(p, key=f"chip_{p}", help="Show inline AOI highlights for this parameter"):
#                 st.session_state.param_choice = p
#         st.markdown('</div>', unsafe_allow_html=True)

#         sel = st.session_state.param_choice
#         if sel:
#             blk = (data.get("per_parameter", {}) or {}).get(sel, {}) or {}
#             st.markdown(f"**{sel} — Score:** {scores.get(sel,'—')}/10")

#             if blk.get("explanation"):
#                 st.markdown("**Why this score**")
#                 st.write(_sanitize_editor_text(blk["explanation"]))

#             if blk.get("weakness") and blk["weakness"] != "Not present":
#                 st.markdown("**Weakness**")
#                 st.write(_sanitize_editor_text(blk["weakness"]))

#             if blk.get("suggestion") and blk["suggestion"] != "Not present":
#                 st.markdown("**Suggestion**")
#                 st.write(_sanitize_editor_text(blk["suggestion"]))

#             aoi = blk.get("areas_of_improvement") or []
#             if aoi:
#                 st.markdown("**Areas of Improvement**")
#                 for i, item in enumerate(aoi, 1):
#                     popover_fn = getattr(st, "popover", None)

#                     aid = f"{sel.replace(' ','_')}-AOI-{i}"
#                     s_e_map = st.session_state.get("aoi_match_ranges", {})
#                     if aid in s_e_map:
#                         s_m, e_m = s_e_map[aid]
#                         matched_line = script_text[s_m:e_m]
#                         line = (matched_line if len(matched_line) <= 320 else matched_line[:300] + "…")
#                     else:
#                         line = _sanitize_editor_text(item.get('quote_verbatim',''))

#                     issue = _sanitize_editor_text(item.get('issue',''))
#                     fix = _sanitize_editor_text(item.get('fix',''))
#                     why = _sanitize_editor_text(item.get('why_this_helps',''))

#                     label = f"Issue {i}"
#                     if callable(popover_fn):
#                         with popover_fn(label):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)
#                     else:
#                         with st.expander(label, expanded=False):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)

#             if blk.get("summary"):
#                 st.markdown("**Summary**")
#                 st.write(_sanitize_editor_text(blk["summary"]))

#     with center:
#         st.subheader("Script with inline highlights")
#         spans = st.session_state.spans_by_param.get(st.session_state.param_choice, []) if st.session_state.param_choice else []

#         if st.session_state.source_docx_path and os.path.exists(st.session_state.source_docx_path):
#             html_code = render_docx_html_with_highlights(
#                 st.session_state.source_docx_path,
#                 merge_overlaps_and_adjacent(script_text, spans)
#             )
#             st.markdown(html_code, unsafe_allow_html=True)
#         else:
#             from html import escape as _esc
#             text = _esc(script_text)
#             spans2 = [s for s in merge_overlaps_and_adjacent(script_text, spans) if s[0] < s[1]]
#             spans2.sort(key=lambda x: x[0])
#             cur = 0; buf: List[str] = []
#             for s,e,c,_aid in spans2:
#                 if s > cur: buf.append(text[cur:s])
#                 buf.append(f'<mark style="background:{c}33;border:1px solid {c};border-radius:3px;padding:0 2px;">{text[s:e]}</mark>')
#                 cur = e
#             if cur < len(text): buf.append(text[cur:])
#             st.markdown(
#                 f'<div class="docxwrap"><p style="white-space:pre-wrap; line-height:1.7; font-family:ui-serif, Georgia, Times New Roman, serif;">{"".join(buf)}</p></div>',
#                 unsafe_allow_html=True
#             )

# # ---------- Router ----------
# mode = st.session_state.ui_mode
# if mode == "recents":
#     _render_recents_centerpane()
# elif mode == "review" and st.session_state.review_ready:
#     render_review()
# else:
#     render_home()





# app_grammarly_ui.py — Mysterious 7 (inline AOI popup + autosizing, no inner scroll)
# -------------------------------------------------------
# UI for reviewing scripts with inline AOI highlights and popups.






























































# # //////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////








# import os, re, glob, json, tempfile, difflib, uuid, datetime, shutil
# from pathlib import Path
# from typing import Dict, Any, List, Tuple, Optional

# import streamlit as st
# import pandas as pd
# import streamlit.components.v1 as components  # for inline HTML/JS popup

# from utils1 import extract_review_json, PARAM_ORDER, load_script_file
# from review_engine_multi import run_review_multi

# # ---- DOCX rendering imports ----
# from docx import Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # =========================
# # RunPod S3 (inline helpers) — FIXED
# # =========================
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError

# # Accept both AWS_* and RUNPOD_* style envs / st.secrets
# def _get_env(key: str, default: str = "") -> str:
#     v = os.getenv(key, "")
#     if v: 
#         return v.strip()
#     try:
#         v2 = st.secrets.get(key)
#         if isinstance(v2, str):
#             return v2.strip()
#     except Exception:
#         pass
#     return (default or "").strip()

# # Primary config
# _RP_ENDPOINT = _get_env("RUNPOD_S3_ENDPOINT")
# _RP_BUCKET   = _get_env("RUNPOD_S3_BUCKET")
# _RP_REGION   = _get_env("RUNPOD_S3_REGION") or _get_env("AWS_DEFAULT_REGION") or ""

# # Credentials: prefer AWS_* if present; else accept RUNPOD_* fallbacks
# _AK = _get_env("AWS_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY")
# _SK = _get_env("AWS_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_KEY")
# _ST = _get_env("AWS_SESSION_TOKEN")  # optional

# # Options
# _FORCE_PATH = (_get_env("RUNPOD_S3_FORCE_PATH_STYLE") or "true").lower() in {"1","true","yes"}
# _USE_SSL    = (_get_env("RUNPOD_S3_USE_SSL") or "true").lower() in {"1","true","yes"}
# _VERIFY_SSL = (_get_env("RUNPOD_S3_VERIFY_SSL") or "true").lower() in {"1","true","yes"}

# def _s3_enabled() -> bool:
#     return bool(_RP_ENDPOINT and _RP_BUCKET and _AK and _SK)

# @st.cache_resource(show_spinner=False)
# def _s3_client():
#     if not _s3_enabled():
#         return None
#     session_kwargs = dict(
#         aws_access_key_id=_AK,
#         aws_secret_access_key=_SK,
#     )
#     if _ST:
#         session_kwargs["aws_session_token"] = _ST

#     # s3v4 + path style are common requirements for S3-compatible services
#     cfg = Config(
#         signature_version="s3v4",
#         s3={"addressing_style": "path" if _FORCE_PATH else "auto"},
#         retries={"max_attempts": 3, "mode": "standard"}
#     )
#     return boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         region_name=_RP_REGION or None,
#         use_ssl=_USE_SSL,
#         verify=_VERIFY_SSL,
#         config=cfg,
#         **session_kwargs,
#     )

# def save_text_key(key: str, text: str) -> str:
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "w", encoding="utf-8") as f:
#             f.write(text)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=text.encode("utf-8"))
#     return f"s3://{_RP_BUCKET}/{key}"

# def save_bytes_key(key: str, data: bytes) -> str:
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "wb") as f:
#             f.write(data)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=data)
#     return f"s3://{_RP_BUCKET}/{key}"

# def read_text_key(key: str, default: str = "") -> str:
#     if not _s3_enabled():
#         try:
#             with open(key, "r", encoding="utf-8") as f:
#                 return f.read()
#         except Exception:
#             return default
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read().decode("utf-8", errors="ignore")
#     except Exception:
#         return default

# def read_bytes_key(key: str) -> Optional[bytes]:
#     if not _s3_enabled():
#         try:
#             with open(key, "rb") as f:
#                 return f.read()
#         except Exception:
#             return None
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read()
#     except Exception:
#         return None

# def list_prefix(prefix: str) -> List[str]:
#     """
#     List object keys under prefix (or local dir paths if not S3).
#     In S3 mode we always return KEYS (not URLs).
#     """
#     if not _s3_enabled():
#         base = prefix if os.path.isdir(prefix) else os.path.dirname(prefix)
#         try:
#             return [os.path.join(base, p) for p in os.listdir(base) if p.endswith(".json")]
#         except Exception:
#             return []

#     out: List[str] = []
#     token = None
#     # Normalize to "dir/" prefix for S3 listing
#     s3_prefix = prefix.rstrip("/") + "/"
#     try:
#         while True:
#             kwargs = {"Bucket": _RP_BUCKET, "Prefix": s3_prefix}
#             if token:
#                 kwargs["ContinuationToken"] = token
#             resp = _s3_client().list_objects_v2(**kwargs)
#             for c in resp.get("Contents", []):
#                 k = c.get("Key", "")
#                 if k.endswith(".json"):
#                     out.append(k)
#             token = resp.get("NextContinuationToken")
#             if not token:
#                 break
#     except (ClientError, EndpointConnectionError, NoCredentialsError):
#         return []
#     return out

# def presigned_url(key: str, expires: int = 3600) -> Optional[str]:
#     if not _s3_enabled():
#         return None
#     try:
#         return _s3_client().generate_presigned_url(
#             "get_object",
#             Params={"Bucket": _RP_BUCKET, "Key": key},
#             ExpiresIn=expires
#         )
#     except ClientError:
#         return None

# def ensure_local_copy(key_or_path: str) -> Optional[str]:
#     """
#     For DOCX/PDF parsing we need a real filesystem path.
#     If S3 mode, download to a temp file and return that path.
#     """
#     if not _s3_enabled():
#         return key_or_path if os.path.exists(key_or_path) else None

#     key = key_or_path
#     if key.startswith("s3://"):
#         # s3://bucket/path/to/file -> path/to/file
#         parts = key.split("/", 3)
#         key = parts[3] if len(parts) >= 4 else ""
#     data = read_bytes_key(key)
#     if data is None:
#         return None
#     fd, tmp = tempfile.mkstemp(suffix=os.path.splitext(key)[1] or "")
#     os.close(fd)
#     with open(tmp, "wb") as f:
#         f.write(data)
#     return tmp

# def _s3_health_summary() -> dict:
#     """
#     Optional tiny health read you can print if needed.
#     Returns a dict; safe to ignore in production.
#     """
#     info = {
#         "enabled": _s3_enabled(),
#         "endpoint": _RP_ENDPOINT,
#         "bucket": _RP_BUCKET,
#         "region": _RP_REGION,
#         "has_keys": bool(_AK and _SK),
#     }
#     if not _s3_enabled():
#         info["status"] = "local-mode"
#         return info
#     try:
#         # Attempt a very cheap list; no exceptions => reachable
#         _ = _s3_client().list_objects_v2(Bucket=_RP_BUCKET, Prefix=(f"{OUTPUT_DIR}/_history/").rstrip("/") + "/",
#                                          MaxKeys=1)
#         info["status"] = "ok"
#     except Exception as e:
#         info["status"] = f"error: {getattr(e, 'response', {}).get('Error', {}).get('Code', str(e))}"
#     return info


# # ---------- Folders ----------
# # SCRIPTS_DIR = "scripts"
# # PROMPTS_DIR = "prompts"
# # OUTPUT_DIR  = "outputs"
# # HISTORY_DIR = os.path.join(OUTPUT_DIR, "_history")
# # ---------- Folders (all under Scriptmodel/) ----------

# BASE_PREFIX = "Scriptmodel"

# SCRIPTS_DIR = f"{BASE_PREFIX}/scripts"
# PROMPTS_DIR = f"{BASE_PREFIX}/prompts"
# OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
# HISTORY_DIR = f"{OUTPUT_DIR}/_history"


# if not _s3_enabled():
#     for p in (SCRIPTS_DIR, PROMPTS_DIR, OUTPUT_DIR, HISTORY_DIR):
#         Path(p).mkdir(parents=True, exist_ok=True)


# # ---------- Colors ----------
# PARAM_COLORS: Dict[str, str] = {
#     "Suspense Building":              "#ff6b6b",
#     "Language/Tone":                  "#6b8cff",
#     "Intro + Main Hook/Cliffhanger":  "#ffb86b",
#     "Story Structure + Flow":         "#a78bfa",
#     "Pacing":                         "#f43f5e",
#     "Mini-Hooks (30–60s)":            "#eab308",
#     "Outro (Ending)":                 "#8b5cf6",
# }

# # ---------- Config ----------
# STRICT_MATCH_ONLY = False

# # ---------- App config ----------
# st.set_page_config(page_title="M7 — Grammarly UI", page_icon="🕵️", layout="wide")
# # ---------- Header patch & CSS ----------
# def render_app_title():
#     st.markdown(
#         '<h1 class="app-title">Viral Script Reviewer</h1>',
#         unsafe_allow_html=True
#     )
#     st.markdown("""
#     <style>
#     html { color-scheme: light dark; }

#     :root{
#       --m7-surface: #eef2f7;
#       --m7-on-surface: #0f172a;
#       --m7-border: rgba(15,23,42,.14);
#       --sep: #e5e7eb;
#     }
#     @media (prefers-color-scheme: dark){
#       :root{
#         --m7-surface: #2f333a;
#         --m7-on-surface: #ffffff;
#         --m7-border: rgba(255,255,255,.18);
#         --sep: #2a2f37;
#       }
#     }

#     .stApp .block-container { padding-top: 4.25rem !important; }
#     .app-title{
#       font-weight: 700; font-size: 2.1rem; line-height: 1.3;
#       margin: 0 0 1rem 0; padding-left: 40px !important; padding-top: .25rem !important;
#       white-space: normal; word-break: break-word; hyphens: auto; overflow: visible;
#       position: relative !important; z-index: 10 !important;
#     }
#     [data-testid="collapsedControl"] { z-index: 6 !important; }
#     header[data-testid="stHeader"], .stAppHeader { background: transparent !important; box-shadow: none !important; }
#     @media (min-width: 992px){ .app-title { padding-left: 0 !important; } }

#     div[data-testid="column"]:nth-of-type(1){position:relative;}
#     div[data-testid="column"]:nth-of-type(1)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     div[data-testid="column"]:nth-of-type(2){position:relative;}
#     div[data-testid="column"]:nth-of-type(2)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}

#     .m7-card{
#       background: var(--m7-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 14px 16px;
#       color: var(--m7-on-surface);
#     }
#     .m7-card, .m7-card * { color: var(--m7-on-surface) !important; }

#     .docxwrap{
#       background: var(--m7-surface);
#       color: var(--m7-on-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 16px 14px 18px;
#     }
#     .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#     .docxwrap .h1, .docxwrap .h2, .docxwrap .h3 { font-weight:700; margin:10px 0 6px; }
#     .docxwrap .h1 { font-size: 1.3rem; border-bottom: 2px solid currentColor; padding-bottom: 4px; }
#     .docxwrap .h2 { font-size: 1.15rem; border-bottom: 1px solid currentColor; padding-bottom: 3px; }
#     .docxwrap .h3 { font-size: 1.05rem; }
#     .docxwrap p { margin: 10px 0; line-height: 1.7; font-family: ui-serif, Georgia, "Times New Roman", serif; }
#     .docxwrap table { border-collapse: collapse; width: 100%; margin: 12px 0; }
#     .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); padding:8px; vertical-align:top; line-height:1.6; }
#     .docxwrap mark{ padding:0 2px; border-radius:3px; border:1px solid var(--m7-border); cursor: pointer; }

#     .rec-card{
#       display:block; text-decoration:none !important;
#       background: var(--m7-surface);
#       border:1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 14px 16px; margin: 10px 0 16px;
#       box-shadow: 0 1px 2px rgba(0,0,0,.06);
#       color: var(--m7-on-surface) !important;
#       transition: filter .1s ease, transform .02s ease;
#     }
#     .rec-card:hover{ filter: brightness(1.02); }
#     .rec-card:active{ transform: translateY(1px); }
#     .rec-card, .rec-card * { color: var(--m7-on-surface) !important; }
#     .rec-title{font-weight:600; margin-bottom:.25rem;}
#     .rec-meta{opacity:.85 !important; font-size:12.5px; margin-bottom:.4rem;}
#     .rec-row{display:flex; align-items:center; justify-content:space-between; gap:12px;}

#     .stTextInput>div>div,
#     .stTextArea>div>div,
#     .stNumberInput>div>div,
#     .stDateInput>div>div,
#     .stTimeInput>div>div,
#     .stFileUploader>div,
#     div[data-baseweb="select"]{
#       background: var(--m7-surface) !important;
#       border: 1px solid var(--m7-border) !important;
#       border-radius: 10px !important;
#       color: var(--m7-on-surface) !important;
#     }
#     .stTextInput input,
#     .stTextArea textarea,
#     .stNumberInput input,
#     .stDateInput input,
#     .stTimeInput input,
#     .stFileUploader div,
#     div[data-baseweb="select"] *{
#       color: var(--m7-on-surface) !important;
#     }
#     .stTextInput input::placeholder,
#     .stTextArea textarea::placeholder{ color: rgba(16,24,39,.55) !important; }
#     @media (prefers-color-scheme: dark){
#       .stTextInput input::placeholder,
#       .stTextArea textarea::placeholder{ color: rgba(255,255,255,.75) !important; }
#     }
#     div[data-testid="stFileUploaderDropzone"] label span { color: var(--m7-on-surface) !important; opacity: 1 !important; }
#     div[data-testid="stFileUploaderDropzone"] label { color: var(--m7-on-surface) !important; }

#     .stMarkdown pre,
#     pre[class*="language-"],
#     .stCodeBlock{
#       background: var(--m7-surface) !important;
#       color: var(--m7-on-surface) !important;
#       border: 1px solid var(--m7-border) !important;
#       border-radius: 12px !important;
#       padding: 12px 14px !important;
#       overflow:auto;
#     }
#     .stMarkdown pre code{ background: transparent !important; color: inherit !important; }

#     div[data-testid="stDataFrame"]{
#       background: var(--m7-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 6px 8px;
#       color: var(--m7-on-surface);
#     }
#     div[data-testid="stDataFrame"] * { color: var(--m7-on-surface) !important; }

#     </style>
#     """, unsafe_allow_html=True)

# render_app_title()

# # ---------- Session ----------
# for key, default in [
#     ("review_ready", False),
#     ("script_text", ""),
#     ("base_stem", ""),
#     ("data", None),
#     ("spans_by_param", {}),
#     ("param_choice", None),
#     ("source_docx_path", None),
#     ("heading_ranges", []),
#     ("flattened_docx_path", None),
#     ("flatten_used", False),
#     ("ui_mode", "home"),
# ]:
#     st.session_state.setdefault(key, default)

# # ---------- helpers for query params (compat across Streamlit versions) ----------
# def _get_query_param(key: str) -> Optional[str]:
#     val = None
#     try:
#         val = st.query_params.get(key)
#     except Exception:
#         q = st.experimental_get_query_params()
#         v = q.get(key)
#         if isinstance(v, list): val = v[0] if v else None
#         else: val = v
#     return val

# def _clear_query_params():
#     try:
#         st.query_params.clear()
#     except Exception:
#         st.experimental_set_query_params()

# # ---------- Sanitizer ----------
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )
# def _sanitize_editor_text(s: Optional[str]) -> str:
#     if not s: return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•)\s*', '', t, flags=re.M)
#     t = re.sub(r'^\s*[-*]\s+', '• ', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# # ---------- DOCX traversal ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# # ---------- Auto-flatten ----------
# def _docx_contains_tables(path: str) -> bool:
#     doc = Document(path)
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Table):
#             return True
#     return False

# def _copy_paragraph(dest_doc: Document, src_para: Paragraph):
#     p = dest_doc.add_paragraph()
#     try:
#         if src_para.style and src_para.style.name:
#             p.style = src_para.style.name
#     except Exception:
#         pass
#     for run in src_para.runs:
#         r = p.add_run(run.text or "")
#         r.bold = run.bold
#         r.italic = run.italic
#         r.underline = run.underline
#     return p

# def flatten_docx_tables_to_longtext(source_path: str) -> str:
#     src = Document(source_path)
#     new = Document()
#     for blk in _iter_docx_blocks(src):
#         if isinstance(blk, Paragraph):
#             _copy_paragraph(new, blk)
#         else:
#             seen_tc_ids = set()
#             for row in blk.rows:
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     if tc_id in seen_tc_ids:
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     for p in cell.paragraphs:
#                         _copy_paragraph(new, p)
#                 new.add_paragraph("")
#             new.add_paragraph("")
#     fd, tmp_path = tempfile.mkstemp(suffix=".docx"); os.close(fd); new.save(tmp_path)
#     return tmp_path

# # ---------- Build plain text + heading ranges ----------
# def build_docx_text_with_meta(docx_path: str) -> Tuple[str, List[Tuple[int,int]]]:
#     doc = Document(docx_path)
#     out: List[str] = []
#     heading_ranges: List[Tuple[int,int]] = []
#     current_offset = 0

#     def _append_and_advance(s: str):
#         nonlocal current_offset
#         out.append(s); current_offset += len(s)

#     seen_tc_ids: set = set()
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             para_text = "".join(run.text or "" for run in blk.runs)
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             if sty.startswith("heading"):
#                 start = current_offset; end = start + len(para_text)
#                 heading_ranges.append((start, end))
#             _append_and_advance(para_text); _append_and_advance("\n")
#         else:
#             for row in blk.rows:
#                 row_cell_tcs = []
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     row_cell_tcs.append((tc_id, cell))
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     if tc_id in seen_tc_ids:
#                         if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     cell_text_parts: List[str] = []
#                     for i, p in enumerate(cell.paragraphs):
#                         t = "".join(r.text or "" for r in p.runs)
#                         sty = (p.style.name or "").lower() if p.style else ""
#                         if sty.startswith("heading"):
#                             hs = current_offset + sum(len(x) for x in cell_text_parts)
#                             he = hs + len(t)
#                             heading_ranges.append((hs, he))
#                         cell_text_parts.append(t)
#                         if i != len(cell.paragraphs) - 1:
#                             cell_text_parts.append("\n")
#                     cell_text = "".join(cell_text_parts)
#                     _append_and_advance(cell_text)
#                     if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                 _append_and_advance("\n")
#             _append_and_advance("\n")

#     return "".join(out), heading_ranges

# def _wrap_inline(safe_text: str, run) -> str:
#     out = safe_text
#     if getattr(run, "underline", False): out = f"<u>{out}</u>"
#     if getattr(run, "italic", False): out = f"<em>{out}</em>"
#     if getattr(run, "bold", False): out = f"<strong>{out}</strong>"
#     return out

# # ---------- Invisible/bridge chars ----------
# _BRIDGE_CHARS = set("\u200b\u200c\u200d\u2060\ufeff\xa0\u00ad")

# # ---------- DOCX -> HTML with highlights (includes data-aid) ----------
# def render_docx_html_with_highlights(docx_path: str,
#                                      highlight_spans: List[Tuple[int,int,str,str]]) -> str:
#     """
#     highlight_spans: list of (start, end, color, aid)
#     """
#     doc = Document(docx_path)
#     spans = [s for s in highlight_spans if s[0] < s[1]]
#     spans.sort(key=lambda x: x[0])

#     cur_span = 0
#     current_offset = 0

#     def esc(s: str) -> str:
#         return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

#     def open_mark_if_needed(html_parts, mark_state, color, end, aid):
#         if not mark_state["open"]:
#             html_parts.append(
#                 f'<mark class="aoi-mark" data-aid="{aid}" '
#                 f'style="background:{color}33;border:1px solid {color};border-radius:3px;padding:0 2px;">'
#             )
#             mark_state.update(open=True, end=end, color=color, aid=aid)

#     def close_mark_if_open(html_parts, mark_state):
#         if mark_state["open"]:
#             html_parts.append('</mark>')
#             mark_state.update(open=False, end=None, color=None, aid=None)

#     def emit_run_text(run_text: str, run, html_parts: List[str], mark_state: Dict[str, Any]):
#         nonlocal cur_span, current_offset
#         t = run_text or ""; i = 0
#         while i < len(t):
#             next_start, next_end, color, next_aid = None, None, None, None
#             if cur_span < len(spans):
#                 next_start, next_end, color, next_aid = spans[cur_span]

#             if not mark_state["open"]:
#                 if cur_span >= len(spans) or current_offset + (len(t) - i) <= next_start:
#                     chunk = t[i:]; html_parts.append(_wrap_inline(esc(chunk), run)); current_offset += len(chunk); break
#                 if current_offset < next_start:
#                     take = next_start - current_offset
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take; continue
#                 open_mark_if_needed(html_parts, mark_state, color, next_end, next_aid)
#             else:
#                 take = min(mark_state["end"] - current_offset, len(t) - i)
#                 if take > 0:
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take
#                 if current_offset >= mark_state["end"]:
#                     close_mark_if_open(html_parts, mark_state)
#                     cur_span += 1

#     html: List[str] = ['<div class="docxwrap">']
#     seen_tc_ids: set = set()

#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             open_tag = '<div class="h1">' if sty.startswith("heading 1") else \
#                        '<div class="h2">' if sty.startswith("heading 2") else \
#                        '<div class="h3">' if sty.startswith("heading 3") else "<p>"
#             close_tag = "</div>" if sty.startswith("heading") else "</p>"
#             html.append(open_tag)
#             for run in blk.runs:
#                 emit_run_text(run.text or "", run, html, mark_state)
#             close_mark_if_open(html, mark_state)
#             html.append(close_tag)
#             current_offset += 1  # '\n'

#         else:
#             html.append("<table>")
#             for row in blk.rows:
#                 html.append("<tr>")
#                 row_cell_tcs = [(id(cell._tc), cell) for cell in row.cells]
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     html.append("<td>")
#                     if tc_id not in seen_tc_ids:
#                         seen_tc_ids.add(tc_id)
#                         for p_idx, p in enumerate(cell.paragraphs):
#                             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#                             html.append("<div>")
#                             for run in p.runs:
#                                 emit_run_text(run.text or "", run, html, mark_state)
#                             close_mark_if_open(html, mark_state)
#                             html.append("</div>")
#                             if p_idx != len(cell.paragraphs) - 1:
#                                 current_offset += 1
#                     html.append("</td>")
#                     if idx != len(row_cell_tcs) - 1: current_offset += 1  # '\t'
#                 html.append("</tr>"); current_offset += 1   # row '\n'
#             html.append("</table>"); current_offset += 1     # extra '\n'

#     html.append("</div>")
#     return "".join(html)

# # ---------- Matching / span utilities ----------
# def _normalize_keep_len(s: str) -> str:
#     trans = {
#         "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
#         "\u2013": "-", "\u2014": "-",
#         "\xa0": " ",
#         "\u200b": " ", "\u200c": " ", "\u200d": " ", "\u2060": " ",
#         "\ufeff": " ", "\u00ad": " ",
#     }
#     return (s or "").translate(str.maketrans(trans))

# def _tokenize(s: str) -> List[str]: return re.findall(r"\w+", (s or "").lower())

# def _iter_sentences_with_spans(text: str) -> List[Tuple[int,int,str]]:
#     spans = []
#     for m in re.finditer(r'[^.!?]+[.!?]+|\Z', text, flags=re.S):
#         s, e = m.start(), m.end()
#         seg = text[s:e]
#         if seg.strip(): spans.append((s, e, seg))
#     return spans

# def _squash_ws(s: str) -> str: return re.sub(r"\s+", " ", s or "").strip()

# def _clean_quote_for_match(q: str) -> str:
#     if not q: return ""
#     q = _normalize_keep_len(q).strip()
#     q = re.sub(r'^[\'"“”‘’\[\(\{<…\-\–\—\s]+', '', q)
#     q = re.sub(r'[\'"“”‘’\]\)\}>…\-\–\—\s]+$', '', q)
#     return _squash_ws(q)

# def _snap_and_bridge_to_word(text: str, start: int, end: int, max_bridge: int = 2) -> Tuple[int,int]:
#     n = len(text); s, e = max(0,start), max(start,end)
#     def _is_inv(ch: str) -> bool: return ch in _BRIDGE_CHARS

#     while s > 0:
#         prev = text[s-1]; cur = text[s] if s < n else ""
#         if prev.isalnum() and cur.isalnum(): s -= 1; continue
#         j = s; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1; 
#         if brid and (s-1)>=0 and text[s-1].isalnum() and (j<n and text[j].isalnum()): s -= 1; continue
#         break

#     while e < n:
#         prev = text[e-1] if e>0 else ""; nxt = text[e]
#         if prev.isalnum() and nxt.isalnum(): e += 1; continue
#         j = e; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (e-1)>=0 and text[e-1].isalnum() and (j<n and text[j].isalnum()): e = j + 1; continue
#         break

#     while e < n and text[e] in ',"”’\')]}': e += 1
#     return s, e

# def _heal_split_word_left(text: str, start: int) -> int:
#     i = start
#     if i <= 1 or i >= len(text): return start
#     if text[i-1] != " ": return start
#     j = i - 2
#     while j >= 0 and text[j].isalpha(): j -= 1
#     prev_token = text[j+1:i-1]
#     if len(prev_token) == 1: return i - 2
#     return start

# def _overlaps_any(s: int, e: int, ranges: List[Tuple[int,int]]) -> bool:
#     for rs, re_ in ranges:
#         if e > rs and s < re_: return True
#     return False

# def _fuzzy_window_span(tl: str, nl: str, start: int, w: int) -> Tuple[float, Optional[Tuple[int,int]]]:
#     window = tl[start:start+w]
#     sm = difflib.SequenceMatcher(a=nl, b=window)
#     blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
#     if not blocks: return 0.0, None
#     coverage = sum(b.size for b in blocks) / max(1, len(nl))
#     first_b = min(blocks, key=lambda b: b.b); last_b = max(blocks, key=lambda b: b.b + b.size)
#     s = start + first_b.b; e = start + last_b.b + last_b.size
#     return coverage, (s, e)

# def find_span_smart(text: str, needle: str) -> Optional[Tuple[int,int]]:
#     if not text or not needle: return None
#     t_orig = text; t_norm = _normalize_keep_len(text); n_norm = _clean_quote_for_match(needle)
#     if not n_norm: return None
#     tl = t_norm.lower(); nl = n_norm.lower()

#     i = tl.find(nl)
#     if i != -1:
#         s, e = _snap_and_bridge_to_word(t_orig, i, i + len(nl)); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     m = re.search(re.escape(nl).replace(r"\ ", r"\s+"), tl, flags=re.IGNORECASE)
#     if m:
#         s, e = _snap_and_bridge_to_word(t_orig, m.start(), m.end()); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     if not STRICT_MATCH_ONLY and len(nl) >= 12:
#         w = max(60, min(240, len(nl) + 80))
#         best_cov, best_span = 0.0, None
#         step = max(1, w // 2)
#         for start in range(0, max(1, len(tl) - w + 1), step):
#             cov, se = _fuzzy_window_span(tl, nl, start, w)
#             if cov > best_cov: best_cov, best_span = cov, se
#         if best_span and best_cov >= 0.65:
#             s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
#             if s > 0 and t_orig[s-1:s+1].lower() in {" the", " a", " an"}: s -= 1
#             s = _heal_split_word_left(t_orig, s)
#             return (s, e)

#     if not STRICT_MATCH_ONLY:
#         keys = [w for w in _tokenize(nl) if len(w) >= 4][:8]
#         if len(keys) >= 2:
#             kset = set(keys)
#             best_score, best_span = 0.0, None
#             for s, e, seg in _iter_sentences_with_spans(t_norm):
#                 toks = set(_tokenize(seg)); ov = len(kset & toks)
#                 if ov == 0: continue
#                 score = ov / max(2, len(kset)); length_pen = min(1.0, 120 / max(20, e - s)); score *= (0.6 + 0.4 * length_pen)
#                 if score > best_score: best_score, best_span = score, (s, min(e, s + 400))
#             if best_span and best_score >= 0.35:
#                 s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1]); s = _heal_split_word_left(t_orig, s)
#                 return (s, e)
#     return None

# def merge_overlaps(spans: List[Tuple[int,int,str,str]]) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans.sort(key=lambda x: x[0]); out = [spans[0]]
#     for s,e,c,aid in spans[1:]:
#         ps,pe,pc,paid = out[-1]
#         if s <= pe and pc == c and e > pe: out[-1] = (ps, e, pc, paid)
#         else: out.append((s,e,c,aid))
#     return out

# _PUNCT_WS = set(" \t\r\n,.;:!?)('\"[]{}-—–…") | _BRIDGE_CHARS

# def merge_overlaps_and_adjacent(base_text: str,
#                                 spans: List[Tuple[int,int,str,str]],
#                                 max_gap: int = 2) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans = sorted(spans, key=lambda x: x[0]); out = [spans[0]]
#     for s, e, c, aid in spans[1:]:
#         ps, pe, pc, paid = out[-1]
#         if c == pc and s <= pe: out[-1] = (ps, max(pe, e), pc, paid); continue
#         if c == pc and s - pe <= max_gap:
#             gap = base_text[max(0, pe):max(0, s)]
#             if all((ch in _PUNCT_WS) for ch in gap): out[-1] = (ps, e, pc, paid); continue
#         out.append((s, e, c, aid))
#     return out

# # ---------- Heading filters ----------
# def _is_heading_like(q: str) -> bool:
#     if not q: return True
#     s = q.strip()
#     if not re.search(r'[.!?]', s):
#         words = re.findall(r"[A-Za-z]+", s)
#         if 1 <= len(words) <= 7:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.8: return True
#         if s.lower() in {"introduction","voiceover","outro","epilogue","prologue","credits","title","hook","horrifying discovery","final decision"}: return True
#         if len(s) <= 3: return True
#     return False

# def _is_heading_context(script_text: str, s: int, e: int) -> bool:
#     left = script_text.rfind("\n", 0, s) + 1
#     right = script_text.find("\n", e); right = len(script_text) if right == -1 else right
#     line = script_text[left:right].strip()
#     if len(line) <= 70 and not re.search(r'[.!?]', line):
#         words = re.findall(r"[A-Za-z]+", line)
#         if 1 <= len(words) <= 8:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.7: return True
#     return False

# def _tighten_to_quote(script_text: str, span: Tuple[int,int], quote: str) -> Tuple[int,int]:
#     if not span or not quote: return span
#     s, e = span
#     if e <= s or s < 0 or e > len(script_text): return span
#     window = script_text[s:e]; win_norm = _normalize_keep_len(window).lower(); q_norm = _clean_quote_for_match(quote).lower()
#     if not q_norm: return span
#     i = win_norm.find(q_norm)
#     if i == -1:
#         m = re.search(re.escape(q_norm).replace(r"\ ", r"\s+"), win_norm, flags=re.IGNORECASE)
#         if not m: return span
#         i, j = m.start(), m.end()
#     else:
#         j = i + len(q_norm)
#     s2, e2 = s + i, s + j
#     s2, e2 = _snap_and_bridge_to_word(script_text, s2, e2); s2 = _heal_split_word_left(script_text, s2)
#     if s2 >= s and e2 <= e and e2 > s2: return (s2, e2)
#     return span

# def build_spans_by_param(script_text: str, data: dict, heading_ranges: Optional[List[Tuple[int,int]]] = None) -> Dict[str, List[Tuple[int,int,str,str]]]:
#     heading_ranges = heading_ranges or []
#     raw = (data or {}).get("per_parameter", {}) or {}
#     per: Dict[str, Dict[str, Any]] = {k:(v or {}) for k,v in raw.items()}
#     spans_map: Dict[str, List[Tuple[int,int,str,str]]] = {p: [] for p in PARAM_ORDER}
#     st.session_state["aoi_match_ranges"] = {}

#     for p in spans_map.keys():
#         color = PARAM_COLORS.get(p, "#ffd54f")
#         blk = per.get(p, {}) or {}
#         aois = blk.get("areas_of_improvement") or []
#         for idx, item in enumerate(aois, start=1):
#             raw_q = (item or {}).get("quote_verbatim", "") or ""
#             q = _sanitize_editor_text(raw_q)
#             clean = _clean_quote_for_match(re.sub(r"^[•\-\d\.\)\s]+", "", q).strip())
#             if not clean: continue
#             if _is_heading_like(clean): continue

#             pos = find_span_smart(script_text, clean)
#             if not pos: continue
#             pos = _tighten_to_quote(script_text, pos, raw_q)
#             s, e = pos

#             if heading_ranges and _overlaps_any(s, e, heading_ranges): continue
#             if _is_heading_context(script_text, s, e): continue

#             aid = f"{p.replace(' ','_')}-AOI-{idx}"
#             spans_map[p].append((s, e, color, aid))
#             st.session_state["aoi_match_ranges"][aid] = (s, e)
#     return spans_map

# # ---------- History (S3-aware) ----------
# def _maybe_copy_docx_to_history(source_docx_path: Optional[str], run_id: str) -> Optional[str]:
#     """Copy the DOCX used for rendering into outputs/_history so Recents can re-render identically."""
#     try:
#         if not source_docx_path:
#             return None
#         # If already an S3 key/url, just store the key in history blob
#         if source_docx_path.startswith("s3://") or (_s3_enabled() and not os.path.exists(source_docx_path)):
#             # ensure it's uploaded under our history key if it's a local temp in S3 mode
#             if os.path.exists(source_docx_path):
#                 with open(source_docx_path, "rb") as f:
#                     save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#                 return f"{HISTORY_DIR}/{run_id}.docx"
#             return source_docx_path

#         # Local file path
#         if _s3_enabled():
#             with open(source_docx_path, "rb") as f:
#                 save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#             return f"{HISTORY_DIR}/{run_id}.docx"
#         else:
#             dst = os.path.join(HISTORY_DIR, f"{run_id}.docx")
#             if os.path.abspath(source_docx_path) != os.path.abspath(dst):
#                 shutil.copyfile(source_docx_path, dst)
#             return dst
#     except Exception:
#         return None

# def _save_history_snapshot(title: str, data: dict, script_text: str,
#                            source_docx_path: Optional[str], heading_ranges: List[Tuple[int,int]],
#                            spans_by_param: Dict[str, List[Tuple[int,int,str,str]]],
#                            aoi_match_ranges: Dict[str, Tuple[int,int]]):
#     run_id = str(uuid.uuid4()); now = datetime.datetime.now()
#     created_at_iso = now.replace(microsecond=0).isoformat()
#     created_at_human = now.strftime("%Y-%m-%d %H:%M:%S")

#     # Copy stable DOCX alongside history JSON (so Recents uses identical renderer)
#     stable_docx_key_or_path = _maybe_copy_docx_to_history(source_docx_path, run_id)

#     blob = {
#         "run_id": run_id, "title": title or "untitled",
#         "created_at": created_at_iso, "created_at_human": created_at_human,
#         "overall_rating": (data or {}).get("overall_rating", ""),
#         "scores": (data or {}).get("scores", {}),
#         "data": data or {}, "script_text": script_text or "",
#         "source_docx_path": stable_docx_key_or_path or source_docx_path,
#         "heading_ranges": heading_ranges or [],
#         "spans_by_param": spans_by_param or {},
#         "aoi_match_ranges": aoi_match_ranges or {},
#     }

#     # Persist JSON (S3 or local)
#     out_name = f"{created_at_iso.replace(':','-')}__{run_id}.json"
#     out_key = f"{HISTORY_DIR}/{out_name}"
#     save_text_key(out_key, json.dumps(blob, ensure_ascii=False, indent=2))

# def _load_all_history() -> List[dict]:
#     out: List[dict] = []
#     if _s3_enabled():
#         keys = sorted(list_prefix(HISTORY_DIR), reverse=True)
#         for key in keys:
#             try:
#                 txt = read_text_key(key, "")
#                 if not txt:
#                     continue
#                 j = json.loads(txt)
#                 j["_key"] = key
#                 if not j.get("created_at_human") and j.get("created_at"):
#                     try:
#                         dt = datetime.datetime.fromisoformat(j["created_at"])
#                         j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#                     except Exception:
#                         pass
#                 out.append(j)
#             except Exception:
#                 continue
#     else:
#         for fp in sorted(glob.glob(os.path.join(HISTORY_DIR, "*.json"))):
#             try:
#                 with open(fp, "r", encoding="utf-8") as f:
#                     j = json.load(f)
#             except Exception:
#                 continue
#             j.setdefault("_path", fp)
#             ca = j.get("created_at")
#             try:
#                 if isinstance(ca, (int, float)):
#                     dt = datetime.datetime.utcfromtimestamp(float(ca))
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.astimezone().strftime("%Y-%m-%d %H:%M:%S")
#                 elif isinstance(ca, str) and ca: pass
#                 else:
#                     mtime = os.path.getmtime(fp); dt = datetime.datetime.fromtimestamp(mtime)
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#             except Exception:
#                 j["created_at"] = str(ca or "")
#             out.append(j)
#     out.sort(key=lambda r: r.get("created_at") or "", reverse=True)
#     return out

# def _open_history_run_by_id(run_id: str) -> bool:
#     """Open a history run by its run_id. Returns True if loaded."""
#     if not run_id: return False
#     recs = _load_all_history()
#     match = next((r for r in recs if r.get("run_id") == run_id), None)
#     if not match: return False

#     # Load JSON content again (S3/local), then set session
#     try:
#         if "_key" in match and _s3_enabled():
#             txt = read_text_key(match["_key"], "")
#             if not txt: return False
#             jj = json.loads(txt)
#         else:
#             path = match.get("_path")
#             if not path or not os.path.exists(path): return False
#             with open(path, "r", encoding="utf-8") as f:
#                 jj = json.load(f)
#     except Exception:
#         return False

#     st.session_state.script_text      = jj.get("script_text","")
#     st.session_state.base_stem        = jj.get("title","untitled")
#     st.session_state.data             = jj.get("data",{})
#     st.session_state.heading_ranges   = jj.get("heading_ranges",[])
#     st.session_state.spans_by_param   = jj.get("spans_by_param",{})
#     st.session_state.param_choice     = None
#     st.session_state.source_docx_path = jj.get("source_docx_path")
#     st.session_state.review_ready     = True
#     st.session_state["aoi_match_ranges"] = jj.get("aoi_match_ranges", {})
#     st.session_state.ui_mode          = "review"
#     return True

# def _render_recents_centerpane():
#     st.subheader("📄 Recents")
#     q = st.text_input("Filter by title…", "")
#     cols = st.columns([1, 4])
#     with cols[0]:
#         if st.button("← Back"):
#             st.session_state.ui_mode = "home"; st.rerun()

#     recs = _load_all_history()
#     ql = q.strip().lower()
#     if ql: recs = [r for r in recs if ql in (r.get("title","").lower())]
#     if not recs: st.caption("No history yet."); st.stop()

#     # Each card is a clickable <a class="rec-card" href="?open=<run_id>">
#     for rec in recs:
#         run_id = rec.get("run_id"); title = rec.get("title") or "(untitled)"
#         created_h = rec.get("created_at_human",""); overall = rec.get("overall_rating","")
#         st.markdown(f"""
#         <a class="rec-card" href="?open={run_id}">
#           <div class="rec-title">{title}</div>
#           <div class="rec-meta">{created_h}</div>
#           <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
#         </a>
#         """, unsafe_allow_html=True)

# # ---------- Sidebar ----------
# with st.sidebar:
#     if st.button("🆕 New review", use_container_width=True):
#         fp = st.session_state.get("flattened_docx_path")
#         if fp and os.path.exists(fp):
#             try: os.remove(fp)
#             except Exception: pass
#         for k in ["review_ready","script_text","base_stem","data","spans_by_param","param_choice",
#                   "source_docx_path","heading_ranges","flattened_docx_path","flatten_used"]:
#             st.session_state[k] = (
#                 False if k=="review_ready"
#                 else "" if k in ("script_text","base_stem")
#                 else {} if k=="spans_by_param"
#                 else [] if k=="heading_ranges"
#                 else None if k in ("source_docx_path","flattened_docx_path")
#                 else False if k=="flatten_used"
#                 else None
#             )
#         st.session_state.ui_mode = "home"
#         _clear_query_params()
#         st.rerun()

#     if st.button("📁 Recents", use_container_width=True):
#         st.session_state.ui_mode = "recents"
#         _clear_query_params()
#         st.rerun()

# # ---------- Input screen ----------
# def render_home():
#     st.subheader("🎬 Script Source")

#     (tab_upload,) = st.tabs(["Upload file"])
#     uploaded_file = None
#     uploaded_name = None
#     uploaded_key  = None

#     with tab_upload:
#         up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
#         if up is not None:
#             file_bytes = up.read()
#             suffix = os.path.splitext(up.name)[1].lower()
#             # Save to S3 (or local) under scripts/
#             uploaded_key = f"{SCRIPTS_DIR}/{up.name}"
#             save_bytes_key(uploaded_key, file_bytes)

#             # Also create a temp local copy for parsing
#             with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
#                 tmp.write(file_bytes)
#                 uploaded_file = tmp.name
#             uploaded_name = os.path.splitext(os.path.basename(up.name))[0] or "uploaded_script"

#     if st.button("🚀 Run Review", type="primary", use_container_width=True):
#         base_stem = "uploaded_script"
#         source_docx_path = None
#         heading_ranges = []

#         if uploaded_file:
#             base_stem = uploaded_name or "uploaded_script"
#             # Prefer key (so we can store it into history); for DOCX rendering we keep a tmp local
#             if uploaded_file.lower().endswith(".docx"):
#                 path_to_use = uploaded_file
#                 if _docx_contains_tables(path_to_use):
#                     flat = flatten_docx_tables_to_longtext(path_to_use)
#                     st.session_state.flattened_docx_path = flat
#                     st.session_state.flatten_used = True
#                     path_to_use = flat
#                 script_text, heading_ranges = build_docx_text_with_meta(path_to_use)

#                 # 🔑 IMPORTANT: set the rendering source to the *flattened* docx
#                 # (History saver will copy this file to Scriptmodel/outputs/_history/)
#                 source_docx_path = path_to_use
#             else:
#                 script_text = load_script_file(uploaded_file)
#                 source_docx_path = uploaded_key or uploaded_file  # keep reference to S3/local
#         else:
#             st.warning("Please upload a script first.")
#             st.stop()

#         if len(script_text.strip()) < 50:
#             st.error("Extracted text looks too short. Check your file extraction.")
#             st.stop()

#         with st.spinner("Running analysis…"):
#             try:
#                 review_text = run_review_multi(
#                     script_text=script_text,
#                     prompts_dir=PROMPTS_DIR,
#                     temperature=0.0
#                 )
#             finally:
#                 # Clean temp upload if we didn't keep it as source_docx_path local
#                 if uploaded_file and not (isinstance(source_docx_path, str) and os.path.exists(source_docx_path)):
#                     try:
#                         os.remove(uploaded_file)
#                     except Exception:
#                         pass

#         data = extract_review_json(review_text)
#         if not data:
#             st.error("JSON not detected in model output.")
#             st.stop()

#         st.session_state.script_text      = script_text
#         st.session_state.base_stem        = base_stem
#         st.session_state.data             = data
#         st.session_state.heading_ranges   = heading_ranges
#         st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
#         st.session_state.param_choice     = None
#         st.session_state.source_docx_path = source_docx_path
#         st.session_state.review_ready     = True
#         st.session_state.ui_mode          = "review"

#         _save_history_snapshot(
#             title=base_stem,
#             data=data,
#             script_text=script_text,
#             source_docx_path=source_docx_path,
#             heading_ranges=heading_ranges,
#             spans_by_param=st.session_state.spans_by_param,
#             aoi_match_ranges=st.session_state.get("aoi_match_ranges", {})
#         )

#         _clear_query_params()
#         st.rerun()

# # ---------- Results screen ----------
# def render_review():
#     script_text     = st.session_state.script_text
#     data            = st.session_state.data
#     spans_by_param  = st.session_state.spans_by_param
#     scores: Dict[str,int] = (data or {}).get("scores", {}) or {}
#     source_docx_path: Optional[str] = st.session_state.source_docx_path

#     # If our source_docx_path is an S3 key/url, ensure we have a local copy for rendering
#     # AFTER (prefer the flattened copy if we have it in-session)
#     docx_local: Optional[str] = None
#     preferred = st.session_state.get("flattened_docx_path") if st.session_state.get("flatten_used") else source_docx_path
#     if preferred:
#         if not os.path.exists(preferred):
#             docx_local = ensure_local_copy(preferred)
#         else:
#             docx_local = preferred

#     left, center, right = st.columns([1.1, 2.7, 1.4], gap="large")

#     with left:
#         st.subheader("Final score")
#         ordered = [p for p in PARAM_ORDER if p in scores]
#         df = pd.DataFrame({"Parameter": ordered, "Score (1–10)": [scores.get(p, "") for p in ordered]})
#         st.dataframe(df, hide_index=True, use_container_width=True)
#         st.markdown(f'**Overall:** {data.get("overall_rating","—")}/10')
#         st.divider()

#         strengths = (data or {}).get("strengths") or []
#         if not strengths:
#             per = (data or {}).get("per_parameter", {}) or {}
#             best = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
#             for name, sc in best:
#                 if sc >= 8 and name in per:
#                     exp = _sanitize_editor_text((per[name] or {}).get("explanation", "") or "")
#                     first = re.split(r"(?<=[.!?])\s+", exp.strip())[0] if exp else f"Consistently strong {name.lower()}."
#                     strengths.append(f"{name}: {first}")
#                 if len(strengths) >= 3: break

#         def _bullets(title: str, items):
#             st.markdown(f"**{title}**")
#             for s in (items or []):
#                 if isinstance(s, str) and s.strip(): st.write("• " + _sanitize_editor_text(s))
#             if not items: st.write("• —")

#         _bullets("Strengths", strengths)
#         _bullets("Weaknesses", data.get("weaknesses"))
#         _bullets("Suggestions", data.get("suggestions"))
#         _bullets("Drop-off Risks", data.get("drop_off_risks"))
#         st.markdown("**Viral Quotient**"); st.write(_sanitize_editor_text(data.get("viral_quotient","—")))

#     with right:
#         st.subheader("Parameters")
#         st.markdown('<div class="param-row">', unsafe_allow_html=True)
#         for p in [p for p in PARAM_ORDER if p in scores]:
#             if st.button(p, key=f"chip_{p}", help="Show inline AOI highlights for this parameter"):
#                 st.session_state.param_choice = p
#         st.markdown('</div>', unsafe_allow_html=True)

#         sel = st.session_state.param_choice
#         if sel:
#             blk = (data.get("per_parameter", {}) or {}).get(sel, {}) or {}
#             st.markdown(f"**{sel} — Score:** {scores.get(sel,'—')}/10")

#             if blk.get("explanation"):
#                 st.markdown("**Why this score**"); st.write(_sanitize_editor_text(blk["explanation"]))
#             if blk.get("weakness") and blk["weakness"] != "Not present":
#                 st.markdown("**Weakness**"); st.write(_sanitize_editor_text(blk["weakness"]))
#             if blk.get("suggestion") and blk["suggestion"] != "Not present":
#                 st.markdown("**Suggestion**"); st.write(_sanitize_editor_text(blk["suggestion"]))

#             aoi = blk.get("areas_of_improvement") or []
#             if aoi:
#                 st.markdown("**Areas of Improvement**")
#                 for i, item in enumerate(aoi, 1):
#                     popover_fn = getattr(st, "popover", None)
#                     aid = f"{sel.replace(' ','_')}-AOI-{i}"
#                     s_e_map = st.session_state.get("aoi_match_ranges", {})
#                     if aid in s_e_map:
#                         s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                         line = (matched_line if len(matched_line) <= 320 else matched_line[:300] + "…")
#                     else:
#                         line = _sanitize_editor_text(item.get('quote_verbatim',''))
#                     issue = _sanitize_editor_text(item.get('issue',''))
#                     fix   = _sanitize_editor_text(item.get('fix',''))
#                     why   = _sanitize_editor_text(item.get('why_this_helps',''))
#                     label = f"Issue {i}"
#                     if callable(popover_fn):
#                         with popover_fn(label):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)
#                     else:
#                         with st.expander(label, expanded=False):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)

#             if blk.get("summary"):
#                 st.markdown("**Summary**"); st.write(_sanitize_editor_text(blk["summary"]))

#     with center:
#         st.subheader("Script with inline highlights")
#         # choose spans for selected parameter (or all if None)
#         spans = st.session_state.spans_by_param.get(st.session_state.param_choice, []) if st.session_state.param_choice else []

#         # Build AOI payload for popup: { aid: {line, issue, fix, why} }
#         aoi_payload: Dict[str, Dict[str, str]] = {}
#         data_per = (data or {}).get("per_parameter") or {}
#         s_e_map = st.session_state.get("aoi_match_ranges", {})
#         sel = st.session_state.param_choice

#         def _mk_line(aid: str, fallback_q: str = "") -> str:
#             if aid in s_e_map:
#                 s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                 return matched_line if len(matched_line) <= 320 else matched_line[:300] + "…"
#             return _sanitize_editor_text(fallback_q or "")

#         def _collect(param_name: str):
#             blk = (data_per.get(param_name) or {})
#             for i, item in enumerate(blk.get("areas_of_improvement") or [], 1):
#                 aid = f"{param_name.replace(' ','_')}-AOI-{i}"
#                 aoi_payload[aid] = {
#                     "line": _mk_line(aid, (item or {}).get("quote_verbatim","")),
#                     "issue": _sanitize_editor_text((item or {}).get("issue","")),
#                     "fix": _sanitize_editor_text((item or {}).get("fix","")),
#                     "why": _sanitize_editor_text((item or {}).get("why_this_helps","")),
#                 }

#         if sel: _collect(sel)
#         else:
#             for pn in [p for p in PARAM_ORDER if p in data_per]:
#                 _collect(pn)

#         payload_json = json.dumps(aoi_payload, ensure_ascii=False)

#         frame_theme_css = """
#         <style>
#           :root{
#             --m7-surface: #eef2f7;
#             --m7-on-surface: #0f172a;
#             --m7-border: rgba(15,23,42,.14);
#           }
#           @media (prefers-color-scheme: dark){
#             :root{
#               --m7-surface: #2f333a;
#               --m7-on-surface: #ffffff;
#               --m7-border: rgba(255,255,255,.18);
#             }
#             body { background: transparent !important; }
#           }

#           .docxwrap{
#             background: var(--m7-surface);
#             color: var(--m7-on-surface);
#             border: 1px solid var(--m7-border);
#             border-radius: 12px;
#             padding: 16px 14px 18px;
#           }
#           .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#           .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); }
#         </style>
#         """

#         tooltip_css = """
#         <style>
#         .aoi-pop {
#           position: absolute; max-width: 520px; min-width: 320px;
#           background: var(--m7-surface); border: 1px solid var(--m7-border); border-radius: 10px;
#           box-shadow: 0 10px 25px rgba(0,0,0,.12);
#           padding: 12px 14px; z-index: 9999; transform: translateY(-8px);
#           color: var(--m7-on-surface);
#         }
#         .aoi-pop h4 { margin: 0 0 .35rem 0; font-size: .95rem; }
#         .aoi-pop p  { margin: .15rem 0; line-height: 1.5; }
#         .aoi-pop .muted { opacity:.85; font-size:.85rem; }
#         .aoi-arrow {
#           position:absolute; left:50%; transform:translateX(-50%);
#           bottom:-7px; width:0;height:0;border-left:7px solid transparent;
#           border-right:7px solid transparent;border-top:7px solid var(--m7-border);
#         }
#         .aoi-arrow::after{
#           content:""; position:absolute; left:-6px; top:-7px; width:0;height:0;
#           border-left:6px solid transparent;border-right:6px solid transparent;border-top:6px solid var(--m7-surface);
#         }
#         </style>
#         """

#         # Select rendering source (DOCX with highlights if we have a local path)
#         if docx_local and os.path.splitext(docx_local)[1].lower() == ".docx":
#             html_core = render_docx_html_with_highlights(
#                 docx_local,
#                 merge_overlaps_and_adjacent(script_text, spans)
#             )
#         else:
#             from html import escape as _esc
#             orig = script_text
#             spans2 = [s for s in merge_overlaps_and_adjacent(orig, spans) if s[0] < s[1]]
#             spans2.sort(key=lambda x: x[0])
#             cur = 0; buf: List[str] = []
#             for s,e,c,aid in spans2:
#                 if s > cur: buf.append(_esc(orig[cur:s]))
#                 buf.append(
#                     f'<mark class="aoi-mark" data-aid="{aid}" '
#                     f'style="background:{c}33;border:1px solid {c};border-radius:3px;padding:0 2px;">'
#                     f'{_esc(orig[s:e])}</mark>'
#                 )
#                 cur = e
#             if cur < len(orig): buf.append(_esc(orig[cur:]))
#             html_core = (
#                 '<div class="docxwrap"><p style="white-space:pre-wrap; '
#                 'line-height:1.7; font-family:ui-serif, Georgia, Times New Roman, serif;">'
#                 + "".join(buf) +
#                 '</p></div>'
#             )

#         # Popup + autosize JS shell
#         html_shell = """
# %%FRAME_THEME_CSS%%
# %%TOOLTIP_CSS%%
# <div id="m7-doc">%%HTML_CORE%%</div>
# <div id="aoi-pop" class="aoi-pop" style="display:none;">
#   <div id="aoi-pop-content"></div>
#   <div class="aoi-arrow"></div>
# </div>
# <script>
# (function(){
#   const AOI = __PAYLOAD__;
#   const wrap = document.getElementById('m7-doc');
#   const pop  = document.getElementById('aoi-pop');
#   const body = document.getElementById('aoi-pop-content');

#   function resizeIframe() {
#     try {
#       const h = Math.max(
#         document.documentElement.scrollHeight,
#         document.body.scrollHeight
#       );
#       if (window.frameElement) {
#         window.frameElement.style.height = (h + 20) + 'px';
#         window.frameElement.style.width  = '100%';
#       }
#     } catch(e) {}
#   }
#   window.addEventListener('load', resizeIframe);
#   window.addEventListener('resize', resizeIframe);

#   function hide(){ pop.style.display='none'; }
#   function showFor(mark){
#     const aid = mark.getAttribute('data-aid');
#     const d = AOI[aid]; if(!d) return;
#     body.innerHTML =
#       (d.line  ? '<p><strong>Line:</strong> '  + d.line  + '</p>' : '') +
#       (d.issue ? '<p><strong>Issue:</strong> ' + d.issue + '</p>' : '') +
#       (d.fix   ? '<p><strong>Fix:</strong> '   + d.fix   + '</p>' : '') +
#       (d.why   ? '<p class="muted">'           + d.why   + '</p>' : '');
#     pop.style.display = 'block';

#     const r = mark.getBoundingClientRect();
#     const scY = window.scrollY || document.documentElement.scrollTop;
#     const scX = window.scrollX || document.documentElement.scrollLeft;
#     let top  = r.top + scY - pop.offsetHeight - 10;
#     let left = r.left + scX + r.width/2 - pop.offsetWidth/2;
#     if (top < 8) top = r.bottom + scY + 10;
#     if (left < 8) left = 8;
#     pop.style.top  = top + 'px';
#     pop.style.left = left + 'px';

#     resizeIframe();
#   }

#   wrap.addEventListener('click', (e) => {
#     const m = e.target.closest('.aoi-mark');
#     if(!m){ hide(); return; }
#     if(pop.style.display === 'block'){ hide(); }
#     showFor(m);
#     e.stopPropagation();
#   });

#   document.addEventListener('click', (e) => {
#     if(!e.target.closest('.aoi-pop') && !e.target.closest('.aoi-mark')) hide();
#   });
# })();
# </script>
# """
#         html_shell = (
#             html_shell
#             .replace("%%FRAME_THEME_CSS%%", frame_theme_css)
#             .replace("%%TOOLTIP_CSS%%", tooltip_css)
#             .replace("%%HTML_CORE%%", html_core)
#             .replace("__PAYLOAD__", payload_json)
#         )

#         components.html(html_shell, height=400, scrolling=False)

# # ---------- Router & query param open ----------
# _open_qp = _get_query_param("open")
# if _open_qp and _open_history_run_by_id(_open_qp):
#     _clear_query_params()  # avoid re-opening on subsequent reruns

# mode = st.session_state.ui_mode
# if mode == "recents":
#     _render_recents_centerpane()
# elif mode == "review" and st.session_state.review_ready:
#     render_review()
# else:
#     render_home()



# # //////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////



# import os, re, glob, json, tempfile, difflib, uuid, datetime, shutil
# from pathlib import Path
# from typing import Dict, Any, List, Tuple, Optional

# import streamlit as st
# import pandas as pd
# import streamlit.components.v1 as components  # for inline HTML/JS popup

# from utils1 import extract_review_json, PARAM_ORDER, load_script_file
# from review_engine_multi import run_review_multi

# # ---- DOCX rendering imports ----
# from docx import Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # =========================
# # RunPod S3 (inline helpers) — FIXED
# # =========================
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError

# # Accept both AWS_* and RUNPOD_* style envs / st.secrets
# def _get_env(key: str, default: str = "") -> str:
#     v = os.getenv(key, "")
#     if v:
#         return v.strip()
#     try:
#         v2 = st.secrets.get(key)
#         if isinstance(v2, str):
#             return v2.strip()
#     except Exception:
#         pass
#     return (default or "").strip()

# # Primary config
# _RP_ENDPOINT = _get_env("RUNPOD_S3_ENDPOINT")
# _RP_BUCKET   = _get_env("RUNPOD_S3_BUCKET")
# _RP_REGION   = _get_env("RUNPOD_S3_REGION") or _get_env("AWS_DEFAULT_REGION") or ""

# # Credentials: prefer AWS_* if present; else accept RUNPOD_* fallbacks
# _AK = _get_env("AWS_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY")
# _SK = _get_env("AWS_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_KEY")
# _ST = _get_env("AWS_SESSION_TOKEN")  # optional

# # Options
# _FORCE_PATH = (_get_env("RUNPOD_S3_FORCE_PATH_STYLE") or "true").lower() in {"1","true","yes"}
# _USE_SSL    = (_get_env("RUNPOD_S3_USE_SSL") or "true").lower() in {"1","true","yes"}
# _VERIFY_SSL = (_get_env("RUNPOD_S3_VERIFY_SSL") or "true").lower() in {"1","true","yes"}

# def _s3_enabled() -> bool:
#     return bool(_RP_ENDPOINT and _RP_BUCKET and _AK and _SK)

# @st.cache_resource(show_spinner=False)
# def _s3_client():
#     if not _s3_enabled():
#         return None
#     session_kwargs = dict(
#         aws_access_key_id=_AK,
#         aws_secret_access_key=_SK,
#     )
#     if _ST:
#         session_kwargs["aws_session_token"] = _ST

#     cfg = Config(
#         signature_version="s3v4",
#         s3={"addressing_style": "path" if _FORCE_PATH else "auto"},
#         retries={"max_attempts": 3, "mode": "standard"}
#     )
#     return boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         region_name=_RP_REGION or None,
#         use_ssl=_USE_SSL,
#         verify=_VERIFY_SSL,
#         config=cfg,
#         **session_kwargs,
#     )

# def save_text_key(key: str, text: str) -> str:
#     if not _s3_enabled():
#         d = os.path.dirname(key)
#         if d:
#             os.makedirs(d, exist_ok=True)
#         with open(key, "w", encoding="utf-8") as f:
#             f.write(text)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=text.encode("utf-8"))
#     return f"s3://{_RP_BUCKET}/{key}"

# def save_bytes_key(key: str, data: bytes) -> str:
#     if not _s3_enabled():
#         d = os.path.dirname(key)
#         if d:
#             os.makedirs(d, exist_ok=True)
#         with open(key, "wb") as f:
#             f.write(data)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=data)
#     return f"s3://{_RP_BUCKET}/{key}"

# def read_text_key(key: str, default: str = "") -> str:
#     if not _s3_enabled():
#         try:
#             with open(key, "r", encoding="utf-8") as f:
#                 return f.read()
#         except Exception:
#             return default
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read().decode("utf-8", errors="ignore")
#     except Exception:
#         return default

# def read_bytes_key(key: str) -> Optional[bytes]:
#     if not _s3_enabled():
#         try:
#             with open(key, "rb") as f:
#                 return f.read()
#         except Exception:
#             return None
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read()
#     except Exception:
#         return None

# def list_prefix(prefix: str) -> List[str]:
#     """
#     List object keys under prefix (or local dir paths if not S3).
#     In S3 mode we always return KEYS (not URLs).
#     """
#     if not _s3_enabled():
#         base = prefix if os.path.isdir(prefix) else os.path.dirname(prefix)
#         try:
#             return [os.path.join(base, p) for p in os.listdir(base) if p.endswith(".json")]
#         except Exception:
#             return []

#     out: List[str] = []
#     token = None
#     s3_prefix = prefix.rstrip("/") + "/"
#     try:
#         while True:
#             kwargs = {"Bucket": _RP_BUCKET, "Prefix": s3_prefix, "MaxKeys": 100}
#             if token:
#                 kwargs["ContinuationToken"] = token
#             resp = _s3_client().list_objects_v2(**kwargs)
#             for c in resp.get("Contents", []):
#                 k = c.get("Key", "")
#                 if k.endswith(".json"):
#                     out.append(k)
#             token = resp.get("NextContinuationToken")
#             if not token:
#                 break
#     except (ClientError, EndpointConnectionError, NoCredentialsError):
#         return []
#     return out

# def presigned_url(key: str, expires: int = 3600) -> Optional[str]:
#     if not _s3_enabled():
#         return None
#     try:
#         return _s3_client().generate_presigned_url(
#             "get_object",
#             Params={"Bucket": _RP_BUCKET, "Key": key},
#             ExpiresIn=expires
#         )
#     except ClientError:
#         return None

# def ensure_local_copy(key_or_path: str) -> Optional[str]:
#     """
#     For DOCX/PDF parsing we need a real filesystem path.
#     If S3 mode, download to a temp file and return that path.
#     """
#     if not _s3_enabled():
#         return key_or_path if os.path.exists(key_or_path) else None

#     key = key_or_path
#     if key.startswith("s3://"):
#         parts = key.split("/", 3)
#         key = parts[3] if len(parts) >= 4 else ""
#     data = read_bytes_key(key)
#     if data is None:
#         return None
#     fd, tmp = tempfile.mkstemp(suffix=os.path.splitext(key)[1] or "")
#     os.close(fd)
#     with open(tmp, "wb") as f:
#         f.write(data)
#     return tmp

# def _s3_health_summary() -> dict:
#     info = {
#         "enabled": _s3_enabled(),
#         "endpoint": _RP_ENDPOINT,
#         "bucket": _RP_BUCKET,
#         "region": _RP_REGION,
#         "has_keys": bool(_AK and _SK),
#     }
#     if not _s3_enabled():
#         info["status"] = "local-mode"
#         return info
#     try:
#         _ = _s3_client().list_objects_v2(
#             Bucket=_RP_BUCKET,
#             Prefix=(f"{OUTPUT_DIR}/_history/").rstrip("/") + "/",
#             MaxKeys=1
#         )
#         info["status"] = "ok"
#     except Exception as e:
#         info["status"] = f"error: {getattr(e, 'response', {}).get('Error', {}).get('Code', str(e))}"
#     return info


# # ---------- Folders (all under Scriptmodel/) ----------
# BASE_PREFIX = "Scriptmodel"
# SCRIPTS_DIR = f"{BASE_PREFIX}/scripts"
# PROMPTS_DIR = f"{BASE_PREFIX}/prompts"
# OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
# HISTORY_DIR = f"{OUTPUT_DIR}/_history"

# if not _s3_enabled():
#     for p in (SCRIPTS_DIR, PROMPTS_DIR, OUTPUT_DIR, HISTORY_DIR):
#         Path(p).mkdir(parents=True, exist_ok=True)

# # ---------- Colors ----------
# PARAM_COLORS: Dict[str, str] = {
#     "Suspense Building":              "#ff6b6b",
#     "Language/Tone":                  "#6b8cff",
#     "Intro + Main Hook/Cliffhanger":  "#ffb86b",
#     "Story Structure + Flow":         "#a78bfa",
#     "Pacing":                         "#f43f5e",
#     "Mini-Hooks (30–60s)":            "#eab308",
#     "Outro (Ending)":                 "#8b5cf6",
# }

# # ---------- Config ----------
# STRICT_MATCH_ONLY = False

# # ---------- App config ----------
# st.set_page_config(page_title="M7 — Grammarly UI", page_icon="🕵️", layout="wide")

# def render_app_title():
#     st.markdown('<h1 class="app-title">Viral Script Reviewer</h1>', unsafe_allow_html=True)
#     st.markdown("""
#     <style>
#     html { color-scheme: light dark; }
#     :root{ --m7-surface:#eef2f7; --m7-on-surface:#0f172a; --m7-border:rgba(15,23,42,.14); --sep:#e5e7eb; }
#     @media (prefers-color-scheme: dark){
#       :root{ --m7-surface:#2f333a; --m7-on-surface:#fff; --m7-border:rgba(255,255,255,.18); --sep:#2a2f37; }
#     }
#     .stApp .block-container { padding-top: 4.25rem !important; }
#     .app-title{ font-weight:700;font-size:2.1rem;margin:0 0 1rem 0;padding-left:40px;padding-top:.25rem; }
#     @media (min-width: 992px){ .app-title { padding-left: 0 !important; } }
#     div[data-testid="column"]:nth-of-type(1){position:relative;}
#     div[data-testid="column"]:nth-of-type(1)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     div[data-testid="column"]:nth-of-type(2){position:relative;}
#     div[data-testid="column"]:nth-of-type(2)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     .m7-card{ background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:14px 16px; }
#     .docxwrap{ background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:16px 14px 18px; }
#     .docxwrap .h1,.docxwrap .h2,.docxwrap .h3{ font-weight:700; margin:10px 0 6px; }
#     .docxwrap .h1{ font-size:1.3rem; border-bottom:2px solid currentColor; padding-bottom:4px; }
#     .docxwrap .h2{ font-size:1.15rem; border-bottom:1px solid currentColor; padding-bottom:3px; }
#     .docxwrap .h3{ font-size:1.05rem; }
#     .docxwrap p{ margin:10px 0; line-height:1.7; font-family: ui-serif, Georgia, "Times New Roman", serif; }
#     .docxwrap table{ border-collapse:collapse; width:100%; margin:12px 0; }
#     .docxwrap th,.docxwrap td{ border:1px solid var(--m7-border); padding:8px; vertical-align:top; line-height:1.6; }
#     .docxwrap mark{ padding:0 2px; border-radius:3px; border:1px solid var(--m7-border); cursor:pointer; }
#     .rec-card{ display:block; text-decoration:none; background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:14px 16px; margin:10px 0 16px; box-shadow:0 1px 2px rgba(0,0,0,.06); }
#     .rec-card:hover{ filter:brightness(1.02); }
#     .stMarkdown pre, .stCodeBlock{ background:var(--m7-surface)!important; border:1px solid var(--m7-border)!important; border-radius:12px!important; padding:12px 14px!important; }
#     div[data-testid="stDataFrame"]{ background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:6px 8px; }
#     </style>
#     """, unsafe_allow_html=True)

# render_app_title()

# # ---------- Session ----------
# for key, default in [
#     ("review_ready", False),
#     ("script_text", ""),
#     ("base_stem", ""),
#     ("data", None),
#     ("spans_by_param", {}),
#     ("param_choice", None),
#     ("source_docx_path", None),
#     ("heading_ranges", []),
#     ("flattened_docx_path", None),
#     ("flatten_used", False),
#     ("ui_mode", "home"),
# ]:
#     st.session_state.setdefault(key, default)

# # ---------- helpers for query params (compat across Streamlit versions) ----------
# def _get_query_param(key: str) -> Optional[str]:
#     val = None
#     try:
#         val = st.query_params.get(key)
#     except Exception:
#         q = st.experimental_get_query_params()
#         v = q.get(key)
#         if isinstance(v, list): val = v[0] if v else None
#         else: val = v
#     return val

# def _clear_query_params():
#     try:
#         st.query_params.clear()
#     except Exception:
#         st.experimental_set_query_params()

# # ---------- Sanitizer ----------
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )
# def _sanitize_editor_text(s: Optional[str]) -> str:
#     if not s: return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•)\s*', '', t, flags=re.M)
#     t = re.sub(r'^\s*[-*]\s+', '• ', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# # ---------- DOCX traversal ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# # ---------- Auto-flatten ----------
# def _docx_contains_tables(path: str) -> bool:
#     doc = Document(path)
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Table):
#             return True
#     return False

# def _copy_paragraph(dest_doc: Document, src_para: Paragraph):
#     p = dest_doc.add_paragraph()
#     try:
#         if src_para.style and src_para.style.name:
#             p.style = src_para.style.name
#     except Exception:
#         pass
#     for run in src_para.runs:
#         r = p.add_run(run.text or "")
#         r.bold = run.bold
#         r.italic = run.italic
#         r.underline = run.underline
#     return p

# def flatten_docx_tables_to_longtext(source_path: str) -> str:
#     src = Document(source_path)
#     new = Document()
#     for blk in _iter_docx_blocks(src):
#         if isinstance(blk, Paragraph):
#             _copy_paragraph(new, blk)
#         else:
#             seen_tc_ids = set()
#             for row in blk.rows:
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     if tc_id in seen_tc_ids:
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     for p in cell.paragraphs:
#                         _copy_paragraph(new, p)
#                 new.add_paragraph("")
#             new.add_paragraph("")
#     fd, tmp_path = tempfile.mkstemp(suffix=".docx"); os.close(fd); new.save(tmp_path)
#     return tmp_path

# # ---------- Build plain text + heading ranges ----------
# def build_docx_text_with_meta(docx_path: str) -> Tuple[str, List[Tuple[int,int]]]:
#     doc = Document(docx_path)
#     out: List[str] = []
#     heading_ranges: List[Tuple[int,int]] = []
#     current_offset = 0

#     def _append_and_advance(s: str):
#         nonlocal current_offset
#         out.append(s); current_offset += len(s)

#     seen_tc_ids: set = set()
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             para_text = "".join(run.text or "" for run in blk.runs)
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             if sty.startswith("heading"):
#                 start = current_offset; end = start + len(para_text)
#                 heading_ranges.append((start, end))
#             _append_and_advance(para_text); _append_and_advance("\n")
#         else:
#             for row in blk.rows:
#                 row_cell_tcs = []
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     row_cell_tcs.append((tc_id, cell))
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     if tc_id in seen_tc_ids:
#                         if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     cell_text_parts: List[str] = []
#                     for i, p in enumerate(cell.paragraphs):
#                         t = "".join(r.text or "" for r in p.runs)
#                         sty = (p.style.name or "").lower() if p.style else ""
#                         if sty.startswith("heading"):
#                             hs = current_offset + sum(len(x) for x in cell_text_parts)
#                             he = hs + len(t)
#                             heading_ranges.append((hs, he))
#                         cell_text_parts.append(t)
#                         if i != len(cell.paragraphs) - 1:
#                             cell_text_parts.append("\n")
#                     cell_text = "".join(cell_text_parts)
#                     _append_and_advance(cell_text)
#                     if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                 _append_and_advance("\n")
#             _append_and_advance("\n")

#     return "".join(out), heading_ranges

# def _wrap_inline(safe_text: str, run) -> str:
#     out = safe_text
#     if getattr(run, "underline", False): out = f"<u>{out}</u>"
#     if getattr(run, "italic", False): out = f"<em>{out}</em>"
#     if getattr(run, "bold", False): out = f"<strong>{out}</strong>"
#     return out

# # ---------- Invisible/bridge chars ----------
# _BRIDGE_CHARS = set("\u200b\u200c\u200d\u2060\ufeff\xa0\u00ad")

# # ---------- DOCX -> HTML with highlights ----------
# def render_docx_html_with_highlights(docx_path: str,
#                                      highlight_spans: List[Tuple[int,int,str,str]]) -> str:
#     doc = Document(docx_path)
#     spans = [s for s in highlight_spans if s[0] < s[1]]
#     spans.sort(key=lambda x: x[0])

#     cur_span = 0
#     current_offset = 0

#     def esc(s: str) -> str:
#         return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

#     def open_mark_if_needed(html_parts, mark_state, color, end, aid):
#         if not mark_state["open"]:
#             html_parts.append(
#                 f'<mark class="aoi-mark" data-aid="{aid}" '
#                 f'style="background:{color}33;border:1px solid {color};border-radius:3px;padding:0 2px;">'
#             )
#             mark_state.update(open=True, end=end, color=color, aid=aid)

#     def close_mark_if_open(html_parts, mark_state):
#         if mark_state["open"]:
#             html_parts.append('</mark>')
#             mark_state.update(open=False, end=None, color=None, aid=None)

#     def emit_run_text(run_text: str, run, html_parts: List[str], mark_state: Dict[str, Any]):
#         nonlocal cur_span, current_offset
#         t = run_text or ""; i = 0
#         while i < len(t):
#             next_start, next_end, color, next_aid = None, None, None, None
#             if cur_span < len(spans):
#                 next_start, next_end, color, next_aid = spans[cur_span]

#             if not mark_state["open"]:
#                 if cur_span >= len(spans) or current_offset + (len(t) - i) <= next_start:
#                     chunk = t[i:]; html_parts.append(_wrap_inline(esc(chunk), run)); current_offset += len(chunk); break
#                 if current_offset < next_start:
#                     take = next_start - current_offset
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take; continue
#                 open_mark_if_needed(html_parts, mark_state, color, next_end, next_aid)
#             else:
#                 take = min(mark_state["end"] - current_offset, len(t) - i)
#                 if take > 0:
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take
#                 if current_offset >= mark_state["end"]:
#                     close_mark_if_open(html_parts, mark_state)
#                     cur_span += 1

#     html: List[str] = ['<div class="docxwrap">']
#     seen_tc_ids: set = set()

#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             open_tag = '<div class="h1">' if sty.startswith("heading 1") else \
#                        '<div class="h2">' if sty.startswith("heading 2") else \
#                        '<div class="h3">' if sty.startswith("heading 3") else "<p>"
#             close_tag = "</div>" if sty.startswith("heading") else "</p>"
#             html.append(open_tag)
#             for run in blk.runs:
#                 emit_run_text(run.text or "", run, html, mark_state)
#             close_mark_if_open(html, mark_state)
#             html.append(close_tag)
#             current_offset += 1  # '\n'

#         else:
#             html.append("<table>")
#             for row in blk.rows:
#                 html.append("<tr>")
#                 row_cell_tcs = [(id(cell._tc), cell) for cell in row.cells]
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     html.append("<td>")
#                     if tc_id not in seen_tc_ids:
#                         seen_tc_ids.add(tc_id)
#                         for p_idx, p in enumerate(cell.paragraphs):
#                             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#                             html.append("<div>")
#                             for run in p.runs:
#                                 emit_run_text(run.text or "", run, html, mark_state)
#                             close_mark_if_open(html, mark_state)
#                             html.append("</div>")
#                             if p_idx != len(cell.paragraphs) - 1:
#                                 current_offset += 1
#                     html.append("</td>")
#                     if idx != len(row_cell_tcs) - 1: current_offset += 1  # '\t'
#                 html.append("</tr>"); current_offset += 1   # row '\n'
#             html.append("</table>"); current_offset += 1     # extra '\n'

#     html.append("</div>")
#     return "".join(html)

# # ---------- Matching / span utilities ----------
# def _normalize_keep_len(s: str) -> str:
#     trans = {
#         "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
#         "\u2013": "-", "\u2014": "-",
#         "\xa0": " ",
#         "\u200b": " ", "\u200c": " ", "\u200d": " ", "\u2060": " ",
#         "\ufeff": " ", "\u00ad": " ",
#     }
#     return (s or "").translate(str.maketrans(trans))

# def _tokenize(s: str) -> List[str]:
#     return re.findall(r"\w+", (s or "").lower())

# def _iter_sentences_with_spans(text: str) -> List[Tuple[int,int,str]]:
#     spans = []
#     for m in re.finditer(r'[^.!?]+[.!?]+|\Z', text, flags=re.S):
#         s, e = m.start(), m.end()
#         seg = text[s:e]
#         if seg.strip(): spans.append((s, e, seg))
#     return spans

# def _squash_ws(s: str) -> str:
#     return re.sub(r"\s+", " ", s or "").strip()

# def _clean_quote_for_match(q: str) -> str:
#     if not q: return ""
#     q = _normalize_keep_len(q).strip()
#     q = re.sub(r'^[\'"“”‘’\[\(\{<…\-\–\—\s]+', '', q)
#     q = re.sub(r'[\'"“”‘’\]\)\}>…\-\–\—\s]+$', '', q)
#     return _squash_ws(q)

# def _snap_and_bridge_to_word(text: str, start: int, end: int, max_bridge: int = 2) -> Tuple[int,int]:
#     n = len(text); s, e = max(0,start), max(start,end)
#     def _is_inv(ch: str) -> bool: return ch in _BRIDGE_CHARS

#     while s > 0:
#         prev = text[s-1]; cur = text[s] if s < n else ""
#         if prev.isalnum() and cur.isalnum(): s -= 1; continue
#         j = s; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (s-1)>=0 and text[s-1].isalnum() and (j<n and text[j].isalnum()): s -= 1; continue
#         break

#     while e < n:
#         prev = text[e-1] if e>0 else ""; nxt = text[e]
#         if prev.isalnum() and nxt.isalnum(): e += 1; continue
#         j = e; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (e-1)>=0 and text[e-1].isalnum() and (j<n and text[j].isalnum()): e = j + 1; continue
#         break

#     while e < n and text[e] in ',"”’\')]}': e += 1
#     return s, e

# def _heal_split_word_left(text: str, start: int) -> int:
#     i = start
#     if i <= 1 or i >= len(text): return start
#     if text[i-1] != " ": return start
#     j = i - 2
#     while j >= 0 and text[j].isalpha(): j -= 1
#     prev_token = text[j+1:i-1]
#     if len(prev_token) == 1: return i - 2
#     return start

# def _overlaps_any(s: int, e: int, ranges: List[Tuple[int,int]]) -> bool:
#     for rs, re_ in ranges:
#         if e > rs and s < re_: return True
#     return False

# def _fuzzy_window_span(tl: str, nl: str, start: int, w: int) -> Tuple[float, Optional[Tuple[int,int]]]:
#     window = tl[start:start+w]
#     sm = difflib.SequenceMatcher(a=nl, b=window)
#     blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
#     if not blocks: return 0.0, None
#     coverage = sum(b.size for b in blocks) / max(1, len(nl))
#     first_b = min(blocks, key=lambda b: b.b); last_b = max(blocks, key=lambda b: b.b + b.size)
#     s = start + first_b.b; e = start + last_b.b + last_b.size
#     return coverage, (s, e)

# def find_span_smart(text: str, needle: str) -> Optional[Tuple[int,int]]:
#     if not text or not needle: return None
#     t_orig = text; t_norm = _normalize_keep_len(text); n_norm = _clean_quote_for_match(needle)
#     if not n_norm: return None
#     tl = t_norm.lower(); nl = n_norm.lower()

#     i = tl.find(nl)
#     if i != -1:
#         s, e = _snap_and_bridge_to_word(t_orig, i, i + len(nl)); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     m = re.search(re.escape(nl).replace(r"\ ", r"\s+"), tl, flags=re.IGNORECASE)
#     if m:
#         s, e = _snap_and_bridge_to_word(t_orig, m.start(), m.end()); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     if not STRICT_MATCH_ONLY and len(nl) >= 12:
#         w = max(60, min(240, len(nl) + 80))
#         best_cov, best_span = 0.0, None
#         step = max(1, w // 2)
#         for start in range(0, max(1, len(tl) - w + 1), step):
#             cov, se = _fuzzy_window_span(tl, nl, start, w)
#             if cov > best_cov: best_cov, best_span = cov, se
#         if best_span and best_cov >= 0.65:
#             s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
#             if s > 0 and t_orig[s-1:s+1].lower() in {" the", " a", " an"}: s -= 1
#             s = _heal_split_word_left(t_orig, s)
#             return (s, e)

#     if not STRICT_MATCH_ONLY:
#         keys = [w for w in _tokenize(nl) if len(w) >= 4][:8]
#         if len(keys) >= 2:
#             kset = set(keys)
#             best_score, best_span = 0.0, None
#             for s, e, seg in _iter_sentences_with_spans(t_norm):
#                 toks = set(_tokenize(seg)); ov = len(kset & toks)
#                 if ov == 0: continue
#                 score = ov / max(2, len(kset)); length_pen = min(1.0, 120 / max(20, e - s)); score *= (0.6 + 0.4 * length_pen)
#                 if score > best_score: best_score, best_span = score, (s, min(e, s + 400))
#             if best_span and best_score >= 0.35:
#                 s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1]); s = _heal_split_word_left(t_orig, s)
#                 return (s, e)
#     return None

# def merge_overlaps(spans: List[Tuple[int,int,str,str]]) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans.sort(key=lambda x: x[0]); out = [spans[0]]
#     for s,e,c,aid in spans[1:]:
#         ps,pe,pc,paid = out[-1]
#         if s <= pe and pc == c and e > pe: out[-1] = (ps, e, pc, paid)
#         else: out.append((s,e,c,aid))
#     return out

# _PUNCT_WS = set(" \t\r\n,.;:!?)('\"[]{}-—–…") | _BRIDGE_CHARS

# def merge_overlaps_and_adjacent(base_text: str,
#                                 spans: List[Tuple[int,int,str,str]],
#                                 max_gap: int = 2) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans = sorted(spans, key=lambda x: x[0]); out = [spans[0]]
#     for s, e, c, aid in spans[1:]:
#         ps, pe, pc, paid = out[-1]
#         if c == pc and s <= pe: out[-1] = (ps, max(pe, e), pc, paid); continue
#         if c == pc and s - pe <= max_gap:
#             gap = base_text[max(0, pe):max(0, s)]
#             if all((ch in _PUNCT_WS) for ch in gap): out[-1] = (ps, e, pc, paid); continue
#         out.append((s, e, c, aid))
#     return out

# # ---------- Heading filters ----------
# def _is_heading_like(q: str) -> bool:
#     if not q: return True
#     s = q.strip()
#     if not re.search(r'[.!?]', s):
#         words = re.findall(r"[A-Za-z]+", s)
#         if 1 <= len(words) <= 7:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.8: return True
#         if s.lower() in {"introduction","voiceover","outro","epilogue","prologue","credits","title","hook","horrifying discovery","final decision"}: return True
#         if len(s) <= 3: return True
#     return False

# def _is_heading_context(script_text: str, s: int, e: int) -> bool:
#     left = script_text.rfind("\n", 0, s) + 1
#     right = script_text.find("\n", e); right = len(script_text) if right == -1 else right
#     line = script_text[left:right].strip()
#     if len(line) <= 70 and not re.search(r'[.!?]', line):
#         words = re.findall(r"[A-Za-z]+", line)
#         if 1 <= len(words) <= 8:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.7: return True
#     return False

# def _tighten_to_quote(script_text: str, span: Tuple[int,int], quote: str) -> Tuple[int,int]:
#     if not span or not quote: return span
#     s, e = span
#     if e <= s or s < 0 or e > len(script_text): return span
#     window = script_text[s:e]; win_norm = _normalize_keep_len(window).lower(); q_norm = _clean_quote_for_match(quote).lower()
#     if not q_norm: return span
#     i = win_norm.find(q_norm)
#     if i == -1:
#         m = re.search(re.escape(q_norm).replace(r"\ ", r"\s+"), win_norm, flags=re.IGNORECASE)
#         if not m: return span
#         i, j = m.start(), m.end()
#     else:
#         j = i + len(q_norm)
#     s2, e2 = s + i, s + j
#     s2, e2 = _snap_and_bridge_to_word(script_text, s2, e2); s2 = _heal_split_word_left(script_text, s2)
#     if s2 >= s and e2 <= e and e2 > s2: return (s2, e2)
#     return span

# def build_spans_by_param(script_text: str, data: dict, heading_ranges: Optional[List[Tuple[int,int]]] = None) -> Dict[str, List[Tuple[int,int,str,str]]]:
#     heading_ranges = heading_ranges or []
#     raw = (data or {}).get("per_parameter", {}) or {}
#     per: Dict[str, Dict[str, Any]] = {k:(v or {}) for k,v in raw.items()}
#     spans_map: Dict[str, List[Tuple[int,int,str,str]]] = {p: [] for p in PARAM_ORDER}
#     st.session_state["aoi_match_ranges"] = {}

#     for p in spans_map.keys():
#         color = PARAM_COLORS.get(p, "#ffd54f")
#         blk = per.get(p, {}) or {}
#         aois = blk.get("areas_of_improvement") or []
#         for idx, item in enumerate(aois, start=1):
#             raw_q = (item or {}).get("quote_verbatim", "") or ""
#             q = _sanitize_editor_text(raw_q)
#             clean = _clean_quote_for_match(re.sub(r"^[•\-\d\.\)\s]+", "", q).strip())
#             if not clean: continue
#             if _is_heading_like(clean): continue

#             pos = find_span_smart(script_text, clean)
#             if not pos: continue
#             pos = _tighten_to_quote(script_text, pos, raw_q)
#             s, e = pos

#             if heading_ranges and _overlaps_any(s, e, heading_ranges): continue
#             if _is_heading_context(script_text, s, e): continue

#             aid = f"{p.replace(' ','_')}-AOI-{idx}"
#             spans_map[p].append((s, e, color, aid))
#             st.session_state["aoi_match_ranges"][aid] = (s, e)
#     return spans_map

# # ---------- History (S3-aware) ----------
# def _maybe_copy_docx_to_history(source_docx_path: Optional[str], run_id: str) -> Optional[str]:
#     """Copy the DOCX used for rendering into outputs/_history so Recents can re-render identically."""
#     try:
#         if not source_docx_path:
#             return None
#         # If already an S3 key/url, just store the key in history blob
#         if source_docx_path.startswith("s3://") or (_s3_enabled() and not os.path.exists(source_docx_path)):
#             # ensure it's uploaded under our history key if it's a local temp in S3 mode
#             if os.path.exists(source_docx_path):
#                 with open(source_docx_path, "rb") as f:
#                     save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#                 return f"{HISTORY_DIR}/{run_id}.docx"
#             return source_docx_path

#         # Local file path
#         if _s3_enabled():
#             with open(source_docx_path, "rb") as f:
#                 save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#             return f"{HISTORY_DIR}/{run_id}.docx"
#         else:
#             dst = os.path.join(HISTORY_DIR, f"{run_id}.docx")
#             if os.path.abspath(source_docx_path) != os.path.abspath(dst):
#                 shutil.copyfile(source_docx_path, dst)
#             return dst
#     except Exception:
#         return None

# @st.cache_data(ttl=5)
# def _cached_history_keys() -> List[str]:
#     """Short-lived cache of history keys to reduce S3 round-trips."""
#     return list_prefix(HISTORY_DIR)

# def _save_history_snapshot(title: str, data: dict, script_text: str,
#                            source_docx_path: Optional[str], heading_ranges: List[Tuple[int,int]],
#                            spans_by_param: Dict[str, List[Tuple[int,int,str,str]]],
#                            aoi_match_ranges: Dict[str, Tuple[int,int]]):
#     run_id = str(uuid.uuid4()); now = datetime.datetime.now()
#     created_at_iso = now.replace(microsecond=0).isoformat()
#     created_at_human = now.strftime("%Y-%m-%d %H:%M:%S")

#     # Copy stable DOCX alongside history JSON (so Recents uses identical renderer)
#     stable_docx_key_or_path = _maybe_copy_docx_to_history(source_docx_path, run_id)

#     blob = {
#         "run_id": run_id, "title": title or "untitled",
#         "created_at": created_at_iso, "created_at_human": created_at_human,
#         "overall_rating": (data or {}).get("overall_rating", ""),
#         "scores": (data or {}).get("scores", {}),
#         "data": data or {}, "script_text": script_text or "",
#         "source_docx_path": stable_docx_key_or_path or source_docx_path,
#         "heading_ranges": heading_ranges or [],
#         "spans_by_param": spans_by_param or {},
#         "aoi_match_ranges": aoi_match_ranges or {},
#     }

#     # Persist JSON (S3 or local)
#     out_name = f"{created_at_iso.replace(':','-')}__{run_id}.json"
#     out_key = f"{HISTORY_DIR}/{out_name}"
#     save_text_key(out_key, json.dumps(blob, ensure_ascii=False, indent=2))

#     # Invalidate cached keys so Recents sees the new run immediately
#     st.cache_data.clear()

# def _load_all_history() -> List[dict]:
#     out: List[dict] = []
#     if _s3_enabled():
#         keys = sorted(_cached_history_keys(), reverse=True)
#         for key in keys:
#             try:
#                 txt = read_text_key(key, "")
#                 if not txt:
#                     continue
#                 j = json.loads(txt)
#                 j["_key"] = key
#                 if not j.get("created_at_human") and j.get("created_at"):
#                     try:
#                         dt = datetime.datetime.fromisoformat(j["created_at"])
#                         j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#                     except Exception:
#                         pass
#                 out.append(j)
#             except Exception:
#                 continue
#     else:
#         for fp in sorted(glob.glob(os.path.join(HISTORY_DIR, "*.json"))):
#             try:
#                 with open(fp, "r", encoding="utf-8") as f:
#                     j = json.load(f)
#             except Exception:
#                 continue
#             j.setdefault("_path", fp)
#             ca = j.get("created_at")
#             try:
#                 if isinstance(ca, (int, float)):
#                     dt = datetime.datetime.utcfromtimestamp(float(ca))
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.astimezone().strftime("%Y-%m-%d %H:%M:%S")
#                 elif isinstance(ca, str) and ca:
#                     pass
#                 else:
#                     mtime = os.path.getmtime(fp); dt = datetime.datetime.fromtimestamp(mtime)
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#             except Exception:
#                 j["created_at"] = str(ca or "")
#             out.append(j)
#     out.sort(key=lambda r: r.get("created_at") or "", reverse=True)
#     return out

# def _open_history_run_by_id(run_id: str) -> bool:
#     """Open a history run by its run_id. Returns True if loaded."""
#     if not run_id: return False
#     recs = _load_all_history()
#     match = next((r for r in recs if r.get("run_id") == run_id), None)
#     if not match: return False

#     # Load JSON content again (S3/local), then set session
#     try:
#         if "_key" in match and _s3_enabled():
#             txt = read_text_key(match["_key"], "")
#             if not txt: return False
#             jj = json.loads(txt)
#         else:
#             path = match.get("_path")
#             if not path or not os.path.exists(path): return False
#             with open(path, "r", encoding="utf-8") as f:
#                 jj = json.load(f)
#     except Exception:
#         return False

#     st.session_state.script_text      = jj.get("script_text","")
#     st.session_state.base_stem        = jj.get("title","untitled")
#     st.session_state.data             = jj.get("data",{})
#     st.session_state.heading_ranges   = jj.get("heading_ranges",[])
#     st.session_state.spans_by_param   = jj.get("spans_by_param",{})
#     st.session_state.param_choice     = None
#     st.session_state.source_docx_path = jj.get("source_docx_path")
#     st.session_state.review_ready     = True
#     st.session_state["aoi_match_ranges"] = jj.get("aoi_match_ranges", {})
#     st.session_state.ui_mode          = "review"
#     return True

# def _render_recents_centerpane():
#     st.subheader("📄 Recents")
#     q = st.text_input("Filter by title…", "")
#     cols = st.columns([1, 4])
#     with cols[0]:
#         if st.button("← Back"):
#             st.session_state.ui_mode = "home"; st.rerun()

#     # 🔄 Spinner while fetching from RunPod S3 (no artificial delay)
#     if _s3_enabled():
#         with st.spinner("Fetching history from RunPod S3…"):
#             recs = _load_all_history()
#     else:
#         with st.spinner("Loading local history…"):
#             recs = _load_all_history()

#     ql = q.strip().lower()
#     if ql: recs = [r for r in recs if ql in (r.get("title","").lower())]
#     if not recs:
#         st.caption("No history yet."); st.stop()

#     # Each card is a clickable <a class="rec-card" href="?open=<run_id>">
#     for rec in recs:
#         run_id = rec.get("run_id"); title = rec.get("title") or "(untitled)"
#         created_h = rec.get("created_at_human",""); overall = rec.get("overall_rating","")
#         st.markdown(f"""
#         <a class="rec-card" href="?open={run_id}">
#           <div class="rec-title">{title}</div>
#           <div class="rec-meta">{created_h}</div>
#           <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
#         </a>
#         """, unsafe_allow_html=True)

# # ---------- Sidebar ----------
# with st.sidebar:
#     if st.button("🆕 New review", use_container_width=True):
#         fp = st.session_state.get("flattened_docx_path")
#         if fp and os.path.exists(fp):
#             try: os.remove(fp)
#             except Exception: pass
#         for k in ["review_ready","script_text","base_stem","data","spans_by_param","param_choice",
#                   "source_docx_path","heading_ranges","flattened_docx_path","flatten_used"]:
#             st.session_state[k] = (
#                 False if k=="review_ready"
#                 else "" if k in ("script_text","base_stem")
#                 else {} if k=="spans_by_param"
#                 else [] if k=="heading_ranges"
#                 else None if k in ("source_docx_path","flattened_docx_path")
#                 else False if k=="flatten_used"
#                 else None
#             )
#         st.session_state.ui_mode = "home"
#         _clear_query_params()
#         st.rerun()

#     if st.button("📁 Recents", use_container_width=True):
#         st.session_state.ui_mode = "recents"
#         _clear_query_params()
#         st.rerun()

#     # 🔎 DEBUGGER: storage panel (lazy — runs only when clicked)
#     with st.sidebar.expander("Debug: storage", expanded=False):
#         if st.button("Run storage diagnostics"):
#             st.write("S3 health summary:")
#             try:
#                 st.json(_s3_health_summary())
#             except Exception as e:
#                 st.write(f"(error reading health: {e})")

#             st.write("History prefix:", HISTORY_DIR)
#             try:
#                 if _s3_enabled():
#                     with st.spinner("Listing RunPod S3 keys…"):
#                         _keys = list_prefix(HISTORY_DIR)
#                 else:
#                     _keys = list_prefix(HISTORY_DIR)

#                 st.write("Found history keys:", len(_keys))
#                 for k in _keys[:10]:
#                     st.write("•", k)

#                 if _keys:
#                     if _s3_enabled():
#                         with st.spinner("Reading first history JSON from S3…"):
#                             txt = read_text_key(_keys[0], "")
#                     else:
#                         txt = read_text_key(_keys[0], "")
#                     st.write("First JSON (truncated):")
#                     st.code((txt[:800] + ("…" if len(txt) > 800 else "")) if txt else "(empty file)")
#             except Exception as e:
#                 st.write(f"(error listing/reading: {e})")
#         else:
#             st.caption("Click the button to run diagnostics")

# # ---------- Input screen ----------
# def render_home():
#     st.subheader("🎬 Script Source")

#     (tab_upload,) = st.tabs(["Upload file"])
#     uploaded_file = None
#     uploaded_name = None
#     uploaded_key  = None

#     with tab_upload:
#         up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
#         if up is not None:
#             file_bytes = up.read()
#             suffix = os.path.splitext(up.name)[1].lower()
#             # Save to S3 (or local) under scripts/
#             uploaded_key = f"{SCRIPTS_DIR}/{up.name}"
#             save_bytes_key(uploaded_key, file_bytes)

#             # Also create a temp local copy for parsing
#             with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
#                 tmp.write(file_bytes)
#                 tmp.flush()
#                 uploaded_file = tmp.name
#             uploaded_name = os.path.splitext(os.path.basename(up.name))[0] or "uploaded_script"

#     if st.button("🚀 Run Review", type="primary", use_container_width=True):
#         base_stem = "uploaded_script"
#         source_docx_path = None
#         heading_ranges = []

#         if uploaded_file:
#             base_stem = uploaded_name or "uploaded_script"
#             # Prefer key (so we can store it into history); for DOCX rendering we keep a tmp local
#             if uploaded_file.lower().endswith(".docx"):
#                 path_to_use = uploaded_file
#                 if _docx_contains_tables(path_to_use):
#                     flat = flatten_docx_tables_to_longtext(path_to_use)
#                     st.session_state.flattened_docx_path = flat
#                     st.session_state.flatten_used = True
#                     path_to_use = flat
#                 script_text, heading_ranges = build_docx_text_with_meta(path_to_use)

#                 # 🔑 IMPORTANT: set the rendering source to the *flattened* docx
#                 # (History saver will copy this file to Scriptmodel/outputs/_history/)
#                 source_docx_path = path_to_use
#             else:
#                 script_text = load_script_file(uploaded_file)
#                 source_docx_path = uploaded_key or uploaded_file  # keep reference to S3/local
#         else:
#             st.warning("Please upload a script first.")
#             st.stop()

#         if len(script_text.strip()) < 50:
#             st.error("Extracted text looks too short. Check your file extraction.")
#             st.stop()

#         with st.spinner("Running analysis…"):
#             try:
#                 review_text = run_review_multi(
#                     script_text=script_text,
#                     prompts_dir=PROMPTS_DIR,
#                     temperature=0.0
#                 )
#             finally:
#                 # Clean temp upload if we didn't keep it as source_docx_path local
#                 if uploaded_file and not (isinstance(source_docx_path, str) and os.path.exists(source_docx_path)):
#                     try:
#                         os.remove(uploaded_file)
#                     except Exception:
#                         pass

#         data = extract_review_json(review_text)
#         if not data:
#             st.error("JSON not detected in model output.")
#             st.stop()

#         st.session_state.script_text      = script_text
#         st.session_state.base_stem        = base_stem
#         st.session_state.data             = data
#         st.session_state.heading_ranges   = heading_ranges
#         st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
#         st.session_state.param_choice     = None
#         st.session_state.source_docx_path = source_docx_path
#         st.session_state.review_ready     = True
#         st.session_state.ui_mode          = "review"

#         _save_history_snapshot(
#             title=base_stem,
#             data=data,
#             script_text=script_text,
#             source_docx_path=source_docx_path,
#             heading_ranges=heading_ranges,
#             spans_by_param=st.session_state.spans_by_param,
#             aoi_match_ranges=st.session_state.get("aoi_match_ranges", {})
#         )

#         _clear_query_params()
#         st.rerun()

# # ---------- Results screen ----------
# def render_review():
#     script_text     = st.session_state.script_text
#     data            = st.session_state.data
#     spans_by_param  = st.session_state.spans_by_param
#     scores: Dict[str,int] = (data or {}).get("scores", {}) or {}
#     source_docx_path: Optional[str] = st.session_state.source_docx_path

#     # If our source_docx_path is an S3 key/url, ensure we have a local copy for rendering
#     # AFTER (prefer the flattened copy if we have it in-session)
#     docx_local: Optional[str] = None
#     preferred = st.session_state.get("flattened_docx_path") if st.session_state.get("flatten_used") else source_docx_path
#     if preferred:
#         if not os.path.exists(preferred):
#             docx_local = ensure_local_copy(preferred)
#         else:
#             docx_local = preferred

#     left, center, right = st.columns([1.1, 2.7, 1.4], gap="large")

#     with left:
#         st.subheader("Final score")
#         ordered = [p for p in PARAM_ORDER if p in scores]
#         df = pd.DataFrame({"Parameter": ordered, "Score (1–10)": [scores.get(p, "") for p in ordered]})
#         st.dataframe(df, hide_index=True, use_container_width=True)
#         st.markdown(f'**Overall:** {data.get("overall_rating","—")}/10')
#         st.divider()

#         strengths = (data or {}).get("strengths") or []
#         if not strengths:
#             per = (data or {}).get("per_parameter", {}) or {}
#             best = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
#             for name, sc in best:
#                 if sc >= 8 and name in per:
#                     exp = _sanitize_editor_text((per[name] or {}).get("explanation", "") or "")
#                     first = re.split(r"(?<=[.!?])\s+", exp.strip())[0] if exp else f"Consistently strong {name.lower()}."
#                     strengths.append(f"{name}: {first}")
#                 if len(strengths) >= 3: break

#         def _bullets(title: str, items):
#             st.markdown(f"**{title}**")
#             for s in (items or []):
#                 if isinstance(s, str) and s.strip(): st.write("• " + _sanitize_editor_text(s))
#             if not items: st.write("• —")

#         _bullets("Strengths", strengths)
#         _bullets("Weaknesses", data.get("weaknesses"))
#         _bullets("Suggestions", data.get("suggestions"))
#         _bullets("Drop-off Risks", data.get("drop_off_risks"))
#         st.markdown("**Viral Quotient**"); st.write(_sanitize_editor_text(data.get("viral_quotient","—")))

#     with right:
#         st.subheader("Parameters")
#         st.markdown('<div class="param-row">', unsafe_allow_html=True)
#         for p in [p for p in PARAM_ORDER if p in scores]:
#             if st.button(p, key=f"chip_{p}", help="Show inline AOI highlights for this parameter"):
#                 st.session_state.param_choice = p
#         st.markdown('</div>', unsafe_allow_html=True)

#         sel = st.session_state.param_choice
#         if sel:
#             blk = (data.get("per_parameter", {}) or {}).get(sel, {}) or {}
#             st.markdown(f"**{sel} — Score:** {scores.get(sel,'—')}/10")

#             if blk.get("explanation"):
#                 st.markdown("**Why this score**"); st.write(_sanitize_editor_text(blk["explanation"]))
#             if blk.get("weakness") and blk["weakness"] != "Not present":
#                 st.markdown("**Weakness**"); st.write(_sanitize_editor_text(blk["weakness"]))
#             if blk.get("suggestion") and blk["suggestion"] != "Not present":
#                 st.markdown("**Suggestion**"); st.write(_sanitize_editor_text(blk["suggestion"]))

#             aoi = blk.get("areas_of_improvement") or []
#             if aoi:
#                 st.markdown("**Areas of Improvement**")
#                 for i, item in enumerate(aoi, 1):
#                     popover_fn = getattr(st, "popover", None)
#                     aid = f"{sel.replace(' ','_')}-AOI-{i}"
#                     s_e_map = st.session_state.get("aoi_match_ranges", {})
#                     if aid in s_e_map:
#                         s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                         line = (matched_line if len(matched_line) <= 320 else matched_line[:300] + "…")
#                     else:
#                         line = _sanitize_editor_text(item.get('quote_verbatim',''))
#                     issue = _sanitize_editor_text(item.get('issue',''))
#                     fix   = _sanitize_editor_text(item.get('fix',''))
#                     why   = _sanitize_editor_text(item.get('why_this_helps',''))
#                     label = f"Issue {i}"
#                     if callable(popover_fn):
#                         with popover_fn(label):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)
#                     else:
#                         with st.expander(label, expanded=False):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)

#             if blk.get("summary"):
#                 st.markdown("**Summary**"); st.write(_sanitize_editor_text(blk["summary"]))

#     with center:
#         st.subheader("Script with inline highlights")
#         # choose spans for selected parameter (or all if None)
#         spans = st.session_state.spans_by_param.get(st.session_state.param_choice, []) if st.session_state.param_choice else []

#         # Build AOI payload for popup: { aid: {line, issue, fix, why} }
#         aoi_payload: Dict[str, Dict[str, str]] = {}
#         data_per = (data or {}).get("per_parameter") or {}
#         s_e_map = st.session_state.get("aoi_match_ranges", {})
#         sel = st.session_state.param_choice

#         def _mk_line(aid: str, fallback_q: str = "") -> str:
#             if aid in s_e_map:
#                 s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                 return matched_line if len(matched_line) <= 320 else matched_line[:300] + "…"
#             return _sanitize_editor_text(fallback_q or "")

#         def _collect(param_name: str):
#             blk = (data_per.get(param_name) or {})
#             for i, item in enumerate(blk.get("areas_of_improvement") or [], 1):
#                 aid = f"{param_name.replace(' ','_')}-AOI-{i}"
#                 aoi_payload[aid] = {
#                     "line": _mk_line(aid, (item or {}).get("quote_verbatim","")),
#                     "issue": _sanitize_editor_text((item or {}).get("issue","")),
#                     "fix": _sanitize_editor_text((item or {}).get("fix","")),
#                     "why": _sanitize_editor_text((item or {}).get("why_this_helps","")),
#                 }

#         if sel: _collect(sel)
#         else:
#             for pn in [p for p in PARAM_ORDER if p in data_per]:
#                 _collect(pn)

#         payload_json = json.dumps(aoi_payload, ensure_ascii=False)

#         frame_theme_css = """
#         <style>
#           :root{ --m7-surface:#eef2f7; --m7-on-surface:#0f172a; --m7-border:rgba(15,23,42,.14); }
#           @media (prefers-color-scheme: dark){
#             :root{ --m7-surface:#2f333a; --m7-on-surface:#ffffff; --m7-border:rgba(255,255,255,.18); }
#             body { background: transparent !important; }
#           }
#           .docxwrap{ background:var(--m7-surface); color:var(--m7-on-surface); border:1px solid var(--m7-border); border-radius:12px; padding:16px 14px 18px; }
#           .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#           .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); }
#         </style>
#         """

#         tooltip_css = """
#         <style>
#         .aoi-pop {
#           position: absolute; max-width: 520px; min-width: 320px;
#           background: var(--m7-surface); border: 1px solid var(--m7-border); border-radius: 10px;
#           box-shadow: 0 10px 25px rgba(0,0,0,.12);
#           padding: 12px 14px; z-index: 9999; transform: translateY(-8px);
#           color: var(--m7-on-surface);
#         }
#         .aoi-pop p  { margin: .15rem 0; line-height: 1.5; }
#         .aoi-pop .muted { opacity:.85; font-size:.85rem; }
#         .aoi-arrow {
#           position:absolute; left:50%; transform:translateX(-50%);
#           bottom:-7px; width:0;height:0;border-left:7px solid transparent;
#           border-right:7px solid transparent;border-top:7px solid var(--m7-border);
#         }
#         .aoi-arrow::after{
#           content:""; position:absolute; left:-6px; top:-7px; width:0;height:0;
#           border-left:6px solid transparent;border-right:6px solid transparent;border-top:6px solid var(--m7-surface);
#         }
#         </style>
#         """

#         # Select rendering source (DOCX with highlights if we have a local path)
#         if docx_local and os.path.splitext(docx_local)[1].lower() == ".docx":
#             html_core = render_docx_html_with_highlights(
#                 docx_local,
#                 merge_overlaps_and_adjacent(script_text, spans)
#             )
#         else:
#             from html import escape as _esc
#             orig = script_text
#             spans2 = [s for s in merge_overlaps_and_adjacent(orig, spans) if s[0] < s[1]]
#             spans2.sort(key=lambda x: x[0])
#             cur = 0; buf: List[str] = []
#             for s,e,c,aid in spans2:
#                 if s > cur: buf.append(_esc(orig[cur:s]))
#                 buf.append(
#                     f'<mark class="aoi-mark" data-aid="{aid}" '
#                     f'style="background:{c}33;border:1px solid {c};border-radius:3px;padding:0 2px;">'
#                     f'{_esc(orig[s:e])}</mark>'
#                 )
#                 cur = e
#             if cur < len(orig): buf.append(_esc(orig[cur:]))
#             html_core = (
#                 '<div class="docxwrap"><p style="white-space:pre-wrap; '
#                 'line-height:1.7; font-family:ui-serif, Georgia, Times New Roman, serif;">'
#                 + "".join(buf) +
#                 '</p></div>'
#             )

#         # Popup + autosize JS shell
#         html_shell = """
# %%FRAME_THEME_CSS%%
# %%TOOLTIP_CSS%%
# <div id="m7-doc">%%HTML_CORE%%</div>
# <div id="aoi-pop" class="aoi-pop" style="display:none;">
#   <div id="aoi-pop-content"></div>
#   <div class="aoi-arrow"></div>
# </div>
# <script>
# (function(){
#   const AOI = __PAYLOAD__;
#   const wrap = document.getElementById('m7-doc');
#   const pop  = document.getElementById('aoi-pop');
#   const body = document.getElementById('aoi-pop-content');

#   function resizeIframe() {
#     try {
#       const h = Math.max(
#         document.documentElement.scrollHeight,
#                document.body.scrollHeight
#       );
#       if (window.frameElement) {
#         window.frameElement.style.height = (h + 20) + 'px';
#         window.frameElement.style.width  = '100%';
#       }
#     } catch(e) {}
#   }
#   window.addEventListener('load', resizeIframe);
#   window.addEventListener('resize', resizeIframe);

#   function hide(){ pop.style.display='none'; }
#   function showFor(mark){
#     const aid = mark.getAttribute('data-aid');
#     const d = AOI[aid]; if(!d) return;
#     body.innerHTML =
#       (d.line  ? '<p><strong>Line:</strong> '  + d.line  + '</p>' : '') +
#       (d.issue ? '<p><strong>Issue:</strong> ' + d.issue + '</p>' : '') +
#       (d.fix   ? '<p><strong>Fix:</strong> '   + d.fix   + '</p>' : '') +
#       (d.why   ? '<p class="muted">'           + d.why   + '</p>' : '');
#     pop.style.display = 'block';

#     const r = mark.getBoundingClientRect();
#     const scY = window.scrollY || document.documentElement.scrollTop;
#     const scX = window.scrollX || document.documentElement.scrollLeft;
#     let top  = r.top + scY - pop.offsetHeight - 10;
#     let left = r.left + scX + r.width/2 - pop.offsetWidth/2;
#     if (top < 8) top = r.bottom + scY + 10;
#     if (left < 8) left = 8;
#     pop.style.top  = top + 'px';
#     pop.style.left = left + 'px';

#     resizeIframe();
#   }

#   wrap.addEventListener('click', (e) => {
#     const m = e.target.closest('.aoi-mark');
#     if(!m){ hide(); return; }
#     if(pop.style.display === 'block'){ hide(); }
#     showFor(m);
#     e.stopPropagation();
#   });

#   document.addEventListener('click', (e) => {
#     if(!e.target.closest('.aoi-pop') && !e.target.closest('.aoi-mark')) hide();
#   });
# })();
# </script>
# """
#         html_shell = (
#             html_shell
#             .replace("%%FRAME_THEME_CSS%%", frame_theme_css)
#             .replace("%%TOOLTIP_CSS%%", tooltip_css)
#             .replace("%%HTML_CORE%%", html_core)
#             .replace("__PAYLOAD__", payload_json)
#         )

#         components.html(html_shell, height=400, scrolling=False)

# # ---------- Router & query param open (with S3 spinner) ----------
# _open_qp = _get_query_param("open")
# if _open_qp:
#     if _s3_enabled():
#         with st.spinner("Fetching run from RunPod S3…"):
#             if _open_history_run_by_id(_open_qp):
#                 _clear_query_params()
#     else:
#         if _open_history_run_by_id(_open_qp):
#             _clear_query_params()

# mode = st.session_state.ui_mode
# if mode == "recents":
#     _render_recents_centerpane()
# elif mode == "review" and st.session_state.review_ready:
#     render_review()
# else:
#     render_home()











##############################################################################











# import os, re, glob, json, tempfile, difflib, uuid, datetime, shutil
# from pathlib import Path
# from typing import Dict, Any, List, Tuple, Optional

# import streamlit as st
# import pandas as pd
# import streamlit.components.v1 as components  # for inline HTML/JS popup

# from utils1 import extract_review_json, PARAM_ORDER, load_script_file
# from review_engine_multi import run_review_multi

# # ---- DOCX rendering imports ----
# from docx import Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # =========================
# # RunPod S3 (inline helpers) — FIXED
# # =========================
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError

# # Accept both AWS_* and RUNPOD_* style envs / st.secrets
# def _get_env(key: str, default: str = "") -> str:
#     v = os.getenv(key, "")
#     if v: 
#         return v.strip()
#     try:
#         v2 = st.secrets.get(key)
#         if isinstance(v2, str):
#             return v2.strip()
#     except Exception:
#         pass
#     return (default or "").strip()

# # Primary config
# _RP_ENDPOINT = _get_env("RUNPOD_S3_ENDPOINT")
# _RP_BUCKET   = _get_env("RUNPOD_S3_BUCKET")
# _RP_REGION   = _get_env("RUNPOD_S3_REGION") or _get_env("AWS_DEFAULT_REGION") or ""

# # Credentials: prefer AWS_* if present; else accept RUNPOD_* fallbacks
# _AK = _get_env("AWS_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY")
# _SK = _get_env("AWS_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_KEY")
# _ST = _get_env("AWS_SESSION_TOKEN")  # optional

# # Options
# _FORCE_PATH = (_get_env("RUNPOD_S3_FORCE_PATH_STYLE") or "true").lower() in {"1","true","yes"}
# _USE_SSL    = (_get_env("RUNPOD_S3_USE_SSL") or "true").lower() in {"1","true","yes"}
# _VERIFY_SSL = (_get_env("RUNPOD_S3_VERIFY_SSL") or "true").lower() in {"1","true","yes"}

# def _s3_enabled() -> bool:
#     return bool(_RP_ENDPOINT and _RP_BUCKET and _AK and _SK)

# @st.cache_resource(show_spinner=False)
# def _s3_client():
#     if not _s3_enabled():
#         return None
#     session_kwargs = dict(
#         aws_access_key_id=_AK,
#         aws_secret_access_key=_SK,
#     )
#     if _ST:
#         session_kwargs["aws_session_token"] = _ST

#     # s3v4 + path style are common requirements for S3-compatible services
#     cfg = Config(
#         signature_version="s3v4",
#         s3={"addressing_style": "path" if _FORCE_PATH else "auto"},
#         retries={"max_attempts": 3, "mode": "standard"}
#     )
#     return boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         region_name=_RP_REGION or None,
#         use_ssl=_USE_SSL,
#         verify=_VERIFY_SSL,
#         config=cfg,
#         **session_kwargs,
#     )

# def save_text_key(key: str, text: str) -> str:
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "w", encoding="utf-8") as f:
#             f.write(text)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=text.encode("utf-8"))
#     return f"s3://{_RP_BUCKET}/{key}"

# def save_bytes_key(key: str, data: bytes) -> str:
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "wb") as f:
#             f.write(data)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=data)
#     return f"s3://{_RP_BUCKET}/{key}"

# def read_text_key(key: str, default: str = "") -> str:
#     if not _s3_enabled():
#         try:
#             with open(key, "r", encoding="utf-8") as f:
#                 return f.read()
#         except Exception:
#             return default
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read().decode("utf-8", errors="ignore")
#     except Exception:
#         return default

# def read_bytes_key(key: str) -> Optional[bytes]:
#     if not _s3_enabled():
#         try:
#             with open(key, "rb") as f:
#                 return f.read()
#         except Exception:
#             return None
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read()
#     except Exception:
#         return None

# def list_prefix(prefix: str) -> List[str]:
#     """
#     List object keys under prefix (or local dir paths if not S3).
#     In S3 mode we always return KEYS (not URLs).
#     """
#     if not _s3_enabled():
#         base = prefix if os.path.isdir(prefix) else os.path.dirname(prefix)
#         try:
#             return [os.path.join(base, p) for p in os.listdir(base) if p.endswith(".json")]
#         except Exception:
#             return []

#     out: List[str] = []
#     token = None
#     # Normalize to "dir/" prefix for S3 listing
#     s3_prefix = prefix.rstrip("/") + "/"
#     try:
#         while True:
#             kwargs = {"Bucket": _RP_BUCKET, "Prefix": s3_prefix}
#             if token:
#                 kwargs["ContinuationToken"] = token
#             resp = _s3_client().list_objects_v2(**kwargs)
#             for c in resp.get("Contents", []):
#                 k = c.get("Key", "")
#                 if k.endswith(".json"):
#                     out.append(k)
#             token = resp.get("NextContinuationToken")
#             if not token:
#                 break
#     except (ClientError, EndpointConnectionError, NoCredentialsError):
#         return []
#     return out

# def presigned_url(key: str, expires: int = 3600) -> Optional[str]:
#     if not _s3_enabled():
#         return None
#     try:
#         return _s3_client().generate_presigned_url(
#             "get_object",
#             Params={"Bucket": _RP_BUCKET, "Key": key},
#             ExpiresIn=expires
#         )
#     except ClientError:
#         return None

# def ensure_local_copy(key_or_path: str) -> Optional[str]:
#     """
#     For DOCX/PDF parsing we need a real filesystem path.
#     If S3 mode, download to a temp file and return that path.
#     """
#     if not _s3_enabled():
#         return key_or_path if os.path.exists(key_or_path) else None

#     key = key_or_path
#     if key.startswith("s3://"):
#         # s3://bucket/path/to/file -> path/to/file
#         parts = key.split("/", 3)
#         key = parts[3] if len(parts) >= 4 else ""
#     data = read_bytes_key(key)
#     if data is None:
#         return None
#     fd, tmp = tempfile.mkstemp(suffix=os.path.splitext(key)[1] or "")
#     os.close(fd)
#     with open(tmp, "wb") as f:
#         f.write(data)
#     return tmp

# def _s3_health_summary() -> dict:
#     """
#     Optional tiny health read you can print if needed.
#     Returns a dict; safe to ignore in production.
#     """
#     info = {
#         "enabled": _s3_enabled(),
#         "endpoint": _RP_ENDPOINT,
#         "bucket": _RP_BUCKET,
#         "region": _RP_REGION,
#         "has_keys": bool(_AK and _SK),
#     }
#     if not _s3_enabled():
#         info["status"] = "local-mode"
#         return info
#     try:
#         # Attempt a very cheap list; no exceptions => reachable
#         _ = _s3_client().list_objects_v2(Bucket=_RP_BUCKET, Prefix=(f"{OUTPUT_DIR}/_history/").rstrip("/") + "/",
#                                          MaxKeys=1)
#         info["status"] = "ok"
#     except Exception as e:
#         info["status"] = f"error: {getattr(e, 'response', {}).get('Error', {}).get('Code', str(e))}"
#     return info


# # ---------- Folders ----------
# # SCRIPTS_DIR = "scripts"
# # PROMPTS_DIR = "prompts"
# # OUTPUT_DIR  = "outputs"
# # HISTORY_DIR = os.path.join(OUTPUT_DIR, "_history")
# # ---------- Folders (all under Scriptmodel/) ----------

# BASE_PREFIX = "Scriptmodel"

# SCRIPTS_DIR = f"{BASE_PREFIX}/scripts"
# PROMPTS_DIR = f"{BASE_PREFIX}/prompts"
# OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
# HISTORY_DIR = f"{OUTPUT_DIR}/_history"


# if not _s3_enabled():
#     for p in (SCRIPTS_DIR, PROMPTS_DIR, OUTPUT_DIR, HISTORY_DIR):
#         Path(p).mkdir(parents=True, exist_ok=True)


# # ---------- Colors ----------
# PARAM_COLORS: Dict[str, str] = {
#     "Suspense Building":              "#ff6b6b",
#     "Language/Tone":                  "#6b8cff",
#     "Intro + Main Hook/Cliffhanger":  "#ffb86b",
#     "Story Structure + Flow":         "#a78bfa",
#     "Pacing":                         "#f43f5e",
#     "Mini-Hooks (30–60s)":            "#eab308",
#     "Outro (Ending)":                 "#8b5cf6",
# }

# # ---------- Config ----------
# STRICT_MATCH_ONLY = False

# # ---------- App config ----------
# st.set_page_config(page_title="M7 — Grammarly UI", page_icon="🕵️", layout="wide")
# # ---------- Header patch & CSS ----------
# def render_app_title():
#     st.markdown(
#         '<h1 class="app-title">Viral Script Reviewer</h1>',
#         unsafe_allow_html=True
#     )
#     st.markdown("""
#     <style>
#     html { color-scheme: light dark; }

#     :root{
#       --m7-surface: #eef2f7;
#       --m7-on-surface: #0f172a;
#       --m7-border: rgba(15,23,42,.14);
#       --sep: #e5e7eb;
#     }
#     @media (prefers-color-scheme: dark){
#       :root{
#         --m7-surface: #2f333a;
#         --m7-on-surface: #ffffff;
#         --m7-border: rgba(255,255,255,.18);
#         --sep: #2a2f37;
#       }
#     }

#     .stApp .block-container { padding-top: 4.25rem !important; }
#     .app-title{
#       font-weight: 700; font-size: 2.1rem; line-height: 1.3;
#       margin: 0 0 1rem 0; padding-left: 40px !important; padding-top: .25rem !important;
#       white-space: normal; word-break: break-word; hyphens: auto; overflow: visible;
#       position: relative !important; z-index: 10 !important;
#     }
#     [data-testid="collapsedControl"] { z-index: 6 !important; }
#     header[data-testid="stHeader"], .stAppHeader { background: transparent !important; box-shadow: none !important; }
#     @media (min-width: 992px){ .app-title { padding-left: 0 !important; } }

#     div[data-testid="column"]:nth-of-type(1){position:relative;}
#     div[data-testid="column"]:nth-of-type(1)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     div[data-testid="column"]:nth-of-type(2){position:relative;}
#     div[data-testid="column"]:nth-of-type(2)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}

#     .m7-card{
#       background: var(--m7-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 14px 16px;
#       color: var(--m7-on-surface);
#     }
#     .m7-card, .m7-card * { color: var(--m7-on-surface) !important; }

#     .docxwrap{
#       background: var(--m7-surface);
#       color: var(--m7-on-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 16px 14px 18px;
#     }
#     .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#     .docxwrap .h1, .docxwrap .h2, .docxwrap .h3 { font-weight:700; margin:10px 0 6px; }
#     .docxwrap .h1 { font-size: 1.3rem; border-bottom: 2px solid currentColor; padding-bottom: 4px; }
#     .docxwrap .h2 { font-size: 1.15rem; border-bottom: 1px solid currentColor; padding-bottom: 3px; }
#     .docxwrap .h3 { font-size: 1.05rem; }
#     .docxwrap p { margin: 10px 0; line-height: 1.7; font-family: ui-serif, Georgia, "Times New Roman", serif; }
#     .docxwrap table { border-collapse: collapse; width: 100%; margin: 12px 0; }
#     .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); padding:8px; vertical-align:top; line-height:1.6; }
#     .docxwrap mark{ padding:0 2px; border-radius:3px; border:1px solid var(--m7-border); cursor: pointer; }

#     .rec-card{
#       display:block; text-decoration:none !important;
#       background: var(--m7-surface);
#       border:1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 14px 16px; margin: 10px 0 16px;
#       box-shadow: 0 1px 2px rgba(0,0,0,.06);
#       color: var(--m7-on-surface) !important;
#       transition: filter .1s ease, transform .02s ease;
#     }
#     .rec-card:hover{ filter: brightness(1.02); }
#     .rec-card:active{ transform: translateY(1px); }
#     .rec-card, .rec-card * { color: var(--m7-on-surface) !important; }
#     .rec-title{font-weight:600; margin-bottom:.25rem;}
#     .rec-meta{opacity:.85 !important; font-size:12.5px; margin-bottom:.4rem;}
#     .rec-row{display:flex; align-items:center; justify-content:space-between; gap:12px;}

#     .stTextInput>div>div,
#     .stTextArea>div>div,
#     .stNumberInput>div>div,
#     .stDateInput>div>div,
#     .stTimeInput>div>div,
#     .stFileUploader>div,
#     div[data-baseweb="select"]{
#       background: var(--m7-surface) !important;
#       border: 1px solid var(--m7-border) !important;
#       border-radius: 10px !important;
#       color: var(--m7-on-surface) !important;
#     }
#     .stTextInput input,
#     .stTextArea textarea,
#     .stNumberInput input,
#     .stDateInput input,
#     .stTimeInput input,
#     .stFileUploader div,
#     div[data-baseweb="select"] *{
#       color: var(--m7-on-surface) !important;
#     }
#     .stTextInput input::placeholder,
#     .stTextArea textarea::placeholder{ color: rgba(16,24,39,.55) !important; }
#     @media (prefers-color-scheme: dark){
#       .stTextInput input::placeholder,
#       .stTextArea textarea::placeholder{ color: rgba(255,255,255,.75) !important; }
#     }
#     div[data-testid="stFileUploaderDropzone"] label span { color: var(--m7-on-surface) !important; opacity: 1 !important; }
#     div[data-testid="stFileUploaderDropzone"] label { color: var(--m7-on-surface) !important; }

#     .stMarkdown pre,
#     pre[class*="language-"],
#     .stCodeBlock{
#       background: var(--m7-surface) !important;
#       color: var(--m7-on-surface) !important;
#       border: 1px solid var(--m7-border) !important;
#       border-radius: 12px !important;
#       padding: 12px 14px !important;
#       overflow:auto;
#     }
#     .stMarkdown pre code{ background: transparent !important; color: inherit !important; }

#     div[data-testid="stDataFrame"]{
#       background: var(--m7-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 6px 8px;
#       color: var(--m7-on-surface);
#     }
#     div[data-testid="stDataFrame"] * { color: var(--m7-on-surface) !important; }

#     </style>
#     """, unsafe_allow_html=True)

# render_app_title()

# # ---------- Session ----------
# for key, default in [
#     ("review_ready", False),
#     ("script_text", ""),
#     ("base_stem", ""),
#     ("data", None),
#     ("spans_by_param", {}),
#     ("param_choice", None),
#     ("source_docx_path", None),
#     ("heading_ranges", []),
#     ("flattened_docx_path", None),
#     ("flatten_used", False),
#     ("ui_mode", "home"),
# ]:
#     st.session_state.setdefault(key, default)

# # ---------- helpers for query params (compat across Streamlit versions) ----------
# def _get_query_param(key: str) -> Optional[str]:
#     val = None
#     try:
#         val = st.query_params.get(key)
#     except Exception:
#         q = st.experimental_get_query_params()
#         v = q.get(key)
#         if isinstance(v, list): val = v[0] if v else None
#         else: val = v
#     return val

# def _clear_query_params():
#     try:
#         st.query_params.clear()
#     except Exception:
#         st.experimental_set_query_params()

# # ---------- Sanitizer ----------
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )
# def _sanitize_editor_text(s: Optional[str]) -> str:
#     if not s: return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•)\s*', '', t, flags=re.M)
#     t = re.sub(r'^\s*[-*]\s+', '• ', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# # ---------- DOCX traversal ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# # ---------- Auto-flatten ----------
# def _docx_contains_tables(path: str) -> bool:
#     doc = Document(path)
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Table):
#             return True
#     return False

# def _copy_paragraph(dest_doc: Document, src_para: Paragraph):
#     p = dest_doc.add_paragraph()
#     try:
#         if src_para.style and src_para.style.name:
#             p.style = src_para.style.name
#     except Exception:
#         pass
#     for run in src_para.runs:
#         r = p.add_run(run.text or "")
#         r.bold = run.bold
#         r.italic = run.italic
#         r.underline = run.underline
#     return p

# def flatten_docx_tables_to_longtext(source_path: str) -> str:
#     src = Document(source_path)
#     new = Document()
#     for blk in _iter_docx_blocks(src):
#         if isinstance(blk, Paragraph):
#             _copy_paragraph(new, blk)
#         else:
#             seen_tc_ids = set()
#             for row in blk.rows:
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     if tc_id in seen_tc_ids:
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     for p in cell.paragraphs:
#                         _copy_paragraph(new, p)
#                 new.add_paragraph("")
#             new.add_paragraph("")
#     fd, tmp_path = tempfile.mkstemp(suffix=".docx"); os.close(fd); new.save(tmp_path)
#     return tmp_path

# # ---------- Build plain text + heading ranges ----------
# def build_docx_text_with_meta(docx_path: str) -> Tuple[str, List[Tuple[int,int]]]:
#     doc = Document(docx_path)
#     out: List[str] = []
#     heading_ranges: List[Tuple[int,int]] = []
#     current_offset = 0

#     def _append_and_advance(s: str):
#         nonlocal current_offset
#         out.append(s); current_offset += len(s)

#     seen_tc_ids: set = set()
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             para_text = "".join(run.text or "" for run in blk.runs)
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             if sty.startswith("heading"):
#                 start = current_offset; end = start + len(para_text)
#                 heading_ranges.append((start, end))
#             _append_and_advance(para_text); _append_and_advance("\n")
#         else:
#             for row in blk.rows:
#                 row_cell_tcs = []
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     row_cell_tcs.append((tc_id, cell))
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     if tc_id in seen_tc_ids:
#                         if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     cell_text_parts: List[str] = []
#                     for i, p in enumerate(cell.paragraphs):
#                         t = "".join(r.text or "" for r in p.runs)
#                         sty = (p.style.name or "").lower() if p.style else ""
#                         if sty.startswith("heading"):
#                             hs = current_offset + sum(len(x) for x in cell_text_parts)
#                             he = hs + len(t)
#                             heading_ranges.append((hs, he))
#                         cell_text_parts.append(t)
#                         if i != len(cell.paragraphs) - 1:
#                             cell_text_parts.append("\n")
#                     cell_text = "".join(cell_text_parts)
#                     _append_and_advance(cell_text)
#                     if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                 _append_and_advance("\n")
#             _append_and_advance("\n")

#     return "".join(out), heading_ranges

# def _wrap_inline(safe_text: str, run) -> str:
#     out = safe_text
#     if getattr(run, "underline", False): out = f"<u>{out}</u>"
#     if getattr(run, "italic", False): out = f"<em>{out}</em>"
#     if getattr(run, "bold", False): out = f"<strong>{out}</strong>"
#     return out

# # ---------- Invisible/bridge chars ----------
# _BRIDGE_CHARS = set("\u200b\u200c\u200d\u2060\ufeff\xa0\u00ad")

# # ---------- DOCX -> HTML with highlights (includes data-aid) ----------
# def render_docx_html_with_highlights(docx_path: str,
#                                      highlight_spans: List[Tuple[int,int,str,str]]) -> str:
#     """
#     highlight_spans: list of (start, end, color, aid)
#     """
#     doc = Document(docx_path)
#     spans = [s for s in highlight_spans if s[0] < s[1]]
#     spans.sort(key=lambda x: x[0])

#     cur_span = 0
#     current_offset = 0

#     def esc(s: str) -> str:
#         return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

#     def open_mark_if_needed(html_parts, mark_state, color, end, aid):
#         if not mark_state["open"]:
#             html_parts.append(
#                 f'<mark class="aoi-mark" data-aid="{aid}" '
#                 f'style="background:{color}33;border:1px solid {color};border-radius:3px;padding:0 2px;">'
#             )
#             mark_state.update(open=True, end=end, color=color, aid=aid)

#     def close_mark_if_open(html_parts, mark_state):
#         if mark_state["open"]:
#             html_parts.append('</mark>')
#             mark_state.update(open=False, end=None, color=None, aid=None)

#     def emit_run_text(run_text: str, run, html_parts: List[str], mark_state: Dict[str, Any]):
#         nonlocal cur_span, current_offset
#         t = run_text or ""; i = 0
#         while i < len(t):
#             next_start, next_end, color, next_aid = None, None, None, None
#             if cur_span < len(spans):
#                 next_start, next_end, color, next_aid = spans[cur_span]

#             if not mark_state["open"]:
#                 if cur_span >= len(spans) or current_offset + (len(t) - i) <= next_start:
#                     chunk = t[i:]; html_parts.append(_wrap_inline(esc(chunk), run)); current_offset += len(chunk); break
#                 if current_offset < next_start:
#                     take = next_start - current_offset
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take; continue
#                 open_mark_if_needed(html_parts, mark_state, color, next_end, next_aid)
#             else:
#                 take = min(mark_state["end"] - current_offset, len(t) - i)
#                 if take > 0:
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take
#                 if current_offset >= mark_state["end"]:
#                     close_mark_if_open(html_parts, mark_state)
#                     cur_span += 1

#     html: List[str] = ['<div class="docxwrap">']
#     seen_tc_ids: set = set()

#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             open_tag = '<div class="h1">' if sty.startswith("heading 1") else \
#                        '<div class="h2">' if sty.startswith("heading 2") else \
#                        '<div class="h3">' if sty.startswith("heading 3") else "<p>"
#             close_tag = "</div>" if sty.startswith("heading") else "</p>"
#             html.append(open_tag)
#             for run in blk.runs:
#                 emit_run_text(run.text or "", run, html, mark_state)
#             close_mark_if_open(html, mark_state)
#             html.append(close_tag)
#             current_offset += 1  # '\n'

#         else:
#             html.append("<table>")
#             for row in blk.rows:
#                 html.append("<tr>")
#                 row_cell_tcs = [(id(cell._tc), cell) for cell in row.cells]
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     html.append("<td>")
#                     if tc_id not in seen_tc_ids:
#                         seen_tc_ids.add(tc_id)
#                         for p_idx, p in enumerate(cell.paragraphs):
#                             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#                             html.append("<div>")
#                             for run in p.runs:
#                                 emit_run_text(run.text or "", run, html, mark_state)
#                             close_mark_if_open(html, mark_state)
#                             html.append("</div>")
#                             if p_idx != len(cell.paragraphs) - 1:
#                                 current_offset += 1
#                     html.append("</td>")
#                     if idx != len(row_cell_tcs) - 1: current_offset += 1  # '\t'
#                 html.append("</tr>"); current_offset += 1   # row '\n'
#             html.append("</table>"); current_offset += 1     # extra '\n'

#     html.append("</div>")
#     return "".join(html)

# # ---------- Matching / span utilities ----------
# def _normalize_keep_len(s: str) -> str:
#     trans = {
#         "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
#         "\u2013": "-", "\u2014": "-",
#         "\xa0": " ",
#         "\u200b": " ", "\u200c": " ", "\u200d": " ", "\u2060": " ",
#         "\ufeff": " ", "\u00ad": " ",
#     }
#     return (s or "").translate(str.maketrans(trans))

# def _tokenize(s: str) -> List[str]: return re.findall(r"\w+", (s or "").lower())

# def _iter_sentences_with_spans(text: str) -> List[Tuple[int,int,str]]:
#     spans = []
#     for m in re.finditer(r'[^.!?]+[.!?]+|\Z', text, flags=re.S):
#         s, e = m.start(), m.end()
#         seg = text[s:e]
#         if seg.strip(): spans.append((s, e, seg))
#     return spans

# def _squash_ws(s: str) -> str: return re.sub(r"\s+", " ", s or "").strip()

# def _clean_quote_for_match(q: str) -> str:
#     if not q: return ""
#     q = _normalize_keep_len(q).strip()
#     q = re.sub(r'^[\'"“”‘’\[\(\{<…\-\–\—\s]+', '', q)
#     q = re.sub(r'[\'"“”‘’\]\)\}>…\-\–\—\s]+$', '', q)
#     return _squash_ws(q)

# def _snap_and_bridge_to_word(text: str, start: int, end: int, max_bridge: int = 2) -> Tuple[int,int]:
#     n = len(text); s, e = max(0,start), max(start,end)
#     def _is_inv(ch: str) -> bool: return ch in _BRIDGE_CHARS

#     while s > 0:
#         prev = text[s-1]; cur = text[s] if s < n else ""
#         if prev.isalnum() and cur.isalnum(): s -= 1; continue
#         j = s; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1; 
#         if brid and (s-1)>=0 and text[s-1].isalnum() and (j<n and text[j].isalnum()): s -= 1; continue
#         break

#     while e < n:
#         prev = text[e-1] if e>0 else ""; nxt = text[e]
#         if prev.isalnum() and nxt.isalnum(): e += 1; continue
#         j = e; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (e-1)>=0 and text[e-1].isalnum() and (j<n and text[j].isalnum()): e = j + 1; continue
#         break

#     while e < n and text[e] in ',"”’\')]}': e += 1
#     return s, e

# def _heal_split_word_left(text: str, start: int) -> int:
#     i = start
#     if i <= 1 or i >= len(text): return start
#     if text[i-1] != " ": return start
#     j = i - 2
#     while j >= 0 and text[j].isalpha(): j -= 1
#     prev_token = text[j+1:i-1]
#     if len(prev_token) == 1: return i - 2
#     return start

# def _overlaps_any(s: int, e: int, ranges: List[Tuple[int,int]]) -> bool:
#     for rs, re_ in ranges:
#         if e > rs and s < re_: return True
#     return False

# def _fuzzy_window_span(tl: str, nl: str, start: int, w: int) -> Tuple[float, Optional[Tuple[int,int]]]:
#     window = tl[start:start+w]
#     sm = difflib.SequenceMatcher(a=nl, b=window)
#     blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
#     if not blocks: return 0.0, None
#     coverage = sum(b.size for b in blocks) / max(1, len(nl))
#     first_b = min(blocks, key=lambda b: b.b); last_b = max(blocks, key=lambda b: b.b + b.size)
#     s = start + first_b.b; e = start + last_b.b + last_b.size
#     return coverage, (s, e)

# def find_span_smart(text: str, needle: str) -> Optional[Tuple[int,int]]:
#     if not text or not needle: return None
#     t_orig = text; t_norm = _normalize_keep_len(text); n_norm = _clean_quote_for_match(needle)
#     if not n_norm: return None
#     tl = t_norm.lower(); nl = n_norm.lower()

#     i = tl.find(nl)
#     if i != -1:
#         s, e = _snap_and_bridge_to_word(t_orig, i, i + len(nl)); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     m = re.search(re.escape(nl).replace(r"\ ", r"\s+"), tl, flags=re.IGNORECASE)
#     if m:
#         s, e = _snap_and_bridge_to_word(t_orig, m.start(), m.end()); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     if not STRICT_MATCH_ONLY and len(nl) >= 12:
#         w = max(60, min(240, len(nl) + 80))
#         best_cov, best_span = 0.0, None
#         step = max(1, w // 2)
#         for start in range(0, max(1, len(tl) - w + 1), step):
#             cov, se = _fuzzy_window_span(tl, nl, start, w)
#             if cov > best_cov: best_cov, best_span = cov, se
#         if best_span and best_cov >= 0.65:
#             s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
#             if s > 0 and t_orig[s-1:s+1].lower() in {" the", " a", " an"}: s -= 1
#             s = _heal_split_word_left(t_orig, s)
#             return (s, e)

#     if not STRICT_MATCH_ONLY:
#         keys = [w for w in _tokenize(nl) if len(w) >= 4][:8]
#         if len(keys) >= 2:
#             kset = set(keys)
#             best_score, best_span = 0.0, None
#             for s, e, seg in _iter_sentences_with_spans(t_norm):
#                 toks = set(_tokenize(seg)); ov = len(kset & toks)
#                 if ov == 0: continue
#                 score = ov / max(2, len(kset)); length_pen = min(1.0, 120 / max(20, e - s)); score *= (0.6 + 0.4 * length_pen)
#                 if score > best_score: best_score, best_span = score, (s, min(e, s + 400))
#             if best_span and best_score >= 0.35:
#                 s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1]); s = _heal_split_word_left(t_orig, s)
#                 return (s, e)
#     return None

# def merge_overlaps(spans: List[Tuple[int,int,str,str]]) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans.sort(key=lambda x: x[0]); out = [spans[0]]
#     for s,e,c,aid in spans[1:]:
#         ps,pe,pc,paid = out[-1]
#         if s <= pe and pc == c and e > pe: out[-1] = (ps, e, pc, paid)
#         else: out.append((s,e,c,aid))
#     return out

# _PUNCT_WS = set(" \t\r\n,.;:!?)('\"[]{}-—–…") | _BRIDGE_CHARS

# def merge_overlaps_and_adjacent(base_text: str,
#                                 spans: List[Tuple[int,int,str,str]],
#                                 max_gap: int = 2) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans = sorted(spans, key=lambda x: x[0]); out = [spans[0]]
#     for s, e, c, aid in spans[1:]:
#         ps, pe, pc, paid = out[-1]
#         if c == pc and s <= pe: out[-1] = (ps, max(pe, e), pc, paid); continue
#         if c == pc and s - pe <= max_gap:
#             gap = base_text[max(0, pe):max(0, s)]
#             if all((ch in _PUNCT_WS) for ch in gap): out[-1] = (ps, e, pc, paid); continue
#         out.append((s, e, c, aid))
#     return out

# # ---------- Heading filters ----------
# def _is_heading_like(q: str) -> bool:
#     if not q: return True
#     s = q.strip()
#     if not re.search(r'[.!?]', s):
#         words = re.findall(r"[A-Za-z]+", s)
#         if 1 <= len(words) <= 7:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.8: return True
#         if s.lower() in {"introduction","voiceover","outro","epilogue","prologue","credits","title","hook","horrifying discovery","final decision"}: return True
#         if len(s) <= 3: return True
#     return False

# def _is_heading_context(script_text: str, s: int, e: int) -> bool:
#     left = script_text.rfind("\n", 0, s) + 1
#     right = script_text.find("\n", e); right = len(script_text) if right == -1 else right
#     line = script_text[left:right].strip()
#     if len(line) <= 70 and not re.search(r'[.!?]', line):
#         words = re.findall(r"[A-Za-z]+", line)
#         if 1 <= len(words) <= 8:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.7: return True
#     return False

# def _tighten_to_quote(script_text: str, span: Tuple[int,int], quote: str) -> Tuple[int,int]:
#     if not span or not quote: return span
#     s, e = span
#     if e <= s or s < 0 or e > len(script_text): return span
#     window = script_text[s:e]; win_norm = _normalize_keep_len(window).lower(); q_norm = _clean_quote_for_match(quote).lower()
#     if not q_norm: return span
#     i = win_norm.find(q_norm)
#     if i == -1:
#         m = re.search(re.escape(q_norm).replace(r"\ ", r"\s+"), win_norm, flags=re.IGNORECASE)
#         if not m: return span
#         i, j = m.start(), m.end()
#     else:
#         j = i + len(q_norm)
#     s2, e2 = s + i, s + j
#     s2, e2 = _snap_and_bridge_to_word(script_text, s2, e2); s2 = _heal_split_word_left(script_text, s2)
#     if s2 >= s and e2 <= e and e2 > s2: return (s2, e2)
#     return span

# def build_spans_by_param(script_text: str, data: dict, heading_ranges: Optional[List[Tuple[int,int]]] = None) -> Dict[str, List[Tuple[int,int,str,str]]]:
#     heading_ranges = heading_ranges or []
#     raw = (data or {}).get("per_parameter", {}) or {}
#     per: Dict[str, Dict[str, Any]] = {k:(v or {}) for k,v in raw.items()}
#     spans_map: Dict[str, List[Tuple[int,int,str,str]]] = {p: [] for p in PARAM_ORDER}
#     st.session_state["aoi_match_ranges"] = {}

#     for p in spans_map.keys():
#         color = PARAM_COLORS.get(p, "#ffd54f")
#         blk = per.get(p, {}) or {}
#         aois = blk.get("areas_of_improvement") or []
#         for idx, item in enumerate(aois, start=1):
#             raw_q = (item or {}).get("quote_verbatim", "") or ""
#             q = _sanitize_editor_text(raw_q)
#             clean = _clean_quote_for_match(re.sub(r"^[•\-\d\.\)\s]+", "", q).strip())
#             if not clean: continue
#             if _is_heading_like(clean): continue

#             pos = find_span_smart(script_text, clean)
#             if not pos: continue
#             pos = _tighten_to_quote(script_text, pos, raw_q)
#             s, e = pos

#             if heading_ranges and _overlaps_any(s, e, heading_ranges): continue
#             if _is_heading_context(script_text, s, e): continue

#             aid = f"{p.replace(' ','_')}-AOI-{idx}"
#             spans_map[p].append((s, e, color, aid))
#             st.session_state["aoi_match_ranges"][aid] = (s, e)
#     return spans_map

# # ---------- History (S3-aware) ----------
# def _maybe_copy_docx_to_history(source_docx_path: Optional[str], run_id: str) -> Optional[str]:
#     """Copy the DOCX used for rendering into outputs/_history so Recents can re-render identically."""
#     try:
#         if not source_docx_path:
#             return None
#         # If already an S3 key/url, just store the key in history blob
#         if source_docx_path.startswith("s3://") or (_s3_enabled() and not os.path.exists(source_docx_path)):
#             # ensure it's uploaded under our history key if it's a local temp in S3 mode
#             if os.path.exists(source_docx_path):
#                 with open(source_docx_path, "rb") as f:
#                     save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#                 return f"{HISTORY_DIR}/{run_id}.docx"
#             return source_docx_path

#         # Local file path
#         if _s3_enabled():
#             with open(source_docx_path, "rb") as f:
#                 save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#             return f"{HISTORY_DIR}/{run_id}.docx"
#         else:
#             dst = os.path.join(HISTORY_DIR, f"{run_id}.docx")
#             if os.path.abspath(source_docx_path) != os.path.abspath(dst):
#                 shutil.copyfile(source_docx_path, dst)
#             return dst
#     except Exception:
#         return None

# def _save_history_snapshot(title: str, data: dict, script_text: str,
#                            source_docx_path: Optional[str], heading_ranges: List[Tuple[int,int]],
#                            spans_by_param: Dict[str, List[Tuple[int,int,str,str]]],
#                            aoi_match_ranges: Dict[str, Tuple[int,int]]):
#     run_id = str(uuid.uuid4()); now = datetime.datetime.now()
#     created_at_iso = now.replace(microsecond=0).isoformat()
#     created_at_human = now.strftime("%Y-%m-%d %H:%M:%S")

#     # Copy stable DOCX alongside history JSON (so Recents uses identical renderer)
#     stable_docx_key_or_path = _maybe_copy_docx_to_history(source_docx_path, run_id)

#     blob = {
#         "run_id": run_id, "title": title or "untitled",
#         "created_at": created_at_iso, "created_at_human": created_at_human,
#         "overall_rating": (data or {}).get("overall_rating", ""),
#         "scores": (data or {}).get("scores", {}),
#         "data": data or {}, "script_text": script_text or "",
#         "source_docx_path": stable_docx_key_or_path or source_docx_path,
#         "heading_ranges": heading_ranges or [],
#         "spans_by_param": spans_by_param or {},
#         "aoi_match_ranges": aoi_match_ranges or {},
#     }

#     # Persist JSON (S3 or local)
#     out_name = f"{created_at_iso.replace(':','-')}__{run_id}.json"
#     out_key = f"{HISTORY_DIR}/{out_name}"
#     save_text_key(out_key, json.dumps(blob, ensure_ascii=False, indent=2))

# def _load_all_history() -> List[dict]:
#     out: List[dict] = []
#     if _s3_enabled():
#         keys = sorted(list_prefix(HISTORY_DIR), reverse=True)
#         for key in keys:
#             try:
#                 txt = read_text_key(key, "")
#                 if not txt:
#                     continue
#                 j = json.loads(txt)
#                 j["_key"] = key
#                 if not j.get("created_at_human") and j.get("created_at"):
#                     try:
#                         dt = datetime.datetime.fromisoformat(j["created_at"])
#                         j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#                     except Exception:
#                         pass
#                 out.append(j)
#             except Exception:
#                 continue
#     else:
#         for fp in sorted(glob.glob(os.path.join(HISTORY_DIR, "*.json"))):
#             try:
#                 with open(fp, "r", encoding="utf-8") as f:
#                     j = json.load(f)
#             except Exception:
#                 continue
#             j.setdefault("_path", fp)
#             ca = j.get("created_at")
#             try:
#                 if isinstance(ca, (int, float)):
#                     dt = datetime.datetime.utcfromtimestamp(float(ca))
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.astimezone().strftime("%Y-%m-%d %H:%M:%S")
#                 elif isinstance(ca, str) and ca: pass
#                 else:
#                     mtime = os.path.getmtime(fp); dt = datetime.datetime.fromtimestamp(mtime)
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#             except Exception:
#                 j["created_at"] = str(ca or "")
#             out.append(j)
#     out.sort(key=lambda r: r.get("created_at") or "", reverse=True)
#     return out

# def _open_history_run_by_id(run_id: str) -> bool:
#     """Open a history run by its run_id. Returns True if loaded."""
#     if not run_id: return False
#     recs = _load_all_history()
#     match = next((r for r in recs if r.get("run_id") == run_id), None)
#     if not match: return False

#     # Load JSON content again (S3/local), then set session
#     try:
#         if "_key" in match and _s3_enabled():
#             txt = read_text_key(match["_key"], "")
#             if not txt: return False
#             jj = json.loads(txt)
#         else:
#             path = match.get("_path")
#             if not path or not os.path.exists(path): return False
#             with open(path, "r", encoding="utf-8") as f:
#                 jj = json.load(f)
#     except Exception:
#         return False

#     st.session_state.script_text      = jj.get("script_text","")
#     st.session_state.base_stem        = jj.get("title","untitled")
#     st.session_state.data             = jj.get("data",{})
#     st.session_state.heading_ranges   = jj.get("heading_ranges",[])
#     st.session_state.spans_by_param   = jj.get("spans_by_param",{})
#     st.session_state.param_choice     = None
#     st.session_state.source_docx_path = jj.get("source_docx_path")
#     st.session_state.review_ready     = True
#     st.session_state["aoi_match_ranges"] = jj.get("aoi_match_ranges", {})
#     st.session_state.ui_mode          = "review"
#     return True

# def _render_recents_centerpane():
#     st.subheader("📄 Recents")
#     q = st.text_input("Filter by title…", "")
#     cols = st.columns([1, 4])
#     with cols[0]:
#         if st.button("← Back"):
#             st.session_state.ui_mode = "home"; st.rerun()

#     recs = _load_all_history()
#     ql = q.strip().lower()
#     if ql: recs = [r for r in recs if ql in (r.get("title","").lower())]
#     if not recs: st.caption("No history yet."); st.stop()

#     # Each card is a clickable <a class="rec-card" href="?open=<run_id>">
#     for rec in recs:
#         run_id = rec.get("run_id"); title = rec.get("title") or "(untitled)"
#         created_h = rec.get("created_at_human",""); overall = rec.get("overall_rating","")
#         st.markdown(f"""
#         <a class="rec-card" href="?open={run_id}">
#           <div class="rec-title">{title}</div>
#           <div class="rec-meta">{created_h}</div>
#           <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
#         </a>
#         """, unsafe_allow_html=True)

# # ---------- Sidebar ----------
# with st.sidebar:
#     if st.button("🆕 New review", use_container_width=True):
#         fp = st.session_state.get("flattened_docx_path")
#         if fp and os.path.exists(fp):
#             try: os.remove(fp)
#             except Exception: pass
#         for k in ["review_ready","script_text","base_stem","data","spans_by_param","param_choice",
#                   "source_docx_path","heading_ranges","flattened_docx_path","flatten_used"]:
#             st.session_state[k] = (
#                 False if k=="review_ready"
#                 else "" if k in ("script_text","base_stem")
#                 else {} if k=="spans_by_param"
#                 else [] if k=="heading_ranges"
#                 else None if k in ("source_docx_path","flattened_docx_path")
#                 else False if k=="flatten_used"
#                 else None
#             )
#         st.session_state.ui_mode = "home"
#         _clear_query_params()
#         st.rerun()

#     if st.button("📁 Recents", use_container_width=True):
#         st.session_state.ui_mode = "recents"
#         _clear_query_params()
#         st.rerun()

# # ---------- Input screen ----------
# def render_home():
#     st.subheader("🎬 Script Source")

#     (tab_upload,) = st.tabs(["Upload file"])
#     uploaded_file = None
#     uploaded_name = None
#     uploaded_key  = None

#     with tab_upload:
#         up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
#         if up is not None:
#             file_bytes = up.read()
#             suffix = os.path.splitext(up.name)[1].lower()
#             # Save to S3 (or local) under scripts/
#             uploaded_key = f"{SCRIPTS_DIR}/{up.name}"
#             save_bytes_key(uploaded_key, file_bytes)

#             # Also create a temp local copy for parsing
#             with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
#                 tmp.write(file_bytes)
#                 uploaded_file = tmp.name
#             uploaded_name = os.path.splitext(os.path.basename(up.name))[0] or "uploaded_script"

#     if st.button("🚀 Run Review", type="primary", use_container_width=True):
#         base_stem = "uploaded_script"
#         source_docx_path = None
#         heading_ranges = []

#         if uploaded_file:
#             base_stem = uploaded_name or "uploaded_script"
#             # Prefer key (so we can store it into history); for DOCX rendering we keep a tmp local
#             if uploaded_file.lower().endswith(".docx"):
#                 path_to_use = uploaded_file
#                 if _docx_contains_tables(path_to_use):
#                     flat = flatten_docx_tables_to_longtext(path_to_use)
#                     st.session_state.flattened_docx_path = flat
#                     st.session_state.flatten_used = True
#                     path_to_use = flat
#                 script_text, heading_ranges = build_docx_text_with_meta(path_to_use)

#                 # 🔑 IMPORTANT: set the rendering source to the *flattened* docx
#                 # (History saver will copy this file to Scriptmodel/outputs/_history/)
#                 source_docx_path = path_to_use
#             else:
#                 script_text = load_script_file(uploaded_file)
#                 source_docx_path = uploaded_key or uploaded_file  # keep reference to S3/local
#         else:
#             st.warning("Please upload a script first.")
#             st.stop()

#         if len(script_text.strip()) < 50:
#             st.error("Extracted text looks too short. Check your file extraction.")
#             st.stop()

#         with st.spinner("Running analysis…"):
#             try:
#                 review_text = run_review_multi(
#                     script_text=script_text,
#                     prompts_dir=PROMPTS_DIR,
#                     temperature=0.0
#                 )
#             finally:
#                 # Clean temp upload if we didn't keep it as source_docx_path local
#                 if uploaded_file and not (isinstance(source_docx_path, str) and os.path.exists(source_docx_path)):
#                     try:
#                         os.remove(uploaded_file)
#                     except Exception:
#                         pass

#         data = extract_review_json(review_text)
#         if not data:
#             st.error("JSON not detected in model output.")
#             st.stop()

#         st.session_state.script_text      = script_text
#         st.session_state.base_stem        = base_stem
#         st.session_state.data             = data
#         st.session_state.heading_ranges   = heading_ranges
#         st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
#         st.session_state.param_choice     = None
#         st.session_state.source_docx_path = source_docx_path
#         st.session_state.review_ready     = True
#         st.session_state.ui_mode          = "review"

#         _save_history_snapshot(
#             title=base_stem,
#             data=data,
#             script_text=script_text,
#             source_docx_path=source_docx_path,
#             heading_ranges=heading_ranges,
#             spans_by_param=st.session_state.spans_by_param,
#             aoi_match_ranges=st.session_state.get("aoi_match_ranges", {})
#         )

#         _clear_query_params()
#         st.rerun()

# # ---------- Results screen ----------
# def render_review():
#     script_text     = st.session_state.script_text
#     data            = st.session_state.data
#     spans_by_param  = st.session_state.spans_by_param
#     scores: Dict[str,int] = (data or {}).get("scores", {}) or {}
#     source_docx_path: Optional[str] = st.session_state.source_docx_path

#     # If our source_docx_path is an S3 key/url, ensure we have a local copy for rendering
#     # AFTER (prefer the flattened copy if we have it in-session)
#     docx_local: Optional[str] = None
#     preferred = st.session_state.get("flattened_docx_path") if st.session_state.get("flatten_used") else source_docx_path
#     if preferred:
#         if not os.path.exists(preferred):
#             docx_local = ensure_local_copy(preferred)
#         else:
#             docx_local = preferred

#     left, center, right = st.columns([1.1, 2.7, 1.4], gap="large")

#     with left:
#         st.subheader("Final score")
#         ordered = [p for p in PARAM_ORDER if p in scores]
#         df = pd.DataFrame({"Parameter": ordered, "Score (1–10)": [scores.get(p, "") for p in ordered]})
#         st.dataframe(df, hide_index=True, use_container_width=True)
#         st.markdown(f'**Overall:** {data.get("overall_rating","—")}/10')
#         st.divider()

#         strengths = (data or {}).get("strengths") or []
#         if not strengths:
#             per = (data or {}).get("per_parameter", {}) or {}
#             best = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
#             for name, sc in best:
#                 if sc >= 8 and name in per:
#                     exp = _sanitize_editor_text((per[name] or {}).get("explanation", "") or "")
#                     first = re.split(r"(?<=[.!?])\s+", exp.strip())[0] if exp else f"Consistently strong {name.lower()}."
#                     strengths.append(f"{name}: {first}")
#                 if len(strengths) >= 3: break

#         def _bullets(title: str, items):
#             st.markdown(f"**{title}**")
#             for s in (items or []):
#                 if isinstance(s, str) and s.strip(): st.write("• " + _sanitize_editor_text(s))
#             if not items: st.write("• —")

#         _bullets("Strengths", strengths)
#         _bullets("Weaknesses", data.get("weaknesses"))
#         _bullets("Suggestions", data.get("suggestions"))
#         _bullets("Drop-off Risks", data.get("drop_off_risks"))
#         st.markdown("**Viral Quotient**"); st.write(_sanitize_editor_text(data.get("viral_quotient","—")))

#     with right:
#         st.subheader("Parameters")
#         st.markdown('<div class="param-row">', unsafe_allow_html=True)
#         for p in [p for p in PARAM_ORDER if p in scores]:
#             if st.button(p, key=f"chip_{p}", help="Show inline AOI highlights for this parameter"):
#                 st.session_state.param_choice = p
#         st.markdown('</div>', unsafe_allow_html=True)

#         sel = st.session_state.param_choice
#         if sel:
#             blk = (data.get("per_parameter", {}) or {}).get(sel, {}) or {}
#             st.markdown(f"**{sel} — Score:** {scores.get(sel,'—')}/10")

#             if blk.get("explanation"):
#                 st.markdown("**Why this score**"); st.write(_sanitize_editor_text(blk["explanation"]))
#             if blk.get("weakness") and blk["weakness"] != "Not present":
#                 st.markdown("**Weakness**"); st.write(_sanitize_editor_text(blk["weakness"]))
#             if blk.get("suggestion") and blk["suggestion"] != "Not present":
#                 st.markdown("**Suggestion**"); st.write(_sanitize_editor_text(blk["suggestion"]))

#             aoi = blk.get("areas_of_improvement") or []
#             if aoi:
#                 st.markdown("**Areas of Improvement**")
#                 for i, item in enumerate(aoi, 1):
#                     popover_fn = getattr(st, "popover", None)
#                     aid = f"{sel.replace(' ','_')}-AOI-{i}"
#                     s_e_map = st.session_state.get("aoi_match_ranges", {})
#                     if aid in s_e_map:
#                         s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                         line = (matched_line if len(matched_line) <= 320 else matched_line[:300] + "…")
#                     else:
#                         line = _sanitize_editor_text(item.get('quote_verbatim',''))
#                     issue = _sanitize_editor_text(item.get('issue',''))
#                     fix   = _sanitize_editor_text(item.get('fix',''))
#                     why   = _sanitize_editor_text(item.get('why_this_helps',''))
#                     label = f"Issue {i}"
#                     if callable(popover_fn):
#                         with popover_fn(label):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)
#                     else:
#                         with st.expander(label, expanded=False):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)

#             if blk.get("summary"):
#                 st.markdown("**Summary**"); st.write(_sanitize_editor_text(blk["summary"]))

#     with center:
#         st.subheader("Script with inline highlights")
#         # choose spans for selected parameter (or all if None)
#         spans = st.session_state.spans_by_param.get(st.session_state.param_choice, []) if st.session_state.param_choice else []

#         # Build AOI payload for popup: { aid: {line, issue, fix, why} }
#         aoi_payload: Dict[str, Dict[str, str]] = {}
#         data_per = (data or {}).get("per_parameter") or {}
#         s_e_map = st.session_state.get("aoi_match_ranges", {})
#         sel = st.session_state.param_choice

#         def _mk_line(aid: str, fallback_q: str = "") -> str:
#             if aid in s_e_map:
#                 s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                 return matched_line if len(matched_line) <= 320 else matched_line[:300] + "…"
#             return _sanitize_editor_text(fallback_q or "")

#         def _collect(param_name: str):
#             blk = (data_per.get(param_name) or {})
#             for i, item in enumerate(blk.get("areas_of_improvement") or [], 1):
#                 aid = f"{param_name.replace(' ','_')}-AOI-{i}"
#                 aoi_payload[aid] = {
#                     "line": _mk_line(aid, (item or {}).get("quote_verbatim","")),
#                     "issue": _sanitize_editor_text((item or {}).get("issue","")),
#                     "fix": _sanitize_editor_text((item or {}).get("fix","")),
#                     "why": _sanitize_editor_text((item or {}).get("why_this_helps","")),
#                 }

#         if sel: _collect(sel)
#         else:
#             for pn in [p for p in PARAM_ORDER if p in data_per]:
#                 _collect(pn)

#         payload_json = json.dumps(aoi_payload, ensure_ascii=False)

#         frame_theme_css = """
#         <style>
#           :root{
#             --m7-surface: #eef2f7;
#             --m7-on-surface: #0f172a;
#             --m7-border: rgba(15,23,42,.14);
#           }
#           @media (prefers-color-scheme: dark){
#             :root{
#               --m7-surface: #2f333a;
#               --m7-on-surface: #ffffff;
#               --m7-border: rgba(255,255,255,.18);
#             }
#             body { background: transparent !important; }
#           }

#           .docxwrap{
#             background: var(--m7-surface);
#             color: var(--m7-on-surface);
#             border: 1px solid var(--m7-border);
#             border-radius: 12px;
#             padding: 16px 14px 18px;
#           }
#           .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#           .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); }
#         </style>
#         """

#         tooltip_css = """
#         <style>
#         .aoi-pop {
#           position: absolute; max-width: 520px; min-width: 320px;
#           background: var(--m7-surface); border: 1px solid var(--m7-border); border-radius: 10px;
#           box-shadow: 0 10px 25px rgba(0,0,0,.12);
#           padding: 12px 14px; z-index: 9999; transform: translateY(-8px);
#           color: var(--m7-on-surface);
#         }
#         .aoi-pop h4 { margin: 0 0 .35rem 0; font-size: .95rem; }
#         .aoi-pop p  { margin: .15rem 0; line-height: 1.5; }
#         .aoi-pop .muted { opacity:.85; font-size:.85rem; }
#         .aoi-arrow {
#           position:absolute; left:50%; transform:translateX(-50%);
#           bottom:-7px; width:0;height:0;border-left:7px solid transparent;
#           border-right:7px solid transparent;border-top:7px solid var(--m7-border);
#         }
#         .aoi-arrow::after{
#           content:""; position:absolute; left:-6px; top:-7px; width:0;height:0;
#           border-left:6px solid transparent;border-right:6px solid transparent;border-top:6px solid var(--m7-surface);
#         }
#         </style>
#         """

#         # Select rendering source (DOCX with highlights if we have a local path)
#         if docx_local and os.path.splitext(docx_local)[1].lower() == ".docx":
#             html_core = render_docx_html_with_highlights(
#                 docx_local,
#                 merge_overlaps_and_adjacent(script_text, spans)
#             )
#         else:
#             from html import escape as _esc
#             orig = script_text
#             spans2 = [s for s in merge_overlaps_and_adjacent(orig, spans) if s[0] < s[1]]
#             spans2.sort(key=lambda x: x[0])
#             cur = 0; buf: List[str] = []
#             for s,e,c,aid in spans2:
#                 if s > cur: buf.append(_esc(orig[cur:s]))
#                 buf.append(
#                     f'<mark class="aoi-mark" data-aid="{aid}" '
#                     f'style="background:{c}33;border:1px solid {c};border-radius:3px;padding:0 2px;">'
#                     f'{_esc(orig[s:e])}</mark>'
#                 )
#                 cur = e
#             if cur < len(orig): buf.append(_esc(orig[cur:]))
#             html_core = (
#                 '<div class="docxwrap"><p style="white-space:pre-wrap; '
#                 'line-height:1.7; font-family:ui-serif, Georgia, Times New Roman, serif;">'
#                 + "".join(buf) +
#                 '</p></div>'
#             )

#         # Popup + autosize JS shell
#         html_shell = """
# %%FRAME_THEME_CSS%%
# %%TOOLTIP_CSS%%
# <div id="m7-doc">%%HTML_CORE%%</div>
# <div id="aoi-pop" class="aoi-pop" style="display:none;">
#   <div id="aoi-pop-content"></div>
#   <div class="aoi-arrow"></div>
# </div>
# <script>
# (function(){
#   const AOI = __PAYLOAD__;
#   const wrap = document.getElementById('m7-doc');
#   const pop  = document.getElementById('aoi-pop');
#   const body = document.getElementById('aoi-pop-content');

#   function resizeIframe() {
#     try {
#       const h = Math.max(
#         document.documentElement.scrollHeight,
#         document.body.scrollHeight
#       );
#       if (window.frameElement) {
#         window.frameElement.style.height = (h + 20) + 'px';
#         window.frameElement.style.width  = '100%';
#       }
#     } catch(e) {}
#   }
#   window.addEventListener('load', resizeIframe);
#   window.addEventListener('resize', resizeIframe);

#   function hide(){ pop.style.display='none'; }
#   function showFor(mark){
#     const aid = mark.getAttribute('data-aid');
#     const d = AOI[aid]; if(!d) return;
#     body.innerHTML =
#       (d.line  ? '<p><strong>Line:</strong> '  + d.line  + '</p>' : '') +
#       (d.issue ? '<p><strong>Issue:</strong> ' + d.issue + '</p>' : '') +
#       (d.fix   ? '<p><strong>Fix:</strong> '   + d.fix   + '</p>' : '') +
#       (d.why   ? '<p class="muted">'           + d.why   + '</p>' : '');
#     pop.style.display = 'block';

#     const r = mark.getBoundingClientRect();
#     const scY = window.scrollY || document.documentElement.scrollTop;
#     const scX = window.scrollX || document.documentElement.scrollLeft;
#     let top  = r.top + scY - pop.offsetHeight - 10;
#     let left = r.left + scX + r.width/2 - pop.offsetWidth/2;
#     if (top < 8) top = r.bottom + scY + 10;
#     if (left < 8) left = 8;
#     pop.style.top  = top + 'px';
#     pop.style.left = left + 'px';

#     resizeIframe();
#   }

#   wrap.addEventListener('click', (e) => {
#     const m = e.target.closest('.aoi-mark');
#     if(!m){ hide(); return; }
#     if(pop.style.display === 'block'){ hide(); }
#     showFor(m);
#     e.stopPropagation();
#   });

#   document.addEventListener('click', (e) => {
#     if(!e.target.closest('.aoi-pop') && !e.target.closest('.aoi-mark')) hide();
#   });
# })();
# </script>
# """
#         html_shell = (
#             html_shell
#             .replace("%%FRAME_THEME_CSS%%", frame_theme_css)
#             .replace("%%TOOLTIP_CSS%%", tooltip_css)
#             .replace("%%HTML_CORE%%", html_core)
#             .replace("__PAYLOAD__", payload_json)
#         )

#         components.html(html_shell, height=400, scrolling=False)

# # ---------- Router & query param open ----------
# _open_qp = _get_query_param("open")
# if _open_qp and _open_history_run_by_id(_open_qp):
#     _clear_query_params()  # avoid re-opening on subsequent reruns

# mode = st.session_state.ui_mode
# if mode == "recents":
#     _render_recents_centerpane()
# elif mode == "review" and st.session_state.review_ready:
#     render_review()
# else:
#     render_home()







##########################################################






# import os, re, glob, json, tempfile, difflib, uuid, datetime, shutil
# from pathlib import Path
# from typing import Dict, Any, List, Tuple, Optional

# import streamlit as st
# import pandas as pd
# import streamlit.components.v1 as components  # for inline HTML/JS popup

# from utils1 import extract_review_json, PARAM_ORDER, load_script_file
# from review_engine_multi import run_review_multi

# # ---- DOCX rendering imports ----
# from docx import Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # =========================
# # RunPod S3 (inline helpers) — FIXED
# # =========================
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError

# # Accept both AWS_* and RUNPOD_* style envs / st.secrets
# def _get_env(key: str, default: str = "") -> str:
#     v = os.getenv(key, "")
#     if v: 
#         return v.strip()
#     try:
#         v2 = st.secrets.get(key)
#         if isinstance(v2, str):
#             return v2.strip()
#     except Exception:
#         pass
#     return (default or "").strip()

# # Primary config
# _RP_ENDPOINT = _get_env("RUNPOD_S3_ENDPOINT")
# _RP_BUCKET   = _get_env("RUNPOD_S3_BUCKET")
# _RP_REGION   = _get_env("RUNPOD_S3_REGION") or _get_env("AWS_DEFAULT_REGION") or ""

# # Credentials: prefer AWS_* if present; else accept RUNPOD_* fallbacks
# _AK = _get_env("AWS_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY")
# _SK = _get_env("AWS_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_KEY")
# _ST = _get_env("AWS_SESSION_TOKEN")  # optional

# # Options
# _FORCE_PATH = (_get_env("RUNPOD_S3_FORCE_PATH_STYLE") or "true").lower() in {"1","true","yes"}
# _USE_SSL    = (_get_env("RUNPOD_S3_USE_SSL") or "true").lower() in {"1","true","yes"}
# _VERIFY_SSL = (_get_env("RUNPOD_S3_VERIFY_SSL") or "true").lower() in {"1","true","yes"}

# def _s3_enabled() -> bool:
#     return bool(_RP_ENDPOINT and _RP_BUCKET and _AK and _SK)

# @st.cache_resource(show_spinner=False)
# def _s3_client():
#     if not _s3_enabled():
#         return None
#     session_kwargs = dict(
#         aws_access_key_id=_AK,
#         aws_secret_access_key=_SK,
#     )
#     if _ST:
#         session_kwargs["aws_session_token"] = _ST

#     # s3v4 + path style are common requirements for S3-compatible services
#     cfg = Config(
#         signature_version="s3v4",
#         s3={"addressing_style": "path" if _FORCE_PATH else "auto"},
#         retries={"max_attempts": 3, "mode": "standard"}
#     )
#     return boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         region_name=_RP_REGION or None,
#         use_ssl=_USE_SSL,
#         verify=_VERIFY_SSL,
#         config=cfg,
#         **session_kwargs,
#     )

# def save_text_key(key: str, text: str) -> str:
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "w", encoding="utf-8") as f:
#             f.write(text)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=text.encode("utf-8"))
#     return f"s3://{_RP_BUCKET}/{key}"

# def save_bytes_key(key: str, data: bytes) -> str:
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "wb") as f:
#             f.write(data)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=data)
#     return f"s3://{_RP_BUCKET}/{key}"

# def read_text_key(key: str, default: str = "") -> str:
#     if not _s3_enabled():
#         try:
#             with open(key, "r", encoding="utf-8") as f:
#                 return f.read()
#         except Exception:
#             return default
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read().decode("utf-8", errors="ignore")
#     except Exception:
#         return default

# def read_bytes_key(key: str) -> Optional[bytes]:
#     if not _s3_enabled():
#         try:
#             with open(key, "rb") as f:
#                 return f.read()
#         except Exception:
#             return None
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read()
#     except Exception:
#         return None

# def list_prefix(prefix: str) -> List[str]:
#     """
#     List object keys under prefix (or local dir paths if not S3).
#     In S3 mode we always return KEYS (not URLs).
#     """
#     if not _s3_enabled():
#         base = prefix if os.path.isdir(prefix) else os.path.dirname(prefix)
#         try:
#             return [os.path.join(base, p) for p in os.listdir(base) if p.endswith(".json")]
#         except Exception:
#             return []

#     out: List[str] = []
#     token = None
#     # Normalize to "dir/" prefix for S3 listing
#     s3_prefix = prefix.rstrip("/") + "/"
#     try:
#         while True:
#             kwargs = {"Bucket": _RP_BUCKET, "Prefix": s3_prefix}
#             if token:
#                 kwargs["ContinuationToken"] = token
#             resp = _s3_client().list_objects_v2(**kwargs)
#             for c in resp.get("Contents", []):
#                 k = c.get("Key", "")
#                 if k.endswith(".json"):
#                     out.append(k)
#             token = resp.get("NextContinuationToken")
#             if not token:
#                 break
#     except (ClientError, EndpointConnectionError, NoCredentialsError):
#         return []
#     return out

# def presigned_url(key: str, expires: int = 3600) -> Optional[str]:
#     if not _s3_enabled():
#         return None
#     try:
#         return _s3_client().generate_presigned_url(
#             "get_object",
#             Params={"Bucket": _RP_BUCKET, "Key": key},
#             ExpiresIn=expires
#         )
#     except ClientError:
#         return None

# def ensure_local_copy(key_or_path: str) -> Optional[str]:
#     """
#     For DOCX/PDF parsing we need a real filesystem path.
#     If S3 mode, download to a temp file and return that path.
#     """
#     if not _s3_enabled():
#         return key_or_path if os.path.exists(key_or_path) else None

#     key = key_or_path
#     if key.startswith("s3://"):
#         # s3://bucket/path/to/file -> path/to/file
#         parts = key.split("/", 3)
#         key = parts[3] if len(parts) >= 4 else ""
#     data = read_bytes_key(key)
#     if data is None:
#         return None
#     fd, tmp = tempfile.mkstemp(suffix=os.path.splitext(key)[1] or "")
#     os.close(fd)
#     with open(tmp, "wb") as f:
#         f.write(data)
#     return tmp

# def _s3_health_summary() -> dict:
#     """
#     Optional tiny health read you can print if needed.
#     Returns a dict; safe to ignore in production.
#     """
#     info = {
#         "enabled": _s3_enabled(),
#         "endpoint": _RP_ENDPOINT,
#         "bucket": _RP_BUCKET,
#         "region": _RP_REGION,
#         "has_keys": bool(_AK and _SK),
#     }
#     if not _s3_enabled():
#         info["status"] = "local-mode"
#         return info
#     try:
#         # Attempt a very cheap list; no exceptions => reachable
#         _ = _s3_client().list_objects_v2(Bucket=_RP_BUCKET, Prefix=(f"{OUTPUT_DIR}/_history/").rstrip("/") + "/",
#                                          MaxKeys=1)
#         info["status"] = "ok"
#     except Exception as e:
#         info["status"] = f"error: {getattr(e, 'response', {}).get('Error', {}).get('Code', str(e))}"
#     return info


# # ---------- Folders ----------
# # SCRIPTS_DIR = "scripts"
# # PROMPTS_DIR = "prompts"
# # OUTPUT_DIR  = "outputs"
# # HISTORY_DIR = os.path.join(OUTPUT_DIR, "_history")
# # ---------- Folders (all under Scriptmodel/) ----------

# BASE_PREFIX = "Scriptmodel"

# SCRIPTS_DIR = f"{BASE_PREFIX}/scripts"
# PROMPTS_DIR = f"{BASE_PREFIX}/prompts"
# OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
# HISTORY_DIR = f"{OUTPUT_DIR}/_history"


# if not _s3_enabled():
#     for p in (SCRIPTS_DIR, PROMPTS_DIR, OUTPUT_DIR, HISTORY_DIR):
#         Path(p).mkdir(parents=True, exist_ok=True)


# # ---------- Colors ----------
# PARAM_COLORS: Dict[str, str] = {
#     "Suspense Building":              "#ff6b6b",
#     "Language/Tone":                  "#6b8cff",
#     "Intro + Main Hook/Cliffhanger":  "#ffb86b",
#     "Story Structure + Flow":         "#a78bfa",
#     "Pacing":                         "#f43f5e",
#     "Mini-Hooks (30–60s)":            "#eab308",
#     "Outro (Ending)":                 "#8b5cf6",
#     "Grammar & Spelling":             "#10b981",  # NEW: distinct highlight color for grammar/spelling
# }

# # ---------- Config ----------
# STRICT_MATCH_ONLY = False

# # ---------- App config ----------
# st.set_page_config(page_title="M7 — Grammarly UI", page_icon="🕵️", layout="wide")
# # ---------- Header patch & CSS ----------
# def render_app_title():
#     st.markdown(
#         '<h1 class="app-title">Viral Script Reviewer</h1>',
#         unsafe_allow_html=True
#     )
#     st.markdown("""
#     <style>
#     html { color-scheme: light dark; }

#     :root{
#       --m7-surface: #eef2f7;
#       --m7-on-surface: #0f172a;
#       --m7-border: rgba(15,23,42,.14);
#       --sep: #e5e7eb;
#     }
#     @media (prefers-color-scheme: dark){
#       :root{
#         --m7-surface: #2f333a;
#         --m7-on-surface: #ffffff;
#         --m7-border: rgba(255,255,255,.18);
#         --sep: #2a2f37;
#       }
#     }

#     .stApp .block-container { padding-top: 4.25rem !important; }
#     .app-title{
#       font-weight: 700; font-size: 2.1rem; line-height: 1.3;
#       margin: 0 0 1rem 0; padding-left: 40px !important; padding-top: .25rem !important;
#       white-space: normal; word-break: break-word; hyphens: auto; overflow: visible;
#       position: relative !important; z-index: 10 !important;
#     }
#     [data-testid="collapsedControl"] { z-index: 6 !important; }
#     header[data-testid="stHeader"], .stAppHeader { background: transparent !important; box-shadow: none !important; }
#     @media (min-width: 992px){ .app-title { padding-left: 0 !important; } }

#     div[data-testid="column"]:nth-of-type(1){position:relative;}
#     div[data-testid="column"]:nth-of-type(1)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     div[data-testid="column"]:nth-of-type(2){position:relative;}
#     div[data-testid="column"]:nth-of-type(2)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}

#     .m7-card{
#       background: var(--m7-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 14px 16px;
#       color: var(--m7-on-surface);
#     }
#     .m7-card, .m7-card * { color: var(--m7-on-surface) !important; }

#     .docxwrap{
#       background: var(--m7-surface);
#       color: var(--m7-on-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 16px 14px 18px;
#     }
#     .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#     .docxwrap .h1, .docxwrap .h2, .docxwrap .h3 { font-weight:700; margin:10px 0 6px; }
#     .docxwrap .h1 { font-size: 1.3rem; border-bottom: 2px solid currentColor; padding-bottom: 4px; }
#     .docxwrap .h2 { font-size: 1.15rem; border-bottom: 1px solid currentColor; padding-bottom: 3px; }
#     .docxwrap .h3 { font-size: 1.05rem; }
#     .docxwrap p { margin: 10px 0; line-height: 1.7; font-family: ui-serif, Georgia, "Times New Roman", serif; }
#     .docxwrap table { border-collapse: collapse; width: 100%; margin: 12px 0; }
#     .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); padding:8px; vertical-align:top; line-height:1.6; }
#     .docxwrap mark{ padding:0 2px; border-radius:3px; border:1px solid var(--m7-border); cursor: pointer; }

#     .rec-card{
#       display:block; text-decoration:none !important;
#       background: var(--m7-surface);
#       border:1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 14px 16px; margin: 10px 0 16px;
#       box-shadow: 0 1px 2px rgba(0,0,0,.06);
#       color: var(--m7-on-surface) !important;
#       transition: filter .1s ease, transform .02s ease;
#     }
#     .rec-card:hover{ filter: brightness(1.02); }
#     .rec-card:active{ transform: translateY(1px); }
#     .rec-card, .rec-card * { color: var(--m7-on-surface) !important; }
#     .rec-title{font-weight:600; margin-bottom:.25rem;}
#     .rec-meta{opacity:.85 !important; font-size:12.5px; margin-bottom:.4rem;}
#     .rec-row{display:flex; align-items:center; justify-content:space-between; gap:12px;}

#     .stTextInput>div>div,
#     .stTextArea>div>div,
#     .stNumberInput>div>div,
#     .stDateInput>div>div,
#     .stTimeInput>div>div,
#     .stFileUploader>div,
#     div[data-baseweb="select"]{
#       background: var(--m7-surface) !important;
#       border: 1px solid var(--m7-border) !important;
#       border-radius: 10px !important;
#       color: var(--m7-on-surface) !important;
#     }
#     .stTextInput input,
#     .stTextArea textarea,
#     .stNumberInput input,
#     .stDateInput input,
#     .stTimeInput input,
#     .stFileUploader div,
#     div[data-baseweb="select"] *{
#       color: var(--m7-on-surface) !important;
#     }
#     .stTextInput input::placeholder,
#     .stTextArea textarea::placeholder{ color: rgba(16,24,39,.55) !important; }
#     @media (prefers-color-scheme: dark){
#       .stTextInput input::placeholder,
#       .stTextArea textarea::placeholder{ color: rgba(255,255,255,.75) !important; }
#     }
#     div[data-testid="stFileUploaderDropzone"] label span { color: var(--m7-on-surface) !important; opacity: 1 !important; }
#     div[data-testid="stFileUploaderDropzone"] label { color: var(--m7-on-surface) !important; }

#     .stMarkdown pre,
#     pre[class*="language-"],
#     .stCodeBlock{
#       background: var(--m7-surface) !important;
#       color: var(--m7-on-surface) !important;
#       border: 1px solid var(--m7-border) !important;
#       border-radius: 12px !important;
#       padding: 12px 14px !important;
#       overflow:auto;
#     }
#     .stMarkdown pre code{ background: transparent !important; color: inherit !important; }

#     div[data-testid="stDataFrame"]{
#       background: var(--m7-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 6px 8px;
#       color: var(--m7-on-surface);
#     }
#     div[data-testid="stDataFrame"] * { color: var(--m7-on-surface) !important; }

#     </style>
#     """, unsafe_allow_html=True)

# render_app_title()

# # ---------- Session ----------
# for key, default in [
#     ("review_ready", False),
#     ("script_text", ""),
#     ("base_stem", ""),
#     ("data", None),
#     ("spans_by_param", {}),
#     ("param_choice", None),
#     ("source_docx_path", None),
#     ("heading_ranges", []),
#     ("flattened_docx_path", None),
#     ("flatten_used", False),
#     ("ui_mode", "home"),
# ]:
#     st.session_state.setdefault(key, default)

# # ---------- helpers for query params (compat across Streamlit versions) ----------
# def _get_query_param(key: str) -> Optional[str]:
#     val = None
#     try:
#         val = st.query_params.get(key)
#     except Exception:
#         q = st.experimental_get_query_params()
#         v = q.get(key)
#         if isinstance(v, list): val = v[0] if v else None
#         else: val = v
#     return val

# def _clear_query_params():
#     try:
#         st.query_params.clear()
#     except Exception:
#         st.experimental_set_query_params()

# # ---------- Sanitizer ----------
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )
# def _sanitize_editor_text(s: Optional[str]) -> str:
#     if not s: return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•)\s*', '', t, flags=re.M)
#     t = re.sub(r'^\s*[-*]\s+', '• ', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# # ---------- DOCX traversal ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# # ---------- Auto-flatten ----------
# def _docx_contains_tables(path: str) -> bool:
#     doc = Document(path)
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Table):
#             return True
#     return False

# def _copy_paragraph(dest_doc: Document, src_para: Paragraph):
#     p = dest_doc.add_paragraph()
#     try:
#         if src_para.style and src_para.style.name:
#             p.style = src_para.style.name
#     except Exception:
#         pass
#     for run in src_para.runs:
#         r = p.add_run(run.text or "")
#         r.bold = run.bold
#         r.italic = run.italic
#         r.underline = run.underline
#     return p

# def flatten_docx_tables_to_longtext(source_path: str) -> str:
#     src = Document(source_path)
#     new = Document()
#     for blk in _iter_docx_blocks(src):
#         if isinstance(blk, Paragraph):
#             _copy_paragraph(new, blk)
#         else:
#             seen_tc_ids = set()
#             for row in blk.rows:
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     if tc_id in seen_tc_ids:
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     for p in cell.paragraphs:
#                         _copy_paragraph(new, p)
#                 new.add_paragraph("")
#             new.add_paragraph("")
#     fd, tmp_path = tempfile.mkstemp(suffix=".docx"); os.close(fd); new.save(tmp_path)
#     return tmp_path

# # ---------- Build plain text + heading ranges ----------
# def build_docx_text_with_meta(docx_path: str) -> Tuple[str, List[Tuple[int,int]]]:
#     doc = Document(docx_path)
#     out: List[str] = []
#     heading_ranges: List[Tuple[int,int]] = []
#     current_offset = 0

#     def _append_and_advance(s: str):
#         nonlocal current_offset
#         out.append(s); current_offset += len(s)

#     seen_tc_ids: set = set()
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             para_text = "".join(run.text or "" for run in blk.runs)
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             if sty.startswith("heading"):
#                 start = current_offset; end = start + len(para_text)
#                 heading_ranges.append((start, end))
#             _append_and_advance(para_text); _append_and_advance("\n")
#         else:
#             for row in blk.rows:
#                 row_cell_tcs = []
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     row_cell_tcs.append((tc_id, cell))
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     if tc_id in seen_tc_ids:
#                         if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     cell_text_parts: List[str] = []
#                     for i, p in enumerate(cell.paragraphs):
#                         t = "".join(r.text or "" for r in p.runs)
#                         sty = (p.style.name or "").lower() if p.style else ""
#                         if sty.startswith("heading"):
#                             hs = current_offset + sum(len(x) for x in cell_text_parts)
#                             he = hs + len(t)
#                             heading_ranges.append((hs, he))
#                         cell_text_parts.append(t)
#                         if i != len(cell.paragraphs) - 1:
#                             cell_text_parts.append("\n")
#                     cell_text = "".join(cell_text_parts)
#                     _append_and_advance(cell_text)
#                     if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                 _append_and_advance("\n")
#             _append_and_advance("\n")

#     return "".join(out), heading_ranges

# def _wrap_inline(safe_text: str, run) -> str:
#     out = safe_text
#     if getattr(run, "underline", False): out = f"<u>{out}</u>"
#     if getattr(run, "italic", False): out = f"<em>{out}</em>"
#     if getattr(run, "bold", False): out = f"<strong>{out}</strong>"
#     return out

# # ---------- Invisible/bridge chars ----------
# _BRIDGE_CHARS = set("\u200b\u200c\u200d\u2060\ufeff\xa0\u00ad")

# # ---------- DOCX -> HTML with highlights (includes data-aid) ----------
# def render_docx_html_with_highlights(docx_path: str,
#                                      highlight_spans: List[Tuple[int,int,str,str]]) -> str:
#     """
#     highlight_spans: list of (start, end, color, aid)
#     """
#     doc = Document(docx_path)
#     spans = [s for s in highlight_spans if s[0] < s[1]]
#     spans.sort(key=lambda x: x[0])

#     cur_span = 0
#     current_offset = 0

#     def esc(s: str) -> str:
#         return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

#     def open_mark_if_needed(html_parts, mark_state, color, end, aid):
#         if not mark_state["open"]:
#             html_parts.append(
#                 f'<mark class="aoi-mark" data-aid="{aid}" '
#                 f'style="background:{color}33;border:1px solid {color};border-radius:3px;padding:0 2px;">'
#             )
#             mark_state.update(open=True, end=end, color=color, aid=aid)

#     def close_mark_if_open(html_parts, mark_state):
#         if mark_state["open"]:
#             html_parts.append('</mark>')
#             mark_state.update(open=False, end=None, color=None, aid=None)

#     def emit_run_text(run_text: str, run, html_parts: List[str], mark_state: Dict[str, Any]):
#         nonlocal cur_span, current_offset
#         t = run_text or ""; i = 0
#         while i < len(t):
#             next_start, next_end, color, next_aid = None, None, None, None
#             if cur_span < len(spans):
#                 next_start, next_end, color, next_aid = spans[cur_span]

#             if not mark_state["open"]:
#                 if cur_span >= len(spans) or current_offset + (len(t) - i) <= next_start:
#                     chunk = t[i:]; html_parts.append(_wrap_inline(esc(chunk), run)); current_offset += len(chunk); break
#                 if current_offset < next_start:
#                     take = next_start - current_offset
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take; continue
#                 open_mark_if_needed(html_parts, mark_state, color, next_end, next_aid)
#             else:
#                 take = min(mark_state["end"] - current_offset, len(t) - i)
#                 if take > 0:
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take
#                 if current_offset >= mark_state["end"]:
#                     close_mark_if_open(html_parts, mark_state)
#                     cur_span += 1

#     html: List[str] = ['<div class="docxwrap">']
#     seen_tc_ids: set = set()

#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             open_tag = '<div class="h1">' if sty.startswith("heading 1") else \
#                        '<div class="h2">' if sty.startswith("heading 2") else \
#                        '<div class="h3">' if sty.startswith("heading 3") else "<p>"
#             close_tag = "</div>" if sty.startswith("heading") else "</p>"
#             html.append(open_tag)
#             for run in blk.runs:
#                 emit_run_text(run.text or "", run, html, mark_state)
#             close_mark_if_open(html, mark_state)
#             html.append(close_tag)
#             current_offset += 1  # '\n'

#         else:
#             html.append("<table>")
#             for row in blk.rows:
#                 html.append("<tr>")
#                 row_cell_tcs = [(id(cell._tc), cell) for cell in row.cells]
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     html.append("<td>")
#                     if tc_id not in seen_tc_ids:
#                         seen_tc_ids.add(tc_id)
#                         for p_idx, p in enumerate(cell.paragraphs):
#                             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#                             html.append("<div>")
#                             for run in p.runs:
#                                 emit_run_text(run.text or "", run, html, mark_state)
#                             close_mark_if_open(html, mark_state)
#                             html.append("</div>")
#                             if p_idx != len(cell.paragraphs) - 1:
#                                 current_offset += 1
#                     html.append("</td>")
#                     if idx != len(row_cell_tcs) - 1: current_offset += 1  # '\t'
#                 html.append("</tr>"); current_offset += 1   # row '\n'
#             html.append("</table>"); current_offset += 1     # extra '\n'

#     html.append("</div>")
#     return "".join(html)

# # ---------- Matching / span utilities ----------
# def _normalize_keep_len(s: str) -> str:
#     trans = {
#         "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
#         "\u2013": "-", "\u2014": "-",
#         "\xa0": " ",
#         "\u200b": " ", "\u200c": " ", "\u200d": " ", "\u2060": " ",
#         "\ufeff": " ", "\u00ad": " ",
#     }
#     return (s or "").translate(str.maketrans(trans))

# def _tokenize(s: str) -> List[str]: return re.findall(r"\w+", (s or "").lower())

# def _iter_sentences_with_spans(text: str) -> List[Tuple[int,int,str]]:
#     spans = []
#     for m in re.finditer(r'[^.!?]+[.!?]+|\Z', text, flags=re.S):
#         s, e = m.start(), m.end()
#         seg = text[s:e]
#         if seg.strip(): spans.append((s, e, seg))
#     return spans

# def _squash_ws(s: str) -> str: return re.sub(r"\s+", " ", s or "").strip()

# def _clean_quote_for_match(q: str) -> str:
#     if not q: return ""
#     q = _normalize_keep_len(q).strip()
#     q = re.sub(r'^[\'"“”‘’\[\(\{<…\-\–\—\s]+', '', q)
#     q = re.sub(r'[\'"“”‘’\]\)\}>…\-\–\—\s]+$', '', q)
#     return _squash_ws(q)

# def _snap_and_bridge_to_word(text: str, start: int, end: int, max_bridge: int = 2) -> Tuple[int,int]:
#     n = len(text); s, e = max(0,start), max(start,end)
#     def _is_inv(ch: str) -> bool: return ch in _BRIDGE_CHARS

#     while s > 0:
#         prev = text[s-1]; cur = text[s] if s < n else ""
#         if prev.isalnum() and cur.isalnum(): s -= 1; continue
#         j = s; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1; 
#         if brid and (s-1)>=0 and text[s-1].isalnum() and (j<n and text[j].isalnum()): s -= 1; continue
#         break

#     while e < n:
#         prev = text[e-1] if e>0 else ""; nxt = text[e]
#         if prev.isalnum() and nxt.isalnum(): e += 1; continue
#         j = e; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (e-1)>=0 and text[e-1].isalnum() and (j<n and text[j].isalnum()): e = j + 1; continue
#         break

#     while e < n and text[e] in ',"”’\')]}': e += 1
#     return s, e

# def _heal_split_word_left(text: str, start: int) -> int:
#     i = start
#     if i <= 1 or i >= len(text): return start
#     if text[i-1] != " ": return start
#     j = i - 2
#     while j >= 0 and text[j].isalpha(): j -= 1
#     prev_token = text[j+1:i-1]
#     if len(prev_token) == 1: return i - 2
#     return start

# def _overlaps_any(s: int, e: int, ranges: List[Tuple[int,int]]) -> bool:
#     for rs, re_ in ranges:
#         if e > rs and s < re_: return True
#     return False

# def _fuzzy_window_span(tl: str, nl: str, start: int, w: int) -> Tuple[float, Optional[Tuple[int,int]]]:
#     window = tl[start:start+w]
#     sm = difflib.SequenceMatcher(a=nl, b=window)
#     blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
#     if not blocks: return 0.0, None
#     coverage = sum(b.size for b in blocks) / max(1, len(nl))
#     first_b = min(blocks, key=lambda b: b.b); last_b = max(blocks, key=lambda b: b.b + b.size)
#     s = start + first_b.b; e = start + last_b.b + last_b.size
#     return coverage, (s, e)

# def find_span_smart(text: str, needle: str) -> Optional[Tuple[int,int]]:
#     if not text or not needle: return None
#     t_orig = text; t_norm = _normalize_keep_len(text); n_norm = _clean_quote_for_match(needle)
#     if not n_norm: return None
#     tl = t_norm.lower(); nl = n_norm.lower()

#     i = tl.find(nl)
#     if i != -1:
#         s, e = _snap_and_bridge_to_word(t_orig, i, i + len(nl)); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     m = re.search(re.escape(nl).replace(r"\ ", r"\s+"), tl, flags=re.IGNORECASE)
#     if m:
#         s, e = _snap_and_bridge_to_word(t_orig, m.start(), m.end()); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     if not STRICT_MATCH_ONLY and len(nl) >= 12:
#         w = max(60, min(240, len(nl) + 80))
#         best_cov, best_span = 0.0, None
#         step = max(1, w // 2)
#         for start in range(0, max(1, len(tl) - w + 1), step):
#             cov, se = _fuzzy_window_span(tl, nl, start, w)
#             if cov > best_cov: best_cov, best_span = cov, se
#         if best_span and best_cov >= 0.65:
#             s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
#             if s > 0 and t_orig[s-1:s+1].lower() in {" the", " a", " an"}: s -= 1
#             s = _heal_split_word_left(t_orig, s)
#             return (s, e)

#     if not STRICT_MATCH_ONLY:
#         keys = [w for w in _tokenize(nl) if len(w) >= 4][:8]
#         if len(keys) >= 2:
#             kset = set(keys)
#             best_score, best_span = 0.0, None
#             for s, e, seg in _iter_sentences_with_spans(t_norm):
#                 toks = set(_tokenize(seg)); ov = len(kset & toks)
#                 if ov == 0: continue
#                 score = ov / max(2, len(kset)); length_pen = min(1.0, 120 / max(20, e - s)); score *= (0.6 + 0.4 * length_pen)
#                 if score > best_score: best_score, best_span = score, (s, min(e, s + 400))
#             if best_span and best_score >= 0.35:
#                 s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1]); s = _heal_split_word_left(t_orig, s)
#                 return (s, e)
#     return None

# def merge_overlaps(spans: List[Tuple[int,int,str,str]]) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans.sort(key=lambda x: x[0]); out = [spans[0]]
#     for s,e,c,aid in spans[1:]:
#         ps,pe,pc,paid = out[-1]
#         if s <= pe and pc == c and e > pe: out[-1] = (ps, e, pc, paid)
#         else: out.append((s,e,c,aid))
#     return out

# _PUNCT_WS = set(" \t\r\n,.;:!?)('\"[]{}-—–…") | _BRIDGE_CHARS

# def merge_overlaps_and_adjacent(base_text: str,
#                                 spans: List[Tuple[int,int,str,str]],
#                                 max_gap: int = 2) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans = sorted(spans, key=lambda x: x[0]); out = [spans[0]]
#     for s, e, c, aid in spans[1:]:
#         ps, pe, pc, paid = out[-1]
#         if c == pc and s <= pe: out[-1] = (ps, max(pe, e), pc, paid); continue
#         if c == pc and s - pe <= max_gap:
#             gap = base_text[max(0, pe):max(0, s)]
#             if all((ch in _PUNCT_WS) for ch in gap): out[-1] = (ps, e, pc, paid); continue
#         out.append((s, e, c, aid))
#     return out

# # ---------- Heading filters ----------
# def _is_heading_like(q: str) -> bool:
#     if not q: return True
#     s = q.strip()
#     if not re.search(r'[.!?]', s):
#         words = re.findall(r"[A-Za-z]+", s)
#         if 1 <= len(words) <= 7:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.8: return True
#         if s.lower() in {"introduction","voiceover","outro","epilogue","prologue","credits","title","hook","horrifying discovery","final decision"}: return True
#         if len(s) <= 3: return True
#     return False

# def _is_heading_context(script_text: str, s: int, e: int) -> bool:
#     left = script_text.rfind("\n", 0, s) + 1
#     right = script_text.find("\n", e); right = len(script_text) if right == -1 else right
#     line = script_text[left:right].strip()
#     if len(line) <= 70 and not re.search(r'[.!?]', line):
#         words = re.findall(r"[A-Za-z]+", line)
#         if 1 <= len(words) <= 8:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.7: return True
#     return False

# def _tighten_to_quote(script_text: str, span: Tuple[int,int], quote: str) -> Tuple[int,int]:
#     if not span or not quote: return span
#     s, e = span
#     if e <= s or s < 0 or e > len(script_text): return span
#     window = script_text[s:e]; win_norm = _normalize_keep_len(window).lower(); q_norm = _clean_quote_for_match(quote).lower()
#     if not q_norm: return span
#     i = win_norm.find(q_norm)
#     if i == -1:
#         m = re.search(re.escape(q_norm).replace(r"\ ", r"\s+"), win_norm, flags=re.IGNORECASE)
#         if not m: return span
#         i, j = m.start(), m.end()
#     else:
#         j = i + len(q_norm)
#     s2, e2 = s + i, s + j
#     s2, e2 = _snap_and_bridge_to_word(script_text, s2, e2); s2 = _heal_split_word_left(script_text, s2)
#     if s2 >= s and e2 <= e and e2 > s2: return (s2, e2)
#     return span

# def build_spans_by_param(script_text: str, data: dict, heading_ranges: Optional[List[Tuple[int,int]]] = None) -> Dict[str, List[Tuple[int,int,str,str]]]:
#     heading_ranges = heading_ranges or []
#     raw = (data or {}).get("per_parameter", {}) or {}
#     per: Dict[str, Dict[str, Any]] = {k:(v or {}) for k,v in raw.items()}
#     spans_map: Dict[str, List[Tuple[int,int,str,str]]] = {p: [] for p in PARAM_ORDER}
#     st.session_state["aoi_match_ranges"] = {}

#     for p in spans_map.keys():
#         color = PARAM_COLORS.get(p, "#ffd54f")
#         blk = per.get(p, {}) or {}
#         aois = blk.get("areas_of_improvement") or []
#         for idx, item in enumerate(aois, start=1):
#             raw_q = (item or {}).get("quote_verbatim", "") or ""
#             q = _sanitize_editor_text(raw_q)
#             clean = _clean_quote_for_match(re.sub(r"^[•\-\d\.\)\s]+", "", q).strip())
#             if not clean: continue
#             if _is_heading_like(clean): continue

#             pos = find_span_smart(script_text, clean)
#             if not pos: continue
#             pos = _tighten_to_quote(script_text, pos, raw_q)
#             s, e = pos

#             if heading_ranges and _overlaps_any(s, e, heading_ranges): continue
#             if _is_heading_context(script_text, s, e): continue

#             aid = f"{p.replace(' ','_')}-AOI-{idx}"
#             spans_map[p].append((s, e, color, aid))
#             st.session_state["aoi_match_ranges"][aid] = (s, e)
#     return spans_map

# # ---------- History (S3-aware) ----------
# def _maybe_copy_docx_to_history(source_docx_path: Optional[str], run_id: str) -> Optional[str]:
#     """Copy the DOCX used for rendering into outputs/_history so Recents can re-render identically."""
#     try:
#         if not source_docx_path:
#             return None
#         # If already an S3 key/url, just store the key in history blob
#         if source_docx_path.startswith("s3://") or (_s3_enabled() and not os.path.exists(source_docx_path)):
#             # ensure it's uploaded under our history key if it's a local temp in S3 mode
#             if os.path.exists(source_docx_path):
#                 with open(source_docx_path, "rb") as f:
#                     save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#                 return f"{HISTORY_DIR}/{run_id}.docx"
#             return source_docx_path

#         # Local file path
#         if _s3_enabled():
#             with open(source_docx_path, "rb") as f:
#                 save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#             return f"{HISTORY_DIR}/{run_id}.docx"
#         else:
#             dst = os.path.join(HISTORY_DIR, f"{run_id}.docx")
#             if os.path.abspath(source_docx_path) != os.path.abspath(dst):
#                 shutil.copyfile(source_docx_path, dst)
#             return dst
#     except Exception:
#         return None

# def _save_history_snapshot(title: str, data: dict, script_text: str,
#                            source_docx_path: Optional[str], heading_ranges: List[Tuple[int,int]],
#                            spans_by_param: Dict[str, List[Tuple[int,int,str,str]]],
#                            aoi_match_ranges: Dict[str, Tuple[int,int]]):
#     run_id = str(uuid.uuid4()); now = datetime.datetime.now()
#     created_at_iso = now.replace(microsecond=0).isoformat()
#     created_at_human = now.strftime("%Y-%m-%d %H:%M:%S")

#     # Copy stable DOCX alongside history JSON (so Recents uses identical renderer)
#     stable_docx_key_or_path = _maybe_copy_docx_to_history(source_docx_path, run_id)

#     blob = {
#         "run_id": run_id, "title": title or "untitled",
#         "created_at": created_at_iso, "created_at_human": created_at_human,
#         "overall_rating": (data or {}).get("overall_rating", ""),
#         "scores": (data or {}).get("scores", {}),
#         "data": data or {}, "script_text": script_text or "",
#         "source_docx_path": stable_docx_key_or_path or source_docx_path,
#         "heading_ranges": heading_ranges or [],
#         "spans_by_param": spans_by_param or {},
#         "aoi_match_ranges": aoi_match_ranges or {},
#     }

#     # Persist JSON (S3 or local)
#     out_name = f"{created_at_iso.replace(':','-')}__{run_id}.json"
#     out_key = f"{HISTORY_DIR}/{out_name}"
#     save_text_key(out_key, json.dumps(blob, ensure_ascii=False, indent=2))

# def _load_all_history() -> List[dict]:
#     out: List[dict] = []
#     if _s3_enabled():
#         keys = sorted(list_prefix(HISTORY_DIR), reverse=True)
#         for key in keys:
#             try:
#                 txt = read_text_key(key, "")
#                 if not txt:
#                     continue
#                 j = json.loads(txt)
#                 j["_key"] = key
#                 if not j.get("created_at_human") and j.get("created_at"):
#                     try:
#                         dt = datetime.datetime.fromisoformat(j["created_at"])
#                         j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#                     except Exception:
#                         pass
#                 out.append(j)
#             except Exception:
#                 continue
#     else:
#         for fp in sorted(glob.glob(os.path.join(HISTORY_DIR, "*.json"))):
#             try:
#                 with open(fp, "r", encoding="utf-8") as f:
#                     j = json.load(f)
#             except Exception:
#                 continue
#             j.setdefault("_path", fp)
#             ca = j.get("created_at")
#             try:
#                 if isinstance(ca, (int, float)):
#                     dt = datetime.datetime.utcfromtimestamp(float(ca))
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.astimezone().strftime("%Y-%m-%d %H:%M:%S")
#                 elif isinstance(ca, str) and ca: pass
#                 else:
#                     mtime = os.path.getmtime(fp); dt = datetime.datetime.fromtimestamp(mtime)
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#             except Exception:
#                 j["created_at"] = str(ca or "")
#             out.append(j)
#     out.sort(key=lambda r: r.get("created_at") or "", reverse=True)
#     return out

# def _open_history_run_by_id(run_id: str) -> bool:
#     """Open a history run by its run_id. Returns True if loaded."""
#     if not run_id: return False
#     recs = _load_all_history()
#     match = next((r for r in recs if r.get("run_id") == run_id), None)
#     if not match: return False

#     # Load JSON content again (S3/local), then set session
#     try:
#         if "_key" in match and _s3_enabled():
#             txt = read_text_key(match["_key"], "")
#             if not txt: return False
#             jj = json.loads(txt)
#         else:
#             path = match.get("_path")
#             if not path or not os.path.exists(path): return False
#             with open(path, "r", encoding="utf-8") as f:
#                 jj = json.load(f)
#     except Exception:
#         return False

#     st.session_state.script_text      = jj.get("script_text","")
#     st.session_state.base_stem        = jj.get("title","untitled")
#     st.session_state.data             = jj.get("data",{})
#     st.session_state.heading_ranges   = jj.get("heading_ranges",[])
#     st.session_state.spans_by_param   = jj.get("spans_by_param",{})
#     st.session_state.param_choice     = None
#     st.session_state.source_docx_path = jj.get("source_docx_path")
#     st.session_state.review_ready     = True
#     st.session_state["aoi_match_ranges"] = jj.get("aoi_match_ranges", {})
#     st.session_state.ui_mode          = "review"
#     return True

# def _render_recents_centerpane():
#     st.subheader("📄 Recents")
#     q = st.text_input("Filter by title…", "")
#     cols = st.columns([1, 4])
#     with cols[0]:
#         if st.button("← Back"):
#             st.session_state.ui_mode = "home"; st.rerun()

#     recs = _load_all_history()
#     ql = q.strip().lower()
#     if ql: recs = [r for r in recs if ql in (r.get("title","").lower())]
#     if not recs: st.caption("No history yet."); st.stop()

#     # Each card is a clickable <a class="rec-card" href="?open=<run_id>">
#     for rec in recs:
#         run_id = rec.get("run_id"); title = rec.get("title") or "(untitled)"
#         created_h = rec.get("created_at_human",""); overall = rec.get("overall_rating","")
#         st.markdown(f"""
#         <a class="rec-card" href="?open={run_id}">
#           <div class="rec-title">{title}</div>
#           <div class="rec-meta">{created_h}</div>
#           <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
#         </a>
#         """, unsafe_allow_html=True)

# # ---------- Sidebar ----------
# with st.sidebar:
#     if st.button("🆕 New review", use_container_width=True):
#         fp = st.session_state.get("flattened_docx_path")
#         if fp and os.path.exists(fp):
#             try: os.remove(fp)
#             except Exception: pass
#         for k in ["review_ready","script_text","base_stem","data","spans_by_param","param_choice",
#                   "source_docx_path","heading_ranges","flattened_docx_path","flatten_used"]:
#             st.session_state[k] = (
#                 False if k=="review_ready"
#                 else "" if k in ("script_text","base_stem")
#                 else {} if k=="spans_by_param"
#                 else [] if k=="heading_ranges"
#                 else None if k in ("source_docx_path","flattened_docx_path")
#                 else False if k=="flatten_used"
#                 else None
#             )
#         st.session_state.ui_mode = "home"
#         _clear_query_params()
#         st.rerun()

#     if st.button("📁 Recents", use_container_width=True):
#         st.session_state.ui_mode = "recents"
#         _clear_query_params()
#         st.rerun()

# # ---------- Input screen ----------
# def render_home():
#     st.subheader("🎬 Script Source")

#     (tab_upload,) = st.tabs(["Upload file"])
#     uploaded_file = None
#     uploaded_name = None
#     uploaded_key  = None

#     with tab_upload:
#         up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
#         if up is not None:
#             file_bytes = up.read()
#             suffix = os.path.splitext(up.name)[1].lower()
#             # Save to S3 (or local) under scripts/
#             uploaded_key = f"{SCRIPTS_DIR}/{up.name}"
#             save_bytes_key(uploaded_key, file_bytes)

#             # Also create a temp local copy for parsing
#             with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
#                 tmp.write(file_bytes)
#                 uploaded_file = tmp.name
#             uploaded_name = os.path.splitext(os.path.basename(up.name))[0] or "uploaded_script"

#     if st.button("🚀 Run Review", type="primary", use_container_width=True):
#         base_stem = "uploaded_script"
#         source_docx_path = None
#         heading_ranges = []

#         if uploaded_file:
#             base_stem = uploaded_name or "uploaded_script"
#             # Prefer key (so we can store it into history); for DOCX rendering we keep a tmp local
#             if uploaded_file.lower().endswith(".docx"):
#                 path_to_use = uploaded_file
#                 if _docx_contains_tables(path_to_use):
#                     flat = flatten_docx_tables_to_longtext(path_to_use)
#                     st.session_state.flattened_docx_path = flat
#                     st.session_state.flatten_used = True
#                     path_to_use = flat
#                 script_text, heading_ranges = build_docx_text_with_meta(path_to_use)

#                 # 🔑 IMPORTANT: set the rendering source to the *flattened* docx
#                 # (History saver will copy this file to Scriptmodel/outputs/_history/)
#                 source_docx_path = path_to_use
#             else:
#                 script_text = load_script_file(uploaded_file)
#                 source_docx_path = uploaded_key or uploaded_file  # keep reference to S3/local
#         else:
#             st.warning("Please upload a script first.")
#             st.stop()

#         if len(script_text.strip()) < 50:
#             st.error("Extracted text looks too short. Check your file extraction.")
#             st.stop()

#         with st.spinner("Running analysis…"):
#             try:
#                 review_text = run_review_multi(
#                     script_text=script_text,
#                     prompts_dir=PROMPTS_DIR,
#                     temperature=0.0
#                 )
#             finally:
#                 # Clean temp upload if we didn't keep it as source_docx_path local
#                 if uploaded_file and not (isinstance(source_docx_path, str) and os.path.exists(source_docx_path)):
#                     try:
#                         os.remove(uploaded_file)
#                     except Exception:
#                         pass

#         data = extract_review_json(review_text)
#         if not data:
#             st.error("JSON not detected in model output.")
#             st.stop()

#         st.session_state.script_text      = script_text
#         st.session_state.base_stem        = base_stem
#         st.session_state.data             = data
#         st.session_state.heading_ranges   = heading_ranges
#         st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
#         st.session_state.param_choice     = None
#         st.session_state.source_docx_path = source_docx_path
#         st.session_state.review_ready     = True
#         st.session_state.ui_mode          = "review"

#         _save_history_snapshot(
#             title=base_stem,
#             data=data,
#             script_text=script_text,
#             source_docx_path=source_docx_path,
#             heading_ranges=heading_ranges,
#             spans_by_param=st.session_state.spans_by_param,
#             aoi_match_ranges=st.session_state.get("aoi_match_ranges", {})
#         )

#         _clear_query_params()
#         st.rerun()

# # ---------- Results screen ----------
# def render_review():
#     script_text     = st.session_state.script_text
#     data            = st.session_state.data
#     spans_by_param  = st.session_state.spans_by_param
#     scores: Dict[str,int] = (data or {}).get("scores", {}) or {}
#     source_docx_path: Optional[str] = st.session_state.source_docx_path

#     # If our source_docx_path is an S3 key/url, ensure we have a local copy for rendering
#     # AFTER (prefer the flattened copy if we have it in-session)
#     docx_local: Optional[str] = None
#     preferred = st.session_state.get("flattened_docx_path") if st.session_state.get("flatten_used") else source_docx_path
#     if preferred:
#         if not os.path.exists(preferred):
#             docx_local = ensure_local_copy(preferred)
#         else:
#             docx_local = preferred

#     left, center, right = st.columns([1.1, 2.7, 1.4], gap="large")

#     with left:
#         st.subheader("Final score")
#         ordered = [p for p in PARAM_ORDER if p in scores]
#         df = pd.DataFrame({"Parameter": ordered, "Score (1–10)": [scores.get(p, "") for p in ordered]})
#         st.dataframe(df, hide_index=True, use_container_width=True)
#         st.markdown(f'**Overall:** {data.get("overall_rating","—")}/10')
#         st.divider()

#         strengths = (data or {}).get("strengths") or []
#         if not strengths:
#             per = (data or {}).get("per_parameter", {}) or {}
#             best = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
#             for name, sc in best:
#                 if sc >= 8 and name in per:
#                     exp = _sanitize_editor_text((per[name] or {}).get("explanation", "") or "")
#                     first = re.split(r"(?<=[.!?])\s+", exp.strip())[0] if exp else f"Consistently strong {name.lower()}."
#                     strengths.append(f"{name}: {first}")
#                 if len(strengths) >= 3: break

#         def _bullets(title: str, items):
#             st.markdown(f"**{title}**")
#             for s in (items or []):
#                 if isinstance(s, str) and s.strip(): st.write("• " + _sanitize_editor_text(s))
#             if not items: st.write("• —")

#         _bullets("Strengths", strengths)
#         _bullets("Weaknesses", data.get("weaknesses"))
#         _bullets("Suggestions", data.get("suggestions"))
#         _bullets("Drop-off Risks", data.get("drop_off_risks"))
#         st.markdown("**Viral Quotient**"); st.write(_sanitize_editor_text(data.get("viral_quotient","—")))

#     with right:
#         st.subheader("Parameters")
#         st.markdown('<div class="param-row">', unsafe_allow_html=True)
#         for p in [p for p in PARAM_ORDER if p in scores]:
#             if st.button(p, key=f"chip_{p}", help="Show inline AOI highlights for this parameter"):
#                 st.session_state.param_choice = p
#         st.markdown('</div>', unsafe_allow_html=True)

#         sel = st.session_state.param_choice
#         if sel:
#             blk = (data.get("per_parameter", {}) or {}).get(sel, {}) or {}
#             st.markdown(f"**{sel} — Score:** {scores.get(sel,'—')}/10")

#             if blk.get("explanation"):
#                 st.markdown("**Why this score**"); st.write(_sanitize_editor_text(blk["explanation"]))
#             if blk.get("weakness") and blk["weakness"] != "Not present":
#                 st.markdown("**Weakness**"); st.write(_sanitize_editor_text(blk["weakness"]))
#             if blk.get("suggestion") and blk["suggestion"] != "Not present":
#                 st.markdown("**Suggestion**"); st.write(_sanitize_editor_text(blk["suggestion"]))

#             aoi = blk.get("areas_of_improvement") or []
#             if aoi:
#                 st.markdown("**Areas of Improvement**")
#                 for i, item in enumerate(aoi, 1):
#                     popover_fn = getattr(st, "popover", None)
#                     aid = f"{sel.replace(' ','_')}-AOI-{i}"
#                     s_e_map = st.session_state.get("aoi_match_ranges", {})
#                     if aid in s_e_map:
#                         s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                         line = (matched_line if len(matched_line) <= 320 else matched_line[:300] + "…")
#                     else:
#                         line = _sanitize_editor_text(item.get('quote_verbatim',''))
#                     issue = _sanitize_editor_text(item.get('issue',''))
#                     fix   = _sanitize_editor_text(item.get('fix',''))
#                     why   = _sanitize_editor_text(item.get('why_this_helps',''))
#                     label = f"Issue {i}"
#                     if callable(popover_fn):
#                         with popover_fn(label):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)
#                     else:
#                         with st.expander(label, expanded=False):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)

#             if blk.get("summary"):
#                 st.markdown("**Summary**"); st.write(_sanitize_editor_text(blk["summary"]))

#     with center:
#         st.subheader("Script with inline highlights")
#         # choose spans for selected parameter (or all if None)
#         spans = st.session_state.spans_by_param.get(st.session_state.param_choice, []) if st.session_state.param_choice else []

#         # Build AOI payload for popup: { aid: {line, issue, fix, why} }
#         aoi_payload: Dict[str, Dict[str, str]] = {}
#         data_per = (data or {}).get("per_parameter") or {}
#         s_e_map = st.session_state.get("aoi_match_ranges", {})
#         sel = st.session_state.param_choice

#         def _mk_line(aid: str, fallback_q: str = "") -> str:
#             if aid in s_e_map:
#                 s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                 return matched_line if len(matched_line) <= 320 else matched_line[:300] + "…"
#             return _sanitize_editor_text(fallback_q or "")

#         def _collect(param_name: str):
#             blk = (data_per.get(param_name) or {})
#             for i, item in enumerate(blk.get("areas_of_improvement") or [], 1):
#                 aid = f"{param_name.replace(' ','_')}-AOI-{i}"
#                 aoi_payload[aid] = {
#                     "line": _mk_line(aid, (item or {}).get("quote_verbatim","")),
#                     "issue": _sanitize_editor_text((item or {}).get("issue","")),
#                     "fix": _sanitize_editor_text((item or {}).get("fix","")),
#                     "why": _sanitize_editor_text((item or {}).get("why_this_helps","")),
#                 }

#         if sel: _collect(sel)
#         else:
#             for pn in [p for p in PARAM_ORDER if p in data_per]:
#                 _collect(pn)

#         payload_json = json.dumps(aoi_payload, ensure_ascii=False)

#         frame_theme_css = """
#         <style>
#           :root{
#             --m7-surface: #eef2f7;
#             --m7-on-surface: #0f172a;
#             --m7-border: rgba(15,23,42,.14);
#           }
#           @media (prefers-color-scheme: dark){
#             :root{
#               --m7-surface: #2f333a;
#               --m7-on-surface: #ffffff;
#               --m7-border: rgba(255,255,255,.18);
#             }
#             body { background: transparent !important; }
#           }

#           .docxwrap{
#             background: var(--m7-surface);
#             color: var(--m7-on-surface);
#             border: 1px solid var(--m7-border);
#             border-radius: 12px;
#             padding: 16px 14px 18px;
#           }
#           .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#           .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); }
#         </style>
#         """

#         tooltip_css = """
#         <style>
#         .aoi-pop {
#           position: absolute; max-width: 520px; min-width: 320px;
#           background: var(--m7-surface); border: 1px solid var(--m7-border); border-radius: 10px;
#           box-shadow: 0 10px 25px rgba(0,0,0,.12);
#           padding: 12px 14px; z-index: 9999; transform: translateY(-8px);
#           color: var(--m7-on-surface);
#         }
#         .aoi-pop h4 { margin: 0 0 .35rem 0; font-size: .95rem; }
#         .aoi-pop p  { margin: .15rem 0; line-height: 1.5; }
#         .aoi-pop .muted { opacity:.85; font-size:.85rem; }
#         .aoi-arrow {
#           position:absolute; left:50%; transform:translateX(-50%);
#           bottom:-7px; width:0;height:0;border-left:7px solid transparent;
#           border-right:7px solid transparent;border-top:7px solid var(--m7-border);
#         }
#         .aoi-arrow::after{
#           content:""; position:absolute; left:-6px; top:-7px; width:0;height:0;
#           border-left:6px solid transparent;border-right:6px solid transparent;border-top:6px solid var(--m7-surface);
#         }
#         </style>
#         """

#         # Select rendering source (DOCX with highlights if we have a local path)
#         if docx_local and os.path.splitext(docx_local)[1].lower() == ".docx":
#             html_core = render_docx_html_with_highlights(
#                 docx_local,
#                 merge_overlaps_and_adjacent(script_text, spans)
#             )
#         else:
#             from html import escape as _esc
#             orig = script_text
#             spans2 = [s for s in merge_overlaps_and_adjacent(orig, spans) if s[0] < s[1]]
#             spans2.sort(key=lambda x: x[0])
#             cur = 0; buf: List[str] = []
#             for s,e,c,aid in spans2:
#                 if s > cur: buf.append(_esc(orig[cur:s]))
#                 buf.append(
#                     f'<mark class="aoi-mark" data-aid="{aid}" '
#                     f'style="background:{c}33;border:1px solid {c};border-radius:3px;padding:0 2px;">'
#                     f'{_esc(orig[s:e])}</mark>'
#                 )
#                 cur = e
#             if cur < len(orig): buf.append(_esc(orig[cur:]))
#             html_core = (
#                 '<div class="docxwrap"><p style="white-space:pre-wrap; '
#                 'line-height:1.7; font-family:ui-serif, Georgia, Times New Roman, serif;">'
#                 + "".join(buf) +
#                 '</p></div>'
#             )

#         # Popup + autosize JS shell
#         html_shell = """
# %%FRAME_THEME_CSS%%
# %%TOOLTIP_CSS%%
# <div id="m7-doc">%%HTML_CORE%%</div>
# <div id="aoi-pop" class="aoi-pop" style="display:none;">
#   <div id="aoi-pop-content"></div>
#   <div class="aoi-arrow"></div>
# </div>
# <script>
# (function(){
#   const AOI = __PAYLOAD__;
#   const wrap = document.getElementById('m7-doc');
#   const pop  = document.getElementById('aoi-pop');
#   const body = document.getElementById('aoi-pop-content');

#   function resizeIframe() {
#     try {
#       const h = Math.max(
#         document.documentElement.scrollHeight,
#         document.body.scrollHeight
#       );
#       if (window.frameElement) {
#         window.frameElement.style.height = (h + 20) + 'px';
#         window.frameElement.style.width  = '100%';
#       }
#     } catch(e) {}
#   }
#   window.addEventListener('load', resizeIframe);
#   window.addEventListener('resize', resizeIframe);

#   function hide(){ pop.style.display='none'; }
#   function showFor(mark){
#     const aid = mark.getAttribute('data-aid');
#     const d = AOI[aid]; if(!d) return;
#     body.innerHTML =
#       (d.line  ? '<p><strong>Line:</strong> '  + d.line  + '</p>' : '') +
#       (d.issue ? '<p><strong>Issue:</strong> ' + d.issue + '</p>' : '') +
#       (d.fix   ? '<p><strong>Fix:</strong> '   + d.fix   + '</p>' : '') +
#       (d.why   ? '<p class="muted">'           + d.why   + '</p>' : '');
#     pop.style.display = 'block';

#     const r = mark.getBoundingClientRect();
#     const scY = window.scrollY || document.documentElement.scrollTop;
#     const scX = window.scrollX || document.documentElement.scrollLeft;
#     let top  = r.top + scY - pop.offsetHeight - 10;
#     let left = r.left + scX + r.width/2 - pop.offsetWidth/2;
#     if (top < 8) top = r.bottom + scY + 10;
#     if (left < 8) left = 8;
#     pop.style.top  = top + 'px';
#     pop.style.left = left + 'px';

#     resizeIframe();
#   }

#   wrap.addEventListener('click', (e) => {
#     const m = e.target.closest('.aoi-mark');
#     if(!m){ hide(); return; }
#     if(pop.style.display === 'block'){ hide(); }
#     showFor(m);
#     e.stopPropagation();
#   });

#   document.addEventListener('click', (e) => {
#     if(!e.target.closest('.aoi-pop') && !e.target.closest('.aoi-mark')) hide();
#   });
# })();
# </script>
# """
#         html_shell = (
#             html_shell
#             .replace("%%FRAME_THEME_CSS%%", frame_theme_css)
#             .replace("%%TOOLTIP_CSS%%", tooltip_css)
#             .replace("%%HTML_CORE%%", html_core)
#             .replace("__PAYLOAD__", payload_json)
#         )

#         components.html(html_shell, height=400, scrolling=False)

# # ---------- Router & query param open ----------
# _open_qp = _get_query_param("open")
# if _open_qp and _open_history_run_by_id(_open_qp):
#     _clear_query_params()  # avoid re-opening on subsequent reruns

# mode = st.session_state.ui_mode
# if mode == "recents":
#     _render_recents_centerpane()
# elif mode == "review" and st.session_state.review_ready:
#     render_review()
# else:
#     render_home()





################################################








##Current working version 










# import os, re, glob, json, tempfile, difflib, uuid, datetime, shutil
# from pathlib import Path
# from typing import Dict, Any, List, Tuple, Optional

# import streamlit as st
# import pandas as pd
# import streamlit.components.v1 as components  # for inline HTML/JS popup

# from utils1 import extract_review_json, PARAM_ORDER, load_script_file
# from review_engine_multi import run_review_multi

# # ---- DOCX rendering imports ----
# from docx import Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # =========================
# # RunPod S3 (inline helpers) — FIXED
# # =========================
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError

# # Accept both AWS_* and RUNPOD_* style envs / st.secrets
# def _get_env(key: str, default: str = "") -> str:
#     v = os.getenv(key, "")
#     if v: 
#         return v.strip()
#     try:
#         v2 = st.secrets.get(key)
#         if isinstance(v2, str):
#             return v2.strip()
#     except Exception:
#         pass
#     return (default or "").strip()

# # Primary config
# _RP_ENDPOINT = _get_env("RUNPOD_S3_ENDPOINT")
# _RP_BUCKET   = _get_env("RUNPOD_S3_BUCKET")
# _RP_REGION   = _get_env("RUNPOD_S3_REGION") or _get_env("AWS_DEFAULT_REGION") or ""

# # Credentials: prefer AWS_* if present; else accept RUNPOD_* fallbacks
# _AK = _get_env("AWS_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY")
# _SK = _get_env("AWS_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_KEY")
# _ST = _get_env("AWS_SESSION_TOKEN")  # optional

# # Options
# _FORCE_PATH = (_get_env("RUNPOD_S3_FORCE_PATH_STYLE") or "true").lower() in {"1","true","yes"}
# _USE_SSL    = (_get_env("RUNPOD_S3_USE_SSL") or "true").lower() in {"1","true","yes"}
# _VERIFY_SSL = (_get_env("RUNPOD_S3_VERIFY_SSL") or "true").lower() in {"1","true","yes"}

# def _s3_enabled() -> bool:
#     return bool(_RP_ENDPOINT and _RP_BUCKET and _AK and _SK)

# @st.cache_resource(show_spinner=False)
# def _s3_client():
#     if not _s3_enabled():
#         return None
#     session_kwargs = dict(
#         aws_access_key_id=_AK,
#         aws_secret_access_key=_SK,
#     )
#     if _ST:
#         session_kwargs["aws_session_token"] = _ST

#     # s3v4 + path style are common requirements for S3-compatible services
#     cfg = Config(
#         signature_version="s3v4",
#         s3={"addressing_style": "path" if _FORCE_PATH else "auto"},
#         retries={"max_attempts": 3, "mode": "standard"}
#     )
#     return boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         region_name=_RP_REGION or None,
#         use_ssl=_USE_SSL,
#         verify=_VERIFY_SSL,
#         config=cfg,
#         **session_kwargs,
#     )

# def save_text_key(key: str, text: str) -> str:
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "w", encoding="utf-8") as f:
#             f.write(text)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=text.encode("utf-8"))
#     return f"s3://{_RP_BUCKET}/{key}"

# def save_bytes_key(key: str, data: bytes) -> str:
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "wb") as f:
#             f.write(data)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=data)
#     return f"s3://{_RP_BUCKET}/{key}"

# def read_text_key(key: str, default: str = "") -> str:
#     if not _s3_enabled():
#         try:
#             with open(key, "r", encoding="utf-8") as f:
#                 return f.read()
#         except Exception:
#             return default
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read().decode("utf-8", errors="ignore")
#     except Exception:
#         return default

# def read_bytes_key(key: str) -> Optional[bytes]:
#     if not _s3_enabled():
#         try:
#             with open(key, "rb") as f:
#                 return f.read()
#         except Exception:
#             return None
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read()
#     except Exception:
#         return None

# def list_prefix(prefix: str) -> List[str]:
#     """
#     List object keys under prefix (or local dir paths if not S3).
#     In S3 mode we always return KEYS (not URLs).
#     """
#     if not _s3_enabled():
#         base = prefix if os.path.isdir(prefix) else os.path.dirname(prefix)
#         try:
#             return [os.path.join(base, p) for p in os.listdir(base) if p.endswith(".json")]
#         except Exception:
#             return []

#     out: List[str] = []
#     token = None
#     # Normalize to "dir/" prefix for S3 listing
#     s3_prefix = prefix.rstrip("/") + "/"
#     try:
#         while True:
#             kwargs = {"Bucket": _RP_BUCKET, "Prefix": s3_prefix}
#             if token:
#                 kwargs["ContinuationToken"] = token
#             resp = _s3_client().list_objects_v2(**kwargs)
#             for c in resp.get("Contents", []):
#                 k = c.get("Key", "")
#                 if k.endswith(".json"):
#                     out.append(k)
#             token = resp.get("NextContinuationToken")
#             if not token:
#                 break
#     except (ClientError, EndpointConnectionError, NoCredentialsError):
#         return []
#     return out

# def presigned_url(key: str, expires: int = 3600) -> Optional[str]:
#     if not _s3_enabled():
#         return None
#     try:
#         return _s3_client().generate_presigned_url(
#             "get_object",
#             Params={"Bucket": _RP_BUCKET, "Key": key},
#             ExpiresIn=expires
#         )
#     except ClientError:
#         return None

# def ensure_local_copy(key_or_path: str) -> Optional[str]:
#     """
#     For DOCX/PDF parsing we need a real filesystem path.
#     If S3 mode, download to a temp file and return that path.
#     """
#     if not _s3_enabled():
#         return key_or_path if os.path.exists(key_or_path) else None

#     key = key_or_path
#     if key.startswith("s3://"):
#         # s3://bucket/path/to/file -> path/to/file
#         parts = key.split("/", 3)
#         key = parts[3] if len(parts) >= 4 else ""
#     data = read_bytes_key(key)
#     if data is None:
#         return None
#     fd, tmp = tempfile.mkstemp(suffix=os.path.splitext(key)[1] or "")
#     os.close(fd)
#     with open(tmp, "wb") as f:
#         f.write(data)
#     return tmp

# def _s3_health_summary() -> dict:
#     """
#     Optional tiny health read you can print if needed.
#     Returns a dict; safe to ignore in production.
#     """
#     info = {
#         "enabled": _s3_enabled(),
#         "endpoint": _RP_ENDPOINT,
#         "bucket": _RP_BUCKET,
#         "region": _RP_REGION,
#         "has_keys": bool(_AK and _SK),
#     }
#     if not _s3_enabled():
#         info["status"] = "local-mode"
#         return info
#     try:
#         # Attempt a very cheap list; no exceptions => reachable
#         _ = _s3_client().list_objects_v2(Bucket=_RP_BUCKET, Prefix=(f"{OUTPUT_DIR}/_history/").rstrip("/") + "/",
#                                          MaxKeys=1)
#         info["status"] = "ok"
#     except Exception as e:
#         info["status"] = f"error: {getattr(e, 'response', {}).get('Error', {}).get('Code', str(e))}"
#     return info


# # ---------- Folders ----------
# # SCRIPTS_DIR = "scripts"
# # PROMPTS_DIR = "prompts"
# # OUTPUT_DIR  = "outputs"
# # HISTORY_DIR = os.path.join(OUTPUT_DIR, "_history")
# # ---------- Folders (all under Scriptmodel/) ----------

# BASE_PREFIX = "Scriptmodel"

# SCRIPTS_DIR = f"{BASE_PREFIX}/scripts"
# PROMPTS_DIR = f"{BASE_PREFIX}/prompts"
# OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
# HISTORY_DIR = f"{OUTPUT_DIR}/_history"


# if not _s3_enabled():
#     for p in (SCRIPTS_DIR, PROMPTS_DIR, OUTPUT_DIR, HISTORY_DIR):
#         Path(p).mkdir(parents=True, exist_ok=True)


# # ---------- Colors ----------
# PARAM_COLORS: Dict[str, str] = {
#     "Suspense Building":              "#ff6b6b",
#     "Language/Tone":                  "#6b8cff",
#     "Intro + Main Hook/Cliffhanger":  "#ffb86b",
#     "Story Structure + Flow":         "#a78bfa",
#     "Pacing":                         "#f43f5e",
#     "Mini-Hooks (30–60s)":            "#eab308",
#     "Outro (Ending)":                 "#8b5cf6",
#     "Grammar & Spelling":             "#10b981",  # NEW: distinct highlight color for grammar/spelling
# }

# # ---------- Config ----------
# STRICT_MATCH_ONLY = False

# # ---------- App config ----------
# st.set_page_config(page_title="M7 — Grammarly UI", page_icon="🕵️", layout="wide")
# # ---------- Header patch & CSS ----------
# def render_app_title():
#     st.markdown(
#         '<h1 class="app-title">Viral Script Reviewer</h1>',
#         unsafe_allow_html=True
#     )
#     st.markdown("""
#     <style>
#     html { color-scheme: light dark; }

#     :root{
#       --m7-surface: #eef2f7;
#       --m7-on-surface: #0f172a;
#       --m7-border: rgba(15,23,42,.14);
#       --sep: #e5e7eb;
#     }
#     @media (prefers-color-scheme: dark){
#       :root{
#         --m7-surface: #2f333a;
#         --m7-on-surface: #ffffff;
#         --m7-border: rgba(255,255,255,.18);
#         --sep: #2a2f37;
#       }
#     }

#     .stApp .block-container { padding-top: 4.25rem !important; }
#     .app-title{
#       font-weight: 700; font-size: 2.1rem; line-height: 1.3;
#       margin: 0 0 1rem 0; padding-left: 40px !important; padding-top: .25rem !important;
#       white-space: normal; word-break: break-word; hyphens: auto; overflow: visible;
#       position: relative !important; z-index: 10 !important;
#     }
#     [data-testid="collapsedControl"] { z-index: 6 !important; }
#     header[data-testid="stHeader"], .stAppHeader { background: transparent !important; box-shadow: none !important; }
#     @media (min-width: 992px){ .app-title { padding-left: 0 !important; } }

#     div[data-testid="column"]:nth-of-type(1){position:relative;}
#     div[data-testid="column"]:nth-of-type(1)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     div[data-testid="column"]:nth-of-type(2){position:relative;}
#     div[data-testid="column"]:nth-of-type(2)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}

#     .m7-card{
#       background: var(--m7-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 14px 16px;
#       color: var(--m7-on-surface);
#     }
#     .m7-card, .m7-card * { color: var(--m7-on-surface) !important; }

#     .docxwrap{
#       background: var(--m7-surface);
#       color: var(--m7-on-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 16px 14px 18px;
#     }
#     .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#     .docxwrap .h1, .docxwrap .h2, .docxwrap .h3 { font-weight:700; margin:10px 0 6px; }
#     .docxwrap .h1 { font-size: 1.3rem; border-bottom: 2px solid currentColor; padding-bottom: 4px; }
#     .docxwrap .h2 { font-size: 1.15rem; border-bottom: 1px solid currentColor; padding-bottom: 3px; }
#     .docxwrap .h3 { font-size: 1.05rem; }
#     .docxwrap p { margin: 10px 0; line-height: 1.7; font-family: ui-serif, Georgia, "Times New Roman", serif; }
#     .docxwrap table { border-collapse: collapse; width: 100%; margin: 12px 0; }
#     .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); padding:8px; vertical-align:top; line-height:1.6; }
#     .docxwrap mark{ padding:0 2px; border-radius:3px; border:1px solid var(--m7-border); cursor: pointer; }

#     .rec-card{
#       display:block; text-decoration:none !important;
#       background: var(--m7-surface);
#       border:1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 14px 16px; margin: 10px 0 16px;
#       box-shadow: 0 1px 2px rgba(0,0,0,.06);
#       color: var(--m7-on-surface) !important;
#       transition: filter .1s ease, transform .02s ease;
#     }
#     .rec-card:hover{ filter: brightness(1.02); }
#     .rec-card:active{ transform: translateY(1px); }
#     .rec-card, .rec-card * { color: var(--m7-on-surface) !important; }
#     .rec-title{font-weight:600; margin-bottom:.25rem;}
#     .rec-meta{opacity:.85 !important; font-size:12.5px; margin-bottom:.4rem;}
#     .rec-row{display:flex; align-items:center; justify-content:space-between; gap:12px;}

#     .stTextInput>div>div,
#     .stTextArea>div>div,
#     .stNumberInput>div>div,
#     .stDateInput>div>div,
#     .stTimeInput>div>div,
#     .stFileUploader>div,
#     div[data-baseweb="select"]{
#       background: var(--m7-surface) !important;
#       border: 1px solid var(--m7-border) !important;
#       border-radius: 10px !important;
#       color: var(--m7-on-surface) !important;
#     }
#     .stTextInput input,
#     .stTextArea textarea,
#     .stNumberInput input,
#     .stDateInput input,
#     .stTimeInput input,
#     .stFileUploader div,
#     div[data-baseweb="select"] *{
#       color: var(--m7-on-surface) !important;
#     }
#     .stTextInput input::placeholder,
#     .stTextArea textarea::placeholder{ color: rgba(16,24,39,.55) !important; }
#     @media (prefers-color-scheme: dark){
#       .stTextInput input::placeholder,
#       .stTextArea textarea::placeholder{ color: rgba(255,255,255,.75) !important; }
#     }
#     div[data-testid="stFileUploaderDropzone"] label span { color: var(--m7-on-surface) !important; opacity: 1 !important; }
#     div[data-testid="stFileUploaderDropzone"] label { color: var(--m7-on-surface) !important; }

#     .stMarkdown pre,
#     pre[class*="language-"],
#     .stCodeBlock{
#       background: var(--m7-surface) !important;
#       color: var(--m7-on-surface) !important;
#       border: 1px solid var(--m7-border) !important;
#       border-radius: 12px !important;
#       padding: 12px 14px !important;
#       overflow:auto;
#     }
#     .stMarkdown pre code{ background: transparent !important; color: inherit !important; }

#     div[data-testid="stDataFrame"]{
#       background: var(--m7-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 6px 8px;
#       color: var(--m7-on-surface);
#     }
#     div[data-testid="stDataFrame"] * { color: var(--m7-on-surface) !important; }

#     </style>
#     """, unsafe_allow_html=True)

# render_app_title()

# # ---------- Session ----------
# for key, default in [
#     ("review_ready", False),
#     ("script_text", ""),
#     ("base_stem", ""),
#     ("data", None),
#     ("spans_by_param", {}),
#     ("param_choice", None),
#     ("source_docx_path", None),
#     ("heading_ranges", []),
#     ("flattened_docx_path", None),
#     ("flatten_used", False),
#     ("ui_mode", "home"),
# ]:
#     st.session_state.setdefault(key, default)

# # ---------- helpers for query params (compat across Streamlit versions) ----------
# def _get_query_param(key: str) -> Optional[str]:
#     val = None
#     try:
#         val = st.query_params.get(key)
#     except Exception:
#         q = st.experimental_get_query_params()
#         v = q.get(key)
#         if isinstance(v, list): val = v[0] if v else None
#         else: val = v
#     return val

# def _clear_query_params():
#     try:
#         st.query_params.clear()
#     except Exception:
#         st.experimental_set_query_params()

# # ---------- Sanitizer ----------
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )
# def _sanitize_editor_text(s: Optional[str]) -> str:
#     if not s: return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•)\s*', '', t, flags=re.M)
#     t = re.sub(r'^\s*[-*]\s+', '• ', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# # ---------- DOCX traversal ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# # ---------- Auto-flatten ----------
# def _docx_contains_tables(path: str) -> bool:
#     doc = Document(path)
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Table):
#             return True
#     return False

# def _copy_paragraph(dest_doc: Document, src_para: Paragraph):
#     p = dest_doc.add_paragraph()
#     try:
#         if src_para.style and src_para.style.name:
#             p.style = src_para.style.name
#     except Exception:
#         pass
#     for run in src_para.runs:
#         r = p.add_run(run.text or "")
#         r.bold = run.bold
#         r.italic = run.italic
#         r.underline = run.underline
#     return p

# def flatten_docx_tables_to_longtext(source_path: str) -> str:
#     src = Document(source_path)
#     new = Document()
#     for blk in _iter_docx_blocks(src):
#         if isinstance(blk, Paragraph):
#             _copy_paragraph(new, blk)
#         else:
#             seen_tc_ids = set()
#             for row in blk.rows:
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     if tc_id in seen_tc_ids:
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     for p in cell.paragraphs:
#                         _copy_paragraph(new, p)
#                 new.add_paragraph("")
#             new.add_paragraph("")
#     fd, tmp_path = tempfile.mkstemp(suffix=".docx"); os.close(fd); new.save(tmp_path)
#     return tmp_path

# # ---------- Build plain text + heading ranges ----------
# def build_docx_text_with_meta(docx_path: str) -> Tuple[str, List[Tuple[int,int]]]:
#     doc = Document(docx_path)
#     out: List[str] = []
#     heading_ranges: List[Tuple[int,int]] = []
#     current_offset = 0

#     def _append_and_advance(s: str):
#         nonlocal current_offset
#         out.append(s); current_offset += len(s)

#     seen_tc_ids: set = set()
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             para_text = "".join(run.text or "" for run in blk.runs)
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             if sty.startswith("heading"):
#                 start = current_offset; end = start + len(para_text)
#                 heading_ranges.append((start, end))
#             _append_and_advance(para_text); _append_and_advance("\n")
#         else:
#             for row in blk.rows:
#                 row_cell_tcs = []
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     row_cell_tcs.append((tc_id, cell))
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     if tc_id in seen_tc_ids:
#                         if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     cell_text_parts: List[str] = []
#                     for i, p in enumerate(cell.paragraphs):
#                         t = "".join(r.text or "" for r in p.runs)
#                         sty = (p.style.name or "").lower() if p.style else ""
#                         if sty.startswith("heading"):
#                             hs = current_offset + sum(len(x) for x in cell_text_parts)
#                             he = hs + len(t)
#                             heading_ranges.append((hs, he))
#                         cell_text_parts.append(t)
#                         if i != len(cell.paragraphs) - 1:
#                             cell_text_parts.append("\n")
#                     cell_text = "".join(cell_text_parts)
#                     _append_and_advance(cell_text)
#                     if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                 _append_and_advance("\n")
#             _append_and_advance("\n")

#     return "".join(out), heading_ranges

# def _wrap_inline(safe_text: str, run) -> str:
#     out = safe_text
#     if getattr(run, "underline", False): out = f"<u>{out}</u>"
#     if getattr(run, "italic", False): out = f"<em>{out}</em>"
#     if getattr(run, "bold", False): out = f"<strong>{out}</strong>"
#     return out

# # ---------- Invisible/bridge chars ----------
# _BRIDGE_CHARS = set("\u200b\u200c\u200d\u2060\ufeff\xa0\u00ad")

# # ---------- DOCX -> HTML with highlights (includes data-aid) ----------
# def render_docx_html_with_highlights(docx_path: str,
#                                      highlight_spans: List[Tuple[int,int,str,str]]) -> str:
#     """
#     highlight_spans: list of (start, end, color, aid)
#     """
#     doc = Document(docx_path)
#     spans = [s for s in highlight_spans if s[0] < s[1]]
#     spans.sort(key=lambda x: x[0])

#     cur_span = 0
#     current_offset = 0

#     def esc(s: str) -> str:
#         return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

#     def open_mark_if_needed(html_parts, mark_state, color, end, aid):
#         if not mark_state["open"]:
#             html_parts.append(
#                 f'<mark class="aoi-mark" data-aid="{aid}" '
#                 f'style="background:{color}33;border:1px solid {color};border-radius:3px;padding:0 2px;">'
#             )
#             mark_state.update(open=True, end=end, color=color, aid=aid)

#     def close_mark_if_open(html_parts, mark_state):
#         if mark_state["open"]:
#             html_parts.append('</mark>')
#             mark_state.update(open=False, end=None, color=None, aid=None)

#     def emit_run_text(run_text: str, run, html_parts: List[str], mark_state: Dict[str, Any]):
#         nonlocal cur_span, current_offset
#         t = run_text or ""; i = 0
#         while i < len(t):
#             next_start, next_end, color, next_aid = None, None, None, None
#             if cur_span < len(spans):
#                 next_start, next_end, color, next_aid = spans[cur_span]

#             if not mark_state["open"]:
#                 if cur_span >= len(spans) or current_offset + (len(t) - i) <= next_start:
#                     chunk = t[i:]; html_parts.append(_wrap_inline(esc(chunk), run)); current_offset += len(chunk); break
#                 if current_offset < next_start:
#                     take = next_start - current_offset
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take; continue
#                 open_mark_if_needed(html_parts, mark_state, color, next_end, next_aid)
#             else:
#                 take = min(mark_state["end"] - current_offset, len(t) - i)
#                 if take > 0:
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take
#                 if current_offset >= mark_state["end"]:
#                     close_mark_if_open(html_parts, mark_state)
#                     cur_span += 1

#     html: List[str] = ['<div class="docxwrap">']
#     seen_tc_ids: set = set()

#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             open_tag = '<div class="h1">' if sty.startswith("heading 1") else \
#                        '<div class="h2">' if sty.startswith("heading 2") else \
#                        '<div class="h3">' if sty.startswith("heading 3") else "<p>"
#             close_tag = "</div>" if sty.startswith("heading") else "</p>"
#             html.append(open_tag)
#             for run in blk.runs:
#                 emit_run_text(run.text or "", run, html, mark_state)
#             close_mark_if_open(html, mark_state)
#             html.append(close_tag)
#             current_offset += 1  # '\n'

#         else:
#             html.append("<table>")
#             for row in blk.rows:
#                 html.append("<tr>")
#                 row_cell_tcs = [(id(cell._tc), cell) for cell in row.cells]
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     html.append("<td>")
#                     if tc_id not in seen_tc_ids:
#                         seen_tc_ids.add(tc_id)
#                         for p_idx, p in enumerate(cell.paragraphs):
#                             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#                             html.append("<div>")
#                             for run in p.runs:
#                                 emit_run_text(run.text or "", run, html, mark_state)
#                             close_mark_if_open(html, mark_state)
#                             html.append("</div>")
#                             if p_idx != len(cell.paragraphs) - 1:
#                                 current_offset += 1
#                     html.append("</td>")
#                     if idx != len(row_cell_tcs) - 1: current_offset += 1  # '\t'
#                 html.append("</tr>"); current_offset += 1   # row '\n'
#             html.append("</table>"); current_offset += 1     # extra '\n'

#     html.append("</div>")
#     return "".join(html)

# # ---------- Matching / span utilities ----------
# def _normalize_keep_len(s: str) -> str:
#     trans = {
#         "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
#         "\u2013": "-", "\u2014": "-",
#         "\xa0": " ",
#         "\u200b": " ", "\u200c": " ", "\u200d": " ", "\u2060": " ",
#         "\ufeff": " ", "\u00ad": " ",
#     }
#     return (s or "").translate(str.maketrans(trans))

# def _tokenize(s: str) -> List[str]: return re.findall(r"\w+", (s or "").lower())

# def _iter_sentences_with_spans(text: str) -> List[Tuple[int,int,str]]:
#     spans = []
#     for m in re.finditer(r'[^.!?]+[.!?]+|\Z', text, flags=re.S):
#         s, e = m.start(), m.end()
#         seg = text[s:e]
#         if seg.strip(): spans.append((s, e, seg))
#     return spans

# def _squash_ws(s: str) -> str: return re.sub(r"\s+", " ", s or "").strip()

# def _clean_quote_for_match(q: str) -> str:
#     if not q: return ""
#     q = _normalize_keep_len(q).strip()
#     q = re.sub(r'^[\'"“”‘’\[\(\{<…\-\–\—\s]+', '', q)
#     q = re.sub(r'[\'"“”‘’\]\)\}>…\-\–\—\s]+$', '', q)
#     return _squash_ws(q)

# def _snap_and_bridge_to_word(text: str, start: int, end: int, max_bridge: int = 2) -> Tuple[int,int]:
#     n = len(text); s, e = max(0,start), max(start,end)
#     def _is_inv(ch: str) -> bool: return ch in _BRIDGE_CHARS

#     while s > 0:
#         prev = text[s-1]; cur = text[s] if s < n else ""
#         if prev.isalnum() and cur.isalnum(): s -= 1; continue
#         j = s; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1; 
#         if brid and (s-1)>=0 and text[s-1].isalnum() and (j<n and text[j].isalnum()): s -= 1; continue
#         break

#     while e < n:
#         prev = text[e-1] if e>0 else ""; nxt = text[e]
#         if prev.isalnum() and nxt.isalnum(): e += 1; continue
#         j = e; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (e-1)>=0 and text[e-1].isalnum() and (j<n and text[j].isalnum()): e = j + 1; continue
#         break

#     while e < n and text[e] in ',"”’\')]}': e += 1
#     return s, e

# def _heal_split_word_left(text: str, start: int) -> int:
#     i = start
#     if i <= 1 or i >= len(text): return start
#     if text[i-1] != " ": return start
#     j = i - 2
#     while j >= 0 and text[j].isalpha(): j -= 1
#     prev_token = text[j+1:i-1]
#     if len(prev_token) == 1: return i - 2
#     return start

# def _overlaps_any(s: int, e: int, ranges: List[Tuple[int,int]]) -> bool:
#     for rs, re_ in ranges:
#         if e > rs and s < re_: return True
#     return False

# def _fuzzy_window_span(tl: str, nl: str, start: int, w: int) -> Tuple[float, Optional[Tuple[int,int]]]:
#     window = tl[start:start+w]
#     sm = difflib.SequenceMatcher(a=nl, b=window)
#     blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
#     if not blocks: return 0.0, None
#     coverage = sum(b.size for b in blocks) / max(1, len(nl))
#     first_b = min(blocks, key=lambda b: b.b); last_b = max(blocks, key=lambda b: b.b + b.size)
#     s = start + first_b.b; e = start + last_b.b + last_b.size
#     return coverage, (s, e)

# def find_span_smart(text: str, needle: str) -> Optional[Tuple[int,int]]:
#     if not text or not needle: return None
#     t_orig = text; t_norm = _normalize_keep_len(text); n_norm = _clean_quote_for_match(needle)
#     if not n_norm: return None
#     tl = t_norm.lower(); nl = n_norm.lower()

#     i = tl.find(nl)
#     if i != -1:
#         s, e = _snap_and_bridge_to_word(t_orig, i, i + len(nl)); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     m = re.search(re.escape(nl).replace(r"\ ", r"\s+"), tl, flags=re.IGNORECASE)
#     if m:
#         s, e = _snap_and_bridge_to_word(t_orig, m.start(), m.end()); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     if not STRICT_MATCH_ONLY and len(nl) >= 12:
#         w = max(60, min(240, len(nl) + 80))
#         best_cov, best_span = 0.0, None
#         step = max(1, w // 2)
#         for start in range(0, max(1, len(tl) - w + 1), step):
#             cov, se = _fuzzy_window_span(tl, nl, start, w)
#             if cov > best_cov: best_cov, best_span = cov, se
#         if best_span and best_cov >= 0.65:
#             s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
#             if s > 0 and t_orig[s-1:s+1].lower() in {" the", " a", " an"}: s -= 1
#             s = _heal_split_word_left(t_orig, s)
#             return (s, e)

#     if not STRICT_MATCH_ONLY:
#         keys = [w for w in _tokenize(nl) if len(w) >= 4][:8]
#         if len(keys) >= 2:
#             kset = set(keys)
#             best_score, best_span = 0.0, None
#             for s, e, seg in _iter_sentences_with_spans(t_norm):
#                 toks = set(_tokenize(seg)); ov = len(kset & toks)
#                 if ov == 0: continue
#                 score = ov / max(2, len(kset)); length_pen = min(1.0, 120 / max(20, e - s)); score *= (0.6 + 0.4 * length_pen)
#                 if score > best_score: best_score, best_span = score, (s, min(e, s + 400))
#             if best_span and best_score >= 0.35:
#                 s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1]); s = _heal_split_word_left(t_orig, s)
#                 return (s, e)
#     return None

# def merge_overlaps(spans: List[Tuple[int,int,str,str]]) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans.sort(key=lambda x: x[0]); out = [spans[0]]
#     for s,e,c,aid in spans[1:]:
#         ps,pe,pc,paid = out[-1]
#         if s <= pe and pc == c and e > pe: out[-1] = (ps, e, pc, paid)
#         else: out.append((s,e,c,aid))
#     return out

# _PUNCT_WS = set(" \t\r\n,.;:!?)('\"[]{}-—–…") | _BRIDGE_CHARS

# def merge_overlaps_and_adjacent(base_text: str,
#                                 spans: List[Tuple[int,int,str,str]],
#                                 max_gap: int = 2) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans = sorted(spans, key=lambda x: x[0]); out = [spans[0]]
#     for s, e, c, aid in spans[1:]:
#         ps, pe, pc, paid = out[-1]
#         if c == pc and s <= pe: out[-1] = (ps, max(pe, e), pc, paid); continue
#         if c == pc and s - pe <= max_gap:
#             gap = base_text[max(0, pe):max(0, s)]
#             if all((ch in _PUNCT_WS) for ch in gap): out[-1] = (ps, e, pc, paid); continue
#         out.append((s, e, c, aid))
#     return out

# # ---------- Heading filters ----------
# def _is_heading_like(q: str) -> bool:
#     if not q: return True
#     s = q.strip()
#     if not re.search(r'[.!?]', s):
#         words = re.findall(r"[A-Za-z]+", s)
#         if 1 <= len(words) <= 7:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.8: return True
#         if s.lower() in {"introduction","voiceover","outro","epilogue","prologue","credits","title","hook","horrifying discovery","final decision"}: return True
#         if len(s) <= 3: return True
#     return False

# def _is_heading_context(script_text: str, s: int, e: int) -> bool:
#     left = script_text.rfind("\n", 0, s) + 1
#     right = script_text.find("\n", e); right = len(script_text) if right == -1 else right
#     line = script_text[left:right].strip()
#     if len(line) <= 70 and not re.search(r'[.!?]', line):
#         words = re.findall(r"[A-Za-z]+", line)
#         if 1 <= len(words) <= 8:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.7: return True
#     return False

# def _tighten_to_quote(script_text: str, span: Tuple[int,int], quote: str) -> Tuple[int,int]:
#     if not span or not quote: return span
#     s, e = span
#     if e <= s or s < 0 or e > len(script_text): return span
#     window = script_text[s:e]; win_norm = _normalize_keep_len(window).lower(); q_norm = _clean_quote_for_match(quote).lower()
#     if not q_norm: return span
#     i = win_norm.find(q_norm)
#     if i == -1:
#         m = re.search(re.escape(q_norm).replace(r"\ ", r"\s+"), win_norm, flags=re.IGNORECASE)
#         if not m: return span
#         i, j = m.start(), m.end()
#     else:
#         j = i + len(q_norm)
#     s2, e2 = s + i, s + j
#     s2, e2 = _snap_and_bridge_to_word(script_text, s2, e2); s2 = _heal_split_word_left(script_text, s2)
#     if s2 >= s and e2 <= e and e2 > s2: return (s2, e2)
#     return span

# def build_spans_by_param(script_text: str, data: dict, heading_ranges: Optional[List[Tuple[int,int]]] = None) -> Dict[str, List[Tuple[int,int,str,str]]]:
#     heading_ranges = heading_ranges or []
#     raw = (data or {}).get("per_parameter", {}) or {}
#     per: Dict[str, Dict[str, Any]] = {k:(v or {}) for k,v in raw.items()}
#     spans_map: Dict[str, List[Tuple[int,int,str,str]]] = {p: [] for p in PARAM_ORDER}
#     st.session_state["aoi_match_ranges"] = {}

#     for p in spans_map.keys():
#         color = PARAM_COLORS.get(p, "#ffd54f")
#         blk = per.get(p, {}) or {}
#         aois = blk.get("areas_of_improvement") or []
#         for idx, item in enumerate(aois, start=1):
#             raw_q = (item or {}).get("quote_verbatim", "") or ""
#             q = _sanitize_editor_text(raw_q)
#             clean = _clean_quote_for_match(re.sub(r"^[•\-\d\.\)\s]+", "", q).strip())
#             if not clean: continue
#             if _is_heading_like(clean): continue

#             pos = find_span_smart(script_text, clean)
#             if not pos: continue
#             pos = _tighten_to_quote(script_text, pos, raw_q)
#             s, e = pos

#             if heading_ranges and _overlaps_any(s, e, heading_ranges): continue
#             if _is_heading_context(script_text, s, e): continue

#             aid = f"{p.replace(' ','_')}-AOI-{idx}"
#             spans_map[p].append((s, e, color, aid))
#             st.session_state["aoi_match_ranges"][aid] = (s, e)
#     return spans_map

# # ---------- History (S3-aware) ----------
# def _maybe_copy_docx_to_history(source_docx_path: Optional[str], run_id: str) -> Optional[str]:
#     """Copy the DOCX used for rendering into outputs/_history so Recents can re-render identically."""
#     try:
#         if not source_docx_path:
#             return None
#         # If already an S3 key/url, just store the key in history blob
#         if source_docx_path.startswith("s3://") or (_s3_enabled() and not os.path.exists(source_docx_path)):
#             # ensure it's uploaded under our history key if it's a local temp in S3 mode
#             if os.path.exists(source_docx_path):
#                 with open(source_docx_path, "rb") as f:
#                     save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#                 return f"{HISTORY_DIR}/{run_id}.docx"
#             return source_docx_path

#         # Local file path
#         if _s3_enabled():
#             with open(source_docx_path, "rb") as f:
#                 save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#             return f"{HISTORY_DIR}/{run_id}.docx"
#         else:
#             dst = os.path.join(HISTORY_DIR, f"{run_id}.docx")
#             if os.path.abspath(source_docx_path) != os.path.abspath(dst):
#                 shutil.copyfile(source_docx_path, dst)
#             return dst
#     except Exception:
#         return None

# def _save_history_snapshot(title: str, data: dict, script_text: str,
#                            source_docx_path: Optional[str], heading_ranges: List[Tuple[int,int]],
#                            spans_by_param: Dict[str, List[Tuple[int,int,str,str]]],
#                            aoi_match_ranges: Dict[str, Tuple[int,int]]):
#     run_id = str(uuid.uuid4()); now = datetime.datetime.now()
#     created_at_iso = now.replace(microsecond=0).isoformat()
#     created_at_human = now.strftime("%Y-%m-%d %H:%M:%S")

#     # Copy stable DOCX alongside history JSON (so Recents uses identical renderer)
#     stable_docx_key_or_path = _maybe_copy_docx_to_history(source_docx_path, run_id)

#     blob = {
#         "run_id": run_id, "title": title or "untitled",
#         "created_at": created_at_iso, "created_at_human": created_at_human,
#         "overall_rating": (data or {}).get("overall_rating", ""),
#         "scores": (data or {}).get("scores", {}),
#         "data": data or {}, "script_text": script_text or "",
#         "source_docx_path": stable_docx_key_or_path or source_docx_path,
#         "heading_ranges": heading_ranges or [],
#         "spans_by_param": spans_by_param or {},
#         "aoi_match_ranges": aoi_match_ranges or {},
#     }

#     # Persist JSON (S3 or local)
#     out_name = f"{created_at_iso.replace(':','-')}__{run_id}.json"
#     out_key = f"{HISTORY_DIR}/{out_name}"
#     save_text_key(out_key, json.dumps(blob, ensure_ascii=False, indent=2))

# def _load_all_history() -> List[dict]:
#     out: List[dict] = []
#     if _s3_enabled():
#         keys = sorted(list_prefix(HISTORY_DIR), reverse=True)
#         for key in keys:
#             try:
#                 txt = read_text_key(key, "")
#                 if not txt:
#                     continue
#                 j = json.loads(txt)
#                 j["_key"] = key
#                 if not j.get("created_at_human") and j.get("created_at"):
#                     try:
#                         dt = datetime.datetime.fromisoformat(j["created_at"])
#                         j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#                     except Exception:
#                         pass
#                 out.append(j)
#             except Exception:
#                 continue
#     else:
#         for fp in sorted(glob.glob(os.path.join(HISTORY_DIR, "*.json"))):
#             try:
#                 with open(fp, "r", encoding="utf-8") as f:
#                     j = json.load(f)
#             except Exception:
#                 continue
#             j.setdefault("_path", fp)
#             ca = j.get("created_at")
#             try:
#                 if isinstance(ca, (int, float)):
#                     dt = datetime.datetime.utcfromtimestamp(float(ca))
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.astimezone().strftime("%Y-%m-%d %H:%M:%S")
#                 elif isinstance(ca, str) and ca: pass
#                 else:
#                     mtime = os.path.getmtime(fp); dt = datetime.datetime.fromtimestamp(mtime)
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#             except Exception:
#                 j["created_at"] = str(ca or "")
#             out.append(j)
#     out.sort(key=lambda r: r.get("created_at") or "", reverse=True)
#     return out

# def _open_history_run_by_id(run_id: str) -> bool:
#     """Open a history run by its run_id. Returns True if loaded."""
#     if not run_id: return False
#     recs = _load_all_history()
#     match = next((r for r in recs if r.get("run_id") == run_id), None)
#     if not match: return False

#     # Load JSON content again (S3/local), then set session
#     try:
#         if "_key" in match and _s3_enabled():
#             txt = read_text_key(match["_key"], "")
#             if not txt: return False
#             jj = json.loads(txt)
#         else:
#             path = match.get("_path")
#             if not path or not os.path.exists(path): return False
#             with open(path, "r", encoding="utf-8") as f:
#                 jj = json.load(f)
#     except Exception:
#         return False

#     st.session_state.script_text      = jj.get("script_text","")
#     st.session_state.base_stem        = jj.get("title","untitled")
#     st.session_state.data             = jj.get("data",{})
#     st.session_state.heading_ranges   = jj.get("heading_ranges",[])
#     st.session_state.spans_by_param   = jj.get("spans_by_param",{})
#     st.session_state.param_choice     = None
#     st.session_state.source_docx_path = jj.get("source_docx_path")
#     st.session_state.review_ready     = True
#     st.session_state["aoi_match_ranges"] = jj.get("aoi_match_ranges", {})
#     st.session_state.ui_mode          = "review"
#     return True

# def _render_recents_centerpane():
#     st.subheader("📄 Recents")
#     q = st.text_input("Filter by title…", "")
#     cols = st.columns([1, 4])
#     with cols[0]:
#         if st.button("← Back"):
#             st.session_state.ui_mode = "home"; st.rerun()

#     recs = _load_all_history()
#     ql = q.strip().lower()
#     if ql: recs = [r for r in recs if ql in (r.get("title","").lower())]
#     if not recs: st.caption("No history yet."); st.stop()

#     # Each card is a clickable <a class="rec-card" href="?open=<run_id>">
#     for rec in recs:
#         run_id = rec.get("run_id"); title = rec.get("title") or "(untitled)"
#         created_h = rec.get("created_at_human",""); overall = rec.get("overall_rating","")
#         st.markdown(f"""
#         <a class="rec-card" href="?open={run_id}">
#           <div class="rec-title">{title}</div>
#           <div class="rec-meta">{created_h}</div>
#           <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
#         </a>
#         """, unsafe_allow_html=True)

# # ---------- Sidebar ----------
# with st.sidebar:
#     if st.button("🆕 New review", use_container_width=True):
#         fp = st.session_state.get("flattened_docx_path")
#         if fp and os.path.exists(fp):
#             try: os.remove(fp)
#             except Exception: pass
#         for k in ["review_ready","script_text","base_stem","data","spans_by_param","param_choice",
#                   "source_docx_path","heading_ranges","flattened_docx_path","flatten_used"]:
#             st.session_state[k] = (
#                 False if k=="review_ready"
#                 else "" if k in ("script_text","base_stem")
#                 else {} if k=="spans_by_param"
#                 else [] if k=="heading_ranges"
#                 else None if k in ("source_docx_path","flattened_docx_path")
#                 else False if k=="flatten_used"
#                 else None
#             )
#         st.session_state.ui_mode = "home"
#         _clear_query_params()
#         st.rerun()

#     if st.button("📁 Recents", use_container_width=True):
#         st.session_state.ui_mode = "recents"
#         _clear_query_params()
#         st.rerun()

# # ---------- Input screen ----------
# def render_home():
#     st.subheader("🎬 Script Source")

#     (tab_upload,) = st.tabs(["Upload file"])
#     uploaded_file = None
#     uploaded_name = None
#     uploaded_key  = None

#     with tab_upload:
#         up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
#         if up is not None:
#             file_bytes = up.read()
#             suffix = os.path.splitext(up.name)[1].lower()
#             # Save to S3 (or local) under scripts/
#             uploaded_key = f"{SCRIPTS_DIR}/{up.name}"
#             save_bytes_key(uploaded_key, file_bytes)

#             # Also create a temp local copy for parsing
#             with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
#                 tmp.write(file_bytes)
#                 uploaded_file = tmp.name
#             uploaded_name = os.path.splitext(os.path.basename(up.name))[0] or "uploaded_script"

#     if st.button("🚀 Run Review", type="primary", use_container_width=True):
#         base_stem = "uploaded_script"
#         source_docx_path = None
#         heading_ranges = []

#         if uploaded_file:
#             base_stem = uploaded_name or "uploaded_script"
#             # Prefer key (so we can store it into history); for DOCX rendering we keep a tmp local
#             if uploaded_file.lower().endswith(".docx"):
#                 path_to_use = uploaded_file
#                 if _docx_contains_tables(path_to_use):
#                     flat = flatten_docx_tables_to_longtext(path_to_use)
#                     st.session_state.flattened_docx_path = flat
#                     st.session_state.flatten_used = True
#                     path_to_use = flat
#                 script_text, heading_ranges = build_docx_text_with_meta(path_to_use)

#                 # 🔑 IMPORTANT: set the rendering source to the *flattened* docx
#                 # (History saver will copy this file to Scriptmodel/outputs/_history/)
#                 source_docx_path = path_to_use
#             else:
#                 script_text = load_script_file(uploaded_file)
#                 source_docx_path = uploaded_key or uploaded_file  # keep reference to S3/local
#         else:
#             st.warning("Please upload a script first.")
#             st.stop()

#         if len(script_text.strip()) < 50:
#             st.error("Extracted text looks too short. Check your file extraction.")
#             st.stop()

#         with st.spinner("Running analysis…"):
#             try:
#                 review_text = run_review_multi(
#                     script_text=script_text,
#                     prompts_dir=PROMPTS_DIR,
#                     temperature=0.0
#                 )
#             finally:
#                 # Clean temp upload if we didn't keep it as source_docx_path local
#                 if uploaded_file and not (isinstance(source_docx_path, str) and os.path.exists(source_docx_path)):
#                     try:
#                         os.remove(uploaded_file)
#                     except Exception:
#                         pass

#         data = extract_review_json(review_text)
#         if not data:
#             st.error("JSON not detected in model output.")
#             st.stop()

#         st.session_state.script_text      = script_text
#         st.session_state.base_stem        = base_stem
#         st.session_state.data             = data
#         st.session_state.heading_ranges   = heading_ranges
#         st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
#         st.session_state.param_choice     = None
#         st.session_state.source_docx_path = source_docx_path
#         st.session_state.review_ready     = True
#         st.session_state.ui_mode          = "review"

#         _save_history_snapshot(
#             title=base_stem,
#             data=data,
#             script_text=script_text,
#             source_docx_path=source_docx_path,
#             heading_ranges=heading_ranges,
#             spans_by_param=st.session_state.spans_by_param,
#             aoi_match_ranges=st.session_state.get("aoi_match_ranges", {})
#         )

#         _clear_query_params()
#         st.rerun()

# # ---------- Results screen ----------
# def render_review():
#     script_text     = st.session_state.script_text
#     data            = st.session_state.data
#     spans_by_param  = st.session_state.spans_by_param
#     scores: Dict[str,int] = (data or {}).get("scores", {}) or {}
#     source_docx_path: Optional[str] = st.session_state.source_docx_path

#     # If our source_docx_path is an S3 key/url, ensure we have a local copy for rendering
#     # AFTER (prefer the flattened copy if we have it in-session)
#     docx_local: Optional[str] = None
#     preferred = st.session_state.get("flattened_docx_path") if st.session_state.get("flatten_used") else source_docx_path
#     if preferred:
#         if not os.path.exists(preferred):
#             docx_local = ensure_local_copy(preferred)
#         else:
#             docx_local = preferred

#     left, center, right = st.columns([1.1, 2.7, 1.4], gap="large")

#     with left:
#         st.subheader("Final score")
#         ordered = [p for p in PARAM_ORDER if p in scores]
#         df = pd.DataFrame({"Parameter": ordered, "Score (1–10)": [scores.get(p, "") for p in ordered]})
#         st.dataframe(df, hide_index=True, use_container_width=True)
#         st.markdown(f'**Overall:** {data.get("overall_rating","—")}/10')
#         st.divider()

#         strengths = (data or {}).get("strengths") or []
#         if not strengths:
#             per = (data or {}).get("per_parameter", {}) or {}
#             best = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
#             for name, sc in best:
#                 if sc >= 8 and name in per:
#                     exp = _sanitize_editor_text((per[name] or {}).get("explanation", "") or "")
#                     first = re.split(r"(?<=[.!?])\s+", exp.strip())[0] if exp else f"Consistently strong {name.lower()}."
#                     strengths.append(f"{name}: {first}")
#                 if len(strengths) >= 3: break

#         def _bullets(title: str, items):
#             st.markdown(f"**{title}**")
#             for s in (items or []):
#                 if isinstance(s, str) and s.strip(): st.write("• " + _sanitize_editor_text(s))
#             if not items: st.write("• —")

#         _bullets("Strengths", strengths)
#         _bullets("Weaknesses", data.get("weaknesses"))
#         _bullets("Suggestions", data.get("suggestions"))
#         _bullets("Drop-off Risks", data.get("drop_off_risks"))
#         st.markdown("**Viral Quotient**"); st.write(_sanitize_editor_text(data.get("viral_quotient","—")))

#     with right:
#         st.subheader("Parameters")
#         st.markdown('<div class="param-row">', unsafe_allow_html=True)
#         for p in [p for p in PARAM_ORDER if p in scores]:
#             if st.button(p, key=f"chip_{p}", help="Show inline AOI highlights for this parameter"):
#                 st.session_state.param_choice = p
#         st.markdown('</div>', unsafe_allow_html=True)

#         sel = st.session_state.param_choice
#         if sel:
#             blk = (data.get("per_parameter", {}) or {}).get(sel, {}) or {}
#             st.markdown(f"**{sel} — Score:** {scores.get(sel,'—')}/10")

#             if blk.get("explanation"):
#                 st.markdown("**Why this score**"); st.write(_sanitize_editor_text(blk["explanation"]))
#             if blk.get("weakness") and blk["weakness"] != "Not present":
#                 st.markdown("**Weakness**"); st.write(_sanitize_editor_text(blk["weakness"]))
#             if blk.get("suggestion") and blk["suggestion"] != "Not present":
#                 st.markdown("**Suggestion**"); st.write(_sanitize_editor_text(blk["suggestion"]))

#             aoi = blk.get("areas_of_improvement") or []
#             if aoi:
#                 st.markdown("**Areas of Improvement**")
#                 for i, item in enumerate(aoi, 1):
#                     popover_fn = getattr(st, "popover", None)
#                     aid = f"{sel.replace(' ','_')}-AOI-{i}"
#                     s_e_map = st.session_state.get("aoi_match_ranges", {})
#                     if aid in s_e_map:
#                         s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                         line = (matched_line if len(matched_line) <= 320 else matched_line[:300] + "…")
#                     else:
#                         line = _sanitize_editor_text(item.get('quote_verbatim',''))
#                     issue = _sanitize_editor_text(item.get('issue',''))
#                     fix   = _sanitize_editor_text(item.get('fix',''))
#                     why   = _sanitize_editor_text(item.get('why_this_helps',''))
#                     label = f"Issue {i}"
#                     if callable(popover_fn):
#                         with popover_fn(label):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)
#                     else:
#                         with st.expander(label, expanded=False):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)

#             if blk.get("summary"):
#                 st.markdown("**Summary**"); st.write(_sanitize_editor_text(blk["summary"]))

#     with center:
#         st.subheader("Script with inline highlights")
#         # choose spans for selected parameter (or all if None)
#         spans = st.session_state.spans_by_param.get(st.session_state.param_choice, []) if st.session_state.param_choice else []

#         # Build AOI payload for popup: { aid: {line, issue, fix, why} }
#         aoi_payload: Dict[str, Dict[str, str]] = {}
#         data_per = (data or {}).get("per_parameter") or {}
#         s_e_map = st.session_state.get("aoi_match_ranges", {})
#         sel = st.session_state.param_choice

#         def _mk_line(aid: str, fallback_q: str = "") -> str:
#             if aid in s_e_map:
#                 s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                 return matched_line if len(matched_line) <= 320 else matched_line[:300] + "…"
#             return _sanitize_editor_text(fallback_q or "")

#         def _collect(param_name: str):
#             blk = (data_per.get(param_name) or {})
#             for i, item in enumerate(blk.get("areas_of_improvement") or [], 1):
#                 aid = f"{param_name.replace(' ','_')}-AOI-{i}"
#                 aoi_payload[aid] = {
#                     "line": _mk_line(aid, (item or {}).get("quote_verbatim","")),
#                     "issue": _sanitize_editor_text((item or {}).get("issue","")),
#                     "fix": _sanitize_editor_text((item or {}).get("fix","")),
#                     "why": _sanitize_editor_text((item or {}).get("why_this_helps","")),
#                 }

#         if sel: _collect(sel)
#         else:
#             for pn in [p for p in PARAM_ORDER if p in data_per]:
#                 _collect(pn)

#         payload_json = json.dumps(aoi_payload, ensure_ascii=False)

#         frame_theme_css = """
#         <style>
#           :root{
#             --m7-surface: #eef2f7;
#             --m7-on-surface: #0f172a;
#             --m7-border: rgba(15,23,42,.14);
#           }
#           @media (prefers-color-scheme: dark){
#             :root{
#               --m7-surface: #2f333a;
#               --m7-on-surface: #ffffff;
#               --m7-border: rgba(255,255,255,.18);
#             }
#             body { background: transparent !important; }
#           }

#           .docxwrap{
#             background: var(--m7-surface);
#             color: var(--m7-on-surface);
#             border: 1px solid var(--m7-border);
#             border-radius: 12px;
#             padding: 16px 14px 18px;
#           }
#           .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#           .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); }
#         </style>
#         """

#         tooltip_css = """
#         <style>
#         .aoi-pop {
#           position: absolute; max-width: 520px; min-width: 320px;
#           background: var(--m7-surface); border: 1px solid var(--m7-border); border-radius: 10px;
#           box-shadow: 0 10px 25px rgba(0,0,0,.12);
#           padding: 12px 14px; z-index: 9999; transform: translateY(-8px);
#           color: var(--m7-on-surface);
#         }
#         .aoi-pop h4 { margin: 0 0 .35rem 0; font-size: .95rem; }
#         .aoi-pop p  { margin: .15rem 0; line-height: 1.5; }
#         .aoi-pop .muted { opacity:.85; font-size:.85rem; }
#         .aoi-arrow {
#           position:absolute; left:50%; transform:translateX(-50%);
#           bottom:-7px; width:0;height:0;border-left:7px solid transparent;
#           border-right:7px solid transparent;border-top:7px solid var(--m7-border);
#         }
#         .aoi-arrow::after{
#           content:""; position:absolute; left:-6px; top:-7px; width:0;height:0;
#           border-left:6px solid transparent;border-right:6px solid transparent;border-top:6px solid var(--m7-surface);
#         }
#         </style>
#         """

#         # Select rendering source (DOCX with highlights if we have a local path)
#         if docx_local and os.path.splitext(docx_local)[1].lower() == ".docx":
#             html_core = render_docx_html_with_highlights(
#                 docx_local,
#                 merge_overlaps_and_adjacent(script_text, spans)
#             )
#         else:
#             from html import escape as _esc
#             orig = script_text
#             spans2 = [s for s in merge_overlaps_and_adjacent(orig, spans) if s[0] < s[1]]
#             spans2.sort(key=lambda x: x[0])
#             cur = 0; buf: List[str] = []
#             for s,e,c,aid in spans2:
#                 if s > cur: buf.append(_esc(orig[cur:s]))
#                 buf.append(
#                     f'<mark class="aoi-mark" data-aid="{aid}" '
#                     f'style="background:{c}33;border:1px solid {c};border-radius:3px;padding:0 2px;">'
#                     f'{_esc(orig[s:e])}</mark>'
#                 )
#                 cur = e
#             if cur < len(orig): buf.append(_esc(orig[cur:]))
#             html_core = (
#                 '<div class="docxwrap"><p style="white-space:pre-wrap; '
#                 'line-height:1.7; font-family:ui-serif, Georgia, Times New Roman, serif;">'
#                 + "".join(buf) +
#                 '</p></div>'
#             )

#         # Popup + autosize JS shell
#         html_shell = """
# %%FRAME_THEME_CSS%%
# %%TOOLTIP_CSS%%
# <div id="m7-doc">%%HTML_CORE%%</div>
# <div id="aoi-pop" class="aoi-pop" style="display:none;">
#   <div id="aoi-pop-content"></div>
#   <div class="aoi-arrow"></div>
# </div>
# <script>
# (function(){
#   const AOI = __PAYLOAD__;
#   const wrap = document.getElementById('m7-doc');
#   const pop  = document.getElementById('aoi-pop');
#   const body = document.getElementById('aoi-pop-content');

#   function resizeIframe() {
#     try {
#       const h = Math.max(
#         document.documentElement.scrollHeight,
#         document.body.scrollHeight
#       );
#       if (window.frameElement) {
#         window.frameElement.style.height = (h + 20) + 'px';
#         window.frameElement.style.width  = '100%';
#       }
#     } catch(e) {}
#   }
#   window.addEventListener('load', resizeIframe);
#   window.addEventListener('resize', resizeIframe);

#   function hide(){ pop.style.display='none'; }
#   function showFor(mark){
#     const aid = mark.getAttribute('data-aid');
#     const d = AOI[aid]; if(!d) return;
#     body.innerHTML =
#       (d.line  ? '<p><strong>Line:</strong> '  + d.line  + '</p>' : '') +
#       (d.issue ? '<p><strong>Issue:</strong> ' + d.issue + '</p>' : '') +
#       (d.fix   ? '<p><strong>Fix:</strong> '   + d.fix   + '</p>' : '') +
#       (d.why   ? '<p class="muted">'           + d.why   + '</p>' : '');
#     pop.style.display = 'block';

#     const r = mark.getBoundingClientRect();
#     const scY = window.scrollY || document.documentElement.scrollTop;
#     const scX = window.scrollX || document.documentElement.scrollLeft;
#     let top  = r.top + scY - pop.offsetHeight - 10;
#     let left = r.left + scX + r.width/2 - pop.offsetWidth/2;
#     if (top < 8) top = r.bottom + scY + 10;
#     if (left < 8) left = 8;
#     pop.style.top  = top + 'px';
#     pop.style.left = left + 'px';

#     resizeIframe();
#   }

#   wrap.addEventListener('click', (e) => {
#     const m = e.target.closest('.aoi-mark');
#     if(!m){ hide(); return; }
#     if(pop.style.display === 'block'){ hide(); }
#     showFor(m);
#     e.stopPropagation();
#   });

#   document.addEventListener('click', (e) => {
#     if(!e.target.closest('.aoi-pop') && !e.target.closest('.aoi-mark')) hide();
#   });
# })();
# </script>
# """
#         html_shell = (
#             html_shell
#             .replace("%%FRAME_THEME_CSS%%", frame_theme_css)
#             .replace("%%TOOLTIP_CSS%%", tooltip_css)
#             .replace("%%HTML_CORE%%", html_core)
#             .replace("__PAYLOAD__", payload_json)
#         )

#         components.html(html_shell, height=400, scrolling=False)

# # ---------- Router & query param open ----------
# _open_qp = _get_query_param("open")
# if _open_qp and _open_history_run_by_id(_open_qp):
#     _clear_query_params()  # avoid re-opening on subsequent reruns

# mode = st.session_state.ui_mode
# if mode == "recents":
#     _render_recents_centerpane()
# elif mode == "review" and st.session_state.review_ready:
#     render_review()
# else:
#     render_home()











############################################################



























# import os, re, glob, json, tempfile, difflib, uuid, datetime, shutil
# from pathlib import Path
# from typing import Dict, Any, List, Tuple, Optional

# import streamlit as st
# import pandas as pd
# import streamlit.components.v1 as components  # for inline HTML/JS popup

# # ---- utils & engine ----
# from utils1 import (
#     extract_review_json,
#     PARAM_ORDER,
#     load_script_file,
#     extract_left_column_script_or_default,  # <-- NEW import for left-column extractor
# )
# from review_engine_multi import run_review_multi

# # ---- DOCX rendering imports ----
# from docx import Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # =========================
# # RunPod S3 (inline helpers) — FIXED
# # =========================
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError

# # Accept both AWS_* and RUNPOD_* style envs / st.secrets
# def _get_env(key: str, default: str = "") -> str:
#     v = os.getenv(key, "")
#     if v: 
#         return v.strip()
#     try:
#         v2 = st.secrets.get(key)
#         if isinstance(v2, str):
#             return v2.strip()
#     except Exception:
#         pass
#     return (default or "").strip()

# # Primary config
# _RP_ENDPOINT = _get_env("RUNPOD_S3_ENDPOINT")
# _RP_BUCKET   = _get_env("RUNPOD_S3_BUCKET")
# _RP_REGION   = _get_env("RUNPOD_S3_REGION") or _get_env("AWS_DEFAULT_REGION") or ""

# # Credentials: prefer AWS_* if present; else accept RUNPOD_* fallbacks
# _AK = _get_env("AWS_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY")
# _SK = _get_env("AWS_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_KEY")
# _ST = _get_env("AWS_SESSION_TOKEN")  # optional

# # Options
# _FORCE_PATH = (_get_env("RUNPOD_S3_FORCE_PATH_STYLE") or "true").lower() in {"1","true","yes"}
# _USE_SSL    = (_get_env("RUNPOD_S3_USE_SSL") or "true").lower() in {"1","true","yes"}
# _VERIFY_SSL = (_get_env("RUNPOD_S3_VERIFY_SSL") or "true").lower() in {"1","true","yes"}

# def _s3_enabled() -> bool:
#     return bool(_RP_ENDPOINT and _RP_BUCKET and _AK and _SK)

# @st.cache_resource(show_spinner=False)
# def _s3_client():
#     if not _s3_enabled():
#         return None
#     session_kwargs = dict(
#         aws_access_key_id=_AK,
#         aws_secret_access_key=_SK,
#     )
#     if _ST:
#         session_kwargs["aws_session_token"] = _ST

#     # s3v4 + path style are common requirements for S3-compatible services
#     cfg = Config(
#         signature_version="s3v4",
#         s3={"addressing_style": "path" if _FORCE_PATH else "auto"},
#         retries={"max_attempts": 3, "mode": "standard"}
#     )
#     return boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         region_name=_RP_REGION or None,
#         use_ssl=_USE_SSL,
#         verify=_VERIFY_SSL,
#         config=cfg,
#         **session_kwargs,
#     )

# def save_text_key(key: str, text: str) -> str:
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "w", encoding="utf-8") as f:
#             f.write(text)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=text.encode("utf-8"))
#     return f"s3://{_RP_BUCKET}/{key}"

# def save_bytes_key(key: str, data: bytes) -> str:
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "wb") as f:
#             f.write(data)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=data)
#     return f"s3://{_RP_BUCKET}/{key}"

# def read_text_key(key: str, default: str = "") -> str:
#     if not _s3_enabled():
#         try:
#             with open(key, "r", encoding="utf-8") as f:
#                 return f.read()
#         except Exception:
#             return default
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read().decode("utf-8", errors="ignore")
#     except Exception:
#         return default

# def read_bytes_key(key: str) -> Optional[bytes]:
#     if not _s3_enabled():
#         try:
#             with open(key, "rb") as f:
#                 return f.read()
#         except Exception:
#             return None
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read()
#     except Exception:
#         return None

# def list_prefix(prefix: str) -> List[str]:
#     """
#     List object keys under prefix (or local dir paths if not S3).
#     In S3 mode we always return KEYS (not URLs).
#     """
#     if not _s3_enabled():
#         base = prefix if os.path.isdir(prefix) else os.path.dirname(prefix)
#         try:
#             return [os.path.join(base, p) for p in os.listdir(base) if p.endswith(".json")]
#         except Exception:
#             return []

#     out: List[str] = []
#     token = None
#     # Normalize to "dir/" prefix for S3 listing
#     s3_prefix = prefix.rstrip("/") + "/"
#     try:
#         while True:
#             kwargs = {"Bucket": _RP_BUCKET, "Prefix": s3_prefix}
#             if token:
#                 kwargs["ContinuationToken"] = token
#             resp = _s3_client().list_objects_v2(**kwargs)
#             for c in resp.get("Contents", []):
#                 k = c.get("Key", "")
#                 if k.endswith(".json"):
#                     out.append(k)
#             token = resp.get("NextContinuationToken")
#             if not token:
#                 break
#     except (ClientError, EndpointConnectionError, NoCredentialsError):
#         return []
#     return out

# def presigned_url(key: str, expires: int = 3600) -> Optional[str]:
#     if not _s3_enabled():
#         return None
#     try:
#         return _s3_client().generate_presigned_url(
#             "get_object",
#             Params={"Bucket": _RP_BUCKET, "Key": key},
#             ExpiresIn=expires
#         )
#     except ClientError:
#         return None

# def ensure_local_copy(key_or_path: str) -> Optional[str]:
#     """
#     For DOCX/PDF parsing we need a real filesystem path.
#     If S3 mode, download to a temp file and return that path.
#     """
#     if not _s3_enabled():
#         return key_or_path if os.path.exists(key_or_path) else None

#     key = key_or_path
#     if key.startswith("s3://"):
#         # s3://bucket/path/to/file -> path/to/file
#         parts = key.split("/", 3)
#         key = parts[3] if len(parts) >= 4 else ""
#     data = read_bytes_key(key)
#     if data is None:
#         return None
#     fd, tmp = tempfile.mkstemp(suffix=os.path.splitext(key)[1] or "")
#     os.close(fd)
#     with open(tmp, "wb") as f:
#         f.write(data)
#     return tmp

# def _s3_health_summary() -> dict:
#     """
#     Optional tiny health read you can print if needed.
#     Returns a dict; safe to ignore in production.
#     """
#     info = {
#         "enabled": _s3_enabled(),
#         "endpoint": _RP_ENDPOINT,
#         "bucket": _RP_BUCKET,
#         "region": _RP_REGION,
#         "has_keys": bool(_AK and _SK),
#     }
#     if not _s3_enabled():
#         info["status"] = "local-mode"
#         return info
#     try:
#         # Attempt a very cheap list; no exceptions => reachable
#         _ = _s3_client().list_objects_v2(Bucket=_RP_BUCKET, Prefix=(f"{OUTPUT_DIR}/_history/").rstrip("/") + "/",
#                                          MaxKeys=1)
#         info["status"] = "ok"
#     except Exception as e:
#         info["status"] = f"error: {getattr(e, 'response', {}).get('Error', {}).get('Code', str(e))}"
#     return info


# # ---------- Folders ----------
# # SCRIPTS_DIR = "scripts"
# # PROMPTS_DIR = "prompts"
# # OUTPUT_DIR  = "outputs"
# # HISTORY_DIR = os.path.join(OUTPUT_DIR, "_history")
# # ---------- Folders (all under Scriptmodel/) ----------

# BASE_PREFIX = "Scriptmodel"

# SCRIPTS_DIR = f"{BASE_PREFIX}/scripts"
# PROMPTS_DIR = f"{BASE_PREFIX}/prompts"
# OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
# HISTORY_DIR = f"{OUTPUT_DIR}/_history"

# # Fix: the previous line had an extra '}' in OUTPUT_DIR; correct it:
# OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
# HISTORY_DIR = f"{OUTPUT_DIR}/_history"

# if not _s3_enabled():
#     for p in (SCRIPTS_DIR, PROMPTS_DIR, OUTPUT_DIR, HISTORY_DIR):
#         Path(p).mkdir(parents=True, exist_ok=True)


# # ---------- Colors ----------
# PARAM_COLORS: Dict[str, str] = {
#     "Suspense Building":              "#ff6b6b",
#     "Language/Tone":                  "#6b8cff",
#     "Intro + Main Hook/Cliffhanger":  "#ffb86b",
#     "Story Structure + Flow":         "#a78bfa",
#     "Pacing":                         "#f43f5e",
#     "Mini-Hooks (30–60s)":            "#eab308",
#     "Outro (Ending)":                 "#8b5cf6",
#     "Grammar & Spelling":             "#10b981",  # NEW: distinct highlight color for grammar/spelling
# }

# # ---------- Config ----------
# STRICT_MATCH_ONLY = False

# # ---------- App config ----------
# st.set_page_config(page_title="M7 — Grammarly UI", page_icon="🕵️", layout="wide")
# # ---------- Header patch & CSS ----------
# def render_app_title():
#     st.markdown(
#         '<h1 class="app-title">Viral Script Reviewer</h1>',
#         unsafe_allow_html=True
#     )
#     st.markdown("""
#     <style>
#     html { color-scheme: light dark; }

#     :root{
#       --m7-surface: #eef2f7;
#       --m7-on-surface: #0f172a;
#       --m7-border: rgba(15,23,42,.14);
#       --sep: #e5e7eb;
#     }
#     @media (prefers-color-scheme: dark){
#       :root{
#         --m7-surface: #2f333a;
#         --m7-on-surface: #ffffff;
#         --m7-border: rgba(255,255,255,.18);
#         --sep: #2a2f37;
#       }
#     }

#     .stApp .block-container { padding-top: 4.25rem !important; }
#     .app-title{
#       font-weight: 700; font-size: 2.1rem; line-height: 1.3;
#       margin: 0 0 1rem 0; padding-left: 40px !important; padding-top: .25rem !important;
#       white-space: normal; word-break: break-word; hyphens: auto; overflow: visible;
#       position: relative !important; z-index: 10 !important;
#     }
#     [data-testid="collapsedControl"] { z-index: 6 !important; }
#     header[data-testid="stHeader"], .stAppHeader { background: transparent !important; box-shadow: none !important; }
#     @media (min-width: 992px){ .app-title { padding-left: 0 !important; } }

#     div[data-testid="column"]:nth-of-type(1){position:relative;}
#     div[data-testid="column"]:nth-of-type(1)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     div[data-testid="column"]:nth-of-type(2){position:relative;}
#     div[data-testid="column"]:nth-of-type(2)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}

#     .m7-card{
#       background: var(--m7-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 14px 16px;
#       color: var(--m7-on-surface);
#     }
#     .m7-card, .m7-card * { color: var(--m7-on-surface) !important; }

#     .docxwrap{
#       background: var(--m7-surface);
#       color: var(--m7-on-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 16px 14px 18px;
#     }
#     .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#     .docxwrap .h1, .docxwrap .h2, .docxwrap .h3 { font-weight:700; margin:10px 0 6px; }
#     .docxwrap .h1 { font-size: 1.3rem; border-bottom: 2px solid currentColor; padding-bottom: 4px; }
#     .docxwrap .h2 { font-size: 1.15rem; border-bottom: 1px solid currentColor; padding-bottom: 3px; }
#     .docxwrap .h3 { font-size: 1.05rem; }
#     .docxwrap p { margin: 10px 0; line-height: 1.7; font-family: ui-serif, Georgia, "Times New Roman", serif; }
#     .docxwrap table { border-collapse: collapse; width: 100%; margin: 12px 0; }
#     .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); padding:8px; vertical-align:top; line-height:1.6; }
#     .docxwrap mark{ padding:0 2px; border-radius:3px; border:1px solid var(--m7-border); cursor: pointer; }

#     .rec-card{
#       display:block; text-decoration:none !important;
#       background: var(--m7-surface);
#       border:1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 14px 16px; margin: 10px 0 16px;
#       box-shadow: 0 1px 2px rgba(0,0,0,.06);
#       color: var(--m7-on-surface) !important;
#       transition: filter .1s ease, transform .02s ease;
#     }
#     .rec-card:hover{ filter: brightness(1.02); }
#     .rec-card:active{ transform: translateY(1px); }
#     .rec-card, .rec-card * { color: var(--m7-on-surface) !important; }
#     .rec-title{font-weight:600; margin-bottom:.25rem;}
#     .rec-meta{opacity:.85 !important; font-size:12.5px; margin-bottom:.4rem;}
#     .rec-row{display:flex; align-items:center; justify-content:space-between; gap:12px;}

#     .stTextInput>div>div,
#     .stTextArea>div>div,
#     .stNumberInput>div>div,
#     .stDateInput>div>div,
#     .stTimeInput>div>div,
#     .stFileUploader>div,
#     div[data-baseweb="select"]{
#       background: var(--m7-surface) !important;
#       border: 1px solid var(--m7-border) !important;
#       border-radius: 10px !important;
#       color: var(--m7-on-surface) !important;
#     }
#     .stTextInput input,
#     .stTextArea textarea,
#     .stNumberInput input,
#     .stDateInput input,
#     .stTimeInput input,
#     .stFileUploader div,
#     div[data-baseweb="select"] *{
#       color: var(--m7-on-surface) !important;
#     }
#     .stTextInput input::placeholder,
#     .stTextArea textarea::placeholder{ color: rgba(16,24,39,.55) !important; }
#     @media (prefers-color-scheme: dark){
#       .stTextInput input::placeholder,
#       .stTextArea textarea::placeholder{ color: rgba(255,255,255,.75) !important; }
#     }
#     div[data-testid="stFileUploaderDropzone"] label span { color: var(--m7-on-surface) !important; opacity: 1 !important; }
#     div[data-testid="stFileUploaderDropzone"] label { color: var(--m7-on-surface) !important; }

#     .stMarkdown pre,
#     pre[class*="language-"],
#     .stCodeBlock{
#       background: var(--m7-surface) !important;
#       color: var(--m7-on-surface) !important;
#       border: 1px solid var(--m7-border) !important;
#       border-radius: 12px !important;
#       padding: 12px 14px !important;
#       overflow:auto;
#     }
#     .stMarkdown pre code{ background: transparent !important; color: inherit !important; }

#     div[data-testid="stDataFrame"]{
#       background: var(--m7-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 6px 8px;
#       color: var(--m7-on-surface);
#     }
#     div[data-testid="stDataFrame"] * { color: var(--m7-on-surface) !important; }

#     </style>
#     """, unsafe_allow_html=True)

# render_app_title()

# # ---------- Session ----------
# for key, default in [
#     ("review_ready", False),
#     ("script_text", ""),
#     ("base_stem", ""),
#     ("data", None),
#     ("spans_by_param", {}),
#     ("param_choice", None),
#     ("source_docx_path", None),
#     ("heading_ranges", []),
#     ("flattened_docx_path", None),
#     ("flatten_used", False),
#     ("ui_mode", "home"),
# ]:
#     st.session_state.setdefault(key, default)

# # ---------- helpers for query params (compat across Streamlit versions) ----------
# def _get_query_param(key: str) -> Optional[str]:
#     val = None
#     try:
#         val = st.query_params.get(key)
#     except Exception:
#         q = st.experimental_get_query_params()
#         v = q.get(key)
#         if isinstance(v, list): val = v[0] if v else None
#         else: val = v
#     return val

# def _clear_query_params():
#     try:
#         st.query_params.clear()
#     except Exception:
#         st.experimental_set_query_params()

# # ---------- Sanitizer ----------
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )
# def _sanitize_editor_text(s: Optional[str]) -> str:
#     if not s: return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•)\s*', '', t, flags=re.M)
#     t = re.sub(r'^\s*[-*]\s+', '• ', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# # ---------- DOCX traversal ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# # ---------- Auto-flatten ----------
# def _docx_contains_tables(path: str) -> bool:
#     doc = Document(path)
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Table):
#             return True
#     return False

# def _copy_paragraph(dest_doc: Document, src_para: Paragraph):
#     p = dest_doc.add_paragraph()
#     try:
#         if src_para.style and src_para.style.name:
#             p.style = src_para.style.name
#     except Exception:
#         pass
#     for run in src_para.runs:
#         r = p.add_run(run.text or "")
#         r.bold = run.bold
#         r.italic = run.italic
#         r.underline = run.underline
#     return p

# def flatten_docx_tables_to_longtext(source_path: str) -> str:
#     src = Document(source_path)
#     new = Document()
#     for blk in _iter_docx_blocks(src):
#         if isinstance(blk, Paragraph):
#             _copy_paragraph(new, blk)
#         else:
#             seen_tc_ids = set()
#             for row in blk.rows:
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     if tc_id in seen_tc_ids:
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     for p in cell.paragraphs:
#                         _copy_paragraph(new, p)
#                 new.add_paragraph("")
#             new.add_paragraph("")
#     fd, tmp_path = tempfile.mkstemp(suffix=".docx"); os.close(fd); new.save(tmp_path)
#     return tmp_path

# # ---------- Build plain text + heading ranges ----------
# def build_docx_text_with_meta(docx_path: str) -> Tuple[str, List[Tuple[int,int]]]:
#     doc = Document(docx_path)
#     out: List[str] = []
#     heading_ranges: List[Tuple[int,int]] = []
#     current_offset = 0

#     def _append_and_advance(s: str):
#         nonlocal current_offset
#         out.append(s); current_offset += len(s)

#     seen_tc_ids: set = set()
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             para_text = "".join(run.text or "" for run in blk.runs)
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             if sty.startswith("heading"):
#                 start = current_offset; end = start + len(para_text)
#                 heading_ranges.append((start, end))
#             _append_and_advance(para_text); _append_and_advance("\n")
#         else:
#             for row in blk.rows:
#                 row_cell_tcs = []
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     row_cell_tcs.append((tc_id, cell))
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     if tc_id in seen_tc_ids:
#                         if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     cell_text_parts: List[str] = []
#                     for i, p in enumerate(cell.paragraphs):
#                         t = "".join(r.text or "" for r in p.runs)
#                         sty = (p.style.name or "").lower() if p.style else ""
#                         if sty.startswith("heading"):
#                             hs = current_offset + sum(len(x) for x in cell_text_parts)
#                             he = hs + len(t)
#                             heading_ranges.append((hs, he))
#                         cell_text_parts.append(t)
#                         if i != len(cell.paragraphs) - 1:
#                             cell_text_parts.append("\n")
#                     cell_text = "".join(cell_text_parts)
#                     _append_and_advance(cell_text)
#                     if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                 _append_and_advance("\n")
#             _append_and_advance("\n")

#     return "".join(out), heading_ranges

# def _wrap_inline(safe_text: str, run) -> str:
#     out = safe_text
#     if getattr(run, "underline", False): out = f"<u>{out}</u>"
#     if getattr(run, "italic", False): out = f"<em>{out}</em>"
#     if getattr(run, "bold", False): out = f"<strong>{out}</strong>"
#     return out

# # ---------- Invisible/bridge chars ----------
# _BRIDGE_CHARS = set("\u200b\u200c\u200d\u2060\ufeff\xa0\u00ad")

# # ---------- DOCX -> HTML with highlights (includes data-aid) ----------
# def render_docx_html_with_highlights(docx_path: str,
#                                      highlight_spans: List[Tuple[int,int,str,str]]) -> str:
#     """
#     highlight_spans: list of (start, end, color, aid)
#     """
#     doc = Document(docx_path)
#     spans = [s for s in highlight_spans if s[0] < s[1]]
#     spans.sort(key=lambda x: x[0])

#     cur_span = 0
#     current_offset = 0

#     def esc(s: str) -> str:
#         return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

#     def open_mark_if_needed(html_parts, mark_state, color, end, aid):
#         if not mark_state["open"]:
#             html_parts.append(
#                 f'<mark class="aoi-mark" data-aid="{aid}" '
#                 f'style="background:{color}33;border:1px solid {color};border-radius:3px;padding:0 2px;">'
#             )
#             mark_state.update(open=True, end=end, color=color, aid=aid)

#     def close_mark_if_open(html_parts, mark_state):
#         if mark_state["open"]:
#             html_parts.append('</mark>')
#             mark_state.update(open=False, end=None, color=None, aid=None)

#     def emit_run_text(run_text: str, run, html_parts: List[str], mark_state: Dict[str, Any]):
#         nonlocal cur_span, current_offset
#         t = run_text or ""; i = 0
#         while i < len(t):
#             next_start, next_end, color, next_aid = None, None, None, None
#             if cur_span < len(spans):
#                 next_start, next_end, color, next_aid = spans[cur_span]

#             if not mark_state["open"]:
#                 if cur_span >= len(spans) or current_offset + (len(t) - i) <= next_start:
#                     chunk = t[i:]; html_parts.append(_wrap_inline(esc(chunk), run)); current_offset += len(chunk); break
#                 if current_offset < next_start:
#                     take = next_start - current_offset
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take; continue
#                 open_mark_if_needed(html_parts, mark_state, color, next_end, next_aid)
#             else:
#                 take = min(mark_state["end"] - current_offset, len(t) - i)
#                 if take > 0:
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take
#                 if current_offset >= mark_state["end"]:
#                     close_mark_if_open(html_parts, mark_state)
#                     cur_span += 1

#     html: List[str] = ['<div class="docxwrap">']
#     seen_tc_ids: set = set()

#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             open_tag = '<div class="h1">' if sty.startswith("heading 1") else \
#                        '<div class="h2">' if sty.startswith("heading 2") else \
#                        '<div class="h3">' if sty.startswith("heading 3") else "<p>"
#             close_tag = "</div>" if sty.startswith("heading") else "</p>"
#             html.append(open_tag)
#             for run in blk.runs:
#                 emit_run_text(run.text or "", run, html, mark_state)
#             close_mark_if_open(html, mark_state)
#             html.append(close_tag)
#             current_offset += 1  # '\n'

#         else:
#             html.append("<table>")
#             for row in blk.rows:
#                 html.append("<tr>")
#                 row_cell_tcs = [(id(cell._tc), cell) for cell in row.cells]
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     html.append("<td>")
#                     if tc_id not in seen_tc_ids:
#                         seen_tc_ids.add(tc_id)
#                         for p_idx, p in enumerate(cell.paragraphs):
#                             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#                             html.append("<div>")
#                             for run in p.runs:
#                                 emit_run_text(run.text or "", run, html, mark_state)
#                             close_mark_if_open(html, mark_state)
#                             html.append("</div>")
#                             if p_idx != len(cell.paragraphs) - 1:
#                                 current_offset += 1
#                     html.append("</td>")
#                     if idx != len(row_cell_tcs) - 1: current_offset += 1  # '\t'
#                 html.append("</tr>"); current_offset += 1   # row '\n'
#             html.append("</table>"); current_offset += 1     # extra '\n'

#     html.append("</div>")
#     return "".join(html)

# # ---------- Matching / span utilities ----------
# def _normalize_keep_len(s: str) -> str:
#     trans = {
#         "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
#         "\u2013": "-", "\u2014": "-",
#         "\xa0": " ",
#         "\u200b": " ", "\u200c": " ", "\u200d": " ", "\u2060": " ",
#         "\ufeff": " ", "\u00ad": " ",
#     }
#     return (s or "").translate(str.maketrans(trans))

# def _tokenize(s: str) -> List[str]: return re.findall(r"\w+", (s or "").lower())

# def _iter_sentences_with_spans(text: str) -> List[Tuple[int,int,str]]:
#     spans = []
#     for m in re.finditer(r'[^.!?]+[.!?]+|\Z', text, flags=re.S):
#         s, e = m.start(), m.end()
#         seg = text[s:e]
#         if seg.strip(): spans.append((s, e, seg))
#     return spans

# def _squash_ws(s: str) -> str: return re.sub(r"\s+", " ", s or "").strip()

# def _clean_quote_for_match(q: str) -> str:
#     if not q: return ""
#     q = _normalize_keep_len(q).strip()
#     q = re.sub(r'^[\'"“”‘’\[\(\{<…\-\–\—\s]+', '', q)
#     q = re.sub(r'[\'"“”‘’\]\)\}>…\-\–\—\s]+$', '', q)
#     return _squash_ws(q)

# def _snap_and_bridge_to_word(text: str, start: int, end: int, max_bridge: int = 2) -> Tuple[int,int]:
#     n = len(text); s, e = max(0,start), max(start,end)
#     def _is_inv(ch: str) -> bool: return ch in _BRIDGE_CHARS

#     while s > 0:
#         prev = text[s-1]; cur = text[s] if s < n else ""
#         if prev.isalnum() and cur.isalnum(): s -= 1; continue
#         j = s; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1; 
#         if brid and (s-1)>=0 and text[s-1].isalnum() and (j<n and text[j].isalnum()): s -= 1; continue
#         break

#     while e < n:
#         prev = text[e-1] if e>0 else ""; nxt = text[e]
#         if prev.isalnum() and nxt.isalnum(): e += 1; continue
#         j = e; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (e-1)>=0 and text[e-1].isalnum() and (j<n and text[j].isalnum()): e = j + 1; continue
#         break

#     while e < n and text[e] in ',"”’\')]}': e += 1
#     return s, e

# def _heal_split_word_left(text: str, start: int) -> int:
#     i = start
#     if i <= 1 or i >= len(text): return start
#     if text[i-1] != " ": return start
#     j = i - 2
#     while j >= 0 and text[j].isalpha(): j -= 1
#     prev_token = text[j+1:i-1]
#     if len(prev_token) == 1: return i - 2
#     return start

# def _overlaps_any(s: int, e: int, ranges: List[Tuple[int,int]]) -> bool:
#     for rs, re_ in ranges:
#         if e > rs and s < re_: return True
#     return False

# def _fuzzy_window_span(tl: str, nl: str, start: int, w: int) -> Tuple[float, Optional[Tuple[int,int]]]:
#     window = tl[start:start+w]
#     sm = difflib.SequenceMatcher(a=nl, b=window)
#     blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
#     if not blocks: return 0.0, None
#     coverage = sum(b.size for b in blocks) / max(1, len(nl))
#     first_b = min(blocks, key=lambda b: b.b); last_b = max(blocks, key=lambda b: b.b + b.size)
#     s = start + first_b.b; e = start + last_b.b + last_b.size
#     return coverage, (s, e)

# def find_span_smart(text: str, needle: str) -> Optional[Tuple[int,int]]:
#     if not text or not needle: return None
#     t_orig = text; t_norm = _normalize_keep_len(text); n_norm = _clean_quote_for_match(needle)
#     if not n_norm: return None
#     tl = t_norm.lower(); nl = n_norm.lower()

#     i = tl.find(nl)
#     if i != -1:
#         s, e = _snap_and_bridge_to_word(t_orig, i, i + len(nl)); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     m = re.search(re.escape(nl).replace(r"\ ", r"\s+"), tl, flags=re.IGNORECASE)
#     if m:
#         s, e = _snap_and_bridge_to_word(t_orig, m.start(), m.end()); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     if not STRICT_MATCH_ONLY and len(nl) >= 12:
#         w = max(60, min(240, len(nl) + 80))
#         best_cov, best_span = 0.0, None
#         step = max(1, w // 2)
#         for start in range(0, max(1, len(tl) - w + 1), step):
#             cov, se = _fuzzy_window_span(tl, nl, start, w)
#             if cov > best_cov: best_cov, best_span = cov, se
#         if best_span and best_cov >= 0.65:
#             s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
#             if s > 0 and t_orig[s-1:s+1].lower() in {" the", " a", " an"}: s -= 1
#             s = _heal_split_word_left(t_orig, s)
#             return (s, e)

#     if not STRICT_MATCH_ONLY:
#         keys = [w for w in _tokenize(nl) if len(w) >= 4][:8]
#         if len(keys) >= 2:
#             kset = set(keys)
#             best_score, best_span = 0.0, None
#             for s, e, seg in _iter_sentences_with_spans(t_norm):
#                 toks = set(_tokenize(seg)); ov = len(kset & toks)
#                 if ov == 0: continue
#                 score = ov / max(2, len(kset)); length_pen = min(1.0, 120 / max(20, e - s)); score *= (0.6 + 0.4 * length_pen)
#                 if score > best_score: best_score, best_span = score, (s, min(e, s + 400))
#             if best_span and best_score >= 0.35:
#                 s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1]); s = _heal_split_word_left(t_orig, s)
#                 return (s, e)
#     return None

# def merge_overlaps(spans: List[Tuple[int,int,str,str]]) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans.sort(key=lambda x: x[0]); out = [spans[0]]
#     for s,e,c,aid in spans[1:]:
#         ps,pe,pc,paid = out[-1]
#         if s <= pe and pc == c and e > pe: out[-1] = (ps, e, pc, paid)
#         else: out.append((s,e,c,aid))
#     return out

# _PUNCT_WS = set(" \t\r\n,.;:!?)('\"[]{}-—–…") | _BRIDGE_CHARS

# def merge_overlaps_and_adjacent(base_text: str,
#                                 spans: List[Tuple[int,int,str,str]],
#                                 max_gap: int = 2) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans = sorted(spans, key=lambda x: x[0]); out = [spans[0]]
#     for s, e, c, aid in spans[1:]:
#         ps, pe, pc, paid = out[-1]
#         if c == pc and s <= pe: out[-1] = (ps, max(pe, e), pc, paid); continue
#         if c == pc and s - pe <= max_gap:
#             gap = base_text[max(0, pe):max(0, s)]
#             if all((ch in _PUNCT_WS) for ch in gap): out[-1] = (ps, e, pc, paid); continue
#         out.append((s, e, c, aid))
#     return out

# # ---------- Heading filters ----------
# def _is_heading_like(q: str) -> bool:
#     if not q: return True
#     s = q.strip()
#     if not re.search(r'[.!?]', s):
#         words = re.findall(r"[A-Za-z]+", s)
#         if 1 <= len(words) <= 7:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.8: return True
#         if s.lower() in {"introduction","voiceover","outro","epilogue","prologue","credits","title","hook","horrifying discovery","final decision"}: return True
#         if len(s) <= 3: return True
#     return False

# def _is_heading_context(script_text: str, s: int, e: int) -> bool:
#     left = script_text.rfind("\n", 0, s) + 1
#     right = script_text.find("\n", e); right = len(script_text) if right == -1 else right
#     line = script_text[left:right].strip()
#     if len(line) <= 70 and not re.search(r'[.!?]', line):
#         words = re.findall(r"[A-Za-z]+", line)
#         if 1 <= len(words) <= 8:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.7: return True
#     return False

# def _tighten_to_quote(script_text: str, span: Tuple[int,int], quote: str) -> Tuple[int,int]:
#     if not span or not quote: return span
#     s, e = span
#     if e <= s or s < 0 or e > len(script_text): return span
#     window = script_text[s:e]; win_norm = _normalize_keep_len(window).lower(); q_norm = _clean_quote_for_match(quote).lower()
#     if not q_norm: return span
#     i = win_norm.find(q_norm)
#     if i == -1:
#         m = re.search(re.escape(q_norm).replace(r"\ ", r"\s+"), win_norm, flags=re.IGNORECASE)
#         if not m: return span
#         i, j = m.start(), m.end()
#     else:
#         j = i + len(q_norm)
#     s2, e2 = s + i, s + j
#     s2, e2 = _snap_and_bridge_to_word(script_text, s2, e2); s2 = _heal_split_word_left(script_text, s2)
#     if s2 >= s and e2 <= e and e2 > s2: return (s2, e2)
#     return span

# def build_spans_by_param(script_text: str, data: dict, heading_ranges: Optional[List[Tuple[int,int]]] = None) -> Dict[str, List[Tuple[int,int,str,str]]]:
#     heading_ranges = heading_ranges or []
#     raw = (data or {}).get("per_parameter", {}) or {}
#     per: Dict[str, Dict[str, Any]] = {k:(v or {}) for k,v in raw.items()}
#     spans_map: Dict[str, List[Tuple[int,int,str,str]]] = {p: [] for p in PARAM_ORDER}
#     st.session_state["aoi_match_ranges"] = {}

#     for p in spans_map.keys():
#         color = PARAM_COLORS.get(p, "#ffd54f")
#         blk = per.get(p, {}) or {}
#         aois = blk.get("areas_of_improvement") or []
#         for idx, item in enumerate(aois, start=1):
#             raw_q = (item or {}).get("quote_verbatim", "") or ""
#             q = _sanitize_editor_text(raw_q)
#             clean = _clean_quote_for_match(re.sub(r"^[•\-\d\.\)\s]+", "", q).strip())
#             if not clean: continue
#             if _is_heading_like(clean): continue

#             pos = find_span_smart(script_text, clean)
#             if not pos: continue
#             pos = _tighten_to_quote(script_text, pos, raw_q)
#             s, e = pos

#             if heading_ranges and _overlaps_any(s, e, heading_ranges): continue
#             if _is_heading_context(script_text, s, e): continue

#             aid = f"{p.replace(' ','_')}-AOI-{idx}"
#             spans_map[p].append((s, e, color, aid))
#             st.session_state["aoi_match_ranges"][aid] = (s, e)
#     return spans_map

# # ---------- History (S3-aware) ----------
# def _maybe_copy_docx_to_history(source_docx_path: Optional[str], run_id: str) -> Optional[str]:
#     """Copy the DOCX used for rendering into outputs/_history so Recents can re-render identically."""
#     try:
#         if not source_docx_path:
#             return None
#         # If already an S3 key/url, just store the key in history blob
#         if source_docx_path.startswith("s3://") or (_s3_enabled() and not os.path.exists(source_docx_path)):
#             # ensure it's uploaded under our history key if it's a local temp in S3 mode
#             if os.path.exists(source_docx_path):
#                 with open(source_docx_path, "rb") as f:
#                     save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#                 return f"{HISTORY_DIR}/{run_id}.docx"
#             return source_docx_path

#         # Local file path
#         if _s3_enabled():
#             with open(source_docx_path, "rb") as f:
#                 save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#             return f"{HISTORY_DIR}/{run_id}.docx"
#         else:
#             dst = os.path.join(HISTORY_DIR, f"{run_id}.docx")
#             if os.path.abspath(source_docx_path) != os.path.abspath(dst):
#                 shutil.copyfile(source_docx_path, dst)
#             return dst
#     except Exception:
#         return None

# def _save_history_snapshot(title: str, data: dict, script_text: str,
#                            source_docx_path: Optional[str], heading_ranges: List[Tuple[int,int]],
#                            spans_by_param: Dict[str, List[Tuple[int,int,str,str]]],
#                            aoi_match_ranges: Dict[str, Tuple[int,int]]):
#     run_id = str(uuid.uuid4()); now = datetime.datetime.now()
#     created_at_iso = now.replace(microsecond=0).isoformat()
#     created_at_human = now.strftime("%Y-%m-%d %H:%M:%S")

#     # Copy stable DOCX alongside history JSON (so Recents uses identical renderer)
#     stable_docx_key_or_path = _maybe_copy_docx_to_history(source_docx_path, run_id)

#     blob = {
#         "run_id": run_id, "title": title or "untitled",
#         "created_at": created_at_iso, "created_at_human": created_at_human,
#         "overall_rating": (data or {}).get("overall_rating", ""),
#         "scores": (data or {}).get("scores", {}),
#         "data": data or {}, "script_text": script_text or "",
#         "source_docx_path": stable_docx_key_or_path or source_docx_path,
#         "heading_ranges": heading_ranges or [],
#         "spans_by_param": spans_by_param or {},
#         "aoi_match_ranges": aoi_match_ranges or {},
#     }

#     # Persist JSON (S3 or local)
#     out_name = f"{created_at_iso.replace(':','-')}__{run_id}.json"
#     out_key = f"{HISTORY_DIR}/{out_name}"
#     save_text_key(out_key, json.dumps(blob, ensure_ascii=False, indent=2))

# def _load_all_history() -> List[dict]:
#     out: List[dict] = []
#     if _s3_enabled():
#         keys = sorted(list_prefix(HISTORY_DIR), reverse=True)
#         for key in keys:
#             try:
#                 txt = read_text_key(key, "")
#                 if not txt:
#                     continue
#                 j = json.loads(txt)
#                 j["_key"] = key
#                 if not j.get("created_at_human") and j.get("created_at"):
#                     try:
#                         dt = datetime.datetime.fromisoformat(j["created_at"])
#                         j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#                     except Exception:
#                         pass
#                 out.append(j)
#             except Exception:
#                 continue
#     else:
#         for fp in sorted(glob.glob(os.path.join(HISTORY_DIR, "*.json"))):
#             try:
#                 with open(fp, "r", encoding="utf-8") as f:
#                     j = json.load(f)
#             except Exception:
#                 continue
#             j.setdefault("_path", fp)
#             ca = j.get("created_at")
#             try:
#                 if isinstance(ca, (int, float)):
#                     dt = datetime.datetime.utcfromtimestamp(float(ca))
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.astimezone().strftime("%Y-%m-%d %H:%M:%S")
#                 elif isinstance(ca, str) and ca: pass
#                 else:
#                     mtime = os.path.getmtime(fp); dt = datetime.datetime.fromtimestamp(mtime)
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#             except Exception:
#                 j["created_at"] = str(ca or "")
#             out.append(j)
#     out.sort(key=lambda r: r.get("created_at") or "", reverse=True)
#     return out

# def _open_history_run_by_id(run_id: str) -> bool:
#     """Open a history run by its run_id. Returns True if loaded."""
#     if not run_id: return False
#     recs = _load_all_history()
#     match = next((r for r in recs if r.get("run_id") == run_id), None)
#     if not match: return False

#     # Load JSON content again (S3/local), then set session
#     try:
#         if "_key" in match and _s3_enabled():
#             txt = read_text_key(match["_key"], "")
#             if not txt: return False
#             jj = json.loads(txt)
#         else:
#             path = match.get("_path")
#             if not path or not os.path.exists(path): return False
#             with open(path, "r", encoding="utf-8") as f:
#                 jj = json.load(f)
#     except Exception:
#         return False

#     st.session_state.script_text      = jj.get("script_text","")
#     st.session_state.base_stem        = jj.get("title","untitled")
#     st.session_state.data             = jj.get("data",{})
#     st.session_state.heading_ranges   = jj.get("heading_ranges",[])
#     st.session_state.spans_by_param   = jj.get("spans_by_param",{})
#     st.session_state.param_choice     = None
#     st.session_state.source_docx_path = jj.get("source_docx_path")
#     st.session_state.review_ready     = True
#     st.session_state["aoi_match_ranges"] = jj.get("aoi_match_ranges", {})
#     st.session_state.ui_mode          = "review"
#     return True

# def _render_recents_centerpane():
#     st.subheader("📄 Recents")
#     q = st.text_input("Filter by title…", "")
#     cols = st.columns([1, 4])
#     with cols[0]:
#         if st.button("← Back"):
#             st.session_state.ui_mode = "home"; st.rerun()

#     recs = _load_all_history()
#     ql = q.strip().lower()
#     if ql: recs = [r for r in recs if ql in (r.get("title","").lower())]
#     if not recs: st.caption("No history yet."); st.stop()

#     # Each card is a clickable <a class="rec-card" href="?open=<run_id>">
#     for rec in recs:
#         run_id = rec.get("run_id"); title = rec.get("title") or "(untitled)"
#         created_h = rec.get("created_at_human",""); overall = rec.get("overall_rating","")
#         st.markdown(f"""
#         <a class="rec-card" href="?open={run_id}">
#           <div class="rec-title">{title}</div>
#           <div class="rec-meta">{created_h}</div>
#           <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
#         </a>
#         """, unsafe_allow_html=True)

# # ---------- Sidebar ----------
# with st.sidebar:
#     if st.button("🆕 New review", use_container_width=True):
#         fp = st.session_state.get("flattened_docx_path")
#         if fp and os.path.exists(fp):
#             try: os.remove(fp)
#             except Exception: pass
#         for k in ["review_ready","script_text","base_stem","data","spans_by_param","param_choice",
#                   "source_docx_path","heading_ranges","flattened_docx_path","flatten_used"]:
#             st.session_state[k] = (
#                 False if k=="review_ready"
#                 else "" if k in ("script_text","base_stem")
#                 else {} if k=="spans_by_param"
#                 else [] if k=="heading_ranges"
#                 else None if k in ("source_docx_path","flattened_docx_path")
#                 else False if k=="flatten_used"
#                 else None
#             )
#         st.session_state.ui_mode = "home"
#         _clear_query_params()
#         st.rerun()

#     if st.button("📁 Recents", use_container_width=True):
#         st.session_state.ui_mode = "recents"
#         _clear_query_params()
#         st.rerun()
# # ---------- Input screen ----------
# def render_home():
#     st.subheader("🎬 Script Source")

#     # Two tabs: Upload OR Paste text
#     tab_upload, tab_paste = st.tabs(["Upload file", "Paste text"])

#     uploaded_file = None
#     uploaded_name = None
#     uploaded_key  = None

#     # helper to create a safe filename for pasted text
#     def _safe_stem(s: str, fallback: str = "pasted_script") -> str:
#         s = (s or "").strip()
#         if not s:
#             return fallback
#         s = re.sub(r"[^A-Za-z0-9._\-]+", "_", s)
#         s = s.strip("._-") or fallback
#         return s

#     # --- Upload tab (unchanged) ---
#     with tab_upload:
#         up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
#         if up is not None:
#             file_bytes = up.read()
#             suffix = os.path.splitext(up.name)[1].lower()
#             # Persist to S3/local under scripts/
#             uploaded_key = f"{SCRIPTS_DIR}/{up.name}"
#             save_bytes_key(uploaded_key, file_bytes)

#             # Local temp for parsing
#             with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
#                 tmp.write(file_bytes)
#                 uploaded_file = tmp.name
#             uploaded_name = os.path.splitext(os.path.basename(up.name))[0] or "uploaded_script"

#     # --- Paste tab (new) ---
#     with tab_paste:
#         paste_title = st.text_input("Title (optional)", placeholder="e.g., my_script")
#         pasted_text = st.text_area(
#             "Paste your script text here",
#             height=360,
#             placeholder="Paste the full script text (we’ll analyze this as-is)."
#         )

#     # Run Review works for BOTH
#     if st.button("🚀 Run Review", type="primary", use_container_width=True):
#         base_stem = "uploaded_script"
#         source_docx_path = None
#         heading_ranges: List[Tuple[int,int]] = []
#         script_text = ""

#         # Prefer pasted text if provided
#         if pasted_text and pasted_text.strip():
#             base_stem = _safe_stem(paste_title, "pasted_script")
#             script_text = pasted_text

#             # Save pasted text so Recents can reopen consistently
#             pasted_key = f"{SCRIPTS_DIR}/{base_stem}.txt"
#             save_text_key(pasted_key, script_text)
#             source_docx_path = pasted_key
#             heading_ranges = []

#         # Fallback to uploaded file
#         elif uploaded_file:
#             base_stem = uploaded_name or "uploaded_script"

#             if uploaded_file.lower().endswith(".docx"):
#                 try:
#                     left_text, used_left = extract_left_column_script_or_default(uploaded_file)
#                 except Exception:
#                     left_text, used_left = "", False

#                 if used_left and left_text.strip():
#                     script_text = left_text
#                     source_docx_path = uploaded_file
#                     heading_ranges = []
#                 else:
#                     path_to_use = uploaded_file
#                     if _docx_contains_tables(path_to_use):
#                         flat = flatten_docx_tables_to_longtext(path_to_use)
#                         st.session_state.flattened_docx_path = flat
#                         st.session_state.flatten_used = True
#                         path_to_use = flat
#                     script_text, heading_ranges = build_docx_text_with_meta(path_to_use)
#                     source_docx_path = path_to_use
#             else:
#                 script_text = load_script_file(uploaded_file)
#                 source_docx_path = uploaded_key or uploaded_file

#         else:
#             st.warning("Please upload a script **or** paste text in the second tab.")
#             st.stop()

#         if len(script_text.strip()) < 50:
#             st.error("Extracted text looks too short. Please check your input.")
#             st.stop()

#         with st.spinner("Running analysis…"):
#             try:
#                 review_text = run_review_multi(
#                     script_text=script_text,
#                     prompts_dir=PROMPTS_DIR,
#                     temperature=0.0
#                 )
#             finally:
#                 if uploaded_file and not (isinstance(source_docx_path, str) and os.path.exists(source_docx_path)):
#                     try: os.remove(uploaded_file)
#                     except Exception: pass

#         data = extract_review_json(review_text)
#         if not data:
#             st.error("JSON not detected in model output.")
#             st.stop()

#         st.session_state.script_text      = script_text
#         st.session_state.base_stem        = base_stem
#         st.session_state.data             = data
#         st.session_state.heading_ranges   = heading_ranges
#         st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
#         st.session_state.param_choice     = None
#         st.session_state.source_docx_path = source_docx_path
#         st.session_state.review_ready     = True
#         st.session_state.ui_mode          = "review"

#         _save_history_snapshot(
#             title=base_stem,
#             data=data,
#             script_text=script_text,
#             source_docx_path=source_docx_path,
#             heading_ranges=heading_ranges,
#             spans_by_param=st.session_state.spans_by_param,
#             aoi_match_ranges=st.session_state.get("aoi_match_ranges", {})
#         )

#         _clear_query_params()
#         st.rerun()

# # ---------- Results screen ----------
# def render_review():
#     script_text     = st.session_state.script_text
#     data            = st.session_state.data
#     spans_by_param  = st.session_state.spans_by_param
#     scores: Dict[str,int] = (data or {}).get("scores", {}) or {}
#     source_docx_path: Optional[str] = st.session_state.source_docx_path

#     # If our source_docx_path is an S3 key/url, ensure we have a local copy for rendering
#     # AFTER (prefer the flattened copy if we have it in-session)
#     docx_local: Optional[str] = None
#     preferred = st.session_state.get("flattened_docx_path") if st.session_state.get("flatten_used") else source_docx_path
#     if preferred:
#         if not os.path.exists(preferred):
#             docx_local = ensure_local_copy(preferred)
#         else:
#             docx_local = preferred

#     left, center, right = st.columns([1.1, 2.7, 1.4], gap="large")

#     with left:
#         st.subheader("Final score")
#         ordered = [p for p in PARAM_ORDER if p in scores]
#         df = pd.DataFrame({"Parameter": ordered, "Score (1–10)": [scores.get(p, "") for p in ordered]})
#         st.dataframe(df, hide_index=True, use_container_width=True)
#         st.markdown(f'**Overall:** {data.get("overall_rating","—")}/10')
#         st.divider()

#         strengths = (data or {}).get("strengths") or []
#         if not strengths:
#             per = (data or {}).get("per_parameter", {}) or {}
#             best = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
#             for name, sc in best:
#                 if sc >= 8 and name in per:
#                     exp = _sanitize_editor_text((per[name] or {}).get("explanation", "") or "")
#                     first = re.split(r"(?<=[.!?])\s+", exp.strip())[0] if exp else f"Consistently strong {name.lower()}."
#                     strengths.append(f"{name}: {first}")
#                 if len(strengths) >= 3: break

#         def _bullets(title: str, items):
#             st.markdown(f"**{title}**")
#             for s in (items or []):
#                 if isinstance(s, str) and s.strip(): st.write("• " + _sanitize_editor_text(s))
#             if not items: st.write("• —")

#         _bullets("Strengths", strengths)
#         _bullets("Weaknesses", data.get("weaknesses"))
#         _bullets("Suggestions", data.get("suggestions"))
#         _bullets("Drop-off Risks", data.get("drop_off_risks"))
#         st.markdown("**Viral Quotient**"); st.write(_sanitize_editor_text(data.get("viral_quotient","—")))

#     with right:
#         st.subheader("Parameters")
#         st.markdown('<div class="param-row">', unsafe_allow_html=True)
#         for p in [p for p in PARAM_ORDER if p in scores]:
#             if st.button(p, key=f"chip_{p}", help="Show inline AOI highlights for this parameter"):
#                 st.session_state.param_choice = p
#         st.markdown('</div>', unsafe_allow_html=True)

#         sel = st.session_state.param_choice
#         if sel:
#             blk = (data.get("per_parameter", {}) or {}).get(sel, {}) or {}
#             st.markdown(f"**{sel} — Score:** {scores.get(sel,'—')}/10")

#             if blk.get("explanation"):
#                 st.markdown("**Why this score**"); st.write(_sanitize_editor_text(blk["explanation"]))
#             if blk.get("weakness") and blk["weakness"] != "Not present":
#                 st.markdown("**Weakness**"); st.write(_sanitize_editor_text(blk["weakness"]))
#             if blk.get("suggestion") and blk["suggestion"] != "Not present":
#                 st.markdown("**Suggestion**"); st.write(_sanitize_editor_text(blk["suggestion"]))

#             aoi = blk.get("areas_of_improvement") or []
#             if aoi:
#                 st.markdown("**Areas of Improvement**")
#                 for i, item in enumerate(aoi, 1):
#                     popover_fn = getattr(st, "popover", None)
#                     aid = f"{sel.replace(' ','_')}-AOI-{i}"
#                     s_e_map = st.session_state.get("aoi_match_ranges", {})
#                     if aid in s_e_map:
#                         s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                         line = (matched_line if len(matched_line) <= 320 else matched_line[:300] + "…")
#                     else:
#                         line = _sanitize_editor_text(item.get('quote_verbatim',''))
#                     issue = _sanitize_editor_text(item.get('issue',''))
#                     fix   = _sanitize_editor_text(item.get('fix',''))
#                     why   = _sanitize_editor_text(item.get('why_this_helps',''))
#                     label = f"Issue {i}"
#                     if callable(popover_fn):
#                         with popover_fn(label):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)
#                     else:
#                         with st.expander(label, expanded=False):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)

#             if blk.get("summary"):
#                 st.markdown("**Summary**"); st.write(_sanitize_editor_text(blk["summary"]))

#     with center:
#         st.subheader("Script with inline highlights")
#         # choose spans for selected parameter (or all if None)
#         spans = st.session_state.spans_by_param.get(st.session_state.param_choice, []) if st.session_state.param_choice else []

#         # Build AOI payload for popup: { aid: {line, issue, fix, why} }
#         aoi_payload: Dict[str, Dict[str, str]] = {}
#         data_per = (data or {}).get("per_parameter") or {}
#         s_e_map = st.session_state.get("aoi_match_ranges", {})
#         sel = st.session_state.param_choice

#         def _mk_line(aid: str, fallback_q: str = "") -> str:
#             if aid in s_e_map:
#                 s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                 return matched_line if len(matched_line) <= 320 else matched_line[:300] + "…"
#             return _sanitize_editor_text(fallback_q or "")

#         def _collect(param_name: str):
#             blk = (data_per.get(param_name) or {})
#             for i, item in enumerate(blk.get("areas_of_improvement") or [], 1):
#                 aid = f"{param_name.replace(' ','_')}-AOI-{i}"
#                 aoi_payload[aid] = {
#                     "line": _mk_line(aid, (item or {}).get("quote_verbatim","")),
#                     "issue": _sanitize_editor_text((item or {}).get("issue","")),
#                     "fix": _sanitize_editor_text((item or {}).get("fix","")),
#                     "why": _sanitize_editor_text((item or {}).get("why_this_helps","")),
#                 }

#         if sel: _collect(sel)
#         else:
#             for pn in [p for p in PARAM_ORDER if p in data_per]:
#                 _collect(pn)

#         payload_json = json.dumps(aoi_payload, ensure_ascii=False)

#         frame_theme_css = """
#         <style>
#           :root{
#             --m7-surface: #eef2f7;
#             --m7-on-surface: #0f172a;
#             --m7-border: rgba(15,23,42,.14);
#           }
#           @media (prefers-color-scheme: dark){
#             :root{
#               --m7-surface: #2f333a;
#               --m7-on-surface: #ffffff;
#               --m7-border: rgba(255,255,255,.18);
#             }
#             body { background: transparent !important; }
#           }

#           .docxwrap{
#             background: var(--m7-surface);
#             color: var(--m7-on-surface);
#             border: 1px solid var(--m7-border);
#             border-radius: 12px;
#             padding: 16px 14px 18px;
#           }
#           .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#           .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); }
#         </style>
#         """

#         tooltip_css = """
#         <style>
#         .aoi-pop {
#           position: absolute; max-width: 520px; min-width: 320px;
#           background: var(--m7-surface); border: 1px solid var(--m7-border); border-radius: 10px;
#           box-shadow: 0 10px 25px rgba(0,0,0,.12);
#           padding: 12px 14px; z-index: 9999; transform: translateY(-8px);
#           color: var(--m7-on-surface);
#         }
#         .aoi-pop h4 { margin: 0 0 .35rem 0; font-size: .95rem; }
#         .aoi-pop p  { margin: .15rem 0; line-height: 1.5; }
#         .aoi-pop .muted { opacity:.85; font-size:.85rem; }
#         .aoi-arrow {
#           position:absolute; left:50%; transform:translateX(-50%);
#           bottom:-7px; width:0;height:0;border-left:7px solid transparent;
#           border-right:7px solid transparent;border-top:7px solid var(--m7-border);
#         }
#         .aoi-arrow::after{
#           content:""; position:absolute; left:-6px; top:-7px; width:0;height:0;
#           border-left:6px solid transparent;border-right:6px solid transparent;border-top:6px solid var(--m7-surface);
#         }
#         </style>
#         """

#         # Select rendering source (DOCX with highlights if we have a local path)
#         if docx_local and os.path.splitext(docx_local)[1].lower() == ".docx":
#             html_core = render_docx_html_with_highlights(
#                 docx_local,
#                 merge_overlaps_and_adjacent(script_text, spans)
#             )
#         else:
#             from html import escape as _esc
#             orig = script_text
#             spans2 = [s for s in merge_overlaps_and_adjacent(orig, spans) if s[0] < s[1]]
#             spans2.sort(key=lambda x: x[0])
#             cur = 0; buf: List[str] = []
#             for s,e,c,aid in spans2:
#                 if s > cur: buf.append(_esc(orig[cur:s]))
#                 buf.append(
#                     f'<mark class="aoi-mark" data-aid="{aid}" '
#                     f'style="background:{c}33;border:1px solid {c};border-radius:3px;padding:0 2px;">'
#                     f'{_esc(orig[s:e])}</mark>'
#                 )
#                 cur = e
#             if cur < len(orig): buf.append(_esc(orig[cur:]))
#             html_core = (
#                 '<div class="docxwrap"><p style="white-space:pre-wrap; '
#                 'line-height:1.7; font-family:ui-serif, Georgia, Times New Roman, serif;">'
#                 + "".join(buf) +
#                 '</p></div>'
#             )

#         # Popup + autosize JS shell
#         html_shell = """
# %%FRAME_THEME_CSS%%
# %%TOOLTIP_CSS%%
# <div id="m7-doc">%%HTML_CORE%%</div>
# <div id="aoi-pop" class="aoi-pop" style="display:none;">
#   <div id="aoi-pop-content"></div>
#   <div class="aoi-arrow"></div>
# </div>
# <script>
# (function(){
#   const AOI = __PAYLOAD__;
#   const wrap = document.getElementById('m7-doc');
#   const pop  = document.getElementById('aoi-pop');
#   const body = document.getElementById('aoi-pop-content');

#   function resizeIframe() {
#     try {
#       const h = Math.max(
#         document.documentElement.scrollHeight,
#         document.body.scrollHeight
#       );
#       if (window.frameElement) {
#         window.frameElement.style.height = (h + 20) + 'px';
#         window.frameElement.style.width  = '100%';
#       }
#     } catch(e) {}
#   }
#   window.addEventListener('load', resizeIframe);
#   window.addEventListener('resize', resizeIframe);

#   function hide(){ pop.style.display='none'; }
#   function showFor(mark){
#     const aid = mark.getAttribute('data-aid');
#     const d = AOI[aid]; if(!d) return;
#     body.innerHTML =
#       (d.line  ? '<p><strong>Line:</strong> '  + d.line  + '</p>' : '') +
#       (d.issue ? '<p><strong>Issue:</strong> ' + d.issue + '</p>' : '') +
#       (d.fix   ? '<p><strong>Fix:</strong> '   + d.fix   + '</p>' : '') +
#       (d.why   ? '<p class="muted">'           + d.why   + '</p>' : '');
#     pop.style.display = 'block';

#     const r = mark.getBoundingClientRect();
#     const scY = window.scrollY || document.documentElement.scrollTop;
#     const scX = window.scrollX || document.documentElement.scrollLeft;
#     let top  = r.top + scY - pop.offsetHeight - 10;
#     let left = r.left + scX + r.width/2 - pop.offsetWidth/2;
#     if (top < 8) top = r.bottom + scY + 10;
#     if (left < 8) left = 8;
#     pop.style.top  = top + 'px';
#     pop.style.left = left + 'px';

#     resizeIframe();
#   }

#   wrap.addEventListener('click', (e) => {
#     const m = e.target.closest('.aoi-mark');
#     if(!m){ hide(); return; }
#     if(pop.style.display === 'block'){ hide(); }
#     showFor(m);
#     e.stopPropagation();
#   });

#   document.addEventListener('click', (e) => {
#     if(!e.target.closest('.aoi-pop') && !e.target.closest('.aoi-mark')) hide();
#   });
# })();
# </script>
# """
#         html_shell = (
#             html_shell
#             .replace("%%FRAME_THEME_CSS%%", frame_theme_css)
#             .replace("%%TOOLTIP_CSS%%", tooltip_css)
#             .replace("%%HTML_CORE%%", html_core)
#             .replace("__PAYLOAD__", payload_json)
#         )

#         components.html(html_shell, height=400, scrolling=False)

# # ---------- Router & query param open ----------
# _open_qp = _get_query_param("open")
# if _open_qp and _open_history_run_by_id(_open_qp):
#     _clear_query_params()  # avoid re-opening on subsequent reruns

# mode = st.session_state.ui_mode
# if mode == "recents":
#     _render_recents_centerpane()
# elif mode == "review" and st.session_state.review_ready:
#     render_review()
# else:
#     render_home()











#########Current working code above############
###############################################











# # app_grammarly_ui.py
# # Complete updated file — Runpod S3 history snapshots + Recents + safe secrets + optional S3 debug

# import os, re, glob, json, tempfile, difflib, uuid, datetime, shutil
# from pathlib import Path
# from typing import Dict, Any, List, Tuple, Optional

# import streamlit as st
# import pandas as pd
# import streamlit.components.v1 as components  # for inline HTML/JS popup

# # ---- utils & engine ----
# from utils1 import (
#     extract_review_json,
#     PARAM_ORDER,
#     load_script_file,
#     extract_left_column_script_or_default,  # left-column extractor
# )
# from review_engine_multi import run_review_multi

# # ---- DOCX rendering imports ----
# from docx import Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # =========================
# # RunPod S3 (inline helpers) — FIXED
# # =========================
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError

# # Accept both AWS_* and RUNPOD_* style envs / st.secrets
# def _get_env(key: str, default: str = "") -> str:
#     v = os.getenv(key, "")
#     if v:
#         return v.strip()
#     # Access st.secrets safely (won't crash if secrets.toml is missing)
#     try:
#         v2 = st.secrets.get(key)
#         if isinstance(v2, str):
#             return v2.strip()
#     except Exception:
#         pass
#     return (default or "").strip()

# # Debug flag (NO direct st.secrets)
# _SHOW_HISTORY_DEBUG = (_get_env("M7_SHOW_HISTORY_DEBUG") or "").lower() in {"1", "true", "yes"}

# # Primary config
# _RP_ENDPOINT = _get_env("RUNPOD_S3_ENDPOINT")
# _RP_BUCKET   = _get_env("RUNPOD_S3_BUCKET")
# _RP_REGION   = _get_env("RUNPOD_S3_REGION") or _get_env("AWS_DEFAULT_REGION") or ""

# # Credentials: prefer AWS_* if present; else accept RUNPOD_* fallbacks
# _AK = _get_env("AWS_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY")
# _SK = _get_env("AWS_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_KEY")
# _ST = _get_env("AWS_SESSION_TOKEN")  # optional

# # Options
# _FORCE_PATH = (_get_env("RUNPOD_S3_FORCE_PATH_STYLE") or "true").lower() in {"1","true","yes"}
# _USE_SSL    = (_get_env("RUNPOD_S3_USE_SSL") or "true").lower() in {"1","true","yes"}
# _VERIFY_SSL = (_get_env("RUNPOD_S3_VERIFY_SSL") or "true").lower() in {"1","true","yes"}

# def _s3_enabled() -> bool:
#     return bool(_RP_ENDPOINT and _RP_BUCKET and _AK and _SK)

# @st.cache_resource(show_spinner=False)
# def _s3_client():
#     if not _s3_enabled():
#         return None
#     session_kwargs = dict(
#         aws_access_key_id=_AK,
#         aws_secret_access_key=_SK,
#     )
#     if _ST:
#         session_kwargs["aws_session_token"] = _ST

#     cfg = Config(
#         signature_version="s3v4",
#         s3={"addressing_style": "path" if _FORCE_PATH else "auto"},
#         retries={"max_attempts": 3, "mode": "standard"}
#     )
#     return boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         region_name=_RP_REGION or None,
#         use_ssl=_USE_SSL,
#         verify=_VERIFY_SSL,
#         config=cfg,
#         **session_kwargs,
#     )

# def save_text_key(key: str, text: str) -> str:
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "w", encoding="utf-8") as f:
#             f.write(text)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=text.encode("utf-8"))
#     return f"s3://{_RP_BUCKET}/{key}"

# def save_bytes_key(key: str, data: bytes) -> str:
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "wb") as f:
#             f.write(data)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=data)
#     return f"s3://{_RP_BUCKET}/{key}"

# def read_text_key(key: str, default: str = "") -> str:
#     if not _s3_enabled():
#         try:
#             with open(key, "r", encoding="utf-8") as f:
#                 return f.read()
#         except Exception:
#             return default
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read().decode("utf-8", errors="ignore")
#     except Exception:
#         return default

# def read_bytes_key(key: str) -> Optional[bytes]:
#     if not _s3_enabled():
#         try:
#             with open(key, "rb") as f:
#                 return f.read()
#         except Exception:
#             return None
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read()
#     except Exception:
#         return None

# def list_prefix(prefix: str) -> List[str]:
#     """
#     List object keys under prefix (or local dir paths if not S3).
#     In S3 mode we always return KEYS (not URLs).
#     Only returns *.json in both modes.
#     """
#     if not _s3_enabled():
#         base = prefix if os.path.isdir(prefix) else os.path.dirname(prefix)
#         try:
#             return [os.path.join(base, p) for p in os.listdir(base) if p.endswith(".json")]
#         except Exception:
#             return []

#     out: List[str] = []
#     token = None
#     s3_prefix = prefix.rstrip("/") + "/"
#     try:
#         while True:
#             kwargs = {"Bucket": _RP_BUCKET, "Prefix": s3_prefix}
#             if token:
#                 kwargs["ContinuationToken"] = token
#             resp = _s3_client().list_objects_v2(**kwargs)
#             for c in resp.get("Contents", []):
#                 k = c.get("Key", "")
#                 if k.endswith(".json"):
#                     out.append(k)
#             token = resp.get("NextContinuationToken")
#             if not token:
#                 break
#     except (ClientError, EndpointConnectionError, NoCredentialsError):
#         return []
#     return out

# def presigned_url(key: str, expires: int = 3600) -> Optional[str]:
#     if not _s3_enabled():
#         return None
#     try:
#         return _s3_client().generate_presigned_url(
#             "get_object",
#             Params={"Bucket": _RP_BUCKET, "Key": key},
#             ExpiresIn=expires
#         )
#     except ClientError:
#         return None

# def ensure_local_copy(key_or_path: str) -> Optional[str]:
#     """
#     For DOCX/PDF parsing we need a real filesystem path.
#     If S3 mode, download to a temp file and return that path.
#     """
#     if not _s3_enabled():
#         return key_or_path if os.path.exists(key_or_path) else None

#     key = key_or_path
#     if key.startswith("s3://"):
#         # s3://bucket/path/to/file -> path/to/file
#         parts = key.split("/", 3)
#         key = parts[3] if len(parts) >= 4 else ""
#     data = read_bytes_key(key)
#     if data is None:
#         return None
#     fd, tmp = tempfile.mkstemp(suffix=os.path.splitext(key)[1] or "")
#     os.close(fd)
#     with open(tmp, "wb") as f:
#         f.write(data)
#     return tmp

# def _s3_health_summary() -> dict:
#     """
#     Optional tiny health read you can print if needed.
#     Returns a dict; safe to ignore in production.
#     """
#     info = {
#         "enabled": _s3_enabled(),
#         "endpoint": _RP_ENDPOINT,
#         "bucket": _RP_BUCKET,
#         "region": _RP_REGION,
#         "has_keys": bool(_AK and _SK),
#     }
#     if not _s3_enabled():
#         info["status"] = "local-mode"
#         return info
#     try:
#         _ = _s3_client().list_objects_v2(Bucket=_RP_BUCKET, Prefix=(f"{OUTPUT_DIR}/_history/").rstrip("/") + "/", MaxKeys=1)
#         info["status"] = "ok"
#     except Exception as e:
#         info["status"] = f"error: {getattr(e, 'response', {}).get('Error', {}).get('Code', str(e))}"
#     return info


# # ---------- Folders (all under Scriptmodel/) ----------
# BASE_PREFIX = "Scriptmodel"

# SCRIPTS_DIR = f"{BASE_PREFIX}/scripts"
# PROMPTS_DIR = f"{BASE_PREFIX}/prompts"
# OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
# HISTORY_DIR = f"{OUTPUT_DIR}/_history"

# # (redefine explicitly to avoid earlier stray brace typo)
# OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
# HISTORY_DIR = f"{OUTPUT_DIR}/_history"

# if not _s3_enabled():
#     for p in (SCRIPTS_DIR, PROMPTS_DIR, OUTPUT_DIR, HISTORY_DIR):
#         Path(p).mkdir(parents=True, exist_ok=True)


# # ---------- Colors ----------
# PARAM_COLORS: Dict[str, str] = {
#     "Suspense Building":              "#ff6b6b",
#     "Language/Tone":                  "#6b8cff",
#     "Intro + Main Hook/Cliffhanger":  "#ffb86b",
#     "Story Structure + Flow":         "#a78bfa",
#     "Pacing":                         "#f43f5e",
#     "Mini-Hooks (30–60s)":            "#eab308",
#     "Outro (Ending)":                 "#8b5cf6",
#     "Grammar & Spelling":             "#10b981",
# }

# # ---------- Config ----------
# STRICT_MATCH_ONLY = False

# # ---------- App config ----------
# st.set_page_config(page_title="M7 — Grammarly UI", page_icon="🕵️", layout="wide")

# # ---------- Header patch & CSS ----------
# def render_app_title():
#     st.markdown('<h1 class="app-title">Viral Script Reviewer</h1>', unsafe_allow_html=True)
#     st.markdown("""
#     <style>
#     html { color-scheme: light dark; }
#     :root{
#       --m7-surface: #eef2f7;
#       --m7-on-surface: #0f172a;
#       --m7-border: rgba(15,23,42,.14);
#       --sep: #e5e7eb;
#     }
#     @media (prefers-color-scheme: dark){
#       :root{ --m7-surface:#2f333a; --m7-on-surface:#ffffff; --m7-border:rgba(255,255,255,.18); --sep:#2a2f37; }
#     }
#     .stApp .block-container { padding-top: 4.25rem !important; }
#     .app-title{
#       font-weight:700; font-size:2.1rem; line-height:1.3;
#       margin:0 0 1rem 0; padding-left:40px !important; padding-top:.25rem !important;
#       white-space:normal; word-break:break-word; hyphens:auto; overflow:visible;
#       position:relative !important; z-index:10 !important;
#     }
#     [data-testid="collapsedControl"] { z-index:6 !important; }
#     header[data-testid="stHeader"], .stAppHeader { background: transparent !important; box-shadow:none !important; }
#     @media (min-width: 992px){ .app-title { padding-left:0 !important; } }
#     div[data-testid="column"]:nth-of-type(1){position:relative;}
#     div[data-testid="column"]:nth-of-type(1)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     div[data-testid="column"]:nth-of-type(2){position:relative;}
#     div[data-testid="column"]:nth-of-type(2)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     .m7-card{ background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:14px 16px; color:var(--m7-on-surface); }
#     .m7-card, .m7-card * { color:var(--m7-on-surface) !important; }
#     .docxwrap{ background:var(--m7-surface); color:var(--m7-on-surface); border:1px solid var(--m7-border); border-radius:12px; padding:16px 14px 18px; }
#     .docxwrap, .docxwrap * { color:var(--m7-on-surface) !important; }
#     .docxwrap .h1, .docxwrap .h2, .docxwrap .h3 { font-weight:700; margin:10px 0 6px; }
#     .docxwrap .h1 { font-size:1.3rem; border-bottom:2px solid currentColor; padding-bottom:4px; }
#     .docxwrap .h2 { font-size:1.15rem; border-bottom:1px solid currentColor; padding-bottom:3px; }
#     .docxwrap .h3 { font-size:1.05rem; }
#     .docxwrap p { margin:10px 0; line-height:1.7; font-family: ui-serif, Georgia, "Times New Roman", serif; }
#     .docxwrap table { border-collapse: collapse; width: 100%; margin: 12px 0; }
#     .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); padding:8px; vertical-align:top; line-height:1.6; }
#     .docxwrap mark{ padding:0 2px; border-radius:3px; border:1px solid var(--m7-border); cursor: pointer; }
#     .rec-card{
#       display:block; text-decoration:none !important;
#       background: var(--m7-surface); border:1px solid var(--m7-border);
#       border-radius: 12px; padding: 14px 16px; margin: 10px 0 16px;
#       box-shadow: 0 1px 2px rgba(0,0,0,.06); color:var(--m7-on-surface) !important;
#       transition: filter .1s ease, transform .02s ease;
#     }
#     .rec-card:hover{ filter: brightness(1.02); }
#     .rec-card:active{ transform: translateY(1px); }
#     .rec-title{font-weight:600; margin-bottom:.25rem;}
#     .rec-meta{opacity:.85 !important; font-size:12.5px; margin-bottom:.4rem;}
#     .rec-row{display:flex; align-items:center; justify-content:space-between; gap:12px;}
#     .stTextInput>div>div, .stTextArea>div>div, .stNumberInput>div>div, .stDateInput>div>div, .stTimeInput>div>div,
#     .stFileUploader>div, div[data-baseweb="select"]{
#       background: var(--m7-surface) !important; border: 1px solid var(--m7-border) !important; border-radius: 10px !important;
#       color: var(--m7-on-surface) !important;
#     }
#     .stTextInput input, .stTextArea textarea, .stNumberInput input, .stDateInput input, .stTimeInput input,
#     .stFileUploader div, div[data-baseweb="select"] *{ color: var(--m7-on-surface) !important; }
#     .stTextInput input::placeholder, .stTextArea textarea::placeholder{ color: rgba(16,24,39,.55) !important; }
#     @media (prefers-color-scheme: dark){
#       .stTextInput input::placeholder, .stTextArea textarea::placeholder{ color: rgba(255,255,255,.75) !important; }
#     }
#     div[data-testid="stFileUploaderDropzone"] label span { color: var(--m7-on-surface) !important; opacity: 1 !important; }
#     div[data-testid="stFileUploaderDropzone"] label { color: var(--m7-on-surface) !important; }
#     .stMarkdown pre, pre[class*="language-"], .stCodeBlock{
#       background: var(--m7-surface) !important; color: var(--m7-on-surface) !important;
#       border: 1px solid var(--m7-border) !important; border-radius: 12px !important; padding: 12px 14px !important; overflow:auto;
#     }
#     .stMarkdown pre code{ background: transparent !important; color: inherit !important; }
#     div[data-testid="stDataFrame"]{ background: var(--m7-surface); border: 1px solid var(--m7-border); border-radius: 12px; padding: 6px 8px; color: var(--m7-on-surface); }
#     div[data-testid="stDataFrame"] * { color: var(--m7-on-surface) !important; }
#     </style>
#     """, unsafe_allow_html=True)

# render_app_title()

# # ---------- Session ----------
# for key, default in [
#     ("review_ready", False),
#     ("script_text", ""),
#     ("base_stem", ""),
#     ("data", None),
#     ("spans_by_param", {}),
#     ("param_choice", None),
#     ("source_docx_path", None),
#     ("heading_ranges", []),
#     ("flattened_docx_path", None),
#     ("flatten_used", False),
#     ("ui_mode", "home"),
# ]:
#     st.session_state.setdefault(key, default)

# # ---------- helpers for query params (compat across Streamlit versions) ----------
# def _get_query_param(key: str) -> Optional[str]:
#     val = None
#     try:
#         val = st.query_params.get(key)
#     except Exception:
#         q = st.experimental_get_query_params()
#         v = q.get(key)
#         if isinstance(v, list): val = v[0] if v else None
#         else: val = v
#     return val

# def _clear_query_params():
#     try:
#         st.query_params.clear()
#     except Exception:
#         st.experimental_set_query_params()

# # ---------- Sanitizer ----------
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )
# def _sanitize_editor_text(s: Optional[str]) -> str:
#     if not s: return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•)\s*', '', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# # ---------- DOCX traversal ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# # ---------- Auto-flatten ----------
# def _docx_contains_tables(path: str) -> bool:
#     doc = Document(path)
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Table):
#             return True
#     return False

# def _copy_paragraph(dest_doc: Document, src_para: Paragraph):
#     p = dest_doc.add_paragraph()
#     try:
#         if src_para.style and src_para.style.name:
#             p.style = src_para.style.name
#     except Exception:
#         pass
#     for run in src_para.runs:
#         r = p.add_run(run.text or "")
#         r.bold = run.bold
#         r.italic = run.italic
#         r.underline = run.underline
#     return p

# def flatten_docx_tables_to_longtext(source_path: str) -> str:
#     src = Document(source_path)
#     new = Document()
#     for blk in _iter_docx_blocks(src):
#         if isinstance(blk, Paragraph):
#             _copy_paragraph(new, blk)
#         else:
#             seen_tc_ids = set()
#             for row in blk.rows:
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     if tc_id in seen_tc_ids:
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     for p in cell.paragraphs:
#                         _copy_paragraph(new, p)
#                 new.add_paragraph("")
#             new.add_paragraph("")
#     fd, tmp_path = tempfile.mkstemp(suffix=".docx"); os.close(fd); new.save(tmp_path)
#     return tmp_path

# # ---------- Build plain text + heading ranges ----------
# def build_docx_text_with_meta(docx_path: str) -> Tuple[str, List[Tuple[int,int]]]:
#     doc = Document(docx_path)
#     out: List[str] = []
#     heading_ranges: List[Tuple[int,int]] = []
#     current_offset = 0

#     def _append_and_advance(s: str):
#         nonlocal current_offset
#         out.append(s); current_offset += len(s)

#     seen_tc_ids: set = set()
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             para_text = "".join(run.text or "" for run in blk.runs)
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             if sty.startswith("heading"):
#                 start = current_offset; end = start + len(para_text)
#                 heading_ranges.append((start, end))
#             _append_and_advance(para_text); _append_and_advance("\n")
#         else:
#             for row in blk.rows:
#                 row_cell_tcs = []
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     row_cell_tcs.append((tc_id, cell))
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     if tc_id in seen_tc_ids:
#                         if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     cell_text_parts: List[str] = []
#                     for i, p in enumerate(cell.paragraphs):
#                         t = "".join(r.text or "" for r in p.runs)
#                         sty = (p.style.name or "").lower() if p.style else ""
#                         if sty.startswith("heading"):
#                             hs = current_offset + sum(len(x) for x in cell_text_parts)
#                             he = hs + len(t)
#                             heading_ranges.append((hs, he))
#                         cell_text_parts.append(t)
#                         if i != len(cell.paragraphs) - 1:
#                             cell_text_parts.append("\n")
#                     cell_text = "".join(cell_text_parts)
#                     _append_and_advance(cell_text)
#                     if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                 _append_and_advance("\n")
#             _append_and_advance("\n")

#     return "".join(out), heading_ranges

# def _wrap_inline(safe_text: str, run) -> str:
#     out = safe_text
#     if getattr(run, "underline", False): out = f"<u>{out}</u>"
#     if getattr(run, "italic", False): out = f"<em>{out}</em>"
#     if getattr(run, "bold", False): out = f"<strong>{out}</strong>"
#     return out

# # ---------- Invisible/bridge chars ----------
# _BRIDGE_CHARS = set("\u200b\u200c\u200d\u2060\ufeff\xa0\u00ad")

# # ---------- DOCX -> HTML with highlights (includes data-aid) ----------
# def render_docx_html_with_highlights(docx_path: str,
#                                      highlight_spans: List[Tuple[int,int,str,str]]) -> str:
#     """
#     highlight_spans: list of (start, end, color, aid)
#     """
#     doc = Document(docx_path)
#     spans = [s for s in highlight_spans if s[0] < s[1]]
#     spans.sort(key=lambda x: x[0])

#     cur_span = 0
#     current_offset = 0

#     def esc(s: str) -> str:
#         return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

#     def open_mark_if_needed(html_parts, mark_state, color, end, aid):
#         if not mark_state["open"]:
#             html_parts.append(
#                 f'<mark class="aoi-mark" data-aid="{aid}" '
#                 f'style="background:{color}33;border:1px solid {color};border-radius:3px;padding:0 2px;">'
#             )
#             mark_state.update(open=True, end=end, color=color, aid=aid)

#     def close_mark_if_open(html_parts, mark_state):
#         if mark_state["open"]:
#             html_parts.append('</mark>')
#             mark_state.update(open=False, end=None, color=None, aid=None)

#     def emit_run_text(run_text: str, run, html_parts: List[str], mark_state: Dict[str, Any]):
#         nonlocal cur_span, current_offset
#         t = run_text or ""; i = 0
#         while i < len(t):
#             next_start, next_end, color, next_aid = None, None, None, None
#             if cur_span < len(spans):
#                 next_start, next_end, color, next_aid = spans[cur_span]

#             if not mark_state["open"]:
#                 if cur_span >= len(spans) or current_offset + (len(t) - i) <= next_start:
#                     chunk = t[i:]; html_parts.append(_wrap_inline(esc(chunk), run)); current_offset += len(chunk); break
#                 if current_offset < next_start:
#                     take = next_start - current_offset
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take; continue
#                 open_mark_if_needed(html_parts, mark_state, color, next_end, next_aid)
#             else:
#                 take = min(mark_state["end"] - current_offset, len(t) - i)
#                 if take > 0:
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take
#                 if current_offset >= mark_state["end"]:
#                     close_mark_if_open(html_parts, mark_state)
#                     cur_span += 1

#     html: List[str] = ['<div class="docxwrap">']
#     seen_tc_ids: set = set()

#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             open_tag = '<div class="h1">' if sty.startswith("heading 1") else \
#                        '<div class="h2">' if sty.startswith("heading 2") else \
#                        '<div class="h3">' if sty.startswith("heading 3") else "<p>"
#             close_tag = "</div>" if sty.startswith("heading") else "</p>"
#             html.append(open_tag)
#             for run in blk.runs:
#                 emit_run_text(run.text or "", run, html, mark_state)
#             close_mark_if_open(html, mark_state)
#             html.append(close_tag)
#             current_offset += 1  # '\n'

#         else:
#             html.append("<table>")
#             for row in blk.rows:
#                 html.append("<tr>")
#                 row_cell_tcs = [(id(cell._tc), cell) for cell in row.cells]
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     html.append("<td>")
#                     if tc_id not in seen_tc_ids:
#                         seen_tc_ids.add(tc_id)
#                         for p_idx, p in enumerate(cell.paragraphs):
#                             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#                             html.append("<div>")
#                             for run in p.runs:
#                                 emit_run_text(run.text or "", run, html, mark_state)
#                             close_mark_if_open(html, mark_state)
#                             html.append("</div>")
#                             if p_idx != len(cell.paragraphs) - 1:
#                                 current_offset += 1
#                     html.append("</td>")
#                     if idx != len(row_cell_tcs) - 1: current_offset += 1  # '\t'
#                 html.append("</tr>"); current_offset += 1   # row '\n'
#             html.append("</table>"); current_offset += 1     # extra '\n'

#     html.append("</div>")
#     return "".join(html)

# # ---------- Matching / span utilities ----------
# def _normalize_keep_len(s: str) -> str:
#     trans = {
#         "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
#         "\u2013": "-", "\u2014": "-",
#         "\xa0": " ",
#         "\u200b": " ", "\u200c": " ", "\u200d": " ", "\u2060": " ",
#         "\ufeff": " ", "\u00ad": " ",
#     }
#     return (s or "").translate(str.maketrans(trans))

# def _tokenize(s: str) -> List[str]: return re.findall(r"\w+", (s or "").lower())

# def _iter_sentences_with_spans(text: str) -> List[Tuple[int,int,str]]:
#     spans = []
#     for m in re.finditer(r'[^.!?]+[.!?]+|\Z', text, flags=re.S):
#         s, e = m.start(), m.end()
#         seg = text[s:e]
#         if seg.strip(): spans.append((s, e, seg))
#     return spans

# def _squash_ws(s: str) -> str: return re.sub(r"\s+", " ", s or "").strip()

# def _clean_quote_for_match(q: str) -> str:
#     if not q: return ""
#     q = _normalize_keep_len(q).strip()
#     q = re.sub(r'^[\'"“”‘’\[\(\{<…\-\–\—\s]+', '', q)
#     q = re.sub(r'[\'"“”‘’\]\)\}>…\-\–\—\s]+$', '', q)
#     return _squash_ws(q)

# def _snap_and_bridge_to_word(text: str, start: int, end: int, max_bridge: int = 2) -> Tuple[int,int]:
#     n = len(text); s, e = max(0,start), max(start,end)
#     def _is_inv(ch: str) -> bool: return ch in _BRIDGE_CHARS

#     while s > 0:
#         prev = text[s-1]; cur = text[s] if s < n else ""
#         if prev.isalnum() and cur.isalnum(): s -= 1; continue
#         j = s; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (s-1)>=0 and text[s-1].isalnum() and (j<n and text[j].isalnum()): s -= 1; continue
#         break

#     while e < n:
#         prev = text[e-1] if e>0 else ""; nxt = text[e]
#         if prev.isalnum() and nxt.isalnum(): e += 1; continue
#         j = e; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (e-1)>=0 and text[e-1].isalnum() and (j<n and text[j].isalnum()): e = j + 1; continue
#         break

#     while e < n and text[e] in ',"”’\')]}': e += 1
#     return s, e

# def _heal_split_word_left(text: str, start: int) -> int:
#     i = start
#     if i <= 1 or i >= len(text): return start
#     if text[i-1] != " ": return start
#     j = i - 2
#     while j >= 0 and text[j].isalpha(): j -= 1
#     prev_token = text[j+1:i-1]
#     if len(prev_token) == 1: return i - 2
#     return start

# def _overlaps_any(s: int, e: int, ranges: List[Tuple[int,int]]) -> bool:
#     for rs, re_ in ranges:
#         if e > rs and s < re_: return True
#     return False

# def _fuzzy_window_span(tl: str, nl: str, start: int, w: int) -> Tuple[float, Optional[Tuple[int,int]]]:
#     window = tl[start:start+w]
#     sm = difflib.SequenceMatcher(a=nl, b=window)
#     blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
#     if not blocks: return 0.0, None
#     coverage = sum(b.size for b in blocks) / max(1, len(nl))
#     first_b = min(blocks, key=lambda b: b.b); last_b = max(blocks, key=lambda b: b.b + b.size)
#     s = start + first_b.b; e = start + last_b.b + last_b.size
#     return coverage, (s, e)

# def find_span_smart(text: str, needle: str) -> Optional[Tuple[int,int]]:
#     if not text or not needle: return None
#     t_orig = text; t_norm = _normalize_keep_len(text); n_norm = _clean_quote_for_match(needle)
#     if not n_norm: return None
#     tl = t_norm.lower(); nl = n_norm.lower()

#     i = tl.find(nl)
#     if i != -1:
#         s, e = _snap_and_bridge_to_word(t_orig, i, i + len(nl)); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     m = re.search(re.escape(nl).replace(r"\ ", r"\s+"), tl, flags=re.IGNORECASE)
#     if m:
#         s, e = _snap_and_bridge_to_word(t_orig, m.start(), m.end()); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     if not STRICT_MATCH_ONLY and len(nl) >= 12:
#         w = max(60, min(240, len(nl) + 80))
#         best_cov, best_span = 0.0, None
#         step = max(1, w // 2)
#         for start in range(0, max(1, len(tl) - w + 1), step):
#             cov, se = _fuzzy_window_span(tl, nl, start, w)
#             if cov > best_cov: best_cov, best_span = cov, se
#         if best_span and best_cov >= 0.65:
#             s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
#             if s > 0 and t_orig[s-1:s+1].lower() in {" the", " a", " an"}: s -= 1
#             s = _heal_split_word_left(t_orig, s)
#             return (s, e)

#     if not STRICT_MATCH_ONLY:
#         keys = [w for w in _tokenize(nl) if len(w) >= 4][:8]
#         if len(keys) >= 2:
#             kset = set(keys)
#             best_score, best_span = 0.0, None
#             for s, e, seg in _iter_sentences_with_spans(t_norm):
#                 toks = set(_tokenize(seg)); ov = len(kset & toks)
#                 if ov == 0: continue
#                 score = ov / max(2, len(kset)); length_pen = min(1.0, 120 / max(20, e - s)); score *= (0.6 + 0.4 * length_pen)
#                 if score > best_score: best_score, best_span = score, (s, min(e, s + 400))
#             if best_span and best_score >= 0.35:
#                 s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1]); s = _heal_split_word_left(t_orig, s)
#                 return (s, e)
#     return None

# def merge_overlaps(spans: List[Tuple[int,int,str,str]]) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans.sort(key=lambda x: x[0]); out = [spans[0]]
#     for s,e,c,aid in spans[1:]:
#         ps,pe,pc,paid = out[-1]
#         if s <= pe and pc == c and e > pe: out[-1] = (ps, e, pc, paid)
#         else: out.append((s,e,c,aid))
#     return out

# _PUNCT_WS = set(" \t\r\n,.;:!?)('\"[]{}-—–…") | _BRIDGE_CHARS

# def merge_overlaps_and_adjacent(base_text: str,
#                                 spans: List[Tuple[int,int,str,str]],
#                                 max_gap: int = 2) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans = sorted(spans, key=lambda x: x[0]); out = [spans[0]]
#     for s, e, c, aid in spans[1:]:
#         ps, pe, pc, paid = out[-1]
#         if c == pc and s <= pe: out[-1] = (ps, max(pe, e), pc, paid); continue
#         if c == pc and s - pe <= max_gap:
#             gap = base_text[max(0, pe):max(0, s)]
#             if all((ch in _PUNCT_WS) for ch in gap): out[-1] = (ps, e, pc, paid); continue
#         out.append((s, e, c, aid))
#     return out

# # ---------- Heading filters ----------
# def _is_heading_like(q: str) -> bool:
#     if not q: return True
#     s = q.strip()
#     if not re.search(r'[.!?]', s):
#         words = re.findall(r"[A-Za-z]+", s)
#         if 1 <= len(words) <= 7:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.8: return True
#         if s.lower() in {"introduction","voiceover","outro","epilogue","prologue","credits","title","hook","horrifying discovery","final decision"}: return True
#         if len(s) <= 3: return True
#     return False

# def _is_heading_context(script_text: str, s: int, e: int) -> bool:
#     left = script_text.rfind("\n", 0, s) + 1
#     right = script_text.find("\n", e); right = len(script_text) if right == -1 else right
#     line = script_text[left:right].strip()
#     if len(line) <= 70 and not re.search(r'[.!?]', line):
#         words = re.findall(r"[A-Za-z]+", line)
#         if 1 <= len(words) <= 8:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.7: return True
#     return False

# def _tighten_to_quote(script_text: str, span: Tuple[int,int], quote: str) -> Tuple[int,int]:
#     if not span or not quote: return span
#     s, e = span
#     if e <= s or s < 0 or e > len(script_text): return span
#     window = script_text[s:e]; win_norm = _normalize_keep_len(window).lower(); q_norm = _clean_quote_for_match(quote).lower()
#     if not q_norm: return span
#     i = win_norm.find(q_norm)
#     if i == -1:
#         m = re.search(re.escape(q_norm).replace(r"\ ", r"\s+"), win_norm, flags=re.IGNORECASE)
#         if not m: return span
#         i, j = m.start(), m.end()
#     else:
#         j = i + len(q_norm)
#     s2, e2 = s + i, s + j
#     s2, e2 = _snap_and_bridge_to_word(script_text, s2, e2); s2 = _heal_split_word_left(script_text, s2)
#     if s2 >= s and e2 <= e and e2 > s2: return (s2, e2)
#     return span

# def build_spans_by_param(script_text: str, data: dict, heading_ranges: Optional[List[Tuple[int,int]]] = None) -> Dict[str, List[Tuple[int,int,str,str]]]:
#     heading_ranges = heading_ranges or []
#     raw = (data or {}).get("per_parameter", {}) or {}
#     per: Dict[str, Dict[str, Any]] = {k:(v or {}) for k,v in raw.items()}
#     spans_map: Dict[str, List[Tuple[int,int,str,str]]] = {p: [] for p in PARAM_ORDER}
#     st.session_state["aoi_match_ranges"] = {}

#     for p in spans_map.keys():
#         color = PARAM_COLORS.get(p, "#ffd54f")
#         blk = per.get(p, {}) or {}
#         aois = blk.get("areas_of_improvement") or []
#         for idx, item in enumerate(aois, start=1):
#             raw_q = (item or {}).get("quote_verbatim", "") or ""
#             q = _sanitize_editor_text(raw_q)
#             clean = _clean_quote_for_match(re.sub(r"^[•\-\d\.\)\s]+", "", q).strip())
#             if not clean: continue
#             if _is_heading_like(clean): continue

#             pos = find_span_smart(script_text, clean)
#             if not pos: continue
#             pos = _tighten_to_quote(script_text, pos, raw_q)
#             s, e = pos

#             if heading_ranges and _overlaps_any(s, e, heading_ranges): continue
#             if _is_heading_context(script_text, s, e): continue

#             aid = f"{p.replace(' ','_')}-AOI-{idx}"
#             spans_map[p].append((s, e, color, aid))
#             st.session_state["aoi_match_ranges"][aid] = (s, e)
#     return spans_map

# # ---------- History (S3-aware) ----------
# def _maybe_copy_docx_to_history(source_docx_path: Optional[str], run_id: str) -> Optional[str]:
#     """Copy the DOCX used for rendering into outputs/_history so Recents can re-render identically."""
#     try:
#         if not source_docx_path:
#             return None
#         # If already an S3 key/url, just store the key in history blob
#         if source_docx_path.startswith("s3://") or (_s3_enabled() and not os.path.exists(source_docx_path)):
#             # ensure it's uploaded under our history key if it's a local temp in S3 mode
#             if os.path.exists(source_docx_path):
#                 with open(source_docx_path, "rb") as f:
#                     save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#                 return f"{HISTORY_DIR}/{run_id}.docx"
#             return source_docx_path

#         # Local file path
#         if _s3_enabled():
#             with open(source_docx_path, "rb") as f:
#                 save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#             return f"{HISTORY_DIR}/{run_id}.docx"
#         else:
#             dst = os.path.join(HISTORY_DIR, f"{run_id}.docx")
#             if os.path.abspath(source_docx_path) != os.path.abspath(dst):
#                 shutil.copyfile(source_docx_path, dst)
#             return dst
#     except Exception:
#         return None

# def _save_history_snapshot(title: str, data: dict, script_text: str,
#                            source_docx_path: Optional[str], heading_ranges: List[Tuple[int,int]],
#                            spans_by_param: Dict[str, List[Tuple[int,int,str,str]]],
#                            aoi_match_ranges: Dict[str, Tuple[int,int]]):
#     run_id = str(uuid.uuid4()); now = datetime.datetime.now()
#     created_at_iso = now.replace(microsecond=0).isoformat()
#     created_at_human = now.strftime("%Y-%m-%d %H:%M:%S")

#     # Copy stable DOCX alongside history JSON (so Recents uses identical renderer)
#     stable_docx_key_or_path = _maybe_copy_docx_to_history(source_docx_path, run_id)

#     blob = {
#         "run_id": run_id, "title": title or "untitled",
#         "created_at": created_at_iso, "created_at_human": created_at_human,
#         "overall_rating": (data or {}).get("overall_rating", ""),
#         "scores": (data or {}).get("scores", {}),
#         "data": data or {}, "script_text": script_text or "",
#         "source_docx_path": stable_docx_key_or_path or source_docx_path,
#         "heading_ranges": heading_ranges or [],
#         "spans_by_param": spans_by_param or {},
#         "aoi_match_ranges": aoi_match_ranges or {},
#     }

#     # Persist JSON (S3 or local)
#     out_name = f"{created_at_iso.replace(':','-')}__{run_id}.json"
#     out_key = f"{HISTORY_DIR}/{out_name}"
#     save_text_key(out_key, json.dumps(blob, ensure_ascii=False, indent=2))

#     # Friendly toast
#     try:
#         st.toast(f"History saved: {out_key}", icon="✅")
#     except Exception:
#         pass

# def _load_all_history(limit: Optional[int] = None) -> List[dict]:
#     """
#     Lightweight history loader:
#     - Lists keys under Scriptmodel/outputs/_history/
#     - Loads only the latest N JSONs (default 50 or M7_HISTORY_LIST_LIMIT)
#     - Parses just the small metadata needed for the Recents cards
#       (title, created_at, created_at_human, overall_rating, run_id)
#     - DOES NOT bring script_text into memory here (loaded on click)
#     """
#     # How many cards to show in Recents
#     try:
#         default_limit = int(os.getenv("M7_HISTORY_LIST_LIMIT", "50"))
#     except Exception:
#         default_limit = 50
#     if limit is None:
#         limit = default_limit

#     out: List[dict] = []

#     if _s3_enabled():
#         # Get all JSON keys (list_prefix already filters to *.json)
#         keys = sorted(list_prefix(HISTORY_DIR), reverse=True)
#         if not keys:
#             return out

#         # Take only the most recent <limit> items by key (keys are prefixed with ISO-like time)
#         keys = keys[:limit]

#         for key in keys:
#             try:
#                 txt = read_text_key(key, "")
#                 if not txt:
#                     continue
#                 j = json.loads(txt)

#                 # Build a compact record only (avoid keeping script_text in memory)
#                 rec = {
#                     "run_id": j.get("run_id"),
#                     "title": (j.get("title") or "untitled"),
#                     "created_at": j.get("created_at", ""),
#                     "created_at_human": j.get("created_at_human", ""),
#                     "overall_rating": j.get("overall_rating", ""),
#                     # keep a pointer to re-open the full JSON on click
#                     "_key": key,
#                 }
#                 out.append(rec)
#             except Exception:
#                 # skip malformed files but don't break the list
#                 continue

#         # Newest first by created_at when available, else by key fallback
#         def _sort_k(r):
#             return (r.get("created_at") or ""), (r.get("_key") or "")
#         out.sort(key=_sort_k, reverse=True)

#     else:
#         # Local mode: read local JSONs from the history dir, but still keep it light
#         paths = sorted(glob.glob(os.path.join(HISTORY_DIR, "*.json")), reverse=True)[:limit]
#         for fp in paths:
#             try:
#                 with open(fp, "r", encoding="utf-8") as f:
#                     j = json.load(f)
#             except Exception:
#                 continue

#             # convert timestamps for local files (keep it small)
#             ca = j.get("created_at")
#             if not j.get("created_at_human"):
#                 try:
#                     if isinstance(ca, (int, float)):
#                         dt = datetime.datetime.utcfromtimestamp(float(ca))
#                         j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                         j["created_at_human"] = dt.astimezone().strftime("%Y-%m-%d %H:%M:%S")
#                     elif isinstance(ca, str) and ca:
#                         # try to normalize to human if present
#                         try:
#                             dt = datetime.datetime.fromisoformat(ca.replace("Z",""))
#                             j["created_at_human"] = dt.astimezone().strftime("%Y-%m-%d %H:%M:%S")
#                         except Exception:
#                             pass
#                     else:
#                         mtime = os.path.getmtime(fp)
#                         dt = datetime.datetime.fromtimestamp(mtime)
#                         j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                         j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#                 except Exception:
#                     pass

#             out.append({
#                 "run_id": j.get("run_id"),
#                 "title": (j.get("title") or "untitled"),
#                 "created_at": j.get("created_at", ""),
#                 "created_at_human": j.get("created_at_human", ""),
#                 "overall_rating": j.get("overall_rating", ""),
#                 "_path": fp,
#             })

#         out.sort(key=lambda r: (r.get("created_at") or "", r.get("_path") or ""), reverse=True)

#     return out

# def _open_history_run_by_id(run_id: str) -> bool:
#     """Open a history run by its run_id. Returns True if loaded."""
#     if not run_id: return False
#     recs = _load_all_history()
#     match = next((r for r in recs if r.get("run_id") == run_id), None)
#     if not match: return False

#     # Load JSON content again (S3/local), then set session
#     try:
#         if "_key" in match and _s3_enabled():
#             txt = read_text_key(match["_key"], "")
#             if not txt: return False
#             jj = json.loads(txt)
#         else:
#             path = match.get("_path")
#             if not path or not os.path.exists(path): return False
#             with open(path, "r", encoding="utf-8") as f:
#                 jj = json.load(f)
#     except Exception:
#         return False

#     st.session_state.script_text      = jj.get("script_text","")
#     st.session_state.base_stem        = jj.get("title","untitled")
#     st.session_state.data             = jj.get("data",{})
#     st.session_state.heading_ranges   = jj.get("heading_ranges",[])
#     st.session_state.spans_by_param   = jj.get("spans_by_param",{})
#     st.session_state.param_choice     = None  # <-- NO HIGHLIGHTS by default
#     st.session_state.source_docx_path = jj.get("source_docx_path")
#     st.session_state.review_ready     = True
#     st.session_state["aoi_match_ranges"] = jj.get("aoi_match_ranges", {})
#     st.session_state.ui_mode          = "review"
#     return True

# def _render_recents_centerpane():
#     st.subheader("📄 Recents")
#     q = st.text_input("Filter by title…", "")
#     cols = st.columns([1, 4])
#     with cols[0]:
#         if st.button("← Back"):
#             st.session_state.ui_mode = "home"; st.rerun()

#     recs = _load_all_history()
#     ql = q.strip().lower()
#     if ql: recs = [r for r in recs if ql in (r.get("title","").lower())]
#     if not recs: st.caption("No history yet."); st.stop()

#     for rec in recs:
#         run_id = rec.get("run_id"); title = rec.get("title") or "(untitled)"
#         created_h = rec.get("created_at_human",""); overall = rec.get("overall_rating","")
#         st.markdown(f"""
#         <a class="rec-card" href="?open={run_id}">
#           <div class="rec-title">{title}</div>
#           <div class="rec-meta">{created_h}</div>
#           <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
#         </a>
#         """, unsafe_allow_html=True)

# # ---------- Optional S3 debug panel ----------
# def _render_history_debug():
#     if not _SHOW_HISTORY_DEBUG:
#         return
#     with st.sidebar.expander("✅ History S3 debug (temporary)", expanded=True):
#         info = _s3_health_summary()
#         st.code(json.dumps({
#             "endpoint": _RP_ENDPOINT,
#             "bucket": _RP_BUCKET,
#             "prefix": f"{HISTORY_DIR}/",
#         }, indent=2))
#         if not _s3_enabled():
#             st.warning("S3 disabled — using local filesystem.")
#             return
#         try:
#             cli = _s3_client()
#             resp = cli.list_objects_v2(Bucket=_RP_BUCKET, Prefix=f"{HISTORY_DIR}/", MaxKeys=1000)
#             count = len(resp.get("Contents", []))
#             st.success(f"List OK. Returned {count} object(s) in first page.")
#             json_keys = [c["Key"] for c in resp.get("Contents", []) if c["Key"].endswith(".json")]
#             for k in json_keys[:50]:
#                 st.code(k)
#             if not json_keys:
#                 st.info("No *.json snapshots under the history prefix. Recents will appear empty until a run saves a JSON snapshot.")
#             else:
#                 # Try reading one JSON
#                 sample = json_keys[0]
#                 st.caption(f"Reading sample JSON: {sample}")
#                 try:
#                     txt = read_text_key(sample, "")
#                     jj = json.loads(txt) if txt else {}
#                     st.write({"title": jj.get("title"), "run_id": jj.get("run_id"), "overall": jj.get("overall_rating")})
#                 except Exception as e:
#                     st.error(f"Unable to read/parse sample JSON: {e}")
#         except Exception as e:
#             st.error(f"List error: {e}")

# # ---------- Sidebar ----------
# with st.sidebar:
#     if st.button("🆕 New review", use_container_width=True):
#         fp = st.session_state.get("flattened_docx_path")
#         if fp and os.path.exists(fp):
#             try: os.remove(fp)
#             except Exception: pass
#         for k in ["review_ready","script_text","base_stem","data","spans_by_param","param_choice",
#                   "source_docx_path","heading_ranges","flattened_docx_path","flatten_used"]:
#             st.session_state[k] = (
#                 False if k=="review_ready"
#                 else "" if k in ("script_text","base_stem")
#                 else {} if k=="spans_by_param"
#                 else [] if k=="heading_ranges"
#                 else None if k in ("source_docx_path","flattened_docx_path")
#                 else False if k=="flatten_used"
#                 else None
#             )
#         st.session_state.ui_mode = "home"
#         _clear_query_params()
#         st.rerun()

#     if st.button("📁 Recents", use_container_width=True):
#         st.session_state.ui_mode = "recents"
#         _clear_query_params()
#         st.rerun()

#     # Optional debug panel
#     _render_history_debug()

# # ---------- Input screen ----------
# def render_home():
#     st.subheader("🎬 Script Source")

#     # Two tabs: Upload OR Paste text
#     tab_upload, tab_paste = st.tabs(["Upload file", "Paste text"])

#     uploaded_file = None
#     uploaded_name = None
#     uploaded_key  = None

#     def _safe_stem(s: str, fallback: str = "pasted_script") -> str:
#         s = (s or "").strip()
#         if not s:
#             return fallback
#         s = re.sub(r"[^A-Za-z0-9._\\-]+", "_", s)
#         s = s.strip("._-") or fallback
#         return s

#     # --- Upload tab ---
#     with tab_upload:
#         up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
#         if up is not None:
#             file_bytes = up.read()
#             suffix = os.path.splitext(up.name)[1].lower()
#             uploaded_key = f"{SCRIPTS_DIR}/{up.name}"
#             save_bytes_key(uploaded_key, file_bytes)
#             with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
#                 tmp.write(file_bytes)
#                 uploaded_file = tmp.name
#             uploaded_name = os.path.splitext(os.path.basename(up.name))[0] or "uploaded_script"

#     # --- Paste tab ---
#     with tab_paste:
#         paste_title = st.text_input("Title (optional)", placeholder="e.g., my_script")
#         pasted_text = st.text_area(
#             "Paste your script text here",
#             height=360,
#             placeholder="Paste the full script text (we’ll analyze this as-is)."
#         )

#     # Run Review works for BOTH
#     if st.button("🚀 Run Review", type="primary", use_container_width=True):
#         base_stem = "uploaded_script"
#         source_docx_path = None
#         heading_ranges: List[Tuple[int,int]] = []
#         script_text = ""

#         if pasted_text and pasted_text.strip():
#             base_stem = _safe_stem(paste_title, "pasted_script")
#             script_text = pasted_text
#             pasted_key = f"{SCRIPTS_DIR}/{base_stem}.txt"
#             save_text_key(pasted_key, script_text)
#             source_docx_path = pasted_key
#             heading_ranges = []

#         elif uploaded_file:
#             base_stem = uploaded_name or "uploaded_script"

#             if uploaded_file.lower().endswith(".docx"):
#                 try:
#                     left_text, used_left = extract_left_column_script_or_default(uploaded_file)
#                 except Exception:
#                     left_text, used_left = "", False

#                 if used_left and left_text.strip():
#                     script_text = left_text
#                     source_docx_path = uploaded_file
#                     heading_ranges = []
#                 else:
#                     path_to_use = uploaded_file
#                     if _docx_contains_tables(path_to_use):
#                         flat = flatten_docx_tables_to_longtext(path_to_use)
#                         st.session_state.flattened_docx_path = flat
#                         st.session_state.flatten_used = True
#                         path_to_use = flat
#                     script_text, heading_ranges = build_docx_text_with_meta(path_to_use)
#                     source_docx_path = path_to_use
#             else:
#                 script_text = load_script_file(uploaded_file)
#                 source_docx_path = uploaded_key or uploaded_file

#         else:
#             st.warning("Please upload a script **or** paste text in the second tab.")
#             st.stop()

#         if len(script_text.strip()) < 50:
#             st.error("Extracted text looks too short. Please check your input.")
#             st.stop()

#         with st.spinner("Running analysis…"):
#             try:
#                 review_text = run_review_multi(
#                     script_text=script_text,
#                     prompts_dir=PROMPTS_DIR,
#                     temperature=0.0
#                 )
#             finally:
#                 if uploaded_file and not (isinstance(source_docx_path, str) and os.path.exists(source_docx_path)):
#                     try: os.remove(uploaded_file)
#                     except Exception: pass

#         data = extract_review_json(review_text)
#         if not data:
#             st.error("JSON not detected in model output.")
#             st.stop()

#         st.session_state.script_text      = script_text
#         st.session_state.base_stem        = base_stem
#         st.session_state.data             = data
#         st.session_state.heading_ranges   = heading_ranges
#         st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
#         st.session_state.param_choice     = None
#         st.session_state.source_docx_path = source_docx_path
#         st.session_state.review_ready     = True
#         st.session_state.ui_mode          = "review"

#         _save_history_snapshot(
#             title=base_stem,
#             data=data,
#             script_text=script_text,
#             source_docx_path=source_docx_path,
#             heading_ranges=heading_ranges,
#             spans_by_param=st.session_state.spans_by_param,
#             aoi_match_ranges=st.session_state.get("aoi_match_ranges", {})
#         )

#         _clear_query_params()
#         st.rerun()

# # ---------- Results screen ----------
# def render_review():
#     script_text     = st.session_state.script_text
#     data            = st.session_state.data
#     spans_by_param  = st.session_state.spans_by_param
#     scores: Dict[str,int] = (data or {}).get("scores", {}) or {}
#     source_docx_path: Optional[str] = st.session_state.source_docx_path

#     # If our source_docx_path is an S3 key/url, ensure we have a local copy for rendering
#     docx_local: Optional[str] = None
#     preferred = st.session_state.get("flattened_docx_path") if st.session_state.get("flatten_used") else source_docx_path
#     if preferred:
#         if not os.path.exists(preferred):
#             docx_local = ensure_local_copy(preferred)
#         else:
#             docx_local = preferred

#     left, center, right = st.columns([1.1, 2.7, 1.4], gap="large")

#     with left:
#         st.subheader("Final score")
#         ordered = [p for p in PARAM_ORDER if p in scores]
#         df = pd.DataFrame({"Parameter": ordered, "Score (1–10)": [scores.get(p, "") for p in ordered]})
#         st.dataframe(df, hide_index=True, use_container_width=True)
#         st.markdown(f'**Overall:** {data.get("overall_rating","—")}/10')
#         st.divider()

#         strengths = (data or {}).get("strengths") or []
#         if not strengths:
#             per = (data or {}).get("per_parameter", {}) or {}
#             best = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
#             for name, sc in best:
#                 if sc >= 8 and name in per:
#                     exp = _sanitize_editor_text((per[name] or {}).get("explanation", "") or "")
#                     first = re.split(r"(?<=[.!?])\s+", exp.strip())[0] if exp else f"Consistently strong {name.lower()}."
#                     strengths.append(f"{name}: {first}")
#                 if len(strengths) >= 3: break

#         def _bullets(title: str, items):
#             st.markdown(f"**{title}**")
#             for s in (items or []):
#                 if isinstance(s, str) and s.strip(): st.write("• " + _sanitize_editor_text(s))
#             if not items: st.write("• —")

#         _bullets("Strengths", strengths)
#         _bullets("Weaknesses", data.get("weaknesses"))
#         _bullets("Suggestions", data.get("suggestions"))
#         _bullets("Drop-off Risks", data.get("drop_off_risks"))
#         st.markdown("**Viral Quotient**"); st.write(_sanitize_editor_text(data.get("viral_quotient","—")))

#     with right:
#         st.subheader("Parameters")
#         st.markdown('<div class="param-row">', unsafe_allow_html=True)
#         for p in [p for p in PARAM_ORDER if p in scores]:
#             if st.button(p, key=f"chip_{p}", help="Show inline AOI highlights for this parameter"):
#                 st.session_state.param_choice = p
#         st.markdown('</div>', unsafe_allow_html=True)

#         sel = st.session_state.param_choice
#         if sel:
#             blk = (data.get("per_parameter", {}) or {}).get(sel, {}) or {}
#             st.markdown(f"**{sel} — Score:** {scores.get(sel,'—')}/10")

#             if blk.get("explanation"):
#                 st.markdown("**Why this score**"); st.write(_sanitize_editor_text(blk["explanation"]))
#             if blk.get("weakness") and blk["weakness"] != "Not present":
#                 st.markdown("**Weakness**"); st.write(_sanitize_editor_text(blk["weakness"]))
#             if blk.get("suggestion") and blk["suggestion"] != "Not present":
#                 st.markdown("**Suggestion**"); st.write(_sanitize_editor_text(blk["suggestion"]))

#             aoi = blk.get("areas_of_improvement") or []
#             if aoi:
#                 st.markdown("**Areas of Improvement**")
#                 for i, item in enumerate(aoi, 1):
#                     popover_fn = getattr(st, "popover", None)
#                     aid = f"{sel.replace(' ','_')}-AOI-{i}"
#                     s_e_map = st.session_state.get("aoi_match_ranges", {})
#                     if aid in s_e_map:
#                         s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                         line = (matched_line if len(matched_line) <= 320 else matched_line[:300] + "…")
#                     else:
#                         line = _sanitize_editor_text(item.get('quote_verbatim',''))
#                     issue = _sanitize_editor_text(item.get('issue',''))
#                     fix   = _sanitize_editor_text(item.get('fix',''))
#                     why   = _sanitize_editor_text(item.get('why_this_helps',''))
#                     label = f"Issue {i}"
#                     if callable(popover_fn):
#                         with popover_fn(label):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)
#                     else:
#                         with st.expander(label, expanded=False):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)

#             if blk.get("summary"):
#                 st.markdown("**Summary**"); st.write(_sanitize_editor_text(blk["summary"]))

#     with center:
#         st.subheader("Script with inline highlights")
#         # choose spans for selected parameter (or none if None)
#         spans = st.session_state.spans_by_param.get(st.session_state.param_choice, []) if st.session_state.param_choice else []

#         # Build AOI payload for popup: { aid: {line, issue, fix, why} }
#         aoi_payload: Dict[str, Dict[str, str]] = {}
#         data_per = (data or {}).get("per_parameter") or {}
#         s_e_map = st.session_state.get("aoi_match_ranges", {})
#         sel = st.session_state.param_choice

#         def _mk_line(aid: str, fallback_q: str = "") -> str:
#             if aid in s_e_map:
#                 s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                 return matched_line if len(matched_line) <= 320 else matched_line[:300] + "…"
#             return _sanitize_editor_text(fallback_q or "")

#         def _collect(param_name: str):
#             blk = (data_per.get(param_name) or {})
#             for i, item in enumerate(blk.get("areas_of_improvement") or [], 1):
#                 aid = f"{param_name.replace(' ','_')}-AOI-{i}"
#                 aoi_payload[aid] = {
#                     "line": _mk_line(aid, (item or {}).get("quote_verbatim","")),
#                     "issue": _sanitize_editor_text((item or {}).get("issue","")),
#                     "fix": _sanitize_editor_text((item or {}).get("fix","")),
#                     "why": _sanitize_editor_text((item or {}).get("why_this_helps","")),
#                 }

#         if sel: _collect(sel)
#         else:
#             for pn in [p for p in PARAM_ORDER if p in data_per]:
#                 _collect(pn)

#         payload_json = json.dumps(aoi_payload, ensure_ascii=False)

#         frame_theme_css = """
#         <style>
#           :root{ --m7-surface:#eef2f7; --m7-on-surface:#0f172a; --m7-border:rgba(15,23,42,.14); }
#           @media (prefers-color-scheme: dark){
#             :root{ --m7-surface:#2f333a; --m7-on-surface:#ffffff; --m7-border:rgba(255,255,255,.18); }
#             body { background: transparent !important; }
#           }
#           .docxwrap{ background:var(--m7-surface); color:var(--m7-on-surface); border:1px solid var(--m7-border); border-radius:12px; padding:16px 14px 18px; }
#           .docxwrap, .docxwrap * { color:var(--m7-on-surface) !important; }
#           .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); }
#         </style>
#         """
#         tooltip_css = """
#         <style>
#         .aoi-pop { position:absolute; max-width:520px; min-width:320px; background:var(--m7-surface); border:1px solid var(--m7-border);
#           border-radius:10px; box-shadow:0 10px 25px rgba(0,0,0,.12); padding:12px 14px; z-index:9999; transform:translateY(-8px);
#           color:var(--m7-on-surface); }
#         .aoi-pop h4 { margin:0 0 .35rem 0; font-size:.95rem; }
#         .aoi-pop p  { margin:.15rem 0; line-height:1.5; }
#         .aoi-pop .muted { opacity:.85; font-size:.85rem; }
#         .aoi-arrow { position:absolute; left:50%; transform:translateX(-50%); bottom:-7px; width:0;height:0;border-left:7px solid transparent;
#           border-right:7px solid transparent;border-top:7px solid var(--m7-border); }
#         .aoi-arrow::after{ content:""; position:absolute; left:-6px; top:-7px; width:0;height:0; border-left:6px solid transparent;
#           border-right:6px solid transparent;border-top:6px solid var(--m7-surface); }
#         </style>
#         """

#         # Select rendering source (DOCX with highlights if we have a local path)
#         if docx_local and os.path.splitext(docx_local)[1].lower() == ".docx":
#             html_core = render_docx_html_with_highlights(
#                 docx_local,
#                 merge_overlaps_and_adjacent(script_text, spans)
#             )
#         else:
#             from html import escape as _esc
#             orig = script_text
#             spans2 = [s for s in merge_overlaps_and_adjacent(orig, spans) if s[0] < s[1]]
#             spans2.sort(key=lambda x: x[0])
#             cur = 0; buf: List[str] = []
#             for s,e,c,aid in spans2:
#                 if s > cur: buf.append(_esc(orig[cur:s]))
#                 buf.append(
#                     f'<mark class="aoi-mark" data-aid="{aid}" '
#                     f'style="background:{c}33;border:1px solid {c};border-radius:3px;padding:0 2px;">'
#                     f'{_esc(orig[s:e])}</mark>'
#                 )
#                 cur = e
#             if cur < len(orig): buf.append(_esc(orig[cur:]))
#             html_core = (
#                 '<div class="docxwrap"><p style="white-space:pre-wrap; '
#                 'line-height:1.7; font-family:ui-serif, Georgia, Times New Roman, serif;">'
#                 + "".join(buf) +
#                 '</p></div>'
#             )

#         html_shell = """
# %%FRAME_THEME_CSS%%
# %%TOOLTIP_CSS%%
# <div id="m7-doc">%%HTML_CORE%%</div>
# <div id="aoi-pop" class="aoi-pop" style="display:none;">
#   <div id="aoi-pop-content"></div>
#   <div class="aoi-arrow"></div>
# </div>
# <script>
# (function(){
#   const AOI = __PAYLOAD__;
#   const wrap = document.getElementById('m7-doc');
#   const pop  = document.getElementById('aoi-pop');
#   const body = document.getElementById('aoi-pop-content');

#   function resizeIframe() {
#     try {
#       const h = Math.max(document.documentElement.scrollHeight, document.body.scrollHeight);
#       if (window.frameElement) { window.frameElement.style.height = (h + 20) + 'px'; window.frameElement.style.width  = '100%'; }
#     } catch(e) {}
#   }
#   window.addEventListener('load', resizeIframe);
#   window.addEventListener('resize', resizeIframe);

#   function hide(){ pop.style.display='none'; }
#   function showFor(mark){
#     const aid = mark.getAttribute('data-aid');
#     const d = AOI[aid]; if(!d) return;
#     body.innerHTML =
#       (d.line  ? '<p><strong>Line:</strong> '  + d.line  + '</p>' : '') +
#       (d.issue ? '<p><strong>Issue:</strong> ' + d.issue + '</p>' : '') +
#       (d.fix   ? '<p><strong>Fix:</strong> '   + d.fix   + '</p>' : '') +
#       (d.why   ? '<p class="muted">'           + d.why   + '</p>' : '');
#     pop.style.display = 'block';

#     const r = mark.getBoundingClientRect();
#     const scY = window.scrollY || document.documentElement.scrollTop;
#     const scX = window.scrollX || document.documentElement.scrollLeft;
#     let top  = r.top + scY - pop.offsetHeight - 10;
#     let left = r.left + scX + r.width/2 - pop.offsetWidth/2;
#     if (top < 8) top = r.bottom + scY + 10;
#     if (left < 8) left = 8;
#     pop.style.top  = top + 'px';
#     pop.style.left = left + 'px';

#     resizeIframe();
#   }

#   wrap.addEventListener('click', (e) => {
#     const m = e.target.closest('.aoi-mark');
#     if(!m){ hide(); return; }
#     if(pop.style.display === 'block'){ hide(); }
#     showFor(m);
#     e.stopPropagation();
#   });

#   document.addEventListener('click', (e) => {
#     if(!e.target.closest('.aoi-pop') && !e.target.closest('.aoi-mark')) hide();
#   });
# })();
# </script>
# """
#         html_shell = (
#             html_shell
#             .replace("%%FRAME_THEME_CSS%%", frame_theme_css)
#             .replace("%%TOOLTIP_CSS%%", tooltip_css)
#             .replace("%%HTML_CORE%%", html_core)
#             .replace("__PAYLOAD__", payload_json)
#         )

#         components.html(html_shell, height=400, scrolling=False)

# # ---------- Router & query param open ----------
# _open_qp = _get_query_param("open")
# if _open_qp and _open_history_run_by_id(_open_qp):
#     _clear_query_params()  # avoid re-opening on subsequent reruns

# mode = st.session_state.ui_mode
# if mode == "recents":
#     _render_recents_centerpane()
# elif mode == "review" and st.session_state.review_ready:
#     render_review()
# else:
#     render_home()








###########################################










# import os, re, glob, json, tempfile, difflib, uuid, datetime, shutil, time
# from pathlib import Path
# from typing import Dict, Any, List, Tuple, Optional

# import streamlit as st
# import pandas as pd
# import streamlit.components.v1 as components  # for inline HTML/JS popup

# # ---- utils & engine ----
# from utils1 import (
#     extract_review_json,
#     PARAM_ORDER,
#     load_script_file,
#     extract_left_column_script_or_default,  # <-- NEW import for left-column extractor
# )
# from review_engine_multi import run_review_multi

# # ---- DOCX rendering imports ----
# from docx import Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # =========================
# # RunPod S3 (S3-ONLY helpers)
# # =========================
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError

# # Accept both AWS_* and RUNPOD_* style envs / st.secrets (prefer secrets if set)
# def _get_env(key: str, default: str = "") -> str:
#     try:
#         v2 = st.secrets.get(key)  # type: ignore[attr-defined]
#         if isinstance(v2, str) and v2.strip():
#             return v2.strip()
#     except Exception:
#         pass
#     v = os.getenv(key, "")
#     if v:
#         return v.strip()
#     return (default or "").strip()

# # Primary config (REQUIRED)
# _RP_ENDPOINT = _get_env("RUNPOD_S3_ENDPOINT", "https://s3.runpod.io")
# _RP_BUCKET   = _get_env("RUNPOD_S3_BUCKET",   "Scriptmodel")
# _RP_REGION   = _get_env("RUNPOD_S3_REGION") or _get_env("AWS_DEFAULT_REGION") or "us-east-1"

# # Credentials: prefer AWS_* if present; else accept RUNPOD_* fallbacks
# _AK = _get_env("AWS_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY")
# _SK = _get_env("AWS_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_KEY")
# _ST = _get_env("AWS_SESSION_TOKEN")  # optional

# # Options
# _FORCE_PATH = (_get_env("RUNPOD_S3_FORCE_PATH_STYLE") or "true").lower() in {"1","true","yes"}
# _USE_SSL    = (_get_env("RUNPOD_S3_USE_SSL") or "true").lower() in {"1","true","yes"}
# _VERIFY_SSL = (_get_env("RUNPOD_S3_VERIFY_SSL") or "true").lower() in {"1","true","yes"}

# def _require_s3_config():
#     missing = []
#     if not _RP_ENDPOINT: missing.append("RUNPOD_S3_ENDPOINT")
#     if not _RP_BUCKET:   missing.append("RUNPOD_S3_BUCKET")
#     if not _AK:          missing.append("AWS_ACCESS_KEY_ID/RUNPOD_S3_ACCESS_KEY_ID")
#     if not _SK:          missing.append("AWS_SECRET_ACCESS_KEY/RUNPOD_S3_SECRET_ACCESS_KEY")
#     if missing:
#         raise RuntimeError("S3 is required but not configured. Missing: " + ", ".join(missing))

# @st.cache_resource(show_spinner=False)
# def _s3_client():
#     _require_s3_config()
#     session_kwargs = dict(
#         aws_access_key_id=_AK,
#         aws_secret_access_key=_SK,
#         region_name=_RP_REGION or None,
#     )
#     if _ST:
#         session_kwargs["aws_session_token"] = _ST

#     # s3v4 + path style are common requirements for S3-compatible services
#     cfg = Config(
#         signature_version="s3v4",
#         s3={"addressing_style": "path" if _FORCE_PATH else "auto"},
#         retries={"max_attempts": 10, "mode": "adaptive"},
#         read_timeout=90,
#         connect_timeout=15,
#         tcp_keepalive=True,
#     )
#     return boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         use_ssl=_USE_SSL,
#         verify=_VERIFY_SSL,
#         config=cfg,
#         **session_kwargs,
#     )

# def _s3_healthcheck() -> None:
#     """Hard fail if S3 isn’t reachable (strict S3-only)."""
#     _require_s3_config()
#     try:
#         _s3_client().list_objects_v2(Bucket=_RP_BUCKET, MaxKeys=1)
#     except Exception as e:
#         raise RuntimeError(f"S3 healthcheck failed for bucket '{_RP_BUCKET}' at '{_RP_ENDPOINT}': {e}")

# def save_text_key(key: str, text: str) -> str:
#     """S3-only put (with backoff). Returns s3:// URL on success."""
#     key = key.lstrip("/")
#     last_err: Optional[Exception] = None
#     for attempt in range(1, 7):
#         try:
#             _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=text.encode("utf-8"), ContentType="text/plain; charset=utf-8")
#             return f"s3://{_RP_BUCKET}/{key}"
#         except (ClientError, EndpointConnectionError) as e:
#             status = getattr(e, "response", {}).get("ResponseMetadata", {}).get("HTTPStatusCode")
#             if isinstance(e, EndpointConnectionError) or status in (500, 502, 503, 504):
#                 last_err = e
#                 time.sleep(min(2 ** attempt, 12))
#                 continue
#             raise
#     raise RuntimeError(f"PutObject failed after retries (Bucket={_RP_BUCKET}, Key={key}). Last error: {last_err}")

# def save_bytes_key(key: str, data: bytes) -> str:
#     """S3-only put (with backoff). Returns s3:// URL on success."""
#     key = key.lstrip("/")
#     last_err: Optional[Exception] = None
#     for attempt in range(1, 7):
#         try:
#             _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=data)
#             return f"s3://{_RP_BUCKET}/{key}"
#         except (ClientError, EndpointConnectionError) as e:
#             status = getattr(e, "response", {}).get("ResponseMetadata", {}).get("HTTPStatusCode")
#             if isinstance(e, EndpointConnectionError) or status in (500, 502, 503, 504):
#                 last_err = e
#                 time.sleep(min(2 ** attempt, 12))
#                 continue
#             raise
#     raise RuntimeError(f"PutObject failed after retries (Bucket={_RP_BUCKET}, Key={key}). Last error: {last_err}")

# def read_text_key(key: str, default: str = "") -> str:
#     key = key.lstrip("/")
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read().decode("utf-8", errors="ignore")
#     except Exception:
#         return default

# def read_bytes_key(key: str) -> Optional[bytes]:
#     key = key.lstrip("/")
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read()
#     except Exception:
#         return None

# def list_prefix(prefix: str) -> List[str]:
#     """
#     List object KEYS under prefix (S3 only).
#     """
#     out: List[str] = []
#     token = None
#     s3_prefix = prefix.rstrip("/") + "/"
#     while True:
#         kwargs = {"Bucket": _RP_BUCKET, "Prefix": s3_prefix, "MaxKeys": 1000}
#         if token:
#             kwargs["ContinuationToken"] = token
#         try:
#             resp = _s3_client().list_objects_v2(**kwargs)
#         except (ClientError, EndpointConnectionError, NoCredentialsError):
#             break
#         for c in resp.get("Contents", []):
#             k = c.get("Key", "")
#             if k.endswith(".json"):
#                 out.append(k)
#         token = resp.get("NextContinuationToken")
#         if not token:
#             break
#     return out

# def presigned_url(key: str, expires: int = 3600) -> Optional[str]:
#     try:
#         return _s3_client().generate_presigned_url(
#             "get_object",
#             Params={"Bucket": _RP_BUCKET, "Key": key.lstrip("/")},
#             ExpiresIn=expires
#         )
#     except ClientError:
#         return None

# def ensure_local_copy(key_or_path: str) -> Optional[str]:
#     """
#     For DOCX/PDF parsing we need a filesystem path.
#     If a local path exists, return it; otherwise treat as S3 key and download to temp.
#     """
#     if key_or_path and os.path.exists(key_or_path):
#         return key_or_path
#     key = key_or_path
#     if key.startswith("s3://"):
#         # s3://bucket/path/to/file -> path/to/file
#         parts = key.split("/", 3)
#         key = parts[3] if len(parts) >= 4 else ""
#     data = read_bytes_key(key)
#     if data is None:
#         return None
#     fd, tmp = tempfile.mkstemp(suffix=os.path.splitext(key)[1] or "")
#     os.close(fd)
#     with open(tmp, "wb") as f:
#         f.write(data)
#     return tmp

# def _s3_health_summary() -> dict:
#     info = {
#         "endpoint": _RP_ENDPOINT,
#         "bucket": _RP_BUCKET,
#         "region": _RP_REGION,
#         "has_keys": bool(_AK and _SK),
#     }
#     try:
#         _s3_client().list_objects_v2(Bucket=_RP_BUCKET, Prefix="health/", MaxKeys=1)
#         info["status"] = "ok"
#     except Exception as e:
#         info["status"] = f"error: {getattr(e, 'response', {}).get('Error', {}).get('Code', str(e))}"
#     return info

# # ---------- Folders (all under Scriptmodel/) ----------
# BASE_PREFIX = "Scriptmodel"
# SCRIPTS_DIR = f"{BASE_PREFIX}/scripts"
# PROMPTS_DIR = f"{BASE_PREFIX}/prompts"
# OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
# HISTORY_DIR = f"{OUTPUT_DIR}/_history"

# # ---------- Colors ----------
# PARAM_COLORS: Dict[str, str] = {
#     "Suspense Building":              "#ff6b6b",
#     "Language/Tone":                  "#6b8cff",
#     "Intro + Main Hook/Cliffhanger":  "#ffb86b",
#     "Story Structure + Flow":         "#a78bfa",
#     "Pacing":                         "#f43f5e",
#     "Mini-Hooks (30–60s)":            "#eab308",
#     "Outro (Ending)":                 "#8b5cf6",
#     "Grammar & Spelling":             "#10b981",  # NEW: distinct highlight color for grammar/spelling
# }

# # ---------- Config ----------
# STRICT_MATCH_ONLY = False

# # ---------- App config ----------
# st.set_page_config(page_title="M7 — Grammarly UI", page_icon="🕵️", layout="wide")

# # Hard S3 health-gate: stop app if S3 not ready
# try:
#     _s3_healthcheck()
# except Exception as e:
#     st.error(f"🔴 Runpod S3 required but not reachable.\n\n{e}")
#     st.stop()

# # ---------- Header patch & CSS ----------
# def render_app_title():
#     st.markdown(
#         '<h1 class="app-title">Viral Script Reviewer</h1>',
#         unsafe_allow_html=True
#     )
#     st.markdown("""
#     <style>
#     html { color-scheme: light dark; }

#     :root{
#       --m7-surface: #eef2f7;
#       --m7-on-surface: #0f172a;
#       --m7-border: rgba(15,23,42,.14);
#       --sep: #e5e7eb;
#     }
#     @media (prefers-color-scheme: dark){
#       :root{
#         --m7-surface: #2f333a;
#         --m7-on-surface: #ffffff;
#         --m7-border: rgba(255,255,255,.18);
#         --sep: #2a2f37;
#       }
#     }

#     .stApp .block-container { padding-top: 4.25rem !important; }
#     .app-title{
#       font-weight: 700; font-size: 2.1rem; line-height: 1.3;
#       margin: 0 0 1rem 0; padding-left: 40px !important; padding-top: .25rem !important;
#       white-space: normal; word-break: break-word; hyphens: auto; overflow: visible;
#       position: relative !important; z-index: 10 !important;
#     }
#     [data-testid="collapsedControl"] { z-index: 6 !important; }
#     header[data-testid="stHeader"], .stAppHeader { background: transparent !important; box-shadow: none !important; }
#     @media (min-width: 992px){ .app-title { padding-left: 0 !important; } }

#     div[data-testid="column"]:nth-of-type(1){position:relative;}
#     div[data-testid="column"]:nth-of-type(1)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     div[data-testid="column"]:nth-of-type(2){position:relative;}
#     div[data-testid="column"]:nth-of-type(2)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}

#     .m7-card{
#       background: var(--m7-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 14px 16px;
#       color: var(--m7-on-surface);
#     }
#     .m7-card, .m7-card * { color: var(--m7-on-surface) !important; }

#     .docxwrap{
#       background: var(--m7-surface);
#       color: var(--m7-on-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 16px 14px 18px;
#     }
#     .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#     .docxwrap .h1, .docxwrap .h2, .docxwrap .h3 { font-weight:700; margin:10px 0 6px; }
#     .docxwrap .h1 { font-size: 1.3rem; border-bottom: 2px solid currentColor; padding-bottom: 4px; }
#     .docxwrap .h2 { font-size: 1.15rem; border-bottom: 1px solid currentColor; padding-bottom: 3px; }
#     .docxwrap .h3 { font-size: 1.05rem; }
#     .docxwrap p { margin: 10px 0; line-height: 1.7; font-family: ui-serif, Georgia, "Times New Roman", serif; }
#     .docxwrap table { border-collapse: collapse; width: 100%; margin: 12px 0; }
#     .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); padding:8px; vertical-align:top; line-height:1.6; }
#     .docxwrap mark{ padding:0 2px; border-radius:3px; border:1px solid var(--m7-border); cursor: pointer; }

#     .rec-card{
#       display:block; text-decoration:none !important;
#       background: var(--m7-surface);
#       border:1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 14px 16px; margin: 10px 0 16px;
#       box-shadow: 0 1px 2px rgba(0,0,0,.06);
#       color: var(--m7-on-surface) !important;
#       transition: filter .1s ease, transform .02s ease;
#     }
#     .rec-card:hover{ filter: brightness(1.02); }
#     .rec-card:active{ transform: translateY(1px); }
#     .rec-card, .rec-card * { color: var(--m7-on-surface) !important; }
#     .rec-title{font-weight:600; margin-bottom:.25rem;}
#     .rec-meta{opacity:.85 !important; font-size:12.5px; margin-bottom:.4rem;}
#     .rec-row{display:flex; align-items:center; justify-content:space-between; gap:12px;}

#     .stTextInput>div>div,
#     .stTextArea>div>div,
#     .stNumberInput>div>div,
#     .stDateInput>div>div,
#     .stTimeInput>div>div,
#     .stFileUploader>div,
#     div[data-baseweb="select"]{
#       background: var(--m7-surface) !important;
#       border: 1px solid var(--m7-border) !important;
#       border-radius: 10px !important;
#       color: var(--m7-on-surface) !important;
#     }
#     .stTextInput input,
#     .stTextArea textarea,
#     .stNumberInput input,
#     .stDateInput input,
#     .stTimeInput input,
#     .stFileUploader div,
#     div[data-baseweb="select"] *{
#       color: var(--m7-on-surface) !important;
#     }
#     .stTextInput input::placeholder,
#     .stTextArea textarea::placeholder{ color: rgba(16,24,39,.55) !important; }
#     @media (prefers-color-scheme: dark){
#       .stTextInput input::placeholder,
#       .stTextArea textarea::placeholder{ color: rgba(255,255,255,.75) !important; }
#     }
#     div[data-testid="stFileUploaderDropzone"] label span { color: var(--m7-on-surface) !important; opacity: 1 !important; }
#     div[data-testid="stFileUploaderDropzone"] label { color: var(--m7-on-surface) !important; }

#     .stMarkdown pre,
#     pre[class*="language-"],
#     .stCodeBlock{
#       background: var(--m7-surface) !important;
#       color: var(--m7-on-surface) !important;
#       border: 1px solid var(--m7-border) !important;
#       border-radius: 12px !important;
#       padding: 12px 14px !important;
#       overflow:auto;
#     }
#     .stMarkdown pre code{ background: transparent !important; color: inherit !important; }

#     div[data-testid="stDataFrame"]{
#       background: var(--m7-surface);
#       border: 1px solid var(--m7-border);
#       border-radius: 12px;
#       padding: 6px 8px;
#       color: var(--m7-on-surface);
#     }
#     div[data-testid="stDataFrame"] * { color: var(--m7-on-surface) !important; }

#     </style>
#     """, unsafe_allow_html=True)

# render_app_title()

# # ---------- Session ----------
# for key, default in [
#     ("review_ready", False),
#     ("script_text", ""),
#     ("base_stem", ""),
#     ("data", None),
#     ("spans_by_param", {}),
#     ("param_choice", None),
#     ("source_docx_path", None),
#     ("heading_ranges", []),
#     ("flattened_docx_path", None),
#     ("flatten_used", False),
#     ("ui_mode", "home"),
# ]:
#     st.session_state.setdefault(key, default)

# # ---------- helpers for query params (compat across Streamlit versions) ----------
# def _get_query_param(key: str) -> Optional[str]:
#     val = None
#     try:
#         val = st.query_params.get(key)  # type: ignore[attr-defined]
#     except Exception:
#         q = st.experimental_get_query_params()
#         v = q.get(key)
#         if isinstance(v, list): val = v[0] if v else None
#         else: val = v
#     return val

# def _clear_query_params():
#     try:
#         st.query_params.clear()  # type: ignore[attr-defined]
#     except Exception:
#         st.experimental_set_query_params()

# # ---------- Sanitizer ----------
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )
# def _sanitize_editor_text(s: Optional[str]) -> str:
#     if not s: return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•)\s*', '', t, flags=re.M)
#     t = re.sub(r'^\s*[-*]\s+', '• ', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# # ---------- DOCX traversal ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# # ---------- Auto-flatten ----------
# def _docx_contains_tables(path: str) -> bool:
#     doc = Document(path)
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Table):
#             return True
#     return False

# def _copy_paragraph(dest_doc: Document, src_para: Paragraph):
#     p = dest_doc.add_paragraph()
#     try:
#         if src_para.style and src_para.style.name:
#             p.style = src_para.style.name
#     except Exception:
#         pass
#     for run in src_para.runs:
#         r = p.add_run(run.text or "")
#         r.bold = run.bold
#         r.italic = run.italic
#         r.underline = run.underline
#     return p

# def flatten_docx_tables_to_longtext(source_path: str) -> str:
#     src = Document(source_path)
#     new = Document()
#     for blk in _iter_docx_blocks(src):
#         if isinstance(blk, Paragraph):
#             _copy_paragraph(new, blk)
#         else:
#             seen_tc_ids = set()
#             for row in blk.rows:
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     if tc_id in seen_tc_ids:
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     for p in cell.paragraphs:
#                         _copy_paragraph(new, p)
#                 new.add_paragraph("")
#             new.add_paragraph("")
#     fd, tmp_path = tempfile.mkstemp(suffix=".docx"); os.close(fd); new.save(tmp_path)
#     return tmp_path

# # ---------- Build plain text + heading ranges ----------
# def build_docx_text_with_meta(docx_path: str) -> Tuple[str, List[Tuple[int,int]]]:
#     doc = Document(docx_path)
#     out: List[str] = []
#     heading_ranges: List[Tuple[int,int]] = []
#     current_offset = 0

#     def _append_and_advance(s: str):
#         nonlocal current_offset
#         out.append(s); current_offset += len(s)

#     seen_tc_ids: set = set()
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             para_text = "".join(run.text or "" for run in blk.runs)
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             if sty.startswith("heading"):
#                 start = current_offset; end = start + len(para_text)
#                 heading_ranges.append((start, end))
#             _append_and_advance(para_text); _append_and_advance("\n")
#         else:
#             for row in blk.rows:
#                 row_cell_tcs = []
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     row_cell_tcs.append((tc_id, cell))
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     if tc_id in seen_tc_ids:
#                         if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     cell_text_parts: List[str] = []
#                     for i, p in enumerate(cell.paragraphs):
#                         t = "".join(r.text or "" for r in p.runs)
#                         sty = (p.style.name or "").lower() if p.style else ""
#                         if sty.startswith("heading"):
#                             hs = current_offset + sum(len(x) for x in cell_text_parts)
#                             he = hs + len(t)
#                             heading_ranges.append((hs, he))
#                         cell_text_parts.append(t)
#                         if i != len(cell.paragraphs) - 1:
#                             cell_text_parts.append("\n")
#                     cell_text = "".join(cell_text_parts)
#                     _append_and_advance(cell_text)
#                     if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                 _append_and_advance("\n")
#             _append_and_advance("\n")

#     return "".join(out), heading_ranges

# def _wrap_inline(safe_text: str, run) -> str:
#     out = safe_text
#     if getattr(run, "underline", False): out = f"<u>{out}</u>"
#     if getattr(run, "italic", False): out = f"<em>{out}</em>"
#     if getattr(run, "bold", False): out = f"<strong>{out}</strong>"
#     return out

# # ---------- Invisible/bridge chars ----------
# _BRIDGE_CHARS = set("\u200b\u200c\u200d\u2060\ufeff\xa0\u00ad")

# # ---------- DOCX -> HTML with highlights (includes data-aid) ----------
# def render_docx_html_with_highlights(docx_path: str,
#                                      highlight_spans: List[Tuple[int,int,str,str]]) -> str:
#     """
#     highlight_spans: list of (start, end, color, aid)
#     """
#     doc = Document(docx_path)
#     spans = [s for s in highlight_spans if s[0] < s[1]]
#     spans.sort(key=lambda x: x[0])

#     cur_span = 0
#     current_offset = 0

#     def esc(s: str) -> str:
#         return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

#     def open_mark_if_needed(html_parts, mark_state, color, end, aid):
#         if not mark_state["open"]:
#             html_parts.append(
#                 f'<mark class="aoi-mark" data-aid="{aid}" '
#                 f'style="background:{color}33;border:1px solid {color};border-radius:3px;padding:0 2px;">'
#             )
#             mark_state.update(open=True, end=end, color=color, aid=aid)

#     def close_mark_if_open(html_parts, mark_state):
#         if mark_state["open"]:
#             html_parts.append('</mark>')
#             mark_state.update(open=False, end=None, color=None, aid=None)

#     def emit_run_text(run_text: str, run, html_parts: List[str], mark_state: Dict[str, Any]):
#         nonlocal cur_span, current_offset
#         t = run_text or ""; i = 0
#         while i < len(t):
#             next_start, next_end, color, next_aid = None, None, None, None
#             if cur_span < len(spans):
#                 next_start, next_end, color, next_aid = spans[cur_span]

#             if not mark_state["open"]:
#                 if cur_span >= len(spans) or current_offset + (len(t) - i) <= next_start:
#                     chunk = t[i:]; html_parts.append(_wrap_inline(esc(chunk), run)); current_offset += len(chunk); break
#                 if current_offset < next_start:
#                     take = next_start - current_offset
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take; continue
#                 open_mark_if_needed(html_parts, mark_state, color, next_end, next_aid)
#             else:
#                 take = min(mark_state["end"] - current_offset, len(t) - i)
#                 if take > 0:
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take
#                 if current_offset >= mark_state["end"]:
#                     close_mark_if_open(html_parts, mark_state)
#                     cur_span += 1

#     html: List[str] = ['<div class="docxwrap">']
#     seen_tc_ids: set = set()

#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             open_tag = '<div class="h1">' if sty.startswith("heading 1") else \
#                        '<div class="h2">' if sty.startswith("heading 2") else \
#                        '<div class="h3">' if sty.startswith("heading 3") else "<p>"
#             close_tag = "</div>" if sty.startswith("heading") else "</p>"
#             html.append(open_tag)
#             for run in blk.runs:
#                 emit_run_text(run.text or "", run, html, mark_state)
#             close_mark_if_open(html, mark_state)
#             html.append(close_tag)
#             current_offset += 1  # '\n'

#         else:
#             html.append("<table>")
#             for row in blk.rows:
#                 html.append("<tr>")
#                 row_cell_tcs = [(id(cell._tc), cell) for cell in row.cells]
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     html.append("<td>")
#                     if tc_id not in seen_tc_ids:
#                         seen_tc_ids.add(tc_id)
#                         for p_idx, p in enumerate(cell.paragraphs):
#                             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#                             html.append("<div>")
#                             for run in p.runs:
#                                 emit_run_text(run.text or "", run, html, mark_state)
#                             close_mark_if_open(html, mark_state)
#                             html.append("</div>")
#                             if p_idx != len(cell.paragraphs) - 1:
#                                 current_offset += 1
#                     html.append("</td>")
#                     if idx != len(row_cell_tcs) - 1: current_offset += 1  # '\t'
#                 html.append("</tr>"); current_offset += 1   # row '\n'
#             html.append("</table>"); current_offset += 1     # extra '\n'

#     html.append("</div>")
#     return "".join(html)

# # ---------- Matching / span utilities ----------
# def _normalize_keep_len(s: str) -> str:
#     trans = {
#         "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
#         "\u2013": "-", "\u2014": "-",
#         "\xa0": " ",
#         "\u200b": " ", "\u200c": " ", "\u200d": " ", "\u2060": " ",
#         "\ufeff": " ", "\u00ad": " ",
#     }
#     return (s or "").translate(str.maketrans(trans))

# def _tokenize(s: str) -> List[str]: return re.findall(r"\w+", (s or "").lower())

# def _iter_sentences_with_spans(text: str) -> List[Tuple[int,int,str]]:
#     spans = []
#     for m in re.finditer(r'[^.!?]+[.!?]+|\Z', text, flags=re.S):
#         s, e = m.start(), m.end()
#         seg = text[s:e]
#         if seg.strip(): spans.append((s, e, seg))
#     return spans

# def _squash_ws(s: str) -> str: return re.sub(r"\s+", " ", s or "").strip()

# def _clean_quote_for_match(q: str) -> str:
#     if not q: return ""
#     q = _normalize_keep_len(q).strip()
#     q = re.sub(r'^[\'"“”‘’\[\(\{<…\-\–\—\s]+', '', q)
#     q = re.sub(r'[\'"“”‘’\]\)\}>…\-\–\—\s]+$', '', q)
#     return _squash_ws(q)

# def _snap_and_bridge_to_word(text: str, start: int, end: int, max_bridge: int = 2) -> Tuple[int,int]:
#     n = len(text); s, e = max(0,start), max(start,end)
#     def _is_inv(ch: str) -> bool: return ch in _BRIDGE_CHARS

#     while s > 0:
#         prev = text[s-1]; cur = text[s] if s < n else ""
#         if prev.isalnum() and cur.isalnum(): s -= 1; continue
#         j = s; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1; 
#         if brid and (s-1)>=0 and text[s-1].isalnum() and (j<n and text[j].isalnum()): s -= 1; continue
#         break

#     while e < n:
#         prev = text[e-1] if e>0 else ""; nxt = text[e]
#         if prev.isalnum() and nxt.isalnum(): e += 1; continue
#         j = e; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (e-1)>=0 and text[e-1].isalnum() and (j<n and text[j].isalnum()): e = j + 1; continue
#         break

#     while e < n and text[e] in ',"”’\')]}': e += 1
#     return s, e

# def _heal_split_word_left(text: str, start: int) -> int:
#     i = start
#     if i <= 1 or i >= len(text): return start
#     if text[i-1] != " ": return start
#     j = i - 2
#     while j >= 0 and text[j].isalpha(): j -= 1
#     prev_token = text[j+1:i-1]
#     if len(prev_token) == 1: return i - 2
#     return start

# def _overlaps_any(s: int, e: int, ranges: List[Tuple[int,int]]) -> bool:
#     for rs, re_ in ranges:
#         if e > rs and s < re_: return True
#     return False

# def _fuzzy_window_span(tl: str, nl: str, start: int, w: int) -> Tuple[float, Optional[Tuple[int,int]]]:
#     window = tl[start:start+w]
#     sm = difflib.SequenceMatcher(a=nl, b=window)
#     blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
#     if not blocks: return 0.0, None
#     coverage = sum(b.size for b in blocks) / max(1, len(nl))
#     first_b = min(blocks, key=lambda b: b.b); last_b = max(blocks, key=lambda b: b.b + b.size)
#     s = start + first_b.b; e = start + last_b.b + last_b.size
#     return coverage, (s, e)

# def find_span_smart(text: str, needle: str) -> Optional[Tuple[int,int]]:
#     if not text or not needle: return None
#     t_orig = text; t_norm = _normalize_keep_len(text); n_norm = _clean_quote_for_match(needle)
#     if not n_norm: return None
#     tl = t_norm.lower(); nl = n_norm.lower()

#     i = tl.find(nl)
#     if i != -1:
#         s, e = _snap_and_bridge_to_word(t_orig, i, i + len(nl)); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     m = re.search(re.escape(nl).replace(r"\\ ", r"\\s+"), tl, flags=re.IGNORECASE)
#     if m:
#         s, e = _snap_and_bridge_to_word(t_orig, m.start(), m.end()); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     if not STRICT_MATCH_ONLY and len(nl) >= 12:
#         w = max(60, min(240, len(nl) + 80))
#         best_cov, best_span = 0.0, None
#         step = max(1, w // 2)
#         for start in range(0, max(1, len(tl) - w + 1), step):
#             cov, se = _fuzzy_window_span(tl, nl, start, w)
#             if cov > best_cov: best_cov, best_span = cov, se
#         if best_span and best_cov >= 0.65:
#             s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
#             if s > 0 and t_orig[s-1:s+1].lower() in {" the", " a", " an"}: s -= 1
#             s = _heal_split_word_left(t_orig, s)
#             return (s, e)

#     if not STRICT_MATCH_ONLY:
#         keys = [w for w in _tokenize(nl) if len(w) >= 4][:8]
#         if len(keys) >= 2:
#             kset = set(keys)
#             best_score, best_span = 0.0, None
#             for s, e, seg in _iter_sentences_with_spans(t_norm):
#                 toks = set(_tokenize(seg)); ov = len(kset & toks)
#                 if ov == 0: continue
#                 score = ov / max(2, len(kset)); length_pen = min(1.0, 120 / max(20, e - s)); score *= (0.6 + 0.4 * length_pen)
#                 if score > best_score: best_score, best_span = score, (s, min(e, s + 400))
#             if best_span and best_score >= 0.35:
#                 s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1]); s = _heal_split_word_left(t_orig, s)
#                 return (s, e)
#     return None

# def merge_overlaps(spans: List[Tuple[int,int,str,str]]) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans.sort(key=lambda x: x[0]); out = [spans[0]]
#     for s,e,c,aid in spans[1:]:
#         ps,pe,pc,paid = out[-1]
#         if s <= pe and pc == c and e > pe: out[-1] = (ps, e, pc, paid)
#         else: out.append((s,e,c,aid))
#     return out

# _PUNCT_WS = set(" \t\r\n,.;:!?)('\"[]{}-—–…") | _BRIDGE_CHARS

# def merge_overlaps_and_adjacent(base_text: str,
#                                 spans: List[Tuple[int,int,str,str]],
#                                 max_gap: int = 2) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans = sorted(spans, key=lambda x: x[0]); out = [spans[0]]
#     for s, e, c, aid in spans[1:]:
#         ps, pe, pc, paid = out[-1]
#         if c == pc and s <= pe: out[-1] = (ps, max(pe, e), pc, paid); continue
#         if c == pc and s - pe <= max_gap:
#             gap = base_text[max(0, pe):max(0, s)]
#             if all((ch in _PUNCT_WS) for ch in gap): out[-1] = (ps, e, pc, paid); continue
#         out.append((s, e, c, aid))
#     return out

# # ---------- Heading filters ----------
# def _is_heading_like(q: str) -> bool:
#     if not q: return True
#     s = q.strip()
#     if not re.search(r'[.!?]', s):
#         words = re.findall(r"[A-Za-z]+", s)
#         if 1 <= len(words) <= 7:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.8: return True
#         if s.lower() in {"introduction","voiceover","outro","epilogue","prologue","credits","title","hook","horrifying discovery","final decision"}: return True
#         if len(s) <= 3: return True
#     return False

# def _is_heading_context(script_text: str, s: int, e: int) -> bool:
#     left = script_text.rfind("\\n", 0, s) + 1
#     right = script_text.find("\\n", e); right = len(script_text) if right == -1 else right
#     line = script_text[left:right].strip()
#     if len(line) <= 70 and not re.search(r'[.!?]', line):
#         words = re.findall(r"[A-Za-z]+", line)
#         if 1 <= len(words) <= 8:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.7: return True
#     return False

# def _tighten_to_quote(script_text: str, span: Tuple[int,int], quote: str) -> Tuple[int,int]:
#     if not span or not quote: return span
#     s, e = span
#     if e <= s or s < 0 or e > len(script_text): return span
#     window = script_text[s:e]; win_norm = _normalize_keep_len(window).lower(); q_norm = _clean_quote_for_match(quote).lower()
#     if not q_norm: return span
#     i = win_norm.find(q_norm)
#     if i == -1:
#         m = re.search(re.escape(q_norm).replace(r"\\ ", r"\\s+"), win_norm, flags=re.IGNORECASE)
#         if not m: return span
#         i, j = m.start(), m.end()
#     else:
#         j = i + len(q_norm)
#     s2, e2 = s + i, s + j
#     s2, e2 = _snap_and_bridge_to_word(script_text, s2, e2); s2 = _heal_split_word_left(script_text, s2)
#     if s2 >= s and e2 <= e and e2 > s2: return (s2, e2)
#     return span

# def build_spans_by_param(script_text: str, data: dict, heading_ranges: Optional[List[Tuple[int,int]]] = None) -> Dict[str, List[Tuple[int,int,str,str]]]:
#     heading_ranges = heading_ranges or []
#     raw = (data or {}).get("per_parameter", {}) or {}
#     per: Dict[str, Dict[str, Any]] = {k:(v or {}) for k,v in raw.items()}
#     spans_map: Dict[str, List[Tuple[int,int,str,str]]] = {p: [] for p in PARAM_ORDER}
#     st.session_state["aoi_match_ranges"] = {}

#     for p in spans_map.keys():
#         color = PARAM_COLORS.get(p, "#ffd54f")
#         blk = per.get(p, {}) or {}
#         aois = blk.get("areas_of_improvement") or []
#         for idx, item in enumerate(aois, start=1):
#             raw_q = (item or {}).get("quote_verbatim", "") or ""
#             q = _sanitize_editor_text(raw_q)
#             clean = _clean_quote_for_match(re.sub(r"^[•\\-\\d\\.\\)\\s]+", "", q).strip())
#             if not clean: continue
#             if _is_heading_like(clean): continue

#             pos = find_span_smart(script_text, clean)
#             if not pos: continue
#             pos = _tighten_to_quote(script_text, pos, raw_q)
#             s, e = pos

#             if heading_ranges and _overlaps_any(s, e, heading_ranges): continue
#             if _is_heading_context(script_text, s, e): continue

#             aid = f"{p.replace(' ','_')}-AOI-{idx}"
#             spans_map[p].append((s, e, color, aid))
#             st.session_state["aoi_match_ranges"][aid] = (s, e)
#     return spans_map

# # ---------- History (S3-only) ----------
# def _maybe_copy_docx_to_history(source_docx_path: Optional[str], run_id: str) -> Optional[str]:
#     """Copy the DOCX used for rendering into outputs/_history so Recents can re-render identically (S3-only)."""
#     try:
#         if not source_docx_path:
#             return None
#         # If we have a local temp, upload it
#         if os.path.exists(source_docx_path):
#             with open(source_docx_path, "rb") as f:
#                 save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#             return f"{HISTORY_DIR}/{run_id}.docx"

#         # Else assume it's an S3 key/URL; if it's an s3:// URL, keep as-is; if it's a key, keep it.
#         return source_docx_path
#     except Exception:
#         return None

# def _save_history_snapshot(title: str, data: dict, script_text: str,
#                            source_docx_path: Optional[str], heading_ranges: List[Tuple[int,int]],
#                            spans_by_param: Dict[str, List[Tuple[int,int,str,str]]],
#                            aoi_match_ranges: Dict[str, Tuple[int,int]]):
#     run_id = str(uuid.uuid4()); now = datetime.datetime.now()
#     created_at_iso = now.replace(microsecond=0).isoformat()
#     created_at_human = now.strftime("%Y-%m-%d %H:%M:%S")

#     # Copy stable DOCX alongside history JSON (so Recents uses identical renderer)
#     stable_docx_key_or_path = _maybe_copy_docx_to_history(source_docx_path, run_id)

#     blob = {
#         "run_id": run_id, "title": title or "untitled",
#         "created_at": created_at_iso, "created_at_human": created_at_human,
#         "overall_rating": (data or {}).get("overall_rating", ""),
#         "scores": (data or {}).get("scores", {}),
#         "data": data or {}, "script_text": script_text or "",
#         "source_docx_path": stable_docx_key_or_path or source_docx_path,
#         "heading_ranges": heading_ranges or [],
#         "spans_by_param": spans_by_param or {},
#         "aoi_match_ranges": aoi_match_ranges or {},
#     }

#     # Persist JSON (S3)
#     out_name = f"{created_at_iso.replace(':','-')}__{run_id}.json"
#     out_key = f"{HISTORY_DIR}/{out_name}"
#     save_text_key(out_key, json.dumps(blob, ensure_ascii=False, indent=2))

# def _load_all_history() -> List[dict]:
#     out: List[dict] = []
#     keys = sorted(list_prefix(HISTORY_DIR), reverse=True)
#     for key in keys:
#         try:
#             txt = read_text_key(key, "")
#             if not txt:
#                 continue
#             j = json.loads(txt)
#             j["_key"] = key
#             if not j.get("created_at_human") and j.get("created_at"):
#                 try:
#                     dt = datetime.datetime.fromisoformat(j["created_at"])
#                     j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#                 except Exception:
#                     pass
#             out.append(j)
#         except Exception:
#             continue
#     out.sort(key=lambda r: r.get("created_at") or "", reverse=True)
#     return out

# def _open_history_run_by_id(run_id: str) -> bool:
#     """Open a history run by its run_id. Returns True if loaded."""
#     if not run_id: return False
#     recs = _load_all_history()
#     match = next((r for r in recs if r.get("run_id") == run_id), None)
#     if not match: return False

#     # Load JSON content again from S3
#     try:
#         txt = read_text_key(match["_key"], "")
#         if not txt: return False
#         jj = json.loads(txt)
#     except Exception:
#         return False

#     st.session_state.script_text      = jj.get("script_text","")
#     st.session_state.base_stem        = jj.get("title","untitled")
#     st.session_state.data             = jj.get("data",{})
#     st.session_state.heading_ranges   = jj.get("heading_ranges",[])
#     st.session_state.spans_by_param   = jj.get("spans_by_param",{})
#     st.session_state.param_choice     = None
#     st.session_state.source_docx_path = jj.get("source_docx_path")
#     st.session_state.review_ready     = True
#     st.session_state["aoi_match_ranges"] = jj.get("aoi_match_ranges", {})
#     st.session_state.ui_mode          = "review"
#     return True

# def _render_recents_centerpane():
#     st.subheader("📄 Recents")
#     q = st.text_input("Filter by title…", "")
#     cols = st.columns([1, 4])
#     with cols[0]:
#         if st.button("← Back"):
#             st.session_state.ui_mode = "home"; st.rerun()

#     recs = _load_all_history()
#     ql = q.strip().lower()
#     if ql: recs = [r for r in recs if ql in (r.get("title","").lower())]
#     if not recs: st.caption("No history yet."); st.stop()

#     # Each card is a clickable <a class="rec-card" href="?open=<run_id>">
#     for rec in recs:
#         run_id = rec.get("run_id"); title = rec.get("title") or "(untitled)"
#         created_h = rec.get("created_at_human",""); overall = rec.get("overall_rating","")
#         st.markdown(f"""
#         <a class="rec-card" href="?open={run_id}">
#           <div class="rec-title">{title}</div>
#           <div class="rec-meta">{created_h}</div>
#           <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
#         </a>
#         """, unsafe_allow_html=True)

# # ---------- Sidebar ----------
# with st.sidebar:
#     if st.button("🆕 New review", use_container_width=True):
#         fp = st.session_state.get("flattened_docx_path")
#         if fp and os.path.exists(fp):
#             try: os.remove(fp)
#             except Exception: pass
#         for k in ["review_ready","script_text","base_stem","data","spans_by_param","param_choice",
#                   "source_docx_path","heading_ranges","flattened_docx_path","flatten_used"]:
#             st.session_state[k] = (
#                 False if k=="review_ready"
#                 else "" if k in ("script_text","base_stem")
#                 else {} if k=="spans_by_param"
#                 else [] if k=="heading_ranges"
#                 else None if k in ("source_docx_path","flattened_docx_path")
#                 else False if k=="flatten_used"
#                 else None
#             )
#         st.session_state.ui_mode = "home"
#         _clear_query_params()
#         st.rerun()

#     if st.button("📁 Recents", use_container_width=True):
#         st.session_state.ui_mode = "recents"
#         _clear_query_params()
#         st.rerun()

# # ---------- Input screen ----------
# def render_home():
#     st.subheader("🎬 Script Source")

#     # Two tabs: Upload OR Paste text
#     tab_upload, tab_paste = st.tabs(["Upload file", "Paste text"])

#     uploaded_file = None
#     uploaded_name = None
#     uploaded_key  = None

#     # helper to create a safe filename for pasted text
#     def _safe_stem(s: str, fallback: str = "pasted_script") -> str:
#         s = (s or "").strip()
#         if not s:
#             return fallback
#         s = re.sub(r"[^A-Za-z0-9._\\-]+", "_", s)
#         s = s.strip("._-") or fallback
#         return s

#     # --- Upload tab ---
#     with tab_upload:
#         up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
#         if up is not None:
#             file_bytes = up.read()
#             suffix = os.path.splitext(up.name)[1].lower()
#             # Persist to S3 under scripts/
#             uploaded_key = f"{SCRIPTS_DIR}/{up.name}"
#             save_bytes_key(uploaded_key, file_bytes)

#             # Local temp for parsing
#             with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
#                 tmp.write(file_bytes)
#                 uploaded_file = tmp.name
#             uploaded_name = os.path.splitext(os.path.basename(up.name))[0] or "uploaded_script"

#     # --- Paste tab ---
#     with tab_paste:
#         paste_title = st.text_input("Title (optional)", placeholder="e.g., my_script")
#         pasted_text = st.text_area(
#             "Paste your script text here",
#             height=360,
#             placeholder="Paste the full script text (we’ll analyze this as-is)."
#         )

#     # Run Review works for BOTH
#     if st.button("🚀 Run Review", type="primary", use_container_width=True):
#         base_stem = "uploaded_script"
#         source_docx_path = None
#         heading_ranges: List[Tuple[int,int]] = []
#         script_text = ""

#         # Prefer pasted text if provided
#         if pasted_text and pasted_text.strip():
#             base_stem = _safe_stem(paste_title, "pasted_script")
#             script_text = pasted_text

#             # Save pasted text so Recents can reopen consistently
#             pasted_key = f"{SCRIPTS_DIR}/{base_stem}.txt"
#             save_text_key(pasted_key, script_text)
#             source_docx_path = pasted_key
#             heading_ranges = []

#         # Else use uploaded file
#         elif uploaded_file:
#             base_stem = uploaded_name or "uploaded_script"

#             if uploaded_file.lower().endswith(".docx"):
#                 try:
#                     left_text, used_left = extract_left_column_script_or_default(uploaded_file)
#                 except Exception:
#                     left_text, used_left = "", False

#                 if used_left and left_text.strip():
#                     script_text = left_text
#                     source_docx_path = uploaded_key or uploaded_file
#                     heading_ranges = []
#                 else:
#                     path_to_use = uploaded_file
#                     if _docx_contains_tables(path_to_use):
#                         flat = flatten_docx_tables_to_longtext(path_to_use)
#                         st.session_state.flattened_docx_path = flat
#                         st.session_state.flatten_used = True
#                         path_to_use = flat
#                     script_text, heading_ranges = build_docx_text_with_meta(path_to_use)
#                     source_docx_path = uploaded_key or uploaded_file
#             else:
#                 script_text = load_script_file(uploaded_file)
#                 source_docx_path = uploaded_key or uploaded_file

#         else:
#             st.warning("Please upload a script **or** paste text in the second tab.")
#             st.stop()

#         if len(script_text.strip()) < 50:
#             st.error("Extracted text looks too short. Please check your input.")
#             st.stop()

#         with st.spinner("Running analysis…"):
#             try:
#                 review_text = run_review_multi(
#                     script_text=script_text,
#                     prompts_dir=PROMPTS_DIR,
#                     temperature=0.0
#                 )
#             finally:
#                 if uploaded_file and os.path.exists(uploaded_file):
#                     try: os.remove(uploaded_file)
#                     except Exception: pass

#         data = extract_review_json(review_text)
#         if not data:
#             st.error("JSON not detected in model output.")
#             st.stop()

#         st.session_state.script_text      = script_text
#         st.session_state.base_stem        = base_stem
#         st.session_state.data             = data
#         st.session_state.heading_ranges   = heading_ranges
#         st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
#         st.session_state.param_choice     = None
#         st.session_state.source_docx_path = source_docx_path
#         st.session_state.review_ready     = True
#         st.session_state.ui_mode          = "review"

#         _save_history_snapshot(
#             title=base_stem,
#             data=data,
#             script_text=script_text,
#             source_docx_path=source_docx_path,
#             heading_ranges=heading_ranges,
#             spans_by_param=st.session_state.spans_by_param,
#             aoi_match_ranges=st.session_state.get("aoi_match_ranges", {})
#         )

#         _clear_query_params()
#         st.rerun()

# # ---------- Results screen ----------
# def render_review():
#     script_text     = st.session_state.script_text
#     data            = st.session_state.data
#     spans_by_param  = st.session_state.spans_by_param
#     scores: Dict[str,int] = (data or {}).get("scores", {}) or {}
#     source_docx_path: Optional[str] = st.session_state.source_docx_path

#     # If our source_docx_path is an S3 key/url, ensure we have a local copy for rendering
#     # AFTER (prefer the flattened copy if we have it in-session)
#     docx_local: Optional[str] = None
#     preferred = st.session_state.get("flattened_docx_path") if st.session_state.get("flatten_used") else source_docx_path
#     if preferred:
#         if not os.path.exists(preferred):
#             docx_local = ensure_local_copy(preferred)
#         else:
#             docx_local = preferred

#     left, center, right = st.columns([1.1, 2.7, 1.4], gap="large")

#     with left:
#         st.subheader("Final score")
#         ordered = [p for p in PARAM_ORDER if p in scores]
#         df = pd.DataFrame({"Parameter": ordered, "Score (1–10)": [scores.get(p, "") for p in ordered]})
#         st.dataframe(df, hide_index=True, use_container_width=True)
#         st.markdown(f'**Overall:** {data.get("overall_rating","—")}/10')
#         st.divider()

#         strengths = (data or {}).get("strengths") or []
#         if not strengths:
#             per = (data or {}).get("per_parameter", {}) or {}
#             best = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
#             for name, sc in best:
#                 if sc >= 8 and name in per:
#                     exp = _sanitize_editor_text((per[name] or {}).get("explanation", "") or "")
#                     first = re.split(r"(?<=[.!?])\\s+", exp.strip())[0] if exp else f"Consistently strong {name.lower()}."
#                     strengths.append(f"{name}: {first}")
#                 if len(strengths) >= 3: break

#         def _bullets(title: str, items):
#             st.markdown(f"**{title}**")
#             for s in (items or []):
#                 if isinstance(s, str) and s.strip(): st.write("• " + _sanitize_editor_text(s))
#             if not items: st.write("• —")

#         _bullets("Strengths", strengths)
#         _bullets("Weaknesses", data.get("weaknesses"))
#         _bullets("Suggestions", data.get("suggestions"))
#         _bullets("Drop-off Risks", data.get("drop_off_risks"))
#         st.markdown("**Viral Quotient**"); st.write(_sanitize_editor_text(data.get("viral_quotient","—")))

#     with right:
#         st.subheader("Parameters")
#         st.markdown('<div class="param-row">', unsafe_allow_html=True)
#         for p in [p for p in PARAM_ORDER if p in scores]:
#             if st.button(p, key=f"chip_{p}", help="Show inline AOI highlights for this parameter"):
#                 st.session_state.param_choice = p
#         st.markdown('</div>', unsafe_allow_html=True)

#         sel = st.session_state.param_choice
#         if sel:
#             blk = (data.get("per_parameter", {}) or {}).get(sel, {}) or {}
#             st.markdown(f"**{sel} — Score:** {scores.get(sel,'—')}/10")

#             if blk.get("explanation"):
#                 st.markdown("**Why this score**"); st.write(_sanitize_editor_text(blk["explanation"]))
#             if blk.get("weakness") and blk["weakness"] != "Not present":
#                 st.markdown("**Weakness**"); st.write(_sanitize_editor_text(blk["weakness"]))
#             if blk.get("suggestion") and blk["suggestion"] != "Not present":
#                 st.markdown("**Suggestion**"); st.write(_sanitize_editor_text(blk["suggestion"]))

#             aoi = blk.get("areas_of_improvement") or []
#             if aoi:
#                 st.markdown("**Areas of Improvement**")
#                 for i, item in enumerate(aoi, 1):
#                     popover_fn = getattr(st, "popover", None)
#                     aid = f"{sel.replace(' ','_')}-AOI-{i}"
#                     s_e_map = st.session_state.get("aoi_match_ranges", {})
#                     if aid in s_e_map:
#                         s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                         line = (matched_line if len(matched_line) <= 320 else matched_line[:300] + "…")
#                     else:
#                         line = _sanitize_editor_text(item.get('quote_verbatim',''))
#                     issue = _sanitize_editor_text(item.get('issue',''))
#                     fix   = _sanitize_editor_text(item.get('fix',''))
#                     why   = _sanitize_editor_text(item.get('why_this_helps',''))
#                     label = f"Issue {i}"
#                     if callable(popover_fn):
#                         with popover_fn(label):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)
#                     else:
#                         with st.expander(label, expanded=False):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)

#             if blk.get("summary"):
#                 st.markdown("**Summary**"); st.write(_sanitize_editor_text(blk["summary"]))

#     with center:
#         st.subheader("Script with inline highlights")
#         # choose spans for selected parameter (or all if None)
#         spans = st.session_state.spans_by_param.get(st.session_state.param_choice, []) if st.session_state.param_choice else []

#         # Build AOI payload for popup: { aid: {line, issue, fix, why} }
#         aoi_payload: Dict[str, Dict[str, str]] = {}
#         data_per = (data or {}).get("per_parameter") or {}
#         s_e_map = st.session_state.get("aoi_match_ranges", {})
#         sel = st.session_state.param_choice

#         def _mk_line(aid: str, fallback_q: str = "") -> str:
#             if aid in s_e_map:
#                 s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                 return matched_line if len(matched_line) <= 320 else matched_line[:300] + "…"
#             return _sanitize_editor_text(fallback_q or "")

#         def _collect(param_name: str):
#             blk = (data_per.get(param_name) or {})
#             for i, item in enumerate(blk.get("areas_of_improvement") or [], 1):
#                 aid = f"{param_name.replace(' ','_')}-AOI-{i}"
#                 aoi_payload[aid] = {
#                     "line": _mk_line(aid, (item or {}).get("quote_verbatim","")),
#                     "issue": _sanitize_editor_text((item or {}).get("issue","")),
#                     "fix": _sanitize_editor_text((item or {}).get("fix","")),
#                     "why": _sanitize_editor_text((item or {}).get("why_this_helps","")),
#                 }

#         if sel: _collect(sel)
#         else:
#             for pn in [p for p in PARAM_ORDER if p in data_per]:
#                 _collect(pn)

#         payload_json = json.dumps(aoi_payload, ensure_ascii=False)

#         frame_theme_css = """
#         <style>
#           :root{
#             --m7-surface: #eef2f7;
#             --m7-on-surface: #0f172a;
#             --m7-border: rgba(15,23,42,.14);
#           }
#           @media (prefers-color-scheme: dark){
#             :root{
#               --m7-surface: #2f333a;
#               --m7-on-surface: #ffffff;
#               --m7-border: rgba(255,255,255,.18);
#             }
#             body { background: transparent !important; }
#           }

#           .docxwrap{
#             background: var(--m7-surface);
#             color: var(--m7-on-surface);
#             border: 1px solid var(--m7-border);
#             border-radius: 12px;
#             padding: 16px 14px 18px;
#           }
#           .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#           .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); }
#         </style>
#         """

#         tooltip_css = """
#         <style>
#         .aoi-pop {
#           position: absolute; max-width: 520px; min-width: 320px;
#           background: var(--m7-surface); border: 1px solid var(--m7-border); border-radius: 10px;
#           box-shadow: 0 10px 25px rgba(0,0,0,.12);
#           padding: 12px 14px; z-index: 9999; transform: translateY(-8px);
#           color: var(--m7-on-surface);
#         }
#         .aoi-pop h4 { margin: 0 0 .35rem 0; font-size: .95rem; }
#         .aoi-pop p  { margin: .15rem 0; line-height: 1.5; }
#         .aoi-pop .muted { opacity:.85; font-size:.85rem; }
#         .aoi-arrow {
#           position:absolute; left:50%; transform:translateX(-50%);
#           bottom:-7px; width:0;height:0;border-left:7px solid transparent;
#           border-right:7px solid transparent;border-top:7px solid var(--m7-border);
#         }
#         .aoi-arrow::after{
#           content:""; position:absolute; left:-6px; top:-7px; width:0;height:0;
#           border-left:6px solid transparent;border-right:6px solid transparent;border-top:6px solid var(--m7-surface);
#         }
#         </style>
#         """

#         # Select rendering source (DOCX with highlights if we have a local path)
#         if docx_local and os.path.splitext(docx_local)[1].lower() == ".docx":
#             html_core = render_docx_html_with_highlights(
#                 docx_local,
#                 merge_overlaps_and_adjacent(script_text, spans)
#             )
#         else:
#             from html import escape as _esc
#             orig = script_text
#             spans2 = [s for s in merge_overlaps_and_adjacent(orig, spans) if s[0] < s[1]]
#             spans2.sort(key=lambda x: x[0])
#             cur = 0; buf: List[str] = []
#             for s,e,c,aid in spans2:
#                 if s > cur: buf.append(_esc(orig[cur:s]))
#                 buf.append(
#                     f'<mark class="aoi-mark" data-aid="{aid}" '
#                     f'style="background:{c}33;border:1px solid {c};border-radius:3px;padding:0 2px;">'
#                     f'{_esc(orig[s:e])}</mark>'
#                 )
#                 cur = e
#             if cur < len(orig): buf.append(_esc(orig[cur:]))
#             html_core = (
#                 '<div class="docxwrap"><p style="white-space:pre-wrap; '
#                 'line-height:1.7; font-family:ui-serif, Georgia, Times New Roman, serif;">'
#                 + "".join(buf) +
#                 '</p></div>'
#             )

#         # Popup + autosize JS shell
#         html_shell = """
# %%FRAME_THEME_CSS%%
# %%TOOLTIP_CSS%%
# <div id="m7-doc">%%HTML_CORE%%</div>
# <div id="aoi-pop" class="aoi-pop" style="display:none;">
#   <div id="aoi-pop-content"></div>
#   <div class="aoi-arrow"></div>
# </div>
# <script>
# (function(){
#   const AOI = __PAYLOAD__;
#   const wrap = document.getElementById('m7-doc');
#   const pop  = document.getElementById('aoi-pop');
#   const body = document.getElementById('aoi-pop-content');

#   function resizeIframe() {
#     try {
#       const h = Math.max(
#         document.documentElement.scrollHeight,
#         document.body.scrollHeight
#       );
#       if (window.frameElement) {
#         window.frameElement.style.height = (h + 20) + 'px';
#         window.frameElement.style.width  = '100%';
#       }
#     } catch(e) {}
#   }
#   window.addEventListener('load', resizeIframe);
#   window.addEventListener('resize', resizeIframe);

#   function hide(){ pop.style.display='none'; }
#   function showFor(mark){
#     const aid = mark.getAttribute('data-aid');
#     const d = AOI[aid]; if(!d) return;
#     body.innerHTML =
#       (d.line  ? '<p><strong>Line:</strong> '  + d.line  + '</p>' : '') +
#       (d.issue ? '<p><strong>Issue:</strong> ' + d.issue + '</p>' : '') +
#       (d.fix   ? '<p><strong>Fix:</strong> '   + d.fix   + '</p>' : '') +
#       (d.why   ? '<p class="muted">'           + d.why   + '</p>' : '');
#     pop.style.display = 'block';

#     const r = mark.getBoundingClientRect();
#     const scY = window.scrollY || document.documentElement.scrollTop;
#     const scX = window.scrollX || document.documentElement.scrollLeft;
#     let top  = r.top + scY - pop.offsetHeight - 10;
#     let left = r.left + scX + r.width/2 - pop.offsetWidth/2;
#     if (top < 8) top = r.bottom + scY + 10;
#     if (left < 8) left = 8;
#     pop.style.top  = top + 'px';
#     pop.style.left = left + 'px';

#     resizeIframe();
#   }

#   wrap.addEventListener('click', (e) => {
#     const m = e.target.closest('.aoi-mark');
#     if(!m){ hide(); return; }
#     if(pop.style.display === 'block'){ hide(); }
#     showFor(m);
#     e.stopPropagation();
#   });

#   document.addEventListener('click', (e) => {
#     if(!e.target.closest('.aoi-pop') && !e.target.closest('.aoi-mark')) hide();
#   });
# })();
# </script>
# """
#         html_shell = (
#             html_shell
#             .replace("%%FRAME_THEME_CSS%%", frame_theme_css)
#             .replace("%%TOOLTIP_CSS%%", tooltip_css)
#             .replace("%%HTML_CORE%%", html_core)
#             .replace("__PAYLOAD__", payload_json)
#         )

#         components.html(html_shell, height=400, scrolling=False)

# # ---------- Router & query param open ----------
# _open_qp = _get_query_param("open")
# if _open_qp and _open_history_run_by_id(_open_qp):
#     _clear_query_params()  # avoid re-opening on subsequent reruns

# mode = st.session_state.ui_mode
# if mode == "recents":
#     _render_recents_centerpane()
# elif mode == "review" and st.session_state.review_ready:
#     render_review()
# else:
#     render_home()











#####################################




### History





# import os, re, glob, json, tempfile, difflib, uuid, datetime, shutil
# from pathlib import Path
# from typing import Dict, Any, List, Tuple, Optional

# import streamlit as st
# import pandas as pd
# import streamlit.components.v1 as components  # for inline HTML/JS popup

# # ---- utils & engine ----
# from utils1 import (
#     extract_review_json,
#     PARAM_ORDER,
#     load_script_file,
#     extract_left_column_script_or_default,  # <-- NEW import for left-column extractor
# )
# from review_engine_multi import run_review_multi

# # ---- DOCX rendering imports ----
# from docx import Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # =========================
# # RunPod S3 (inline helpers) — FIXED
# # =========================
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError

# # Accept both AWS_* and RUNPOD_* style envs / st.secrets
# def _get_env(key: str, default: str = "") -> str:
#     v = os.getenv(key, "")
#     if v:
#         return v.strip()
#     try:
#         v2 = st.secrets.get(key)
#         if isinstance(v2, str):
#             return v2.strip()
#     except Exception:
#         pass
#     return (default or "").strip()

# # Primary config
# _RP_ENDPOINT = _get_env("RUNPOD_S3_ENDPOINT")
# _RP_BUCKET   = _get_env("RUNPOD_S3_BUCKET")
# _RP_REGION   = _get_env("RUNPOD_S3_REGION") or _get_env("AWS_DEFAULT_REGION") or ""

# # Credentials: prefer AWS_* if present; else accept RUNPOD_* fallbacks
# _AK = _get_env("AWS_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY")
# _SK = _get_env("AWS_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_KEY")
# _ST = _get_env("AWS_SESSION_TOKEN")  # optional

# # Options
# _FORCE_PATH = (_get_env("RUNPOD_S3_FORCE_PATH_STYLE") or "true").lower() in {"1","true","yes"}
# _USE_SSL    = (_get_env("RUNPOD_S3_USE_SSL") or "true").lower() in {"1","true","yes"}
# _VERIFY_SSL = (_get_env("RUNPOD_S3_VERIFY_SSL") or "true").lower() in {"1","true","yes"}

# def _s3_enabled() -> bool:
#     return bool(_RP_ENDPOINT and _RP_BUCKET and _AK and _SK)

# @st.cache_resource(show_spinner=False)
# def _s3_client():
#     if not _s3_enabled():
#         return None
#     session_kwargs = dict(
#         aws_access_key_id=_AK,
#         aws_secret_access_key=_SK,
#     )
#     if _ST:
#         session_kwargs["aws_session_token"] = _ST

#     cfg = Config(
#         signature_version="s3v4",
#         s3={"addressing_style": "path" if _FORCE_PATH else "auto"},
#         retries={"max_attempts": 3, "mode": "standard"}
#     )
#     return boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         region_name=_RP_REGION or None,
#         use_ssl=_USE_SSL,
#         verify=_VERIFY_SSL,
#         config=cfg,
#         **session_kwargs,
#     )

# def save_text_key(key: str, text: str) -> str:
#     key = key.lstrip("/")  # ensure no leading slash in S3 key
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "w", encoding="utf-8") as f:
#             f.write(text)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=text.encode("utf-8"))
#     return f"s3://{_RP_BUCKET}/{key}"

# def save_bytes_key(key: str, data: bytes) -> str:
#     key = key.lstrip("/")  # ensure no leading slash in S3 key
#     if not _s3_enabled():
#         os.makedirs(os.path.dirname(key), exist_ok=True)
#         with open(key, "wb") as f:
#             f.write(data)
#         return key
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=data)
#     return f"s3://{_RP_BUCKET}/{key}"

# def read_text_key(key: str, default: str = "") -> str:
#     if not _s3_enabled():
#         try:
#             with open(key, "r", encoding="utf-8") as f:
#                 return f.read()
#         except Exception:
#             return default
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read().decode("utf-8", errors="ignore")
#     except Exception:
#         return default

# def read_bytes_key(key: str) -> Optional[bytes]:
#     if not _s3_enabled():
#         try:
#             with open(key, "rb") as f:
#                 return f.read()
#         except Exception:
#             return None
#     try:
#         resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#         return resp["Body"].read()
#     except Exception:
#         return None

# def list_prefix(prefix: str) -> List[str]:
#     """
#     List object keys under prefix (or local dir paths if not S3).
#     In S3 mode we always return KEYS (not URLs).
#     """
#     if not _s3_enabled():
#         base = prefix if os.path.isdir(prefix) else os.path.dirname(prefix)
#         try:
#             return [os.path.join(base, p) for p in os.listdir(base) if p.endswith(".json")]
#         except Exception:
#             return []

#     out: List[str] = []
#     token = None
#     s3_prefix = prefix.rstrip("/") + "/"
#     try:
#         while True:
#             kwargs = {"Bucket": _RP_BUCKET, "Prefix": s3_prefix}
#             if token:
#                 kwargs["ContinuationToken"] = token
#             resp = _s3_client().list_objects_v2(**kwargs)
#             for c in resp.get("Contents", []):
#                 k = c.get("Key", "")
#                 if k.endswith(".json"):
#                     out.append(k)
#             token = resp.get("NextContinuationToken")
#             if not token:
#                 break
#     except (ClientError, EndpointConnectionError, NoCredentialsError):
#         return []
#     return out

# def presigned_url(key: str, expires: int = 3600) -> Optional[str]:
#     if not _s3_enabled():
#         return None
#     try:
#         return _s3_client().generate_presigned_url(
#             "get_object",
#             Params={"Bucket": _RP_BUCKET, "Key": key},
#             ExpiresIn=expires
#         )
#     except ClientError:
#         return None

# def ensure_local_copy(key_or_path: str) -> Optional[str]:
#     """
#     For DOCX/PDF parsing we need a real filesystem path.
#     If S3 mode, download to a temp file and return that path.
#     """
#     if not _s3_enabled():
#         return key_or_path if os.path.exists(key_or_path) else None

#     key = key_or_path
#     if key.startswith("s3://"):
#         parts = key.split("/", 3)
#         key = parts[3] if len(parts) >= 4 else ""
#     data = read_bytes_key(key)
#     if data is None:
#         return None
#     fd, tmp = tempfile.mkstemp(suffix=os.path.splitext(key)[1] or "")
#     os.close(fd)
#     with open(tmp, "wb") as f:
#         f.write(data)
#     return tmp

# def _s3_health_summary() -> dict:
#     info = {
#         "enabled": _s3_enabled(),
#         "endpoint": _RP_ENDPOINT,
#         "bucket": _RP_BUCKET,
#         "region": _RP_REGION,
#         "has_keys": bool(_AK and _SK),
#     }
#     if not _s3_enabled():
#         info["status"] = "local-mode"
#         return info
#     try:
#         _ = _s3_client().list_objects_v2(Bucket=_RP_BUCKET, Prefix=(f"{OUTPUT_DIR}/_history/").rstrip("/") + "/",
#                                          MaxKeys=1)
#         info["status"] = "ok"
#     except Exception as e:
#         info["status"] = f"error: {getattr(e, 'response', {}).get('Error', {}).get('Code', str(e))}"
#     return info

# # ---------- Folders (all under Scriptmodel/) ----------
# BASE_PREFIX = "Scriptmodel"

# SCRIPTS_DIR = f"{BASE_PREFIX}/scripts"
# PROMPTS_DIR = f"{BASE_PREFIX}/prompts"
# OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
# HISTORY_DIR = f"{OUTPUT_DIR}/_history"

# if not _s3_enabled():
#     for p in (SCRIPTS_DIR, PROMPTS_DIR, OUTPUT_DIR, HISTORY_DIR):
#         Path(p).mkdir(parents=True, exist_ok=True)

# # --- History manifest (to support Recents without ListBucket) ---
# _MANIFEST_KEY = f"{HISTORY_DIR}/_manifest.json"

# def _history_manifest_read() -> List[dict]:
#     """
#     Returns an array of short entries:
#     [{run_id, key, title, created_at, created_at_human, overall_rating}, ...]
#     """
#     txt = read_text_key(_MANIFEST_KEY, "")
#     if not txt:
#         return []
#     try:
#         j = json.loads(txt)
#         return j if isinstance(j, list) else []
#     except Exception:
#         return []

# def _history_manifest_append(entry: dict) -> None:
#     """
#     Appends an entry to the manifest and trims to last 1000 items.
#     This avoids requiring s3:ListBucket for Recents.
#     """
#     arr = _history_manifest_read()
#     arr.append(entry)
#     if len(arr) > 1000:
#         arr = arr[-1000:]
#     save_text_key(_MANIFEST_KEY, json.dumps(arr, ensure_ascii=False))

# # ---------- Colors ----------
# PARAM_COLORS: Dict[str, str] = {
#     "Suspense Building":              "#ff6b6b",
#     "Language/Tone":                  "#6b8cff",
#     "Intro + Main Hook/Cliffhanger":  "#ffb86b",
#     "Story Structure + Flow":         "#a78bfa",
#     "Pacing":                         "#f43f5e",
#     "Mini-Hooks (30–60s)":            "#eab308",
#     "Outro (Ending)":                 "#8b5cf6",
#     "Grammar & Spelling":             "#10b981",
# }

# # ---------- Config ----------
# STRICT_MATCH_ONLY = False

# # ---------- App config ----------
# st.set_page_config(page_title="M7 — Grammarly UI", page_icon="🕵️", layout="wide")

# def render_app_title():
#     st.markdown('<h1 class="app-title">Viral Script Reviewer</h1>', unsafe_allow_html=True)
#     # (CSS omitted for brevity; cosmetic only)

# render_app_title()

# # ---------- Session ----------
# for key, default in [
#     ("review_ready", False),
#     ("script_text", ""),
#     ("base_stem", ""),
#     ("data", None),
#     ("spans_by_param", {}),
#     ("param_choice", None),
#     ("source_docx_path", None),
#     ("heading_ranges", []),
#     ("flattened_docx_path", None),
#     ("flatten_used", False),
#     ("ui_mode", "home"),
# ]:
#     st.session_state.setdefault(key, default)

# # ---------- helpers for query params ----------
# def _get_query_param(key: str) -> Optional[str]:
#     val = None
#     try:
#         val = st.query_params.get(key)
#     except Exception:
#         q = st.experimental_get_query_params()
#         v = q.get(key)
#         if isinstance(v, list): val = v[0] if v else None
#         else: val = v
#     return val

# def _clear_query_params():
#     try:
#         st.query_params.clear()
#     except Exception:
#         st.experimental_set_query_params()

# # ---------- Sanitizer ----------
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )
# def _sanitize_editor_text(s: Optional[str]) -> str:
#     if not s: return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•)\s*', '', t, flags=re.M)
#     t = re.sub(r'^\s*[-*]\s+', '• ', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# # ---------- DOCX traversal ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# # ---------- Auto-flatten ----------
# def _docx_contains_tables(path: str) -> bool:
#     doc = Document(path)
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Table):
#             return True
#     return False

# def _copy_paragraph(dest_doc: Document, src_para: Paragraph):
#     p = dest_doc.add_paragraph()
#     try:
#         if src_para.style and src_para.style.name:
#             p.style = src_para.style.name
#     except Exception:
#         pass
#     for run in src_para.runs:
#         r = p.add_run(run.text or "")
#         r.bold = run.bold
#         r.italic = run.italic
#         r.underline = run.underline
#     return p

# def flatten_docx_tables_to_longtext(source_path: str) -> str:
#     src = Document(source_path)
#     new = Document()
#     for blk in _iter_docx_blocks(src):
#         if isinstance(blk, Paragraph):
#             _copy_paragraph(new, blk)
#         else:
#             seen_tc_ids = set()
#             for row in blk.rows:
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     if tc_id in seen_tc_ids:
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     for p in cell.paragraphs:
#                         _copy_paragraph(new, p)
#                 new.add_paragraph("")
#             new.add_paragraph("")
#     fd, tmp_path = tempfile.mkstemp(suffix=".docx"); os.close(fd); new.save(tmp_path)
#     return tmp_path

# # ---------- Build plain text + heading ranges ----------
# def build_docx_text_with_meta(docx_path: str) -> Tuple[str, List[Tuple[int,int]]]:
#     doc = Document(docx_path)
#     out: List[str] = []
#     heading_ranges: List[Tuple[int,int]] = []
#     current_offset = 0

#     def _append_and_advance(s: str):
#         nonlocal current_offset
#         out.append(s); current_offset += len(s)

#     seen_tc_ids: set = set()
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             para_text = "".join(run.text or "" for run in blk.runs)
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             if sty.startswith("heading"):
#                 start = current_offset; end = start + len(para_text)
#                 heading_ranges.append((start, end))
#             _append_and_advance(para_text); _append_and_advance("\n")
#         else:
#             for row in blk.rows:
#                 row_cell_tcs = []
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     row_cell_tcs.append((tc_id, cell))
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     if tc_id in seen_tc_ids:
#                         if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     cell_text_parts: List[str] = []
#                     for i, p in enumerate(cell.paragraphs):
#                         t = "".join(r.text or "" for r in p.runs)
#                         sty = (p.style.name or "").lower() if p.style else ""
#                         if sty.startswith("heading"):
#                             hs = current_offset + sum(len(x) for x in cell_text_parts)
#                             he = hs + len(t)
#                             heading_ranges.append((hs, he))
#                         cell_text_parts.append(t)
#                         if i != len(cell.paragraphs) - 1:
#                             cell_text_parts.append("\n")
#                     cell_text = "".join(cell_text_parts)
#                     _append_and_advance(cell_text)
#                     if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                 _append_and_advance("\n")
#             _append_and_advance("\n")

#     return "".join(out), heading_ranges

# def _wrap_inline(safe_text: str, run) -> str:
#     out = safe_text
#     if getattr(run, "underline", False): out = f"<u>{out}</u>"
#     if getattr(run, "italic", False): out = f"<em>{out}</em>"
#     if getattr(run, "bold", False): out = f"<strong>{out}</strong>"
#     return out

# # ---------- Invisible/bridge chars ----------
# _BRIDGE_CHARS = set("\u200b\u200c\u200d\u2060\ufeff\xa0\u00ad")

# # ---------- DOCX -> HTML with highlights (includes data-aid) ----------
# def render_docx_html_with_highlights(docx_path: str,
#                                      highlight_spans: List[Tuple[int,int,str,str]]) -> str:
#     doc = Document(docx_path)
#     spans = [s for s in highlight_spans if s[0] < s[1]]
#     spans.sort(key=lambda x: x[0])

#     cur_span = 0
#     current_offset = 0

#     def esc(s: str) -> str:
#         return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

#     def open_mark_if_needed(html_parts, mark_state, color, end, aid):
#         if not mark_state["open"]:
#             html_parts.append(
#                 f'<mark class="aoi-mark" data-aid="{aid}" '
#                 f'style="background:{color}33;border:1px solid {color};border-radius:3px;padding:0 2px;">'
#             )
#             mark_state.update(open=True, end=end, color=color, aid=aid)

#     def close_mark_if_open(html_parts, mark_state):
#         if mark_state["open"]:
#             html_parts.append('</mark>')
#             mark_state.update(open=False, end=None, color=None, aid=None)

#     def emit_run_text(run_text: str, run, html_parts: List[str], mark_state: Dict[str, Any]):
#         nonlocal cur_span, current_offset
#         t = run_text or ""; i = 0
#         while i < len(t):
#             next_start, next_end, color, next_aid = None, None, None, None
#             if cur_span < len(spans):
#                 next_start, next_end, color, next_aid = spans[cur_span]

#             if not mark_state["open"]:
#                 if cur_span >= len(spans) or current_offset + (len(t) - i) <= next_start:
#                     chunk = t[i:]; html_parts.append(_wrap_inline(esc(chunk), run)); current_offset += len(chunk); break
#                 if current_offset < next_start:
#                     take = next_start - current_offset
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take; continue
#                 open_mark_if_needed(html_parts, mark_state, color, next_end, next_aid)
#             else:
#                 take = min(mark_state["end"] - current_offset, len(t) - i)
#                 if take > 0:
#                     chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                     current_offset += take; i += take
#                 if current_offset >= mark_state["end"]:
#                     close_mark_if_open(html_parts, mark_state)
#                     cur_span += 1

#     html: List[str] = ['<div class="docxwrap">']
#     seen_tc_ids: set = set()

#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             open_tag = '<div class="h1">' if sty.startswith("heading 1") else \
#                        '<div class="h2">' if sty.startswith("heading 2") else \
#                        '<div class="h3">' if sty.startswith("heading 3") else "<p>"
#             close_tag = "</div>" if sty.startswith("heading") else "</p>"
#             html.append(open_tag)
#             for run in blk.runs:
#                 emit_run_text(run.text or "", run, html, mark_state)
#             close_mark_if_open(html, mark_state)
#             html.append(close_tag)
#             current_offset += 1  # '\n'

#         else:
#             html.append("<table>")
#             for row in blk.rows:
#                 html.append("<tr>")
#                 row_cell_tcs = [(id(cell._tc), cell) for cell in row.cells]
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     html.append("<td>")
#                     if tc_id not in seen_tc_ids:
#                         seen_tc_ids.add(tc_id)
#                         for p_idx, p in enumerate(cell.paragraphs):
#                             mark_state = {"open": False, "end": None, "color": None, "aid": None}
#                             html.append("<div>")
#                             for run in p.runs:
#                                 emit_run_text(run.text or "", run, html, mark_state)
#                             close_mark_if_open(html, mark_state)
#                             html.append("</div>")
#                             if p_idx != len(cell.paragraphs) - 1:
#                                 current_offset += 1
#                     html.append("</td>")
#                     if idx != len(row_cell_tcs) - 1: current_offset += 1  # '\t'
#                 html.append("</tr>"); current_offset += 1   # row '\n'
#             html.append("</table>"); current_offset += 1     # extra '\n'

#     html.append("</div>")
#     return "".join(html)

# # ---------- Matching / span utilities ----------
# def _normalize_keep_len(s: str) -> str:
#     trans = {
#         "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
#         "\u2013": "-", "\u2014": "-",
#         "\xa0": " ",
#         "\u200b": " ", "\u200c": " ", "\u200d": " ", "\u2060": " ",
#         "\ufeff": " ", "\u00ad": " ",
#     }
#     return (s or "").translate(str.maketrans(trans))

# def _tokenize(s: str) -> List[str]: return re.findall(r"\w+", (s or "").lower())

# def _iter_sentences_with_spans(text: str) -> List[Tuple[int,int,str]]:
#     spans = []
#     for m in re.finditer(r'[^.!?]+[.!?]+|\Z', text, flags=re.S):
#         s, e = m.start(), m.end()
#         seg = text[s:e]
#         if seg.strip(): spans.append((s, e, seg))
#     return spans

# def _squash_ws(s: str) -> str: return re.sub(r"\s+", " ", s or "").strip()

# def _clean_quote_for_match(q: str) -> str:
#     if not q: return ""
#     q = _normalize_keep_len(q).strip()
#     q = re.sub(r'^[\'"“”‘’\[\(\{<…\-\–\—\s]+', '', q)
#     q = re.sub(r'[\'"“”‘’\]\)\}>…\-\–\—\s]+$', '', q)
#     return _squash_ws(q)

# def _snap_and_bridge_to_word(text: str, start: int, end: int, max_bridge: int = 2) -> Tuple[int,int]:
#     n = len(text); s, e = max(0,start), max(start,end)
#     def _is_inv(ch: str) -> bool: return ch in _BRIDGE_CHARS

#     while s > 0:
#         prev = text[s-1]; cur = text[s] if s < n else ""
#         if prev.isalnum() and cur.isalnum(): s -= 1; continue
#         j = s; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1;
#         if brid and (s-1)>=0 and text[s-1].isalnum() and (j<n and text[j].isalnum()): s -= 1; continue
#         break

#     while e < n:
#         prev = text[e-1] if e>0 else ""; nxt = text[e]
#         if prev.isalnum() and nxt.isalnum(): e += 1; continue
#         j = e; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (e-1)>=0 and text[e-1].isalnum() and (j<n and text[j].isalnum()): e = j + 1; continue
#         break

#     while e < n and text[e] in ',"”’\')]}': e += 1
#     return s, e

# def _heal_split_word_left(text: str, start: int) -> int:
#     i = start
#     if i <= 1 or i >= len(text): return start
#     if text[i-1] != " ": return start
#     j = i - 2
#     while j >= 0 and text[j].isalpha(): j -= 1
#     prev_token = text[j+1:i-1]
#     if len(prev_token) == 1: return i - 2
#     return start

# def _overlaps_any(s: int, e: int, ranges: List[Tuple[int,int]]) -> bool:
#     for rs, re_ in ranges:
#         if e > rs and s < re_: return True
#     return False

# def _fuzzy_window_span(tl: str, nl: str, start: int, w: int) -> Tuple[float, Optional[Tuple[int,int]]]:
#     window = tl[start:start+w]
#     sm = difflib.SequenceMatcher(a=nl, b=window)
#     blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
#     if not blocks: return 0.0, None
#     coverage = sum(b.size for b in blocks) / max(1, len(nl))
#     first_b = min(blocks, key=lambda b: b.b); last_b = max(blocks, key=lambda b: b.b + b.size)
#     s = start + first_b.b; e = start + last_b.b + last_b.size
#     return coverage, (s, e)

# def find_span_smart(text: str, needle: str) -> Optional[Tuple[int,int]]:
#     if not text or not needle: return None
#     t_orig = text; t_norm = _normalize_keep_len(text); n_norm = _clean_quote_for_match(needle)
#     if not n_norm: return None
#     tl = t_norm.lower(); nl = n_norm.lower()

#     i = tl.find(nl)
#     if i != -1:
#         s, e = _snap_and_bridge_to_word(t_orig, i, i + len(nl)); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     m = re.search(re.escape(nl).replace(r"\ ", r"\s+"), tl, flags=re.IGNORECASE)
#     if m:
#         s, e = _snap_and_bridge_to_word(t_orig, m.start(), m.end()); s = _heal_split_word_left(t_orig, s)
#         return (s, e)

#     if not STRICT_MATCH_ONLY and len(nl) >= 12:
#         w = max(60, min(240, len(nl) + 80))
#         best_cov, best_span = 0.0, None
#         step = max(1, w // 2)
#         for start in range(0, max(1, len(tl) - w + 1), step):
#             cov, se = _fuzzy_window_span(tl, nl, start, w)
#             if cov > best_cov: best_cov, best_span = cov, se
#         if best_span and best_cov >= 0.65:
#             s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
#             if s > 0 and t_orig[s-1:s+1].lower() in {" the", " a", " an"}: s -= 1
#             s = _heal_split_word_left(t_orig, s)
#             return (s, e)

#     if not STRICT_MATCH_ONLY:
#         keys = [w for w in _tokenize(nl) if len(w) >= 4][:8]
#         if len(keys) >= 2:
#             kset = set(keys)
#             best_score, best_span = 0.0, None
#             for s, e, seg in _iter_sentences_with_spans(t_norm):
#                 toks = set(_tokenize(seg)); ov = len(kset & toks)
#                 if ov == 0: continue
#                 score = ov / max(2, len(kset)); length_pen = min(1.0, 120 / max(20, e - s)); score *= (0.6 + 0.4 * length_pen)
#                 if score > best_score: best_score, best_span = score, (s, min(e, s + 400))
#             if best_span and best_score >= 0.35:
#                 s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1]); s = _heal_split_word_left(t_orig, s)
#                 return (s, e)
#     return None

# def merge_overlaps(spans: List[Tuple[int,int,str,str]]) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans.sort(key=lambda x: x[0]); out = [spans[0]]
#     for s,e,c,aid in spans[1:]:
#         ps,pe,pc,paid = out[-1]
#         if s <= pe and pc == c and e > pe: out[-1] = (ps, e, pc, paid)
#         else: out.append((s,e,c,aid))
#     return out

# _PUNCT_WS = set(" \t\r\n,.;:!?)('\"[]{}-—–…") | _BRIDGE_CHARS

# def merge_overlaps_and_adjacent(base_text: str,
#                                 spans: List[Tuple[int,int,str,str]],
#                                 max_gap: int = 2) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans = sorted(spans, key=lambda x: x[0]); out = [spans[0]]
#     for s, e, c, aid in spans[1:]:
#         ps, pe, pc, paid = out[-1]
#         if c == pc and s <= pe: out[-1] = (ps, max(pe, e), pc, paid); continue
#         if c == pc and s - pe <= max_gap:
#             gap = base_text[max(0, pe):max(0, s)]
#             if all((ch in _PUNCT_WS) for ch in gap): out[-1] = (ps, e, pc, paid); continue
#         out.append((s, e, c, aid))
#     return out

# # ---------- Heading filters ----------
# def _is_heading_like(q: str) -> bool:
#     if not q: return True
#     s = q.strip()
#     if not re.search(r'[.!?]', s):
#         words = re.findall(r"[A-Za-z]+", s)
#         if 1 <= len(words) <= 7:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.8: return True
#         if s.lower() in {"introduction","voiceover","outro","epilogue","prologue","credits","title","hook","horrifying discovery","final decision"}: return True
#         if len(s) <= 3: return True
#     return False

# def _is_heading_context(script_text: str, s: int, e: int) -> bool:
#     left = script_text.rfind("\n", 0, s) + 1
#     right = script_text.find("\n", e); right = len(script_text) if right == -1 else right
#     line = script_text[left:right].strip()
#     if len(line) <= 70 and not re.search(r'[.!?]', line):
#         words = re.findall(r"[A-Za-z]+", line)
#         if 1 <= len(words) <= 8:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.7: return True
#     return False

# def _tighten_to_quote(script_text: str, span: Tuple[int,int], quote: str) -> Tuple[int,int]:
#     if not span or not quote: return span
#     s, e = span
#     if e <= s or s < 0 or e > len(script_text): return span
#     window = script_text[s:e]; win_norm = _normalize_keep_len(window).lower(); q_norm = _clean_quote_for_match(quote).lower()
#     if not q_norm: return span
#     i = win_norm.find(q_norm)
#     if i == -1:
#         m = re.search(re.escape(q_norm).replace(r"\ ", r"\s+"), win_norm, flags=re.IGNORECASE)
#         if not m: return span
#         i, j = m.start(), m.end()
#     else:
#         j = i + len(q_norm)
#     s2, e2 = s + i, s + j
#     s2, e2 = _snap_and_bridge_to_word(script_text, s2, e2); s2 = _heal_split_word_left(script_text, s2)
#     if s2 >= s and e2 <= e and e2 > s2: return (s2, e2)
#     return span

# def build_spans_by_param(script_text: str, data: dict, heading_ranges: Optional[List[Tuple[int,int]]] = None) -> Dict[str, List[Tuple[int,int,str,str]]]:
#     heading_ranges = heading_ranges or []
#     raw = (data or {}).get("per_parameter", {}) or {}
#     per: Dict[str, Dict[str, Any]] = {k:(v or {}) for k,v in raw.items()}
#     spans_map: Dict[str, List[Tuple[int,int,str,str]]] = {p: [] for p in PARAM_ORDER}
#     st.session_state["aoi_match_ranges"] = {}

#     for p in spans_map.keys():
#         color = PARAM_COLORS.get(p, "#ffd54f")
#         blk = per.get(p, {}) or {}
#         aois = blk.get("areas_of_improvement") or []
#         for idx, item in enumerate(aois, start=1):
#             raw_q = (item or {}).get("quote_verbatim", "") or ""
#             q = _sanitize_editor_text(raw_q)
#             clean = _clean_quote_for_match(re.sub(r"^[•\-\d\.\)\s]+", "", q).strip())
#             if not clean: continue
#             if _is_heading_like(clean): continue

#             pos = find_span_smart(script_text, clean)
#             if not pos: continue
#             pos = _tighten_to_quote(script_text, pos, raw_q)
#             s, e = pos

#             if heading_ranges and _overlaps_any(s, e, heading_ranges): continue
#             if _is_heading_context(script_text, s, e): continue

#             aid = f"{p.replace(' ','_')}-AOI-{idx}"
#             spans_map[p].append((s, e, color, aid))
#             st.session_state["aoi_match_ranges"][aid] = (s, e)
#     return spans_map

# # ---------- History (S3-aware) ----------
# def _maybe_copy_docx_to_history(source_docx_path: Optional[str], run_id: str) -> Optional[str]:
#     """Copy the DOCX used for rendering into outputs/_history so Recents can re-render identically."""
#     try:
#         if not source_docx_path:
#             return None
#         if source_docx_path.startswith("s3://") or (_s3_enabled() and not os.path.exists(source_docx_path)):
#             if os.path.exists(source_docx_path):
#                 with open(source_docx_path, "rb") as f:
#                     save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#                 return f"{HISTORY_DIR}/{run_id}.docx"
#             return source_docx_path

#         if _s3_enabled():
#             with open(source_docx_path, "rb") as f:
#                 save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#             return f"{HISTORY_DIR}/{run_id}.docx"
#         else:
#             dst = os.path.join(HISTORY_DIR, f"{run_id}.docx")
#             if os.path.abspath(source_docx_path) != os.path.abspath(dst):
#                 shutil.copyfile(source_docx_path, dst)
#             return dst
#     except Exception:
#         return None

# def _save_history_snapshot(title: str, data: dict, script_text: str,
#                            source_docx_path: Optional[str], heading_ranges: List[Tuple[int,int]],
#                            spans_by_param: Dict[str, List[Tuple[int,int,str,str]]],
#                            aoi_match_ranges: Dict[str, Tuple[int,int]]):
#     run_id = str(uuid.uuid4()); now = datetime.datetime.now()
#     created_at_iso = now.replace(microsecond=0).isoformat()
#     created_at_human = now.strftime("%Y-%m-%d %H:%M:%S")

#     stable_docx_key_or_path = _maybe_copy_docx_to_history(source_docx_path, run_id)

#     blob = {
#         "run_id": run_id, "title": title or "untitled",
#         "created_at": created_at_iso, "created_at_human": created_at_human,
#         "overall_rating": (data or {}).get("overall_rating", ""),
#         "scores": (data or {}).get("scores", {}),
#         "data": data or {}, "script_text": script_text or "",
#         "source_docx_path": stable_docx_key_or_path or source_docx_path,
#         "heading_ranges": heading_ranges or [],
#         "spans_by_param": spans_by_param or {},
#         "aoi_match_ranges": aoi_match_ranges or {},
#     }

#     out_name = f"{created_at_iso.replace(':','-')}__{run_id}.json"
#     out_key = f"{HISTORY_DIR}/{out_name}"
#     save_text_key(out_key, json.dumps(blob, ensure_ascii=False, indent=2))

#     # Manifest entry so Recents can work without ListBucket
#     _history_manifest_append({
#         "run_id": run_id,
#         "key": out_key,
#         "title": blob["title"],
#         "created_at": created_at_iso,
#         "created_at_human": created_at_human,
#         "overall_rating": blob.get("overall_rating", ""),
#     })

# def _load_all_history() -> List[dict]:
#     out: List[dict] = []
#     if _s3_enabled():
#         keys = sorted(list_prefix(HISTORY_DIR), reverse=True)

#         # Fallback to manifest when ListBucket is denied or nothing is found
#         if not keys:
#             man = _history_manifest_read()
#             if not man:
#                 return []
#             man = sorted(man, key=lambda m: m.get("created_at",""), reverse=True)
#             for m in man:
#                 k = m.get("key")
#                 if not k:
#                     continue
#                 try:
#                     txt = read_text_key(k, "")
#                     if not txt:
#                         continue
#                     j = json.loads(txt)
#                     j["_key"] = k
#                     if not j.get("created_at_human") and j.get("created_at"):
#                         try:
#                             dt = datetime.datetime.fromisoformat(j["created_at"])
#                             j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#                         except Exception:
#                             pass
#                     out.append(j)
#                 except Exception:
#                     continue
#             return out

#         # Normal path (ListBucket works)
#         for key in keys:
#             try:
#                 txt = read_text_key(key, "")
#                 if not txt:
#                     continue
#                 j = json.loads(txt)
#                 j["_key"] = key
#                 if not j.get("created_at_human") and j.get("created_at"):
#                     try:
#                         dt = datetime.datetime.fromisoformat(j["created_at"])
#                         j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#                     except Exception:
#                         pass
#                 out.append(j)
#             except Exception:
#                 continue
#     else:
#         for fp in sorted(glob.glob(os.path.join(HISTORY_DIR, "*.json"))):
#             try:
#                 with open(fp, "r", encoding="utf-8") as f:
#                     j = json.load(f)
#             except Exception:
#                 continue
#             j.setdefault("_path", fp)
#             ca = j.get("created_at")
#             try:
#                 if isinstance(ca, (int, float)):
#                     dt = datetime.datetime.utcfromtimestamp(float(ca))
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.astimezone().strftime("%Y-%m-%d %H:%M:%S")
#                 elif isinstance(ca, str) and ca: pass
#                 else:
#                     mtime = os.path.getmtime(fp); dt = datetime.datetime.fromtimestamp(mtime)
#                     j["created_at"] = dt.replace(microsecond=0).isoformat() + "Z"
#                     if not j.get("created_at_human"): j["created_at_human"] = dt.strftime("%Y-%m-%d %H:%M:%S")
#             except Exception:
#                 j["created_at"] = str(ca or "")
#             out.append(j)
#     out.sort(key=lambda r: r.get("created_at") or "", reverse=True)
#     return out

# def _open_history_run_by_id(run_id: str) -> bool:
#     """Open a history run by its run_id. Returns True if loaded."""
#     if not run_id: return False
#     recs = _load_all_history()
#     match = next((r for r in recs if r.get("run_id") == run_id), None)
#     if not match: return False

#     try:
#         if "_key" in match and _s3_enabled():
#             txt = read_text_key(match["_key"], "")
#             if not txt: return False
#             jj = json.loads(txt)
#         else:
#             path = match.get("_path")
#             if not path or not os.path.exists(path): return False
#             with open(path, "r", encoding="utf-8") as f:
#                 jj = json.load(f)
#     except Exception:
#         return False

#     st.session_state.script_text      = jj.get("script_text","")
#     st.session_state.base_stem        = jj.get("title","untitled")
#     st.session_state.data             = jj.get("data",{})
#     st.session_state.heading_ranges   = jj.get("heading_ranges",[])
#     st.session_state.spans_by_param   = jj.get("spans_by_param",{})
#     st.session_state.param_choice     = None
#     st.session_state.source_docx_path = jj.get("source_docx_path")
#     st.session_state.review_ready     = True
#     st.session_state["aoi_match_ranges"] = jj.get("aoi_match_ranges", {})
#     st.session_state.ui_mode          = "review"
#     return True

# def _render_recents_centerpane():
#     st.subheader("📄 Recents")
#     q = st.text_input("Filter by title…", "")
#     cols = st.columns([1, 4])
#     with cols[0]:
#         if st.button("← Back"):
#             st.session_state.ui_mode = "home"; st.rerun()

#     recs = _load_all_history()
#     ql = q.strip().lower()
#     if ql: recs = [r for r in recs if ql in (r.get("title","").lower())]
#     if not recs: st.caption("No history yet."); st.stop()

#     for rec in recs:
#         run_id = rec.get("run_id"); title = rec.get("title") or "(untitled)"
#         created_h = rec.get("created_at_human",""); overall = rec.get("overall_rating","")
#         st.markdown(f"""
#         <a class="rec-card" href="?open={run_id}">
#           <div class="rec-title">{title}</div>
#           <div class="rec-meta">{created_h}</div>
#           <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
#         </a>
#         """, unsafe_allow_html=True)

# # ---------- Sidebar ----------
# with st.sidebar:
#     if st.button("🆕 New review", use_container_width=True):
#         fp = st.session_state.get("flattened_docx_path")
#         if fp and os.path.exists(fp):
#             try: os.remove(fp)
#             except Exception: pass
#         for k in ["review_ready","script_text","base_stem","data","spans_by_param","param_choice",
#                   "source_docx_path","heading_ranges","flattened_docx_path","flatten_used"]:
#             st.session_state[k] = (
#                 False if k=="review_ready"
#                 else "" if k in ("script_text","base_stem")
#                 else {} if k=="spans_by_param"
#                 else [] if k=="heading_ranges"
#                 else None if k in ("source_docx_path","flattened_docx_path")
#                 else False if k=="flatten_used"
#                 else None
#             )
#         st.session_state.ui_mode = "home"
#         _clear_query_params()
#         st.rerun()

#     if st.button("📁 Recents", use_container_width=True):
#         st.session_state.ui_mode = "recents"
#         _clear_query_params()
#         st.rerun()

# # ---------- Input screen ----------
# def render_home():
#     st.subheader("🎬 Script Source")

#     tab_upload, tab_paste = st.tabs(["Upload file", "Paste text"])

#     uploaded_file = None
#     uploaded_name = None
#     uploaded_key  = None

#     def _safe_stem(s: str, fallback: str = "pasted_script") -> str:
#         s = (s or "").strip()
#         if not s:
#             return fallback
#         s = re.sub(r"[^A-Za-z0-9._\-]+", "_", s)
#         s = s.strip("._-") or fallback
#         return s

#     # --- Upload tab ---
#     with tab_upload:
#         up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
#         if up is not None:
#             file_bytes = up.read()
#             suffix = os.path.splitext(up.name)[1].lower()
#             uploaded_key = f"{SCRIPTS_DIR}/{up.name}"
#             save_bytes_key(uploaded_key, file_bytes)

#             with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
#                 tmp.write(file_bytes)
#                 uploaded_file = tmp.name
#             uploaded_name = os.path.splitext(os.path.basename(up.name))[0] or "uploaded_script"

#     # --- Paste tab ---
#     with tab_paste:
#         paste_title = st.text_input("Title (optional)", placeholder="e.g., my_script")
#         pasted_text = st.text_area(
#             "Paste your script text here",
#             height=360,
#             placeholder="Paste the full script text (we’ll analyze this as-is)."
#         )

#     # Run Review works for BOTH
#     if st.button("🚀 Run Review", type="primary", use_container_width=True):
#         base_stem = "uploaded_script"
#         source_docx_path = None
#         heading_ranges: List[Tuple[int,int]] = []
#         script_text = ""

#         if pasted_text and pasted_text.strip():
#             base_stem = _safe_stem(paste_title, "pasted_script")
#             script_text = pasted_text

#             pasted_key = f"{SCRIPTS_DIR}/{base_stem}.txt"
#             save_text_key(pasted_key, script_text)
#             source_docx_path = pasted_key
#             heading_ranges = []

#         elif uploaded_file:
#             base_stem = uploaded_name or "uploaded_script"

#             if uploaded_file.lower().endswith(".docx"):
#                 try:
#                     left_text, used_left = extract_left_column_script_or_default(uploaded_file)
#                 except Exception:
#                     left_text, used_left = "", False

#                 if used_left and left_text.strip():
#                     script_text = left_text
#                     source_docx_path = uploaded_file
#                     heading_ranges = []
#                 else:
#                     path_to_use = uploaded_file
#                     if _docx_contains_tables(path_to_use):
#                         flat = flatten_docx_tables_to_longtext(path_to_use)
#                         st.session_state.flattened_docx_path = flat
#                         st.session_state.flatten_used = True
#                         path_to_use = flat
#                     script_text, heading_ranges = build_docx_text_with_meta(path_to_use)
#                     source_docx_path = path_to_use
#             else:
#                 script_text = load_script_file(uploaded_file)
#                 source_docx_path = uploaded_key or uploaded_file

#         else:
#             st.warning("Please upload a script **or** paste text in the second tab.")
#             st.stop()

#         if len(script_text.strip()) < 50:
#             st.error("Extracted text looks too short. Please check your input.")
#             st.stop()

#         with st.spinner("Running analysis…"):
#             try:
#                 review_text = run_review_multi(
#                     script_text=script_text,
#                     prompts_dir=PROMPTS_DIR,
#                     temperature=0.0
#                 )
#             finally:
#                 if uploaded_file and not (isinstance(source_docx_path, str) and os.path.exists(source_docx_path)):
#                     try: os.remove(uploaded_file)
#                     except Exception: pass

#         data = extract_review_json(review_text)
#         if not data:
#             st.error("JSON not detected in model output.")
#             st.stop()

#         st.session_state.script_text      = script_text
#         st.session_state.base_stem        = base_stem
#         st.session_state.data             = data
#         st.session_state.heading_ranges   = heading_ranges
#         st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
#         st.session_state.param_choice     = None
#         st.session_state.source_docx_path = source_docx_path
#         st.session_state.review_ready     = True
#         st.session_state.ui_mode          = "review"

#         _save_history_snapshot(
#             title=base_stem,
#             data=data,
#             script_text=script_text,
#             source_docx_path=source_docx_path,
#             heading_ranges=heading_ranges,
#             spans_by_param=st.session_state.spans_by_param,
#             aoi_match_ranges=st.session_state.get("aoi_match_ranges", {})
#         )

#         _clear_query_params()
#         st.rerun()

# # ---------- Results screen ----------
# def render_review():
#     script_text     = st.session_state.script_text
#     data            = st.session_state.data
#     spans_by_param  = st.session_state.spans_by_param
#     scores: Dict[str,int] = (data or {}).get("scores", {}) or {}
#     source_docx_path: Optional[str] = st.session_state.source_docx_path

#     # If our source_docx_path is an S3 key/url, ensure we have a local copy for rendering
#     docx_local: Optional[str] = None
#     preferred = st.session_state.get("flattened_docx_path") if st.session_state.get("flatten_used") else source_docx_path
#     if preferred:
#         if not os.path.exists(preferred):
#             docx_local = ensure_local_copy(preferred)
#         else:
#             docx_local = preferred

#     left, center, right = st.columns([1.1, 2.7, 1.4], gap="large")

#     with left:
#         st.subheader("Final score")
#         ordered = [p for p in PARAM_ORDER if p in scores]
#         df = pd.DataFrame({"Parameter": ordered, "Score (1–10)": [scores.get(p, "") for p in ordered]})
#         st.dataframe(df, hide_index=True, use_container_width=True)
#         st.markdown(f'**Overall:** {data.get("overall_rating","—")}/10')
#         st.divider()

#         strengths = (data or {}).get("strengths") or []
#         if not strengths:
#             per = (data or {}).get("per_parameter", {}) or {}
#             best = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
#             for name, sc in best:
#                 if sc >= 8 and name in per:
#                     exp = _sanitize_editor_text((per[name] or {}).get("explanation", "") or "")
#                     first = re.split(r"(?<=[.!?])\s+", exp.strip())[0] if exp else f"Consistently strong {name.lower()}."
#                     strengths.append(f"{name}: {first}")
#                 if len(strengths) >= 3: break

#         def _bullets(title: str, items):
#             st.markdown(f"**{title}**")
#             for s in (items or []):
#                 if isinstance(s, str) and s.strip(): st.write("• " + _sanitize_editor_text(s))
#             if not items: st.write("• —")

#         _bullets("Strengths", strengths)
#         _bullets("Weaknesses", data.get("weaknesses"))
#         _bullets("Suggestions", data.get("suggestions"))
#         _bullets("Drop-off Risks", data.get("drop_off_risks"))
#         st.markdown("**Viral Quotient**"); st.write(_sanitize_editor_text(data.get("viral_quotient","—")))

#     with right:
#         st.subheader("Parameters")
#         st.markdown('<div class="param-row">', unsafe_allow_html=True)
#         for p in [p for p in PARAM_ORDER if p in scores]:
#             if st.button(p, key=f"chip_{p}", help="Show inline AOI highlights for this parameter"):
#                 st.session_state.param_choice = p
#         st.markdown('</div>', unsafe_allow_html=True)

#         sel = st.session_state.param_choice
#         if sel:
#             blk = (data.get("per_parameter", {}) or {}).get(sel, {}) or {}
#             st.markdown(f"**{sel} — Score:** {scores.get(sel,'—')}/10")

#             if blk.get("explanation"):
#                 st.markdown("**Why this score**"); st.write(_sanitize_editor_text(blk["explanation"]))
#             if blk.get("weakness") and blk["weakness"] != "Not present":
#                 st.markdown("**Weakness**"); st.write(_sanitize_editor_text(blk["weakness"]))
#             if blk.get("suggestion") and blk["suggestion"] != "Not present":
#                 st.markdown("**Suggestion**"); st.write(_sanitize_editor_text(blk["suggestion"]))

#             aoi = blk.get("areas_of_improvement") or []
#             if aoi:
#                 st.markdown("**Areas of Improvement**")
#                 for i, item in enumerate(aoi, 1):
#                     popover_fn = getattr(st, "popover", None)
#                     aid = f"{sel.replace(' ','_')}-AOI-{i}"
#                     s_e_map = st.session_state.get("aoi_match_ranges", {})
#                     if aid in s_e_map:
#                         s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                         line = (matched_line if len(matched_line) <= 320 else matched_line[:300] + "…")
#                     else:
#                         line = _sanitize_editor_text(item.get('quote_verbatim',''))
#                     issue = _sanitize_editor_text(item.get('issue',''))
#                     fix   = _sanitize_editor_text(item.get('fix',''))
#                     why   = _sanitize_editor_text(item.get('why','') or item.get('why_this_helps',''))
#                     label = f"Issue {i}"
#                     if callable(popover_fn):
#                         with popover_fn(label):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)
#                     else:
#                         with st.expander(label, expanded=False):
#                             if line: st.markdown(f"**Line:** {line}")
#                             if issue: st.markdown(f"**Issue:** {issue}")
#                             if fix: st.markdown(f"**Fix:** {fix}")
#                             if why: st.caption(why)

#             if blk.get("summary"):
#                 st.markdown("**Summary**"); st.write(_sanitize_editor_text(blk["summary"]))

#     with center:
#         st.subheader("Script with inline highlights")
#         spans = st.session_state.spans_by_param.get(st.session_state.param_choice, []) if st.session_state.param_choice else []

#         aoi_payload: Dict[str, Dict[str, str]] = {}
#         data_per = (data or {}).get("per_parameter") or {}
#         s_e_map = st.session_state.get("aoi_match_ranges", {})
#         sel = st.session_state.param_choice

#         def _mk_line(aid: str, fallback_q: str = "") -> str:
#             if aid in s_e_map:
#                 s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                 return matched_line if len(matched_line) <= 320 else matched_line[:300] + "…"
#             return _sanitize_editor_text(fallback_q or "")

#         def _collect(param_name: str):
#             blk = (data_per.get(param_name) or {})
#             for i, item in enumerate(blk.get("areas_of_improvement") or [], 1):
#                 aid = f"{param_name.replace(' ','_')}-AOI-{i}"
#                 aoi_payload[aid] = {
#                     "line": _mk_line(aid, (item or {}).get("quote_verbatim","")),
#                     "issue": _sanitize_editor_text((item or {}).get("issue","")),
#                     "fix": _sanitize_editor_text((item or {}).get("fix","")),
#                     "why": _sanitize_editor_text((item or {}).get("why","") or (item or {}).get("why_this_helps","")),
#                 }

#         if sel: _collect(sel)
#         else:
#             for pn in [p for p in PARAM_ORDER if p in data_per]:
#                 _collect(pn)

#         payload_json = json.dumps(aoi_payload, ensure_ascii=False)

#         frame_theme_css = """
#         <style>
#           :root{
#             --m7-surface: #eef2f7;
#             --m7-on-surface: #0f172a;
#             --m7-border: rgba(15,23,42,.14);
#           }
#           @media (prefers-color-scheme: dark){
#             :root{
#               --m7-surface: #2f333a;
#               --m7-on-surface: #ffffff;
#               --m7-border: rgba(255,255,255,.18);
#             }
#             body { background: transparent !important; }
#           }
#           .docxwrap{ background: var(--m7-surface); color: var(--m7-on-surface); border: 1px solid var(--m7-border); border-radius: 12px; padding: 16px 14px 18px; }
#           .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#           .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); }
#         </style>
#         """

#         tooltip_css = """
#         <style>
#         .aoi-pop {
#           position: absolute; max-width: 520px; min-width: 320px;
#           background: var(--m7-surface); border: 1px solid var(--m7-border); border-radius: 10px;
#           box-shadow: 0 10px 25px rgba(0,0,0,.12);
#           padding: 12px 14px; z-index: 9999; transform: translateY(-8px);
#           color: var(--m7-on-surface);
#         }
#         .aoi-pop h4 { margin: 0 0 .35rem 0; font-size: .95rem; }
#         .aoi-pop p  { margin: .15rem 0; line-height: 1.5; }
#         .aoi-pop .muted { opacity:.85; font-size:.85rem; }
#         .aoi-arrow {
#           position:absolute; left:50%; transform:translateX(-50%);
#           bottom:-7px; width:0;height:0;border-left:7px solid transparent;
#           border-right:7px solid transparent;border-top:7px solid var(--m7-border);
#         }
#         .aoi-arrow::after{
#           content:""; position:absolute; left:-6px; top:-7px; width:0;height:0;
#           border-left:6px solid transparent;border-right:6px solid transparent;border-top:6px solid var(--m7-surface);
#         }
#         </style>
#         """

#         if docx_local and os.path.splitext(docx_local)[1].lower() == ".docx":
#             html_core = render_docx_html_with_highlights(
#                 docx_local,
#                 merge_overlaps_and_adjacent(script_text, spans)
#             )
#         else:
#             from html import escape as _esc
#             orig = script_text
#             spans2 = [s for s in merge_overlaps_and_adjacent(orig, spans) if s[0] < s[1]]
#             spans2.sort(key=lambda x: x[0])
#             cur = 0; buf: List[str] = []
#             for s,e,c,aid in spans2:
#                 if s > cur: buf.append(_esc(orig[cur:s]))
#                 buf.append(
#                     f'<mark class="aoi-mark" data-aid="{aid}" '
#                     f'style="background:{c}33;border:1px solid {c};border-radius:3px;padding:0 2px;">'
#                     f'{_esc(orig[s:e])}</mark>'
#                 )
#                 cur = e
#             if cur < len(orig): buf.append(_esc(orig[cur:]))
#             html_core = (
#                 '<div class="docxwrap"><p style="white-space:pre-wrap; '
#                 'line-height:1.7; font-family:ui-serif, Georgia, Times New Roman, serif;">'
#                 + "".join(buf) +
#                 '</p></div>'
#             )

#         html_shell = """
# %%FRAME_THEME_CSS%%
# %%TOOLTIP_CSS%%
# <div id="m7-doc">%%HTML_CORE%%</div>
# <div id="aoi-pop" class="aoi-pop" style="display:none;">
#   <div id="aoi-pop-content"></div>
#   <div class="aoi-arrow"></div>
# </div>
# <script>
# (function(){
#   const AOI = __PAYLOAD__;
#   const wrap = document.getElementById('m7-doc');
#   const pop  = document.getElementById('aoi-pop');
#   const body = document.getElementById('aoi-pop-content');

#   function resizeIframe() {
#     try {
#       const h = Math.max(
#         document.documentElement.scrollHeight,
#         document.body.scrollHeight
#       );
#       if (window.frameElement) {
#         window.frameElement.style.height = (h + 20) + 'px';
#         window.frameElement.style.width  = '100%';
#       }
#     } catch(e) {}
#   }
#   window.addEventListener('load', resizeIframe);
#   window.addEventListener('resize', resizeIframe);

#   function hide(){ pop.style.display='none'; }
#   function showFor(mark){
#     const aid = mark.getAttribute('data-aid');
#     const d = AOI[aid]; if(!d) return;
#     body.innerHTML =
#       (d.line  ? '<p><strong>Line:</strong> '  + d.line  + '</p>' : '') +
#       (d.issue ? '<p><strong>Issue:</strong> ' + d.issue + '</p>' : '') +
#       (d.fix   ? '<p><strong>Fix:</strong> '   + d.fix   + '</p>' : '') +
#       (d.why   ? '<p class="muted">'           + d.why   + '</p>' : '');
#     pop.style.display = 'block';

#     const r = mark.getBoundingClientRect();
#     const scY = window.scrollY || document.documentElement.scrollTop;
#     const scX = window.scrollX || document.documentElement.scrollLeft;
#     let top  = r.top + scY - pop.offsetHeight - 10;
#     let left = r.left + scX + r.width/2 - pop.offsetWidth/2;
#     if (top < 8) top = r.bottom + scY + 10;
#     if (left < 8) left = 8;
#     pop.style.top  = top + 'px';
#     pop.style.left = left + 'px';

#     resizeIframe();
#   }

#   wrap.addEventListener('click', (e) => {
#     const m = e.target.closest('.aoi-mark');
#     if(!m){ hide(); return; }
#     if(pop.style.display === 'block'){ hide(); }
#     showFor(m);
#     e.stopPropagation();
#   });

#   document.addEventListener('click', (e) => {
#     if(!e.target.closest('.aoi-pop') && !e.target.closest('.aoi-mark')) hide();
#   });
# })();
# </script>
# """
#         html_shell = (
#             html_shell
#             .replace("%%FRAME_THEME_CSS%%", frame_theme_css)
#             .replace("%%TOOLTIP_CSS%%", tooltip_css)
#             .replace("%%HTML_CORE%%", html_core)
#             .replace("__PAYLOAD__", payload_json)
#         )

#         components.html(html_shell, height=400, scrolling=False)

# # ---------- Router & query param open ----------
# _open_qp = _get_query_param("open")
# if _open_qp and _open_history_run_by_id(_open_qp):
#     _clear_query_params()  # avoid re-opening on subsequent reruns

# mode = st.session_state.ui_mode
# if mode == "recents":
#     _render_recents_centerpane()
# elif mode == "review" and st.session_state.review_ready:
#     render_review()
# else:
#     render_home()





##############################################








###Final History 







# # app_grammarly_ui.py — Runpod S3-only + Stable Recents + In-place open

# import os, re, glob, json, tempfile, difflib, uuid, datetime, shutil, time
# from pathlib import Path
# from typing import Dict, Any, List, Tuple, Optional

# import streamlit as st
# import pandas as pd
# import streamlit.components.v1 as components  # for inline HTML/JS popup

# # ---- utils & engine ----
# from utils1 import (
#     extract_review_json,
#     PARAM_ORDER,
#     load_script_file,
#     extract_left_column_script_or_default,  # <-- left-column extractor for DOCX tables
# )
# from review_engine_multi import run_review_multi

# # ---- DOCX rendering imports ----
# from docx import Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # =========================
# # RunPod S3 (S3-only helpers)
# # =========================
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError

# def _get_env(key: str, default: str = "") -> str:
#     v = os.getenv(key, "")
#     if v:
#         return v.strip()
#     try:
#         v2 = st.secrets.get(key)
#         if isinstance(v2, str):
#             return v2.strip()
#     except Exception:
#         pass
#     return (default or "").strip()

# # Primary config
# _RP_ENDPOINT = _get_env("RUNPOD_S3_ENDPOINT")
# _RP_BUCKET   = _get_env("RUNPOD_S3_BUCKET")
# _RP_REGION   = _get_env("RUNPOD_S3_REGION") or _get_env("AWS_DEFAULT_REGION") or ""

# # Credentials: prefer AWS_* if present; else RUNPOD_* fallbacks
# _AK = _get_env("AWS_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY")
# _SK = _get_env("AWS_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_KEY")
# _ST = _get_env("AWS_SESSION_TOKEN")  # optional

# # Options
# _FORCE_PATH = (_get_env("RUNPOD_S3_FORCE_PATH_STYLE") or "true").lower() in {"1","true","yes"}
# _USE_SSL    = (_get_env("RUNPOD_S3_USE_SSL") or "true").lower() in {"1","true","yes"}
# _VERIFY_SSL = (_get_env("RUNPOD_S3_VERIFY_SSL") or "true").lower() in {"1","true","yes"}

# def _s3_enabled() -> bool:
#     return bool(_RP_ENDPOINT and _RP_BUCKET and _AK and _SK)

# @st.cache_resource(show_spinner=False)
# def _s3_client():
#     if not _s3_enabled():
#         return None
#     session_kwargs = dict(
#         aws_access_key_id=_AK,
#         aws_secret_access_key=_SK,
#     )
#     if _ST:
#         session_kwargs["aws_session_token"] = _ST

#     cfg = Config(
#         signature_version="s3v4",
#         s3={"addressing_style": "path" if _FORCE_PATH else "auto"},
#         retries={"max_attempts": 3, "mode": "standard"}
#     )
#     return boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         region_name=_RP_REGION or None,
#         use_ssl=_USE_SSL,
#         verify=_VERIFY_SSL,
#         config=cfg,
#         **session_kwargs,
#     )

# # ---------- S3 I/O (S3-only) ----------
# def save_text_key(key: str, text: str) -> str:
#     key = key.lstrip("/")
#     if not _s3_enabled():
#         # hard fail: S3-only mode
#         raise RuntimeError("S3 is not configured (RUNPOD_* / AWS_* envs).")
#     kwargs = {
#         "Bucket": _RP_BUCKET,
#         "Key": key,
#         "Body": text.encode("utf-8"),
#     }
#     if key.endswith(".json"):
#         kwargs["ContentType"] = "application/json"
#         kwargs["CacheControl"] = "no-store"
#     _s3_client().put_object(**kwargs)
#     return f"s3://{_RP_BUCKET}/{key}"

# def save_bytes_key(key: str, data: bytes) -> str:
#     key = key.lstrip("/")
#     if not _s3_enabled():
#         raise RuntimeError("S3 is not configured (RUNPOD_* / AWS_* envs).")
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=data)
#     return f"s3://{_RP_BUCKET}/{key}"

# def read_text_key(key: str, default: str = "") -> str:
#     if not _s3_enabled():
#         return default
#     tries = 4
#     for k in range(tries):
#         try:
#             resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#             body = resp["Body"].read().decode("utf-8", errors="ignore")
#             if body.strip() == "" and k < tries - 1:
#                 time.sleep(0.25 * (k + 1))
#                 continue
#             return body
#         except Exception:
#             if k < tries - 1:
#                 time.sleep(0.25 * (k + 1))
#                 continue
#             return default
#     return default

# def read_bytes_key(key: str) -> Optional[bytes]:
#     if not _s3_enabled():
#         return None
#     tries = 4
#     for k in range(tries):
#         try:
#             resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#             return resp["Body"].read()
#         except Exception:
#             if k < tries - 1:
#                 time.sleep(0.25 * (k + 1))
#                 continue
#             return None
#     return None

# def list_prefix(prefix: str) -> List[str]:
#     if not _s3_enabled():
#         return []
#     out: List[str] = []
#     token = None
#     s3_prefix = prefix.rstrip("/") + "/"
#     try:
#         while True:
#             kwargs = {"Bucket": _RP_BUCKET, "Prefix": s3_prefix}
#             if token:
#                 kwargs["ContinuationToken"] = token
#             resp = _s3_client().list_objects_v2(**kwargs)
#             for c in resp.get("Contents", []):
#                 k = c.get("Key", "")
#                 if k.endswith(".json"):
#                     out.append(k)
#             token = resp.get("NextContinuationToken")
#             if not token:
#                 break
#     except (ClientError, EndpointConnectionError, NoCredentialsError):
#         return []
#     return out

# def presigned_url(key: str, expires: int = 3600) -> Optional[str]:
#     if not _s3_enabled():
#         return None
#     try:
#         return _s3_client().generate_presigned_url(
#             "get_object",
#             Params={"Bucket": _RP_BUCKET, "Key": key},
#             ExpiresIn=expires
#         )
#     except ClientError:
#         return None

# def ensure_local_copy(key_or_keyurl: str) -> Optional[str]:
#     """
#     Always download to a temp file for parsing (DOCX/PDF).
#     """
#     if not _s3_enabled():
#         return None
#     key = key_or_keyurl
#     if key.startswith("s3://"):
#         parts = key.split("/", 3)
#         key = parts[3] if len(parts) >= 4 else ""
#     data = read_bytes_key(key)
#     if data is None:
#         return None
#     fd, tmp = tempfile.mkstemp(suffix=os.path.splitext(key)[1] or "")
#     os.close(fd)
#     with open(tmp, "wb") as f:
#         f.write(data)
#     return tmp

# def _s3_health_summary() -> dict:
#     info = {
#         "enabled": _s3_enabled(),
#         "endpoint": _RP_ENDPOINT,
#         "bucket": _RP_BUCKET,
#         "region": _RP_REGION,
#         "has_keys": bool(_AK and _SK),
#     }
#     if not _s3_enabled():
#         info["status"] = "disabled"
#         return info
#     try:
#         _ = _s3_client().list_objects_v2(Bucket=_RP_BUCKET, Prefix="Scriptmodel/outputs/_history/", MaxKeys=1)
#         info["status"] = "ok"
#     except Exception as e:
#         info["status"] = f"error: {getattr(e, 'response', {}).get('Error', {}).get('Code', str(e))}"
#     return info

# # ---------- Folders (all under Scriptmodel/) ----------
# BASE_PREFIX = "Scriptmodel"
# SCRIPTS_DIR = f"{BASE_PREFIX}/scripts"
# PROMPTS_DIR = f"{BASE_PREFIX}/prompts"
# OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
# HISTORY_DIR = f"{OUTPUT_DIR}/_history"

# # ---------- Colors ----------
# PARAM_COLORS: Dict[str, str] = {
#     "Suspense Building":              "#ff6b6b",
#     "Language/Tone":                  "#6b8cff",
#     "Intro + Main Hook/Cliffhanger":  "#ffb86b",
#     "Story Structure + Flow":         "#a78bfa",
#     "Pacing":                         "#f43f5e",
#     "Mini-Hooks (30–60s)":            "#eab308",
#     "Outro (Ending)":                 "#8b5cf6",
#     "Grammar & Spelling":             "#10b981",
# }

# STRICT_MATCH_ONLY = False

# # ---------- App config ----------
# st.set_page_config(page_title="M7 — Grammarly UI", page_icon="🕵️", layout="wide")

# # ---------- Header patch & CSS ----------
# def render_app_title():
#     st.markdown('<h1 class="app-title">Viral Script Reviewer</h1>', unsafe_allow_html=True)
#     st.markdown("""
#     <style>
#     html { color-scheme: light dark; }
#     :root{ --m7-surface:#eef2f7; --m7-on-surface:#0f172a; --m7-border:rgba(15,23,42,.14); --sep:#e5e7eb; }
#     @media (prefers-color-scheme: dark){
#       :root{ --m7-surface:#2f333a; --m7-on-surface:#ffffff; --m7-border:rgba(255,255,255,.18); --sep:#2a2f37; }
#     }
#     .stApp .block-container { padding-top: 4.25rem !important; }
#     .app-title{ font-weight:700; font-size:2.1rem; line-height:1.3; margin:0 0 1rem 0; padding-left:40px!important; padding-top:.25rem!important; }
#     [data-testid="collapsedControl"] { z-index: 6 !important; }
#     header[data-testid="stHeader"], .stAppHeader { background: transparent !important; box-shadow:none!important; }
#     @media (min-width: 992px){ .app-title { padding-left: 0 !important; } }
#     div[data-testid="column"]:nth-of-type(1){position:relative;}
#     div[data-testid="column"]:nth-of-type(1)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     div[data-testid="column"]:nth-of-type(2){position:relative;}
#     div[data-testid="column"]:nth-of-type(2)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     .m7-card{ background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:14px 16px; color:var(--m7-on-surface); }
#     .m7-card, .m7-card * { color:var(--m7-on-surface)!important; }
#     .docxwrap{ background:var(--m7-surface); color:var(--m7-on-surface); border:1px solid var(--m7-border); border-radius:12px; padding:16px 14px 18px; }
#     .docxwrap .h1,.docxwrap .h2,.docxwrap .h3 { font-weight:700; margin:10px 0 6px; }
#     .docxwrap .h1{font-size:1.3rem; border-bottom:2px solid currentColor; padding-bottom:4px;}
#     .docxwrap .h2{font-size:1.15rem; border-bottom:1px solid currentColor; padding-bottom:3px;}
#     .docxwrap .h3{font-size:1.05rem;}
#     .docxwrap p{ margin:10px 0; line-height:1.7; font-family: ui-serif, Georgia, "Times New Roman", serif; }
#     .docxwrap table{ border-collapse:collapse; width:100%; margin:12px 0; }
#     .docxwrap th,.docxwrap td{ border:1px solid var(--m7-border); padding:8px; vertical-align:top; line-height:1.6; }
#     .docxwrap mark{ padding:0 2px; border-radius:3px; border:1px solid var(--m7-border); cursor:pointer; }
#     .rec-card{ display:block; text-decoration:none!important; background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:14px 16px; margin:10px 0 16px; box-shadow:0 1px 2px rgba(0,0,0,.06); color:var(--m7-on-surface)!important; transition: filter .1s ease, transform .02s ease; }
#     .rec-card:hover{ filter:brightness(1.02); }
#     .rec-card:active{ transform: translateY(1px); }
#     .rec-title{font-weight:600; margin-bottom:.25rem;}
#     .rec-meta{opacity:.85!important; font-size:12.5px; margin-bottom:.4rem;}
#     .rec-row{display:flex; align-items:center; justify-content:space-between; gap:12px;}
#     .stTextInput>div>div, .stTextArea>div>div, .stNumberInput>div>div, .stDateInput>div>div, .stTimeInput>div>div, .stFileUploader>div, div[data-baseweb="select"]{ background:var(--m7-surface)!important; border:1px solid var(--m7-border)!important; border-radius:10px!important; color:var(--m7-on-surface)!important; }
#     .stTextInput input,.stTextArea textarea,.stNumberInput input,.stDateInput input,.stTimeInput input,.stFileUploader div,div[data-baseweb="select"] *{ color:var(--m7-on-surface)!important; }
#     .stTextInput input::placeholder,.stTextArea textarea::placeholder{ color:rgba(16,24,39,.55)!important; }
#     @media (prefers-color-scheme: dark){ .stTextInput input::placeholder,.stTextArea textarea::placeholder{ color:rgba(255,255,255,.75)!important; } }
#     div[data-testid="stFileUploaderDropzone"] label span { color: var(--m7-on-surface) !important; opacity:1!important; }
#     div[data-testid="stDataFrame"]{ background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:6px 8px; color:var(--m7-on-surface); }
#     .stMarkdown pre, pre[class*="language-"], .stCodeBlock{ background:var(--m7-surface)!important; color:var(--m7-on-surface)!important; border:1px solid var(--m7-border)!important; border-radius:12px!important; padding:12px 14px!important; overflow:auto; }
#     .stMarkdown pre code{ background:transparent!important; color:inherit!important; }
#     </style>
#     """, unsafe_allow_html=True)

# render_app_title()

# # ---------- Session defaults ----------
# for key, default in [
#     ("review_ready", False),
#     ("script_text", ""),
#     ("base_stem", ""),
#     ("data", None),
#     ("spans_by_param", {}),
#     ("param_choice", None),
#     ("source_docx_path", None),
#     ("heading_ranges", []),
#     ("flattened_docx_path", None),
#     ("flatten_used", False),
#     ("ui_mode", "home"),
# ]:
#     st.session_state.setdefault(key, default)

# # Recents stability helpers
# st.session_state.setdefault("_last_history_cache", [])
# st.session_state.setdefault("_open_run_key", None)
# st.session_state.setdefault("_open_run_id", None)

# # ---------- helpers for query params ----------
# def _get_query_param(key: str) -> Optional[str]:
#     val = None
#     try:
#         val = st.query_params.get(key)
#     except Exception:
#         q = st.experimental_get_query_params()
#         v = q.get(key)
#         if isinstance(v, list): val = v[0] if v else None
#         else: val = v
#     return val

# def _clear_query_params():
#     try:
#         st.query_params.clear()
#     except Exception:
#         st.experimental_set_query_params()

# # ---------- Sanitizer ----------
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )
# def _sanitize_editor_text(s: Optional[str]) -> str:
#     if not s: return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•)\s*', '', t, flags=re.M)
#     t = re.sub(r'^\s*[-*]\s+', '• ', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# # ---------- DOCX traversal ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# # ---------- Auto-flatten ----------
# def _docx_contains_tables(path: str) -> bool:
#     doc = Document(path)
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Table):
#             return True
#     return False

# def _copy_paragraph(dest_doc: Document, src_para: Paragraph):
#     p = dest_doc.add_paragraph()
#     try:
#         if src_para.style and src_para.style.name:
#             p.style = src_para.style.name
#     except Exception:
#         pass
#     for run in src_para.runs:
#         r = p.add_run(run.text or "")
#         r.bold = run.bold
#         r.italic = run.italic
#         r.underline = run.underline
#     return p

# def flatten_docx_tables_to_longtext(source_path: str) -> str:
#     src = Document(source_path)
#     new = Document()
#     for blk in _iter_docx_blocks(src):
#         if isinstance(blk, Paragraph):
#             _copy_paragraph(new, blk)
#         else:
#             seen_tc_ids = set()
#             for row in blk.rows:
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     if tc_id in seen_tc_ids:
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     for p in cell.paragraphs:
#                         _copy_paragraph(new, p)
#                 new.add_paragraph("")
#             new.add_paragraph("")
#     fd, tmp_path = tempfile.mkstemp(suffix=".docx"); os.close(fd); new.save(tmp_path)
#     return tmp_path

# # ---------- Build plain text + heading ranges ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# def build_docx_text_with_meta(docx_path: str) -> Tuple[str, List[Tuple[int,int]]]:
#     doc = Document(docx_path)
#     out: List[str] = []
#     heading_ranges: List[Tuple[int,int]] = []
#     current_offset = 0

#     def _append_and_advance(s: str):
#         nonlocal current_offset
#         out.append(s); current_offset += len(s)

#     seen_tc_ids: set = set()
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             para_text = "".join(run.text or "" for run in blk.runs)
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             if sty.startswith("heading"):
#                 start = current_offset; end = start + len(para_text)
#                 heading_ranges.append((start, end))
#             _append_and_advance(para_text); _append_and_advance("\n")
#         else:
#             for row in blk.rows:
#                 row_cell_tcs = []
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     row_cell_tcs.append((tc_id, cell))
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     if tc_id in seen_tc_ids:
#                         if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     cell_text_parts: List[str] = []
#                     for i, p in enumerate(cell.paragraphs):
#                         t = "".join(r.text or "" for r in p.runs)
#                         sty = (p.style.name or "").lower() if p.style else ""
#                         if sty.startswith("heading"):
#                             hs = current_offset + sum(len(x) for x in cell_text_parts)
#                             he = hs + len(t)
#                             heading_ranges.append((hs, he))
#                         cell_text_parts.append(t)
#                         if i != len(cell.paragraphs) - 1:
#                             cell_text_parts.append("\n")
#                     cell_text = "".join(cell_text_parts)
#                     _append_and_advance(cell_text)
#                     if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                 _append_and_advance("\n")
#             _append_and_advance("\n")

#     return "".join(out), heading_ranges

# # ---------- Matching / spans (unchanged core logic) ----------
# _BRIDGE_CHARS = set("\u200b\u200c\u200d\u2060\ufeff\xa0\u00ad")
# def _normalize_keep_len(s: str) -> str:
#     trans = {
#         "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
#         "\u2013": "-", "\u2014": "-",
#         "\xa0": " ",
#         "\u200b": " ", "\u200c": " ", "\u200d": " ", "\u2060": " ",
#         "\ufeff": " ", "\u00ad": " ",
#     }
#     return (s or "").translate(str.maketrans(trans))
# def _tokenize(s: str) -> List[str]: return re.findall(r"\w+", (s or "").lower())
# def _iter_sentences_with_spans(text: str) -> List[Tuple[int,int,str]]:
#     spans = []
#     for m in re.finditer(r'[^.!?]+[.!?]+|\Z', text, flags=re.S):
#         s, e = m.start(), m.end()
#         seg = text[s:e]
#         if seg.strip(): spans.append((s, e, seg))
#     return spans
# def _squash_ws(s: str) -> str: return re.sub(r"\s+", " ", s or "").strip()
# def _clean_quote_for_match(q: str) -> str:
#     if not q: return ""
#     q = _normalize_keep_len(q).strip()
#     q = re.sub(r'^[\'"“”‘’\[\(\{<…\-\–\—\s]+', '', q)
#     q = re.sub(r'[\'"“”‘’\]\)\}>…\-\–\—\s]+$', '', q)
#     return _squash_ws(q)
# def _snap_and_bridge_to_word(text: str, start: int, end: int, max_bridge: int = 2) -> Tuple[int,int]:
#     n = len(text); s, e = max(0,start), max(start,end)
#     def _is_inv(ch: str) -> bool: return ch in _BRIDGE_CHARS
#     while s > 0:
#         prev = text[s-1]; cur = text[s] if s < n else ""
#         if prev.isalnum() and cur.isalnum(): s -= 1; continue
#         j = s; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (s-1)>=0 and text[s-1].isalnum() and (j<n and text[j].isalnum()): s -= 1; continue
#         break
#     while e < n:
#         prev = text[e-1] if e>0 else ""; nxt = text[e]
#         if prev.isalnum() and nxt.isalnum(): e += 1; continue
#         j = e; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (e-1)>=0 and text[e-1].isalnum() and (j<n and text[j].isalnum()): e = j + 1; continue
#         break
#     while e < n and text[e] in ',"”’\')]}': e += 1
#     return s, e
# def _heal_split_word_left(text: str, start: int) -> int:
#     i = start
#     if i <= 1 or i >= len(text): return start
#     if text[i-1] != " ": return start
#     j = i - 2
#     while j >= 0 and text[j].isalpha(): j -= 1
#     prev_token = text[j+1:i-1]
#     if len(prev_token) == 1: return i - 2
#     return start
# def _overlaps_any(s: int, e: int, ranges: List[Tuple[int,int]]) -> bool:
#     for rs, re_ in ranges:
#         if e > rs and s < re_: return True
#     return False
# def _fuzzy_window_span(tl: str, nl: str, start: int, w: int) -> Tuple[float, Optional[Tuple[int,int]]]:
#     window = tl[start:start+w]
#     sm = difflib.SequenceMatcher(a=nl, b=window)
#     blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
#     if not blocks: return 0.0, None
#     coverage = sum(b.size for b in blocks) / max(1, len(nl))
#     first_b = min(blocks, key=lambda b: b.b); last_b = max(blocks, key=lambda b: b.b + b.size)
#     s = start + first_b.b; e = start + last_b.b + last_b.size
#     return coverage, (s, e)
# def find_span_smart(text: str, needle: str) -> Optional[Tuple[int,int]]:
#     if not text or not needle: return None
#     t_orig = text; t_norm = _normalize_keep_len(text); n_norm = _clean_quote_for_match(needle)
#     if not n_norm: return None
#     tl = t_norm.lower(); nl = n_norm.lower()
#     i = tl.find(nl)
#     if i != -1:
#         s, e = _snap_and_bridge_to_word(t_orig, i, i + len(nl)); s = _heal_split_word_left(t_orig, s)
#         return (s, e)
#     m = re.search(re.escape(nl).replace(r"\ ", r"\s+"), tl, flags=re.IGNORECASE)
#     if m:
#         s, e = _snap_and_bridge_to_word(t_orig, m.start(), m.end()); s = _heal_split_word_left(t_orig, s)
#         return (s, e)
#     if not STRICT_MATCH_ONLY and len(nl) >= 12:
#         w = max(60, min(240, len(nl) + 80))
#         best_cov, best_span = 0.0, None
#         step = max(1, w // 2)
#         for start in range(0, max(1, len(tl) - w + 1), step):
#             cov, se = _fuzzy_window_span(tl, nl, start, w)
#             if cov > best_cov: best_cov, best_span = cov, se
#         if best_span and best_cov >= 0.65:
#             s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
#             if s > 0 and t_orig[s-1:s+1].lower() in {" the", " a", " an"}: s -= 1
#             s = _heal_split_word_left(t_orig, s)
#             return (s, e)
#     if not STRICT_MATCH_ONLY:
#         keys = [w for w in _tokenize(nl) if len(w) >= 4][:8]
#         if len(keys) >= 2:
#             kset = set(keys)
#             best_score, best_span = 0.0, None
#             for s, e, seg in _iter_sentences_with_spans(t_norm):
#                 toks = set(_tokenize(seg)); ov = len(kset & toks)
#                 if ov == 0: continue
#                 score = ov / max(2, len(kset)); length_pen = min(1.0, 120 / max(20, e - s)); score *= (0.6 + 0.4 * length_pen)
#                 if score > best_score: best_score, best_span = score, (s, min(e, s + 400))
#             if best_span and best_score >= 0.35:
#                 s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1]); s = _heal_split_word_left(t_orig, s)
#                 return (s, e)
#     return None

# def merge_overlaps_and_adjacent(base_text: str,
#                                 spans: List[Tuple[int,int,str,str]],
#                                 max_gap: int = 2) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans = sorted(spans, key=lambda x: x[0]); out = [spans[0]]
#     _PUNCT_WS = set(" \t\r\n,.;:!?)('\"[]{}-—–…") | _BRIDGE_CHARS
#     for s, e, c, aid in spans[1:]:
#         ps, pe, pc, paid = out[-1]
#         if c == pc and s <= pe: out[-1] = (ps, max(pe, e), pc, paid); continue
#         if c == pc and s - pe <= max_gap:
#             gap = base_text[max(0, pe):max(0, s)]
#             if all((ch in _PUNCT_WS) for ch in gap): out[-1] = (ps, e, pc, paid); continue
#         out.append((s, e, c, aid))
#     return out

# def _is_heading_like(q: str) -> bool:
#     if not q: return True
#     s = q.strip()
#     if not re.search(r'[.!?]', s):
#         words = re.findall(r"[A-Za-z]+", s)
#         if 1 <= len(words) <= 7:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.8: return True
#         if s.lower() in {"introduction","voiceover","outro","epilogue","prologue","credits","title","hook","horrifying discovery","final decision"}: return True
#         if len(s) <= 3: return True
#     return False

# def _is_heading_context(script_text: str, s: int, e: int) -> bool:
#     left = script_text.rfind("\n", 0, s) + 1
#     right = script_text.find("\n", e); right = len(script_text) if right == -1 else right
#     line = script_text[left:right].strip()
#     if len(line) <= 70 and not re.search(r'[.!?]', line):
#         words = re.findall(r"[A-Za-z]+", line)
#         if 1 <= len(words) <= 8:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.7: return True
#     return False

# def _tighten_to_quote(script_text: str, span: Tuple[int,int], quote: str) -> Tuple[int,int]:
#     if not span or not quote: return span
#     s, e = span
#     if e <= s or s < 0 or e > len(script_text): return span
#     window = script_text[s:e]; win_norm = _normalize_keep_len(window).lower(); q_norm = _clean_quote_for_match(quote).lower()
#     if not q_norm: return span
#     i = win_norm.find(q_norm)
#     if i == -1:
#         m = re.search(re.escape(q_norm).replace(r"\ ", r"\s+"), win_norm, flags=re.IGNORECASE)
#         if not m: return span
#         i, j = m.start(), m.end()
#     else:
#         j = i + len(q_norm)
#     s2, e2 = s + i, s + j
#     s2, e2 = _snap_and_bridge_to_word(script_text, s2, e2); s2 = _heal_split_word_left(script_text, s2)
#     if s2 >= s and e2 <= e and e2 > s2: return (s2, e2)
#     return span

# def build_spans_by_param(script_text: str, data: dict, heading_ranges: Optional[List[Tuple[int,int]]] = None) -> Dict[str, List[Tuple[int,int,str,str]]]:
#     heading_ranges = heading_ranges or []
#     raw = (data or {}).get("per_parameter", {}) or {}
#     per: Dict[str, Dict[str, Any]] = {k:(v or {}) for k,v in raw.items()}
#     spans_map: Dict[str, List[Tuple[int,int,str,str]]] = {p: [] for p in PARAM_ORDER}
#     st.session_state["aoi_match_ranges"] = {}

#     for p in spans_map.keys():
#         color = PARAM_COLORS.get(p, "#ffd54f")
#         blk = per.get(p, {}) or {}
#         aois = blk.get("areas_of_improvement") or []
#         for idx, item in enumerate(aois, start=1):
#             raw_q = (item or {}).get("quote_verbatim", "") or ""
#             q = _sanitize_editor_text(raw_q)
#             clean = _clean_quote_for_match(re.sub(r"^[•\-\d\.\)\s]+", "", q).strip())
#             if not clean: continue
#             if _is_heading_like(clean): continue
#             pos = find_span_smart(script_text, clean)
#             if not pos: continue
#             pos = _tighten_to_quote(script_text, pos, raw_q)
#             s, e = pos
#             if heading_ranges and _overlaps_any(s, e, heading_ranges): continue
#             if _is_heading_context(script_text, s, e): continue
#             aid = f"{p.replace(' ','_')}-AOI-{idx}"
#             spans_map[p].append((s, e, color, aid))
#             st.session_state["aoi_match_ranges"][aid] = (s, e)
#     return spans_map

# # ---------- History (S3-aware + manifest + cache) ----------
# _MANIFEST_KEY = f"{HISTORY_DIR}/_manifest.json"

# def _manifest_read() -> List[dict]:
#     txt = read_text_key(_MANIFEST_KEY, default="")
#     if not txt.strip():
#         return []
#     try:
#         arr = json.loads(txt)
#         if isinstance(arr, list):
#             return arr
#     except Exception:
#         return []
#     return []

# def _manifest_append(entry: dict):
#     # read-modify-write with small retry
#     for k in range(3):
#         cur = _manifest_read()
#         cur.append(entry)
#         try:
#             save_text_key(_MANIFEST_KEY, json.dumps(cur, ensure_ascii=False, indent=2))
#             return
#         except Exception:
#             if k < 2:
#                 time.sleep(0.2 * (k + 1))
#             else:
#                 return

# def _maybe_copy_docx_to_history(source_docx_path: Optional[str], run_id: str) -> Optional[str]:
#     """
#     If source_docx_path is a local temp (downloaded), upload it under _history so Recents can re-render.
#     If it's already an S3 key/url, just return that key/url.
#     """
#     try:
#         if not source_docx_path:
#             return None
#         # If path exists locally (temp), push to S3 history
#         if os.path.exists(source_docx_path):
#             with open(source_docx_path, "rb") as f:
#                 save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#             return f"{HISTORY_DIR}/{run_id}.docx"
#         # If it's an S3 reference already
#         return source_docx_path
#     except Exception:
#         return None

# def _save_history_snapshot(title: str, data: dict, script_text: str,
#                            source_docx_path: Optional[str], heading_ranges: List[Tuple[int,int]],
#                            spans_by_param: Dict[str, List[Tuple[int,int,str,str]]],
#                            aoi_match_ranges: Dict[str, Tuple[int,int]]):
#     run_id = str(uuid.uuid4()); now = datetime.datetime.now()
#     created_at_iso = now.replace(microsecond=0).isoformat()
#     created_at_human = now.strftime("%Y-%m-%d %H:%M:%S")

#     stable_docx_key_or_path = _maybe_copy_docx_to_history(source_docx_path, run_id)

#     blob = {
#         "run_id": run_id, "title": title or "untitled",
#         "created_at": created_at_iso, "created_at_human": created_at_human,
#         "overall_rating": (data or {}).get("overall_rating", ""),
#         "scores": (data or {}).get("scores", {}),
#         "data": data or {}, "script_text": script_text or "",
#         "source_docx_path": stable_docx_key_or_path or source_docx_path,
#         "heading_ranges": heading_ranges or [],
#         "spans_by_param": spans_by_param or {},
#         "aoi_match_ranges": aoi_match_ranges or {},
#     }

#     out_name = f"{created_at_iso.replace(':','-')}__{run_id}.json"
#     out_key = f"{HISTORY_DIR}/{out_name}"
#     save_text_key(out_key, json.dumps(blob, ensure_ascii=False, indent=2))

#     # Append manifest entry (for list without ListBucket)
#     _manifest_append({
#         "run_id": run_id,
#         "key": out_key,  # exact S3 key to open
#         "title": blob["title"],
#         "created_at": blob["created_at"],
#         "created_at_human": blob["created_at_human"],
#         "overall_rating": blob["overall_rating"],
#     })

# def _load_all_history() -> List[dict]:
#     out: List[dict] = []

#     # Prefer manifest (works even without ListBucket)
#     man = _manifest_read()
#     if man:
#         # newest first by created_at
#         man_sorted = sorted(man, key=lambda r: r.get("created_at",""), reverse=True)
#         for m in man_sorted:
#             # we don't read the whole JSON here (fast listing); _open_history_by_key loads full
#             out.append({
#                 "run_id": m.get("run_id"),
#                 "title": m.get("title") or "(untitled)",
#                 "created_at": m.get("created_at"),
#                 "created_at_human": m.get("created_at_human", ""),
#                 "overall_rating": m.get("overall_rating", ""),
#                 "key": m.get("key"),  # exact S3 key
#             })

#     # Optional: also list from S3 if allowed, to backfill older runs
#     try:
#         keys = list_prefix(HISTORY_DIR)
#         for key in keys:
#             if key.endswith("_manifest.json"):
#                 continue
#             if any(x.get("key") == key for x in out):
#                 continue  # already present from manifest
#             try:
#                 txt = read_text_key(key, "")
#                 if not txt:
#                     continue
#                 j = json.loads(txt)
#                 out.append({
#                     "run_id": j.get("run_id"),
#                     "title": j.get("title","untitled"),
#                     "created_at": j.get("created_at") or "",
#                     "created_at_human": j.get("created_at_human",""),
#                     "overall_rating": j.get("overall_rating",""),
#                     "_key": key,  # loader-provided key
#                 })
#             except Exception:
#                 continue
#     except Exception:
#         pass

#     out.sort(key=lambda r: r.get("created_at") or "", reverse=True)

#     # Last-known-good cache to avoid flicker on transient failures
#     if out:
#         st.session_state["_last_history_cache"] = out
#         return out
#     else:
#         if st.session_state.get("_last_history_cache"):
#             return st.session_state["_last_history_cache"]
#         return out

# def _open_history_run_by_id(run_id: str) -> bool:
#     """Back-compat: open by run_id by searching manifest/list (less reliable than by key)."""
#     if not run_id:
#         return False
#     recs = _load_all_history()
#     match = next((r for r in recs if r.get("run_id") == run_id), None)
#     if not match:
#         return False
#     key = match.get("_key") or match.get("key")
#     if key:
#         return _open_history_by_key(key)
#     return False

# def _open_history_by_key(key: str) -> bool:
#     """
#     Open a history run by exact S3 key. Returns True if loaded.
#     """
#     if not key:
#         return False
#     try:
#         txt = read_text_key(key, "")
#         if not txt:
#             return False
#         jj = json.loads(txt)
#     except Exception:
#         return False

#     st.session_state.script_text      = jj.get("script_text","")
#     st.session_state.base_stem        = jj.get("title","untitled")
#     st.session_state.data             = jj.get("data",{})
#     st.session_state.heading_ranges   = jj.get("heading_ranges",[])
#     st.session_state.spans_by_param   = jj.get("spans_by_param",{})
#     st.session_state.param_choice     = None
#     st.session_state.source_docx_path = jj.get("source_docx_path")
#     st.session_state.review_ready     = True
#     st.session_state["aoi_match_ranges"] = jj.get("aoi_match_ranges", {})
#     st.session_state.ui_mode          = "review"
#     return True

# # def _render_recents_centerpane():
# #     st.subheader("📄 Recents")
# #     q = st.text_input("Filter by title…", "")

# #     cols = st.columns([1, 4])
# #     with cols[0]:
# #         if st.button("← Back"):
# #             st.session_state.ui_mode = "home"
# #             _clear_query_params()
# #             st.rerun()

# #     recs = _load_all_history()
# #     ql = q.strip().lower()
# #     if ql:
# #         recs = [r for r in recs if ql in (r.get("title","").lower())]

# #     if not recs:
# #         st.caption("No history yet.")
# #         return

# #     for rec in recs:
# #         run_id = rec.get("run_id")
# #         title = rec.get("title") or "(untitled)"
# #         created_h = rec.get("created_at_human","")
# #         overall = rec.get("overall_rating","")
# #         key = rec.get("_key") or rec.get("key")

# #         colA, colB = st.columns([6, 1])
# #         with colA:
# #             st.markdown(
# #                 f"""
# #                 <div class="rec-card">
# #                   <div class="rec-title">{title}</div>
# #                   <div class="rec-meta">{created_h}</div>
# #                   <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
# #                 </div>
# #                 """,
# #                 unsafe_allow_html=True
# #             )
# #         with colB:
# #             if st.button("Open", key=f"open_{run_id}", use_container_width=True):
# #                 st.session_state["_open_run_key"] = key
# #                 st.session_state["_open_run_id"] = run_id
# #                 st.rerun()


# def _render_recents_centerpane():
#     st.subheader("📄 Recents")
#     q = st.text_input("Filter by title…", "")

#     cols = st.columns([1, 4])
#     with cols[0]:
#         if st.button("← Back"):
#             st.session_state.ui_mode = "home"
#             _clear_query_params()
#             st.rerun()

#     recs = _load_all_history()
#     ql = q.strip().lower()
#     if ql:
#         recs = [r for r in recs if ql in (r.get("title","").lower())]

#     if not recs:
#         st.caption("No history yet.")
#         return

#     # Inline styles so you don't need to modify your global CSS
#     card_css = """
#     <style>
#       .rec-card { position:relative; display:block; text-decoration:none!important;
#         background:var(--m7-surface); border:1px solid var(--m7-border);
#         border-radius:12px; padding:14px 16px; margin:10px 0 16px;
#         box-shadow:0 1px 2px rgba(0,0,0,.06); color:var(--m7-on-surface)!important;
#         transition: filter .1s ease, transform .02s ease; }
#       .rec-card:hover{ filter:brightness(1.02); }
#       .rec-card:active{ transform: translateY(1px); }
#       .rec-row{ display:flex; align-items:center; justify-content:space-between; gap:12px; }
#       .rec-title{ font-weight:600; margin-bottom:.25rem; }
#       .rec-meta{ opacity:.85!important; font-size:12.5px; margin-bottom:.4rem; }
#       .rec-open{ margin-left:auto; display:inline-block; padding:6px 12px;
#         border:1px solid var(--m7-border); border-radius:10px;
#         text-decoration:none; font-weight:600; opacity:.95; }
#       .rec-open:hover{ filter:brightness(1.05); }
#     </style>
#     """
#     st.markdown(card_css, unsafe_allow_html=True)

#     for rec in recs:
#         run_id    = rec.get("run_id")
#         title     = rec.get("title") or "(untitled)"
#         created_h = rec.get("created_at_human","")
#         overall   = rec.get("overall_rating","")

#         st.markdown(
#             f"""
#             <a class="rec-card" href="?open={run_id}" target="_self" rel="noopener">
#             <div class="rec-row">
#                 <div>
#                 <div class="rec-title">{title}</div>
#                 <div class="rec-meta">{created_h}</div>
#                 <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
#                 </div>
#                 <span class="rec-open">Open</span>
#             </div>
#             </a>
#             """,
#             unsafe_allow_html=True
#         )

# # ---------- Sidebar ----------
# with st.sidebar:
#     if st.button("🆕 New review", use_container_width=True):
#         fp = st.session_state.get("flattened_docx_path")
#         if fp and os.path.exists(fp):
#             try: os.remove(fp)
#             except Exception: pass
#         for k in ["review_ready","script_text","base_stem","data","spans_by_param","param_choice",
#                   "source_docx_path","heading_ranges","flattened_docx_path","flatten_used"]:
#             st.session_state[k] = (
#                 False if k=="review_ready"
#                 else "" if k in ("script_text","base_stem")
#                 else {} if k=="spans_by_param"
#                 else [] if k=="heading_ranges"
#                 else None if k in ("source_docx_path","flattened_docx_path")
#                 else False if k=="flatten_used"
#                 else None
#             )
#         st.session_state.ui_mode = "home"
#         _clear_query_params()
#         st.rerun()

#     if st.button("📁 Recents", use_container_width=True):
#         st.session_state.ui_mode = "recents"
#         _clear_query_params()
#         st.rerun()

# # # ---------- Input screen ----------
# # def render_home():
# #     st.subheader("🎬 Script Source")

# #     tab_upload, tab_paste = st.tabs(["Upload file", "Paste text"])

# #     uploaded_file = None
# #     uploaded_name = None
# #     uploaded_key  = None

# #     def _safe_stem(s: str, fallback: str = "pasted_script") -> str:
# #         s = (s or "").strip()
# #         if not s:
# #             return fallback
# #         s = re.sub(r"[^A-Za-z0-9._\-]+", "_", s)
# #         s = s.strip("._-") or fallback
# #         return s

# #     with tab_upload:
# #         up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
# #         if up is not None:
# #             file_bytes = up.read()
# #             suffix = os.path.splitext(up.name)[1].lower()
# #             uploaded_key = f"{SCRIPTS_DIR}/{up.name}"
# #             save_bytes_key(uploaded_key, file_bytes)
# #             with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
# #                 tmp.write(file_bytes)
# #                 uploaded_file = tmp.name
# #             uploaded_name = os.path.splitext(os.path.basename(up.name))[0] or "uploaded_script"

# #     with tab_paste:
# #         paste_title = st.text_input("Title (optional)", placeholder="e.g., my_script")
# #         pasted_text = st.text_area(
# #             "Paste your script text here",
# #             height=360,
# #             placeholder="Paste the full script text (we’ll analyze this as-is)."
# #         )

# #     if st.button("🚀 Run Review", type="primary", use_container_width=True):
# #         base_stem = "uploaded_script"
# #         source_docx_path = None
# #         heading_ranges: List[Tuple[int,int]] = []
# #         script_text = ""

# #         if pasted_text and pasted_text.strip():
# #             base_stem = _safe_stem(paste_title, "pasted_script")
# #             script_text = pasted_text
# #             pasted_key = f"{SCRIPTS_DIR}/{base_stem}.txt"
# #             save_text_key(pasted_key, script_text)
# #             source_docx_path = pasted_key
# #             heading_ranges = []
# #         elif uploaded_file:
# #             base_stem = uploaded_name or "uploaded_script"
# #             if uploaded_file.lower().endswith(".docx"):
# #                 try:
# #                     left_text, used_left = extract_left_column_script_or_default(uploaded_file)
# #                 except Exception:
# #                     left_text, used_left = "", False

# #                 if used_left and left_text.strip():
# #                     script_text = left_text
# #                     source_docx_path = uploaded_key  # keep S3 key (we saved uploaded file)
# #                     heading_ranges = []
# #                 else:
# #                     path_to_use = uploaded_file
# #                     if _docx_contains_tables(path_to_use):
# #                         flat = flatten_docx_tables_to_longtext(path_to_use)
# #                         st.session_state.flattened_docx_path = flat
# #                         st.session_state.flatten_used = True
# #                         path_to_use = flat
# #                     script_text, heading_ranges = build_docx_text_with_meta(path_to_use)
# #                     source_docx_path = uploaded_key
# #             else:
# #                 script_text = load_script_file(uploaded_file)
# #                 source_docx_path = uploaded_key
# #         else:
# #             st.warning("Please upload a script **or** paste text in the second tab.")
# #             st.stop()

# #         if len(script_text.strip()) < 50:
# #             st.error("Extracted text looks too short. Please check your input.")
# #             st.stop()

# #         with st.spinner("Running analysis…"):
# #             try:
# #                 review_text = run_review_multi(
# #                     script_text=script_text,
# #                     prompts_dir=PROMPTS_DIR,  # treated as S3 prefix by review_engine_multi
# #                     temperature=0.0
# #                 )
# #             finally:
# #                 if uploaded_file and os.path.exists(uploaded_file):
# #                     try: os.remove(uploaded_file)
# #                     except Exception: pass

# #         data = extract_review_json(review_text)
# #         if not data:
# #             st.error("JSON not detected in model output.")
# #             st.stop()

# #         st.session_state.script_text      = script_text
# #         st.session_state.base_stem        = base_stem
# #         st.session_state.data             = data
# #         st.session_state.heading_ranges   = heading_ranges
# #         st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
# #         st.session_state.param_choice     = None
# #         st.session_state.source_docx_path = source_docx_path
# #         st.session_state.review_ready     = True
# #         st.session_state.ui_mode          = "review"

# # ---------- Input screen ----------
# def render_home():
#     st.subheader("🎬 Script Source")

#     tab_upload, tab_paste = st.tabs(["Upload file", "Paste text"])

#     uploaded_file = None
#     uploaded_name = None
#     uploaded_key  = None

#     def _safe_stem(s: str, fallback: str = "pasted_script") -> str:
#         s = (s or "").strip()
#         if not s:
#             return fallback
#         s = re.sub(r"[^A-Za-z0-9._\-]+", "_", s)
#         s = s.strip("._-") or fallback
#         return s

#     with tab_upload:
#         up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
#         if up is not None:
#             file_bytes = up.read()
#             suffix = os.path.splitext(up.name)[1].lower()
#             uploaded_key = f"{SCRIPTS_DIR}/{up.name}"
#             save_bytes_key(uploaded_key, file_bytes)
#             with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
#                 tmp.write(file_bytes)
#                 uploaded_file = tmp.name
#             uploaded_name = os.path.splitext(os.path.basename(up.name))[0] or "uploaded_script"

#     with tab_paste:
#         paste_title = st.text_input("Title (optional)", placeholder="e.g., my_script")
#         pasted_text = st.text_area(
#             "Paste your script text here",
#             height=360,
#             placeholder="Paste the full script text (we’ll analyze this as-is)."
#         )

#     if st.button("🚀 Run Review", type="primary", use_container_width=True):
#         base_stem = "uploaded_script"
#         source_docx_path = None
#         heading_ranges: List[Tuple[int,int]] = []
#         script_text = ""

#         if pasted_text and pasted_text.strip():
#             base_stem = _safe_stem(paste_title, "pasted_script")
#             script_text = pasted_text
#             pasted_key = f"{SCRIPTS_DIR}/{base_stem}.txt"
#             save_text_key(pasted_key, script_text)
#             source_docx_path = pasted_key
#             heading_ranges = []
#             # Render as plain text for pasted input
#             st.session_state["render_plain_from_docx"] = True

#         elif uploaded_file:
#             base_stem = uploaded_name or "uploaded_script"
#             if uploaded_file.lower().endswith(".docx"):
#                 try:
#                     left_text, used_left = extract_left_column_script_or_default(uploaded_file)
#                 except Exception:
#                     left_text, used_left = "", False

#                 if used_left and left_text.strip():
#                     # Two-column script detected → use ONLY left VO column and render plain
#                     script_text = left_text
#                     source_docx_path = uploaded_key  # keep S3 key (we saved uploaded file)
#                     heading_ranges = []
#                     st.session_state["render_plain_from_docx"] = True
#                 else:
#                     # Regular DOCX: flatten if tables exist; otherwise build plain text+meta
#                     path_to_use = uploaded_file
#                     if _docx_contains_tables(path_to_use):
#                         flat = flatten_docx_tables_to_longtext(path_to_use)
#                         st.session_state.flattened_docx_path = flat
#                         st.session_state.flatten_used = True
#                         path_to_use = flat
#                     script_text, heading_ranges = build_docx_text_with_meta(path_to_use)
#                     source_docx_path = uploaded_key
#                     # For non-left-column DOCX, prefer DOCX render unless we flattened
#                     st.session_state["render_plain_from_docx"] = bool(st.session_state.get("flatten_used"))
#             else:
#                 # txt/pdf → always render as plain text
#                 script_text = load_script_file(uploaded_file)
#                 source_docx_path = uploaded_key
#                 st.session_state["render_plain_from_docx"] = True
#         else:
#             st.warning("Please upload a script **or** paste text in the second tab.")
#             st.stop()

#         if len(script_text.strip()) < 50:
#             st.error("Extracted text looks too short. Please check your input.")
#             st.stop()

#         with st.spinner("Running analysis…"):
#             try:
#                 review_text = run_review_multi(
#                     script_text=script_text,
#                     prompts_dir=PROMPTS_DIR,  # treated as S3 prefix by review_engine_multi
#                     temperature=0.0
#                 )
#             finally:
#                 if uploaded_file and os.path.exists(uploaded_file):
#                     try:
#                         os.remove(uploaded_file)
#                     except Exception:
#                         pass

#         data = extract_review_json(review_text)
#         if not data:
#             st.error("JSON not detected in model output.")
#             st.stop()

#         st.session_state.script_text      = script_text
#         st.session_state.base_stem        = base_stem
#         st.session_state.data             = data
#         st.session_state.heading_ranges   = heading_ranges
#         st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
#         st.session_state.param_choice     = None
#         st.session_state.source_docx_path = source_docx_path
#         st.session_state.review_ready     = True
#         st.session_state.ui_mode          = "review"

#         _save_history_snapshot(
#             title=base_stem,
#             data=data,
#             script_text=script_text,
#             source_docx_path=source_docx_path,
#             heading_ranges=heading_ranges,
#             spans_by_param=st.session_state.spans_by_param,
#             aoi_match_ranges=st.session_state.get("aoi_match_ranges", {})
#         )

#         _clear_query_params()
#         st.rerun()

# # ---------- Results screen ----------
# def render_review():
#     script_text     = st.session_state.script_text
#     data            = st.session_state.data
#     spans_by_param  = st.session_state.spans_by_param
#     scores: Dict[str,int] = (data or {}).get("scores", {}) or {}
#     source_docx_path: Optional[str] = st.session_state.source_docx_path

#     # # If our source_docx_path is an S3 key/url, ensure we have a local copy for rendering
#     # docx_local: Optional[str] = None
#     # preferred = st.session_state.get("flattened_docx_path") if st.session_state.get("flatten_used") else None
#     # if not preferred and source_docx_path:
#     #     if source_docx_path.endswith(".docx"):
#     #         docx_local = ensure_local_copy(source_docx_path)

#     # If our source_docx_path is an S3 key/url, ensure we have a local copy for rendering
#     # BUT skip DOCX rendering when we used the left-column extractor.
#     docx_local: Optional[str] = None
#     render_plain = bool(st.session_state.get("render_plain_from_docx"))
#     preferred = st.session_state.get("flattened_docx_path") if st.session_state.get("flatten_used") else None
#     if not render_plain and not preferred and source_docx_path:
#         if source_docx_path.endswith(".docx"):
#             docx_local = ensure_local_copy(source_docx_path)


#     left, center, right = st.columns([1.1, 2.7, 1.4], gap="large")

#     with left:
#         st.subheader("Final score")
#         ordered = [p for p in PARAM_ORDER if p in scores]
#         df = pd.DataFrame({"Parameter": ordered, "Score (1–10)": [scores.get(p, "") for p in ordered]})
#         st.dataframe(df, hide_index=True, use_container_width=True)
#         st.markdown(f'**Overall:** {data.get("overall_rating","—")}/10')
#         st.divider()

#         strengths = (data or {}).get("strengths") or []
#         if not strengths:
#             per = (data or {}).get("per_parameter", {}) or {}
#             best = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
#             for name, sc in best:
#                 if sc >= 8 and name in per:
#                     exp = _sanitize_editor_text((per[name] or {}).get("explanation", "") or "")
#                     first = re.split(r"(?<=[.!?])\s+", exp.strip())[0] if exp else f"Consistently strong {name.lower()}."
#                     strengths.append(f"{name}: {first}")
#                 if len(strengths) >= 3: break

#         def _bullets(title: str, items):
#             st.markdown(f"**{title}**")
#             for s in (items or []):
#                 if isinstance(s, str) and s.strip(): st.write("• " + _sanitize_editor_text(s))
#             if not items: st.write("• —")

#         _bullets("Strengths", strengths)
#         _bullets("Weaknesses", data.get("weaknesses"))
#         _bullets("Suggestions", data.get("suggestions"))
#         _bullets("Drop-off Risks", data.get("drop_off_risks"))
#         st.markdown("**Viral Quotient**"); st.write(_sanitize_editor_text(data.get("viral_quotient","—")))

#     with right:
#         st.subheader("Parameters")
#         st.markdown('<div class="param-row">', unsafe_allow_html=True)
#         for p in [p for p in PARAM_ORDER if p in scores]:
#             if st.button(p, key=f"chip_{p}", help="Show inline AOI highlights for this parameter"):
#                 st.session_state.param_choice = p
#         st.markdown('</div>', unsafe_allow_html=True)

#         sel = st.session_state.param_choice
#         if sel:
#             blk = (data.get("per_parameter", {}) or {}).get(sel, {}) or {}
#             st.markdown(f"**{sel} — Score:** {scores.get(sel,'—')}/10")

#             if blk.get("explanation"):
#                 st.markdown("**Why this score**"); st.write(_sanitize_editor_text(blk["explanation"]))
#             if blk.get("weakness") and blk["weakness"] != "Not present":
#                 st.markdown("**Weakness**"); st.write(_sanitize_editor_text(blk["weakness"]))
#             if blk.get("suggestion") and blk["suggestion"] != "Not present":
#                 st.markdown("**Suggestion**"); st.write(_sanitize_editor_text(blk["suggestion"]))

#             # aoi = blk.get("areas_of_improvement") or []
#             # if aoi:
#             #     st.markdown("**Areas of Improvement**")
#             #     for i, item in enumerate(aoi, 1):
#             #         popover_fn = getattr(st, "popover", None)
#             #         aid = f"{sel.replace(' ','_')}-AOI-{i}"
#             #         s_e_map = st.session_state.get("aoi_match_ranges", {})
#             #         if aid in s_e_map:
#             #             s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#             #             line = (matched_line if len(matched_line) <= 320 else matched_line[:300] + "…")
#             #         else:
#             #             line = _sanitize_editor_text(item.get('quote_verbatim',''))
#             #         issue = _sanitize_editor_text(item.get('issue',''))
#             #         fix   = _sanitize_editor_text(item.get('fix',''))
#             #         why   = _sanitize_editor_text(item.get('why_this_helps',''))
#             #         label = f"Issue {i}"
#             #         if callable(popover_fn):
#             #             with popover_fn(label):
#             #                 if line: st.markdown(f"**Line:** {line}")
#             #                 if issue: st.markdown(f"**Issue:** {issue}")
#             #                 if fix: st.markdown(f"**Fix:** {fix}")
#             #                 if why: st.caption(why)
#             #         else:
#             #             with st.expander(label, expanded=False):
#             #                 if line: st.markdown(f"**Line:** {line}")
#             #                 if issue: st.markdown(f"**Issue:** {issue}")
#             #                 if fix: st.markdown(f"**Fix:** {fix}")
#             #                 if why: st.caption(why)

#             if blk.get("summary"):
#                 st.markdown("**Summary**"); st.write(_sanitize_editor_text(blk["summary"]))

#     with center:
#         st.subheader("Script with inline highlights")
#         spans = st.session_state.spans_by_param.get(st.session_state.param_choice, []) if st.session_state.param_choice else []

#         aoi_payload: Dict[str, Dict[str, str]] = {}
#         data_per = (data or {}).get("per_parameter") or {}
#         s_e_map = st.session_state.get("aoi_match_ranges", {})
#         sel = st.session_state.param_choice

#         def _mk_line(aid: str, fallback_q: str = "") -> str:
#             if aid in s_e_map:
#                 s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                 return matched_line if len(matched_line) <= 320 else matched_line[:300] + "…"
#             return _sanitize_editor_text(fallback_q or "")

#         def _collect(param_name: str):
#             blk = (data_per.get(param_name) or {})
#             for i, item in enumerate(blk.get("areas_of_improvement") or [], 1):
#                 aid = f"{param_name.replace(' ','_')}-AOI-{i}"
#                 aoi_payload[aid] = {
#                     "line": _mk_line(aid, (item or {}).get("quote_verbatim","")),
#                     "issue": _sanitize_editor_text((item or {}).get("issue","")),
#                     "fix": _sanitize_editor_text((item or {}).get("fix","")),
#                     "why": _sanitize_editor_text((item or {}).get("why_this_helps","")),
#                 }

#         if sel: _collect(sel)
#         else:
#             for pn in [p for p in PARAM_ORDER if p in data_per]:
#                 _collect(pn)

#         payload_json = json.dumps(aoi_payload, ensure_ascii=False)

#         frame_theme_css = """
#         <style>
#           :root{
#             --m7-surface: #eef2f7;
#             --m7-on-surface: #0f172a;
#             --m7-border: rgba(15,23,42,.14);
#           }
#           @media (prefers-color-scheme: dark){
#             :root{
#               --m7-surface: #2f333a;
#               --m7-on-surface: #ffffff;
#               --m7-border: rgba(255,255,255,.18);
#             }
#             body { background: transparent !important; }
#           }
#           .docxwrap{ background: var(--m7-surface); color: var(--m7-on-surface); border: 1px solid var(--m7-border); border-radius: 12px; padding: 16px 14px 18px; }
#           .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#           .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); }
#         </style>
#         """

#         tooltip_css = """
#         <style>
#         .aoi-pop { position: absolute; max-width: 520px; min-width: 320px; background: var(--m7-surface); border: 1px solid var(--m7-border); border-radius: 10px;
#           box-shadow: 0 10px 25px rgba(0,0,0,.12); padding: 12px 14px; z-index: 9999; transform: translateY(-8px); color: var(--m7-on-surface); }
#         .aoi-pop h4 { margin: 0 0 .35rem 0; font-size: .95rem; }
#         .aoi-pop p  { margin: .15rem 0; line-height: 1.5; }
#         .aoi-pop .muted { opacity:.85; font-size:.85rem; }
#         .aoi-arrow { position:absolute; left:50%; transform:translateX(-50%); bottom:-7px; width:0;height:0;border-left:7px solid transparent; border-right:7px solid transparent;border-top:7px solid var(--m7-border); }
#         .aoi-arrow::after{ content:""; position:absolute; left:-6px; top:-7px; width:0;height:0; border-left:6px solid transparent;border-right:6px solid transparent;border-top:6px solid var(--m7-surface); }
#         </style>
#         """

#         # Choose rendering source
#         # if docx_local and os.path.splitext(docx_local)[1].lower() == ".docx":
#         if (not render_plain) and docx_local and os.path.splitext(docx_local)[1].lower() == ".docx":
#             def render_docx_html_with_highlights(docx_path: str, highlight_spans: List[Tuple[int,int,str,str]]) -> str:
#                 doc = Document(docx_path)
#                 spans = [s for s in highlight_spans if s[0] < s[1]]
#                 spans.sort(key=lambda x: x[0])
#                 cur_span = 0
#                 current_offset = 0
#                 def esc(s: str) -> str:
#                     return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
#                 def open_mark_if_needed(html_parts, mark_state, color, end, aid):
#                     if not mark_state["open"]:
#                         html_parts.append(
#                             f'<mark class="aoi-mark" data-aid="{aid}" style="background:{color}33;border:1px solid {color};border-radius:3px;padding:0 2px;">'
#                         )
#                         mark_state.update(open=True, end=end, color=color, aid=aid)
#                 def close_mark_if_open(html_parts, mark_state):
#                     if mark_state["open"]:
#                         html_parts.append('</mark>')
#                         mark_state.update(open=False, end=None, color=None, aid=None)
#                 def _wrap_inline(safe_text: str, run) -> str:
#                     out = safe_text
#                     if getattr(run, "underline", False): out = f"<u>{out}</u>"
#                     if getattr(run, "italic", False): out = f"<em>{out}</em>"
#                     if getattr(run, "bold", False): out = f"<strong>{out}</strong>"
#                     return out
#                 def emit_run_text(run_text: str, run, html_parts: List[str], mark_state: Dict[str, Any]):
#                     nonlocal cur_span, current_offset
#                     t = run_text or ""; i = 0
#                     while i < len(t):
#                         next_start, next_end, color, next_aid = None, None, None, None
#                         if cur_span < len(spans):
#                             next_start, next_end, color, next_aid = spans[cur_span]
#                         if not mark_state["open"]:
#                             if cur_span >= len(spans) or current_offset + (len(t) - i) <= next_start:
#                                 chunk = t[i:]; html_parts.append(_wrap_inline(esc(chunk), run)); current_offset += len(chunk); break
#                             if current_offset < next_start:
#                                 take = next_start - current_offset
#                                 chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                                 current_offset += take; i += take; continue
#                             open_mark_if_needed(html_parts, mark_state, color, next_end, next_aid)
#                         else:
#                             take = min(mark_state["end"] - current_offset, len(t) - i)
#                             if take > 0:
#                                 chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                                 current_offset += take; i += take
#                             if current_offset >= mark_state["end"]:
#                                 close_mark_if_open(html_parts, mark_state)
#                                 cur_span += 1
#                 html: List[str] = ['<div class="docxwrap">']
#                 seen_tc_ids: set = set()
#                 for blk in _iter_docx_blocks(doc):
#                     if isinstance(blk, Paragraph):
#                         mark_state = {"open": False, "end": None, "color": None, "aid": None}
#                         sty = (blk.style.name or "").lower() if blk.style else ""
#                         open_tag = '<div class="h1">' if sty.startswith("heading 1") else \
#                                    '<div class="h2">' if sty.startswith("heading 2") else \
#                                    '<div class="h3">' if sty.startswith("heading 3") else "<p>"
#                         close_tag = "</div>" if sty.startswith("heading") else "</p>"
#                         html.append(open_tag)
#                         for run in blk.runs:
#                             emit_run_text(run.text or "", run, html, mark_state)
#                         close_mark_if_open(html, mark_state)
#                         html.append(close_tag)
#                         current_offset += 1
#                     else:
#                         html.append("<table>")
#                         for row in blk.rows:
#                             html.append("<tr>")
#                             row_cell_tcs = [(id(cell._tc), cell) for cell in row.cells]
#                             for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                                 html.append("<td>")
#                                 if tc_id not in seen_tc_ids:
#                                     seen_tc_ids.add(tc_id)
#                                     for p_idx, p in enumerate(cell.paragraphs):
#                                         mark_state = {"open": False, "end": None, "color": None, "aid": None}
#                                         html.append("<div>")
#                                         for run in p.runs:
#                                             emit_run_text(run.text or "", run, html, mark_state)
#                                         close_mark_if_open(html, mark_state)
#                                         html.append("</div>")
#                                         if p_idx != len(cell.paragraphs) - 1:
#                                             current_offset += 1
#                                 html.append("</td>")
#                                 if idx != len(row_cell_tcs) - 1: current_offset += 1
#                             html.append("</tr>"); current_offset += 1
#                         html.append("</table>"); current_offset += 1
#                 html.append("</div>")
#                 return "".join(html)

#             html_core = render_docx_html_with_highlights(
#                 docx_local,
#                 merge_overlaps_and_adjacent(script_text, spans)
#             )
#         else:
#             from html import escape as _esc
#             orig = script_text
#             spans2 = [s for s in merge_overlaps_and_adjacent(orig, spans) if s[0] < s[1]]
#             spans2.sort(key=lambda x: x[0])
#             cur = 0; buf: List[str] = []
#             for s,e,c,aid in spans2:
#                 if s > cur: buf.append(_esc(orig[cur:s]))
#                 buf.append(
#                     f'<mark class="aoi-mark" data-aid="{aid}" '
#                     f'style="background:{c}33;border:1px solid {c};border-radius:3px;padding:0 2px;">'
#                     f'{_esc(orig[s:e])}</mark>'
#                 )
#                 cur = e
#             if cur < len(orig): buf.append(_esc(orig[cur:]))
#             html_core = (
#                 '<div class="docxwrap"><p style="white-space:pre-wrap; '
#                 'line-height:1.7; font-family:ui-serif, Georgia, Times New Roman, serif;">'
#                 + "".join(buf) +
#                 '</p></div>'
#             )

#         html_shell = """
# %%FRAME_THEME_CSS%%
# %%TOOLTIP_CSS%%
# <div id="m7-doc">%%HTML_CORE%%</div>
# <div id="aoi-pop" class="aoi-pop" style="display:none;">
#   <div id="aoi-pop-content"></div>
#   <div class="aoi-arrow"></div>
# </div>
# <script>
# (function(){
#   const AOI = __PAYLOAD__;
#   const wrap = document.getElementById('m7-doc');
#   const pop  = document.getElementById('aoi-pop');
#   const body = document.getElementById('aoi-pop-content');

#   function resizeIframe() {
#     try {
#       const h = Math.max(
#         document.documentElement.scrollHeight,
#         document.body.scrollHeight
#       );
#       if (window.frameElement) {
#         window.frameElement.style.height = (h + 20) + 'px';
#         window.frameElement.style.width  = '100%';
#       }
#     } catch(e) {}
#   }
#   window.addEventListener('load', resizeIframe);
#   window.addEventListener('resize', resizeIframe);

#   function hide(){ pop.style.display='none'; }
#   function showFor(mark){
#     const aid = mark.getAttribute('data-aid');
#     const d = AOI[aid]; if(!d) return;
#     body.innerHTML =
#       (d.line  ? '<p><strong>Line:</strong> '  + d.line  + '</p>' : '') +
#       (d.issue ? '<p><strong>Issue:</strong> ' + d.issue + '</p>' : '') +
#       (d.fix   ? '<p><strong>Fix:</strong> '   + d.fix   + '</p>' : '') +
#       (d.why   ? '<p class="muted">'           + d.why   + '</p>' : '');
#     pop.style.display = 'block';

#     const r = mark.getBoundingClientRect();
#     const scY = window.scrollY || document.documentElement.scrollTop;
#     const scX = window.scrollX || document.documentElement.scrollLeft;
#     let top  = r.top + scY - pop.offsetHeight - 10;
#     let left = r.left + scX + r.width/2 - pop.offsetWidth/2;
#     if (top < 8) top = r.bottom + scY + 10;
#     if (left < 8) left = 8;
#     pop.style.top  = top + 'px';
#     pop.style.left = left + 'px';

#     resizeIframe();
#   }

#   wrap.addEventListener('click', (e) => {
#     const m = e.target.closest('.aoi-mark');
#     if(!m){ hide(); return; }
#     if(pop.style.display === 'block'){ hide(); }
#     showFor(m);
#     e.stopPropagation();
#   });

#   document.addEventListener('click', (e) => {
#     if(!e.target.closest('.aoi-pop') && !e.target.closest('.aoi-mark')) hide();
#   });
# })();
# </script>
# """
#         html_shell = (
#             html_shell
#             .replace("%%FRAME_THEME_CSS%%", frame_theme_css)
#             .replace("%%TOOLTIP_CSS%%", tooltip_css)
#             .replace("%%HTML_CORE%%", html_core)
#             .replace("__PAYLOAD__", payload_json)
#         )
#         components.html(html_shell, height=400, scrolling=False)

# # ---------- Router & query param open ----------
# _open_qp = _get_query_param("open")
# # keep legacy query-param open if present (will try via run_id fallback)
# if _open_qp and _open_history_run_by_id(_open_qp):
#     _clear_query_params()

# # Handle in-place open requests from Recents buttons FIRST
# if st.session_state.get("_open_run_key") or st.session_state.get("_open_run_id"):
#     key = st.session_state.pop("_open_run_key", None)
#     rid = st.session_state.pop("_open_run_id", None)

#     opened = False
#     if key:
#         opened = _open_history_by_key(key)  # most reliable
#     if not opened and rid:
#         opened = _open_history_run_by_id(rid)  # fallback via search/manifest
#     if opened:
#         _clear_query_params()
#         st.rerun()

# mode = st.session_state.ui_mode
# if mode == "recents":
#     _render_recents_centerpane()
# elif mode == "review" and st.session_state.review_ready:
#     render_review()
# else:
#     render_home()








###############################################












# # app_grammarly_ui.py — Runpod S3-only + Stable Recents + In-place open

# import os, re, glob, json, tempfile, difflib, uuid, datetime, shutil, time
# from pathlib import Path
# from typing import Dict, Any, List, Tuple, Optional

# import streamlit as st
# import pandas as pd
# import streamlit.components.v1 as components  # for inline HTML/JS popup

# # ---- utils & engine ----
# from utils1 import (
#     extract_review_json,
#     PARAM_ORDER,
#     load_script_file,
#     extract_left_column_script_or_default,  # <-- left-column extractor for DOCX tables
# )
# from review_engine_multi import run_review_multi

# # ---- DOCX rendering imports ----
# from docx import Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.text.paragraph import Paragraph
# from docx.table import Table

# # =========================
# # RunPod S3 (S3-only helpers)
# # =========================
# import boto3
# from botocore.config import Config
# from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError

# def _get_env(key: str, default: str = "") -> str:
#     v = os.getenv(key, "")
#     if v:
#         return v.strip()
#     try:
#         v2 = st.secrets.get(key)
#         if isinstance(v2, str):
#             return v2.strip()
#     except Exception:
#         pass
#     return (default or "").strip()

# # Primary config
# _RP_ENDPOINT = _get_env("RUNPOD_S3_ENDPOINT")
# _RP_BUCKET   = _get_env("RUNPOD_S3_BUCKET")
# _RP_REGION   = _get_env("RUNPOD_S3_REGION") or _get_env("AWS_DEFAULT_REGION") or ""

# # Credentials: prefer AWS_* if present; else RUNPOD_* fallbacks
# _AK = _get_env("AWS_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY_ID") or _get_env("RUNPOD_S3_ACCESS_KEY")
# _SK = _get_env("AWS_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_ACCESS_KEY") or _get_env("RUNPOD_S3_SECRET_KEY")
# _ST = _get_env("AWS_SESSION_TOKEN")  # optional

# # Options
# _FORCE_PATH = (_get_env("RUNPOD_S3_FORCE_PATH_STYLE") or "true").lower() in {"1","true","yes"}
# _USE_SSL    = (_get_env("RUNPOD_S3_USE_SSL") or "true").lower() in {"1","true","yes"}
# _VERIFY_SSL = (_get_env("RUNPOD_S3_VERIFY_SSL") or "true").lower() in {"1","true","yes"}

# def _s3_enabled() -> bool:
#     return bool(_RP_ENDPOINT and _RP_BUCKET and _AK and _SK)

# @st.cache_resource(show_spinner=False)
# def _s3_client():
#     if not _s3_enabled():
#         return None
#     session_kwargs = dict(
#         aws_access_key_id=_AK,
#         aws_secret_access_key=_SK,
#     )
#     if _ST:
#         session_kwargs["aws_session_token"] = _ST

#     cfg = Config(
#         signature_version="s3v4",
#         s3={"addressing_style": "path" if _FORCE_PATH else "auto"},
#         retries={"max_attempts": 3, "mode": "standard"}
#     )
#     return boto3.client(
#         "s3",
#         endpoint_url=_RP_ENDPOINT,
#         region_name=_RP_REGION or None,
#         use_ssl=_USE_SSL,
#         verify=_VERIFY_SSL,
#         config=cfg,
#         **session_kwargs,
#     )

# # ---------- S3 I/O (S3-only) ----------
# def save_text_key(key: str, text: str) -> str:
#     key = key.lstrip("/")
#     if not _s3_enabled():
#         # hard fail: S3-only mode
#         raise RuntimeError("S3 is not configured (RUNPOD_* / AWS_* envs).")
#     kwargs = {
#         "Bucket": _RP_BUCKET,
#         "Key": key,
#         "Body": text.encode("utf-8"),
#     }
#     if key.endswith(".json"):
#         kwargs["ContentType"] = "application/json"
#         kwargs["CacheControl"] = "no-store"
#     _s3_client().put_object(**kwargs)
#     return f"s3://{_RP_BUCKET}/{key}"

# def save_bytes_key(key: str, data: bytes) -> str:
#     key = key.lstrip("/")
#     if not _s3_enabled():
#         raise RuntimeError("S3 is not configured (RUNPOD_* / AWS_* envs).")
#     _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=data)
#     return f"s3://{_RP_BUCKET}/{key}"

# def read_text_key(key: str, default: str = "") -> str:
#     if not _s3_enabled():
#         return default
#     tries = 4
#     for k in range(tries):
#         try:
#             resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#             body = resp["Body"].read().decode("utf-8", errors="ignore")
#             if body.strip() == "" and k < tries - 1:
#                 time.sleep(0.25 * (k + 1))
#                 continue
#             return body
#         except Exception:
#             if k < tries - 1:
#                 time.sleep(0.25 * (k + 1))
#                 continue
#             return default
#     return default

# def read_bytes_key(key: str) -> Optional[bytes]:
#     if not _s3_enabled():
#         return None
#     tries = 4
#     for k in range(tries):
#         try:
#             resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
#             return resp["Body"].read()
#         except Exception:
#             if k < tries - 1:
#                 time.sleep(0.25 * (k + 1))
#                 continue
#             return None
#     return None

# def list_prefix(prefix: str) -> List[str]:
#     if not _s3_enabled():
#         return []
#     out: List[str] = []
#     token = None
#     s3_prefix = prefix.rstrip("/") + "/"
#     try:
#         while True:
#             kwargs = {"Bucket": _RP_BUCKET, "Prefix": s3_prefix}
#             if token:
#                 kwargs["ContinuationToken"] = token
#             resp = _s3_client().list_objects_v2(**kwargs)
#             for c in resp.get("Contents", []):
#                 k = c.get("Key", "")
#                 if k.endswith(".json"):
#                     out.append(k)
#             token = resp.get("NextContinuationToken")
#             if not token:
#                 break
#     except (ClientError, EndpointConnectionError, NoCredentialsError):
#         return []
#     return out

# def presigned_url(key: str, expires: int = 3600) -> Optional[str]:
#     if not _s3_enabled():
#         return None
#     try:
#         return _s3_client().generate_presigned_url(
#             "get_object",
#             Params={"Bucket": _RP_BUCKET, "Key": key},
#             ExpiresIn=expires
#         )
#     except ClientError:
#         return None

# def ensure_local_copy(key_or_keyurl: str) -> Optional[str]:
#     """
#     Always download to a temp file for parsing (DOCX/PDF).
#     """
#     if not _s3_enabled():
#         return None
#     key = key_or_keyurl
#     if key.startswith("s3://"):
#         parts = key.split("/", 3)
#         key = parts[3] if len(parts) >= 4 else ""
#     data = read_bytes_key(key)
#     if data is None:
#         return None
#     fd, tmp = tempfile.mkstemp(suffix=os.path.splitext(key)[1] or "")
#     os.close(fd)
#     with open(tmp, "wb") as f:
#         f.write(data)
#     return tmp

# def _s3_health_summary() -> dict:
#     info = {
#         "enabled": _s3_enabled(),
#         "endpoint": _RP_ENDPOINT,
#         "bucket": _RP_BUCKET,
#         "region": _RP_REGION,
#         "has_keys": bool(_AK and _SK),
#     }
#     if not _s3_enabled():
#         info["status"] = "disabled"
#         return info
#     try:
#         _ = _s3_client().list_objects_v2(Bucket=_RP_BUCKET, Prefix="Scriptmodel/outputs/_history/", MaxKeys=1)
#         info["status"] = "ok"
#     except Exception as e:
#         info["status"] = f"error: {getattr(e, 'response', {}).get('Error', {}).get('Code', str(e))}"
#     return info

# # ---------- Folders (all under Scriptmodel/) ----------
# BASE_PREFIX = "Scriptmodel"
# SCRIPTS_DIR = f"{BASE_PREFIX}/scripts"
# PROMPTS_DIR = f"{BASE_PREFIX}/prompts"
# OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
# HISTORY_DIR = f"{OUTPUT_DIR}/_history"

# # ---------- Colors ----------
# PARAM_COLORS: Dict[str, str] = {
#     "Suspense Building":              "#ff6b6b",
#     "Language/Tone":                  "#6b8cff",
#     "Intro + Main Hook/Cliffhanger":  "#ffb86b",
#     "Story Structure + Flow":         "#a78bfa",
#     "Pacing":                         "#f43f5e",
#     "Mini-Hooks (30–60s)":            "#eab308",
#     "Outro (Ending)":                 "#8b5cf6",
#     "Grammar & Spelling":             "#10b981",
# }

# STRICT_MATCH_ONLY = False

# # ---------- App config ----------
# st.set_page_config(page_title="M7 — Grammarly UI", page_icon="🕵️", layout="wide")

# # ---------- Header patch & CSS ----------
# def render_app_title():
#     st.markdown('<h1 class="app-title">Viral Script Reviewer</h1>', unsafe_allow_html=True)
#     st.markdown("""
#     <style>
#     html { color-scheme: light dark; }
#     :root{ --m7-surface:#eef2f7; --m7-on-surface:#0f172a; --m7-border:rgba(15,23,42,.14); --sep:#e5e7eb; }
#     @media (prefers-color-scheme: dark){
#       :root{ --m7-surface:#2f333a; --m7-on-surface:#ffffff; --m7-border:rgba(255,255,255,.18); --sep:#2a2f37; }
#     }
#     .stApp .block-container { padding-top: 4.25rem !important; }
#     .app-title{ font-weight:700; font-size:2.1rem; line-height:1.3; margin:0 0 1rem 0; padding-left:40px!important; padding-top:.25rem!important; }
#     [data-testid="collapsedControl"] { z-index: 6 !important; }
#     header[data-testid="stHeader"], .stAppHeader { background: transparent !important; box-shadow:none!important; }
#     @media (min-width: 992px){ .app-title { padding-left: 0 !important; } }
#     div[data-testid="column"]:nth-of-type(1){position:relative;}
#     div[data-testid="column"]:nth-of-type(1)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     div[data-testid="column"]:nth-of-type(2){position:relative;}
#     div[data-testid="column"]:nth-of-type(2)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
#     .m7-card{ background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:14px 16px; color:var(--m7-on-surface); }
#     .m7-card, .m7-card * { color:var(--m7-on-surface)!important; }
#     .docxwrap{ background:var(--m7-surface); color:var(--m7-on-surface); border:1px solid var(--m7-border); border-radius:12px; padding:16px 14px 18px; }
#     .docxwrap .h1,.docxwrap .h2,.docxwrap .h3 { font-weight:700; margin:10px 0 6px; }
#     .docxwrap .h1{font-size:1.3rem; border-bottom:2px solid currentColor; padding-bottom:4px;}
#     .docxwrap .h2{font-size:1.15rem; border-bottom:1px solid currentColor; padding-bottom:3px;}
#     .docxwrap .h3{font-size:1.05rem;}
#     .docxwrap p{ margin:10px 0; line-height:1.7; font-family: ui-serif, Georgia, "Times New Roman", serif; }
#     .docxwrap table{ border-collapse:collapse; width:100%; margin:12px 0; }
#     .docxwrap th,.docxwrap td{ border:1px solid var(--m7-border); padding:8px; vertical-align:top; line-height:1.6; }
#     .docxwrap mark{ padding:0 2px; border-radius:3px; border:1px solid var(--m7-border); cursor:pointer; }
#     .rec-card{ display:block; text-decoration:none!important; background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:14px 16px; margin:10px 0 16px; box-shadow:0 1px 2px rgba(0,0,0,.06); color:var(--m7-on-surface)!important; transition: filter .1s ease, transform .02s ease; }
#     .rec-card:hover{ filter:brightness(1.02); }
#     .rec-card:active{ transform: translateY(1px); }
#     .rec-title{font-weight:600; margin-bottom:.25rem;}
#     .rec-meta{opacity:.85!important; font-size:12.5px; margin-bottom:.4rem;}
#     .rec-row{display:flex; align-items:center; justify-content:space-between; gap:12px;}
#     .stTextInput>div>div, .stTextArea>div>div, .stNumberInput>div>div, .stDateInput>div>div, .stTimeInput>div>div, .stFileUploader>div, div[data-baseweb="select"]{ background:var(--m7-surface)!important; border:1px solid var(--m7-border)!important; border-radius:10px!important; color:var(--m7-on-surface)!important; }
#     .stTextInput input,.stTextArea textarea,.stNumberInput input,.stDateInput input,.stTimeInput input,.stFileUploader div,div[data-baseweb="select"] *{ color:var(--m7-on-surface)!important; }
#     .stTextInput input::placeholder,.stTextArea textarea::placeholder{ color:rgba(16,24,39,.55)!important; }
#     @media (prefers-color-scheme: dark){ .stTextInput input::placeholder,.stTextArea textarea::placeholder{ color:rgba(255,255,255,.75)!important; } }
#     div[data-testid="stFileUploaderDropzone"] label span { color: var(--m7-on-surface) !important; opacity:1!important; }
#     div[data-testid="stDataFrame"]{ background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:6px 8px; color:var(--m7-on-surface); }
#     .stMarkdown pre, pre[class*="language-"], .stCodeBlock{ background:var(--m7-surface)!important; color:var(--m7-on-surface)!important; border:1px solid var(--m7-border)!important; border-radius:12px!important; padding:12px 14px!important; overflow:auto; }
#     .stMarkdown pre code{ background:transparent!important; color:inherit!important; }
#     </style>
#     """, unsafe_allow_html=True)

# render_app_title()

# # ---------- Session defaults ----------
# for key, default in [
#     ("review_ready", False),
#     ("script_text", ""),
#     ("base_stem", ""),
#     ("data", None),
#     ("spans_by_param", {}),
#     ("param_choice", None),
#     ("source_docx_path", None),
#     ("heading_ranges", []),
#     ("flattened_docx_path", None),
#     ("flatten_used", False),
#     ("ui_mode", "home"),
#     ("render_plain_from_docx", False),   # NEW: persist render intent for Recents
# ]:
#     st.session_state.setdefault(key, default)

# # Recents stability helpers
# st.session_state.setdefault("_last_history_cache", [])
# st.session_state.setdefault("_open_run_key", None)
# st.session_state.setdefault("_open_run_id", None)

# # ---------- helpers for query params ----------
# def _get_query_param(key: str) -> Optional[str]:
#     val = None
#     try:
#         val = st.query_params.get(key)
#     except Exception:
#         q = st.experimental_get_query_params()
#         v = q.get(key)
#         if isinstance(v, list): val = v[0] if v else None
#         else: val = v
#     return val

# def _clear_query_params():
#     try:
#         st.query_params.clear()
#     except Exception:
#         st.experimental_set_query_params()

# # ---------- Sanitizer ----------
# _EMOJI_RE = re.compile(
#     r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
#     flags=re.UNICODE
# )
# def _sanitize_editor_text(s: Optional[str]) -> str:
#     if not s: return ""
#     t = str(s)
#     t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
#     t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
#     t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•)\s*', '', t, flags=re.M)
#     t = re.sub(r'^\s*[-*]\s+', '• ', t, flags=re.M)
#     t = _EMOJI_RE.sub('', t)
#     t = re.sub(r'[ \t]+', ' ', t)
#     t = re.sub(r'\n{3,}', '\n\n', t)
#     return t.strip()

# # ---------- DOCX traversal ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# # ---------- Auto-flatten ----------
# def _docx_contains_tables(path: str) -> bool:
#     doc = Document(path)
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Table):
#             return True
#     return False

# def _copy_paragraph(dest_doc: Document, src_para: Paragraph):
#     p = dest_doc.add_paragraph()
#     try:
#         if src_para.style and src_para.style.name:
#             p.style = src_para.style.name
#     except Exception:
#         pass
#     for run in src_para.runs:
#         r = p.add_run(run.text or "")
#         r.bold = run.bold
#         r.italic = run.italic
#         r.underline = run.underline
#     return p

# def flatten_docx_tables_to_longtext(source_path: str) -> str:
#     src = Document(source_path)
#     new = Document()
#     for blk in _iter_docx_blocks(src):
#         if isinstance(blk, Paragraph):
#             _copy_paragraph(new, blk)
#         else:
#             seen_tc_ids = set()
#             for row in blk.rows:
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     if tc_id in seen_tc_ids:
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     for p in cell.paragraphs:
#                         _copy_paragraph(new, p)
#                 new.add_paragraph("")
#             new.add_paragraph("")
#     fd, tmp_path = tempfile.mkstemp(suffix=".docx"); os.close(fd); new.save(tmp_path)
#     return tmp_path

# # ---------- Build plain text + heading ranges ----------
# def _iter_docx_blocks(document: Document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)

# def build_docx_text_with_meta(docx_path: str) -> Tuple[str, List[Tuple[int,int]]]:
#     doc = Document(docx_path)
#     out: List[str] = []
#     heading_ranges: List[Tuple[int,int]] = []
#     current_offset = 0

#     def _append_and_advance(s: str):
#         nonlocal current_offset
#         out.append(s); current_offset += len(s)

#     seen_tc_ids: set = set()
#     for blk in _iter_docx_blocks(doc):
#         if isinstance(blk, Paragraph):
#             para_text = "".join(run.text or "" for run in blk.runs)
#             sty = (blk.style.name or "").lower() if blk.style else ""
#             if sty.startswith("heading"):
#                 start = current_offset; end = start + len(para_text)
#                 heading_ranges.append((start, end))
#             _append_and_advance(para_text); _append_and_advance("\n")
#         else:
#             for row in blk.rows:
#                 row_cell_tcs = []
#                 for cell in row.cells:
#                     tc_id = id(cell._tc)
#                     row_cell_tcs.append((tc_id, cell))
#                 for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                     if tc_id in seen_tc_ids:
#                         if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                         continue
#                     seen_tc_ids.add(tc_id)
#                     cell_text_parts: List[str] = []
#                     for i, p in enumerate(cell.paragraphs):
#                         t = "".join(r.text or "" for r in p.runs)
#                         sty = (p.style.name or "").lower() if p.style else ""
#                         if sty.startswith("heading"):
#                             hs = current_offset + sum(len(x) for x in cell_text_parts)
#                             he = hs + len(t)
#                             heading_ranges.append((hs, he))
#                         cell_text_parts.append(t)
#                         if i != len(cell.paragraphs) - 1:
#                             cell_text_parts.append("\n")
#                     cell_text = "".join(cell_text_parts)
#                     _append_and_advance(cell_text)
#                     if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
#                 _append_and_advance("\n")
#             _append_and_advance("\n")

#     return "".join(out), heading_ranges

# # ---------- Matching / spans (unchanged core logic) ----------
# _BRIDGE_CHARS = set("\u200b\u200c\u200d\u2060\ufeff\xa0\u00ad")
# def _normalize_keep_len(s: str) -> str:
#     trans = {
#         "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
#         "\u2013": "-", "\u2014": "-",
#         "\xa0": " ",
#         "\u200b": " ", "\u200c": " ", "\u200d": " ", "\u2060": " ",
#         "\ufeff": " ", "\u00ad": " ",
#     }
#     return (s or "").translate(str.maketrans(trans))
# def _tokenize(s: str) -> List[str]: return re.findall(r"\w+", (s or "").lower())
# def _iter_sentences_with_spans(text: str) -> List[Tuple[int,int,str]]:
#     spans = []
#     for m in re.finditer(r'[^.!?]+[.!?]+|\Z', text, flags=re.S):
#         s, e = m.start(), m.end()
#         seg = text[s:e]
#         if seg.strip(): spans.append((s, e, seg))
#     return spans
# def _squash_ws(s: str) -> str: return re.sub(r"\s+", " ", s or "").strip()
# def _clean_quote_for_match(q: str) -> str:
#     if not q: return ""
#     q = _normalize_keep_len(q).strip()
#     q = re.sub(r'^[\'"“”‘’\[\(\{<…\-\–\—\s]+', '', q)
#     q = re.sub(r'[\'"“”‘’\]\)\}>…\-\–\—\s]+$', '', q)
#     return _squash_ws(q)
# def _snap_and_bridge_to_word(text: str, start: int, end: int, max_bridge: int = 2) -> Tuple[int,int]:
#     n = len(text); s, e = max(0,start), max(start,end)
#     def _is_inv(ch: str) -> bool: return ch in _BRIDGE_CHARS
#     while s > 0:
#         prev = text[s-1]; cur = text[s] if s < n else ""
#         if prev.isalnum() and cur.isalnum(): s -= 1; continue
#         j = s; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (s-1)>=0 and text[s-1].isalnum() and (j<n and text[j].isalnum()): s -= 1; continue
#         break
#     while e < n:
#         prev = text[e-1] if e>0 else ""; nxt = text[e]
#         if prev.isalnum() and nxt.isalnum(): e += 1; continue
#         j = e; brid = 0
#         while j < n and _is_inv(text[j]): brid += 1; j += 1
#         if brid and (e-1)>=0 and text[e-1].isalnum() and (j<n and text[j].isalnum()): e = j + 1; continue
#         break
#     while e < n and text[e] in ',"”’\')]}': e += 1
#     return s, e
# def _heal_split_word_left(text: str, start: int) -> int:
#     i = start
#     if i <= 1 or i >= len(text): return start
#     if text[i-1] != " ": return start
#     j = i - 2
#     while j >= 0 and text[j].isalpha(): j -= 1
#     prev_token = text[j+1:i-1]
#     if len(prev_token) == 1: return i - 2
#     return start
# def _overlaps_any(s: int, e: int, ranges: List[Tuple[int,int]]) -> bool:
#     for rs, re_ in ranges:
#         if e > rs and s < re_: return True
#     return False
# def _fuzzy_window_span(tl: str, nl: str, start: int, w: int) -> Tuple[float, Optional[Tuple[int,int]]]:
#     window = tl[start:start+w]
#     sm = difflib.SequenceMatcher(a=nl, b=window)
#     blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
#     if not blocks: return 0.0, None
#     coverage = sum(b.size for b in blocks) / max(1, len(nl))
#     first_b = min(blocks, key=lambda b: b.b); last_b = max(blocks, key=lambda b: b.b + b.size)
#     s = start + first_b.b; e = start + last_b.b + last_b.size
#     return coverage, (s, e)
# def find_span_smart(text: str, needle: str) -> Optional[Tuple[int,int]]:
#     if not text or not needle: return None
#     t_orig = text; t_norm = _normalize_keep_len(text); n_norm = _clean_quote_for_match(needle)
#     if not n_norm: return None
#     tl = t_norm.lower(); nl = n_norm.lower()
#     i = tl.find(nl)
#     if i != -1:
#         s, e = _snap_and_bridge_to_word(t_orig, i, i + len(nl)); s = _heal_split_word_left(t_orig, s)
#         return (s, e)
#     m = re.search(re.escape(nl).replace(r"\ ", r"\s+"), tl, flags=re.IGNORECASE)
#     if m:
#         s, e = _snap_and_bridge_to_word(t_orig, m.start(), m.end()); s = _heal_split_word_left(t_orig, s)
#         return (s, e)
#     if not STRICT_MATCH_ONLY and len(nl) >= 12:
#         w = max(60, min(240, len(nl) + 80))
#         best_cov, best_span = 0.0, None
#         step = max(1, w // 2)
#         for start in range(0, max(1, len(tl) - w + 1), step):
#             cov, se = _fuzzy_window_span(tl, nl, start, w)
#             if cov > best_cov: best_cov, best_span = cov, se
#         if best_span and best_cov >= 0.65:
#             s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
#             if s > 0 and t_orig[s-1:s+1].lower() in {" the", " a", " an"}: s -= 1
#             s = _heal_split_word_left(t_orig, s)
#             return (s, e)
#     if not STRICT_MATCH_ONLY:
#         keys = [w for w in _tokenize(nl) if len(w) >= 4][:8]
#         if len(keys) >= 2:
#             kset = set(keys)
#             best_score, best_span = 0.0, None
#             for s, e, seg in _iter_sentences_with_spans(t_norm):
#                 toks = set(_tokenize(seg)); ov = len(kset & toks)
#                 if ov == 0: continue
#                 score = ov / max(2, len(kset)); length_pen = min(1.0, 120 / max(20, e - s)); score *= (0.6 + 0.4 * length_pen)
#                 if score > best_score: best_score, best_span = score, (s, min(e, s + 400))
#             if best_span and best_score >= 0.35:
#                 s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1]); s = _heal_split_word_left(t_orig, s)
#                 return (s, e)
#     return None

# def merge_overlaps_and_adjacent(base_text: str,
#                                 spans: List[Tuple[int,int,str,str]],
#                                 max_gap: int = 2) -> List[Tuple[int,int,str,str]]:
#     if not spans: return []
#     spans = sorted(spans, key=lambda x: x[0]); out = [spans[0]]
#     _PUNCT_WS = set(" \t\r\n,.;:!?)('\"[]{}-—–…") | _BRIDGE_CHARS
#     for s, e, c, aid in spans[1:]:
#         ps, pe, pc, paid = out[-1]
#         if c == pc and s <= pe: out[-1] = (ps, max(pe, e), pc, paid); continue
#         if c == pc and s - pe <= max_gap:
#             gap = base_text[max(0, pe):max(0, s)]
#             if all((ch in _PUNCT_WS) for ch in gap): out[-1] = (ps, e, pc, paid); continue
#         out.append((s, e, c, aid))
#     return out

# def _is_heading_like(q: str) -> bool:
#     if not q: return True
#     s = q.strip()
#     if not re.search(r'[.!?]', s):
#         words = re.findall(r"[A-Za-z]+", s)
#         if 1 <= len(words) <= 7:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.8: return True
#         if s.lower() in {"introduction","voiceover","outro","epilogue","prologue","credits","title","hook","horrifying discovery","final decision"}: return True
#         if len(s) <= 3: return True
#     return False

# def _is_heading_context(script_text: str, s: int, e: int) -> bool:
#     left = script_text.rfind("\n", 0, s) + 1
#     right = script_text.find("\n", e); right = len(script_text) if right == -1 else right
#     line = script_text[left:right].strip()
#     if len(line) <= 70 and not re.search(r'[.!?]', line):
#         words = re.findall(r"[A-Za-z]+", line)
#         if 1 <= len(words) <= 8:
#             caps = sum(1 for w in words if w and w[0].isupper())
#             if caps / max(1, len(words)) >= 0.7: return True
#     return False

# def _tighten_to_quote(script_text: str, span: Tuple[int,int], quote: str) -> Tuple[int,int]:
#     if not span or not quote: return span
#     s, e = span
#     if e <= s or s < 0 or e > len(script_text): return span
#     window = script_text[s:e]; win_norm = _normalize_keep_len(window).lower(); q_norm = _clean_quote_for_match(quote).lower()
#     if not q_norm: return span
#     i = win_norm.find(q_norm)
#     if i == -1:
#         m = re.search(re.escape(q_norm).replace(r"\ ", r"\s+"), win_norm, flags=re.IGNORECASE)
#         if not m: return span
#         i, j = m.start(), m.end()
#     else:
#         j = i + len(q_norm)
#     s2, e2 = s + i, s + j
#     s2, e2 = _snap_and_bridge_to_word(script_text, s2, e2); s2 = _heal_split_word_left(script_text, s2)
#     if s2 >= s and e2 <= e and e2 > s2: return (s2, e2)
#     return span

# def build_spans_by_param(script_text: str, data: dict, heading_ranges: Optional[List[Tuple[int,int]]] = None) -> Dict[str, List[Tuple[int,int,str,str]]]:
#     heading_ranges = heading_ranges or []
#     raw = (data or {}).get("per_parameter", {}) or {}
#     per: Dict[str, Dict[str, Any]] = {k:(v or {}) for k,v in raw.items()}
#     spans_map: Dict[str, List[Tuple[int,int,str,str]]] = {p: [] for p in PARAM_ORDER}
#     st.session_state["aoi_match_ranges"] = {}

#     for p in spans_map.keys():
#         color = PARAM_COLORS.get(p, "#ffd54f")
#         blk = per.get(p, {}) or {}
#         aois = blk.get("areas_of_improvement") or []
#         for idx, item in enumerate(aois, start=1):
#             raw_q = (item or {}).get("quote_verbatim", "") or ""
#             q = _sanitize_editor_text(raw_q)
#             clean = _clean_quote_for_match(re.sub(r"^[•\-\d\.\)\s]+", "", q).strip())
#             if not clean: continue
#             if _is_heading_like(clean): continue
#             pos = find_span_smart(script_text, clean)
#             if not pos: continue
#             pos = _tighten_to_quote(script_text, pos, raw_q)
#             s, e = pos
#             if heading_ranges and _overlaps_any(s, e, heading_ranges): continue
#             if _is_heading_context(script_text, s, e): continue
#             aid = f"{p.replace(' ','_')}-AOI-{idx}"
#             spans_map[p].append((s, e, color, aid))
#             st.session_state["aoi_match_ranges"][aid] = (s, e)
#     return spans_map

# # ---------- History (S3-aware + manifest + cache) ----------
# _MANIFEST_KEY = f"{HISTORY_DIR}/_manifest.json"

# def _manifest_read() -> List[dict]:
#     txt = read_text_key(_MANIFEST_KEY, default="")
#     if not txt.strip():
#         return []
#     try:
#         arr = json.loads(txt)
#         if isinstance(arr, list):
#             return arr
#     except Exception:
#         return []
#     return []

# def _manifest_append(entry: dict):
#     # read-modify-write with small retry
#     for k in range(3):
#         cur = _manifest_read()
#         cur.append(entry)
#         try:
#             save_text_key(_MANIFEST_KEY, json.dumps(cur, ensure_ascii=False, indent=2))
#             return
#         except Exception:
#             if k < 2:
#                 time.sleep(0.2 * (k + 1))
#             else:
#                 return

# def _maybe_copy_docx_to_history(source_docx_path: Optional[str], run_id: str) -> Optional[str]:
#     """
#     If source_docx_path is a local temp (downloaded), upload it under _history so Recents can re-render.
#     If it's already an S3 key/url, just return that key/url.
#     """
#     try:
#         if not source_docx_path:
#             return None
#         # If path exists locally (temp), push to S3 history
#         if os.path.exists(source_docx_path):
#             with open(source_docx_path, "rb") as f:
#                 save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
#             return f"{HISTORY_DIR}/{run_id}.docx"
#         # If it's an S3 reference already
#         return source_docx_path
#     except Exception:
#         return None

# def _save_history_snapshot(title: str, data: dict, script_text: str,
#                            source_docx_path: Optional[str], heading_ranges: List[Tuple[int,int]],
#                            spans_by_param: Dict[str, List[Tuple[int,int,str,str]]],
#                            aoi_match_ranges: Dict[str, Tuple[int,int]],
#                            used_left_column: bool = False):  # NEW
#     run_id = str(uuid.uuid4()); now = datetime.datetime.now()
#     created_at_iso = now.replace(microsecond=0).isoformat()
#     created_at_human = now.strftime("%Y-%m-%d %H:%M:%S")

#     stable_docx_key_or_path = _maybe_copy_docx_to_history(source_docx_path, run_id)

#     blob = {
#         "run_id": run_id, "title": title or "untitled",
#         "created_at": created_at_iso, "created_at_human": created_at_human,
#         "overall_rating": (data or {}).get("overall_rating", ""),
#         "scores": (data or {}).get("scores", {}),
#         "data": data or {}, "script_text": script_text or "",
#         "source_docx_path": stable_docx_key_or_path or source_docx_path,
#         "heading_ranges": heading_ranges or [],
#         "spans_by_param": spans_by_param or {},
#         "aoi_match_ranges": aoi_match_ranges or {},
#         # persist the render intent so Recents shows the same thing
#         "used_left_column": bool(used_left_column),
#         "render_plain_from_docx": bool(used_left_column),  # mirror for forward/back compat
#     }

#     out_name = f"{created_at_iso.replace(':','-')}__{run_id}.json"
#     out_key = f"{HISTORY_DIR}/{out_name}"
#     save_text_key(out_key, json.dumps(blob, ensure_ascii=False, indent=2))

#     # Append manifest entry (for list without ListBucket)
#     _manifest_append({
#         "run_id": run_id,
#         "key": out_key,  # exact S3 key to open
#         "title": blob["title"],
#         "created_at": blob["created_at"],
#         "created_at_human": blob["created_at_human"],
#         "overall_rating": blob["overall_rating"],
#     })

# def _load_all_history() -> List[dict]:
#     out: List[dict] = []

#     # Prefer manifest (works even without ListBucket)
#     man = _manifest_read()
#     if man:
#         # newest first by created_at
#         man_sorted = sorted(man, key=lambda r: r.get("created_at",""), reverse=True)
#         for m in man_sorted:
#             # we don't read the whole JSON here (fast listing); _open_history_by_key loads full
#             out.append({
#                 "run_id": m.get("run_id"),
#                 "title": m.get("title") or "(untitled)",
#                 "created_at": m.get("created_at"),
#                 "created_at_human": m.get("created_at_human", ""),
#                 "overall_rating": m.get("overall_rating", ""),
#                 "key": m.get("key"),  # exact S3 key
#             })

#     # Optional: also list from S3 if allowed, to backfill older runs
#     try:
#         keys = list_prefix(HISTORY_DIR)
#         for key in keys:
#             if key.endswith("_manifest.json"):
#                 continue
#             if any(x.get("key") == key for x in out):
#                 continue  # already present from manifest
#             try:
#                 txt = read_text_key(key, "")
#                 if not txt:
#                     continue
#                 j = json.loads(txt)
#                 out.append({
#                     "run_id": j.get("run_id"),
#                     "title": j.get("title","untitled"),
#                     "created_at": j.get("created_at") or "",
#                     "created_at_human": j.get("created_at_human",""),
#                     "overall_rating": j.get("overall_rating",""),
#                     "_key": key,  # loader-provided key
#                 })
#             except Exception:
#                 continue
#     except Exception:
#         pass

#     out.sort(key=lambda r: r.get("created_at") or "", reverse=True)

#     # Last-known-good cache to avoid flicker on transient failures
#     if out:
#         st.session_state["_last_history_cache"] = out
#         return out
#     else:
#         if st.session_state.get("_last_history_cache"):
#             return st.session_state["_last_history_cache"]
#         return out

# def _open_history_run_by_id(run_id: str) -> bool:
#     """Back-compat: open by run_id by searching manifest/list (less reliable than by key)."""
#     if not run_id:
#         return False
#     recs = _load_all_history()
#     match = next((r for r in recs if r.get("run_id") == run_id), None)
#     if not match:
#         return False
#     key = match.get("_key") or match.get("key")
#     if key:
#         return _open_history_by_key(key)
#     return False

# def _open_history_by_key(key: str) -> bool:
#     """
#     Open a history run by exact S3 key. Returns True if loaded.
#     """
#     if not key:
#         return False
#     try:
#         txt = read_text_key(key, "")
#         if not txt:
#             return False
#         jj = json.loads(txt)
#     except Exception:
#         return False

#     # Respect saved render intent
#     used_left = bool(jj.get("used_left_column", False) or jj.get("render_plain_from_docx", False))
#     st.session_state["render_plain_from_docx"] = used_left

#     st.session_state.script_text      = jj.get("script_text","")
#     st.session_state.base_stem        = jj.get("title","untitled")
#     st.session_state.data             = jj.get("data",{})
#     st.session_state.heading_ranges   = jj.get("heading_ranges",[])
#     st.session_state.spans_by_param   = jj.get("spans_by_param",{})
#     st.session_state.param_choice     = None
#     # If we should render plain, ensure we don't try to re-render the DOCX table
#     if used_left:
#         st.session_state.source_docx_path = None
#     else:
#         st.session_state.source_docx_path = jj.get("source_docx_path")
#     st.session_state.review_ready     = True
#     st.session_state["aoi_match_ranges"] = jj.get("aoi_match_ranges", {})
#     st.session_state.ui_mode          = "review"
#     return True

# def _render_recents_centerpane():
#     st.subheader("📄 Recents")
#     q = st.text_input("Filter by title…", "")

#     cols = st.columns([1, 4])
#     with cols[0]:
#         if st.button("← Back"):
#             st.session_state.ui_mode = "home"
#             _clear_query_params()
#             st.rerun()

#     recs = _load_all_history()
#     ql = q.strip().lower()
#     if ql:
#         recs = [r for r in recs if ql in (r.get("title","").lower())]

#     if not recs:
#         st.caption("No history yet.")
#         return

#     # Inline styles so you don't need to modify your global CSS
#     card_css = """
#     <style>
#       .rec-card { position:relative; display:block; text-decoration:none!important;
#         background:var(--m7-surface); border:1px solid var(--m7-border);
#         border-radius:12px; padding:14px 16px; margin:10px 0 16px;
#         box-shadow:0 1px 2px rgba(0,0,0,.06); color:var(--m7-on-surface)!important;
#         transition: filter .1s ease, transform .02s ease; }
#       .rec-card:hover{ filter:brightness(1.02); }
#       .rec-card:active{ transform: translateY(1px); }
#       .rec-row{ display:flex; align-items:center; justify-content:space-between; gap:12px; }
#       .rec-title{ font-weight:600; margin-bottom:.25rem; }
#       .rec-meta{ opacity:.85!important; font-size:12.5px; margin-bottom:.4rem; }
#       .rec-open{ margin-left:auto; display:inline-block; padding:6px 12px;
#         border:1px solid var(--m7-border); border-radius:10px;
#         text-decoration:none; font-weight:600; opacity:.95; }
#       .rec-open:hover{ filter:brightness(1.05); }
#     </style>
#     """
#     st.markdown(card_css, unsafe_allow_html=True)

#     for rec in recs:
#         run_id    = rec.get("run_id")
#         title     = rec.get("title") or "(untitled)"
#         created_h = rec.get("created_at_human","")
#         overall   = rec.get("overall_rating","")

#         st.markdown(
#             f"""
#             <a class="rec-card" href="?open={run_id}" target="_self" rel="noopener">
#             <div class="rec-row">
#                 <div>
#                 <div class="rec-title">{title}</div>
#                 <div class="rec-meta">{created_h}</div>
#                 <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
#                 </div>
#                 <span class="rec-open">Open</span>
#             </div>
#             </a>
#             """,
#             unsafe_allow_html=True
#         )

# # ---------- Sidebar ----------
# with st.sidebar:
#     if st.button("🆕 New review", use_container_width=True):
#         fp = st.session_state.get("flattened_docx_path")
#         if fp and os.path.exists(fp):
#             try: os.remove(fp)
#             except Exception: pass
#         for k in ["review_ready","script_text","base_stem","data","spans_by_param","param_choice",
#                   "source_docx_path","heading_ranges","flattened_docx_path","flatten_used","render_plain_from_docx"]:
#             st.session_state[k] = (
#                 False if k=="review_ready"
#                 else "" if k in ("script_text","base_stem")
#                 else {} if k=="spans_by_param"
#                 else [] if k=="heading_ranges"
#                 else None if k in ("source_docx_path","flattened_docx_path")
#                 else False if k in ("flatten_used","render_plain_from_docx")
#                 else None
#             )
#         st.session_state.ui_mode = "home"
#         _clear_query_params()
#         st.rerun()

#     if st.button("📁 Recents", use_container_width=True):
#         st.session_state.ui_mode = "recents"
#         _clear_query_params()
#         st.rerun()

# # ---------- Input screen ----------
# def render_home():
#     st.subheader("🎬 Script Source")

#     tab_upload, tab_paste = st.tabs(["Upload file", "Paste text"])

#     uploaded_file = None
#     uploaded_name = None
#     uploaded_key  = None

#     def _safe_stem(s: str, fallback: str = "pasted_script") -> str:
#         s = (s or "").strip()
#         if not s:
#             return fallback
#         s = re.sub(r"[^A-Za-z0-9._\-]+", "_", s)
#         s = s.strip("._-") or fallback
#         return s

#     with tab_upload:
#         up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
#         if up is not None:
#             file_bytes = up.read()
#             suffix = os.path.splitext(up.name)[1].lower()
#             uploaded_key = f"{SCRIPTS_DIR}/{up.name}"
#             save_bytes_key(uploaded_key, file_bytes)
#             with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
#                 tmp.write(file_bytes)
#                 uploaded_file = tmp.name
#             uploaded_name = os.path.splitext(os.path.basename(up.name))[0] or "uploaded_script"

#     with tab_paste:
#         paste_title = st.text_input("Title (optional)", placeholder="e.g., my_script")
#         pasted_text = st.text_area(
#             "Paste your script text here",
#             height=360,
#             placeholder="Paste the full script text (we’ll analyze this as-is)."
#         )

#     if st.button("🚀 Run Review", type="primary", use_container_width=True):
#         base_stem = "uploaded_script"
#         source_docx_path = None
#         heading_ranges: List[Tuple[int,int]] = []
#         script_text = ""

#         if pasted_text and pasted_text.strip():
#             base_stem = _safe_stem(paste_title, "pasted_script")
#             script_text = pasted_text
#             pasted_key = f"{SCRIPTS_DIR}/{base_stem}.txt"
#             save_text_key(pasted_key, script_text)
#             source_docx_path = pasted_key
#             heading_ranges = []
#             # Render as plain text for pasted input
#             st.session_state["render_plain_from_docx"] = True

#         elif uploaded_file:
#             base_stem = uploaded_name or "uploaded_script"
#             if uploaded_file.lower().endswith(".docx"):
#                 try:
#                     left_text, used_left = extract_left_column_script_or_default(uploaded_file)
#                 except Exception:
#                     left_text, used_left = "", False

#                 if used_left and left_text.strip():
#                     # Two-column script detected → use ONLY left VO column and render plain
#                     script_text = left_text
#                     source_docx_path = uploaded_key  # keep S3 key (we saved uploaded file)
#                     heading_ranges = []
#                     st.session_state["render_plain_from_docx"] = True
#                 else:
#                     # Regular DOCX: flatten if tables exist; otherwise build plain text+meta
#                     path_to_use = uploaded_file
#                     if _docx_contains_tables(path_to_use):
#                         flat = flatten_docx_tables_to_longtext(path_to_use)
#                         st.session_state.flattened_docx_path = flat
#                         st.session_state.flatten_used = True
#                         path_to_use = flat
#                     script_text, heading_ranges = build_docx_text_with_meta(path_to_use)
#                     source_docx_path = uploaded_key
#                     # For non-left-column DOCX, prefer DOCX render unless we flattened
#                     st.session_state["render_plain_from_docx"] = bool(st.session_state.get("flatten_used"))
#             else:
#                 # txt/pdf → always render as plain text
#                 script_text = load_script_file(uploaded_file)
#                 source_docx_path = uploaded_key
#                 st.session_state["render_plain_from_docx"] = True
#         else:
#             st.warning("Please upload a script **or** paste text in the second tab.")
#             st.stop()

#         if len(script_text.strip()) < 50:
#             st.error("Extracted text looks too short. Please check your input.")
#             st.stop()

#         with st.spinner("Running analysis…"):
#             try:
#                 review_text = run_review_multi(
#                     script_text=script_text,
#                     prompts_dir=PROMPTS_DIR,  # treated as S3 prefix by review_engine_multi
#                     temperature=0.0
#                 )
#             finally:
#                 if uploaded_file and os.path.exists(uploaded_file):
#                     try:
#                         os.remove(uploaded_file)
#                     except Exception:
#                         pass

#         data = extract_review_json(review_text)
#         if not data:
#             st.error("JSON not detected in model output.")
#             st.stop()

#         st.session_state.script_text      = script_text
#         st.session_state.base_stem        = base_stem
#         st.session_state.data             = data
#         st.session_state.heading_ranges   = heading_ranges
#         st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
#         st.session_state.param_choice     = None
#         st.session_state.source_docx_path = source_docx_path
#         st.session_state.review_ready     = True
#         st.session_state.ui_mode          = "review"

#         _save_history_snapshot(
#             title=base_stem,
#             data=data,
#             script_text=script_text,
#             source_docx_path=source_docx_path,
#             heading_ranges=heading_ranges,
#             spans_by_param=st.session_state.spans_by_param,
#             aoi_match_ranges=st.session_state.get("aoi_match_ranges", {}),
#             used_left_column=bool(st.session_state.get("render_plain_from_docx", False)),  # NEW
#         )

#         _clear_query_params()
#         st.rerun()

# # ---------- Results screen ----------
# def render_review():
#     script_text     = st.session_state.script_text
#     data            = st.session_state.data
#     spans_by_param  = st.session_state.spans_by_param
#     scores: Dict[str,int] = (data or {}).get("scores", {}) or {}
#     source_docx_path: Optional[str] = st.session_state.source_docx_path

#     # If our source_docx_path is an S3 key/url, ensure we have a local copy for rendering
#     # BUT skip DOCX rendering when we used the left-column extractor or forced plain.
#     docx_local: Optional[str] = None
#     render_plain = bool(st.session_state.get("render_plain_from_docx"))
#     preferred = st.session_state.get("flattened_docx_path") if st.session_state.get("flatten_used") else None
#     if not render_plain and not preferred and source_docx_path:
#         if source_docx_path.endswith(".docx"):
#             docx_local = ensure_local_copy(source_docx_path)

#     left, center, right = st.columns([1.1, 2.7, 1.4], gap="large")

#     with left:
#         st.subheader("Final score")
#         ordered = [p for p in PARAM_ORDER if p in scores]
#         df = pd.DataFrame({"Parameter": ordered, "Score (1–10)": [scores.get(p, "") for p in ordered]})
#         st.dataframe(df, hide_index=True, use_container_width=True)
#         st.markdown(f'**Overall:** {data.get("overall_rating","—")}/10')
#         st.divider()

#         strengths = (data or {}).get("strengths") or []
#         if not strengths:
#             per = (data or {}).get("per_parameter", {}) or {}
#             best = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
#             for name, sc in best:
#                 if sc >= 8 and name in per:
#                     exp = _sanitize_editor_text((per[name] or {}).get("explanation", "") or "")
#                     first = re.split(r"(?<=[.!?])\s+", exp.strip())[0] if exp else f"Consistently strong {name.lower()}."
#                     strengths.append(f"{name}: {first}")
#                 if len(strengths) >= 3: break

#         def _bullets(title: str, items):
#             st.markdown(f"**{title}**")
#             for s in (items or []):
#                 if isinstance(s, str) and s.strip(): st.write("• " + _sanitize_editor_text(s))
#             if not items: st.write("• —")

#         _bullets("Strengths", strengths)
#         _bullets("Weaknesses", data.get("weaknesses"))
#         _bullets("Suggestions", data.get("suggestions"))
#         _bullets("Drop-off Risks", data.get("drop_off_risks"))
#         st.markdown("**Viral Quotient**"); st.write(_sanitize_editor_text(data.get("viral_quotient","—")))

#     with right:
#         st.subheader("Parameters")
#         st.markdown('<div class="param-row">', unsafe_allow_html=True)
#         for p in [p for p in PARAM_ORDER if p in scores]:
#             if st.button(p, key=f"chip_{p}", help="Show inline AOI highlights for this parameter"):
#                 st.session_state.param_choice = p
#         st.markdown('</div>', unsafe_allow_html=True)

#         sel = st.session_state.param_choice
#         if sel:
#             blk = (data.get("per_parameter", {}) or {}).get(sel, {}) or {}
#             st.markdown(f"**{sel} — Score:** {scores.get(sel,'—')}/10")

#             if blk.get("explanation"):
#                 st.markdown("**Why this score**"); st.write(_sanitize_editor_text(blk["explanation"]))
#             if blk.get("weakness") and blk["weakness"] != "Not present":
#                 st.markdown("**Weakness**"); st.write(_sanitize_editor_text(blk["weakness"]))
#             if blk.get("suggestion") and blk["suggestion"] != "Not present":
#                 st.markdown("**Suggestion**"); st.write(_sanitize_editor_text(blk["suggestion"]))

#             if blk.get("summary"):
#                 st.markdown("**Summary**"); st.write(_sanitize_editor_text(blk["summary"]))

#     with center:
#         st.subheader("Script with inline highlights")
#         spans = st.session_state.spans_by_param.get(st.session_state.param_choice, []) if st.session_state.param_choice else []

#         aoi_payload: Dict[str, Dict[str, str]] = {}
#         data_per = (data or {}).get("per_parameter") or {}
#         s_e_map = st.session_state.get("aoi_match_ranges", {})
#         sel = st.session_state.param_choice

#         def _mk_line(aid: str, fallback_q: str = "") -> str:
#             if aid in s_e_map:
#                 s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
#                 return matched_line if len(matched_line) <= 320 else matched_line[:300] + "…"
#             return _sanitize_editor_text(fallback_q or "")

#         def _collect(param_name: str):
#             blk = (data_per.get(param_name) or {})
#             for i, item in enumerate(blk.get("areas_of_improvement") or [], 1):
#                 aid = f"{param_name.replace(' ','_')}-AOI-{i}"
#                 aoi_payload[aid] = {
#                     "line": _mk_line(aid, (item or {}).get("quote_verbatim","")),
#                     "issue": _sanitize_editor_text((item or {}).get("issue","")),
#                     "fix": _sanitize_editor_text((item or {}).get("fix","")),
#                     "why": _sanitize_editor_text((item or {}).get("why_this_helps","")),
#                 }

#         if sel: _collect(sel)
#         else:
#             for pn in [p for p in PARAM_ORDER if p in data_per]:
#                 _collect(pn)

#         payload_json = json.dumps(aoi_payload, ensure_ascii=False)

#         frame_theme_css = """
#         <style>
#           :root{
#             --m7-surface: #eef2f7;
#             --m7-on-surface: #0f172a;
#             --m7-border: rgba(15,23,42,.14);
#           }
#           @media (prefers-color-scheme: dark){
#             :root{
#               --m7-surface: #2f333a;
#               --m7-on-surface: #ffffff;
#               --m7-border: rgba(255,255,255,.18);
#             }
#             body { background: transparent !important; }
#           }
#           .docxwrap{ background: var(--m7-surface); color: var(--m7-on-surface); border: 1px solid var(--m7-border); border-radius: 12px; padding: 16px 14px 18px; }
#           .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
#           .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); }
#         </style>
#         """

#         tooltip_css = """
#         <style>
#         .aoi-pop { position: absolute; max-width: 520px; min-width: 320px; background: var(--m7-surface); border: 1px solid var(--m7-border); border-radius: 10px;
#           box-shadow: 0 10px 25px rgba(0,0,0,.12); padding: 12px 14px; z-index: 9999; transform: translateY(-8px); color: var(--m7-on-surface); }
#         .aoi-pop h4 { margin: 0 0 .35rem 0; font-size: .95rem; }
#         .aoi-pop p  { margin: .15rem 0; line-height: 1.5; }
#         .aoi-pop .muted { opacity:.85; font-size:.85rem; }
#         .aoi-arrow { position:absolute; left:50%; transform:translateX(-50%); bottom:-7px; width:0;height:0;border-left:7px solid transparent; border-right:7px solid transparent;border-top:7px solid var(--m7-border); }
#         .aoi-arrow::after{ content:""; position:absolute; left:-6px; top:-7px; width:0;height:0; border-left:6px solid transparent;border-right:6px solid transparent;border-top:6px solid var(--m7-surface); }
#         </style>
#         """

#         # Choose rendering source
#         if (not render_plain) and docx_local and os.path.splitext(docx_local)[1].lower() == ".docx":
#             def render_docx_html_with_highlights(docx_path: str, highlight_spans: List[Tuple[int,int,str,str]]) -> str:
#                 doc = Document(docx_path)
#                 spans = [s for s in highlight_spans if s[0] < s[1]]
#                 spans.sort(key=lambda x: x[0])
#                 cur_span = 0
#                 current_offset = 0
#                 def esc(s: str) -> str:
#                     return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
#                 def open_mark_if_needed(html_parts, mark_state, color, end, aid):
#                     if not mark_state["open"]:
#                         html_parts.append(
#                             f'<mark class="aoi-mark" data-aid="{aid}" style="background:{color}33;border:1px solid {color};border-radius:3px;padding:0 2px;">'
#                         )
#                         mark_state.update(open=True, end=end, color=color, aid=aid)
#                 def close_mark_if_open(html_parts, mark_state):
#                     if mark_state["open"]:
#                         html_parts.append('</mark>')
#                         mark_state.update(open=False, end=None, color=None, aid=None)
#                 def _wrap_inline(safe_text: str, run) -> str:
#                     out = safe_text
#                     if getattr(run, "underline", False): out = f"<u>{out}</u>"
#                     if getattr(run, "italic", False): out = f"<em>{out}</em>"
#                     if getattr(run, "bold", False): out = f"<strong>{out}</strong>"
#                     return out
#                 def emit_run_text(run_text: str, run, html_parts: List[str], mark_state: Dict[str, Any]):
#                     nonlocal cur_span, current_offset
#                     t = run_text or ""; i = 0
#                     while i < len(t):
#                         next_start, next_end, color, next_aid = None, None, None, None
#                         if cur_span < len(spans):
#                             next_start, next_end, color, next_aid = spans[cur_span]
#                         if not mark_state["open"]:
#                             if cur_span >= len(spans) or current_offset + (len(t) - i) <= next_start:
#                                 chunk = t[i:]; html_parts.append(_wrap_inline(esc(chunk), run)); current_offset += len(chunk); break
#                             if current_offset < next_start:
#                                 take = next_start - current_offset
#                                 chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                                 current_offset += take; i += take; continue
#                             open_mark_if_needed(html_parts, mark_state, color, next_end, next_aid)
#                         else:
#                             take = min(mark_state["end"] - current_offset, len(t) - i)
#                             if take > 0:
#                                 chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
#                                 current_offset += take; i += take
#                             if current_offset >= mark_state["end"]:
#                                 close_mark_if_open(html_parts, mark_state)
#                                 cur_span += 1
#                 html: List[str] = ['<div class="docxwrap">']
#                 seen_tc_ids: set = set()
#                 for blk in _iter_docx_blocks(doc):
#                     if isinstance(blk, Paragraph):
#                         mark_state = {"open": False, "end": None, "color": None, "aid": None}
#                         sty = (blk.style.name or "").lower() if blk.style else ""
#                         open_tag = '<div class="h1">' if sty.startswith("heading 1") else \
#                                    '<div class="h2">' if sty.startswith("heading 2") else \
#                                    '<div class="h3">' if sty.startswith("heading 3") else "<p>"
#                         close_tag = "</div>" if sty.startswith("heading") else "</p>"
#                         html.append(open_tag)
#                         for run in blk.runs:
#                             emit_run_text(run.text or "", run, html, mark_state)
#                         close_mark_if_open(html, mark_state)
#                         html.append(close_tag)
#                         current_offset += 1
#                     else:
#                         html.append("<table>")
#                         for row in blk.rows:
#                             html.append("<tr>")
#                             row_cell_tcs = [(id(cell._tc), cell) for cell in row.cells]
#                             for idx, (tc_id, cell) in enumerate(row_cell_tcs):
#                                 html.append("<td>")
#                                 if tc_id not in seen_tc_ids:
#                                     seen_tc_ids.add(tc_id)
#                                     for p_idx, p in enumerate(cell.paragraphs):
#                                         mark_state = {"open": False, "end": None, "color": None, "aid": None}
#                                         html.append("<div>")
#                                         for run in p.runs:
#                                             emit_run_text(run.text or "", run, html, mark_state)
#                                         close_mark_if_open(html, mark_state)
#                                         html.append("</div>")
#                                         if p_idx != len(cell.paragraphs) - 1:
#                                             current_offset += 1
#                                 html.append("</td>")
#                                 if idx != len(row_cell_tcs) - 1: current_offset += 1
#                             html.append("</tr>"); current_offset += 1
#                         html.append("</table>"); current_offset += 1
#                 html.append("</div>")
#                 return "".join(html)

#             html_core = render_docx_html_with_highlights(
#                 docx_local,
#                 merge_overlaps_and_adjacent(script_text, spans)
#             )
#         else:
#             from html import escape as _esc
#             orig = script_text
#             spans2 = [s for s in merge_overlaps_and_adjacent(orig, spans) if s[0] < s[1]]
#             spans2.sort(key=lambda x: x[0])
#             cur = 0; buf: List[str] = []
#             for s,e,c,aid in spans2:
#                 if s > cur: buf.append(_esc(orig[cur:s]))
#                 buf.append(
#                     f'<mark class="aoi-mark" data-aid="{aid}" '
#                     f'style="background:{c}33;border:1px solid {c};border-radius:3px;padding:0 2px;">'
#                     f'{_esc(orig[s:e])}</mark>'
#                 )
#                 cur = e
#             if cur < len(orig): buf.append(_esc(orig[cur:]))
#             html_core = (
#                 '<div class="docxwrap"><p style="white-space:pre-wrap; '
#                 'line-height:1.7; font-family:ui-serif, Georgia, Times New Roman, serif;">'
#                 + "".join(buf) +
#                 '</p></div>'
#             )

#         html_shell = """
# %%FRAME_THEME_CSS%%
# %%TOOLTIP_CSS%%
# <div id="m7-doc">%%HTML_CORE%%</div>
# <div id="aoi-pop" class="aoi-pop" style="display:none;">
#   <div id="aoi-pop-content"></div>
#   <div class="aoi-arrow"></div>
# </div>
# <script>
# (function(){
#   const AOI = __PAYLOAD__;
#   const wrap = document.getElementById('m7-doc');
#   const pop  = document.getElementById('aoi-pop');
#   const body = document.getElementById('aoi-pop-content');

#   function resizeIframe() {
#     try {
#       const h = Math.max(
#         document.documentElement.scrollHeight,
#         document.body.scrollHeight
#       );
#       if (window.frameElement) {
#         window.frameElement.style.height = (h + 20) + 'px';
#         window.frameElement.style.width  = '100%';
#       }
#     } catch(e) {}
#   }
#   window.addEventListener('load', resizeIframe);
#   window.addEventListener('resize', resizeIframe);

#   function hide(){ pop.style.display='none'; }
#   function showFor(mark){
#     const aid = mark.getAttribute('data-aid');
#     const d = AOI[aid]; if(!d) return;
#     body.innerHTML =
#       (d.line  ? '<p><strong>Line:</strong> '  + d.line  + '</p>' : '') +
#       (d.issue ? '<p><strong>Issue:</strong> ' + d.issue + '</p>' : '') +
#       (d.fix   ? '<p><strong>Fix:</strong> '   + d.fix   + '</p>' : '') +
#       (d.why   ? '<p class="muted">'           + d.why   + '</p>' : '');
#     pop.style.display = 'block';

#     const r = mark.getBoundingClientRect();
#     const scY = window.scrollY || document.documentElement.scrollTop;
#     const scX = window.scrollX || document.documentElement.scrollLeft;
#     let top  = r.top + scY - pop.offsetHeight - 10;
#     let left = r.left + scX + r.width/2 - pop.offsetWidth/2;
#     if (top < 8) top = r.bottom + scY + 10;
#     if (left < 8) left = 8;
#     pop.style.top  = top + 'px';
#     pop.style.left = left + 'px';

#     resizeIframe();
#   }

#   wrap.addEventListener('click', (e) => {
#     const m = e.target.closest('.aoi-mark');
#     if(!m){ hide(); return; }
#     if(pop.style.display === 'block'){ hide(); }
#     showFor(m);
#     e.stopPropagation();
#   });

#   document.addEventListener('click', (e) => {
#     if(!e.target.closest('.aoi-pop') && !e.target.closest('.aoi-mark')) hide();
#   });
# })();
# </script>
# """
#         html_shell = (
#             html_shell
#             .replace("%%FRAME_THEME_CSS%%", frame_theme_css)
#             .replace("%%TOOLTIP_CSS%%", tooltip_css)
#             .replace("%%HTML_CORE%%", html_core)
#             .replace("__PAYLOAD__", payload_json)
#         )
#         components.html(html_shell, height=400, scrolling=False)

# # ---------- Router & query param open ----------
# _open_qp = _get_query_param("open")
# # keep legacy query-param open if present (will try via run_id fallback)
# if _open_qp and _open_history_run_by_id(_open_qp):
#     _clear_query_params()

# # Handle in-place open requests from Recents buttons FIRST
# if st.session_state.get("_open_run_key") or st.session_state.get("_open_run_id"):
#     key = st.session_state.pop("_open_run_key", None)
#     rid = st.session_state.pop("_open_run_id", None)

#     opened = False
#     if key:
#         opened = _open_history_by_key(key)  # most reliable
#     if not opened and rid:
#         opened = _open_history_run_by_id(rid)  # fallback via search/manifest
#     if opened:
#         _clear_query_params()
#         st.rerun()

# mode = st.session_state.ui_mode
# if mode == "recents":
#     _render_recents_centerpane()
# elif mode == "review" and st.session_state.review_ready:
#     render_review()
# else:
#     render_home()




















#########################################################




























# app_grammarly_ui.py — Runpod S3-only + Stable Recents + In-place open

import os, re, glob, json, tempfile, difflib, uuid, datetime, shutil, time
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional

import streamlit as st
import pandas as pd
import streamlit.components.v1 as components  # for inline HTML/JS popup

# ---- utils & engine ----
from utils1 import (
    extract_review_json,
    PARAM_ORDER,
    load_script_file,
    extract_left_column_script_or_default,  # <-- left-column extractor for DOCX tables
)
from review_engine_multi import run_review_multi

# ---- DOCX rendering imports ----
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

# =========================
# RunPod S3 (S3-only helpers)
# =========================
import boto3
from botocore.config import Config
from botocore.exceptions import ClientError, EndpointConnectionError, NoCredentialsError

# .env loader (so env vars from .env are available)
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

# ---------- Secrets loader (_sget) ----------
def _sget(key: str, default: str = "", section: Optional[str] = None) -> str:
    """
    Secrets getter with precedence:
      env var > st.secrets[section][key] > st.secrets[key] > default
    """
    v = os.getenv(key, "")
    if v:
        return v.strip()
    try:
        if section:
            sect = st.secrets.get(section, {})
            if isinstance(sect, dict):
                v2 = sect.get(key)
                if isinstance(v2, str) and v2.strip():
                    return v2.strip()
        v3 = st.secrets.get(key)
        if isinstance(v3, str) and v3.strip():
            return v3.strip()
    except Exception:
        pass
    return (default or "").strip()

# Primary config (supports .env and/or [runpod_s3] secrets; env wins)
_RP_ENDPOINT = _sget("RUNPOD_S3_ENDPOINT", section="runpod_s3")
_RP_BUCKET   = _sget("RUNPOD_S3_BUCKET", section="runpod_s3")
_RP_REGION   = _sget("RUNPOD_S3_REGION", section="runpod_s3") or _sget("AWS_DEFAULT_REGION")

# Credentials: prefer AWS_*; fall back to Runpod-style (both *_ID and legacy names)
_AK = (
    _sget("AWS_ACCESS_KEY_ID")
    or _sget("RUNPOD_S3_ACCESS_KEY_ID", section="runpod_s3")
    or _sget("RUNPOD_S3_ACCESS_KEY", section="runpod_s3")
)
_SK = (
    _sget("AWS_SECRET_ACCESS_KEY")
    or _sget("RUNPOD_S3_SECRET_ACCESS_KEY", section="runpod_s3")
    or _sget("RUNPOD_S3_SECRET_KEY", section="runpod_s3")
)
_ST = _sget("AWS_SESSION_TOKEN")  # optional

# Options (boolean-ish strings)
_FORCE_PATH = (_sget("RUNPOD_S3_FORCE_PATH_STYLE", "true", section="runpod_s3") or "true").lower() in {"1","true","yes"}
_USE_SSL    = (_sget("RUNPOD_S3_USE_SSL",          "true", section="runpod_s3") or "true").lower() in {"1","true","yes"}
_VERIFY_SSL = (_sget("RUNPOD_S3_VERIFY_SSL",       "true", section="runpod_s3") or "true").lower() in {"1","true","yes"}

def _s3_enabled() -> bool:
    return bool(_RP_ENDPOINT and _RP_BUCKET and _AK and _SK)

@st.cache_resource(show_spinner=False)
def _s3_client():
    if not _s3_enabled():
        return None
    session_kwargs = dict(
        aws_access_key_id=_AK,
        aws_secret_access_key=_SK,
    )
    if _ST:
        session_kwargs["aws_session_token"] = _ST

    cfg = Config(
        signature_version="s3v4",
        s3={"addressing_style": "path" if _FORCE_PATH else "auto"},
        retries={"max_attempts": 3, "mode": "standard"}
    )
    return boto3.client(
        "s3",
        endpoint_url=_RP_ENDPOINT,
        region_name=_RP_REGION or None,
        use_ssl=_USE_SSL,
        verify=_VERIFY_SSL,
        config=cfg,
        **session_kwargs,
    )

def _assert_s3_ready():
    """
    Hard-require S3. If not ready or not reachable, show a friendly error and stop the app
    BEFORE any upload/save happens. This prevents deep runtime crashes.
    """
    if not _s3_enabled():
        missing = []
        if not _RP_ENDPOINT: missing.append("RUNPOD_S3_ENDPOINT")
        if not _RP_BUCKET:   missing.append("RUNPOD_S3_BUCKET")
        if not _AK:         missing.append("AWS_ACCESS_KEY_ID / RUNPOD_S3_ACCESS_KEY_ID / RUNPOD_S3_ACCESS_KEY")
        if not _SK:         missing.append("AWS_SECRET_ACCESS_KEY / RUNPOD_S3_SECRET_ACCESS_KEY / RUNPOD_S3_SECRET_KEY")
        st.error("S3 (Runpod) is required but not configured.")
        if missing:
            st.write("**Missing keys:**", ", ".join(missing))
        st.info(
            "Set these in your `.env` or in `st.secrets` (optionally under a `[runpod_s3]` section):\n\n"
            "- RUNPOD_S3_ENDPOINT (e.g., `https://s3api-eu-ro-1.runpod.io`)\n"
            "- RUNPOD_S3_BUCKET\n"
            "- RUNPOD_S3_REGION (optional if your endpoint doesn’t require it)\n"
            "- RUNPOD_S3_ACCESS_KEY / RUNPOD_S3_ACCESS_KEY_ID\n"
            "- RUNPOD_S3_SECRET_KEY / RUNPOD_S3_SECRET_ACCESS_KEY\n"
            "- (optional) AWS_SESSION_TOKEN"
        )
        st.stop()

    # Try a tiny list to verify connectivity/creds
    try:
        _s3_client().list_objects_v2(Bucket=_RP_BUCKET, Prefix="Scriptmodel/", MaxKeys=1)
    except Exception as e:
        st.error("Cannot reach Runpod S3 with the provided settings.")
        st.code(str(e))
        st.stop()

# Call early so the app never proceeds without a healthy S3
_assert_s3_ready()

# ---------- S3 I/O (S3-only) ----------
def save_text_key(key: str, text: str) -> str:
    key = key.lstrip("/")
    if not _s3_enabled():
        raise RuntimeError("S3 is not configured (RUNPOD_* / AWS_* envs).")
    kwargs = {
        "Bucket": _RP_BUCKET,
        "Key": key,
        "Body": text.encode("utf-8"),
    }
    if key.endswith(".json"):
        kwargs["ContentType"] = "application/json"
        kwargs["CacheControl"] = "no-store"
    _s3_client().put_object(**kwargs)
    return f"s3://{_RP_BUCKET}/{key}"

def save_bytes_key(key: str, data: bytes) -> str:
    key = key.lstrip("/")
    if not _s3_enabled():
        raise RuntimeError("S3 is not configured (RUNPOD_* / AWS_* envs).")
    _s3_client().put_object(Bucket=_RP_BUCKET, Key=key, Body=data)
    return f"s3://{_RP_BUCKET}/{key}"

def read_text_key(key: str, default: str = "") -> str:
    if not _s3_enabled():
        return default
    tries = 4
    for k in range(tries):
        try:
            resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
            body = resp["Body"].read().decode("utf-8", errors="ignore")
            if body.strip() == "" and k < tries - 1:
                time.sleep(0.25 * (k + 1))
                continue
            return body
        except Exception:
            if k < tries - 1:
                time.sleep(0.25 * (k + 1))
                continue
            return default
    return default

def read_bytes_key(key: str) -> Optional[bytes]:
    if not _s3_enabled():
        return None
    tries = 4
    for k in range(tries):
        try:
            resp = _s3_client().get_object(Bucket=_RP_BUCKET, Key=key)
            return resp["Body"].read()
        except Exception:
            if k < tries - 1:
                time.sleep(0.25 * (k + 1))
                continue
            return None
    return None

def list_prefix(prefix: str) -> List[str]:
    if not _s3_enabled():
        return []
    out: List[str] = []
    token = None
    s3_prefix = prefix.rstrip("/") + "/"
    try:
        while True:
            kwargs = {"Bucket": _RP_BUCKET, "Prefix": s3_prefix}
            if token:
                kwargs["ContinuationToken"] = token
            resp = _s3_client().list_objects_v2(**kwargs)
            for c in resp.get("Contents", []):
                k = c.get("Key", "")
                if k.endswith(".json"):
                    out.append(k)
            token = resp.get("NextContinuationToken")
            if not token:
                break
    except (ClientError, EndpointConnectionError, NoCredentialsError):
        return []
    return out

def presigned_url(key: str, expires: int = 3600) -> Optional[str]:
    if not _s3_enabled():
        return None
    try:
        return _s3_client().generate_presigned_url(
            "get_object",
            Params={"Bucket": _RP_BUCKET, "Key": key},
            ExpiresIn=expires
        )
    except ClientError:
        return None

def ensure_local_copy(key_or_keyurl: str) -> Optional[str]:
    """
    Always download to a temp file for parsing (DOCX/PDF).
    """
    if not _s3_enabled():
        return None
    key = key_or_keyurl
    if key.startswith("s3://"):
        parts = key.split("/", 3)
        key = parts[3] if len(parts) >= 4 else ""
    data = read_bytes_key(key)
    if data is None:
        return None
    fd, tmp = tempfile.mkstemp(suffix=os.path.splitext(key)[1] or "")
    os.close(fd)
    with open(tmp, "wb") as f:
        f.write(data)
    return tmp

def _s3_health_summary() -> dict:
    info = {
        "enabled": _s3_enabled(),
        "endpoint": _RP_ENDPOINT,
        "bucket": _RP_BUCKET,
        "region": _RP_REGION,
        "has_keys": bool(_AK and _SK),
    }
    if not _s3_enabled():
        info["status"] = "disabled"
        return info
    try:
        _ = _s3_client().list_objects_v2(Bucket=_RP_BUCKET, Prefix="Scriptmodel/outputs/_history/", MaxKeys=1)
        info["status"] = "ok"
    except Exception as e:
        info["status"] = f"error: {getattr(e, 'response', {}).get('Error', {}).get('Code', str(e))}"
    return info

# ---------- Folders (all under Scriptmodel/) ----------
BASE_PREFIX = "Scriptmodel"
SCRIPTS_DIR = f"{BASE_PREFIX}/scripts"
PROMPTS_DIR = f"{BASE_PREFIX}/prompts"
OUTPUT_DIR  = f"{BASE_PREFIX}/outputs"
HISTORY_DIR = f"{OUTPUT_DIR}/_history"

# ---------- Colors ----------
PARAM_COLORS: Dict[str, str] = {
    "Suspense Building":              "#ff6b6b",
    "Language/Tone":                  "#6b8cff",
    "Intro + Main Hook/Cliffhanger":  "#ffb86b",
    "Story Structure + Flow":         "#a78bfa",
    "Pacing":                         "#f43f5e",
    "Mini-Hooks (30–60s)":            "#eab308",
    "Outro (Ending)":                 "#8b5cf6",
    "Grammar & Spelling":             "#10b981",
}

STRICT_MATCH_ONLY = False

# ---------- App config ----------
st.set_page_config(page_title="M7 — Grammarly UI", page_icon="🕵️", layout="wide")

# ---------- Header patch & CSS ----------
def render_app_title():
    st.markdown('<h1 class="app-title">Viral Script Reviewer</h1>', unsafe_allow_html=True)
    st.markdown("""
    <style>
    html { color-scheme: light dark; }
    :root{ --m7-surface:#eef2f7; --m7-on-surface:#0f172a; --m7-border:rgba(15,23,42,.14); --sep:#e5e7eb; }
    @media (prefers-color-scheme: dark){
      :root{ --m7-surface:#2f333a; --m7-on-surface:#ffffff; --m7-border:rgba(255,255,255,.18); --sep:#2a2f37; }
    }
    .stApp .block-container { padding-top: 4.25rem !important; }
    .app-title{ font-weight:700; font-size:2.1rem; line-height:1.3; margin:0 0 1rem 0; padding-left:40px!important; padding-top:.25rem!important; }
    [data-testid="collapsedControl"] { z-index: 6 !important; }
    header[data-testid="stHeader"], .stAppHeader { background: transparent !important; box-shadow:none!important; }
    @media (min-width: 992px){ .app-title { padding-left: 0 !important; } }
    div[data-testid="column"]:nth-of-type(1){position:relative;}
    div[data-testid="column"]:nth-of-type(1)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
    div[data-testid="column"]:nth-of-type(2){position:relative;}
    div[data-testid="column"]:nth-of-type(2)::after{content:"";position:absolute;top:0;right:0;width:1px;height:100%;background:var(--sep);}
    .m7-card{ background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:14px 16px; color:var(--m7-on-surface); }
    .m7-card, .m7-card * { color:var(--m7-on-surface)!important; }
    .docxwrap{ background:var(--m7-surface); color:var(--m7-on-surface); border:1px solid var(--m7-border); border-radius:12px; padding:16px 14px 18px; }
    .docxwrap .h1,.docxwrap .h2,.docxwrap .h3 { font-weight:700; margin:10px 0 6px; }
    .docxwrap .h1{font-size:1.3rem; border-bottom:2px solid currentColor; padding-bottom:4px;}
    .docxwrap .h2{font-size:1.15rem; border-bottom:1px solid currentColor; padding-bottom:3px;}
    .docxwrap .h3{font-size:1.05rem;}
    .docxwrap p{ margin:10px 0; line-height:1.7; font-family: ui-serif, Georgia, "Times New Roman", serif; }
    .docxwrap table{ border-collapse:collapse; width:100%; margin:12px 0; }
    .docxwrap th,.docxwrap td{ border:1px solid var(--m7-border); padding:8px; vertical-align:top; line-height:1.6; }
    .docxwrap mark{ padding:0 2px; border-radius:3px; border:1px solid var(--m7-border); cursor:pointer; }
    .rec-card{ display:block; text-decoration:none!important; background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:14px 16px; margin:10px 0 16px; box-shadow:0 1px 2px rgba(0,0,0,.06); color:var(--m7-on-surface)!important; transition: filter .1s ease, transform .02s ease; }
    .rec-card:hover{ filter:brightness(1.02); }
    .rec-card:active{ transform: translateY(1px); }
    .rec-title{font-weight:600; margin-bottom:.25rem;}
    .rec-meta{opacity:.85!important; font-size:12.5px; margin-bottom:.4rem;}
    .rec-row{display:flex; align-items:center; justify-content:space-between; gap:12px;}
    .stTextInput>div>div, .stTextArea>div>div, .stNumberInput>div>div, .stDateInput>div>div, .stTimeInput>div>div, .stFileUploader>div, div[data-baseweb="select"]{ background:var(--m7-surface)!important; border:1px solid var(--m7-border)!important; border-radius:10px!important; color:var(--m7-on-surface)!important; }
    .stTextInput input,.stTextArea textarea,.stNumberInput input,.stDateInput input,.stTimeInput input,.stFileUploader div,div[data-baseweb="select"] *{ color:var(--m7-on-surface)!important; }
    .stTextInput input::placeholder,.stTextArea textarea::placeholder{ color:rgba(16,24,39,.55)!important; }
    @media (prefers-color-scheme: dark){ .stTextInput input::placeholder,.stTextArea textarea::placeholder{ color:rgba(255,255,255,.75)!important; } }
    div[data-testid="stFileUploaderDropzone"] label span { color: var(--m7-on-surface) !important; opacity:1!important; }
    div[data-testid="stDataFrame"]{ background:var(--m7-surface); border:1px solid var(--m7-border); border-radius:12px; padding:6px 8px; color:var(--m7-on-surface); }
    .stMarkdown pre, pre[class*="language-"], .stCodeBlock{ background:var(--m7-surface)!important; color:var(--m7-on-surface)!important; border:1px solid var(--m7-border)!important; border-radius:12px!important; padding:12px 14px!important; overflow:auto; }
    .stMarkdown pre code{ background:transparent!important; color:inherit!important; }
    </style>
    """, unsafe_allow_html=True)

render_app_title()

# ---------- Session defaults ----------
for key, default in [
    ("review_ready", False),
    ("script_text", ""),
    ("base_stem", ""),
    ("data", None),
    ("spans_by_param", {}),
    ("param_choice", None),
    ("source_docx_path", None),
    ("heading_ranges", []),
    ("flattened_docx_path", None),
    ("flatten_used", False),
    ("ui_mode", "home"),
    ("render_plain_from_docx", False),   # NEW: persist render intent for Recents
]:
    st.session_state.setdefault(key, default)

# Recents stability helpers
st.session_state.setdefault("_last_history_cache", [])
st.session_state.setdefault("_open_run_key", None)
st.session_state.setdefault("_open_run_id", None)

# ---------- helpers for query params ----------
def _get_query_param(key: str) -> Optional[str]:
    val = None
    try:
        val = st.query_params.get(key)
    except Exception:
        q = st.experimental_get_query_params()
        v = q.get(key)
        if isinstance(v, list): val = v[0] if v else None
        else: val = v
    return val

def _clear_query_params():
    try:
        st.query_params.clear()
    except Exception:
        st.experimental_set_query_params()

# ---------- Sanitizer ----------
_EMOJI_RE = re.compile(
    r'[\U0001F1E0-\U0001F6FF\U0001F900-\U0001FAFF\U00002700-\U000027BF\U0001F300-\U0001F5FF]',
    flags=re.UNICODE
)
def _sanitize_editor_text(s: Optional[str]) -> str:
    if not s: return ""
    t = str(s)
    t = re.sub(r'\bDecision\s*:\s*', '', t, flags=re.I)
    t = re.sub(r'\bScore\s*[:\-]?\s*\d+(\.\d+)?\b', '', t, flags=re.I)
    t = re.sub(r'^\s*(\(?\d+[\)\.]|\-|\•)\s*', '', t, flags=re.M)
    t = re.sub(r'^\s*[-*]\s+', '• ', t, flags=re.M)
    t = _EMOJI_RE.sub('', t)
    t = re.sub(r'[ \t]+', ' ', t)
    t = re.sub(r'\n{3,}', '\n\n', t)
    return t.strip()

# ---------- DOCX traversal ----------
def _iter_docx_blocks(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)

# ---------- Auto-flatten ----------
def _docx_contains_tables(path: str) -> bool:
    doc = Document(path)
    for blk in _iter_docx_blocks(doc):
        if isinstance(blk, Table):
            return True
    return False

def _copy_paragraph(dest_doc: Document, src_para: Paragraph):
    p = dest_doc.add_paragraph()
    try:
        if src_para.style and src_para.style.name:
            p.style = src_para.style.name
    except Exception:
        pass
    for run in src_para.runs:
        r = p.add_run(run.text or "")
        r.bold = run.bold
        r.italic = run.italic
        r.underline = run.underline
    return p

def flatten_docx_tables_to_longtext(source_path: str) -> str:
    src = Document(source_path)
    new = Document()
    for blk in _iter_docx_blocks(src):
        if isinstance(blk, Paragraph):
            _copy_paragraph(new, blk)
        else:
            seen_tc_ids = set()
            for row in blk.rows:
                for cell in row.cells:
                    tc_id = id(cell._tc)
                    if tc_id in seen_tc_ids:
                        continue
                    seen_tc_ids.add(tc_id)
                    for p in cell.paragraphs:
                        _copy_paragraph(new, p)
                new.add_paragraph("")
            new.add_paragraph("")
    fd, tmp_path = tempfile.mkstemp(suffix=".docx"); os.close(fd); new.save(tmp_path)
    return tmp_path

# ---------- Build plain text + heading ranges ----------
def _iter_docx_blocks(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)

def build_docx_text_with_meta(docx_path: str) -> Tuple[str, List[Tuple[int,int]]]:
    doc = Document(docx_path)
    out: List[str] = []
    heading_ranges: List[Tuple[int,int]] = []
    current_offset = 0

    def _append_and_advance(s: str):
        nonlocal current_offset
        out.append(s); current_offset += len(s)

    seen_tc_ids: set = set()
    for blk in _iter_docx_blocks(doc):
        if isinstance(blk, Paragraph):
            para_text = "".join(run.text or "" for run in blk.runs)
            sty = (blk.style.name or "").lower() if blk.style else ""
            if sty.startswith("heading"):
                start = current_offset; end = start + len(para_text)
                heading_ranges.append((start, end))
            _append_and_advance(para_text); _append_and_advance("\n")
        else:
            for row in blk.rows:
                row_cell_tcs = []
                for cell in row.cells:
                    tc_id = id(cell._tc)
                    row_cell_tcs.append((tc_id, cell))
                for idx, (tc_id, cell) in enumerate(row_cell_tcs):
                    if tc_id in seen_tc_ids:
                        if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
                        continue
                    seen_tc_ids.add(tc_id)
                    cell_text_parts: List[str] = []
                    for i, p in enumerate(cell.paragraphs):
                        t = "".join(r.text or "" for r in p.runs)
                        sty = (p.style.name or "").lower() if p.style else ""
                        if sty.startswith("heading"):
                            hs = current_offset + sum(len(x) for x in cell_text_parts)
                            he = hs + len(t)
                            heading_ranges.append((hs, he))
                        cell_text_parts.append(t)
                        if i != len(cell.paragraphs) - 1:
                            cell_text_parts.append("\n")
                    cell_text = "".join(cell_text_parts)
                    _append_and_advance(cell_text)
                    if idx != len(row_cell_tcs) - 1: _append_and_advance("\t")
                _append_and_advance("\n")
            _append_and_advance("\n")

    return "".join(out), heading_ranges

# ---------- Matching / spans (unchanged core logic) ----------
_BRIDGE_CHARS = set("\u200b\u200c\u200d\u2060\ufeff\xa0\u00ad")
def _normalize_keep_len(s: str) -> str:
    trans = {
        "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
        "\u2013": "-", "\u2014": "-",
        "\xa0": " ",
        "\u200b": " ", "\u200c": " ", "\u200d": " ", "\u2060": " ",
        "\ufeff": " ", "\u00ad": " ",
    }
    return (s or "").translate(str.maketrans(trans))
def _tokenize(s: str) -> List[str]: return re.findall(r"\w+", (s or "").lower())
def _iter_sentences_with_spans(text: str) -> List[Tuple[int,int,str]]:
    spans = []
    for m in re.finditer(r'[^.!?]+[.!?]+|\Z', text, flags=re.S):
        s, e = m.start(), m.end()
        seg = text[s:e]
        if seg.strip(): spans.append((s, e, seg))
    return spans
def _squash_ws(s: str) -> str: return re.sub(r"\s+", " ", s or "").strip()
def _clean_quote_for_match(q: str) -> str:
    if not q: return ""
    q = _normalize_keep_len(q).strip()
    q = re.sub(r'^[\'"“”‘’\[\(\{<…\-\–\—\s]+', '', q)
    q = re.sub(r'[\'"“”‘’\]\)\}>…\-\–\—\s]+$', '', q)
    return _squash_ws(q)
def _snap_and_bridge_to_word(text: str, start: int, end: int, max_bridge: int = 2) -> Tuple[int,int]:
    n = len(text); s, e = max(0,start), max(start,end)
    def _is_inv(ch: str) -> bool: return ch in _BRIDGE_CHARS
    while s > 0:
        prev = text[s-1]; cur = text[s] if s < n else ""
        if prev.isalnum() and cur.isalnum(): s -= 1; continue
        j = s; brid = 0
        while j < n and _is_inv(text[j]): brid += 1; j += 1
        if brid and (s-1)>=0 and text[s-1].isalnum() and (j<n and text[j].isalnum()): s -= 1; continue
        break
    while e < n:
        prev = text[e-1] if e>0 else ""; nxt = text[e]
        if prev.isalnum() and nxt.isalnum(): e += 1; continue
        j = e; brid = 0
        while j < n and _is_inv(text[j]): brid += 1; j += 1
        if brid and (e-1)>=0 and text[e-1].isalnum() and (j<n and text[j].isalnum()): e = j + 1; continue
        break
    while e < n and text[e] in ',"”’\')]}': e += 1
    return s, e
def _heal_split_word_left(text: str, start: int) -> int:
    i = start
    if i <= 1 or i >= len(text): return start
    if text[i-1] != " ": return start
    j = i - 2
    while j >= 0 and text[j].isalpha(): j -= 1
    prev_token = text[j+1:i-1]
    if len(prev_token) == 1: return i - 2
    return start
def _overlaps_any(s: int, e: int, ranges: List[Tuple[int,int]]) -> bool:
    for rs, re_ in ranges:
        if e > rs and s < re_: return True
    return False
def _fuzzy_window_span(tl: str, nl: str, start: int, w: int) -> Tuple[float, Optional[Tuple[int,int]]]:
    window = tl[start:start+w]
    sm = difflib.SequenceMatcher(a=nl, b=window)
    blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
    if not blocks: return 0.0, None
    coverage = sum(b.size for b in blocks) / max(1, len(nl))
    first_b = min(blocks, key=lambda b: b.b); last_b = max(blocks, key=lambda b: b.b + b.size)
    s = start + first_b.b; e = start + last_b.b + last_b.size
    return coverage, (s, e)
def find_span_smart(text: str, needle: str) -> Optional[Tuple[int,int]]:
    if not text or not needle: return None
    t_orig = text; t_norm = _normalize_keep_len(text); n_norm = _clean_quote_for_match(needle)
    if not n_norm: return None
    tl = t_norm.lower(); nl = n_norm.lower()
    i = tl.find(nl)
    if i != -1:
        s, e = _snap_and_bridge_to_word(t_orig, i, i + len(nl)); s = _heal_split_word_left(t_orig, s)
        return (s, e)
    m = re.search(re.escape(nl).replace(r"\ ", r"\s+"), tl, flags=re.IGNORECASE)
    if m:
        s, e = _snap_and_bridge_to_word(t_orig, m.start(), m.end()); s = _heal_split_word_left(t_orig, s)
        return (s, e)
    if not STRICT_MATCH_ONLY and len(nl) >= 12:
        w = max(60, min(240, len(nl) + 80))
        best_cov, best_span = 0.0, None
        step = max(1, w // 2)
        for start in range(0, max(1, len(tl) - w + 1), step):
            cov, se = _fuzzy_window_span(tl, nl, start, w)
            if cov > best_cov: best_cov, best_span = cov, se
        if best_span and best_cov >= 0.65:
            s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1])
            if s > 0 and t_orig[s-1:s+1].lower() in {" the", " a", " an"}: s -= 1
            s = _heal_split_word_left(t_orig, s)
            return (s, e)
    if not STRICT_MATCH_ONLY:
        keys = [w for w in _tokenize(nl) if len(w) >= 4][:8]
        if len(keys) >= 2:
            kset = set(keys)
            best_score, best_span = 0.0, None
            for s, e, seg in _iter_sentences_with_spans(t_norm):
                toks = set(_tokenize(seg)); ov = len(kset & toks)
                if ov == 0: continue
                score = ov / max(2, len(kset)); length_pen = min(1.0, 120 / max(20, e - s)); score *= (0.6 + 0.4 * length_pen)
                if score > best_score: best_score, best_span = score, (s, min(e, s + 400))
            if best_span and best_score >= 0.35:
                s, e = _snap_and_bridge_to_word(t_orig, best_span[0], best_span[1]); s = _heal_split_word_left(t_orig, s)
                return (s, e)
    return None

def merge_overlaps_and_adjacent(base_text: str,
                                spans: List[Tuple[int,int,str,str]],
                                max_gap: int = 2) -> List[Tuple[int,int,str,str]]:
    if not spans: return []
    spans = sorted(spans, key=lambda x: x[0]); out = [spans[0]]
    _PUNCT_WS = set(" \t\r\n,.;:!?)('\"[]{}-—–…") | _BRIDGE_CHARS
    for s, e, c, aid in spans[1:]:
        ps, pe, pc, paid = out[-1]
        if c == pc and s <= pe: out[-1] = (ps, max(pe, e), pc, paid); continue
        if c == pc and s - pe <= max_gap:
            gap = base_text[max(0, pe):max(0, s)]
            if all((ch in _PUNCT_WS) for ch in gap): out[-1] = (ps, e, pc, paid); continue
        out.append((s, e, c, aid))
    return out

def _is_heading_like(q: str) -> bool:
    if not q: return True
    s = q.strip()
    if not re.search(r'[.!?]', s):
        words = re.findall(r"[A-Za-z]+", s)
        if 1 <= len(words) <= 7:
            caps = sum(1 for w in words if w and w[0].isupper())
            if caps / max(1, len(words)) >= 0.8: return True
        if s.lower() in {"introduction","voiceover","outro","epilogue","prologue","credits","title","hook","horrifying discovery","final decision"}: return True
        if len(s) <= 3: return True
    return False

def _is_heading_context(script_text: str, s: int, e: int) -> bool:
    left = script_text.rfind("\n", 0, s) + 1
    right = script_text.find("\n", e); right = len(script_text) if right == -1 else right
    line = script_text[left:right].strip()
    if len(line) <= 70 and not re.search(r'[.!?]', line):
        words = re.findall(r"[A-Za-z]+", line)
        if 1 <= len(words) <= 8:
            caps = sum(1 for w in words if w and w[0].isupper())
            if caps / max(1, len(words)) >= 0.7: return True
    return False

def _tighten_to_quote(script_text: str, span: Tuple[int,int], quote: str) -> Tuple[int,int]:
    if not span or not quote: return span
    s, e = span
    if e <= s or s < 0 or e > len(script_text): return span
    window = script_text[s:e]; win_norm = _normalize_keep_len(window).lower(); q_norm = _clean_quote_for_match(quote).lower()
    if not q_norm: return span
    i = win_norm.find(q_norm)
    if i == -1:
        m = re.search(re.escape(q_norm).replace(r"\ ", r"\s+"), win_norm, flags=re.IGNORECASE)
        if not m: return span
        i, j = m.start(), m.end()
    else:
        j = i + len(q_norm)
    s2, e2 = s + i, s + j
    s2, e2 = _snap_and_bridge_to_word(script_text, s2, e2); s2 = _heal_split_word_left(script_text, s2)
    if s2 >= s and e2 <= e and e2 > s2: return (s2, e2)
    return span

def build_spans_by_param(script_text: str, data: dict, heading_ranges: Optional[List[Tuple[int,int]]] = None) -> Dict[str, List[Tuple[int,int,str,str]]]:
    heading_ranges = heading_ranges or []
    raw = (data or {}).get("per_parameter", {}) or {}
    per: Dict[str, Dict[str, Any]] = {k:(v or {}) for k,v in raw.items()}
    spans_map: Dict[str, List[Tuple[int,int,str,str]]] = {p: [] for p in PARAM_ORDER}
    st.session_state["aoi_match_ranges"] = {}

    for p in spans_map.keys():
        color = PARAM_COLORS.get(p, "#ffd54f")
        blk = per.get(p, {}) or {}
        aois = blk.get("areas_of_improvement") or []
        for idx, item in enumerate(aois, start=1):
            raw_q = (item or {}).get("quote_verbatim", "") or ""
            q = _sanitize_editor_text(raw_q)
            clean = _clean_quote_for_match(re.sub(r"^[•\-\d\.\)\s]+", "", q).strip())
            if not clean: continue
            if _is_heading_like(clean): continue
            pos = find_span_smart(script_text, clean)
            if not pos: continue
            pos = _tighten_to_quote(script_text, pos, raw_q)
            s, e = pos
            if heading_ranges and _overlaps_any(s, e, heading_ranges): continue
            if _is_heading_context(script_text, s, e): continue
            aid = f"{p.replace(' ','_')}-AOI-{idx}"
            spans_map[p].append((s, e, color, aid))
            st.session_state["aoi_match_ranges"][aid] = (s, e)
    return spans_map

# ---------- History (S3-aware + manifest + cache) ----------
_MANIFEST_KEY = f"{HISTORY_DIR}/_manifest.json"

def _manifest_read() -> List[dict]:
    txt = read_text_key(_MANIFEST_KEY, default="")
    if not txt.strip():
        return []
    try:
        arr = json.loads(txt)
        if isinstance(arr, list):
            return arr
    except Exception:
        return []
    return []

def _manifest_append(entry: dict):
    # read-modify-write with small retry
    for k in range(3):
        cur = _manifest_read()
        cur.append(entry)
        try:
            save_text_key(_MANIFEST_KEY, json.dumps(cur, ensure_ascii=False, indent=2))
            return
        except Exception:
            if k < 2:
                time.sleep(0.2 * (k + 1))
            else:
                return

def _maybe_copy_docx_to_history(source_docx_path: Optional[str], run_id: str) -> Optional[str]:
    """
    If source_docx_path is a local temp (downloaded), upload it under _history so Recents can re-render.
    If it's already an S3 key/url, just return that key/url.
    """
    try:
        if not source_docx_path:
            return None
        # If path exists locally (temp), push to S3 history
        if os.path.exists(source_docx_path):
            with open(source_docx_path, "rb") as f:
                save_bytes_key(f"{HISTORY_DIR}/{run_id}.docx", f.read())
            return f"{HISTORY_DIR}/{run_id}.docx"
        # If it's an S3 reference already
        return source_docx_path
    except Exception:
        return None

def _save_history_snapshot(title: str, data: dict, script_text: str,
                           source_docx_path: Optional[str], heading_ranges: List[Tuple[int,int]],
                           spans_by_param: Dict[str, List[Tuple[int,int,str,str]]],
                           aoi_match_ranges: Dict[str, Tuple[int,int]],
                           used_left_column: bool = False):  # NEW
    run_id = str(uuid.uuid4()); now = datetime.datetime.now()
    created_at_iso = now.replace(microsecond=0).isoformat()
    created_at_human = now.strftime("%Y-%m-%d %H:%M:%S")

    stable_docx_key_or_path = _maybe_copy_docx_to_history(source_docx_path, run_id)

    blob = {
        "run_id": run_id, "title": title or "untitled",
        "created_at": created_at_iso, "created_at_human": created_at_human,
        "overall_rating": (data or {}).get("overall_rating", ""),
        "scores": (data or {}).get("scores", {}),
        "data": data or {}, "script_text": script_text or "",
        "source_docx_path": stable_docx_key_or_path or source_docx_path,
        "heading_ranges": heading_ranges or [],
        "spans_by_param": spans_by_param or {},
        "aoi_match_ranges": aoi_match_ranges or {},
        # persist the render intent so Recents shows the same thing
        "used_left_column": bool(used_left_column),
        "render_plain_from_docx": bool(used_left_column),  # mirror for forward/back compat
    }

    out_name = f"{created_at_iso.replace(':','-')}__{run_id}.json"
    out_key = f"{HISTORY_DIR}/{out_name}"
    save_text_key(out_key, json.dumps(blob, ensure_ascii=False, indent=2))

    # Append manifest entry (for list without ListBucket)
    _manifest_append({
        "run_id": run_id,
        "key": out_key,  # exact S3 key to open
        "title": blob["title"],
        "created_at": blob["created_at"],
        "created_at_human": blob["created_at_human"],
        "overall_rating": blob["overall_rating"],
    })

def _load_all_history() -> List[dict]:
    out: List[dict] = []

    # Prefer manifest (works even without ListBucket)
    man = _manifest_read()
    if man:
        # newest first by created_at
        man_sorted = sorted(man, key=lambda r: r.get("created_at",""), reverse=True)
        for m in man_sorted:
            # we don't read the whole JSON here (fast listing); _open_history_by_key loads full
            out.append({
                "run_id": m.get("run_id"),
                "title": m.get("title") or "(untitled)",
                "created_at": m.get("created_at"),
                "created_at_human": m.get("created_at_human", ""),
                "overall_rating": m.get("overall_rating", ""),
                "key": m.get("key"),  # exact S3 key
            })

    # Optional: also list from S3 if allowed, to backfill older runs
    try:
        keys = list_prefix(HISTORY_DIR)
        for key in keys:
            if key.endswith("_manifest.json"):
                continue
            if any(x.get("key") == key for x in out):
                continue  # already present from manifest
            try:
                txt = read_text_key(key, "")
                if not txt:
                    continue
                j = json.loads(txt)
                out.append({
                    "run_id": j.get("run_id"),
                    "title": j.get("title","untitled"),
                    "created_at": j.get("created_at") or "",
                    "created_at_human": j.get("created_at_human",""),
                    "overall_rating": j.get("overall_rating",""),
                    "_key": key,  # loader-provided key
                })
            except Exception:
                continue
    except Exception:
        pass

    out.sort(key=lambda r: r.get("created_at") or "", reverse=True)

    # Last-known-good cache to avoid flicker on transient failures
    if out:
        st.session_state["_last_history_cache"] = out
        return out
    else:
        if st.session_state.get("_last_history_cache"):
            return st.session_state["_last_history_cache"]
        return out

def _open_history_run_by_id(run_id: str) -> bool:
    """Back-compat: open by run_id by searching manifest/list (less reliable than by key)."""
    if not run_id:
        return False
    recs = _load_all_history()
    match = next((r for r in recs if r.get("run_id") == run_id), None)
    if not match:
        return False
    key = match.get("_key") or match.get("key")
    if key:
        return _open_history_by_key(key)
    return False

def _open_history_by_key(key: str) -> bool:
    """
    Open a history run by exact S3 key. Returns True if loaded.
    """
    if not key:
        return False
    try:
        txt = read_text_key(key, "")
        if not txt:
            return False
        jj = json.loads(txt)
    except Exception:
        return False

    # Respect saved render intent
    used_left = bool(jj.get("used_left_column", False) or jj.get("render_plain_from_docx", False))
    st.session_state["render_plain_from_docx"] = used_left

    st.session_state.script_text      = jj.get("script_text","")
    st.session_state.base_stem        = jj.get("title","untitled")
    st.session_state.data             = jj.get("data",{})
    st.session_state.heading_ranges   = jj.get("heading_ranges",[])
    st.session_state.spans_by_param   = jj.get("spans_by_param",{})
    st.session_state.param_choice     = None
    # If we should render plain, ensure we don't try to re-render the DOCX table
    if used_left:
        st.session_state.source_docx_path = None
    else:
        st.session_state.source_docx_path = jj.get("source_docx_path")
    st.session_state.review_ready     = True
    st.session_state["aoi_match_ranges"] = jj.get("aoi_match_ranges", {})
    st.session_state.ui_mode          = "review"
    return True

def _render_recents_centerpane():
    st.subheader("📄 Recents")
    q = st.text_input("Filter by title…", "")

    cols = st.columns([1, 4])
    with cols[0]:
        if st.button("← Back"):
            st.session_state.ui_mode = "home"
            _clear_query_params()
            st.rerun()

    recs = _load_all_history()
    ql = q.strip().lower()
    if ql:
        recs = [r for r in recs if ql in (r.get("title","").lower())]

    if not recs:
        st.caption("No history yet.")
        return

    # Inline styles so you don't need to modify your global CSS
    card_css = """
    <style>
      .rec-card { position:relative; display:block; text-decoration:none!important;
        background:var(--m7-surface); border:1px solid var(--m7-border);
        border-radius:12px; padding:14px 16px; margin:10px 0 16px;
        box-shadow:0 1px 2px rgba(0,0,0,.06); color:var(--m7-on-surface)!important;
        transition: filter .1s ease, transform .02s ease; }
      .rec-card:hover{ filter:brightness(1.02); }
      .rec-card:active{ transform: translateY(1px); }
      .rec-row{ display:flex; align-items:center; justify-content:space-between; gap:12px; }
      .rec-title{ font-weight:600; margin-bottom:.25rem; }
      .rec-meta{ opacity:.85!important; font-size:12.5px; margin-bottom:.4rem; }
      .rec-open{ margin-left:auto; display:inline-block; padding:6px 12px;
        border:1px solid var(--m7-border); border-radius:10px;
        text-decoration:none; font-weight:600; opacity:.95; }
      .rec-open:hover{ filter:brightness(1.05); }
    </style>
    """
    st.markdown(card_css, unsafe_allow_html=True)

    for rec in recs:
        run_id    = rec.get("run_id")
        title     = rec.get("title") or "(untitled)"
        created_h = rec.get("created_at_human","")
        overall   = rec.get("overall_rating","")

        st.markdown(
            f"""
            <a class="rec-card" href="?open={run_id}" target="_self" rel="noopener">
            <div class="rec-row">
                <div>
                <div class="rec-title">{title}</div>
                <div class="rec-meta">{created_h}</div>
                <div><strong>Overall:</strong> {overall if overall != "" else "—"}/10</div>
                </div>
                <span class="rec-open">Open</span>
            </div>
            </a>
            """,
            unsafe_allow_html=True
        )

# ---------- Sidebar ----------
with st.sidebar:
    if st.button("🆕 New review", use_container_width=True):
        fp = st.session_state.get("flattened_docx_path")
        if fp and os.path.exists(fp):
            try: os.remove(fp)
            except Exception: pass
        for k in ["review_ready","script_text","base_stem","data","spans_by_param","param_choice",
                  "source_docx_path","heading_ranges","flattened_docx_path","flatten_used","render_plain_from_docx"]:
            st.session_state[k] = (
                False if k=="review_ready"
                else "" if k in ("script_text","base_stem")
                else {} if k=="spans_by_param"
                else [] if k=="heading_ranges"
                else None if k in ("source_docx_path","flattened_docx_path")
                else False if k in ("flatten_used","render_plain_from_docx")
                else None
            )
        st.session_state.ui_mode = "home"
        _clear_query_params()
        st.rerun()

    if st.button("📁 Recents", use_container_width=True):
        st.session_state.ui_mode = "recents"
        _clear_query_params()
        st.rerun()

# ---------- Input screen ----------
def render_home():
    st.subheader("🎬 Script Source")

    tab_upload, tab_paste = st.tabs(["Upload file", "Paste text"])

    uploaded_file = None
    uploaded_name = None
    uploaded_key  = None

    def _safe_stem(s: str, fallback: str = "pasted_script") -> str:
        s = (s or "").strip()
        if not s:
            return fallback
        s = re.sub(r"[^A-Za-z0-9._\-]+", "_", s)
        s = s.strip("._-") or fallback
        return s

    with tab_upload:
        up = st.file_uploader("Upload .pdf / .docx / .txt", type=["pdf","docx","txt"])
        if up is not None:
            file_bytes = up.read()
            suffix = os.path.splitext(up.name)[1].lower()
            uploaded_key = f"{SCRIPTS_DIR}/{up.name}"
            save_bytes_key(uploaded_key, file_bytes)  # strict S3-only
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(file_bytes)
                uploaded_file = tmp.name
            uploaded_name = os.path.splitext(os.path.basename(up.name))[0] or "uploaded_script"

    with tab_paste:
        paste_title = st.text_input("Title (optional)", placeholder="e.g., my_script")
        pasted_text = st.text_area(
            "Paste your script text here",
            height=360,
            placeholder="Paste the full script text (we’ll analyze this as-is)."
        )

    if st.button("🚀 Run Review", type="primary", use_container_width=True):
        base_stem = "uploaded_script"
        source_docx_path = None
        heading_ranges: List[Tuple[int,int]] = []
        script_text = ""

        if pasted_text and pasted_text.strip():
            base_stem = _safe_stem(paste_title, "pasted_script")
            script_text = pasted_text
            pasted_key = f"{SCRIPTS_DIR}/{base_stem}.txt"
            save_text_key(pasted_key, script_text)
            source_docx_path = pasted_key
            heading_ranges = []
            # Render as plain text for pasted input
            st.session_state["render_plain_from_docx"] = True

        elif uploaded_file:
            base_stem = uploaded_name or "uploaded_script"
            if uploaded_file.lower().endswith(".docx"):
                try:
                    left_text, used_left = extract_left_column_script_or_default(uploaded_file)
                except Exception:
                    left_text, used_left = "", False

                if used_left and left_text.strip():
                    # Two-column script detected → use ONLY left VO column and render plain
                    script_text = left_text
                    source_docx_path = uploaded_key  # S3 key
                    heading_ranges = []
                    st.session_state["render_plain_from_docx"] = True
                else:
                    # Regular DOCX: flatten if tables exist; otherwise build plain text+meta
                    path_to_use = uploaded_file
                    if _docx_contains_tables(path_to_use):
                        flat = flatten_docx_tables_to_longtext(path_to_use)
                        st.session_state.flattened_docx_path = flat
                        st.session_state.flatten_used = True
                        path_to_use = flat
                    script_text, heading_ranges = build_docx_text_with_meta(path_to_use)
                    source_docx_path = uploaded_key
                    # For non-left-column DOCX, prefer DOCX render unless we flattened
                    st.session_state["render_plain_from_docx"] = bool(st.session_state.get("flatten_used"))
            else:
                # txt/pdf → always render as plain text
                script_text = load_script_file(uploaded_file)
                source_docx_path = uploaded_key
                st.session_state["render_plain_from_docx"] = True
        else:
            st.warning("Please upload a script **or** paste text in the second tab.")
            st.stop()

        if len(script_text.strip()) < 50:
            st.error("Extracted text looks too short. Please check your input.")
            st.stop()

        with st.spinner("Running analysis…"):
            try:
                review_text = run_review_multi(
                    script_text=script_text,
                    prompts_dir=PROMPTS_DIR,  # treated as S3 prefix by review_engine_multi
                    temperature=0.0
                )
            finally:
                if uploaded_file and os.path.exists(uploaded_file):
                    try:
                        os.remove(uploaded_file)
                    except Exception:
                        pass

        data = extract_review_json(review_text)
        if not data:
            st.error("JSON not detected in model output.")
            st.stop()

        st.session_state.script_text      = script_text
        st.session_state.base_stem        = base_stem
        st.session_state.data             = data
        st.session_state.heading_ranges   = heading_ranges
        st.session_state.spans_by_param   = build_spans_by_param(script_text, data, heading_ranges)
        st.session_state.param_choice     = None
        st.session_state.source_docx_path = source_docx_path
        st.session_state.review_ready     = True
        st.session_state.ui_mode          = "review"

        _save_history_snapshot(
            title=base_stem,
            data=data,
            script_text=script_text,
            source_docx_path=source_docx_path,
            heading_ranges=heading_ranges,
            spans_by_param=st.session_state.spans_by_param,
            aoi_match_ranges=st.session_state.get("aoi_match_ranges", {}),
            used_left_column=bool(st.session_state.get("render_plain_from_docx", False)),  # NEW
        )

        _clear_query_params()
        st.rerun()

# ---------- Results screen ----------
def render_review():
    script_text     = st.session_state.script_text
    data            = st.session_state.data
    spans_by_param  = st.session_state.spans_by_param
    scores: Dict[str,int] = (data or {}).get("scores", {}) or {}
    source_docx_path: Optional[str] = st.session_state.source_docx_path

    # If our source_docx_path is an S3 key/url, ensure we have a local copy for rendering
    # BUT skip DOCX rendering when we used the left-column extractor or forced plain.
    docx_local: Optional[str] = None
    render_plain = bool(st.session_state.get("render_plain_from_docx"))
    preferred = st.session_state.get("flattened_docx_path") if st.session_state.get("flatten_used") else None
    if not render_plain and not preferred and source_docx_path:
        if str(source_docx_path).endswith(".docx"):
            docx_local = ensure_local_copy(source_docx_path)

    left, center, right = st.columns([1.1, 2.7, 1.4], gap="large")

    with left:
        st.subheader("Final score")
        ordered = [p for p in PARAM_ORDER if p in scores]
        df = pd.DataFrame({"Parameter": ordered, "Score (1–10)": [scores.get(p, "") for p in ordered]})
        st.dataframe(df, hide_index=True, use_container_width=True)
        st.markdown(f'**Overall:** {data.get("overall_rating","—")}/10')
        st.divider()

        strengths = (data or {}).get("strengths") or []
        if not strengths:
            per = (data or {}).get("per_parameter", {}) or {}
            best = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
            for name, sc in best:
                if sc >= 8 and name in per:
                    exp = _sanitize_editor_text((per[name] or {}).get("explanation", "") or "")
                    first = re.split(r"(?<=[.!?])\s+", exp.strip())[0] if exp else f"Consistently strong {name.lower()}."
                    strengths.append(f"{name}: {first}")
                if len(strengths) >= 3: break

        def _bullets(title: str, items):
            st.markdown(f"**{title}**")
            for s in (items or []):
                if isinstance(s, str) and s.strip(): st.write("• " + _sanitize_editor_text(s))
            if not items: st.write("• —")

        _bullets("Strengths", strengths)
        _bullets("Weaknesses", data.get("weaknesses"))
        _bullets("Suggestions", data.get("suggestions"))
        _bullets("Drop-off Risks", data.get("drop_off_risks"))
        st.markdown("**Viral Quotient**"); st.write(_sanitize_editor_text(data.get("viral_quotient","—")))

    with right:
        st.subheader("Parameters")
        st.markdown('<div class="param-row">', unsafe_allow_html=True)
        for p in [p for p in PARAM_ORDER if p in scores]:
            if st.button(p, key=f"chip_{p}", help="Show inline AOI highlights for this parameter"):
                st.session_state.param_choice = p
        st.markdown('</div>', unsafe_allow_html=True)

        sel = st.session_state.param_choice
        if sel:
            blk = (data.get("per_parameter", {}) or {}).get(sel, {}) or {}
            st.markdown(f"**{sel} — Score:** {scores.get(sel,'—')}/10")

            if blk.get("explanation"):
                st.markdown("**Why this score**"); st.write(_sanitize_editor_text(blk["explanation"]))
            if blk.get("weakness") and blk["weakness"] != "Not present":
                st.markdown("**Weakness**"); st.write(_sanitize_editor_text(blk["weakness"]))
            if blk.get("suggestion") and blk["suggestion"] != "Not present":
                st.markdown("**Suggestion**"); st.write(_sanitize_editor_text(blk["suggestion"]))

            if blk.get("summary"):
                st.markdown("**Summary**"); st.write(_sanitize_editor_text(blk["summary"]))

    with center:
        st.subheader("Script with inline highlights")
        spans = st.session_state.spans_by_param.get(st.session_state.param_choice, []) if st.session_state.param_choice else []

        aoi_payload: Dict[str, Dict[str, str]] = {}
        data_per = (data or {}).get("per_parameter") or {}
        s_e_map = st.session_state.get("aoi_match_ranges", {})
        sel = st.session_state.param_choice

        def _mk_line(aid: str, fallback_q: str = "") -> str:
            if aid in s_e_map:
                s_m, e_m = s_e_map[aid]; matched_line = script_text[s_m:e_m]
                return matched_line if len(matched_line) <= 320 else matched_line[:300] + "…"
            return _sanitize_editor_text(fallback_q or "")

        def _collect(param_name: str):
            blk = (data_per.get(param_name) or {})
            for i, item in enumerate(blk.get("areas_of_improvement") or [], 1):
                aid = f"{param_name.replace(' ','_')}-AOI-{i}"
                aoi_payload[aid] = {
                    "line": _mk_line(aid, (item or {}).get("quote_verbatim","")),
                    "issue": _sanitize_editor_text((item or {}).get("issue","")),
                    "fix": _sanitize_editor_text((item or {}).get("fix","")),
                    "why": _sanitize_editor_text((item or {}).get("why_this_helps","")),
                }

        if sel: _collect(sel)
        else:
            for pn in [p for p in PARAM_ORDER if p in data_per]:
                _collect(pn)

        payload_json = json.dumps(aoi_payload, ensure_ascii=False)

        frame_theme_css = """
        <style>
          :root{
            --m7-surface: #eef2f7;
            --m7-on-surface: #0f172a;
            --m7-border: rgba(15,23,42,.14);
          }
          @media (prefers-color-scheme: dark){
            :root{
              --m7-surface: #2f333a;
              --m7-on-surface: #ffffff;
              --m7-border: rgba(255,255,255,.18);
            }
            body { background: transparent !important; }
          }
          .docxwrap{ background: var(--m7-surface); color: var(--m7-on-surface); border: 1px solid var(--m7-border); border-radius: 12px; padding: 16px 14px 18px; }
          .docxwrap, .docxwrap * { color: var(--m7-on-surface) !important; }
          .docxwrap th, .docxwrap td { border:1px solid var(--m7-border); }
        </style>
        """

        tooltip_css = """
        <style>
        .aoi-pop { position: absolute; max-width: 520px; min-width: 320px; background: var(--m7-surface); border: 1px solid var(--m7-border); border-radius: 10px;
          box-shadow: 0 10px 25px rgba(0,0,0,.12); padding: 12px 14px; z-index: 9999; transform: translateY(-8px); color: var(--m7-on-surface); }
        .aoi-pop h4 { margin: 0 0 .35rem 0; font-size: .95rem; }
        .aoi-pop p  { margin: .15rem 0; line-height: 1.5; }
        .aoi-pop .muted { opacity:.85; font-size:.85rem; }
        .aoi-arrow { position:absolute; left:50%; transform:translateX(-50%); bottom:-7px; width:0;height:0;border-left:7px solid transparent; border-right:7px solid transparent;border-top:7px solid var(--m7-border); }
        .aoi-arrow::after{ content:""; position:absolute; left:-6px; top:-7px; width:0;height:0; border-left:6px solid transparent;border-right:6px solid transparent;border-top:6px solid var(--m7-surface); }
        </style>
        """

        # Choose rendering source
        if (not render_plain) and docx_local and os.path.splitext(docx_local)[1].lower() == ".docx":
            def render_docx_html_with_highlights(docx_path: str, highlight_spans: List[Tuple[int,int,str,str]]) -> str:
                doc = Document(docx_path)
                spans = [s for s in highlight_spans if s[0] < s[1]]
                spans.sort(key=lambda x: x[0])
                cur_span = 0
                current_offset = 0
                def esc(s: str) -> str:
                    return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                def open_mark_if_needed(html_parts, mark_state, color, end, aid):
                    if not mark_state["open"]:
                        html_parts.append(
                            f'<mark class="aoi-mark" data-aid="{aid}" style="background:{color}33;border:1px solid {color};border-radius:3px;padding:0 2px;">'
                        )
                        mark_state.update(open=True, end=end, color=color, aid=aid)
                def close_mark_if_open(html_parts, mark_state):
                    if mark_state["open"]:
                        html_parts.append('</mark>')
                        mark_state.update(open=False, end=None, color=None, aid=None)
                def _wrap_inline(safe_text: str, run) -> str:
                    out = safe_text
                    if getattr(run, "underline", False): out = f"<u>{out}</u>"
                    if getattr(run, "italic", False): out = f"<em>{out}</em>"
                    if getattr(run, "bold", False): out = f"<strong>{out}</strong>"
                    return out
                def emit_run_text(run_text: str, run, html_parts: List[str], mark_state: Dict[str, Any]):
                    nonlocal cur_span, current_offset
                    t = run_text or ""; i = 0
                    while i < len(t):
                        next_start, next_end, color, next_aid = None, None, None, None
                        if cur_span < len(spans):
                            next_start, next_end, color, next_aid = spans[cur_span]
                        if not mark_state["open"]:
                            if cur_span >= len(spans) or current_offset + (len(t) - i) <= next_start:
                                chunk = t[i:]; html_parts.append(_wrap_inline(esc(chunk), run)); current_offset += len(chunk); break
                            if current_offset < next_start:
                                take = next_start - current_offset
                                chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
                                current_offset += take; i += take; continue
                            open_mark_if_needed(html_parts, mark_state, color, next_end, next_aid)
                        else:
                            take = min(mark_state["end"] - current_offset, len(t) - i)
                            if take > 0:
                                chunk = t[i:i+take]; html_parts.append(_wrap_inline(esc(chunk), run))
                                current_offset += take; i += take
                            if current_offset >= mark_state["end"]:
                                close_mark_if_open(html_parts, mark_state)
                                cur_span += 1
                html: List[str] = ['<div class="docxwrap">']
                seen_tc_ids: set = set()
                for blk in _iter_docx_blocks(doc):
                    if isinstance(blk, Paragraph):
                        mark_state = {"open": False, "end": None, "color": None, "aid": None}
                        sty = (blk.style.name or "").lower() if blk.style else ""
                        open_tag = '<div class="h1">' if sty.startswith("heading 1") else \
                                   '<div class="h2">' if sty.startswith("heading 2") else \
                                   '<div class="h3">' if sty.startswith("heading 3") else "<p>"
                        close_tag = "</div>" if sty.startswith("heading") else "</p>"
                        html.append(open_tag)
                        for run in blk.runs:
                            emit_run_text(run.text or "", run, html, mark_state)
                        close_mark_if_open(html_parts, mark_state)
                        html.append(close_tag)
                        current_offset += 1
                    else:
                        html.append("<table>")
                        for row in blk.rows:
                            html.append("<tr>")
                            row_cell_tcs = [(id(cell._tc), cell) for cell in row.cells]
                            for idx, (tc_id, cell) in enumerate(row_cell_tcs):
                                html.append("<td>")
                                if tc_id not in seen_tc_ids:
                                    seen_tc_ids.add(tc_id)
                                    for p_idx, p in enumerate(cell.paragraphs):
                                        mark_state = {"open": False, "end": None, "color": None, "aid": None}
                                        html.append("<div>")
                                        for run in p.runs:
                                            emit_run_text(run.text or "", run, html, mark_state)
                                        close_mark_if_open(html_parts, mark_state)
                                        html.append("</div>")
                                        if p_idx != len(cell.paragraphs) - 1:
                                            current_offset += 1
                                html.append("</td>")
                                if idx != len(row_cell_tcs) - 1: current_offset += 1
                            html.append("</tr>"); current_offset += 1
                        html.append("</table>"); current_offset += 1
                html.append("</div>")
                return "".join(html)

            html_core = render_docx_html_with_highlights(
                docx_local,
                merge_overlaps_and_adjacent(script_text, spans)
            )
        else:
            from html import escape as _esc
            orig = script_text
            spans2 = [s for s in merge_overlaps_and_adjacent(orig, spans) if s[0] < s[1]]
            spans2.sort(key=lambda x: x[0])
            cur = 0; buf: List[str] = []
            for s,e,c,aid in spans2:
                if s > cur: buf.append(_esc(orig[cur:s]))
                buf.append(
                    f'<mark class="aoi-mark" data-aid="{aid}" '
                    f'style="background:{c}33;border:1px solid {c};border-radius:3px;padding:0 2px;">'
                    f'{_esc(orig[s:e])}</mark>'
                )
                cur = e
            if cur < len(orig): buf.append(_esc(orig[cur:]))
            html_core = (
                '<div class="docxwrap"><p style="white-space:pre-wrap; '
                'line-height:1.7; font-family:ui-serif, Georgia, Times New Roman, serif;">'
                + "".join(buf) +
                '</p></div>'
            )

        html_shell = """
%%FRAME_THEME_CSS%%
%%TOOLTIP_CSS%%
<div id="m7-doc">%%HTML_CORE%%</div>
<div id="aoi-pop" class="aoi-pop" style="display:none;">
  <div id="aoi-pop-content"></div>
  <div class="aoi-arrow"></div>
</div>
<script>
(function(){
  const AOI = __PAYLOAD__;
  const wrap = document.getElementById('m7-doc');
  const pop  = document.getElementById('aoi-pop');
  const body = document.getElementById('aoi-pop-content');

  function resizeIframe() {
    try {
      const h = Math.max(
        document.documentElement.scrollHeight,
               document.body.scrollHeight
      );
      if (window.frameElement) {
        window.frameElement.style.height = (h + 20) + 'px';
        window.frameElement.style.width  = '100%';
      }
    } catch(e) {}
  }
  window.addEventListener('load', resizeIframe);
  window.addEventListener('resize', resizeIframe);

  function hide(){ pop.style.display='none'; }
  function showFor(mark){
    const aid = mark.getAttribute('data-aid');
    const d = AOI[aid]; if(!d) return;
    body.innerHTML =
      (d.line  ? '<p><strong>Line:</strong> '  + d.line  + '</p>' : '') +
      (d.issue ? '<p><strong>Issue:</strong> ' + d.issue + '</p>' : '') +
      (d.fix   ? '<p><strong>Fix:</strong> '   + d.fix   + '</p>' : '') +
      (d.why   ? '<p class="muted">'           + d.why   + '</p>' : '');
    pop.style.display = 'block';

    const r = mark.getBoundingClientRect();
    const scY = window.scrollY || document.documentElement.scrollTop;
    const scX = window.scrollX || document.documentElement.scrollLeft;
    let top  = r.top + scY - pop.offsetHeight - 10;
    let left = r.left + scX + r.width/2 - pop.offsetWidth/2;
    if (top < 8) top = r.bottom + scY + 10;
    if (left < 8) left = 8;
    pop.style.top  = top + 'px';
    pop.style.left = left + 'px';

    resizeIframe();
  }

  wrap.addEventListener('click', (e) => {
    const m = e.target.closest('.aoi-mark');
    if(!m){ hide(); return; }
    if(pop.style.display === 'block'){ hide(); }
    showFor(m);
    e.stopPropagation();
  });

  document.addEventListener('click', (e) => {
    if(!e.target.closest('.aoi-pop') && !e.target.closest('.aoi-mark')) hide();
  });
})();
</script>
"""
        html_shell = (
            html_shell
            .replace("%%FRAME_THEME_CSS%%", frame_theme_css)
            .replace("%%TOOLTIP_CSS%%", tooltip_css)
            .replace("%%HTML_CORE%%", html_core)
            .replace("__PAYLOAD__", payload_json)
        )
        components.html(html_shell, height=400, scrolling=False)

# ---------- Router & query param open ----------
_open_qp = _get_query_param("open")
# keep legacy query-param open if present (will try via run_id fallback)
if _open_qp and _open_history_run_by_id(_open_qp):
    _clear_query_params()

# Handle in-place open requests from Recents buttons FIRST
if st.session_state.get("_open_run_key") or st.session_state.get("_open_run_id"):
    key = st.session_state.pop("_open_run_key", None)
    rid = st.session_state.pop("_open_run_id", None)

    opened = False
    if key:
        opened = _open_history_by_key(key)  # most reliable
    if not opened and rid:
        opened = _open_history_run_by_id(rid)  # fallback via search/manifest
    if opened:
        _clear_query_params()
        st.rerun()

mode = st.session_state.ui_mode
if mode == "recents":
    _render_recents_centerpane()
elif mode == "review" and st.session_state.review_ready:
    render_review()
else:
    render_home()

